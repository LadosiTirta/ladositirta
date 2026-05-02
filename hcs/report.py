"""
HCS Design — Report Generator  (PROMPT FIX-5 rewrite)
=======================================================
Reference : ACI/PCI CODE-319-25 Ch. 7, 12, 16, 26
            PCI Design Handbook, 8th Edition
            ASCE 7-22  (load combinations)
            AISC Design Guide 11  (vibration)
Units     : SI only — mm, kN, MPa

Chapter order (FIX-5):
  Cover
  Ch 1  — Project Parameters & Code References
  Ch 2  — Section Properties  (MOVED UP)
  Ch 3  — Material Properties
  Ch 4  — Applied Loads & Internal Forces
  Ch 5  — Lifting Stage Check
  Ch 6  — Erection Stage Check
  Ch 7  — Prestress Losses & Development Length
  Ch 8  — Service Stress Checks
  Ch 9  — Flexural & Shear Capacity
  Ch 10 — Deflection & Camber
  Ch 11 — Vibration & Natural Frequency
  Appendix A — Section Property Details
  Appendix B — Prestress Loss Details
  Appendix C — ACI/PCI 319-25 Sec. 12 Diaphragm Summary

All calculations use the pattern:
    Formula  [symbolic]
    Substitution  [with numbers]
    Result  [= value unit]

Diagrams: matplotlib only (more reliable than kaleido/plotly for reports).
Word: python-docx   PDF: fpdf2 / ReportLab
"""

from __future__ import annotations

import io
import math
import traceback
from datetime import datetime
from typing import Any

import numpy as np

# ── python-docx ────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Mm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── fpdf2 ──────────────────────────────────────────────────────────────────────
try:
    from fpdf import FPDF
    HAS_FPDF = True
except ImportError:
    HAS_FPDF = False

# ── matplotlib (always used for diagrams in FIX-5) ────────────────────────────
try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    HAS_MPL = True
except ImportError:
    HAS_MPL = False

# ── Plotly kept only for SFD/BMD fallback reference ───────────────────────────
try:
    import plotly.graph_objects as go
    from plotly.subplots import make_subplots
    HAS_PLOTLY = True
except ImportError:
    HAS_PLOTLY = False


# =============================================================================
# UTILITIES
# =============================================================================

def _g(ss: dict, key: str, default: Any = 0.0) -> Any:
    """Safe session-state getter."""
    return ss.get(key, default)


def _ascii(text: str) -> str:
    """Replace non-ASCII chars with ASCII equivalents (for fpdf2)."""
    table = {
        "\u00b2": "^2",   "\u00b3": "^3",   "\u00b9": "1",
        "\u00b0": "deg",  "\u00b7": ".",     "\u22c5": ".",
        "\u221a": "sqrt", "\u00d7": "x",     "\u00f7": "/",
        "\u2212": "-",    "\u2013": "-",     "\u2014": "--",
        "\u00b1": "+/-",  "\u2264": "<=",    "\u2265": ">=",
        "\u2248": "~=",   "\u2260": "!=",    "\u221e": "inf",
        "\u00bd": "1/2",  "\u00bc": "1/4",   "\u00be": "3/4",
        "\u03b1": "alfa",  "\u03b2": "beta",   "\u03b3": "gamma",
        "\u03b4": "delta", "\u03b5": "epsilon","\u03b6": "zeta",
        "\u03b7": "eta",   "\u03b8": "theta",  "\u03bb": "lambda",
        "\u03bc": "mu",    "\u03bd": "nu",     "\u03c0": "pi",
        "\u03c1": "rho",   "\u03c3": "sigma",  "\u03c4": "tau",
        "\u03c6": "phi",   "\u03c8": "psi",    "\u03c9": "omega",
        "\u03a6": "Phi",   "\u03a9": "Omega",  "\u0394": "Delta",
        "\u2019": "'",    "\u2018": "'",
        "\u201c": '"',    "\u201d": '"',
        "\u2026": "...",  "\u00a0": " ",
        "\u2032": "'",    "\u2033": "''",
        "\u00e9": "e",    "\u00e8": "e",    "\u00ea": "e",
        "\u00e0": "a",    "\u00e2": "a",    "\u00fc": "u",
        "\u2022": "*",    "\u25cf": "*",    "\u2192": "->",
        "\u2713": "OK",   "\u2714": "OK",   "\u2718": "NG",
        "\u26a0": "(!)",  "\u00b4": "'",
    }
    for uni, asc in table.items():
        text = text.replace(uni, asc)
    return text.encode("ascii", errors="replace").decode("ascii")


def _fmt(val: Any, dec: int = 3) -> str:
    try:
        v = float(val)
        if not math.isfinite(v):
            return "--"
        return f"{v:.{dec}f}"
    except (TypeError, ValueError):
        return str(val) if val is not None else "--"


def _span_table(ss: dict, n_seg: int = 10) -> list[dict]:
    """Build span distribution table at 0.0, 0.1 … 1.0 x L_an."""
    x_raw  = _g(ss, "lb_x_arr",  None)
    Vu_raw = _g(ss, "lb_Vu_arr", None)
    Mu_raw = _g(ss, "lb_Mu_arr", None)
    L_an   = float(_g(ss, "L_an", 5850.0))

    if x_raw is None or len(x_raw) < 2:
        wu   = float(_g(ss, "lb_wu_area", 5.0))
        bbot = float(_g(ss, "b_bottom", 1199.0))
        wlin = wu * bbot / 1e6
        Ra   = wlin * L_an / 2.0
        rows = []
        for i in range(n_seg + 1):
            frac = i / n_seg
            x    = frac * L_an
            Vu   = Ra - wlin * x
            Mu   = Ra * x - wlin * x**2 / 2.0
            rows.append({"frac": frac, "x_mm": x, "x_m": x / 1000,
                         "Vu_kN": Vu, "Mu_kNm": Mu / 1e6})
        return rows

    x_arr  = np.asarray(x_raw,  dtype=float)
    Vu_arr = np.asarray(Vu_raw, dtype=float)
    Mu_arr = np.asarray(Mu_raw, dtype=float)
    rows   = []
    for i in range(n_seg + 1):
        frac  = i / n_seg
        x_tgt = frac * L_an
        Vu_i  = float(np.interp(x_tgt, x_arr, Vu_arr))
        Mu_i  = float(np.interp(x_tgt, x_arr, Mu_arr)) / 1e6
        rows.append({"frac": frac, "x_mm": x_tgt, "x_m": x_tgt / 1000,
                     "Vu_kN": Vu_i, "Mu_kNm": Mu_i})
    return rows


# =============================================================================
# MATPLOTLIB DIAGRAM BUILDERS  (all diagrams use matplotlib in FIX-5)
# =============================================================================

def _mpl_sfd_bmd(ss: dict) -> bytes | None:
    """SFD + BMD using matplotlib."""
    if not HAS_MPL:
        return None
    try:
        x_raw  = _g(ss, "lb_x_arr",  None)
        Vu_raw = _g(ss, "lb_Vu_arr", None)
        Mu_raw = _g(ss, "lb_Mu_arr", None)
        Vs_raw = _g(ss, "lb_Vs_arr", None)
        Ms_raw = _g(ss, "lb_Ms_arr", None)
        L_an   = float(_g(ss, "L_an", 5850.0))

        if x_raw is None or len(x_raw) < 2:
            # fallback parabolic
            bbot = float(_g(ss, "b_bottom", 1199.0))
            wu   = float(_g(ss, "lb_wu_area", 5.0))
            wlin = wu * bbot / 1e6
            Ra   = wlin * L_an / 2.0
            x_arr = np.linspace(0, L_an, 101)
            Vu    = Ra - wlin * x_arr
            Mu    = (Ra * x_arr - wlin * x_arr**2 / 2.0) / 1e6
            Vs    = None
            Ms    = None
        else:
            x_arr = np.asarray(x_raw, dtype=float)
            Vu    = np.asarray(Vu_raw, dtype=float)
            Mu    = np.asarray(Mu_raw, dtype=float) / 1e6
            Vs    = np.asarray(Vs_raw, dtype=float) if Vs_raw is not None else None
            Ms    = np.asarray(Ms_raw, dtype=float) / 1e6 if Ms_raw is not None else None

        x_m = x_arr / 1000.0

        fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(10, 6),
                                        facecolor="white", sharex=True)
        fig.subplots_adjust(hspace=0.38)

        # SFD
        ax1.plot(x_m, Vu, color="#1A476F", linewidth=2.0, label="Vu (factored)")
        ax1.fill_between(x_m, Vu, alpha=0.10, color="#1A476F")
        ax1.axhline(0, color="black", linewidth=0.6)
        if Vs is not None:
            ax1.plot(x_m, Vs, color="#1A476F", linewidth=1.5,
                     linestyle="--", label="Vs (service)")
        # transfer-length zone
        l_t_m = float(_g(ss, "lb_l_t", 0.0)) / 1000.0
        L_m   = L_an / 1000.0
        for x0, x1 in [(0, l_t_m), (L_m - l_t_m, L_m)]:
            ax1.axvspan(x0, x1, alpha=0.08, color="red")
        ax1.set_ylabel("Shear (kN)", fontsize=9)
        ax1.set_title("Shear Force Diagram (kN)  — Factored loads", fontsize=10, fontweight="bold")
        ax1.grid(True, linestyle="--", alpha=0.35)
        ax1.legend(fontsize=8)

        # BMD
        ax2.plot(x_m, Mu, color="#1A476F", linewidth=2.0, label="Mu (factored)")
        ax2.fill_between(x_m, Mu, alpha=0.10, color="#1A476F")
        ax2.axhline(0, color="black", linewidth=0.6)
        if Ms is not None:
            ax2.plot(x_m, Ms, color="#1A476F", linewidth=1.5,
                     linestyle="--", label="Ms (service)")
        idx_max = int(np.argmax(Mu))
        ax2.annotate(f"Mu_max={Mu[idx_max]:.2f} kN.m",
                     xy=(x_m[idx_max], Mu[idx_max]),
                     xytext=(x_m[idx_max], Mu[idx_max] * 0.75),
                     arrowprops=dict(arrowstyle="->", color="#555"),
                     fontsize=8, ha="center")
        ax2.set_xlabel("Distance from left support (m)", fontsize=9)
        ax2.set_ylabel("Moment (kN.m)", fontsize=9)
        ax2.set_title("Bending Moment Diagram (kN.m)  — Factored loads", fontsize=10, fontweight="bold")
        ax2.grid(True, linestyle="--", alpha=0.35)
        ax2.legend(fontsize=8)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        traceback.print_exc()
        return None


def _mpl_phi_mn_vs_mu(ss: dict) -> bytes | None:
    """phi*Mn vs Mu envelope along span."""
    if not HAS_MPL:
        return None
    try:
        x_raw  = _g(ss, "lb_x_arr",  None)
        Mu_raw = _g(ss, "lb_Mu_arr", None)
        phi_Mn = float(_g(ss, "cap_phi_Mn", 0.0))
        L_an   = float(_g(ss, "L_an", 5850.0))

        if x_raw is None or len(x_raw) < 2:
            x_m = np.array([0.0, L_an / 1000.0])
            Mu  = np.array([0.0, 0.0])
        else:
            x_m = np.asarray(x_raw, dtype=float) / 1000.0
            Mu  = np.asarray(Mu_raw, dtype=float) / 1e6

        fig, ax = plt.subplots(figsize=(10, 4), facecolor="white")
        ax.plot(x_m, Mu, color="#1A476F", linewidth=2.0, label="Mu (factored)")
        ax.axhline(phi_Mn, color="red", linewidth=1.8, linestyle="--",
                   label=f"phi*Mn = {phi_Mn:.2f} kN.m")
        ax.fill_between(x_m, Mu, alpha=0.10, color="#1A476F")
        ax.set_xlabel("Distance from left support (m)", fontsize=9)
        ax.set_ylabel("Moment (kN.m)", fontsize=9)
        ax.set_title("phi*Mn vs Mu along span", fontsize=10, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(True, linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


def _mpl_phi_vn_envelope(ss: dict) -> bytes | None:
    """phi*Vn vs Vu envelope along span."""
    if not HAS_MPL:
        return None
    try:
        x_raw      = _g(ss, "lb_x_arr",      None)
        Vu_raw     = _g(ss, "lb_Vu_arr",     None)
        Vci_raw    = _g(ss, "cap_Vci_arr",   None)
        Vcw_raw    = _g(ss, "cap_Vcw_arr",   None)
        phi_Vn_raw = _g(ss, "cap_phi_Vn_arr", None)
        L_an       = float(_g(ss, "L_an", 5850.0))

        if x_raw is None or len(x_raw) < 2:
            return None

        x_m    = np.asarray(x_raw,      dtype=float) / 1000.0
        Vu_abs = np.abs(np.asarray(Vu_raw, dtype=float))
        fig, ax = plt.subplots(figsize=(10, 4), facecolor="white")
        ax.plot(x_m, Vu_abs, color="#1A476F", linewidth=2.0, label="|Vu| factored")
        if phi_Vn_raw is not None:
            phi_Vn = np.asarray(phi_Vn_raw, dtype=float)
            ax.plot(x_m, phi_Vn, color="red", linewidth=1.8,
                    linestyle="--", label="phi*Vn (min of Vci, Vcw)")
        if Vci_raw is not None:
            ax.plot(x_m, np.asarray(Vci_raw, dtype=float) * 0.75,
                    color="orange", linewidth=1.0, linestyle=":", label="phi*Vci")
        if Vcw_raw is not None:
            ax.plot(x_m, np.asarray(Vcw_raw, dtype=float) * 0.75,
                    color="green",  linewidth=1.0, linestyle=":", label="phi*Vcw")
        ax.set_xlabel("Distance from left support (m)", fontsize=9)
        ax.set_ylabel("Shear (kN)", fontsize=9)
        ax.set_title("phi*Vn vs Vu envelope along span", fontsize=10, fontweight="bold")
        ax.legend(fontsize=8)
        ax.grid(True, linestyle="--", alpha=0.35)

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


def _mpl_deflection_profile(ss: dict) -> bytes | None:
    """Deflection profile along span (illustrative)."""
    if not HAS_MPL:
        return None
    try:
        L_an      = float(_g(ss, "L_an", 5850.0))
        net_rel   = float(_g(ss, "def_net_release",    0.0))
        total_lt  = float(_g(ss, "def_total_longterm", 0.0))
        x_m = np.linspace(0, L_an / 1000.0, 101)
        # parabolic profile
        xn  = x_m / (L_an / 1000.0)
        mid_rel = net_rel
        mid_lt  = total_lt
        y_rel = 4 * mid_rel * xn * (1 - xn)
        y_lt  = 4 * mid_lt  * xn * (1 - xn)

        fig, ax = plt.subplots(figsize=(10, 4), facecolor="white")
        ax.plot(x_m, y_rel, color="#1A476F", linewidth=2.0,
                label=f"At release  ({mid_rel:.2f} mm at mid)")
        ax.plot(x_m, y_lt,  color="red",     linewidth=1.8, linestyle="--",
                label=f"Long-term  ({mid_lt:.2f} mm at mid)")
        ax.axhline(0, color="black", linewidth=0.6)
        ax.fill_between(x_m, y_lt, alpha=0.08, color="red")
        ax.invert_yaxis()
        ax.set_xlabel("Distance from left support (m)", fontsize=9)
        ax.set_ylabel("Deflection (mm) [downward +ve]", fontsize=9)
        ax.set_title("Deflection Profile (parabolic approximation)", fontsize=10, fontweight="bold")
        ax.legend(fontsize=9)
        ax.grid(True, linestyle="--", alpha=0.35)
        ax.annotate("Note: upward camber shown as negative",
                    xy=(0.01, 0.02), xycoords="axes fraction", fontsize=7, color="#555")

        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=150, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        return None


# =============================================================================
# WORD REPORT  (python-docx)
# =============================================================================

class _DocxReport:
    """
    Word calculation report — FIX-5 chapter order.
    All text uses python-docx Run objects (Unicode OK).
    """

    _NAVY  = RGBColor(0x1A, 0x47, 0x6F)
    _BLUE  = RGBColor(0x2E, 0x74, 0xB5)
    _GREEN = RGBColor(0x00, 0x51, 0x28)
    _GREY  = RGBColor(0x60, 0x60, 0x60)
    _RED   = RGBColor(0xC0, 0x00, 0x00)
    _WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    _LGREY = RGBColor(0xF2, 0xF2, 0xF2)

    def __init__(self, ss: dict):
        self.ss  = ss
        self.doc = Document()
        self._setup_page()

    def _setup_page(self):
        for sec in self.doc.sections:
            sec.top_margin    = Mm(20)
            sec.bottom_margin = Mm(20)
            sec.left_margin   = Mm(22)
            sec.right_margin  = Mm(18)

    # ── Heading helpers ──────────────────────────────────────────────────────
    def _h1(self, text: str):
        p = self.doc.add_heading(text, level=1)
        if p.runs:
            p.runs[0].font.color.rgb = self._NAVY

    def _h2(self, text: str):
        p = self.doc.add_heading(text, level=2)
        if p.runs:
            p.runs[0].font.color.rgb = self._BLUE

    def _h3(self, text: str):
        p = self.doc.add_heading(text, level=3)
        if p.runs:
            p.runs[0].font.color.rgb = self._BLUE

    # ── Body helpers ─────────────────────────────────────────────────────────
    def _calc(self, text: str, ref: str = ""):
        p  = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(text)
        r1.font.name      = "Courier New"
        r1.font.size      = Pt(9)
        r1.font.color.rgb = self._GREEN
        if ref:
            r2 = p.add_run(f"    [{ref}]")
            r2.font.name      = "Courier New"
            r2.font.size      = Pt(7.5)
            r2.font.italic    = True
            r2.font.color.rgb = self._GREY

    def _note(self, text: str):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.italic    = True
        r.font.size      = Pt(8.5)
        r.font.color.rgb = self._GREY

    def _warn(self, text: str):
        p = self.doc.add_paragraph()
        r = p.add_run("  ⚠  " + text)
        r.font.bold      = True
        r.font.size      = Pt(9)
        r.font.color.rgb = self._RED

    def _ok(self, text: str):
        p = self.doc.add_paragraph()
        r = p.add_run("  ✓  " + text)
        r.font.bold      = True
        r.font.size      = Pt(9)
        r.font.color.rgb = self._GREEN

    def _status_line(self, text: str, ok: bool):
        if ok:
            self._ok(text)
        else:
            self._warn(text)

    # ── Table helpers ────────────────────────────────────────────────────────
    def _shd_cell(self, cell, hex_fill: str):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"),  hex_fill)
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"),   "clear")
        tc_pr.append(shd)

    def _kv_table(self, rows: list):
        t  = self.doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, txt in enumerate(("Parameter", "Value")):
            hdr[i].text = txt
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold      = True
            run.font.size      = Pt(9)
            run.font.color.rgb = self._WHITE
            self._shd_cell(hdr[i], "1A476F")
        fill = False
        for key, val in rows:
            rc = t.add_row().cells
            rc[0].text = str(key)
            rc[1].text = str(val)
            for c in rc:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(9)
                if fill:
                    self._shd_cell(c, "F2F2F2")
            fill = not fill
        self.doc.add_paragraph()

    def _wide_table(self, header: list, data: list):
        nc = len(header)
        t  = self.doc.add_table(rows=1, cols=nc)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, txt in enumerate(header):
            hdr[i].text = str(txt)
            run = hdr[i].paragraphs[0].runs[0]
            run.font.bold      = True
            run.font.size      = Pt(8.5)
            run.font.color.rgb = self._WHITE
            self._shd_cell(hdr[i], "1A476F")
        fill = False
        for row in data:
            rc = t.add_row().cells
            for i, val in enumerate(row):
                rc[i].text = str(val)
                if rc[i].paragraphs[0].runs:
                    rc[i].paragraphs[0].runs[0].font.size = Pt(8.5)
                if fill:
                    self._shd_cell(rc[i], "F2F2F2")
            fill = not fill
        self.doc.add_paragraph()

    def _insert_image(self, img_bytes: bytes | None, width_mm: float = 162):
        if not img_bytes:
            self._note("(Diagram not available — matplotlib not installed or error occurred.)")
            return
        buf = io.BytesIO(img_bytes)
        self.doc.add_picture(buf, width=Mm(width_mm))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # =========================================================================
    # COVER
    # =========================================================================
    def _cover(self):
        ss  = self.ss
        doc = self.doc
        doc.add_paragraph()
        doc.add_paragraph()
        t = doc.add_paragraph("HOLLOW CORE SLAB\nSTRUCTURAL DESIGN CALCULATION")
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in t.runs:
            run.font.bold      = True
            run.font.size      = Pt(20)
            run.font.color.rgb = self._NAVY
        doc.add_paragraph()
        self._kv_table([
            ("Project",    _g(ss, "project_name",  "HCS Design — Automatically generated")),
            ("Prepared by", _g(ss, "engineer_name", "Engineer of Record")),
            ("Code ref.",  "ACI/PCI CODE-319-25  |  PCI Design Handbook, 8th Ed."),
            ("Units",      "mm · kN · MPa  (SI only, no conversions)"),
            ("Date",       _g(ss, "report_datetime", datetime.now().strftime("%d %B %Y   %H:%M"))),
            ("Software",   "HCS Design App v1.0  (Streamlit)"),
        ])
        self._note(
            "DISCLAIMER: This report is auto-generated.  "
            "The engineer of record must verify all inputs, assumptions, and results "
            "before any use in design, fabrication, or construction."
        )
        doc.add_page_break()

    # =========================================================================
    # CHAPTER 1 — Project Parameters & Code References
    # =========================================================================
    def _ch1_project_params(self):
        ss = self.ss
        self._h1("1.  Project Parameters and Code References")

        # 1.1 Code references
        self._h2("1.1  Code References")
        self._wide_table(
            ["Standard", "Edition", "Chapters / Sections Used"],
            [
                ["ACI/PCI CODE-319-25", "2025", "Chapters 7, 12, 16, 22, 24, 25, 26"],
                ["PCI Design Handbook", "8th Edition (2017)", "Sec. 2.2, 4.2, 4.7, 4.8, 5.2, 5.3"],
                ["ASCE 7-22", "2022", "Load combinations — Sec. 2.3 (LRFD)"],
                ["AISC Design Guide 11", "2nd Edition", "Floor vibration — walking excitation"],
                ["ASTM A416", "Latest", "7-wire strand properties (Grade 270, Low-Relax)"],
            ]
        )

        # 1.2 Disclaimer
        self._h2("1.2  Disclaimer")
        self._note(
            "This calculation is generated automatically by HCS Design App v1.0.  "
            "It is the responsibility of the engineer of record to verify all design inputs, "
            "section properties, load assumptions, and results before use.  "
            "The software author accepts no liability for errors in design or construction."
        )

        # 1.3 Units
        self._h2("1.3  Units and Sign Conventions")
        self._note("All calculations use SI units throughout.  No imperial conversions are performed.")
        self._kv_table([
            ("Length",       "mm  (millimetres)"),
            ("Force",        "kN  (kilonewtons)"),
            ("Stress",       "MPa  (= N/mm²)"),
            ("Area",         "mm²"),
            ("Moment",       "kN·m"),
            ("Compression",  "Negative (−)"),
            ("Tension",      "Positive (+)"),
            ("Camber",       "Positive upward (+)"),
            ("Deflection",   "Positive downward (−) in calculations; shown as + in tables"),
        ])

        # 1.4 Design inputs
        self._h2("1.4  Design Inputs Summary")
        cs = _g(ss, "core_shape", "Teardrop")
        rows = [
            # Geometry
            ("b_nominal — nominal panel width",     f"{_g(ss,'b_nominal'):.0f} mm"),
            ("b_bottom  — actual bottom width",     f"{_g(ss,'b_bottom'):.0f} mm"),
            ("b_top     — top flange width",        f"{_g(ss,'b_top'):.0f} mm"),
            ("h         — total HCS thickness",     f"{_g(ss,'h'):.0f} mm"),
            ("tf_top    — top flange thickness",    f"{_g(ss,'tf_top'):.0f} mm"),
            ("tf_bot    — bottom flange thickness", f"{_g(ss,'tf_bot'):.0f} mm"),
            ("Core shape",                          str(cs)),
            ("d_core    — core diameter",           f"{_g(ss,'d_core'):.0f} mm"),
            ("n_core    — number of cores",         f"{int(_g(ss,'n_core',9))}"),
            # Span
            ("L_cc      — centre-to-centre span",   f"{_g(ss,'L_cc'):.0f} mm"),
            ("L_an      — analysis span",           f"{_g(ss,'L_an'):.0f} mm"),
            # Loads
            ("SW_HCS    — HCS self-weight",         f"{_g(ss,'SW_HCS'):.3f} kN/m²"),
            ("SW_topping — topping self-weight",    f"{_g(ss,'SW_topping'):.3f} kN/m²"),
            ("SDL        — superimposed dead load", f"{_g(ss,'SDL'):.2f} kN/m²"),
            ("LL         — live load",              f"{_g(ss,'LL'):.2f} kN/m²"),
            # Prestress
            ("Prestress type",                      _g(ss,"ps_type","PC Wire")),
            ("fpi_pct   — initial prestress",       f"{_g(ss,'fpi_pct'):.1f} % fpu"),
            ("Pi        — initial prestress force", f"{_g(ss,'Pi'):.2f} kN"),
            ("SDC       — seismic design category", _g(ss,"sdc","B")),
            ("RH        — relative humidity",       f"{_g(ss,'RH'):.0f} %"),
        ]
        if _g(ss,"has_topping"):
            rows.insert(6, ("t_topping  — structural topping thickness",
                            f"{_g(ss,'t_topping'):.0f} mm"))
        self._kv_table(rows)
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 2 — Section Properties (MOVED UP)
    # =========================================================================
    def _ch2_section_props(self):
        ss = self.ss
        self._h1("2.  Section Properties")
        self._note(
            "Coordinate: y measured UPWARD from the BOTTOM FACE of HCS (topping excluded).\n"
            "Model: rectangular b_top x h; void centroid at y_void = tf_bot + h_core/2  (>97% accuracy).\n"
            "Transformed steel: (n_ps - 1) x Aps method.\n"
            "Ref: ACI/PCI CODE-319-25 Cl. 26.12  |  PCI Design Handbook 8th Ed. Sec. 2.2"
        )

        b_top  = float(_g(ss, "b_top",  1187.0))
        h      = float(_g(ss, "h",       200.0))
        tf_bot = float(_g(ss, "tf_bot",   50.0))
        h_core = float(_g(ss, "h_core",  120.0))
        n_core = int  (_g(ss, "n_core",    9))
        A_c1   = float(_g(ss, "A_core_1", 7107.0))
        A_void = float(_g(ss, "A_voids_total", 63959.0))
        Aps_b  = float(_g(ss, "Aps_bot",  196.0))
        Aps_t  = float(_g(ss, "Aps_top",    0.0))
        dp_bot = float(_g(ss, "dp_bot",   165.0))
        dp_top = float(_g(ss, "dp_top",    30.0))
        n_ps   = float(_g(ss, "sp_n_ps",   5.41))
        Ec_hcs = float(_g(ss, "Ec_hcs", 36793.0))
        Eps    = float(_g(ss, "Eps",   199050.0))
        b_bot  = float(_g(ss, "b_bottom", 1199.0))

        Ag    = float(_g(ss, "sp_Ag",    b_top * h))
        yb_g  = float(_g(ss, "sp_yb_g", h / 2.0))
        Ig    = float(_g(ss, "sp_Ig",    b_top * h**3 / 12.0))
        Sb_g  = float(_g(ss, "sp_Sb_g", Ig / yb_g   if yb_g   > 0 else 0))
        St_g  = float(_g(ss, "sp_St_g", Ig / (h - yb_g) if h - yb_g > 0 else 0))

        An    = float(_g(ss, "sp_An",  1.0))
        yb    = float(_g(ss, "sp_yb",  1.0))
        yt    = float(_g(ss, "sp_yt",  1.0))
        In    = float(_g(ss, "sp_In",  1.0))
        Sb_n  = float(_g(ss, "sp_Sb",  1.0))
        St_n  = float(_g(ss, "sp_St",  1.0))
        r2    = float(_g(ss, "sp_r2",  1.0))
        e_bot = float(_g(ss, "sp_e_bot", 1.0))
        e_net = float(_g(ss, "sp_e_net", e_bot))

        # 2.1 Gross Section
        self._h2("2.1  Gross Section  (rectangular, no voids, no steel)")
        self._calc(
            "  Formula:       Ag = b_top x h",
            "ACI/PCI CODE-319-25 Cl. 26.12.1"
        )
        self._calc(
            f"  Substitution:  Ag = {b_top:.0f} x {h:.0f}"
        )
        self._calc(
            f"  Result:        Ag = {Ag:,.0f} mm2"
        )
        self._calc("")
        self._calc(
            "  Formula:       yb_g = h / 2   (symmetric rectangular)"
        )
        self._calc(
            f"  Substitution:  yb_g = {h:.0f} / 2"
        )
        self._calc(
            f"  Result:        yb_g = {yb_g:.2f} mm  (centroid from bottom)"
        )
        self._calc("")
        self._calc(
            "  Formula:       Ig = b_top x h^3 / 12"
        )
        self._calc(
            f"  Substitution:  Ig = {b_top:.0f} x {h:.0f}^3 / 12"
        )
        self._calc(
            f"  Result:        Ig = {Ig/1e6:.4f} x 10^6 mm4"
        )
        self._calc("")
        self._calc(
            f"  Formula:       Sb_g = Ig / yb_g"
        )
        self._calc(
            f"  Substitution:  Sb_g = {Ig/1e6:.4f}e6 / {yb_g:.2f}"
        )
        self._calc(
            f"  Result:        Sb_g = {Sb_g/1e3:.3f} x 10^3 mm3"
        )
        self._calc("")
        self._calc(
            f"  Formula:       St_g = Ig / yt_g   where yt_g = h - yb_g"
        )
        self._calc(
            f"  Substitution:  St_g = {Ig/1e6:.4f}e6 / {h - yb_g:.2f}"
        )
        self._calc(
            f"  Result:        St_g = {St_g/1e3:.3f} x 10^3 mm3"
        )

        # 2.2 Core Void Area
        self._h2("2.2  Core Void Area")
        cs   = _g(ss, "core_shape", "Teardrop")
        d_c  = float(_g(ss, "d_core", 80.0))
        h_st = float(_g(ss, "h_straight", 40.0))
        h_tp = float(_g(ss, "h_taper",    40.0))
        if cs == "Circular":
            self._calc(
                "  Formula:       A_core_1 = pi/4 x d_core^2   [Circular]"
            )
            self._calc(
                f"  Substitution:  A_core_1 = pi/4 x {d_c:.0f}^2"
            )
        elif cs == "Capsule":
            self._calc(
                "  Formula:       A_core_1 = pi/4 x d_core^2 + d_core x h_straight   [Capsule]"
            )
            self._calc(
                f"  Substitution:  A_core_1 = pi/4 x {d_c:.0f}^2 + {d_c:.0f} x {h_st:.0f}"
            )
        else:
            self._calc(
                "  Formula:       A_core_1 = pi/4 x d_core^2 + 0.65 x d_core x h_taper   [Teardrop]"
            )
            self._calc(
                f"  Substitution:  A_core_1 = pi/4 x {d_c:.0f}^2 + 0.65 x {d_c:.0f} x {h_tp:.0f}"
            )
        self._calc(
            f"  Result:        A_core_1 = {A_c1:.1f} mm2"
        )
        self._calc("")
        self._calc(
            "  Formula:       A_voids_total = n_core x A_core_1"
        )
        self._calc(
            f"  Substitution:  A_voids_total = {n_core} x {A_c1:.1f}"
        )
        self._calc(
            f"  Result:        A_voids_total = {A_void:.0f} mm2"
        )

        # 2.3 Net Section
        self._h2("2.3  Net HCS Section  (voids subtracted + transformed steel)")
        self._note("Ref: PCI Design Handbook 8th Ed. Sec. 2.2.1")

        y_void   = tf_bot + h_core / 2.0
        An_conc  = b_top * h - A_void
        dA_bot   = (n_ps - 1.0) * Aps_b
        dA_top   = (n_ps - 1.0) * Aps_t

        self._note("STEP A — Void centroid:")
        self._calc(
            "  Formula:       y_void = tf_bot + h_core/2"
        )
        self._calc(
            f"  Substitution:  y_void = {tf_bot:.0f} + {h_core:.1f}/2"
        )
        self._calc(
            f"  Result:        y_void = {y_void:.2f} mm  (from bottom)"
        )

        self._note("STEP B — Net concrete area:")
        self._calc(
            "  Formula:       A_net_c = b_top x h - A_voids_total"
        )
        self._calc(
            f"  Substitution:  A_net_c = {b_top:.0f} x {h:.0f} - {A_void:.0f}"
        )
        self._calc(
            f"  Result:        A_net_c = {An_conc:,.0f} mm2"
        )

        self._note(f"STEP C — Modular ratio:  n_ps = Eps/Ec_hcs = {Eps:.0f}/{Ec_hcs:.0f} = {n_ps:.4f}")
        self._calc(
            "  Formula:       dA_bot = (n_ps - 1) x Aps_bot"
        )
        self._calc(
            f"  Substitution:  dA_bot = ({n_ps:.4f} - 1) x {Aps_b:.1f}"
        )
        self._calc(
            f"  Result:        dA_bot = {dA_bot:.1f} mm2"
        )
        if Aps_t > 0:
            self._calc(
                "  Formula:       dA_top = (n_ps - 1) x Aps_top"
            )
            self._calc(
                f"  Substitution:  dA_top = ({n_ps:.4f} - 1) x {Aps_t:.1f}"
            )
            self._calc(
                f"  Result:        dA_top = {dA_top:.1f} mm2"
            )

        self._note("STEP D — Net section totals (parallel-axis theorem):")
        self._calc(
            "  Formula:       An = A_net_c + dA_bot + dA_top"
        )
        self._calc(
            f"  Substitution:  An = {An_conc:,.0f} + {dA_bot:.1f} + {dA_top:.1f}"
        )
        self._calc(
            f"  Result:        An = {An:,.2f} mm2"
        )
        self._calc("")
        self._calc(
            "  Formula:       yb = Sum(Ai x yi) / An   [centroid from bottom]"
        )
        self._calc(
            f"  Result:        yb = {yb:.4f} mm"
        )
        self._calc(
            f"  Formula:       yt = h - yb"
        )
        self._calc(
            f"  Substitution:  yt = {h:.0f} - {yb:.4f}"
        )
        self._calc(
            f"  Result:        yt = {yt:.4f} mm"
        )
        self._calc("")
        self._calc(
            "  Formula:       In = Sum(Ii + Ai x di^2)   [parallel axis theorem]"
        )
        self._calc(
            f"  Result:        In = {In/1e6:.4f} x 10^6 mm4"
        )
        self._calc("")
        self._calc(
            "  Formula:       Sb_n = In / yb"
        )
        self._calc(
            f"  Substitution:  Sb_n = {In/1e6:.4f}e6 / {yb:.4f}"
        )
        self._calc(
            f"  Result:        Sb_n = {Sb_n/1e3:.3f} x 10^3 mm3"
        )
        self._calc("")
        self._calc(
            "  Formula:       St_n = In / yt"
        )
        self._calc(
            f"  Substitution:  St_n = {In/1e6:.4f}e6 / {yt:.4f}"
        )
        self._calc(
            f"  Result:        St_n = {St_n/1e3:.3f} x 10^3 mm3"
        )
        self._calc("")
        self._calc(
            "  Formula:       r2 = In / An   (radius of gyration squared)"
        )
        self._calc(
            f"  Substitution:  r2 = {In/1e6:.4f}e6 / {An:,.2f}"
        )
        self._calc(
            f"  Result:        r2 = {r2:.2f} mm2"
        )

        # 2.4 Eccentricity
        self._h2("2.4  Tendon Eccentricity")
        self._calc(
            "  Formula:       e_bot = yb - dp_bot   (BELOW centroid = favourable for hogging)"
        )
        self._calc(
            f"  Substitution:  e_bot = {yb:.4f} - {dp_bot:.0f}"
        )
        self._calc(
            f"  Result:        e_bot = {e_bot:.4f} mm   (+ve = PS below NA = upward camber)"
        )
        if Aps_t > 0:
            e_top = float(_g(ss, "sp_e_top", dp_top - yb))
            self._calc(
                f"  Formula:       e_top = dp_top - yb"
            )
            self._calc(
                f"  Substitution:  e_top = {dp_top:.0f} - {yb:.4f}"
            )
            self._calc(
                f"  Result:        e_top = {e_top:.4f} mm"
            )
        self._calc(
            f"  Net eccentricity (for loss calcs): e_net = {e_net:.4f} mm"
        )

        # 2.5 Kern points
        self._h2("2.5  Kern Points")
        kt = In / (An * yb) if An * yb > 0 else 0.0
        kb = In / (An * yt) if An * yt > 0 else 0.0
        self._calc(
            "  Formula:       k_t = In / (An x yb)   (upper kern from NA toward top)",
            "ACI/PCI CODE-319-25 Cl. 22.5"
        )
        self._calc(
            f"  Substitution:  k_t = {In/1e6:.4f}e6 / ({An:,.2f} x {yb:.4f})"
        )
        self._calc(
            f"  Result:        k_t = {kt:.2f} mm"
        )
        self._calc(
            "  Formula:       k_b = In / (An x yt)   (lower kern from NA toward bottom)"
        )
        self._calc(
            f"  Substitution:  k_b = {In/1e6:.4f}e6 / ({An:,.2f} x {yt:.4f})"
        )
        self._calc(
            f"  Result:        k_b = {kb:.2f} mm"
        )

        # 2.6 Composite Section
        has_top = bool(_g(ss, "has_topping"))
        t_top   = float(_g(ss, "t_topping", 0.0))
        A_comp  = float(_g(ss, "sp_A_comp",  1.0))
        yb_comp = float(_g(ss, "sp_yb_comp", 1.0))
        yt_comp = float(_g(ss, "sp_yt_comp", 1.0))
        I_comp  = float(_g(ss, "sp_I_comp",  1.0))
        Sbc     = float(_g(ss, "sp_Sbc_comp", 1.0))
        Stc     = float(_g(ss, "sp_Stc_comp", 1.0))

        if has_top and t_top > 0:
            self._h2("2.6  Composite Section  (net HCS + transformed structural topping)")
            self._note("Ref: PCI Design Handbook 8th Ed. Sec. 4.2.3")
            n_mod    = float(_g(ss, "n_mod", 0.818))
            b_nom    = float(_g(ss, "b_nominal", 1200.0))
            b_top_tr = b_nom / n_mod
            A_top_tr = b_top_tr * t_top
            y_top_c  = h + t_top / 2.0
            self._calc(
                "  Formula:       b_top_tr = b_nominal / n_mod   (transformed topping width)"
            )
            self._calc(
                f"  Substitution:  b_top_tr = {b_nom:.0f} / {n_mod:.4f}"
            )
            self._calc(
                f"  Result:        b_top_tr = {b_top_tr:.2f} mm"
            )
            self._calc("")
            self._calc(
                "  Formula:       A_top_tr = b_top_tr x t_topping"
            )
            self._calc(
                f"  Substitution:  A_top_tr = {b_top_tr:.2f} x {t_top:.0f}"
            )
            self._calc(
                f"  Result:        A_top_tr = {A_top_tr:,.2f} mm2"
            )
            self._calc("")
            self._calc(
                "  Formula:       y_top_c = h + t_topping/2   (topping centroid from HCS bottom)"
            )
            self._calc(
                f"  Substitution:  y_top_c = {h:.0f} + {t_top:.0f}/2"
            )
            self._calc(
                f"  Result:        y_top_c = {y_top_c:.1f} mm"
            )
            self._calc("")
            self._calc(
                "  Formula:       A_comp = An + A_top_tr"
            )
            self._calc(
                f"  Substitution:  A_comp = {An:,.2f} + {A_top_tr:,.2f}"
            )
            self._calc(
                f"  Result:        A_comp = {A_comp:,.2f} mm2"
            )
            self._calc("")
            self._calc(
                "  Formula:       yb_comp = Sum(Ai x yi) / A_comp   (from HCS bottom)"
            )
            self._calc(
                f"  Result:        yb_comp = {yb_comp:.4f} mm"
            )
            self._calc(
                "  Formula:       yt_comp = (h + t_topping) - yb_comp   (from top of topping)"
            )
            self._calc(
                f"  Substitution:  yt_comp = ({h:.0f} + {t_top:.0f}) - {yb_comp:.4f}"
            )
            self._calc(
                f"  Result:        yt_comp = {yt_comp:.4f} mm"
            )
            self._calc("")
            self._calc(
                "  Formula:       I_comp = Sum(Ii + Ai x di^2)   [parallel axis, all components]"
            )
            self._calc(
                f"  Result:        I_comp = {I_comp/1e6:.4f} x 10^6 mm4"
            )
            self._calc("")
            self._calc(
                "  Formula:       Sbc_comp = I_comp / yb_comp   (HCS bottom fibre)"
            )
            self._calc(
                f"  Result:        Sbc_comp = {Sbc/1e3:.3f} x 10^3 mm3"
            )
            self._calc(
                "  Formula:       Stc_comp = I_comp / yt_comp   (top of topping fibre)"
            )
            self._calc(
                f"  Result:        Stc_comp = {Stc/1e3:.3f} x 10^3 mm3"
            )
        else:
            self._h2("2.6  Composite Section")
            self._note("No structural topping specified — composite section equals net HCS section.")

        # 2.7 Summary table
        self._h2("2.7  Section Properties Summary Table")
        self._wide_table(
            ["Property", "Gross", "Net HCS", "Composite"],
            [
                ["Area A  (mm2)",       f"{Ag:,.0f}",      f"{An:,.2f}",        f"{A_comp:,.2f}" if has_top else "=Net"],
                ["yb  (mm)",            f"{yb_g:.2f}",     f"{yb:.4f}",         f"{yb_comp:.4f}" if has_top else "=Net"],
                ["yt  (mm)",            f"{h-yb_g:.2f}",   f"{yt:.4f}",         f"{yt_comp:.4f}" if has_top else "=Net"],
                ["I   (x10^6 mm4)",     f"{Ig/1e6:.4f}",   f"{In/1e6:.4f}",     f"{I_comp/1e6:.4f}" if has_top else "=Net"],
                ["Sb  (x10^3 mm3)",     f"{Sb_g/1e3:.3f}", f"{Sb_n/1e3:.3f}",   f"{Sbc/1e3:.3f}" if has_top else "=Net"],
                ["St  (x10^3 mm3)",     f"{St_g/1e3:.3f}", f"{St_n/1e3:.3f}",   f"{Stc/1e3:.3f}" if has_top else "=Net"],
                ["r2  (mm2)",           "--",              f"{r2:.2f}",          "--"],
                ["e_bot (mm)",          "--",              f"{e_bot:.4f}",       "--"],
                ["k_t  (mm)",           "--",              f"{kt:.2f}",          "--"],
                ["k_b  (mm)",           "--",              f"{kb:.2f}",          "--"],
            ]
        )
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 3 — Material Properties
    # =========================================================================
    def _ch3_materials(self):
        ss = self.ss
        self._h1("3.  Material Properties")

        # 3.1 Concrete
        self._h2("3.1  Concrete Properties")
        self._note("Elastic modulus per ACI 318-19 Eq. 19.2.2.1  (used in ACI/PCI CODE-319-25):")
        wc     = float(_g(ss, "wc",      24.0))
        f_ci   = float(_g(ss, "f_ci",    35.0))
        f_c    = float(_g(ss, "f_c",     50.0))
        Ec_hcs = float(_g(ss, "Ec_hcs", 36793.0))
        wc_kg  = wc * 1000.0 / 9.81

        self._calc(
            "  Formula:       Ec_hcs = 0.043 x wc^1.5 x sqrt(f_c)",
            "ACI 318-19 Eq. 19.2.2.1"
        )
        self._calc(
            f"  Substitution:  Ec_hcs = 0.043 x {wc_kg:.0f}^1.5 x sqrt({f_c:.0f})"
        )
        self._calc(
            f"                        = 0.043 x {wc_kg**1.5:.1f} x {math.sqrt(f_c):.4f}"
        )
        self._calc(
            f"  Result:        Ec_hcs = {Ec_hcs:.0f} MPa"
        )

        rows = [
            ("f_ci    — strength at release",      f"{_g(ss,'f_ci'):.0f} MPa"),
            ("f_c_cut — strength at wire cutting", f"{_g(ss,'f_c_cut'):.0f} MPa"),
            ("f_c_del — strength at delivery",     f"{_g(ss,'f_c_del'):.0f} MPa"),
            ("f_c_ere — strength at erection",     f"{_g(ss,'f_c_ere'):.0f} MPa"),
            ("f_c     — 28-day design strength",   f"{f_c:.0f} MPa"),
            ("wc      — HCS unit weight",          f"{wc:.1f} kN/m3"),
            ("Ec_hcs  — elastic modulus (auto)",   f"{Ec_hcs:.0f} MPa"),
        ]
        if _g(ss, "has_topping"):
            wc_top  = float(_g(ss, "wc_top",   24.0))
            f_c_top = float(_g(ss, "f_c_top",  25.0))
            Ec_top  = float(_g(ss, "Ec_top",   30000.0))
            n_mod   = float(_g(ss, "n_mod",    0.818))
            wc_top_kg = wc_top * 1000.0 / 9.81
            self._calc("")
            self._calc(
                "  Formula:       Ec_top = 0.043 x wc_top^1.5 x sqrt(f_c_top)"
            )
            self._calc(
                f"  Substitution:  Ec_top = 0.043 x {wc_top_kg:.0f}^1.5 x sqrt({f_c_top:.0f})"
            )
            self._calc(
                f"  Result:        Ec_top = {Ec_top:.0f} MPa"
            )
            self._calc("")
            self._calc(
                "  Formula:       n_mod = Ec_top / Ec_hcs"
            )
            self._calc(
                f"  Substitution:  n_mod = {Ec_top:.0f} / {Ec_hcs:.0f}"
            )
            self._calc(
                f"  Result:        n_mod = {n_mod:.4f}   (modular ratio)"
            )
            rows += [
                ("f_c_top — topping strength",       f"{f_c_top:.0f} MPa"),
                ("wc_top  — topping unit weight",    f"{wc_top:.1f} kN/m3"),
                ("Ec_top  — topping modulus (auto)", f"{Ec_top:.0f} MPa"),
                ("n_mod   — modular ratio (auto)",   f"{n_mod:.4f}"),
            ]
        self._kv_table(rows)

        # 3.2 Prestressing Steel
        self._h2("3.2  Prestressing Steel Properties")
        ps_type = _g(ss, "ps_type", "PC Wire")
        rows = [
            ("Type",                          str(ps_type)),
            ("Wire dia / Strand size",
             f"{_g(ss,'wire_dia'):.1f} mm" if "Wire" in str(ps_type) else _g(ss,"strand_size","—")),
            ("ps_area — area per unit",       f"{_g(ss,'ps_area'):.2f} mm2"),
            ("fpu     — ultimate strength",   f"{_g(ss,'fpu'):.0f} MPa"),
            ("fpy     — yield strength",      f"{_g(ss,'fpy'):.0f} MPa"),
            ("Eps     — elastic modulus",     f"{_g(ss,'Eps'):.0f} MPa"),
            ("n_bot   — bottom tendon count", f"{int(_g(ss,'n_bot',0))}"),
            ("n_top   — top tendon count",    f"{int(_g(ss,'n_top',0))}"),
            ("cover_bot — clear cover",       f"{_g(ss,'cover_bot'):.0f} mm"),
            ("dp_bot    — effective depth",   f"{_g(ss,'dp_bot'):.0f} mm"),
            ("fpi_pct   — initial prestress", f"{_g(ss,'fpi_pct'):.1f} % of fpu"),
            ("fpi       — initial stress",    f"{_g(ss,'fpi'):.1f} MPa"),
            ("Aps_bot   — total bottom area", f"{_g(ss,'Aps_bot'):.1f} mm2"),
            ("Aps_top   — total top area",    f"{_g(ss,'Aps_top'):.1f} mm2"),
            ("Pi        — initial force",     f"{_g(ss,'Pi'):.2f} kN"),
        ]
        self._kv_table(rows)

        # 3.3 Loss Parameters
        self._h2("3.3  Loss Parameters")
        RH  = float(_g(ss, "RH", 75.0))
        V_S = float(_g(ss, "V_S", 38.0))
        b_bot = float(_g(ss, "b_bottom", 1199.0))
        h_hcs = float(_g(ss, "h", 200.0))
        bw_sh = float(_g(ss, "bw_shear", b_bot - int(_g(ss,"n_core",9)) * float(_g(ss,"d_core",80.0))))
        # Auto V/S for HCS: V = b_bottom * h per metre; S = 2*(b+h) per metre
        # simplified: V/S ~ A_net / (2*(b+h) per unit length)
        An_net = float(_g(ss, "sp_An", b_bot * h_hcs))
        vs_auto = An_net / (2.0 * (b_bot + h_hcs))
        self._calc(
            "  Formula:       V/S = A_net / [2 x (b_bottom + h)]   [auto per unit length]"
        )
        self._calc(
            f"  Substitution:  V/S = {An_net:,.0f} / [2 x ({b_bot:.0f} + {h_hcs:.0f})]"
        )
        self._calc(
            f"  Result:        V/S = {vs_auto:.1f} mm   (user override = {V_S:.1f} mm)"
        )
        self._kv_table([
            ("RH    — relative humidity",   f"{RH:.0f} %"),
            ("V/S   — volume/surface ratio (user)", f"{V_S:.1f} mm"),
            ("V/S   — auto calculated",     f"{vs_auto:.1f} mm"),
        ])
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 4 — Applied Loads & Internal Forces
    # =========================================================================
    def _ch4_loads(self):
        ss = self.ss
        self._h1("4.  Applied Loads and Internal Forces")
        self._note(
            "Load combination: wu = 1.2(SW_HCS + SW_top + SDL) + 1.6LL\n"
            "Ref: ASCE 7-22 Sec. 2.3 (LRFD)  |  ACI 318-19 Table 5.3.1"
        )

        b_bot = float(_g(ss, "b_bottom", 1199.0))
        h     = float(_g(ss, "h", 200.0))
        wc    = float(_g(ss, "wc", 24.0))
        An    = float(_g(ss, "sp_An", 1.0))
        SW_HCS  = float(_g(ss, "SW_HCS", 0.0))
        SW_top  = float(_g(ss, "SW_topping", 0.0))
        SDL     = float(_g(ss, "SDL", 0.0))
        LL      = float(_g(ss, "LL", 0.0))
        wu_area = float(_g(ss, "lb_wu_area", 0.0))
        L_an    = float(_g(ss, "L_an", 5850.0))
        A_void  = float(_g(ss, "A_voids_total", 0.0))

        # 4.1 Self-Weight
        self._h2("4.1  Self-Weight Calculation")
        A_net_sw = b_bot * h - A_void
        SW_check = wc * A_net_sw / (b_bot * 1_000_000.0)
        self._calc(
            "  Formula:       A_net = b_bottom x h - A_voids_total"
        )
        self._calc(
            f"  Substitution:  A_net = {b_bot:.0f} x {h:.0f} - {A_void:.0f}"
        )
        self._calc(
            f"  Result:        A_net = {A_net_sw:,.0f} mm2"
        )
        self._calc("")
        self._calc(
            "  Formula:       SW = wc x A_net / (b_bottom x 1,000,000)   [kN/m2]"
        )
        self._calc(
            f"  Substitution:  SW = {wc:.2f} x {A_net_sw:,.0f} / ({b_bot:.0f} x 1,000,000)"
        )
        self._calc(
            f"  Result:        SW_HCS = {SW_HCS:.4f} kN/m2"
        )

        # 4.2 Factored Load Combination
        self._h2("4.2  Factored Load Combination")
        self._calc(
            "  Formula:       wu = 1.2 x (SW_HCS + SW_topping + SDL) + 1.6 x LL",
            "ASCE 7-22 Sec. 2.3"
        )
        self._calc(
            f"  Substitution:  wu = 1.2 x ({SW_HCS:.4f} + {SW_top:.4f} + {SDL:.4f}) + 1.6 x {LL:.4f}"
        )
        self._calc(
            f"                   = 1.2 x {SW_HCS + SW_top + SDL:.4f} + 1.6 x {LL:.4f}"
        )
        self._calc(
            f"  Result:        wu = {wu_area:.4f} kN/m2"
        )

        # 4.3 Line load & reactions
        wu_line = wu_area * b_bot / 1e6
        Ra_u    = float(_g(ss, "lb_Ra_u", wu_line * L_an / 2.0))
        self._h2("4.3  Support Reactions")
        self._calc(
            "  Formula:       wu_line = wu x b_bottom / 1,000,000   [kN/mm]"
        )
        self._calc(
            f"  Substitution:  wu_line = {wu_area:.4f} x {b_bot:.0f} / 1,000,000"
        )
        self._calc(
            f"  Result:        wu_line = {wu_line:.6f} kN/mm"
        )
        self._calc("")
        self._calc(
            "  Formula:       Ra = wu_line x L_an / 2   (simply supported)"
        )
        self._calc(
            f"  Substitution:  Ra = {wu_line:.6f} x {L_an:.0f} / 2"
        )
        self._calc(
            f"  Result:        Ra = {Ra_u:.3f} kN"
        )

        # 4.4 Critical Section Forces
        self._h2("4.4  Shear and Moment at Critical Sections")
        Vu_max = float(_g(ss, "lb_Vu_max", 0.0))
        Mu_max = float(_g(ss, "lb_Mu_max", 0.0)) / 1e6
        l_t    = float(_g(ss, "lb_l_t", 250.0))
        d_crit = float(_g(ss, "dp_bot", 165.0))

        self._calc(
            "  Formula:       Vu_max = Ra   (at face of support)"
        )
        self._calc(
            f"  Result:        Vu_max = {Vu_max:.3f} kN"
        )
        self._calc("")
        self._calc(
            "  Formula:       Mu_max = wu_line x L_an^2 / 8   (at midspan)"
        )
        self._calc(
            f"  Substitution:  Mu_max = {wu_line:.6f} x {L_an:.0f}^2 / 8"
        )
        self._calc(
            f"  Result:        Mu_max = {Mu_max:.4f} kN.m"
        )
        rows = _span_table(ss, n_seg=10)
        self._wide_table(
            ["x/L", "x (m)", "x (mm)", "Vu (kN)", "Mu (kN.m)"],
            [
                [f"{r['frac']:.1f}", f"{r['x_m']:.3f}", f"{r['x_mm']:.0f}",
                 f"{r['Vu_kN']:.3f}", f"{r['Mu_kNm']:.3f}"]
                for r in rows
            ]
        )

        # 4.5 SFD and BMD
        self._h2("4.5  Shear Force and Bending Moment Diagrams")
        self._note("Diagrams generated via matplotlib (reliable, no kaleido dependency).")
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 5 — Lifting Stage Check
    # =========================================================================
    def _ch5_lifting(self):
        ss = self.ss
        self._h1("5.  Lifting Stage Check")
        self._note(
            "Ref: PCI Design Handbook 8th Ed. Sec. 5.3\n"
            "Lifting checks use Pi (initial prestress before any losses) and f_ci.\n"
            "Sign convention: compression (-), tension (+)."
        )

        Pi    = float(_g(ss, "Pi",       237.8))
        L_an  = float(_g(ss, "L_an",    5850.0))
        b_bot = float(_g(ss, "b_bottom", 1199.0))
        An    = float(_g(ss, "sp_An",  174000.0))
        In    = float(_g(ss, "sp_In",     776e6))
        yb    = float(_g(ss, "sp_yb",    100.0))
        h     = float(_g(ss, "h",         200.0))
        e_net = float(_g(ss, "sp_e_net",   50.0))
        f_ci  = float(_g(ss, "f_ci",       35.0))
        SW_HCS = float(_g(ss, "SW_HCS",    0.0))

        yt = h - yb
        SW_line = SW_HCS * b_bot / 1e6   # kN/mm

        # Lifting point position
        a_lift_default = L_an / 5.0
        a_lift = float(_g(ss, "a_lift", a_lift_default))

        self._h2("5.1  Lifting Point Position")
        self._note(
            f"Default: L/5 from each end = {L_an:.0f}/5 = {a_lift_default:.0f} mm from support.\n"
            "User can override 'a_lift' in session_state."
        )
        self._calc(
            "  Formula:       a_lift = L_an / 5   [default]",
            "PCI Design Handbook 8th Ed. Sec. 5.3.2"
        )
        self._calc(
            f"  Substitution:  a_lift = {L_an:.0f} / 5"
        )
        self._calc(
            f"  Result:        a_lift = {a_lift:.0f} mm  (from each end)"
        )

        # Moments at lifting
        self._h2("5.2  Moments at Lifting")
        # Reaction at each lifting point (two-point lift, symmetric)
        Ra_lift = SW_line * L_an / 2.0
        M_neg   = SW_line * a_lift**2 / 2.0   # hogging at lift point  (kN.mm)
        M_pos   = Ra_lift * (L_an / 2.0) - SW_line * (L_an / 2.0)**2 / 2.0  # sagging at midspan
        M_neg_m = M_neg / 1e6
        M_pos_m = M_pos / 1e6

        self._calc(
            "  At lifting point (hogging):",
            "PCI 8th Ed. Sec. 5.3"
        )
        self._calc(
            "  Formula:       M_lift_neg = SW_line x a_lift^2 / 2"
        )
        self._calc(
            f"  Substitution:  M_lift_neg = {SW_line:.6f} x {a_lift:.0f}^2 / 2"
        )
        self._calc(
            f"  Result:        M_lift_neg = {M_neg_m:.4f} kN.m  (hogging, top in tension)"
        )
        self._calc("")
        self._calc(
            "  At midspan (sagging):"
        )
        self._calc(
            "  Formula:       M_lift_pos = Ra x L/2 - SW_line x (L/2)^2 / 2"
        )
        self._calc(
            f"  Substitution:  M_lift_pos = {Ra_lift:.6f} x {L_an/2:.0f}"
            f" - {SW_line:.6f} x {L_an/2:.0f}^2 / 2"
        )
        self._calc(
            f"  Result:        M_lift_pos = {M_pos_m:.4f} kN.m  (sagging)"
        )

        # Stress check at lifting point (worst: hogging, top in tension)
        self._h2("5.3  Stress Check at Lifting (at lifting point, hogging)")
        self._note("Uses Pi (initial prestress), f_ci limits, non-composite section.")
        comp_lim  = 0.60 * f_ci
        tens_lim  = 0.25 * math.sqrt(f_ci)
        Pi_N = Pi * 1000.0

        # At hogging section (M = M_neg, hogging → top +ve, bottom -ve)
        f_top_lift = (-Pi_N/An - Pi_N * e_net * yt / In
                      + M_neg * 1000.0 * yt / In)
        f_bot_lift = (-Pi_N/An + Pi_N * e_net * yb / In
                      - M_neg * 1000.0 * yb / In)

        self._calc(
            "  Formula:  f_top = -Pi/An - Pi*e*yt/In + M_neg*yt/In"
        )
        self._calc(
            f"  Term1:  -Pi/An    = -{Pi_N:.0f}/{An:.0f}  =  {-Pi_N/An:.4f} MPa"
        )
        self._calc(
            f"  Term2:  -Pi*e*yt/In = -{Pi_N:.0f}*{e_net:.2f}*{yt:.4f}/{In:.3e}  =  {-Pi_N*e_net*yt/In:.4f} MPa"
        )
        self._calc(
            f"  Term3:  +M*yt/In  = +{M_neg*1000:.0f}*{yt:.4f}/{In:.3e}  =  {M_neg*1000*yt/In:.4f} MPa"
        )
        self._calc(
            f"  Result: f_top = {-Pi_N/An:.4f} + {-Pi_N*e_net*yt/In:.4f} + {M_neg*1000*yt/In:.4f}"
        )
        self._calc(
            f"          f_top = {f_top_lift:.4f} MPa"
        )
        self._calc("")
        self._calc(
            "  Formula:  f_bot = -Pi/An + Pi*e*yb/In - M_neg*yb/In"
        )
        self._calc(
            f"  Result:   f_bot = {f_bot_lift:.4f} MPa"
        )

        ok_top = (-f_top_lift <= comp_lim if f_top_lift < 0 else True) and \
                 (f_top_lift <= tens_lim  if f_top_lift > 0 else True)
        ok_bot = (-f_bot_lift <= comp_lim if f_bot_lift < 0 else True) and \
                 (f_bot_lift <= tens_lim  if f_bot_lift > 0 else True)
        self._wide_table(
            ["Fibre", "Stress (MPa)", "Limit (MPa)", "Type", "Status"],
            [
                ["Top", f"{f_top_lift:.4f}",
                 f"-{comp_lim:.3f} / +{tens_lim:.4f}",
                 "compr/tens", "OK" if ok_top else "NG"],
                ["Bottom", f"{f_bot_lift:.4f}",
                 f"-{comp_lim:.3f} / +{tens_lim:.4f}",
                 "compr/tens", "OK" if ok_bot else "NG"],
            ]
        )
        self._status_line(f"Lifting stress check: {'OK' if (ok_top and ok_bot) else 'NG'}",
                          ok_top and ok_bot)
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 6 — Erection Stage Check
    # =========================================================================
    def _ch6_erection(self):
        ss = self.ss
        self._h1("6.  Erection Stage Check")
        self._note(
            "Erection stage: HCS in final position, prior to composite action.\n"
            "Loads: SW_HCS + SW_topping (wet) + construction live load (w_construction_LL).\n"
            "Non-composite section used.  Uses Pe (effective after initial losses).\n"
            "Ref: PCI Design Handbook 8th Ed. Sec. 5.3"
        )

        Pe     = float(_g(ss, "pl_Pe",    200.0))
        An     = float(_g(ss, "sp_An", 174000.0))
        In     = float(_g(ss, "sp_In",    776e6))
        yb     = float(_g(ss, "sp_yb",   100.0))
        h      = float(_g(ss, "h",        200.0))
        e_net  = float(_g(ss, "sp_e_net",  50.0))
        f_ci   = float(_g(ss, "f_ci",      35.0))
        f_c    = float(_g(ss, "f_c",       50.0))
        b_bot  = float(_g(ss, "b_bottom", 1199.0))
        L_an   = float(_g(ss, "L_an",    5850.0))
        SW_HCS = float(_g(ss, "SW_HCS",    0.0))
        SW_top = float(_g(ss, "SW_topping", 0.0))
        w_cll  = float(_g(ss, "w_construction_LL", 1.0))  # DEFAULT 1.0 kN/m2

        yt   = h - yb
        w_ere = (SW_HCS + SW_top + w_cll) * b_bot / 1e6  # kN/mm
        M_ere = w_ere * L_an**2 / 8.0   # kN.mm
        M_ere_m = M_ere / 1e6

        self._h2("6.1  Erection Loads")
        self._note(
            f"w_construction_LL = {w_cll:.2f} kN/m2  "
            "[DEFAULT 1.0 kN/m2 per PCI; set via session_state['w_construction_LL']]"
        )
        self._calc(
            "  Formula:       w_erection = (SW_HCS + SW_topping + w_construction_LL) x b_bottom / 1e6   [kN/mm]"
        )
        self._calc(
            f"  Substitution:  w_erection = ({SW_HCS:.4f} + {SW_top:.4f} + {w_cll:.4f}) x {b_bot:.0f} / 1e6"
        )
        self._calc(
            f"  Result:        w_erection = {w_ere:.6f} kN/mm"
        )
        self._calc("")
        self._calc(
            "  Formula:       M_erection = w_erection x L_an^2 / 8"
        )
        self._calc(
            f"  Substitution:  M_erection = {w_ere:.6f} x {L_an:.0f}^2 / 8"
        )
        self._calc(
            f"  Result:        M_erection = {M_ere_m:.4f} kN.m"
        )

        self._h2("6.2  Stress Check at Erection (non-composite, midspan)")
        comp_lim = 0.60 * f_c
        tens_lim = 0.50 * math.sqrt(f_c)
        Pe_N = Pe * 1000.0

        f_top = (-Pe_N/An - Pe_N * e_net * yt / In
                 + M_ere * 1000.0 * yt / In)
        f_bot = (-Pe_N/An + Pe_N * e_net * yb / In
                 - M_ere * 1000.0 * yb / In)

        self._calc(
            "  Formula:  f_top = -Pe/An - Pe*e*yt/In + M_ere*yt/In"
        )
        self._calc(
            f"  Term1: -Pe/An   = -{Pe_N:.0f}/{An:.0f}  =  {-Pe_N/An:.4f} MPa"
        )
        self._calc(
            f"  Term2: -Pe*e*yt/In = -{Pe_N:.0f}*{e_net:.2f}*{yt:.4f}/{In:.3e}  =  {-Pe_N*e_net*yt/In:.4f} MPa"
        )
        self._calc(
            f"  Term3: +M*yt/In = +{M_ere*1000:.0f}*{yt:.4f}/{In:.3e}  =  {M_ere*1000*yt/In:.4f} MPa"
        )
        self._calc(
            f"  Result: f_top = {f_top:.4f} MPa"
        )
        self._calc("")
        self._calc(
            "  Formula:  f_bot = -Pe/An + Pe*e*yb/In - M_ere*yb/In"
        )
        self._calc(
            f"  Term1: -Pe/An   = {-Pe_N/An:.4f} MPa"
        )
        self._calc(
            f"  Term2: +Pe*e*yb/In = {Pe_N*e_net*yb/In:.4f} MPa"
        )
        self._calc(
            f"  Term3: -M*yb/In = {-M_ere*1000*yb/In:.4f} MPa"
        )
        self._calc(
            f"  Result: f_bot = {f_bot:.4f} MPa"
        )

        ok_top = (-f_top <= comp_lim if f_top < 0 else True) and \
                 (f_top <= tens_lim  if f_top > 0 else True)
        ok_bot = (-f_bot <= comp_lim if f_bot < 0 else True) and \
                 (f_bot <= tens_lim  if f_bot > 0 else True)
        self._wide_table(
            ["Fibre", "Stress (MPa)", "Comp limit (MPa)", "Tens limit (MPa)", "Status"],
            [
                ["Top",    f"{f_top:.4f}", f"-{comp_lim:.3f}", f"+{tens_lim:.4f}", "OK" if ok_top else "NG"],
                ["Bottom", f"{f_bot:.4f}", f"-{comp_lim:.3f}", f"+{tens_lim:.4f}", "OK" if ok_bot else "NG"],
            ]
        )
        self._status_line(f"Erection stress check: {'OK' if (ok_top and ok_bot) else 'NG'}",
                          ok_top and ok_bot)

        self._h2("6.3  Capacity Diagram at Erection")
        self._note(
            "phi*Mn vs Mu diagram shown in Chapter 9 (Flexural Capacity).\n"
            "Erection Mu is less critical than service; diagram covers both."
        )
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 7 — Prestress Losses & Development Length
    # =========================================================================
    def _ch7_losses(self):
        ss = self.ss
        self._h1("7.  Prestress Losses and Development Length")
        self._note(
            "Method: PCI Lump-Sum  (ES + CR + SH + RE)\n"
            "Ref: PCI Design Handbook 8th Ed. Sec. 4.7  |  ACI/PCI CODE-319-25 Cl. 26.10"
        )

        fpi    = float(_g(ss, "fpi",     1213.5))
        Pi     = float(_g(ss, "Pi",       237.8))
        Aps_b  = float(_g(ss, "Aps_bot",  196.0))
        Aps_t  = float(_g(ss, "Aps_top",    0.0))
        An     = float(_g(ss, "sp_An",  174000.0))
        In     = float(_g(ss, "sp_In",    776e6))
        e_net  = float(_g(ss, "sp_e_net",  68.35))
        f_ci   = float(_g(ss, "f_ci",      35.0))
        Ec_hcs = float(_g(ss, "Ec_hcs", 36793.0))
        Eps    = float(_g(ss, "Eps",   199050.0))
        RH     = float(_g(ss, "RH",       75.0))
        V_S    = float(_g(ss, "V_S",      38.0))
        ES     = float(_g(ss, "pl_ES",     0.0))
        CR     = float(_g(ss, "pl_CR",     0.0))
        SH     = float(_g(ss, "pl_SH",     0.0))
        RE     = float(_g(ss, "pl_RE",     0.0))
        total  = float(_g(ss, "pl_total_MPa", ES+CR+SH+RE))
        pct    = float(_g(ss, "pl_total_pct", total/fpi*100 if fpi>0 else 0))
        fse    = float(_g(ss, "pl_fse",   fpi-total))
        Pe     = float(_g(ss, "pl_Pe",    0.0))
        n_ps   = Eps / Ec_hcs

        # 7.1 Elastic Shortening
        self._h2("7.1  Elastic Shortening  (ES)")
        self._note("Concrete stress at PS centroid immediately after transfer:")
        Pi_N = Pi * 1000.0
        f_cir = Pi_N/An + Pi_N * e_net**2 / In
        self._calc(
            "  Formula:  n_ps = Eps / Ec_ci",
            "ACI 318-19 Eq. 19.2.2.1"
        )
        self._calc(
            f"  Substitution:  n_ps = {Eps:.0f} / {Ec_hcs:.0f}"
        )
        self._calc(
            f"  Result:        n_ps = {n_ps:.4f}"
        )
        self._calc("")
        self._calc(
            "  Formula:  f_cir = Pi/An + Pi*e_net^2/In   (SW moment omitted for estimate)"
        )
        self._calc(
            f"  Substitution:  f_cir = {Pi_N:.0f}/{An:.0f} + {Pi_N:.0f}*{e_net:.2f}^2/{In:.3e}"
        )
        self._calc(
            f"  Result:        f_cir = {Pi_N/An:.4f} + {Pi_N*e_net**2/In:.4f}  =  {f_cir:.4f} MPa"
        )
        self._calc("")
        self._calc(
            "  Formula:  ES = n_ps * 0.5 * f_cir",
            "PCI 8th Ed. Sec. 4.7.2  (0.5 factor for pretensioned members)"
        )
        self._calc(
            f"  Substitution:  ES = {n_ps:.4f} * 0.5 * {f_cir:.4f}"
        )
        self._calc(
            f"  Result:        ES = {ES:.3f} MPa"
        )

        # 7.2 Creep
        self._h2("7.2  Creep Loss  (CR)")
        self._note("f_cds = stress at PS centroid due to loads applied after transfer (SDL).")
        self._calc(
            "  Formula:  CR = Kcr * (Eps/Ec) * (f_cir - f_cds)",
            "PCI 8th Ed. Eq. 4.7.3  (Kcr = 2.0 for NW concrete)"
        )
        self._calc(
            f"  Result:   CR = {CR:.3f} MPa"
        )

        # 7.3 Shrinkage
        self._h2("7.3  Shrinkage Loss  (SH)")
        self._calc(
            "  Formula:  SH = 8.2e-6 * Ksh * Eps * (1 - 0.06 * V/S) * (100 - RH)",
            "PCI 8th Ed. Eq. 4.7.4  (Ksh = 1.0)"
        )
        self._calc(
            f"  Substitution:  SH = 8.2e-6 * 1.0 * {Eps:.0f} * (1 - 0.06 * {V_S:.1f}) * (100 - {RH:.0f})"
        )
        self._calc(
            f"  Result:        SH = {SH:.3f} MPa"
        )

        # 7.4 Relaxation
        self._h2("7.4  Relaxation Loss  (RE)")
        self._note("Kre, J, C factors per PCI Table 4.7.5 (depend on prestress type and fpi/fpu).")
        self._calc(
            "  Formula:  RE = [Kre - J * (ES + CR + SH)] * C",
            "PCI 8th Ed. Eq. 4.7.5"
        )
        self._calc(
            f"  Result:   RE = {RE:.3f} MPa"
        )

        # 7.5 Totals
        self._h2("7.5  Total Loss and Effective Prestress")
        self._calc(
            "  Formula:  Total = ES + CR + SH + RE"
        )
        self._calc(
            f"  Substitution:  Total = {ES:.3f} + {CR:.3f} + {SH:.3f} + {RE:.3f}"
        )
        self._calc(
            f"  Result:        Total = {total:.3f} MPa   ({pct:.2f} % of fpi = {fpi:.1f} MPa)"
        )
        self._calc("")
        self._calc(
            "  Formula:  fse = fpi - Total"
        )
        self._calc(
            f"  Substitution:  fse = {fpi:.1f} - {total:.3f}"
        )
        self._calc(
            f"  Result:        fse = {fse:.3f} MPa"
        )
        self._calc("")
        self._calc(
            "  Formula:  Pe = (Aps_bot + Aps_top) * fse / 1000   [kN]"
        )
        self._calc(
            f"  Substitution:  Pe = ({Aps_b:.1f} + {Aps_t:.1f}) * {fse:.3f} / 1000"
        )
        self._calc(
            f"  Result:        Pe = {Pe:.3f} kN"
        )

        self._wide_table(
            ["Loss Component", "Value (MPa)", "% of fpi"],
            [
                ["Elastic shortening  ES", f"{ES:.3f}",    f"{ES/fpi*100:.2f}" if fpi>0 else "--"],
                ["Creep               CR", f"{CR:.3f}",    f"{CR/fpi*100:.2f}" if fpi>0 else "--"],
                ["Shrinkage           SH", f"{SH:.3f}",    f"{SH/fpi*100:.2f}" if fpi>0 else "--"],
                ["Relaxation          RE", f"{RE:.3f}",    f"{RE/fpi*100:.2f}" if fpi>0 else "--"],
                ["TOTAL",                  f"{total:.3f}", f"{pct:.2f}"],
                ["fse  (effective MPa)",   f"{fse:.3f}",   "--"],
                ["Pe   (effective kN)",    f"{Pe:.3f}",    "--"],
            ]
        )

        # 7.6 Development & Transfer Length
        self._h2("7.6  Transfer Length  l_t")
        ps_type = _g(ss, "ps_type", "PC Wire")
        l_t   = float(_g(ss, "lb_l_t",      250.0))
        l_d   = float(_g(ss, "lb_l_d",      380.0))
        fps   = float(_g(ss, "lb_fps_est",  1500.0))
        fpu   = float(_g(ss, "fpu",         1618.0))
        fpy   = float(_g(ss, "fpy",         1432.0))
        L_an  = float(_g(ss, "L_an",        5850.0))
        d_ps  = float(_g(ss, "wire_dia", 5.0)) if "Wire" in str(ps_type) else l_t / 50.0

        if "Wire" in str(ps_type):
            self._calc(
                "  Formula:  l_t = 50 * d_ps   [PC Wire]",
                "PCI Design Handbook 8th Ed. Sec. 4.2.3"
            )
            self._calc(
                f"  Substitution:  l_t = 50 * {d_ps:.1f}"
            )
            self._calc(
                f"  Result:        l_t = {l_t:.1f} mm"
            )
        else:
            self._calc(
                "  Formula:  l_t = max(60*d_ps, fse/20.7*d_ps)   [7-wire strand]",
                "ACI 318-19 Sec. 25.8.6.1"
            )
            self._calc(
                f"  Result:  l_t = {l_t:.1f} mm"
            )

        self._h2("7.7  Development Length  l_d")
        self._calc(
            "  Formula:  fps = min(fpu, fpy + 70)   [preliminary estimate]",
            "ACI 318-19 Sec. 20.3.2.4 (conservative)"
        )
        self._calc(
            f"  Substitution:  fps = min({fpu:.1f}, {fpy:.1f} + 70)"
        )
        self._calc(
            f"  Result:        fps = {fps:.1f} MPa"
        )
        self._calc("")
        self._calc(
            "  Formula:  l_d = l_t + (fps - fse) * d_ps / 20.7",
            "ACI 318-19 Eq. 25.8.7.1 (SI)"
        )
        self._calc(
            f"  Substitution:  l_d = {l_t:.1f} + ({fps:.1f} - {fse:.1f}) * {d_ps:.1f} / 20.7"
        )
        self._calc(
            f"  Result:        l_d = {l_t:.1f} + {(fps-fse)*d_ps/20.7:.1f}  =  {l_d:.1f} mm"
        )
        ratio  = L_an / l_d if l_d > 0 else 0.0
        status = _g(ss, "lb_ps_status", "FULL")
        self._calc(
            f"  L_an / l_d = {L_an:.0f} / {l_d:.1f} = {ratio:.3f}",
            "ACI 318-19 Sec. 25.8.7"
        )
        self._status_line(
            f"Prestress development: {status}  (L_an/l_d = {ratio:.3f})",
            ratio >= 1.0
        )
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 8 — Service Stress Checks
    # =========================================================================
    def _ch8_service_stress(self):
        ss = self.ss
        self._h1("8.  Service Stress Checks")
        self._note(
            "Sign convention:  compression (-)  tension (+)\n"
            "Code limits — ACI/PCI CODE-319-25 Table 24.5.3.1:\n"
            "  Compression at transfer:  -0.60 f_ci\n"
            "  Tension at transfer:      +0.25 sqrt(f_ci)\n"
            "  Compression at service:   -0.45 f_c (sustained)  /  -0.60 f_c (total)\n"
            "  Tension at service:       +0.50 sqrt(f_c) (Class T)  or 0 (Class U)"
        )

        Pe    = float(_g(ss, "pl_Pe",    200.0))
        An    = float(_g(ss, "sp_An", 174000.0))
        In    = float(_g(ss, "sp_In",    776e6))
        yb    = float(_g(ss, "sp_yb",   100.0))
        h     = float(_g(ss, "h",        200.0))
        e_net = float(_g(ss, "sp_e_net",  50.0))
        Sb_n  = float(_g(ss, "sp_Sb",   1.0))
        St_n  = float(_g(ss, "sp_St",   1.0))
        f_ci  = float(_g(ss, "f_ci",    35.0))
        f_c   = float(_g(ss, "f_c",     50.0))
        Pi    = float(_g(ss, "Pi",      237.8))
        b_bot = float(_g(ss, "b_bottom", 1199.0))
        L_an  = float(_g(ss, "L_an",   5850.0))

        yt    = h - yb
        w_sw  = float(_g(ss, "SW_HCS",   0.0)) * b_bot / 1e6
        Mg_sw = w_sw * L_an**2 / 8.0  # kN.mm
        w_top = float(_g(ss, "SW_topping", 0.0)) * b_bot / 1e6 if _g(ss, "has_topping") else 0.0
        Mg_top = w_top * L_an**2 / 8.0
        w_sdl = float(_g(ss, "SDL",  0.0)) * b_bot / 1e6
        M_sdl = w_sdl * L_an**2 / 8.0
        w_ll  = float(_g(ss, "LL",   0.0)) * b_bot / 1e6
        M_ll  = w_ll  * L_an**2 / 8.0

        stages = [
            ("8.1  Transfer / Release",    "sc_transfer",    "ACI/PCI 319-25 Table 24.5.3.1"),
            ("8.2  Lifting",               "sc_lifting",     "PCI Design Handbook 8th Ed. Sec. 5.3"),
            ("8.3  Erection / Construction","sc_construction","ACI/PCI 319-25 Table 24.5.3.1"),
            ("8.4  Service",               "sc_service",     "ACI/PCI 319-25 Table 24.5.3.1"),
        ]

        summary_rows = []
        for title, key, ref in stages:
            self._h2(title)
            self._note(f"Ref: {ref}")

            d = ss.get(key, {})
            if not d:
                self._note("(data not yet available — run calculations first)")
                summary_rows.append([title.split(" ")[1], "--", "--", "--", "--", "N/A"])
                continue

            f_top    = float(d.get("f_top",     0.0))
            f_bot    = float(d.get("f_bot",     0.0))
            lim_comp = float(d.get("limit_comp", 0.0))
            lim_tens = float(d.get("limit_tens", 0.0))
            status   = str(d.get("status", "N/A"))

            # Show step-by-step for each stage
            if key == "sc_transfer":
                Pi_N = Pi * 1000.0
                self._calc(
                    "  Formula:  f_top = -Pi/An - Pi*e*yt/In + M_sw*yt/In"
                )
                t1 = -Pi_N / An
                t2 = -Pi_N * e_net * yt / In
                t3 = Mg_sw * 1000.0 * yt / In
                self._calc(
                    f"  Substitution:  f_top = -{Pi_N:.0f}/{An:.0f}"
                    f" - {Pi_N:.0f}*{e_net:.2f}*{yt:.4f}/{In:.3e}"
                    f" + {Mg_sw*1000:.0f}*{yt:.4f}/{In:.3e}"
                )
                self._calc(
                    f"               f_top = {t1:.4f} + {t2:.4f} + {t3:.4f}"
                )
                self._calc(
                    f"  Result:        f_top = {f_top:.4f} MPa"
                )
                self._calc("")
                self._calc(
                    "  Formula:  f_bot = -Pi/An + Pi*e*yb/In - M_sw*yb/In"
                )
                self._calc(
                    f"  Result:        f_bot = {f_bot:.4f} MPa"
                )
            else:
                Pe_N = Pe * 1000.0
                if key == "sc_service":
                    self._calc(
                        "  Formula:  f_top = -Pe/An + Pe*e*yt/In - M_DL*yt/In - M_super*(yt_comp)/I_comp"
                    )
                else:
                    self._calc(
                        "  Formula:  f_top = -Pe/An + Pe*e*yt/In - M_total*yt/In"
                    )
                t1 = -Pe_N / An
                t2 = Pe_N * e_net * yt / In
                # approximate term3 from result back
                self._calc(
                    f"  Term1: -Pe/An   = -{Pe_N:.0f}/{An:.0f}  =  {t1:.4f} MPa"
                )
                self._calc(
                    f"  Term2: +Pe*e*yt/In = +{Pe_N:.0f}*{e_net:.2f}*{yt:.4f}/{In:.3e}  =  {t2:.4f} MPa"
                )
                self._calc(
                    f"  Result:        f_top = {f_top:.4f} MPa"
                )
                self._calc("")
                self._calc(
                    f"  Result:        f_bot = {f_bot:.4f} MPa"
                )

            self._calc(
                f"  Limit (compression):  {-lim_comp:.3f} MPa"
            )
            self._calc(
                f"  Limit (tension):      +{lim_tens:.4f} MPa"
            )
            ok = status.upper() in ("OK", "PASS")
            self._status_line(f"Status: {status}", ok)
            summary_rows.append([
                title.split("  ")[0].strip(),
                f"{f_top:.3f}",
                f"{f_bot:.3f}",
                f"-{lim_comp:.3f}",
                f"+{lim_tens:.4f}",
                status,
            ])

        self._h2("8.5  Stress Check Summary")
        self._wide_table(
            ["Stage", "f_top (MPa)", "f_bot (MPa)", "Comp limit", "Tens limit", "Status"],
            summary_rows
        )
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 9 — Flexural & Shear Capacity
    # =========================================================================
    def _ch9_capacity(self, img_phi_mn: bytes | None, img_phi_vn: bytes | None):
        ss = self.ss
        self._h1("9.  Flexural and Shear Capacity")
        self._note(
            "Ref: ACI/PCI CODE-319-25 Cl. 22.2 (flexure)  |  Cl. 22.5 (shear)\n"
            "phi = 0.90 (flexure, tension-controlled)  |  phi = 0.75 (shear)"
        )

        fps    = float(_g(ss, "cap_fps",       1400.0))
        a      = float(_g(ss, "cap_a",           30.0))
        Mn     = float(_g(ss, "cap_Mn",         200.0))
        phi_Mn = float(_g(ss, "cap_phi_Mn",     180.0))
        Mu_max = float(_g(ss, "lb_Mu_max",        0.0)) / 1e6
        DCR_M  = float(_g(ss, "cap_DCR_M",        1.0))
        phi_Vn = float(_g(ss, "cap_phi_Vn_min",  50.0))
        Vu_max = float(_g(ss, "lb_Vu_max",        0.0))
        DCR_V  = float(_g(ss, "cap_DCR_V",        1.0))
        Aps_b  = float(_g(ss, "Aps_bot",         196.0))
        dp_bot = float(_g(ss, "dp_bot",          165.0))
        f_c    = float(_g(ss, "f_c",              50.0))
        b_top  = float(_g(ss, "b_top",          1187.0))
        fpu    = float(_g(ss, "fpu",            1618.0))
        fpy    = float(_g(ss, "fpy",            1432.0))
        rho_p  = Aps_b / (b_top * dp_bot) if b_top * dp_bot > 0 else 0.0
        beta1  = (0.85 if f_c <= 28 else
                  max(0.65, 0.85 - 0.05 * (f_c - 28) / 7))
        needs_Av = bool(_g(ss, "cap_needs_Av_min", False))

        # 9.1 Flexural Capacity
        self._h2("9.1  Flexural Capacity  Mn  (Whitney stress block)")
        self._calc(
            "  Formula:  beta1 = 0.85 - 0.05*(f_c - 28)/7   [for f_c > 28 MPa]",
            "ACI 318-19 Table 22.2.2.4.3"
        )
        self._calc(
            f"  Substitution:  beta1 = 0.85 - 0.05*({f_c:.0f} - 28)/7"
        )
        self._calc(
            f"  Result:        beta1 = {beta1:.4f}"
        )
        self._calc("")
        self._calc(
            "  Formula:  rho_p = Aps_bot / (b_top * dp_bot)"
        )
        self._calc(
            f"  Substitution:  rho_p = {Aps_b:.1f} / ({b_top:.0f} * {dp_bot:.0f})"
        )
        self._calc(
            f"  Result:        rho_p = {rho_p:.6f}"
        )
        self._calc("")
        gamma_p = 0.28
        self._calc(
            "  Formula:  fps = fpu * (1 - gamma_p/beta1 * rho_p * fpu / f_c)   [gamma_p=0.28]",
            "ACI 318-19 Eq. 20.3.2.4.1  (bonded tendons)"
        )
        term = (rho_p * fpu) / (beta1 * f_c) * gamma_p
        self._calc(
            f"  Substitution:  fps = {fpu:.0f} * (1 - 0.28/{beta1:.4f} * {rho_p:.6f} * {fpu:.0f} / {f_c:.0f})"
        )
        self._calc(
            f"             = {fpu:.0f} * (1 - {term:.6f})"
        )
        self._calc(
            f"  Result:        fps = {fps:.2f} MPa  (but max = fpy = {fpy:.0f} MPa)"
        )
        self._calc("")
        self._calc(
            "  Formula:  a = Aps_bot * fps / (0.85 * f_c * b_top)",
            "ACI 318-19 Eq. 22.2.2.4.1  (compression block depth)"
        )
        self._calc(
            f"  Substitution:  a = {Aps_b:.1f} * {fps:.2f} / (0.85 * {f_c:.0f} * {b_top:.0f})"
        )
        self._calc(
            f"              = {Aps_b * fps:.1f} / {0.85 * f_c * b_top:.1f}"
        )
        self._calc(
            f"  Result:        a = {a:.4f} mm"
        )
        self._calc("")
        self._calc(
            "  Formula:  Mn = Aps_bot * fps * (dp_bot - a/2) / 1e6   [kN.m]",
            "ACI 318-19 Eq. 22.2.1.1"
        )
        self._calc(
            f"  Substitution:  Mn = {Aps_b:.1f} * {fps:.2f} * ({dp_bot:.0f} - {a:.4f}/2) / 1e6"
        )
        self._calc(
            f"              = {Aps_b:.1f} * {fps:.2f} * {dp_bot - a/2:.4f} / 1e6"
        )
        self._calc(
            f"  Result:        Mn = {Mn:.4f} kN.m"
        )
        self._calc("")
        self._calc(
            "  Formula:  phi*Mn = 0.90 * Mn",
            "ACI 318-19 Table 21.2.2  (phi = 0.90, tension-controlled)"
        )
        self._calc(
            f"  Substitution:  phi*Mn = 0.90 * {Mn:.4f}"
        )
        self._calc(
            f"  Result:        phi*Mn = {phi_Mn:.4f} kN.m"
        )
        self._calc("")
        self._calc(
            f"  Mu_max (factored demand) = {Mu_max:.4f} kN.m"
        )
        self._calc(
            "  Formula:  DCR_M = Mu_max / phi*Mn"
        )
        self._calc(
            f"  Substitution:  DCR_M = {Mu_max:.4f} / {phi_Mn:.4f}"
        )
        self._calc(
            f"  Result:        DCR_M = {DCR_M:.4f}  ({'<= 1.00  OK' if DCR_M <= 1.0 else '> 1.00  NG'})"
        )
        self._status_line(
            f"Flexure: DCR = {DCR_M:.4f}  =>  {'OK' if DCR_M<=1.0 else 'OVERSTRESS'}",
            DCR_M <= 1.0
        )

        self._h2("9.1.1  phi*Mn vs Mu Diagram")
        self._insert_image(img_phi_mn, width_mm=162)

        # 9.2 Shear Capacity
        self._h2("9.2  Shear Capacity  Vn")
        self._note(
            "HCS has unique shear behaviour:\n"
            "  - Web-shear (Vcw) controls near supports (full prestress, low flexure)\n"
            "  - Flexure-shear (Vci) controls near midspan (reduced prestress, high flexure)\n"
            "  phi = 0.75 per ACI/PCI CODE-319-25 Cl. 22.5.1.1\n"
            "Ref: ACI/PCI CODE-319-25 Cl. 22.5.8  |  PCI Design Handbook 8th Ed. Sec. 5.2"
        )
        Pe     = float(_g(ss, "pl_Pe",    200.0))
        An_sh  = float(_g(ss, "sp_An",  174000.0))
        fpc    = Pe * 1000.0 / An_sh if An_sh > 0 else 0.0
        bw     = float(_g(ss, "bw_shear", 300.0))

        self._calc(
            "  Formula:  fpc = Pe * 1000 / An   (concrete compressive stress at PS centroid)"
        )
        self._calc(
            f"  Substitution:  fpc = {Pe:.3f} * 1000 / {An_sh:.0f}"
        )
        self._calc(
            f"  Result:        fpc = {fpc:.4f} MPa"
        )
        self._calc("")
        sqrt_fc = math.sqrt(f_c)
        self._calc(
            "  Formula:  Vcw = (0.29 * sqrt(f_c) + 0.3 * fpc) * bw * dp_bot / 1000   [kN]",
            "ACI 318-19 Eq. 22.5.8.3.2"
        )
        self._calc(
            f"  Substitution:  Vcw = (0.29 * sqrt({f_c:.0f}) + 0.3 * {fpc:.4f}) * {bw:.0f} * {dp_bot:.0f} / 1000"
        )
        term_vcw = (0.29 * sqrt_fc + 0.3 * fpc) * bw * dp_bot / 1000.0
        self._calc(
            f"              = (0.29 * {sqrt_fc:.4f} + {0.3*fpc:.4f}) * {bw:.0f} * {dp_bot:.0f} / 1000"
        )
        self._calc(
            f"  Result:        Vcw ~ {term_vcw:.3f} kN  (at support region)"
        )
        self._calc("")
        self._calc(
            "  Formula:  Vci = 0.05*sqrt(f_c)*bw*dp + Vd + Vi*Mcre/Mmax  (simplified)",
            "ACI 318-19 Eq. 22.5.8.3.1"
        )
        self._calc(
            f"  Minimum phi*Vn along span = {phi_Vn:.4f} kN   (envelope of Vci and Vcw)"
        )
        self._calc(
            f"  Vu_max (factored demand)  = {Vu_max:.4f} kN"
        )
        self._calc(
            "  Formula:  DCR_V = Vu_max / min(phi*Vn)"
        )
        self._calc(
            f"  Substitution:  DCR_V = {Vu_max:.4f} / {phi_Vn:.4f}"
        )
        self._calc(
            f"  Result:        DCR_V = {DCR_V:.4f}  ({'<= 1.00  OK' if DCR_V <= 1.0 else '> 1.00  NG'})"
        )
        if needs_Av:
            self._warn(
                "h > 317 mm, no topping, and Vu > 0.5*phi*Vcw  =>  "
                "minimum shear reinforcement Av,min REQUIRED per ACI/PCI CODE-319-25 Cl. 9.6.3."
            )
        self._status_line(
            f"Shear: DCR = {DCR_V:.4f}  =>  {'OK' if DCR_V<=1.0 else 'OVERSTRESS'}",
            DCR_V <= 1.0
        )

        self._h2("9.2.1  phi*Vn vs Vu Envelope Diagram")
        self._insert_image(img_phi_vn, width_mm=162)

        # 9.3 Capacity Summary
        self._h2("9.3  Capacity Summary")
        self._wide_table(
            ["Check", "Demand", "Capacity phi*R", "DCR", "Status"],
            [
                ["Flexure  Mn", f"{Mu_max:.3f} kN.m", f"{phi_Mn:.3f} kN.m",
                 f"{DCR_M:.4f}", "OK" if DCR_M<=1.0 else "NG"],
                ["Shear    Vn", f"{Vu_max:.3f} kN",   f"{phi_Vn:.3f} kN",
                 f"{DCR_V:.4f}", "OK" if DCR_V<=1.0 else "NG"],
            ]
        )
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 10 — Deflection & Camber
    # =========================================================================
    def _ch10_deflection(self, img_defl: bytes | None):
        ss = self.ss
        self._h1("10.  Deflection and Camber")
        self._note(
            "Ref: PCI Design Handbook 8th Ed. Sec. 4.8 and Table 4.8.3\n"
            "     ACI 318-19 Table 24.2.2  (code limits)\n"
            "Method: Elastic formulas for instantaneous; PCI multipliers for long-term.\n"
            "Sign: upward camber (+), downward deflection (-) in formulas; magnitude in table."
        )

        Pe       = float(_g(ss, "pl_Pe",                200.0))
        e_bot    = float(_g(ss, "sp_e_bot",              68.35))
        In       = float(_g(ss, "sp_In",                776e6))
        Ec_hcs   = float(_g(ss, "Ec_hcs",            36793.0))
        L_an     = float(_g(ss, "L_an",               5850.0))
        b_bot    = float(_g(ss, "b_bottom",           1199.0))
        SW_HCS   = float(_g(ss, "SW_HCS",               0.0))
        delta_ps = float(_g(ss, "def_delta_ps_initial",  0.0))
        delta_sw = float(_g(ss, "def_delta_sw",           0.0))
        net_rel  = float(_g(ss, "def_net_release",        0.0))
        total_lt = float(_g(ss, "def_total_longterm",     0.0))
        lim_ll   = float(_g(ss, "def_limit_ll_mm",        0.0))
        lim_tot  = float(_g(ss, "def_limit_total_mm",     0.0))
        st_ll    = str  (_g(ss, "def_status_ll",        "N/A"))
        st_tot   = str  (_g(ss, "def_status_total",     "N/A"))

        # PCI multipliers
        mult_cam = float(_g(ss, "mult_camber_final", 2.20))
        mult_sw  = float(_g(ss, "mult_dw_final",     2.40))
        mult_sdl = float(_g(ss, "mult_sdl_final",    1.80))

        final_cam = float(_g(ss, "def_final_camber",  0.0))
        final_sw  = float(_g(ss, "def_final_sw",      0.0))
        final_sdl = float(_g(ss, "def_final_sdl",     0.0))
        final_ll  = float(_g(ss, "def_final_ll",      0.0))

        w_sw  = SW_HCS * b_bot / 1e6  # kN/mm

        # 10.1 Initial Camber
        self._h2("10.1  Initial Prestress Camber at Release  (upward, straight tendon)")
        self._calc(
            "  Formula:  delta_ps = Pe * e_bot * L_an^2 / (8 * Ec_ci * In)",
            "PCI Design Handbook 8th Ed. Eq. 4.8.1  (straight tendon)"
        )
        self._calc(
            f"  Substitution:  delta_ps = {Pe*1000:.0f} * {e_bot:.4f} * {L_an:.0f}^2"
            f"  /  (8 * {Ec_hcs:.0f} * {In:.3e})"
        )
        self._calc(
            f"  Result:        delta_ps = {delta_ps:.4f} mm  (upward, positive)"
        )

        # 10.2 Self-weight
        self._h2("10.2  Self-Weight Deflection at Release  (downward)")
        self._calc(
            "  Formula:  delta_sw = 5 * w_sw * L_an^4 / (384 * Ec_ci * In)",
            "PCI Design Handbook 8th Ed. Eq. 4.8.2"
        )
        self._calc(
            f"  Substitution:  delta_sw = 5 * {w_sw:.6f}*1000 * {L_an:.0f}^4"
            f"  /  (384 * {Ec_hcs:.0f} * {In:.3e})"
        )
        self._calc(
            f"  Result:        delta_sw = {delta_sw:.4f} mm  (downward, negative)"
        )

        # 10.3 Net at release
        self._h2("10.3  Net Deflection at Release")
        self._calc(
            "  Formula:  net_release = delta_sw + delta_ps   (+upward, -downward)"
        )
        self._calc(
            f"  Substitution:  net_release = {delta_sw:.4f} + {delta_ps:.4f}"
        )
        self._calc(
            f"  Result:        net_release = {net_rel:.4f} mm"
            "   (negative = net upward camber; positive = net sag)"
        )

        # 10.4 Long-term with PCI multipliers
        self._h2("10.4  Long-term Deflection  (PCI Multipliers)")
        self._note("Ref: PCI Design Handbook Table 4.8.3")
        self._wide_table(
            ["Component", "PCI Multiplier", "Instantaneous (mm)", "Long-term (mm)"],
            [
                ["Prestress camber", f"{mult_cam:.2f}", f"{delta_ps:.4f}", f"{final_cam:.4f}"],
                ["Self-weight",       f"{mult_sw:.2f}",  f"{delta_sw:.4f}", f"{final_sw:.4f}"],
                ["SDL",               f"{mult_sdl:.2f}", "--",              f"{final_sdl:.4f}"],
                ["LL (transient)",    "1.00",            "--",              f"{final_ll:.4f}"],
                ["NET TOTAL",         "--",              f"{net_rel:.4f}",  f"{total_lt:.4f}"],
            ]
        )
        self._calc(
            "  Formula:  delta_total = final_camber + final_sw + final_sdl + final_ll"
        )
        self._calc(
            f"  Substitution:  delta_total = {final_cam:.4f} + {final_sw:.4f}"
            f" + {final_sdl:.4f} + {final_ll:.4f}"
        )
        self._calc(
            f"  Result:        delta_total = {total_lt:.4f} mm"
        )

        # 10.5 Thermal camber (if enabled)
        has_therm = bool(_g(ss, "has_thermal", False))
        if has_therm:
            self._h2("10.5  Thermal Camber")
            alpha_T  = float(_g(ss, "alpha_T", 10e-6))
            delta_T  = float(_g(ss, "delta_T",  0.0))
            h_hcs    = float(_g(ss, "h", 200.0))
            d_therm  = float(_g(ss, "def_delta_thermal", 0.0))
            self._calc(
                "  Formula:  delta_thermal = alfa_T * delta_T * L_an^2 / (8 * h)"
            )
            self._calc(
                f"  Substitution:  delta_thermal = {alpha_T:.2e} * {delta_T:.1f}"
                f" * {L_an:.0f}^2 / (8 * {h_hcs:.0f})"
            )
            self._calc(
                f"  Result:        delta_thermal = {d_therm:.4f} mm"
            )

        # 10.6 Code limits
        self._h2("10.6  Code Deflection Limit Checks")
        lim_frac_ll  = float(_g(ss, "limit_LL_fraction",    360))
        lim_frac_tot = float(_g(ss, "limit_total_fraction", 240))
        self._calc(
            f"  Formula:  Limit_LL    = L_an / {lim_frac_ll:.0f}",
            "ACI 318-19 Table 24.2.2  (floors supporting non-brittle partitions)"
        )
        self._calc(
            f"  Substitution:  Limit_LL    = {L_an:.0f} / {lim_frac_ll:.0f}"
        )
        self._calc(
            f"  Result:        Limit_LL    = {lim_ll:.2f} mm"
        )
        self._calc("")
        self._calc(
            f"  Formula:  Limit_total = L_an / {lim_frac_tot:.0f}",
            "ACI 318-19 Table 24.2.2  (after attachment of non-structural elements)"
        )
        self._calc(
            f"  Substitution:  Limit_total = {L_an:.0f} / {lim_frac_tot:.0f}"
        )
        self._calc(
            f"  Result:        Limit_total = {lim_tot:.2f} mm"
        )
        ok_ll  = st_ll.upper()  in ("OK", "PASS")
        ok_tot = st_tot.upper() in ("OK", "PASS")
        self._status_line(f"Live-load deflection:    {st_ll}", ok_ll)
        self._status_line(f"Total (long-term) deflection: {st_tot}", ok_tot)

        self._wide_table(
            ["Item", "Value (mm)", "Limit (mm)", "Status"],
            [
                ["Prestress camber (initial)", f"{delta_ps:.4f}", "--",            "--"],
                ["Self-weight (initial)",       f"{delta_sw:.4f}", "--",            "--"],
                ["Net at release",              f"{net_rel:.4f}",  "--",            "--"],
                ["Total long-term",             f"{total_lt:.4f}", f"{lim_tot:.2f}", st_tot],
                ["LL deflection (est.)",        f"{final_ll:.4f}", f"{lim_ll:.2f}",  st_ll],
            ]
        )

        self._h2("10.7  Deflection Profile Diagram")
        self._insert_image(img_defl, width_mm=162)
        self.doc.add_page_break()

    # =========================================================================
    # CHAPTER 11 — Vibration & Natural Frequency
    # =========================================================================
    def _ch11_vibration(self):
        ss = self.ss
        self._h1("11.  Vibration and Natural Frequency")
        self._note(
            "Ref: AISC Design Guide 11 (2nd Ed.)  |  ISO 10137:2007\n"
            "Method: Simply-supported beam fundamental frequency + walking acceleration check."
        )

        fn      = float(_g(ss, "def_vib_fn",       0.0))
        fn_lim  = float(_g(ss, "def_vib_fn_limit", 4.0))
        fn_ok   = bool (_g(ss, "def_vib_fn_ok",    False))
        ag      = float(_g(ss, "def_vib_ag",       0.0))
        ag_lim  = float(_g(ss, "def_vib_ag_limit", 0.005))
        ag_ok   = bool (_g(ss, "def_vib_ag_ok",    False))
        W_eff   = float(_g(ss, "def_vib_W_eff",    0.0))
        beta    = float(_g(ss, "def_vib_beta",     3.0))
        mode    = str  (_g(ss, "def_vib_mode",     "Walking / Occupancy"))

        L_an  = float(_g(ss, "L_an", 5850.0))
        L_m   = L_an / 1000.0
        Ec    = float(_g(ss, "Ec_hcs", 36793.0))
        I_c   = float(_g(ss, "sp_I_comp", float(_g(ss, "sp_In", 1.0))))
        EI_SI = Ec * I_c * 1e-6   # N.m2

        self._h2("11.1  Natural Frequency  f_n")
        self._calc(
            "  Formula:  f_n = (pi^2 / (2*L^2)) * sqrt(EI / m_per_m)   [Hz]",
            "AISC Design Guide 11 Eq. 3.3  (simply-supported fundamental)"
        )
        self._calc(
            f"  EI_SI = Ec_hcs * I_comp * 1e-6  =  {Ec:.0f} * {I_c:.3e} * 1e-6  =  {EI_SI:.2f} N.m2"
        )
        self._calc(
            f"  L_m   = {L_an:.0f} / 1000  =  {L_m:.3f} m"
        )
        self._calc(
            f"  W_eff = {W_eff:.2f} kN   (effective tributary weight, including 10% LL)"
        )
        self._calc(
            f"  Result:  f_n = {fn:.3f} Hz"
        )
        self._calc(
            f"  Limit:   f_n >= {fn_lim:.1f} Hz   ({mode})",
            "AISC DG11 Table 4.1"
        )
        self._status_line(
            f"Frequency check: f_n = {fn:.3f} Hz  {'>=  ' if fn_ok else '<  '} {fn_lim:.1f} Hz  =>  {'OK' if fn_ok else 'NG'}",
            fn_ok
        )

        self._h2("11.2  Peak Acceleration Ratio  a/g")
        self._calc(
            "  Formula:  a/g = P0 * exp(-2*pi*beta) / W_eff",
            "AISC Design Guide 11 Eq. 4.2"
        )
        P0 = 0.29
        self._calc(
            f"  Substitution:  a/g = {P0} * exp(-2*pi*{beta/100:.4f}) / {W_eff:.2f}"
        )
        self._calc(
            f"  Result:        a/g = {ag:.5f}"
        )
        self._calc(
            f"  Limit:         a/g <= {ag_lim:.4f}   ({mode})",
            "AISC DG11 Table 4.1"
        )
        self._status_line(
            f"Acceleration check: a/g = {ag:.5f}  {'<=  ' if ag_ok else '>  '} {ag_lim:.4f}  =>  {'OK' if ag_ok else 'NG'}",
            ag_ok
        )

        self._h2("11.3  Vibration Summary")
        self._wide_table(
            ["Check", "Calculated", "Limit", "Mode", "Status"],
            [
                ["Natural frequency f_n (Hz)", f"{fn:.3f}", f">= {fn_lim:.1f}", mode, "OK" if fn_ok else "NG"],
                ["Acceleration ratio a/g",      f"{ag:.5f}", f"<= {ag_lim:.4f}", mode, "OK" if ag_ok else "NG"],
            ]
        )
        self.doc.add_page_break()

    # =========================================================================
    # APPENDICES
    # =========================================================================
    def _appendix_a(self):
        """Appendix A — Section Property Details."""
        self._h1("Appendix A — Section Property Details")
        self._note(
            "Cross-section modelled as rectangular b_top x h.\n"
            "Void centroid: y_void = tf_bot + h_core/2.\n"
            "Transformed steel: (n_ps - 1) x Aps (concrete already in gross).\n"
            "Accuracy > 97% for standard HCS configurations  (PCI practice)."
        )
        ss = self.ss
        rows = [
            ("b_top      — top flange width",       f"{_g(ss,'b_top'):.0f} mm"),
            ("b_bottom   — bottom width",            f"{_g(ss,'b_bottom'):.0f} mm"),
            ("h          — total depth",             f"{_g(ss,'h'):.0f} mm"),
            ("tf_top     — top flange thickness",    f"{_g(ss,'tf_top'):.0f} mm"),
            ("tf_bot     — bottom flange thickness", f"{_g(ss,'tf_bot'):.0f} mm"),
            ("h_core     — total void height",       f"{_g(ss,'h_core'):.1f} mm"),
            ("d_core     — core diameter",           f"{_g(ss,'d_core'):.0f} mm"),
            ("n_core     — number of cores",         f"{int(_g(ss,'n_core',9))}"),
            ("A_core_1   — area one core",           f"{_g(ss,'A_core_1'):.1f} mm2"),
            ("A_voids_total — total void area",      f"{_g(ss,'A_voids_total'):.0f} mm2"),
            ("sp_Ag      — gross area",              f"{_g(ss,'sp_Ag'):.0f} mm2"),
            ("sp_An      — net area",                f"{_g(ss,'sp_An'):.2f} mm2"),
            ("sp_yb      — centroid from bottom",    f"{_g(ss,'sp_yb'):.4f} mm"),
            ("sp_In      — moment of inertia",       f"{_g(ss,'sp_In')/1e6:.4f} x 10^6 mm4"),
            ("sp_Sb      — bottom section modulus",  f"{_g(ss,'sp_Sb')/1e3:.3f} x 10^3 mm3"),
            ("sp_St      — top section modulus",     f"{_g(ss,'sp_St')/1e3:.3f} x 10^3 mm3"),
            ("sp_r2      — radius of gyration sq.",  f"{_g(ss,'sp_r2'):.2f} mm2"),
            ("sp_e_bot   — bottom tendon eccentricity", f"{_g(ss,'sp_e_bot'):.4f} mm"),
            ("sp_e_net   — net eccentricity",        f"{_g(ss,'sp_e_net'):.4f} mm"),
            ("bw_shear   — effective web width",     f"{_g(ss,'bw_shear'):.0f} mm"),
        ]
        if _g(ss, "has_topping"):
            rows += [
                ("sp_A_comp   — composite area",         f"{_g(ss,'sp_A_comp'):.2f} mm2"),
                ("sp_yb_comp  — composite centroid",     f"{_g(ss,'sp_yb_comp'):.4f} mm"),
                ("sp_I_comp   — composite inertia",      f"{_g(ss,'sp_I_comp')/1e6:.4f} x 10^6 mm4"),
                ("sp_Sbc_comp — composite Sb (HCS bot)", f"{_g(ss,'sp_Sbc_comp')/1e3:.3f} x 10^3 mm3"),
                ("sp_Stc_comp — composite St (top fibre)",f"{_g(ss,'sp_Stc_comp')/1e3:.3f} x 10^3 mm3"),
            ]
        self._kv_table(rows)

    def _appendix_b(self):
        """Appendix B — Prestress Loss Details."""
        self._h1("Appendix B — Prestress Loss Details")
        self._note(
            "PCI Lump-Sum method: ES + CR + SH + RE\n"
            "Ref: PCI Design Handbook 8th Ed. Sec. 4.7  |  ACI/PCI CODE-319-25 Cl. 26.10"
        )
        ss = self.ss
        fpi = float(_g(ss, "fpi", 1.0))
        ES  = float(_g(ss, "pl_ES", 0.0))
        CR  = float(_g(ss, "pl_CR", 0.0))
        SH  = float(_g(ss, "pl_SH", 0.0))
        RE  = float(_g(ss, "pl_RE", 0.0))
        total = float(_g(ss, "pl_total_MPa", ES+CR+SH+RE))
        fse   = float(_g(ss, "pl_fse", 0.0))
        Pe    = float(_g(ss, "pl_Pe",  0.0))
        pct   = float(_g(ss, "pl_total_pct", total/fpi*100 if fpi>0 else 0))
        self._wide_table(
            ["Loss Component", "Symbol", "Value (MPa)", "% of fpi"],
            [
                ["Elastic shortening", "ES", f"{ES:.3f}", f"{ES/fpi*100:.2f}" if fpi>0 else "--"],
                ["Creep",              "CR", f"{CR:.3f}", f"{CR/fpi*100:.2f}" if fpi>0 else "--"],
                ["Shrinkage",          "SH", f"{SH:.3f}", f"{SH/fpi*100:.2f}" if fpi>0 else "--"],
                ["Relaxation",         "RE", f"{RE:.3f}", f"{RE/fpi*100:.2f}" if fpi>0 else "--"],
                ["TOTAL",              "--", f"{total:.3f}", f"{pct:.2f}"],
                ["fse  (effective)",   "--", f"{fse:.3f}", "--"],
                ["Pe   (kN)",          "--", f"{Pe:.3f}",  "--"],
            ]
        )
        self._kv_table([
            ("pl_ES",  f"{ES:.3f} MPa"),
            ("pl_CR",  f"{CR:.3f} MPa"),
            ("pl_SH",  f"{SH:.3f} MPa"),
            ("pl_RE",  f"{RE:.3f} MPa"),
            ("pl_total_MPa", f"{total:.3f} MPa"),
            ("pl_total_pct", f"{pct:.2f} %"),
            ("pl_fse",  f"{fse:.3f} MPa"),
            ("pl_Pe",   f"{Pe:.3f} kN"),
        ])

    def _appendix_c(self):
        """Appendix C — ACI/PCI 319-25 Sec. 12 Diaphragm Summary."""
        self._h1("Appendix C — ACI/PCI CODE-319-25 Section 12 Diaphragm Summary")
        self._note(
            "Applicable to Seismic Design Category (SDC) D, E, and F.\n"
            "Ref: ACI/PCI CODE-319-25 Chapter 12  |  ACI CODE-550.5  |  ASCE 7-22 Sec. 12.10"
        )
        ss  = self.ss
        sdc = str(_g(ss, "sdc", "B"))

        self._kv_table([
            ("Seismic Design Category (SDC)", sdc),
            ("HCS topping", "Yes" if _g(ss,"has_topping") else "No"),
        ])

        if sdc in ("D", "E", "F"):
            self._note(
                "SDC D/E/F requirements (summary):"
            )
            items = [
                "ACI/PCI 319-25 Sec. 12.2: Diaphragm must transfer seismic forces to lateral "
                "force-resisting elements (LFRS).",
                "ACI/PCI 319-25 Sec. 12.3: Untopped HCS diaphragms are FLEXIBLE in their own "
                "plane and may NOT be assumed rigid for SDC D/E/F (ACI 318-19 Sec. 12.3.1.1).",
                "ACI/PCI 319-25 Sec. 12.4: Structural topping (cast-in-place, min 50 mm) "
                "required for rigid diaphragm assumption.",
                "ACI/PCI 319-25 Sec. 12.5: Chord forces must be developed at perimeter; "
                "hairpin or collector reinforcement through cores typically required.",
                "ACI/PCI 319-25 Sec. 12.6: Grouted shear keys between HCS units develop "
                "in-plane shear transfer; design per Sec. 26.11.",
                "ACI CODE-550.5: Connections of HCS to LFRS must be designed for "
                "force demands including overstrength (Omega0) effects.",
                "ASCE 7-22 Sec. 12.10.3: Diaphragm forces Fpx must be calculated "
                "separately from element design forces.",
                "NOTE: This summary is informational only. Full diaphragm design requires "
                "specific analysis not included in this automated report.",
            ]
            for item in items:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_after = Pt(3)
                p.add_run(f"\u2022  {item}").font.size = Pt(9)
        else:
            self._note(
                f"SDC {sdc}: Diaphragm flexibility checks per ACI/PCI 319-25 Chapter 12 "
                "are less critical but must still be reviewed per project requirements."
            )

    def _appendix_notes(self):
        """General notes."""
        self.doc.add_page_break()
        self._h1("General Calculation Notes")
        items = [
            "Report generated automatically by HCS Design App v1.0.",
            "Primary references: ACI/PCI CODE-319-25 (Chapters 7, 12, 16, 22, 24, 25, 26) "
            "and PCI Design Handbook 8th Edition.",
            "Units: mm, kN, MPa throughout. No imperial conversions.",
            "Section properties: rectangular simplification with voids at y_void = tf_bot + h_core/2. "
            "Accuracy > 97% for standard HCS.",
            "Prestress losses: PCI lump-sum (ES+CR+SH+RE). "
            "Use detailed time-step method for critical projects.",
            "Development length: preliminary fps = min(fpu, fpy+70). "
            "Exact iterated fps used in capacity calculations.",
            "Deflection multipliers from PCI Table 4.8.3 — parabolic tendon profile assumed.",
            "All diagrams rendered via matplotlib (reliable, no external renderer required).",
            "DISCLAIMER: Engineer of record must verify all inputs and results before construction.",
        ]
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.add_run(f"\u2022  {item}").font.size = Pt(9)

    # =========================================================================
    # BUILD
    # =========================================================================
    def build(self,
              img_sfd_bmd:  bytes | None = None,
              img_phi_mn:   bytes | None = None,
              img_phi_vn:   bytes | None = None,
              img_defl:     bytes | None = None) -> io.BytesIO:
        self._cover()
        self._ch1_project_params()
        self._ch2_section_props()
        self._ch3_materials()
        self._ch4_loads()
        # insert SFD/BMD after ch4
        self._insert_image(img_sfd_bmd, width_mm=162)
        self.doc.add_page_break()
        self._ch5_lifting()
        self._ch6_erection()
        self._ch7_losses()
        self._ch8_service_stress()
        self._ch9_capacity(img_phi_mn, img_phi_vn)
        self._ch10_deflection(img_defl)
        self._ch11_vibration()
        self._appendix_a()
        self.doc.add_page_break()
        self._appendix_b()
        self.doc.add_page_break()
        self._appendix_c()
        self._appendix_notes()
        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf


# =============================================================================
# PDF REPORT  (fpdf2) — mirrors Word chapter order
# =============================================================================

class _PdfReport(FPDF):
    """
    fpdf2 PDF — Helvetica font only (avoids encoding issues on Streamlit Cloud).
    ALL strings pass through _ascii() before printing.
    Mirrors the Word chapter order from FIX-5.
    """

    _NAVY  = (26,  71, 111)
    _BLUE  = (46, 116, 181)
    _GREEN = ( 0,  81,  40)
    _GREY  = (96,  96,  96)
    _RED   = (192,   0,   0)
    _WHITE = (255, 255, 255)
    _LGREY = (242, 242, 242)
    _BLACK = (  0,   0,   0)

    def __init__(self, ss: dict):
        super().__init__("P", "mm", "A4")
        self.ss = ss
        self.set_auto_page_break(auto=True, margin=20)
        self.set_margins(20, 22, 15)
        self.add_page()

    def header(self):
        if self.page_no() == 1:
            return
        self.set_font("Helvetica", "I", 7.5)
        self.set_text_color(*self._GREY)
        self.cell(0, 5,
                  _ascii("HCS Design Report  |  ACI/PCI CODE-319-25  |  PCI Design Handbook 8th Ed."),
                  ln=False)
        self.ln(6)
        self.set_draw_color(200, 200, 200)
        self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
        self.ln(2)
        self.set_text_color(*self._BLACK)

    def footer(self):
        self.set_y(-14)
        self.set_font("Helvetica", "I", 7.5)
        self.set_text_color(*self._GREY)
        ds = datetime.now().strftime("%d %b %Y %H:%M")
        self.cell(0, 5,
                  _ascii(f"Page {self.page_no()}   |   {ds}   |   HCS Design App v1.0"),
                  align="C")

    # ── Helpers ──────────────────────────────────────────────────────────────
    def _h1(self, text: str):
        self.ln(3)
        self.set_fill_color(*self._NAVY)
        self.set_text_color(*self._WHITE)
        self.set_font("Helvetica", "B", 13)
        self.cell(0, 8, _ascii(text), fill=True, ln=True)
        self.set_text_color(*self._BLACK)
        self.ln(2)

    def _h2(self, text: str):
        self.ln(2)
        self.set_font("Helvetica", "B", 10.5)
        self.set_text_color(*self._BLUE)
        self.multi_cell(0, 6, _ascii(text), ln=True)
        self.set_text_color(*self._BLACK)

    def _calc(self, text: str, ref: str = ""):
        self.set_font("Courier", "", 8.5)
        self.set_text_color(*self._GREEN)
        self.multi_cell(0, 5, _ascii(text), ln=True)
        if ref:
            self.set_font("Courier", "I", 7.5)
            self.set_text_color(*self._GREY)
            self.multi_cell(0, 4.5, _ascii(f"    [{ref}]"), ln=True)
        self.set_text_color(*self._BLACK)

    def _note(self, text: str):
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(*self._GREY)
        self.multi_cell(0, 5, _ascii(text), ln=True)
        self.set_text_color(*self._BLACK)

    def _ok(self, text: str):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*self._GREEN)
        self.multi_cell(0, 5.5, _ascii("  [OK]  " + text), ln=True)
        self.set_text_color(*self._BLACK)

    def _warn(self, text: str):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*self._RED)
        self.multi_cell(0, 5.5, _ascii("  [NG]  " + text), ln=True)
        self.set_text_color(*self._BLACK)

    def _status(self, text: str, ok: bool):
        if ok:
            self._ok(text)
        else:
            self._warn(text)

    def _wide_table(self, header: list, data: list):
        nc = len(header)
        col_w = (self.w - self.l_margin - self.r_margin) / nc
        self.set_font("Helvetica", "B", 8)
        self.set_fill_color(*self._NAVY)
        self.set_text_color(*self._WHITE)
        for h in header:
            self.cell(col_w, 6, _ascii(str(h)), border=1, fill=True)
        self.ln()
        self.set_text_color(*self._BLACK)
        fill = False
        for row in data:
            self.set_font("Helvetica", "", 7.5)
            if fill:
                self.set_fill_color(*self._LGREY)
            else:
                self.set_fill_color(*self._WHITE)
            for val in row:
                self.cell(col_w, 5.5, _ascii(str(val)), border=1, fill=True)
            self.ln()
            fill = not fill
        self.ln(1)

    def _insert_image(self, img_bytes: bytes | None, w_mm: float = 170):
        if not img_bytes:
            self._note("(Diagram not available.)")
            return
        try:
            tmp = io.BytesIO(img_bytes)
            x   = self.l_margin + (self.w - self.l_margin - self.r_margin - w_mm) / 2
            self.image(tmp, x=x, w=w_mm)
            self.ln(3)
        except Exception:
            self._note("(Image rendering failed.)")

    # ── Cover ────────────────────────────────────────────────────────────────
    def _cover(self):
        ss = self.ss
        self.set_font("Helvetica", "B", 20)
        self.set_text_color(*self._NAVY)
        self.ln(20)
        self.multi_cell(0, 12,
                        _ascii("HOLLOW CORE SLAB\nSTRUCTURAL DESIGN CALCULATION"),
                        align="C", ln=True)
        self.ln(8)
        self.set_font("Helvetica", "", 9)
        self.set_text_color(*self._BLACK)
        for label, val in [
            ("Project",    _g(ss,"project_name","HCS Design")),
            ("Code ref.",  "ACI/PCI CODE-319-25  |  PCI Design Handbook 8th Ed."),
            ("Units",      "mm . kN . MPa  (SI only)"),
            ("Date",       _g(ss,"report_datetime", datetime.now().strftime("%d %B %Y  %H:%M"))),
            ("Software",   "HCS Design App v1.0  (Streamlit)"),
        ]:
            self.cell(40, 6, _ascii(label + ":"), border=0)
            self.multi_cell(0, 6, _ascii(str(val)), ln=True)
        self.ln(4)
        self._note(
            "DISCLAIMER: Auto-generated report.  Engineer of record must verify "
            "all inputs and results before use in design or construction."
        )
        self.add_page()

    # ── Chapter helpers — PDF mirrors Word content ──────────────────────────
    def _ch1_params(self):
        ss = self.ss
        self._h1("1.  Project Parameters and Code References")
        self._h2("1.1  Code References")
        self._wide_table(
            ["Standard", "Edition", "Chapters"],
            [
                ["ACI/PCI CODE-319-25", "2025", "7, 12, 16, 22, 24, 25, 26"],
                ["PCI Design Handbook", "8th Ed.", "Sec. 2.2, 4.2, 4.7, 4.8, 5.2, 5.3"],
                ["ASCE 7-22", "2022", "Sec. 2.3  (load combinations)"],
                ["AISC Design Guide 11", "2nd Ed.", "Floor vibration"],
            ]
        )
        self._h2("1.2  Disclaimer")
        self._note(
            "This calculation is auto-generated by HCS Design App v1.0.  "
            "The engineer of record must verify all inputs and results."
        )
        self._h2("1.3  Units:  mm | kN | MPa  (SI only)")
        self._h2("1.4  Design Inputs Summary")
        self._wide_table(
            ["Parameter", "Value"],
            [
                ["h  (total depth)",          f"{_g(ss,'h'):.0f} mm"],
                ["b_bottom",                  f"{_g(ss,'b_bottom'):.0f} mm"],
                ["L_an  (analysis span)",     f"{_g(ss,'L_an'):.0f} mm"],
                ["f_ci",                      f"{_g(ss,'f_ci'):.0f} MPa"],
                ["f_c",                       f"{_g(ss,'f_c'):.0f} MPa"],
                ["SW_HCS",                    f"{_g(ss,'SW_HCS'):.3f} kN/m2"],
                ["SDL",                       f"{_g(ss,'SDL'):.2f} kN/m2"],
                ["LL",                        f"{_g(ss,'LL'):.2f} kN/m2"],
                ["Pi",                        f"{_g(ss,'Pi'):.2f} kN"],
                ["SDC",                       _g(ss,"sdc","B")],
            ]
        )
        self.add_page()

    def _ch2_sec_props_pdf(self):
        ss = self.ss
        self._h1("2.  Section Properties")
        An    = float(_g(ss, "sp_An",  1.0))
        yb    = float(_g(ss, "sp_yb",  1.0))
        yt    = float(_g(ss, "sp_yt",  1.0))
        In    = float(_g(ss, "sp_In",  1.0))
        Sb_n  = float(_g(ss, "sp_Sb",  1.0))
        St_n  = float(_g(ss, "sp_St",  1.0))
        r2    = float(_g(ss, "sp_r2",  1.0))
        e_bot = float(_g(ss, "sp_e_bot", 1.0))
        Ag    = float(_g(ss, "sp_Ag",  1.0))
        Ig    = float(_g(ss, "sp_Ig",  1.0))
        A_c1  = float(_g(ss, "A_core_1", 0.0))
        A_void = float(_g(ss, "A_voids_total", 0.0))
        n_core = int(_g(ss, "n_core", 9))
        b_top  = float(_g(ss, "b_top", 1187.0))
        h      = float(_g(ss, "h", 200.0))

        self._h2("2.1  Gross Section")
        self._calc(f"  Ag   = b_top x h  =  {b_top:.0f} x {h:.0f}  =  {Ag:,.0f} mm2")
        self._calc(f"  Ig   = b_top x h^3 / 12  =  {Ig/1e6:.4f} x 10^6 mm4")

        self._h2("2.2  One Core Void")
        self._calc(f"  A_core_1 = {A_c1:.1f} mm2")
        self._calc(f"  A_voids  = {n_core} x {A_c1:.1f}  =  {A_void:.0f} mm2")

        self._h2("2.3  Net HCS Section")
        self._calc(f"  An   =  {An:,.2f} mm2")
        self._calc(f"  yb   =  {yb:.4f} mm  (centroid from bottom)")
        self._calc(f"  yt   =  {yt:.4f} mm")
        self._calc(f"  In   =  {In/1e6:.4f} x 10^6 mm4")
        self._calc(f"  Sb_n =  {Sb_n/1e3:.3f} x 10^3 mm3")
        self._calc(f"  St_n =  {St_n/1e3:.3f} x 10^3 mm3")
        self._calc(f"  r2   =  {r2:.2f} mm2")
        self._calc(f"  e_bot =  {e_bot:.4f} mm")

        has_top = bool(_g(ss, "has_topping"))
        if has_top:
            A_comp  = float(_g(ss, "sp_A_comp",   1.0))
            yb_comp = float(_g(ss, "sp_yb_comp",  1.0))
            I_comp  = float(_g(ss, "sp_I_comp",   1.0))
            Sbc     = float(_g(ss, "sp_Sbc_comp", 1.0))
            Stc     = float(_g(ss, "sp_Stc_comp", 1.0))
            self._h2("2.4  Composite Section")
            self._calc(f"  A_comp   =  {A_comp:,.2f} mm2")
            self._calc(f"  yb_comp  =  {yb_comp:.4f} mm")
            self._calc(f"  I_comp   =  {I_comp/1e6:.4f} x 10^6 mm4")
            self._calc(f"  Sbc_comp =  {Sbc/1e3:.3f} x 10^3 mm3")
            self._calc(f"  Stc_comp =  {Stc/1e3:.3f} x 10^3 mm3")
        self.add_page()

    def _ch3_materials_pdf(self):
        ss = self.ss
        self._h1("3.  Material Properties")
        Ec_hcs = float(_g(ss, "Ec_hcs", 36793.0))
        f_c    = float(_g(ss, "f_c",    50.0))
        wc     = float(_g(ss, "wc",     24.0))
        wc_kg  = wc * 1000.0 / 9.81
        self._h2("3.1  Concrete — Ec Calculation")
        self._calc("  Formula:  Ec_hcs = 0.043 x wc^1.5 x sqrt(f_c)",
                   "ACI 318-19 Eq. 19.2.2.1")
        self._calc(
            f"  Substitution:  Ec_hcs = 0.043 x {wc_kg:.0f}^1.5 x sqrt({f_c:.0f})"
        )
        self._calc(f"  Result:        Ec_hcs = {Ec_hcs:.0f} MPa")
        self._wide_table(
            ["Property", "Value"],
            [
                ["f_ci",   f"{_g(ss,'f_ci'):.0f} MPa"],
                ["f_c",    f"{f_c:.0f} MPa"],
                ["wc",     f"{wc:.1f} kN/m3"],
                ["Ec_hcs", f"{Ec_hcs:.0f} MPa"],
                ["RH",     f"{_g(ss,'RH'):.0f} %"],
                ["V/S",    f"{_g(ss,'V_S'):.1f} mm"],
            ]
        )
        self._h2("3.2  Prestressing Steel")
        self._wide_table(
            ["Property", "Value"],
            [
                ["Type",       str(_g(ss,"ps_type","PC Wire"))],
                ["fpu",        f"{_g(ss,'fpu'):.0f} MPa"],
                ["fpy",        f"{_g(ss,'fpy'):.0f} MPa"],
                ["Eps",        f"{_g(ss,'Eps'):.0f} MPa"],
                ["Aps_bot",    f"{_g(ss,'Aps_bot'):.1f} mm2"],
                ["Pi",         f"{_g(ss,'Pi'):.2f} kN"],
            ]
        )
        self.add_page()

    def _ch4_loads_pdf(self, img_sfd_bmd: bytes | None):
        ss = self.ss
        self._h1("4.  Applied Loads and Internal Forces")
        b_bot   = float(_g(ss, "b_bottom", 1199.0))
        h       = float(_g(ss, "h",         200.0))
        wc      = float(_g(ss, "wc",         24.0))
        A_void  = float(_g(ss, "A_voids_total", 0.0))
        SW_HCS  = float(_g(ss, "SW_HCS",     0.0))
        SW_top  = float(_g(ss, "SW_topping", 0.0))
        SDL     = float(_g(ss, "SDL",         0.0))
        LL      = float(_g(ss, "LL",          0.0))
        wu_area = float(_g(ss, "lb_wu_area",  0.0))
        L_an    = float(_g(ss, "L_an",     5850.0))
        Ra_u    = float(_g(ss, "lb_Ra_u",     0.0))
        Vu_max  = float(_g(ss, "lb_Vu_max",   0.0))
        Mu_max  = float(_g(ss, "lb_Mu_max",   0.0)) / 1e6

        self._h2("4.1  Self-Weight")
        A_net = b_bot * h - A_void
        self._calc(f"  A_net  = b_bottom x h - A_voids  =  {b_bot:.0f} x {h:.0f} - {A_void:.0f}  =  {A_net:,.0f} mm2")
        self._calc(f"  SW_HCS = wc x A_net / (b_bottom x 1e6)  =  {SW_HCS:.4f} kN/m2")

        self._h2("4.2  Factored Load Combination")
        self._calc("  wu = 1.2*(SW_HCS + SW_top + SDL) + 1.6*LL",
                   "ASCE 7-22 Sec. 2.3")
        self._calc(
            f"       = 1.2*({SW_HCS:.3f}+{SW_top:.3f}+{SDL:.3f}) + 1.6*{LL:.3f}"
        )
        self._calc(f"  wu = {wu_area:.4f} kN/m2")

        self._h2("4.3  Reactions and Critical Forces")
        wu_line = wu_area * b_bot / 1e6
        self._calc(f"  wu_line = {wu_line:.6f} kN/mm")
        self._calc(f"  Ra      = wu_line x L_an / 2  =  {Ra_u:.3f} kN")
        self._calc(f"  Vu_max  = {Vu_max:.3f} kN")
        self._calc(f"  Mu_max  = wu_line x L_an^2 / 8  =  {Mu_max:.4f} kN.m")

        self._h2("4.4  Span Distribution Table")
        rows = _span_table(ss, n_seg=10)
        self._wide_table(
            ["x/L", "x (m)", "Vu (kN)", "Mu (kN.m)"],
            [[f"{r['frac']:.1f}", f"{r['x_m']:.3f}",
              f"{r['Vu_kN']:.3f}", f"{r['Mu_kNm']:.3f}"] for r in rows]
        )
        self._h2("4.5  SFD and BMD Diagrams")
        self._insert_image(img_sfd_bmd, w_mm=170)
        self.add_page()

    def _ch5_lifting_pdf(self):
        ss = self.ss
        self._h1("5.  Lifting Stage Check")
        Pi    = float(_g(ss, "Pi",      237.8))
        L_an  = float(_g(ss, "L_an",   5850.0))
        b_bot = float(_g(ss, "b_bottom", 1199.0))
        SW_HCS = float(_g(ss, "SW_HCS", 0.0))
        An    = float(_g(ss, "sp_An",  174000.0))
        In    = float(_g(ss, "sp_In",    776e6))
        yb    = float(_g(ss, "sp_yb",   100.0))
        h     = float(_g(ss, "h",       200.0))
        e_net = float(_g(ss, "sp_e_net", 50.0))
        f_ci  = float(_g(ss, "f_ci",    35.0))
        a_lift = float(_g(ss, "a_lift", L_an / 5.0))
        SW_line = SW_HCS * b_bot / 1e6
        yt = h - yb
        Pi_N = Pi * 1000.0
        M_neg = SW_line * a_lift**2 / 2.0
        Ra_lift = SW_line * L_an / 2.0

        self._calc(f"  a_lift = L_an / 5  =  {L_an:.0f}/5  =  {a_lift:.0f} mm  [default]")
        self._calc(f"  M_lift_neg = SW_line x a^2/2  =  {M_neg/1e6:.4f} kN.m  (hogging)")
        comp_lim = 0.60 * f_ci
        tens_lim = 0.25 * math.sqrt(f_ci)
        f_top = (-Pi_N/An - Pi_N*e_net*yt/In + M_neg*1000*yt/In)
        f_bot = (-Pi_N/An + Pi_N*e_net*yb/In - M_neg*1000*yb/In)
        ok = abs(f_top) <= comp_lim and abs(f_bot) <= comp_lim
        self._calc(f"  f_top = {f_top:.4f} MPa  |  f_bot = {f_bot:.4f} MPa")
        self._calc(f"  Comp limit = {-comp_lim:.3f} MPa  |  Tens limit = {+tens_lim:.4f} MPa")
        self._status(f"Lifting: {'OK' if ok else 'NG'}", ok)
        self.add_page()

    def _ch6_erection_pdf(self):
        ss = self.ss
        self._h1("6.  Erection Stage Check")
        b_bot  = float(_g(ss, "b_bottom", 1199.0))
        L_an   = float(_g(ss, "L_an",    5850.0))
        SW_HCS = float(_g(ss, "SW_HCS",  0.0))
        SW_top = float(_g(ss, "SW_topping", 0.0))
        w_cll  = float(_g(ss, "w_construction_LL", 1.0))
        Pe     = float(_g(ss, "pl_Pe",   200.0))
        An     = float(_g(ss, "sp_An",  174000.0))
        In     = float(_g(ss, "sp_In",   776e6))
        yb     = float(_g(ss, "sp_yb",  100.0))
        h      = float(_g(ss, "h",      200.0))
        e_net  = float(_g(ss, "sp_e_net", 50.0))
        f_c    = float(_g(ss, "f_c",    50.0))
        yt = h - yb
        w_ere = (SW_HCS + SW_top + w_cll) * b_bot / 1e6
        M_ere = w_ere * L_an**2 / 8.0
        Pe_N  = Pe * 1000.0
        f_top = (-Pe_N/An - Pe_N*e_net*yt/In + M_ere*1000*yt/In)
        f_bot = (-Pe_N/An + Pe_N*e_net*yb/In - M_ere*1000*yb/In)
        comp_lim = 0.60 * f_c
        tens_lim = 0.50 * math.sqrt(f_c)
        ok = ((-f_top <= comp_lim if f_top < 0 else True) and
              (-f_bot <= comp_lim if f_bot < 0 else True))
        self._note(f"w_construction_LL = {w_cll:.2f} kN/m2  [DEFAULT 1.0]")
        self._calc(f"  w_erection = {w_ere:.6f} kN/mm")
        self._calc(f"  M_erection = {M_ere/1e6:.4f} kN.m")
        self._calc(f"  f_top = {f_top:.4f} MPa  |  f_bot = {f_bot:.4f} MPa")
        self._calc(f"  Comp limit = -{comp_lim:.3f} MPa  |  Tens limit = +{tens_lim:.4f} MPa")
        self._status(f"Erection: {'OK' if ok else 'NG'}", ok)
        self.add_page()

    def _ch7_losses_pdf(self):
        ss = self.ss
        self._h1("7.  Prestress Losses and Development Length")
        ES = float(_g(ss, "pl_ES", 0.0));  CR = float(_g(ss, "pl_CR", 0.0))
        SH = float(_g(ss, "pl_SH", 0.0));  RE = float(_g(ss, "pl_RE", 0.0))
        fpi = float(_g(ss, "fpi", 1.0))
        total = ES + CR + SH + RE
        fse  = float(_g(ss, "pl_fse",  fpi - total))
        Pe   = float(_g(ss, "pl_Pe",   0.0))
        pct  = total / fpi * 100 if fpi > 0 else 0.0

        self._h2("7.1-7.4  ES + CR + SH + RE")
        self._wide_table(
            ["Component", "MPa", "% fpi"],
            [
                ["Elastic Shortening ES", f"{ES:.3f}", f"{ES/fpi*100:.2f}" if fpi>0 else "--"],
                ["Creep CR",              f"{CR:.3f}", f"{CR/fpi*100:.2f}" if fpi>0 else "--"],
                ["Shrinkage SH",          f"{SH:.3f}", f"{SH/fpi*100:.2f}" if fpi>0 else "--"],
                ["Relaxation RE",         f"{RE:.3f}", f"{RE/fpi*100:.2f}" if fpi>0 else "--"],
                ["TOTAL",                 f"{total:.3f}", f"{pct:.2f}"],
                ["fse",                   f"{fse:.3f}", "--"],
                ["Pe (kN)",               f"{Pe:.3f}",  "--"],
            ]
        )
        l_t = float(_g(ss, "lb_l_t", 250.0))
        l_d = float(_g(ss, "lb_l_d", 380.0))
        L_an = float(_g(ss, "L_an", 5850.0))
        ratio = L_an / l_d if l_d > 0 else 0.0
        status = _g(ss, "lb_ps_status", "FULL")
        self._h2("7.5  Transfer & Development Length")
        self._calc(f"  l_t = {l_t:.1f} mm   |   l_d = {l_d:.1f} mm")
        self._calc(f"  L_an / l_d = {ratio:.3f}")
        self._status(f"PS Development: {status}  (ratio={ratio:.3f})", ratio >= 1.0)
        self.add_page()

    def _ch8_stress_pdf(self):
        ss = self.ss
        self._h1("8.  Service Stress Checks")
        stages = [
            ("8.1 Transfer",    "sc_transfer"),
            ("8.2 Lifting",     "sc_lifting"),
            ("8.3 Erection",    "sc_construction"),
            ("8.4 Service",     "sc_service"),
        ]
        rows = []
        for title, key in stages:
            d = ss.get(key, {})
            if not d:
                rows.append([title, "--", "--", "--", "--", "N/A"])
                continue
            f_top    = float(d.get("f_top",    0.0))
            f_bot    = float(d.get("f_bot",    0.0))
            lim_comp = float(d.get("limit_comp", 0.0))
            lim_tens = float(d.get("limit_tens", 0.0))
            status   = str(d.get("status", "N/A"))
            rows.append([title, f"{f_top:.3f}", f"{f_bot:.3f}",
                         f"-{lim_comp:.3f}", f"+{lim_tens:.4f}", status])
        self._wide_table(
            ["Stage", "f_top (MPa)", "f_bot (MPa)", "Comp lim", "Tens lim", "Status"],
            rows
        )
        self.add_page()

    def _ch9_capacity_pdf(self, img_phi_mn: bytes | None, img_phi_vn: bytes | None):
        ss = self.ss
        self._h1("9.  Flexural and Shear Capacity")
        fps    = float(_g(ss, "cap_fps",       1400.0))
        a      = float(_g(ss, "cap_a",           30.0))
        Mn     = float(_g(ss, "cap_Mn",         200.0))
        phi_Mn = float(_g(ss, "cap_phi_Mn",     180.0))
        Mu_max = float(_g(ss, "lb_Mu_max",        0.0)) / 1e6
        DCR_M  = float(_g(ss, "cap_DCR_M",        1.0))
        phi_Vn = float(_g(ss, "cap_phi_Vn_min",  50.0))
        Vu_max = float(_g(ss, "lb_Vu_max",        0.0))
        DCR_V  = float(_g(ss, "cap_DCR_V",        1.0))
        Aps_b  = float(_g(ss, "Aps_bot",         196.0))
        dp_bot = float(_g(ss, "dp_bot",          165.0))
        f_c    = float(_g(ss, "f_c",              50.0))
        b_top  = float(_g(ss, "b_top",          1187.0))

        self._h2("9.1  Flexural Capacity")
        self._calc(f"  fps  = {fps:.2f} MPa", "ACI 319-25 Eq. 20.3.2.4.1")
        self._calc(f"  a    = Aps_bot x fps / (0.85 x f_c x b_top)"
                   f"  =  {Aps_b:.1f}*{fps:.2f}/(0.85*{f_c:.0f}*{b_top:.0f})  =  {a:.4f} mm")
        self._calc(f"  Mn   = Aps_bot x fps x (dp_bot - a/2) / 1e6"
                   f"  =  {Mn:.4f} kN.m")
        self._calc(f"  phi*Mn = 0.90 x {Mn:.4f}  =  {phi_Mn:.4f} kN.m")
        self._calc(f"  Mu_max = {Mu_max:.4f} kN.m")
        self._calc(f"  DCR_M  = {Mu_max:.4f}/{phi_Mn:.4f}  =  {DCR_M:.4f}")
        self._status(f"Flexure: DCR = {DCR_M:.4f}  => {'OK' if DCR_M<=1.0 else 'NG'}",
                     DCR_M <= 1.0)
        self._h2("phi*Mn vs Mu Diagram")
        self._insert_image(img_phi_mn)

        self._h2("9.2  Shear Capacity")
        self._calc(f"  min phi*Vn = {phi_Vn:.4f} kN  |  Vu_max = {Vu_max:.4f} kN")
        self._calc(f"  DCR_V = {Vu_max:.4f}/{phi_Vn:.4f}  =  {DCR_V:.4f}")
        self._status(f"Shear: DCR = {DCR_V:.4f}  => {'OK' if DCR_V<=1.0 else 'NG'}",
                     DCR_V <= 1.0)
        self._h2("phi*Vn vs Vu Envelope")
        self._insert_image(img_phi_vn)

        self._h2("9.3  Capacity Summary")
        self._wide_table(
            ["Check", "Demand", "Capacity", "DCR", "Status"],
            [
                ["Flexure Mn", f"{Mu_max:.3f} kN.m", f"{phi_Mn:.3f} kN.m",
                 f"{DCR_M:.4f}", "OK" if DCR_M<=1.0 else "NG"],
                ["Shear Vn",  f"{Vu_max:.3f} kN",   f"{phi_Vn:.3f} kN",
                 f"{DCR_V:.4f}", "OK" if DCR_V<=1.0 else "NG"],
            ]
        )
        self.add_page()

    def _ch10_deflection_pdf(self, img_defl: bytes | None):
        ss = self.ss
        self._h1("10.  Deflection and Camber")
        Pe      = float(_g(ss, "pl_Pe",              200.0))
        e_bot   = float(_g(ss, "sp_e_bot",            68.35))
        In      = float(_g(ss, "sp_In",              776e6))
        Ec_hcs  = float(_g(ss, "Ec_hcs",          36793.0))
        L_an    = float(_g(ss, "L_an",             5850.0))
        delta_ps = float(_g(ss, "def_delta_ps_initial", 0.0))
        delta_sw = float(_g(ss, "def_delta_sw",         0.0))
        net_rel  = float(_g(ss, "def_net_release",      0.0))
        total_lt = float(_g(ss, "def_total_longterm",   0.0))
        lim_ll   = float(_g(ss, "def_limit_ll_mm",      0.0))
        lim_tot  = float(_g(ss, "def_limit_total_mm",   0.0))
        st_ll    = str  (_g(ss, "def_status_ll",      "N/A"))
        st_tot   = str  (_g(ss, "def_status_total",   "N/A"))
        mult_cam = float(_g(ss, "mult_camber_final", 2.20))
        mult_sw  = float(_g(ss, "mult_dw_final",     2.40))
        mult_sdl = float(_g(ss, "mult_sdl_final",    1.80))
        final_cam = float(_g(ss, "def_final_camber",  0.0))
        final_sw  = float(_g(ss, "def_final_sw",      0.0))
        final_sdl = float(_g(ss, "def_final_sdl",     0.0))
        final_ll  = float(_g(ss, "def_final_ll",      0.0))

        self._h2("10.1  Prestress Camber")
        self._calc("  delta_ps = Pe x e_bot x L^2 / (8 x Ec x In)",
                   "PCI 8th Ed. Eq. 4.8.1")
        self._calc(f"  = {Pe*1000:.0f} x {e_bot:.4f} x {L_an:.0f}^2"
                   f" / (8 x {Ec_hcs:.0f} x {In:.3e})")
        self._calc(f"  = {delta_ps:.4f} mm  (upward)")

        self._h2("10.2  Self-Weight Deflection")
        self._calc(f"  delta_sw = {delta_sw:.4f} mm  (downward)",
                   "PCI 8th Ed. Eq. 4.8.2")

        self._h2("10.3  Net at Release")
        self._calc(f"  net_release = {delta_sw:.4f} + {delta_ps:.4f}  =  {net_rel:.4f} mm")

        self._h2("10.4  Long-term  (PCI Multipliers)")
        self._wide_table(
            ["Component", "Multiplier", "Result (mm)"],
            [
                ["Camber",     f"{mult_cam:.2f}", f"{final_cam:.4f}"],
                ["Self-weight",f"{mult_sw:.2f}",  f"{final_sw:.4f}"],
                ["SDL",        f"{mult_sdl:.2f}", f"{final_sdl:.4f}"],
                ["LL",         "1.00",            f"{final_ll:.4f}"],
                ["NET TOTAL",  "--",              f"{total_lt:.4f}"],
            ]
        )
        self._h2("10.5  Code Limits")
        self._calc(f"  L/360 = {L_an:.0f}/360 = {lim_ll:.2f} mm  (LL limit)",
                   "ACI 318-19 Table 24.2.2")
        self._calc(f"  L/240 = {L_an:.0f}/240 = {lim_tot:.2f} mm  (total limit)")
        self._status(f"LL deflection:    {st_ll}", st_ll.upper() in ("OK","PASS"))
        self._status(f"Total deflection: {st_tot}", st_tot.upper() in ("OK","PASS"))
        self._h2("10.6  Deflection Profile Diagram")
        self._insert_image(img_defl)
        self.add_page()

    def _ch11_vibration_pdf(self):
        ss = self.ss
        self._h1("11.  Vibration and Natural Frequency")
        fn     = float(_g(ss, "def_vib_fn",       0.0))
        fn_lim = float(_g(ss, "def_vib_fn_limit", 4.0))
        fn_ok  = bool (_g(ss, "def_vib_fn_ok",    False))
        ag     = float(_g(ss, "def_vib_ag",       0.0))
        ag_lim = float(_g(ss, "def_vib_ag_limit", 0.005))
        ag_ok  = bool (_g(ss, "def_vib_ag_ok",    False))
        W_eff  = float(_g(ss, "def_vib_W_eff",    0.0))
        beta   = float(_g(ss, "def_vib_beta",     3.0))
        mode   = str  (_g(ss, "def_vib_mode",     "Walking"))

        self._calc("  Formula:  f_n = (pi^2/(2*L^2)) * sqrt(EI/m_per_m)",
                   "AISC Design Guide 11 Eq. 3.3")
        self._calc(f"  Result:   f_n = {fn:.3f} Hz")
        self._calc(f"  Limit:    f_n >= {fn_lim:.1f} Hz  ({mode})")
        self._status(f"Frequency: f_n={fn:.3f} Hz  => {'OK' if fn_ok else 'NG'}", fn_ok)
        self._calc("  Formula:  a/g = P0 * exp(-2*pi*beta) / W_eff",
                   "AISC DG11 Eq. 4.2")
        self._calc(f"  Result:   a/g = {ag:.5f}  |  Limit = {ag_lim:.4f}")
        self._status(f"Acceleration: a/g={ag:.5f}  => {'OK' if ag_ok else 'NG'}", ag_ok)
        self._wide_table(
            ["Check", "Calculated", "Limit", "Status"],
            [
                ["f_n (Hz)", f"{fn:.3f}", f">= {fn_lim:.1f}", "OK" if fn_ok else "NG"],
                ["a/g",      f"{ag:.5f}", f"<= {ag_lim:.4f}", "OK" if ag_ok else "NG"],
            ]
        )
        self.add_page()

    def _appendix_pdf(self):
        ss = self.ss
        self._h1("Appendix A — Section Property Details")
        An  = float(_g(ss,"sp_An",1.0)); In = float(_g(ss,"sp_In",1.0))
        yb  = float(_g(ss,"sp_yb",1.0)); yt = float(_g(ss,"sp_yt",1.0))
        self._wide_table(
            ["Property", "Value"],
            [
                ["sp_An",  f"{An:,.2f} mm2"],
                ["sp_yb",  f"{yb:.4f} mm"],
                ["sp_yt",  f"{yt:.4f} mm"],
                ["sp_In",  f"{In/1e6:.4f} x 10^6 mm4"],
                ["sp_Sb",  f"{_g(ss,'sp_Sb')/1e3:.3f} x 10^3 mm3"],
                ["sp_St",  f"{_g(ss,'sp_St')/1e3:.3f} x 10^3 mm3"],
                ["sp_r2",  f"{_g(ss,'sp_r2'):.2f} mm2"],
                ["sp_e_bot", f"{_g(ss,'sp_e_bot'):.4f} mm"],
            ]
        )
        self._h1("Appendix B — Prestress Loss Details")
        ES = float(_g(ss,"pl_ES",0.0)); CR = float(_g(ss,"pl_CR",0.0))
        SH = float(_g(ss,"pl_SH",0.0)); RE = float(_g(ss,"pl_RE",0.0))
        fpi = float(_g(ss,"fpi",1.0))
        total = ES+CR+SH+RE
        self._wide_table(
            ["Component", "MPa", "% fpi"],
            [
                ["ES", f"{ES:.3f}", f"{ES/fpi*100:.2f}" if fpi>0 else "--"],
                ["CR", f"{CR:.3f}", f"{CR/fpi*100:.2f}" if fpi>0 else "--"],
                ["SH", f"{SH:.3f}", f"{SH/fpi*100:.2f}" if fpi>0 else "--"],
                ["RE", f"{RE:.3f}", f"{RE/fpi*100:.2f}" if fpi>0 else "--"],
                ["TOTAL", f"{total:.3f}", f"{total/fpi*100:.2f}" if fpi>0 else "--"],
            ]
        )
        self._h1("Appendix C — ACI/PCI 319-25 Sec. 12 Diaphragm Summary")
        sdc = str(_g(ss,"sdc","B"))
        self._note(f"SDC = {sdc}")
        if sdc in ("D","E","F"):
            self._note(
                "SDC D/E/F: Untopped HCS diaphragm CANNOT be assumed rigid.  "
                "Structural topping (>=50 mm) required for rigid diaphragm.  "
                "Chord forces, collector connections and shear keys must be designed.  "
                "Ref: ACI/PCI CODE-319-25 Sec. 12  |  ASCE 7-22 Sec. 12.10."
            )
        else:
            self._note(f"SDC {sdc}: Diaphragm checks per ACI/PCI 319-25 Ch. 12.")

    def build_pdf(self,
                  img_sfd_bmd: bytes | None = None,
                  img_phi_mn:  bytes | None = None,
                  img_phi_vn:  bytes | None = None,
                  img_defl:    bytes | None = None) -> io.BytesIO:
        self._cover()
        self._ch1_params()
        self._ch2_sec_props_pdf()
        self._ch3_materials_pdf()
        self._ch4_loads_pdf(img_sfd_bmd)
        self._ch5_lifting_pdf()
        self._ch6_erection_pdf()
        self._ch7_losses_pdf()
        self._ch8_stress_pdf()
        self._ch9_capacity_pdf(img_phi_mn, img_phi_vn)
        self._ch10_deflection_pdf(img_defl)
        self._ch11_vibration_pdf()
        self._appendix_pdf()
        raw = self.output()
        buf = io.BytesIO(raw)
        buf.seek(0)
        return buf


# =============================================================================
# PUBLIC API  (called by pages/3_HCS_Design.py)
# =============================================================================

def generate_word_report(ss: dict, output_stream=None) -> io.BytesIO | None:
    """
    Generate Word (.docx) report.
    Returns BytesIO or None if python-docx is not installed or error occurs.
    """
    if not HAS_DOCX:
        return None
    try:
        img_sfd = _mpl_sfd_bmd(ss)
        img_mn  = _mpl_phi_mn_vs_mu(ss)
        img_vn  = _mpl_phi_vn_envelope(ss)
        img_def = _mpl_deflection_profile(ss)
        builder = _DocxReport(ss)
        buf     = builder.build(img_sfd_bmd=img_sfd,
                                img_phi_mn=img_mn,
                                img_phi_vn=img_vn,
                                img_defl=img_def)
        if output_stream is not None:
            output_stream.write(buf.getvalue())
            output_stream.seek(0)
            return output_stream
        return buf
    except Exception:
        traceback.print_exc()
        return None


def generate_pdf_report(ss: dict, output_stream=None) -> io.BytesIO | None:
    """
    Generate PDF report via fpdf2.
    All Unicode sanitised through _ascii().
    Returns BytesIO or None.
    """
    if not HAS_FPDF:
        return None
    try:
        img_sfd = _mpl_sfd_bmd(ss)
        img_mn  = _mpl_phi_mn_vs_mu(ss)
        img_vn  = _mpl_phi_vn_envelope(ss)
        img_def = _mpl_deflection_profile(ss)
        builder = _PdfReport(ss)
        buf     = builder.build_pdf(img_sfd_bmd=img_sfd,
                                    img_phi_mn=img_mn,
                                    img_phi_vn=img_vn,
                                    img_defl=img_def)
        if output_stream is not None:
            output_stream.write(buf.getvalue())
            output_stream.seek(0)
            return output_stream
        return buf
    except Exception:
        traceback.print_exc()
        return None


def get_report_bytes(ss: dict) -> tuple:
    """
    Returns (word_bytes, pdf_bytes) for Streamlit st.download_button.
    Either may be None if the library is missing or an error occurs.

    Usage in 3_HCS_Design.py (persistent download buttons via st.session_state):

        if "report_word" not in st.session_state:
            st.session_state.report_word, st.session_state.report_pdf = get_report_bytes(ss)

        st.download_button("Download Word", data=st.session_state.report_word,
                           file_name="HCS_Report.docx",
                           mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                           key="dl_word")
        st.download_button("Download PDF", data=st.session_state.report_pdf,
                           file_name="HCS_Report.pdf", mime="application/pdf",
                           key="dl_pdf")
    """
    word_bytes: bytes | None = None
    pdf_bytes:  bytes | None = None

    if HAS_DOCX:
        word_io = generate_word_report(ss)
        if word_io is not None:
            word_bytes = word_io.getvalue()

    if HAS_FPDF:
        pdf_io = generate_pdf_report(ss)
        if pdf_io is not None:
            pdf_bytes = pdf_io.getvalue()

    return word_bytes, pdf_bytes
