"""
pilecap/report_pilecap.py
=========================
Generator laporan perhitungan pilecap dalam format Word (.docx) dan PDF.

Kedua format menghasilkan isi yang IDENTIK melalui fungsi generate_content()
yang mengembalikan struktur data laporan, lalu di-render ke masing-masing format.

Struktur laporan:
  0. Header proyek
  1. Data Input
  2. Sketsa Denah Pilecap
  3. Perhitungan Gaya Tiang
  4. Efisiensi Grup Tiang
  5. Cek Geser (satu arah & pons)
  6. Penulangan Pilecap
  7. Kesimpulan

Standar: SNI 2847:2019, SNI 8460:2017
Dibuat oleh: Ladosi Engineering
"""

import io
import math
import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from typing import List, Dict, Tuple, Any, Optional

# python-docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# reportlab
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    Image as RLImage, HRFlowable, KeepTogether,
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

# Modul pilecap
from pilecap.geometry import PilecapGeometry
from pilecap.pile_forces import BebanTotal, HasilGayaTiang
from pilecap.group_efficiency import HasilEfisiensiGrup
from pilecap.shear_check import DEfektif, HasilGeserSatuArah, HasilPons
from pilecap.reinforcement import HasilTulangan


# ---------------------------------------------------------------------------
# Warna tema
# ---------------------------------------------------------------------------
WARNA_BIRU_TUA  = RGBColor(0x1e, 0x3a, 0x5f)
WARNA_BIRU      = RGBColor(0x1e, 0x40, 0xaf)
WARNA_HIJAU     = RGBColor(0x16, 0xa3, 0x4a)
WARNA_MERAH     = RGBColor(0xdc, 0x26, 0x26)
WARNA_ABU       = RGBColor(0x6b, 0x72, 0x80)
WARNA_KUNING_BG = RGBColor(0xff, 0xfb, 0xeb)

RL_BIRU_TUA  = colors.HexColor("#1e3a5f")
RL_BIRU      = colors.HexColor("#1e40af")
RL_HIJAU     = colors.HexColor("#16a34a")
RL_MERAH     = colors.HexColor("#dc2626")
RL_ABU       = colors.HexColor("#f1f5f9")
RL_ABU_TUA   = colors.HexColor("#374151")
RL_KUNING    = colors.HexColor("#fef9c3")


# ---------------------------------------------------------------------------
# Fungsi bantu: buat gambar denah (BytesIO)
# ---------------------------------------------------------------------------
def _buat_gambar_denah(
    geom        : PilecapGeometry,
    hasil_tiang : List[HasilGayaTiang],
    ukuran      : Tuple[float, float] = (7.5, 7.5),
) -> io.BytesIO:
    """
    Buat gambar denah pilecap lengkap untuk laporan.
    Return BytesIO PNG 150 dpi.
    """
    fig, ax = plt.subplots(figsize=ukuran)
    ax.set_aspect('equal')

    # Pilecap
    ax.add_patch(plt.Rectangle((0, 0), geom.Lx, geom.Ly,
                                lw=2.0, edgecolor='#1a1a2e', facecolor='#f0f4f8', zorder=1))

    # Kolom
    for kol in geom.kolom_list:
        ax.add_patch(plt.Rectangle((kol.xk - kol.bk/2, kol.yk - kol.hk/2),
                                    kol.bk, kol.hk, lw=1.5, edgecolor='#374151',
                                    facecolor='#9ca3af', zorder=3))
        ax.text(kol.xk, kol.yk, f"K{kol.id_kolom}\n{kol.bk*100:.0f}×{kol.hk*100:.0f}",
                ha='center', va='center', fontsize=7, color='white', fontweight='bold', zorder=4)

    # Tiang
    r = geom.D_m / 2
    for h in hasil_tiang:
        warna = '#1e40af' if h.Pi >= 0 else '#dc2626'
        ax.add_patch(plt.Circle((h.x, h.y), r, lw=1.5,
                                 edgecolor='#1e3a5f', facecolor=warna, alpha=0.85, zorder=5))
        ax.text(h.x, h.y, str(h.no_tiang), ha='center', va='center',
                fontsize=8, color='white', fontweight='bold', zorder=6)

    # Centroid
    xb, yb = geom.centroid_grup()
    ax.plot(xb, yb, '+', ms=14, mew=2.5, color='#16a34a', zorder=7)

    # Dimensi Lx
    off = max(geom.Ly * 0.09, 0.18)
    ax.annotate('', xy=(geom.Lx, -off), xytext=(0, -off),
                 arrowprops=dict(arrowstyle='<->', color='#374151', lw=1.1))
    ax.text(geom.Lx/2, -off - 0.06, f"Lx = {geom.Lx:.2f} m",
            ha='center', va='top', fontsize=8.5, color='#374151')

    # Dimensi Ly
    ax.annotate('', xy=(-off, geom.Ly), xytext=(-off, 0),
                 arrowprops=dict(arrowstyle='<->', color='#374151', lw=1.1))
    ax.text(-off - 0.06, geom.Ly/2, f"Ly = {geom.Ly:.2f} m",
            ha='right', va='center', fontsize=8.5, color='#374151', rotation=90)

    # Arah sumbu
    mx = 0.10
    ax.annotate('X (+)', xy=(geom.Lx + mx + 0.18, 0.08), xytext=(geom.Lx + mx, 0.08),
                 arrowprops=dict(arrowstyle='->', color='#dc2626', lw=1.4),
                 fontsize=8, color='#dc2626', fontweight='bold')
    ax.annotate('Y (+)', xy=(0.08, geom.Ly + mx + 0.18), xytext=(0.08, geom.Ly + mx),
                 arrowprops=dict(arrowstyle='->', color='#dc2626', lw=1.4),
                 fontsize=8, color='#dc2626', fontweight='bold')

    # Keterangan konvensi tulangan
    keterangan = (
        "Konvensi tulangan:\n"
        "  Arah X: batang membentang → arah X, berbaris ke arah Y\n"
        "  Arah Y: batang membentang ↑ arah Y, berbaris ke arah X"
    )
    ax.text(0.01, 0.01, keterangan, transform=ax.transAxes,
            fontsize=7.5, color='#374151', va='bottom',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#fffbeb', alpha=0.85))

    # Legenda
    handles = [
        mpatches.Patch(facecolor='#f0f4f8', edgecolor='#1a1a2e',
                        label=f'Pilecap {geom.Lx:.2f}×{geom.Ly:.2f}×{geom.t:.2f} m'),
        mpatches.Patch(facecolor='#1e40af', label=f'Tiang TEKAN Ø{geom.diameter_pile:.0f} mm'),
        mpatches.Patch(facecolor='#dc2626', label='Tiang TARIK'),
        mpatches.Patch(facecolor='#9ca3af', edgecolor='#374151', label='Kolom'),
        mpatches.Patch(facecolor='#16a34a', label='Centroid grup tiang'),
    ]
    ax.legend(handles=handles, loc='upper right', fontsize=7, framealpha=0.92)

    pad = max(geom.Lx, geom.Ly) * 0.22 + 0.3
    ax.set_xlim(-pad, geom.Lx + pad)
    ax.set_ylim(-pad, geom.Ly + pad)
    ax.set_xlabel("Arah X (m)", fontsize=9); ax.set_ylabel("Arah Y (m)", fontsize=9)
    ax.set_title("Denah Pilecap — Layout Tiang & Kolom", fontsize=11, fontweight='bold', pad=10)
    ax.grid(True, linestyle='--', alpha=0.35, color='#94a3b8')
    ax.tick_params(labelsize=8)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight')
    buf.seek(0)
    plt.close(fig)
    return buf


# ---------------------------------------------------------------------------
# BAGIAN A — GENERATOR KONTEN (struktur data, dipakai Word & PDF)
# ---------------------------------------------------------------------------
def generate_content(
    geom          : PilecapGeometry,
    beban         : BebanTotal,
    hasil_tiang   : List[HasilGayaTiang],
    hasil_grup    : HasilEfisiensiGrup,
    d_efektif     : DEfektif,
    geser_x       : HasilGeserSatuArah,
    geser_y       : HasilGeserSatuArah,
    list_pons     : List[HasilPons],
    hasil_tul     : Dict[str, HasilTulangan],
    nama_proyek   : str = "",
    no_dokumen    : str = "",
) -> Dict:
    """
    Kumpulkan semua konten laporan ke dalam satu dict terstruktur.
    Dipakai bersama oleh render_word() dan render_pdf().
    """
    tanggal = datetime.date.today().strftime("%d %B %Y")
    xb, yb  = geom.centroid_grup()

    # --- Ringkasan status ---
    URUTAN_TUL = ["Bawah-X","Bawah-Y","Atas-X","Atas-Y","Badan-X","Badan-Y"]
    semua_cek = []

    # Geser
    semua_cek.append(("Geser 1 arah X", geser_x.status,
                       f"Vu={geser_x.Vu:.1f} kN ≤ φVn={geser_x.phi_Vn:.1f} kN (rasio={geser_x.rasio:.3f})"))
    semua_cek.append(("Geser 1 arah Y", geser_y.status,
                       f"Vu={geser_y.Vu:.1f} kN ≤ φVn={geser_y.phi_Vn:.1f} kN (rasio={geser_y.rasio:.3f})"))
    for p in list_pons:
        semua_cek.append((f"Pons Kolom {p.id_kolom}", p.status,
                           f"Vu={p.Vu_pons:.1f} kN ≤ φVn={p.phi_Vn:.1f} kN (rasio={p.rasio:.3f})"))

    # Efisiensi grup
    semua_cek.append(("Efisiensi grup (Cek 1 Pmax)", hasil_grup.cek_Pmax,
                       f"Pmax={hasil_grup.Pmax:.1f} kN ≤ η×P_ijin={hasil_grup.eta_pakai*hasil_grup.Pmax/max(hasil_grup.Pmax,0.001):.0f} kN"))
    semua_cek.append(("Efisiensi grup (Cek 2 ΣP)", hasil_grup.cek_grup,
                       f"ΣPtekan={hasil_grup.SigmaPtekan:.1f} kN ≤ P_grup={hasil_grup.P_grup_efektif:.1f} kN"))

    # Penulangan
    for pos in URUTAN_TUL:
        if pos in hasil_tul:
            h = hasil_tul[pos]
            semua_cek.append((f"Tulangan {pos}", "OK" if h.OK else "NG",
                               f"{h.label_notasi}  As_pasang={h.As_pasang:.0f} mm² ≥ As_perlu={h.As_perlu:.0f} mm²"))

    # --- Tabel gaya tiang ---
    rows_tiang = []
    for h in hasil_tiang:
        rows_tiang.append({
            "no": h.no_tiang, "x": h.x, "y": h.y, "xi": h.xi, "yi": h.yi,
            "Pi": h.Pi, "Hxi": h.Hxi, "Hyi": h.Hyi, "Hi": h.Hi,
            "aksial": h.status_aksial, "cek_a": h.status_aksial_cek,
            "cek_l": h.status_lateral,
        })

    # --- Tabel penulangan rekap ---
    rows_tul = []
    for pos in URUTAN_TUL:
        if pos in hasil_tul:
            h = hasil_tul[pos]
            arah = pos.split("-")[1]
            berbaris = {"X": "ke arah Y", "Y": "ke arah X"}.get(arah, "ke arah Z")
            rows_tul.append({
                "posisi"    : pos,
                "notasi"    : h.label_notasi,
                "As_perlu"  : f"{h.As_perlu:.1f}",
                "As_pasang" : f"{h.As_pasang:.1f}",
                "rasio"     : f"{h.rasio_As:.3f}",
                "status"    : "OK" if h.OK else "NG",
                "penjelasan": h.penjelasan,
            })

    return {
        "meta": {
            "nama_proyek" : nama_proyek or "—",
            "no_dokumen"  : no_dokumen  or "—",
            "tanggal"     : tanggal,
            "dibuat_oleh" : "Ladosi Engineering Portal",
            "standar"     : "SNI 2847:2019 & SNI 8460:2017",
        },
        "input": {
            "pilecap": {
                "Lx (m)": geom.Lx, "Ly (m)": geom.Ly, "t (m)": geom.t,
                "f'c (MPa)": geom.fc, "fy (MPa)": geom.fy,
                "Selimut beton (mm)": geom.cover,
                "Diameter tiang (mm)": geom.diameter_pile,
                "Jumlah tiang": geom.jumlah_tiang,
                "h galian (m)": geom.galian.h_galian,
                "γ tanah (kN/m³)": geom.galian.gamma_tanah,
                "h muka air (m)": geom.galian.h_muka_air,
            },
            "pile_coords": geom.pile_coords,
            "kolom_list" : geom.kolom_list,
        },
        "beban": {
            "SigmaNu"   : beban.SigmaNu,
            "W_pilecap" : beban.W_pilecap,
            "W_tanah"   : beban.W_tanah,
            "F_uplift"  : beban.F_uplift,
            "SigmaMuy"  : beban.SigmaMuy,
            "SigmaMux"  : beban.SigmaMux,
            "SigmaVux"  : beban.SigmaVux,
            "SigmaVuy"  : beban.SigmaVuy,
            "Nu_list"   : beban.Nu_kolom_list,
        },
        "tiang": {
            "rows"  : rows_tiang,
            "xbar"  : round(xb, 4),
            "ybar"  : round(yb, 4),
            "Ix"    : round(geom.Ix_grup(), 4),
            "Iy"    : round(geom.Iy_grup(), 4),
        },
        "grup": {
            "eta_CL"    : hasil_grup.eta_CL,
            "eta_Feld"  : hasil_grup.eta_Feld,
            "P_grup"    : hasil_grup.P_grup,
            "P_blok"    : hasil_grup.P_blok,
            "P_efektif" : hasil_grup.P_grup_efektif,
            "cek_Pmax"  : hasil_grup.cek_Pmax,
            "cek_grup"  : hasil_grup.cek_grup,
            "metode_CL" : hasil_grup.metode_CL,
            "metode_F"  : hasil_grup.metode_Feld,
        },
        "geser": {
            "dx"      : d_efektif.dx,
            "dy"      : d_efektif.dy,
            "d_pakai" : d_efektif.d_pakai,
            "gx_Vu"   : geser_x.Vu, "gx_phiVn": geser_x.phi_Vn,
            "gx_rasio": geser_x.rasio, "gx_status": geser_x.status,
            "gy_Vu"   : geser_y.Vu, "gy_phiVn": geser_y.phi_Vn,
            "gy_rasio": geser_y.rasio, "gy_status": geser_y.status,
            "pons"    : [{"id":p.id_kolom,"bo":p.bo,"Vu":p.Vu_pons,
                           "Vc1":p.Vc1,"Vc2":p.Vc2,"Vc3":p.Vc3,
                           "Vc_min":p.Vc_min,"phiVn":p.phi_Vn,
                           "rasio":p.rasio,"status":p.status} for p in list_pons],
        },
        "tul": {
            "rows"   : rows_tul,
            "d_eff"  : {"dx": d_efektif.dx, "dy": d_efektif.dy},
            "detail" : {pos: hasil_tul[pos] for pos in URUTAN_TUL if pos in hasil_tul},
        },
        "kesimpulan": semua_cek,
    }


# ===========================================================================
# BAGIAN B — RENDER WORD (.docx)
# ===========================================================================
def _set_cell_bg(cell, hex_color: str):
    """Set latar belakang sel Word."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_tabel_word(doc: Document, headers: List[str], rows: List[List[str]],
                    col_widths: List[float] = None) -> None:
    """Tambahkan tabel berformat ke dokumen Word."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    # Header
    hr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _set_cell_bg(cell, "1e3a5f")
    # Data
    for ri, row in enumerate(rows):
        tr = tbl.rows[ri + 1]
        bg = "f8fafc" if ri % 2 == 0 else "ffffff"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            _set_cell_bg(cell, bg)
    # Lebar kolom
    if col_widths:
        for ri2, row_obj in enumerate(tbl.rows):
            for ci2, w in enumerate(col_widths):
                row_obj.cells[ci2].width = Cm(w)
    doc.add_paragraph()


def _heading_word(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = WARNA_BIRU_TUA
    return p


def _para_word(doc: Document, text: str, bold: bool = False, italic: bool = False,
               size: int = 10, color: RGBColor = None, indent: float = 0.0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    if indent > 0:
        p.paragraph_format.left_indent = Cm(indent)
    return p


def render_word(
    content : Dict,
    geom    : PilecapGeometry,
    hasil_tiang: List[HasilGayaTiang],
    gambar_buf : io.BytesIO,
) -> io.BytesIO:
    """
    Render laporan ke format Word (.docx).
    Return BytesIO buffer .docx.
    """
    doc = Document()

    # Ukuran halaman A4
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(2.5)
    section.top_margin  = section.bottom_margin = Cm(2.0)

    meta = content["meta"]

    # ── HEADER ──────────────────────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run("LADOSI ENGINEERING")
    run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = WARNA_BIRU_TUA

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p_sub.add_run("Perhitungan Pondasi Pilecap (Pile Cap)")
    run2.font.size = Pt(13); run2.font.bold = True; run2.font.color.rgb = WARNA_BIRU

    doc.add_paragraph()
    _add_tabel_word(doc,
        ["Proyek", "No. Dokumen", "Tanggal", "Standar"],
        [[meta["nama_proyek"], meta["no_dokumen"], meta["tanggal"], meta["standar"]]],
        col_widths=[5.5, 3.5, 3.5, 3.5]
    )

    # ── 1. DATA INPUT ────────────────────────────────────────────────────
    _heading_word(doc, "1. Data Input", 1)
    _heading_word(doc, "1.1 Dimensi & Material Pilecap", 2)
    inp = content["input"]["pilecap"]
    _add_tabel_word(doc,
        ["Parameter", "Nilai", "Satuan"],
        [[k, str(v), ""] for k, v in inp.items()],
        col_widths=[7.0, 4.0, 2.0]
    )

    _heading_word(doc, "1.2 Posisi Tiang Pancang", 2)
    pc = content["input"]["pile_coords"]
    _add_tabel_word(doc,
        ["No. Tiang", "x (m)", "y (m)"],
        [[str(i+1), f"{x:.3f}", f"{y:.3f}"] for i, (x, y) in enumerate(pc)],
        col_widths=[3.0, 4.5, 4.5]
    )

    _heading_word(doc, "1.3 Data Kolom & Beban Terfaktor", 2)
    _para_word(doc, "Catatan: Semua beban adalah beban terfaktor (LRFD) — SNI 2847:2019 Ps. 5.3",
               italic=True, size=9, color=WARNA_ABU)
    kol_list = content["input"]["kolom_list"]
    _add_tabel_word(doc,
        ["No.", "xk (m)", "yk (m)", "bk (m)", "hk (m)",
         "Nu (kN)", "Vux (kN)", "Vuy (kN)", "Mux (kNm)", "Muy (kNm)"],
        [[str(k.id_kolom), f"{k.xk:.3f}", f"{k.yk:.3f}", f"{k.bk:.2f}", f"{k.hk:.2f}",
          f"{k.Nu:.1f}", f"{k.Vux:.1f}", f"{k.Vuy:.1f}", f"{k.Mux:.2f}", f"{k.Muy:.2f}"]
         for k in kol_list],
        col_widths=[0.8, 1.4, 1.4, 1.2, 1.2, 1.5, 1.5, 1.5, 1.8, 1.8]
    )

    # ── 2. SKETSA DENAH ─────────────────────────────────────────────────
    _heading_word(doc, "2. Sketsa Denah Pilecap", 1)
    gambar_buf.seek(0)
    doc.add_picture(gambar_buf, width=Cm(14))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_word(doc, "Gambar 1. Denah pilecap — layout tiang (biru=tekan, merah=tarik) dan kolom",
               italic=True, size=9, color=WARNA_ABU)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── 3. GAYA TIANG ───────────────────────────────────────────────────
    _heading_word(doc, "3. Perhitungan Gaya Tiang", 1)
    b = content["beban"]
    Nu_str = " + ".join(f"{v:.2f}" for v in b["Nu_list"])

    calc_lines = [
        "Beban total ke pilecap:",
        f"  W_pc  = Lx × Ly × t × γbeton",
        f"  W_pc  = {geom.Lx:.2f} × {geom.Ly:.2f} × {geom.t:.2f} × 25",
        f"  W_pc  = {b['W_pilecap']:.2f} kN",
        "",
        f"  W_tanah = {b['W_tanah']:.2f} kN  (tanah urug di atas pilecap)",
        f"  F_uplift = {b['F_uplift']:.2f} kN  (tekanan air ke atas)",
        "",
        f"  ΣNu = ΣNu_kolom + W_pc + W_tanah − F_uplift",
        f"  ΣNu = ({Nu_str}) + {b['W_pilecap']:.2f} + {b['W_tanah']:.2f} − {b['F_uplift']:.2f}",
        f"  ΣNu = {b['SigmaNu']:.2f} kN",
        "",
        f"  ΣMuy_total = {b['SigmaMuy']:.2f} kNm",
        f"  ΣMux_total = {b['SigmaMux']:.2f} kNm",
        "",
        f"Gaya aksial tiang ke-i:",
        f"  Pi = ΣNu/n  ±  ΣMuy×xi/Iy  ±  ΣMux×yi/Ix",
        f"  n={geom.jumlah_tiang}, Iy=Σxi²={content['tiang']['Iy']:.4f} m², Ix=Σyi²={content['tiang']['Ix']:.4f} m²",
    ]

    # Tambahkan detail per tiang
    for h in hasil_tiang:
        P_rata = b['SigmaNu'] / geom.jumlah_tiang
        Iy = content['tiang']['Iy']; Ix = content['tiang']['Ix']
        dP_y = b['SigmaMuy'] * h.xi / Iy if Iy > 1e-10 else 0
        dP_x = b['SigmaMux'] * h.yi / Ix if Ix > 1e-10 else 0
        tanda_y = "+" if dP_y >= 0 else "−"
        tanda_x = "+" if dP_x >= 0 else "−"
        calc_lines.append("")
        calc_lines.append(f"  Tiang {h.no_tiang}: xi={h.xi:+.4f} m, yi={h.yi:+.4f} m")
        calc_lines.append(f"  P{h.no_tiang} = {P_rata:.4f} {tanda_y} {abs(dP_y):.4f} {tanda_x} {abs(dP_x):.4f}")
        calc_lines.append(f"  P{h.no_tiang} = {h.Pi:.2f} kN  → {h.status_aksial}")

    for line in calc_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(line)
        run.font.size = Pt(8.5)
        run.font.name = 'Courier New'

    doc.add_paragraph()
    _heading_word(doc, "Rekapitulasi Gaya Per Tiang", 2)
    _add_tabel_word(doc,
        ["No.", "x (m)", "y (m)", "Pi (kN)", "Hx (kN)", "Hy (kN)", "H (kN)", "Aksial", "Status"],
        [[str(r["no"]), f"{r['x']:.3f}", f"{r['y']:.3f}", f"{r['Pi']:.2f}",
          f"{r['Hxi']:.2f}", f"{r['Hyi']:.2f}", f"{r['Hi']:.2f}", r["aksial"],
          "OK" if r["cek_a"]=="OK" and r["cek_l"]=="OK" else "NG"]
         for r in content["tiang"]["rows"]],
        col_widths=[0.7, 1.5, 1.5, 1.8, 1.5, 1.5, 1.5, 1.5, 1.2]
    )

    # ── 4. EFISIENSI GRUP ────────────────────────────────────────────────
    _heading_word(doc, "4. Efisiensi Grup Tiang", 1)
    g = content["grup"]
    for line in [
        g["metode_CL"],
        g["metode_F"],
        f"P_grup = eta × n × P_ijin_tekan = {g['P_grup']:.2f} kN",
        f"P_blok = {g['P_blok']:.2f} kN",
        f"P kapasitas menentukan = {g['P_efektif']:.2f} kN",
        f"Cek 1 (Pmax <= eta x P_ijin): {g['cek_Pmax']}",
        f"Cek 2 (SigmaP <= P_grup): {g['cek_grup']}",
    ]:
        _para_word(doc, line, size=9, indent=0.5)

    # ── 5. CEK GESER ─────────────────────────────────────────────────────
    _heading_word(doc, "5. Cek Geser Pilecap (SNI 2847:2019)", 1)
    gs = content["geser"]
    _heading_word(doc, "5.1 Tinggi Efektif (d)", 2)
    _add_tabel_word(doc,
        ["Parameter", "Nilai (mm)"],
        [["dx (arah X)", f"{gs['dx']:.1f}"],
         ["dy (arah Y)", f"{gs['dy']:.1f}"],
         ["d_pakai (cek geser)", f"{gs['d_pakai']:.1f}"]],
        col_widths=[6.0, 5.0]
    )

    _heading_word(doc, "5.2 Geser Satu Arah", 2)
    for arah, Vu, phiVn, rasio, status in [
        ("X", gs["gx_Vu"], gs["gx_phiVn"], gs["gx_rasio"], gs["gx_status"]),
        ("Y", gs["gy_Vu"], gs["gy_phiVn"], gs["gy_rasio"], gs["gy_status"]),
    ]:
        _para_word(doc, f"Arah {arah}:", bold=True, size=9)
        for line in [
            f"  phiVn = phi x 0.17 x lambda x sqrt(f'c) x b x d",
            f"  Vu_{arah} = {Vu:.2f} kN  |  phiVn = {phiVn:.2f} kN",
            f"  Rasio = {rasio:.3f}  ->  {'AMAN' if status=='OK' else 'TIDAK AMAN'}",
        ]:
            _para_word(doc, line, size=9, indent=0.5)

    _heading_word(doc, "5.3 Geser Dua Arah (Pons) — SNI 2847:2019 Ps. 22.6", 2)
    _add_tabel_word(doc,
        ["Kolom", "bo (mm)", "Vu_pons (kN)", "Vc1 (kN)", "Vc2 (kN)", "Vc3 (kN)",
         "Vc_min (kN)", "phiVn (kN)", "Rasio", "Status"],
        [[str(p["id"]), f"{p['bo']:.1f}", f"{p['Vu']:.2f}", f"{p['Vc1']:.2f}",
          f"{p['Vc2']:.2f}", f"{p['Vc3']:.2f}", f"{p['Vc_min']:.2f}",
          f"{p['phiVn']:.2f}", f"{p['rasio']:.3f}", p["status"]]
         for p in gs["pons"]],
        col_widths=[0.8, 1.5, 1.8, 1.5, 1.5, 1.5, 1.8, 1.8, 1.2, 1.2]
    )

    # ── 6. PENULANGAN ────────────────────────────────────────────────────
    _heading_word(doc, "6. Penulangan Pilecap (SNI 2847:2019)", 1)

    # Catatan konvensi wajib
    p_box = doc.add_paragraph()
    run_box = p_box.add_run(
        "PERHATIAN NOTASI: D16-200 arah X berarti batang baja Ø16 mm yang membentang "
        "ke arah X (panjang batang ≈ Lx), disusun berbaris ke arah Y dengan jarak 200 mm. "
        "BUKAN sebaliknya. Lihat sketsa denah untuk konfirmasi visual."
    )
    run_box.font.bold = True; run_box.font.size = Pt(9)
    run_box.font.color.rgb = RGBColor(0x78, 0x35, 0x0f)
    doc.add_paragraph()

    # Detail per posisi
    for pos, ht in content["tul"]["detail"].items():
        _heading_word(doc, f"6.{list(content['tul']['detail'].keys()).index(pos)+1} Tulangan {pos}", 2)
        if ht.Mu > 0:
            for line in [
                f"Mu = {ht.Mu:.2f} kNm  |  b = {ht.b_mm:.0f} mm  |  d = {ht.d_mm:.1f} mm",
                f"Rn = Mu / (phi x b x d^2) = {ht.Rn:.4f} N/mm^2",
                f"rho = {ht.rho:.6f}  |  rho_min = {ht.rho_min:.6f}  |  rho_pakai = {ht.rho_pakai:.6f}",
                f"As_perlu = rho x b x d = {ht.As_perlu:.2f} mm^2",
            ]:
                _para_word(doc, line, size=9, indent=0.5)
        else:
            _para_word(doc, f"As_perlu = {ht.As_perlu:.2f} mm² (dari susut/tarik minimum)", size=9, indent=0.5)

        _para_word(doc, f"Pilih: {ht.label_notasi}", bold=True, size=10, color=WARNA_BIRU)
        _para_word(doc, ht.penjelasan, italic=True, size=9, color=WARNA_ABU)
        _para_word(doc, f"As_pasang = {ht.As_pasang:.2f} mm²  |  Rasio = {ht.rasio_As:.3f}  |  {'OK' if ht.OK else 'NG'}",
                   size=9, color=WARNA_HIJAU if ht.OK else WARNA_MERAH)

    doc.add_paragraph()
    _heading_word(doc, "Rekap Tabel Penulangan", 2)
    _add_tabel_word(doc,
        ["Posisi", "Notasi", "Arah Batang", "As_perlu (mm²)", "As_pasang (mm²)", "Rasio", "Status"],
        [[r["posisi"], r["notasi"],
          "Membentang X\nBerbaris ke Y" if "X" in r["posisi"].split("-")[1] else "Membentang Y\nBerbaris ke X",
          r["As_perlu"], r["As_pasang"], r["rasio"], r["status"]]
         for r in content["tul"]["rows"]],
        col_widths=[2.0, 1.8, 2.5, 2.2, 2.2, 1.5, 1.2]
    )

    # ── 7. KESIMPULAN ────────────────────────────────────────────────────
    _heading_word(doc, "7. Kesimpulan", 1)
    semua_ok = all(s == "OK" for _, s, _ in content["kesimpulan"])
    _para_word(doc,
               "Seluruh cek memenuhi persyaratan SNI." if semua_ok
               else "PERHATIAN: Ada cek yang tidak memenuhi — tinjau ulang desain.",
               bold=True, size=10,
               color=WARNA_HIJAU if semua_ok else WARNA_MERAH)
    doc.add_paragraph()
    _add_tabel_word(doc,
        ["Item Cek", "Status", "Keterangan"],
        [[nama, status, ket] for nama, status, ket in content["kesimpulan"]],
        col_widths=[5.0, 2.0, 9.0]
    )
    _para_word(doc, f"Dibuat oleh: {meta['dibuat_oleh']}  |  Tanggal: {meta['tanggal']}",
               italic=True, size=8, color=WARNA_ABU)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ===========================================================================
# BAGIAN C — RENDER PDF (reportlab)
# ===========================================================================
def render_pdf(
    content    : Dict,
    geom       : PilecapGeometry,
    hasil_tiang: List[HasilGayaTiang],
    gambar_buf : io.BytesIO,
) -> io.BytesIO:
    """
    Render laporan ke format PDF menggunakan ReportLab.
    Return BytesIO buffer .pdf.
    """
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.2*cm, rightMargin=2.2*cm,
        topMargin=2.0*cm, bottomMargin=2.0*cm,
    )

    styles = getSampleStyleSheet()
    PAGE_W = A4[0] - 4.4*cm

    # Style kustom
    sty_h1 = ParagraphStyle('H1', parent=styles['Heading1'],
                              fontSize=13, textColor=RL_BIRU_TUA,
                              spaceAfter=6, spaceBefore=12, fontName='Helvetica-Bold')
    sty_h2 = ParagraphStyle('H2', parent=styles['Heading2'],
                              fontSize=11, textColor=RL_BIRU,
                              spaceAfter=4, spaceBefore=8, fontName='Helvetica-Bold')
    sty_body = ParagraphStyle('Body', parent=styles['Normal'],
                               fontSize=9, leading=13, spaceAfter=3)
    sty_mono = ParagraphStyle('Mono', parent=styles['Normal'],
                               fontName='Courier', fontSize=8.5, leading=12,
                               leftIndent=0.5*cm, spaceAfter=1)
    sty_caption = ParagraphStyle('Caption', parent=styles['Normal'],
                                  fontSize=8, textColor=colors.grey,
                                  alignment=TA_CENTER, spaceAfter=6)
    sty_warn = ParagraphStyle('Warn', parent=styles['Normal'],
                               fontSize=9, textColor=colors.HexColor("#78350f"),
                               backColor=RL_KUNING, borderPadding=6,
                               leading=14, spaceAfter=6)

    def tbl_rl(headers, rows, col_widths=None):
        """Buat Table ReportLab berformat."""
        data = [[Paragraph(f"<b>{h}</b>", ParagraphStyle('th', fontSize=8.5,
                  textColor=colors.white, fontName='Helvetica-Bold',
                  alignment=TA_CENTER)) for h in headers]]
        for row in rows:
            data.append([Paragraph(str(v), ParagraphStyle('td', fontSize=8,
                          alignment=TA_CENTER, leading=11)) for v in row])

        if col_widths:
            cw = [w*cm for w in col_widths]
        else:
            w_each = PAGE_W / len(headers)
            cw = [w_each] * len(headers)

        t = Table(data, colWidths=cw, repeatRows=1)
        style = TableStyle([
            ('BACKGROUND', (0,0), (-1,0), RL_BIRU_TUA),
            ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
            ('ROWBACKGROUNDS', (0,1), (-1,-1), [RL_ABU, colors.white]),
            ('GRID',  (0,0), (-1,-1), 0.4, colors.HexColor("#cbd5e1")),
            ('TOPPADDING',    (0,0), (-1,-1), 3),
            ('BOTTOMPADDING', (0,0), (-1,-1), 3),
            ('LEFTPADDING',   (0,0), (-1,-1), 4),
            ('RIGHTPADDING',  (0,0), (-1,-1), 4),
            ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ])
        t.setStyle(style)
        return t

    meta = content["meta"]
    story = []

    # ── HEADER ──────────────────────────────────────────────────────────
    story.append(Paragraph("<b>LADOSI ENGINEERING</b>",
                            ParagraphStyle('title', fontSize=16, alignment=TA_CENTER,
                                           textColor=RL_BIRU_TUA, fontName='Helvetica-Bold', spaceAfter=4)))
    story.append(Paragraph("Perhitungan Pondasi Pilecap (Pile Cap)",
                            ParagraphStyle('sub', fontSize=12, alignment=TA_CENTER,
                                           textColor=RL_BIRU, fontName='Helvetica-Bold', spaceAfter=8)))
    story.append(tbl_rl(
        ["Proyek", "No. Dokumen", "Tanggal", "Standar"],
        [[meta["nama_proyek"], meta["no_dokumen"], meta["tanggal"], meta["standar"]]],
        col_widths=[5.0, 3.5, 3.5, 5.2]
    ))
    story.append(Spacer(1, 8*mm))

    # ── 1. DATA INPUT ────────────────────────────────────────────────────
    story.append(Paragraph("1. Data Input", sty_h1))
    story.append(Paragraph("1.1 Dimensi & Material Pilecap", sty_h2))
    inp = content["input"]["pilecap"]
    story.append(tbl_rl(["Parameter","Nilai","Satuan"],
                          [[k, str(v), ""] for k, v in inp.items()],
                          col_widths=[7.0, 4.0, 2.0]))

    story.append(Paragraph("1.2 Posisi Tiang Pancang", sty_h2))
    pc = content["input"]["pile_coords"]
    story.append(tbl_rl(["No.","x (m)","y (m)"],
                          [[str(i+1), f"{x:.3f}", f"{y:.3f}"] for i,(x,y) in enumerate(pc)],
                          col_widths=[2.0, 5.0, 5.0]))

    story.append(Paragraph("1.3 Data Kolom & Beban Terfaktor", sty_h2))
    story.append(Paragraph("Semua beban terfaktor (LRFD) — SNI 2847:2019 Ps. 5.3",
                             ParagraphStyle('note', fontSize=8, textColor=colors.grey, spaceAfter=4)))
    kol_list = content["input"]["kolom_list"]
    story.append(tbl_rl(
        ["No.","xk","yk","bk","hk","Nu (kN)","Vux","Vuy","Mux","Muy"],
        [[str(k.id_kolom), f"{k.xk:.3f}", f"{k.yk:.3f}", f"{k.bk:.2f}", f"{k.hk:.2f}",
          f"{k.Nu:.1f}", f"{k.Vux:.1f}", f"{k.Vuy:.1f}", f"{k.Mux:.2f}", f"{k.Muy:.2f}"]
         for k in kol_list],
        col_widths=[0.8,1.3,1.3,1.2,1.2,1.5,1.3,1.3,1.6,1.7]
    ))
    story.append(Spacer(1, 5*mm))

    # ── 2. SKETSA DENAH ─────────────────────────────────────────────────
    story.append(Paragraph("2. Sketsa Denah Pilecap", sty_h1))
    gambar_buf.seek(0)
    img = RLImage(gambar_buf, width=13*cm, height=13*cm)
    story.append(img)
    story.append(Paragraph("Gambar 1. Denah pilecap — biru=tekan, merah=tarik", sty_caption))
    story.append(Spacer(1, 4*mm))

    # ── 3. GAYA TIANG ───────────────────────────────────────────────────
    story.append(Paragraph("3. Perhitungan Gaya Tiang", sty_h1))
    b = content["beban"]
    Nu_str = " + ".join(f"{v:.2f}" for v in b["Nu_list"])

    calc_lines = [
        "Berat sendiri pilecap:",
        f"  W_pc = Lx x Ly x t x gbeton = {geom.Lx:.2f} x {geom.Ly:.2f} x {geom.t:.2f} x 25 = {b['W_pilecap']:.2f} kN",
        f"  W_tanah = {b['W_tanah']:.2f} kN  |  F_uplift = {b['F_uplift']:.2f} kN",
        "",
        f"  SigmaNu = ({Nu_str}) + {b['W_pilecap']:.2f} + {b['W_tanah']:.2f} - {b['F_uplift']:.2f} = {b['SigmaNu']:.2f} kN",
        f"  SigmaMuy = {b['SigmaMuy']:.2f} kNm  |  SigmaMux = {b['SigmaMux']:.2f} kNm",
        "",
        "Gaya aksial tiang ke-i:",
        f"  Pi = SigmaNu/n +/- SigmaMuy*xi/Iy +/- SigmaMux*yi/Ix",
        f"  n={geom.jumlah_tiang}, Iy={content['tiang']['Iy']:.4f} m2, Ix={content['tiang']['Ix']:.4f} m2",
    ]
    for h in hasil_tiang:
        P_rata = b['SigmaNu'] / geom.jumlah_tiang
        Iy = content['tiang']['Iy']; Ix = content['tiang']['Ix']
        dP_y = b['SigmaMuy'] * h.xi / Iy if Iy > 1e-10 else 0
        dP_x = b['SigmaMux'] * h.yi / Ix if Ix > 1e-10 else 0
        calc_lines.append(f"  Tiang {h.no_tiang}: P{h.no_tiang} = {P_rata:.4f} + {dP_y:.4f} + {dP_x:.4f} = {h.Pi:.2f} kN -> {h.status_aksial}")
    for line in calc_lines:
        story.append(Paragraph(line.replace(" ", "&nbsp;") if line.startswith("  ") else line, sty_mono))

    story.append(Spacer(1, 4*mm))
    story.append(Paragraph("Rekapitulasi Gaya Per Tiang:", sty_h2))
    story.append(tbl_rl(
        ["No.","x (m)","y (m)","Pi (kN)","Hx (kN)","Hy (kN)","H (kN)","Aksial","Status"],
        [[str(r["no"]),f"{r['x']:.3f}",f"{r['y']:.3f}",f"{r['Pi']:.2f}",
          f"{r['Hxi']:.2f}",f"{r['Hyi']:.2f}",f"{r['Hi']:.2f}",r["aksial"],
          "OK" if r["cek_a"]=="OK" and r["cek_l"]=="OK" else "NG"]
         for r in content["tiang"]["rows"]],
        col_widths=[0.8,1.5,1.5,1.8,1.5,1.5,1.5,1.5,1.1]
    ))

    # ── 4. EFISIENSI GRUP ────────────────────────────────────────────────
    story.append(Paragraph("4. Efisiensi Grup Tiang", sty_h1))
    g = content["grup"]
    for line in [g["metode_CL"], g["metode_F"],
                  f"P_grup = {g['P_grup']:.2f} kN  |  P_blok = {g['P_blok']:.2f} kN",
                  f"Kapasitas menentukan = {g['P_efektif']:.2f} kN",
                  f"Cek 1 (Pmax <= eta x P_ijin): {g['cek_Pmax']}",
                  f"Cek 2 (SigmaP <= P_grup): {g['cek_grup']}"]:
        story.append(Paragraph(line, sty_mono))

    # ── 5. CEK GESER ─────────────────────────────────────────────────────
    story.append(Paragraph("5. Cek Geser Pilecap (SNI 2847:2019)", sty_h1))
    gs = content["geser"]
    story.append(Paragraph("5.1 Tinggi Efektif (d)", sty_h2))
    story.append(tbl_rl(["Parameter","Nilai (mm)"],
                          [["dx (arah X)",f"{gs['dx']:.1f}"],
                           ["dy (arah Y)",f"{gs['dy']:.1f}"],
                           ["d_pakai (geser)",f"{gs['d_pakai']:.1f}"]],
                          col_widths=[6.0, 5.0]))

    story.append(Paragraph("5.2 Geser Satu Arah", sty_h2))
    story.append(tbl_rl(
        ["Arah", "Vu (kN)", "phiVn (kN)", "Rasio", "Status"],
        [["X", f"{gs['gx_Vu']:.2f}", f"{gs['gx_phiVn']:.2f}", f"{gs['gx_rasio']:.3f}", gs['gx_status']],
         ["Y", f"{gs['gy_Vu']:.2f}", f"{gs['gy_phiVn']:.2f}", f"{gs['gy_rasio']:.3f}", gs['gy_status']]],
        col_widths=[2.5, 3.0, 3.0, 2.5, 2.2]
    ))
    story.append(Paragraph("5.3 Geser Dua Arah (Pons) — SNI 2847:2019 Ps. 22.6", sty_h2))
    story.append(tbl_rl(
        ["Kolom","bo (mm)","Vu (kN)","Vc1","Vc2","Vc3","Vc_min","phiVn","Rasio","Status"],
        [[str(p["id"]),f"{p['bo']:.1f}",f"{p['Vu']:.2f}",f"{p['Vc1']:.2f}",
          f"{p['Vc2']:.2f}",f"{p['Vc3']:.2f}",f"{p['Vc_min']:.2f}",
          f"{p['phiVn']:.2f}",f"{p['rasio']:.3f}",p["status"]] for p in gs["pons"]],
        col_widths=[1.0,1.5,1.5,1.5,1.5,1.5,1.5,1.5,1.2,1.0]
    ))

    # ── 6. PENULANGAN ────────────────────────────────────────────────────
    story.append(Paragraph("6. Penulangan Pilecap (SNI 2847:2019)", sty_h1))
    story.append(Paragraph(
        "<b>PERHATIAN NOTASI:</b> D16-200 arah X berarti batang baja Ø16 mm yang membentang "
        "ke arah X (panjang batang ≈ Lx), disusun berbaris ke arah Y dengan jarak 200 mm. "
        "BUKAN sebaliknya. Lihat sketsa denah untuk konfirmasi visual.",
        sty_warn
    ))

    for i, (pos, ht) in enumerate(content["tul"]["detail"].items()):
        story.append(Paragraph(f"6.{i+1} Tulangan {pos}", sty_h2))
        if ht.Mu > 0:
            for line in [
                f"Mu = {ht.Mu:.2f} kNm  |  b = {ht.b_mm:.0f} mm  |  d = {ht.d_mm:.1f} mm",
                f"Rn = {ht.Rn:.4f} N/mm2  |  rho = {ht.rho:.6f}  |  rho_min = {ht.rho_min:.6f}",
                f"rho_pakai = {ht.rho_pakai:.6f}  |  As_perlu = {ht.As_perlu:.2f} mm2",
            ]:
                story.append(Paragraph(line, sty_mono))
        else:
            story.append(Paragraph(f"As_perlu = {ht.As_perlu:.2f} mm2 (susut/tarik minimum)", sty_mono))

        warna_status = RL_HIJAU if ht.OK else RL_MERAH
        story.append(Paragraph(
            f"<b>Pilih: {ht.label_notasi}</b>  |  As_pasang={ht.As_pasang:.2f} mm2  "
            f"|  Rasio={ht.rasio_As:.3f}  |  {'OK' if ht.OK else 'NG'}",
            ParagraphStyle('res', fontSize=10, textColor=warna_status,
                            fontName='Helvetica-Bold', spaceAfter=2)
        ))
        story.append(Paragraph(ht.penjelasan,
                                ParagraphStyle('pj', fontSize=8.5, textColor=RL_ABU_TUA,
                                               fontName='Helvetica-Oblique', spaceAfter=4)))

    story.append(Spacer(1, 4*mm))
    story.append(Paragraph("Rekap Tabel Penulangan", sty_h2))
    story.append(tbl_rl(
        ["Posisi","Notasi","Arah Batang","As_perlu\n(mm2)","As_pasang\n(mm2)","Rasio","Status"],
        [[r["posisi"], r["notasi"],
          "Membentang X / Berbaris Y" if "X" in r["posisi"].split("-")[1] else "Membentang Y / Berbaris X",
          r["As_perlu"], r["As_pasang"], r["rasio"], r["status"]]
         for r in content["tul"]["rows"]],
        col_widths=[2.0, 1.8, 3.5, 2.0, 2.0, 1.4, 1.1]
    ))

    # ── 7. KESIMPULAN ────────────────────────────────────────────────────
    story.append(Paragraph("7. Kesimpulan", sty_h1))
    semua_ok = all(s == "OK" for _, s, _ in content["kesimpulan"])
    story.append(Paragraph(
        "Seluruh cek memenuhi persyaratan SNI 2847:2019 & SNI 8460:2017." if semua_ok
        else "PERHATIAN: Ada cek yang TIDAK MEMENUHI — tinjau ulang desain.",
        ParagraphStyle('ok', fontSize=10, fontName='Helvetica-Bold',
                        textColor=RL_HIJAU if semua_ok else RL_MERAH, spaceAfter=6)
    ))

    # Beri warna baris kesimpulan
    kes_data = [["Item Cek","Status","Keterangan"]]
    for nama, status, ket in content["kesimpulan"]:
        kes_data.append([
            Paragraph(nama, ParagraphStyle('k', fontSize=8.5)),
            Paragraph(f"<b>{status}</b>",
                       ParagraphStyle('s', fontSize=8.5, fontName='Helvetica-Bold',
                                      textColor=RL_HIJAU if status=="OK" else RL_MERAH)),
            Paragraph(ket, ParagraphStyle('kt', fontSize=7.5)),
        ])
    t_kes = Table(kes_data, colWidths=[4.5*cm, 2.0*cm, 10.7*cm], repeatRows=1)
    t_kes.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), RL_BIRU_TUA),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [RL_ABU, colors.white]),
        ('GRID', (0,0), (-1,-1), 0.4, colors.HexColor("#cbd5e1")),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('TOPPADDING', (0,0), (-1,-1), 3),
        ('BOTTOMPADDING', (0,0), (-1,-1), 3),
        ('LEFTPADDING', (0,0), (-1,-1), 4),
    ]))
    story.append(t_kes)
    story.append(Spacer(1, 6*mm))
    story.append(Paragraph(
        f"Dibuat oleh: {meta['dibuat_oleh']}  |  Tanggal: {meta['tanggal']}",
        ParagraphStyle('foot', fontSize=8, textColor=colors.grey, alignment=TA_RIGHT)
    ))

    doc.build(story)
    buf.seek(0)
    return buf


# ===========================================================================
# FUNGSI UTAMA — dipanggil dari Streamlit
# ===========================================================================
def buat_laporan(
    geom          : PilecapGeometry,
    beban         : BebanTotal,
    hasil_tiang   : List[HasilGayaTiang],
    hasil_grup    : HasilEfisiensiGrup,
    d_efektif     : DEfektif,
    geser_x       : HasilGeserSatuArah,
    geser_y       : HasilGeserSatuArah,
    list_pons     : List[HasilPons],
    hasil_tul     : Dict[str, HasilTulangan],
    nama_proyek   : str = "",
    no_dokumen    : str = "",
) -> Tuple[io.BytesIO, io.BytesIO]:
    """
    Buat laporan Word dan PDF sekaligus.

    Return
    ------
    (buf_word, buf_pdf) — keduanya BytesIO, siap untuk st.download_button
    """
    # Buat gambar sekali, dipakai keduanya
    gambar_buf = _buat_gambar_denah(geom, hasil_tiang)

    # Generate konten (struktur data bersama)
    content = generate_content(
        geom, beban, hasil_tiang, hasil_grup,
        d_efektif, geser_x, geser_y, list_pons, hasil_tul,
        nama_proyek, no_dokumen,
    )

    # Render Word
    gambar_buf.seek(0)
    buf_word = render_word(content, geom, hasil_tiang, gambar_buf)

    # Render PDF
    gambar_buf.seek(0)
    buf_pdf = render_pdf(content, geom, hasil_tiang, gambar_buf)

    return buf_word, buf_pdf
