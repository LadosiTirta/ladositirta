# =============================================================================
# uditch/ui_output.py
# Output Module — U-Ditch & Cover Precast Structural Analysis
#
# PUBLIC API:
#   render_output(design_result, force_result, inp, lang)
#       Master renderer. Called from 11_UDitch_CU.py tab_output.
#
# INTERNAL sections:
#   _section_summary_metrics(...)   → st.metric cards
#   _section_cross_section(...)     → matplotlib cross-section drawing
#   _section_force_diagrams(...)    → plotly bending/shear/axial diagrams
#   _section_stress_block(...)      → matplotlib Whitney stress-block sketch
#   _section_pm_curve(...)          → plotly P-M interaction diagram
#   _section_calc_steps(...)        → step-by-step LaTeX table
#   _section_exports(...)           → Word & PDF download buttons
#
# EXPORT HELPERS:
#   export_docx(design_result, force_result, inp, lang) → bytes
#   export_pdf(design_result, force_result, inp, lang)  → bytes
# =============================================================================

from __future__ import annotations

import io
import math
import base64
import tempfile
import os
from datetime import datetime
from typing import Optional

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import FancyArrowPatch
import plotly.graph_objects as go
import streamlit as st

from uditch.lang_dict import t
from uditch.calc_engine import (
    DesignResult, ForceResult, PMCurve, CapacityResult, CalcStep
)

# ── Colour palette (consistent across all charts) ────────────────────────────
C = {
    "concrete":  "#B0BEC5",
    "cu":        "#78909C",
    "steel":     "#E53935",
    "comp_steel":"#1E88E5",
    "soil":      "#A5D6A7",
    "load":      "#D32F2F",
    "moment":    "#1565C0",
    "shear":     "#2E7D32",
    "axial":     "#E65100",
    "pm_curve":  "#1565C0",
    "pm_phi":    "#42A5F5",
    "pm_demand": "#D32F2F",
    "ok":        "#2E7D32",
    "ng":        "#C62828",
    "neutral":   "#455A64",
    "bg":        "#FAFAFA",
    "grid":      "#ECEFF1",
}

_DPI = 120   # matplotlib figure DPI

# =============================================================================
# 1. SUMMARY METRICS
# =============================================================================

def _section_summary_metrics(
    dr: DesignResult,
    fr: ForceResult,
    lang: str,
) -> None:
    """Top-row metric cards: key forces and checks."""

    ok_str  = "✅ OK"  if lang == "ID" else "✅ OK"
    ng_str  = "❌ NG"  if lang == "ID" else "❌ NG"

    st.markdown(f"### {'📊 Ringkasan Hasil' if lang=='ID' else '📊 Results Summary'}")

    wb = dr.wall_base_cap
    if wb:
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("Mu (kN·m/m)",  f"{wb.Mu_kNmm:.2f}")
        c2.metric("φMn (kN·m/m)", f"{wb.phi_Mn_kNm:.2f}",
                  delta=ok_str if wb.flexure_ok else ng_str,
                  delta_color="normal" if wb.flexure_ok else "inverse")
        c3.metric("Vu (kN/m)",    f"{wb.Vu_kNm:.2f}")
        c4.metric("φVn (kN/m)",   f"{wb.phi_Vn_kNm:.2f}",
                  delta=ok_str if wb.shear_ok else ng_str,
                  delta_color="normal" if wb.shear_ok else "inverse")
        c5.metric("As,req (mm²/m)", f"{wb.As_req_mm2:.0f}")

    if fr.condition == "Kondisi 1":
        st.info(
            f"🔧 {'Gap' if lang=='EN' else 'Celah'}: **{fr.gap_mm} mm** | "
            f"δ_cant = **{fr.delta_cant_mm:.2f} mm** | "
            + ("Strut aktif (propped)" if fr.gap_closed else "Kantilever murni")
        )
    if fr.condition == "Kondisi 2" and dr.pm_curve:
        pm = dr.pm_curve
        inside_label = ("✅ Dalam kurva P-M — Aman" if pm.inside_curve
                        else "❌ Di luar kurva P-M — Tidak Aman")
        st.info(f"📐 **P-M Check:** {inside_label}")


# =============================================================================
# 2. CROSS-SECTION DRAWING  (matplotlib)
# =============================================================================

def _fig_cross_section(inp: dict, dr: DesignResult, fr: ForceResult) -> plt.Figure:
    """
    Detailed cross-section with:
    • UD body (walls + base)
    • CU cover (if present)
    • Gap zone
    • Rebar dots at correct positions
    • Dimension annotations
    """
    # Support both old and new field names
    B_i  = inp.get("ud_inner_width",    600) / 1000
    H_i  = inp.get("ud_inner_height",   600) / 1000
    t_w  = inp.get("ud_wall_thick_top", inp.get("ud_wall_thickness", 80)) / 1000
    t_b  = inp.get("ud_slab_thick",     inp.get("ud_base_thickness", 120)) / 1000
    t_cu = inp.get("cu_thick_centre",   inp.get("cu_thickness", 100)) / 1000
    gap  = inp.get("cu_gap",            inp.get("gap_cu_ud", 20)) / 1000
    # te-cu = ta + gap (new geometry: CU sits fully on ta, no separate lap)
    te_cu = t_w + gap      # = ta + gap (in metres)
    lap   = te_cu          # compatibility: 'lap' used in y_cu_bot calculation
    cov  = inp.get("cover_clear", 30) / 1000
    cond = fr.condition
    show_cu = cond in ("Kondisi 1", "Kondisi 2")

    B_o = B_i + 2 * t_w

    fig, ax = plt.subplots(figsize=(7, 6), facecolor=C["bg"])
    ax.set_facecolor(C["bg"])
    ax.set_aspect("equal")

    # ── Coordinate origin: bottom-left of UD outer ────────────────────────────
    # y=0 at bottom of UD base, x=0 at left outer face
    # CU sits on top of UD walls with `lap` overlap (CU bottom = UD top - lap)
    y_ud_base_bot = 0
    y_ud_base_top = t_b
    y_ud_wall_top = t_b + H_i
    y_cu_bot      = y_ud_wall_top - lap  if show_cu else None
    y_cu_top      = y_cu_bot + t_cu      if show_cu else None
    y_gap_top     = y_cu_bot             if show_cu else None
    y_gap_bot     = y_cu_bot - gap       if show_cu else None

    def rect(ax, x, y, w, h, fc, ec="#546E7A", lw=1.2, alpha=1.0, hatch=None):
        kw = dict(linewidth=lw, edgecolor=ec, facecolor=fc, alpha=alpha)
        if hatch:
            kw["hatch"] = hatch
        ax.add_patch(patches.Rectangle((x, y), w, h, **kw))

    # ── UD base ───────────────────────────────────────────────────────────────
    rect(ax, 0, y_ud_base_bot, B_o, t_b, C["concrete"])
    # ── UD left wall ──────────────────────────────────────────────────────────
    rect(ax, 0, y_ud_base_top, t_w, H_i, C["concrete"])
    # ── UD right wall ─────────────────────────────────────────────────────────
    rect(ax, B_o - t_w, y_ud_base_top, t_w, H_i, C["concrete"])
    # Interior void fill (light)
    rect(ax, t_w, y_ud_base_top, B_i, H_i, "#ECEFF1", ec="none")

    # ── CU ────────────────────────────────────────────────────────────────────
    if show_cu:
        rect(ax, 0, y_cu_bot, B_o, t_cu, C["cu"], alpha=0.85)
        # Gap zones (left & right)
        rect(ax, 0, y_gap_bot, t_w, gap, "#FFF9C4", ec="#F9A825",
             lw=0.8, hatch="...", alpha=0.7)
        rect(ax, B_o - t_w, y_gap_bot, t_w, gap, "#FFF9C4", ec="#F9A825",
             lw=0.8, hatch="...", alpha=0.7)
        ax.text(t_w / 2, y_gap_bot + gap / 2,
                f"gap\n{inp.get('cu_gap', inp.get('gap_cu_ud', 20))}mm",
                ha="center", va="center", fontsize=5.5, color="#E65100")

    # ── Rebar (schematic dots) ────────────────────────────────────────────────
    r_dot = 0.004   # m radius
    # Wall: outer tension (bottom of outer face)
    for face_x in [cov + t_w * 0.0, B_o - cov]:   # left outer, right outer
        ax.add_patch(plt.Circle((face_x, y_ud_base_top + cov + 0.008),
                                r_dot, color=C["steel"], zorder=5))
    # Wall: inner tension
    for face_x in [t_w - cov - 0.008, B_o - t_w + cov + 0.008]:
        ax.add_patch(plt.Circle((face_x, y_ud_base_top + cov + 0.008),
                                r_dot, color=C["steel"], alpha=0.7, zorder=5))
    # Wall: compression (top, outer face)
    for face_x in [cov, B_o - cov]:
        ax.add_patch(plt.Circle((face_x, y_ud_wall_top - cov - 0.008),
                                r_dot * 0.7, color=C["comp_steel"], zorder=5))
    # CU: bottom tension
    if show_cu:
        for x in [t_w + cov, B_o - t_w - cov]:
            ax.add_patch(plt.Circle((x, y_cu_bot + cov + 0.008),
                                    r_dot, color=C["steel"], zorder=5))

    # ── Dimension annotations ─────────────────────────────────────────────────
    ann_kw = dict(fontsize=7, color=C["neutral"], ha="center")
    dim_kw = dict(arrowprops=dict(arrowstyle="<->", color=C["neutral"], lw=0.8),
                  fontsize=6.5, color=C["neutral"], ha="center", va="center")

    def horiz_dim(ax, x1, x2, y, label, dy=0.02):
        xm = (x1 + x2) / 2
        ax.annotate("", xy=(x2, y - dy), xytext=(x1, y - dy),
                    arrowprops=dict(arrowstyle="<->", color=C["neutral"], lw=0.8))
        ax.text(xm, y - dy - 0.015, label, **ann_kw)

    def vert_dim(ax, x, y1, y2, label, dx=0.025):
        ym = (y1 + y2) / 2
        ax.annotate("", xy=(x + dx, y2), xytext=(x + dx, y1),
                    arrowprops=dict(arrowstyle="<->", color=C["neutral"], lw=0.8))
        ax.text(x + dx + 0.018, ym, label, fontsize=6.5, color=C["neutral"],
                va="center", rotation=90)

    Bo_mm = int(B_o * 1000)
    Bi_mm = inp.get("ud_inner_width", int(B_i * 1000))
    Hi_mm = inp.get("ud_inner_height", int(H_i * 1000))
    tcu_mm = inp.get("cu_thick_centre", inp.get("cu_thickness", int(t_cu * 1000)))

    horiz_dim(ax, 0,   B_o, -0.02, f"B_o={Bo_mm}mm")
    horiz_dim(ax, t_w, B_o - t_w, -0.05, f"B_i={Bi_mm}mm")
    vert_dim(ax, B_o, y_ud_base_top, y_ud_wall_top, f"H_i={Hi_mm}mm")
    if show_cu:
        vert_dim(ax, B_o + 0.04, y_cu_bot, y_cu_top, f"t_CU={tcu_mm}mm")

    # ── Labels ────────────────────────────────────────────────────────────────
    ax.text(B_o / 2, y_ud_base_top + H_i / 2, "U-Ditch (UD)",
            ha="center", va="center", fontsize=9, color="#37474F",
            fontweight="bold", alpha=0.6)
    if show_cu:
        ax.text(B_o / 2, y_cu_bot + t_cu / 2, "Cover (CU)",
                ha="center", va="center", fontsize=8, color="white",
                fontweight="bold")

    ax.set_xlim(-0.12, B_o + 0.18)
    ax.set_ylim(-0.10, (y_cu_top + 0.08) if show_cu else (y_ud_wall_top + 0.08))
    ax.axis("off")
    ax.set_title(f"Penampang Melintang — {fr.condition}", fontsize=10,
                 color=C["neutral"], pad=8)
    fig.tight_layout()
    return fig


# =============================================================================
# 3. FORCE DIAGRAMS  (plotly)
# =============================================================================

def _fig_force_diagrams(fr: ForceResult) -> go.Figure:
    """
    Bending moment, shear, and axial force diagrams for the UD wall.
    Uses wall height as the y-axis (0 = base, H = top).
    """
    wb = fr.wall_base
    wt = fr.wall_top
    if wb is None:
        return go.Figure()

    H = fr.inp_echo.get("ud_inner_height", 600) / 1000   # m
    cond = fr.condition

    # Build (y, value) pairs — simplified linear diagrams
    y  = [0,   H]
    Mu = [wb.Mu, 0.0]
    Vu = [wb.Vu, wt.Vu if wt else 0.0]
    Nu = [wb.Nu, wt.Nu if wt else 0.0]

    if cond == "Kondisi 1" and fr.gap_closed and wt:
        # Propped cantilever: peak moment somewhere in span
        # (approx location at V=0)
        pass  # linear approx is conservative for display

    fig = go.Figure()

    # Moment diagram (horizontal axis = moment, y = height)
    fig.add_trace(go.Scatter(
        x=Mu + [0], y=y + [0],
        mode="lines", fill="tozerox",
        line=dict(color=C["moment"], width=2),
        fillcolor="rgba(21,101,192,0.15)",
        name="Mu (kN·m/m)",
    ))
    # Shear
    fig.add_trace(go.Scatter(
        x=Vu + [0], y=y + [0],
        mode="lines", fill="tozerox",
        line=dict(color=C["shear"], width=2),
        fillcolor="rgba(46,125,50,0.15)",
        name="Vu (kN/m)",
    ))
    # Axial
    fig.add_trace(go.Scatter(
        x=Nu + [0], y=y + [0],
        mode="lines", fill="tozerox",
        line=dict(color=C["axial"], width=2, dash="dash"),
        fillcolor="rgba(230,81,0,0.10)",
        name="Nu (kN/m)",
    ))

    # Annotation: values at base
    fig.add_annotation(x=Mu[0], y=0, text=f"Mu={Mu[0]:.2f}",
                       showarrow=True, arrowhead=2, font=dict(size=10, color=C["moment"]))
    fig.add_annotation(x=Vu[0], y=0, text=f"Vu={Vu[0]:.2f}",
                       showarrow=True, arrowhead=2, ay=-20,
                       font=dict(size=10, color=C["shear"]))

    fig.update_layout(
        title=dict(text=f"Diagram Gaya Dalam — {cond}", font=dict(size=13)),
        xaxis_title="Force / Moment",
        yaxis_title="Tinggi dinding (m) / Wall height (m)",
        legend=dict(orientation="h", y=1.08),
        height=420,
        plot_bgcolor=C["bg"],
        paper_bgcolor=C["bg"],
        margin=dict(l=60, r=30, t=70, b=40),
    )
    fig.update_xaxes(gridcolor=C["grid"], zeroline=True,
                     zerolinecolor="#90A4AE", zerolinewidth=1.5)
    fig.update_yaxes(gridcolor=C["grid"])
    return fig


# =============================================================================
# 4. STRESS-STRAIN BLOCK  (matplotlib)
# =============================================================================

def _fig_stress_block(cap: CapacityResult, inp: dict) -> plt.Figure:
    """
    Whitney rectangular stress-block diagram for the section at wall base.
    Shows: section | strain diagram | stress block | forces
    """
    # Support both old and new field names
    h   = inp.get("ud_wall_thick_top", inp.get("ud_wall_thickness", 80))   # mm
    d   = inp.get("d_eff_tension", inp.get("d_eff_outer", 38))             # mm
    dia_c = inp.get("rebar_comp_dia", inp.get("rebar_comp_dia", 10))
    d_p = inp.get("cover_clear", 30) + 8 + dia_c / 2                       # mm comp
    fc  = inp.get("fc_prime", 30.0)
    fy  = inp.get("fy_main", 420.0)
    As  = cap.As_prov_mm2 if cap.As_prov_mm2 > 0 else cap.As_req_mm2

    # Derived geometry
    b1  = max(0.65, min(0.85, 0.85 - 0.05 * (fc - 28) / 7))
    a   = As * fy / (0.85 * fc * 1000)   # mm (b=1000mm/m)
    c   = a / b1
    et  = 0.003 * (d - c) / c if c > 0 else 0.003
    eu  = 0.003

    # ── Figure layout: 4 sub-panels ──────────────────────────────────────────
    fig, axes = plt.subplots(1, 4, figsize=(10, 5), facecolor=C["bg"])
    fig.subplots_adjust(wspace=0.5)
    for ax in axes:
        ax.set_facecolor(C["bg"])

    H_sc = h      # mm  (section height for drawing)
    b_sc = 60     # mm  (arbitrary section width for drawing)

    # Panel 0: Cross-section
    ax0 = axes[0]
    rect0 = patches.Rectangle((0, 0), b_sc, H_sc,
                               fc=C["concrete"], ec=C["neutral"], lw=1.5)
    ax0.add_patch(rect0)
    # Tension bar (bottom)
    ax0.add_patch(plt.Circle((b_sc / 2, H_sc - d + d_p * 0.4), 4,
                             color=C["steel"], zorder=5))
    # Compression bar (top)
    ax0.add_patch(plt.Circle((b_sc / 2, d_p), 3,
                             color=C["comp_steel"], zorder=5))
    # dim lines
    ax0.annotate("", xy=(b_sc + 10, H_sc - d), xytext=(b_sc + 10, H_sc),
                 arrowprops=dict(arrowstyle="<->", color=C["neutral"], lw=0.8))
    ax0.text(b_sc + 22, H_sc - d / 2, f"d={d:.0f}mm", fontsize=7,
             va="center", color=C["neutral"], rotation=90)
    ax0.set_xlim(-5, b_sc + 40); ax0.set_ylim(-10, H_sc + 10)
    ax0.axis("off"); ax0.set_title("Penampang", fontsize=8)

    # Panel 1: Strain diagram
    ax1 = axes[1]
    eu_draw = 0.003
    et_draw = et
    ax1.plot([eu_draw, 0, et_draw], [0, c, H_sc - d + d_p * 0.4],
             color=C["neutral"], lw=1.8, marker="o", ms=4)
    ax1.axhline(c, color=C["neutral"], lw=0.6, ls="--")
    ax1.axvline(0, color=C["neutral"], lw=0.8)
    ax1.text(eu_draw, 2, f"εu={eu_draw}", fontsize=7, color="#C62828")
    ax1.text(et_draw if et_draw > 0 else -0.001,
             H_sc - d + d_p * 0.4 + 3, f"εt={et:.4f}", fontsize=7,
             color=C["steel"] if et >= 0.005 else C["axial"])
    ax1.text(-0.002, c + 2, f"c={c:.1f}mm", fontsize=7, color=C["neutral"])
    ax1.set_xlim(-0.005, max(eu_draw, abs(et_draw)) + 0.002)
    ax1.set_ylim(-5, H_sc + 8)
    ax1.set_xlabel("Regangan / Strain", fontsize=7)
    ax1.set_title("Diagram Regangan", fontsize=8)
    ax1.tick_params(labelsize=6); ax1.invert_yaxis()

    # Panel 2: Stress block
    ax2 = axes[2]
    stress_w = 0.85 * fc
    # Compression block (top, height = a)
    rect2 = patches.Rectangle((0, 0), stress_w, a,
                               fc="#1E88E5", ec=C["neutral"], alpha=0.5, lw=1)
    ax2.add_patch(rect2)
    ax2.annotate("", xy=(stress_w + 5, 0), xytext=(stress_w + 5, a),
                 arrowprops=dict(arrowstyle="<->", color=C["neutral"], lw=0.8))
    ax2.text(stress_w + 12, a / 2, f"a={a:.1f}mm", fontsize=7, va="center",
             color=C["neutral"], rotation=90)
    ax2.text(stress_w / 2, a / 2, f"0.85f'c\n={stress_w:.1f}MPa",
             ha="center", va="center", fontsize=6.5, color="white")
    ax2.set_xlim(-5, stress_w + 40)
    ax2.set_ylim(-5, H_sc + 8)
    ax2.set_xlabel("Tegangan / Stress (MPa)", fontsize=7)
    ax2.set_title("Blok Tegangan", fontsize=8)
    ax2.tick_params(labelsize=6); ax2.invert_yaxis()

    # Panel 3: Force resultants
    ax3 = axes[3]
    Cc = 0.85 * fc * a * 1.0    # kN per mm width → ×1000 for per m
    T  = As * fy / 1000          # kN per m (As mm²/m, fy MPa → kN/m)
    Mn = As * fy * (d - a / 2) / 1e6   # kN·m/m

    ax3.annotate("", xy=(0.5, a / 2), xytext=(0.5, 0),
                 arrowprops=dict(arrowstyle="->", color=C["comp_steel"], lw=2))
    ax3.text(0.7, a / 3, f"Cc={Cc * 1000:.0f}kN/m", fontsize=7, color=C["comp_steel"])

    ax3.annotate("", xy=(0.5, H_sc - d + d_p * 0.4), xytext=(0.5, H_sc),
                 arrowprops=dict(arrowstyle="->", color=C["steel"], lw=2))
    ax3.text(0.7, H_sc - d / 2, f"T={T:.0f}kN/m", fontsize=7, color=C["steel"])

    ax3.text(0.5, H_sc / 2,
             f"Mn={Mn:.2f}\nkN·m/m\n\nφMn={cap.phi_Mn_kNm:.2f}\nkN·m/m",
             ha="center", va="center", fontsize=8, color=C["neutral"],
             bbox=dict(boxstyle="round,pad=0.3", fc="white", ec=C["neutral"], alpha=0.7))
    ax3.set_xlim(0, 1.5); ax3.set_ylim(-5, H_sc + 8)
    ax3.axis("off"); ax3.set_title("Resultan Gaya", fontsize=8)
    ax3.invert_yaxis()

    fig.suptitle(f"Blok Tegangan Whitney — {cap.label}",
                 fontsize=10, color=C["neutral"], y=1.02)
    return fig


# =============================================================================
# 5. P-M INTERACTION DIAGRAM  (plotly)
# =============================================================================

def _fig_pm_curve(pm: PMCurve) -> go.Figure:
    """
    P-M interaction diagram with:
    • Nominal curve (grey)
    • φ-reduced design curve (blue)
    • Control points labelled
    • Demand point (red star)
    """
    # Nominal curve
    Mn_nom = [p.Mn for p in pm.points]
    Pn_nom = [p.Pn for p in pm.points]
    # Design curve
    Mn_des = [p.phi_Mn for p in pm.points]
    Pn_des = [p.phi_Pn for p in pm.points]

    fig = go.Figure()

    fig.add_trace(go.Scatter(
        x=Mn_nom, y=Pn_nom,
        mode="lines",
        line=dict(color="#90A4AE", width=1.5, dash="dot"),
        name="Nominal (Pn-Mn)",
    ))
    fig.add_trace(go.Scatter(
        x=Mn_des, y=Pn_des,
        mode="lines", fill="toself",
        line=dict(color=C["pm_curve"], width=2.5),
        fillcolor="rgba(21,101,192,0.08)",
        name="Design (φPn-φMn)",
    ))

    # Control points
    ctrl = [
        (0,         pm.Pn_max,      "Pn,max (Pure Comp.)", "square"),
        (pm.Mb,     pm.Pb,          "Balanced",            "diamond"),
        (pm.Mn_pure, 0,             "Pure Flexure",        "circle"),
        (0,         pm.Pn_tension,  "Pure Tension",        "triangle-down"),
    ]
    for mx, px, lbl, sym in ctrl:
        fig.add_trace(go.Scatter(
            x=[mx], y=[px], mode="markers+text",
            marker=dict(symbol=sym, size=10,
                        color=C["pm_curve"], line=dict(color="white", width=1.5)),
            text=[lbl], textposition="middle right",
            textfont=dict(size=10),
            name=lbl, showlegend=False,
        ))

    # Demand point
    inside_col = C["ok"] if pm.inside_curve else C["ng"]
    inside_sym = "star" if pm.inside_curve else "x"
    fig.add_trace(go.Scatter(
        x=[pm.Mu_demand], y=[pm.Nu_demand],
        mode="markers+text",
        marker=dict(symbol=inside_sym, size=14, color=inside_col,
                    line=dict(color="white", width=1.5)),
        text=[f"Demand\n({pm.Mu_demand:.1f},{pm.Nu_demand:.1f})"],
        textposition="top right",
        textfont=dict(size=10, color=inside_col),
        name="Demand (Mu, Nu)",
    ))

    fig.update_layout(
        title=dict(
            text=(f"Diagram Interaksi P-M — Dinding UD (Kondisi 2)<br>"
                  f"<sub>{'✅ Demand dalam kurva' if pm.inside_curve else '❌ Demand di luar kurva'}</sub>"),
            font=dict(size=13),
        ),
        xaxis_title="Mn / φMn  (kN·m/m)",
        yaxis_title="Pn / φPn  (kN/m)  [+ = Tekan]",
        legend=dict(orientation="v", x=1.01),
        height=520,
        plot_bgcolor=C["bg"],
        paper_bgcolor=C["bg"],
        margin=dict(l=70, r=160, t=90, b=50),
    )
    fig.update_xaxes(gridcolor=C["grid"], zeroline=True,
                     zerolinecolor="#90A4AE", rangemode="tozero")
    fig.update_yaxes(gridcolor=C["grid"], zeroline=True,
                     zerolinecolor="#90A4AE")
    return fig


# =============================================================================
# 6. CAPACITY CHECK BAR CHART  (plotly)
# =============================================================================

def _fig_capacity_bars(dr: DesignResult) -> go.Figure:
    """
    Horizontal bar chart: Demand vs Capacity for all checked sections.
    """
    labels, demand, capacity, colors = [], [], [], []

    sections = [
        (dr.wall_base_cap,  "Wall Base Flexure", "Mu",  "phi_Mn_kNm"),
        (dr.wall_base_cap,  "Wall Base Shear",   "Vu",  "phi_Vn_kNm"),
        (dr.wall_base_dbl,  "Wall Base (2-rebar)","Mu_kNmm","phi_Mn_kNm"),
        (dr.cu_midspan_cap, "CU Midspan Flex.",  "Mu_kNmm","phi_Mn_kNm"),
        (dr.cu_support_cap, "CU Support Shear",  "Vu_kNm", "phi_Vn_kNm"),
        (dr.base_slab_cap,  "Base Slab Flexure", "Mu_kNmm","phi_Mn_kNm"),
        (dr.base_slab_cap,  "Base Slab Shear",   "Vu_kNm", "phi_Vn_kNm"),
    ]
    for cap, lbl, dem_attr, cap_attr in sections:
        if cap is None:
            continue
        d_val = getattr(cap, dem_attr, 0.0)
        c_val = getattr(cap, cap_attr, 0.0)
        if d_val == 0.0 and c_val == 0.0:
            continue
        labels.append(lbl)
        demand.append(abs(d_val))
        capacity.append(abs(c_val))
        colors.append(C["ok"] if c_val >= d_val else C["ng"])

    fig = go.Figure()
    fig.add_trace(go.Bar(
        x=capacity, y=labels, orientation="h",
        marker_color=colors, name="φ Kapasitas",
        opacity=0.75,
    ))
    fig.add_trace(go.Bar(
        x=demand, y=labels, orientation="h",
        marker_color=["rgba(0,0,0,0.6)"] * len(labels),
        marker_pattern_shape="/", name="Demand (Mu/Vu)",
        opacity=0.6,
    ))
    fig.update_layout(
        barmode="overlay",
        title="Kapasitas vs Gaya Dalam / Capacity vs Demand",
        xaxis_title="kN·m/m  atau  kN/m",
        height=max(300, len(labels) * 48 + 100),
        plot_bgcolor=C["bg"], paper_bgcolor=C["bg"],
        legend=dict(orientation="h", y=1.08),
        margin=dict(l=180, r=30, t=70, b=40),
    )
    fig.update_xaxes(gridcolor=C["grid"])
    return fig


# =============================================================================
# 7. STEP-BY-STEP CALCULATION DISPLAY
# =============================================================================


# =============================================================================
# TAHAP C — BOOK-STYLE CALCULATION DISPLAY
# =============================================================================
# Format tiap langkah:
#   [Header seksi]
#   Simbol — Deskripsi
#   Formula (LaTeX)
#   = Substitusi → Hasil  Satuan  [Ref. Code]
#
# Urutan tampilan:
#   A. Data & Geometri
#   B. Tekanan Tanah & Kendaraan  (+ diagram beban)
#   C. Gaya Dalam Mu, Vu, Nu     (+ diagram gaya dalam)
#   D. Kapasitas Penampang        (+ diagram blok tegangan)
#   E. Kontrol & Kesimpulan
# =============================================================================

# ── Section dividers used as marker keys in CalcStep.symbol ──────────────────
_SEC_MARKERS = {
    "Ka":            "B",   # start of section B
    "Mu,cant":       "C",   # start of section C
    "As,min":        "D",   # start of section D
    "Cek Lentur":    "E",   # start of section E (check)
    "Cek Geser Final":"E",
    "Cek P-M":       "E",
}

_SEC_LABELS = {
    "ID": {
        "A": "📐 A. Data & Geometri",
        "B": "🌍 B. Tekanan Tanah & Beban Kendaraan",
        "C": "⚡ C. Gaya Dalam — Mu, Vu, Nu",
        "D": "🔩 D. Kapasitas Penampang",
        "E": "✅ E. Kontrol & Kesimpulan",
    },
    "EN": {
        "A": "📐 A. Data & Geometry",
        "B": "🌍 B. Earth & Vehicle Pressure",
        "C": "⚡ C. Internal Forces — Mu, Vu, Nu",
        "D": "🔩 D. Section Capacity",
        "E": "✅ E. Checks & Conclusions",
    },
}


def _current_section(symbol: str, current: str) -> str:
    """Detect section transitions from CalcStep symbol."""
    return _SEC_MARKERS.get(symbol, current)


def _render_one_step(s, idx: int) -> None:
    """
    Render a single CalcStep in book style inside Streamlit.
    Format:
        Simbol — Deskripsi                    [Ref. Code]
        LaTeX formula
        = Substitusi → Result  unit
    """
    # Determine background for OK/NG rows
    is_ok = "✅" in s.substitution or "OK" in s.substitution
    is_ng = "❌" in s.substitution or "NG" in s.substitution

    if is_ok:
        bg = "#E8F5E9"
    elif is_ng:
        bg = "#FFEBEE"
    else:
        bg = "#F8F9FA" if idx % 2 == 0 else "white"

    # Box container per step
    with st.container():
        st.markdown(
            f'<div style="background:{bg};border-left:3px solid #1A237E;'
            f'padding:8px 14px;margin-bottom:6px;border-radius:4px;">'
            f'<span style="font-weight:700;color:#1A237E;font-size:13px;">'
            f'{s.symbol}</span>'
            f'<span style="color:#546E7A;font-size:12px;"> — {s.description}</span>'
            f'<span style="float:right;color:#9C27B0;font-size:10px;font-style:italic;">'
            f'{s.code_ref}</span>'
            f'</div>',
            unsafe_allow_html=True,
        )
        # LaTeX formula
        try:
            if s.formula and s.formula.strip():
                st.latex(s.formula)
        except Exception:
            st.code(s.formula, language=None)

        # Substitution → Result line
        result_str = f"{s.result:.4f}".rstrip("0").rstrip(".") if s.result != 0 else "0"
        # Truncate long substitution strings for readability
        sub_disp = s.substitution if len(s.substitution) <= 120 else s.substitution[:117] + "…"

        ok_badge = ""
        if is_ok:
            ok_badge = ' <span style="color:#2E7D32;font-weight:700;">✅ OK</span>'
        elif is_ng:
            ok_badge = ' <span style="color:#C62828;font-weight:700;">❌ NG</span>'

        st.markdown(
            f'<div style="padding:2px 14px 8px 14px;">'
            f'<code style="background:#EDE7F6;padding:2px 6px;border-radius:3px;'
            f'font-size:12px;">{sub_disp}</code>'
            f' → <strong style="color:#1565C0;font-size:13px;">{result_str}</strong>'
            f' <span style="color:#455A64;font-size:11px;">{s.unit}</span>'
            f'{ok_badge}</div>',
            unsafe_allow_html=True,
        )


def _render_section_header(label: str) -> None:
    st.markdown(
        f'<div style="background:#1A237E;color:white;padding:10px 16px;'
        f'margin:18px 0 10px 0;border-radius:6px;font-size:14px;font-weight:700;">'
        f'{label}</div>',
        unsafe_allow_html=True,
    )


def _render_calc_steps_book(steps: list, lang: str,
                             fr: "ForceResult" = None,
                             inp: dict = None,
                             dr: "DesignResult" = None) -> None:
    """
    Main Tahap C renderer.
    Shows FULL CHAIN:
    A. Data & Geometri
    B. Tekanan lateral (derivation showing WHERE pressures come from)
       → Load pressure diagram
    C. Gaya Dalam (derivation showing HOW Mu/Vu/Nu are obtained)
       → Mu/Vu/Nu diagrams along wall height
    D. Kapasitas Penampang
       → Stress block diagram
    E. Kontrol & Kesimpulan
    """
    if not steps:
        st.info("Tidak ada langkah perhitungan." if lang == "ID"
                else "No calculation steps available.")
        return

    sec_labels = _SEC_LABELS.get(lang, _SEC_LABELS["ID"])
    current_sec = "A"
    _render_section_header(sec_labels["A"])

    # Show brief data summary from CalcStep list (section A steps)
    for i, s in enumerate(steps):
        new_sec = _current_section(s.symbol, current_sec)

        if new_sec != current_sec:
            current_sec = new_sec
            _render_section_header(sec_labels.get(current_sec, current_sec))

            # ── After B header: full load derivation + pressure diagram ──────
            if current_sec == "B" and fr is not None and inp is not None:
                _render_load_derivation(fr, inp, lang)
                st.markdown(
                    f"**{'📊 Diagram Distribusi Tekanan Lateral' if lang=='ID' else '📊 Lateral Pressure Distribution Diagram'}**"
                )
                try:
                    fig_p = _fig_pressure_full(fr, inp, lang)
                    st.pyplot(fig_p, use_container_width=True)
                    plt.close(fig_p)
                except Exception as e:
                    st.caption(f"Diagram tidak tersedia: {e}")
                st.divider()

            # ── After C header: force derivation + Mu/Vu/Nu diagrams ─────────
            if current_sec == "C" and fr is not None and inp is not None:
                _render_force_derivation(fr, inp, lang)
                st.markdown(
                    f"**{'📊 Diagram Gaya Dalam Sepanjang Tinggi Dinding' if lang=='ID' else '📊 Internal Forces Diagram Along Wall Height'}**"
                )
                try:
                    fig_f = _fig_mvu_full(fr, inp, lang)
                    st.pyplot(fig_f, use_container_width=True)
                    plt.close(fig_f)
                except Exception as e:
                    st.caption(f"Diagram tidak tersedia: {e}")
                st.divider()

            # ── After D header: stress block diagram ─────────────────────────
            if current_sec == "D" and dr is not None and inp is not None:
                cap = dr.wall_base_cap or dr.wall_base_dbl
                if cap:
                    st.markdown(
                        f"**{'📊 Blok Tegangan Whitney' if lang=='ID' else '📊 Whitney Stress Block'}**"
                    )
                    try:
                        fig_sb = _fig_stress_block(cap, inp)
                        st.pyplot(fig_sb, use_container_width=True)
                        plt.close(fig_sb)
                    except Exception as e:
                        st.caption(f"Diagram tidak tersedia: {e}")
                st.divider()

        _render_one_step(s, i)

    # ── P-M diagram at end (Kondisi 2) ────────────────────────────────────────
    if fr is not None and fr.condition == "Kondisi 2" and dr is not None and dr.pm_curve:
        _render_section_header(
            "📈 Diagram Interaksi P-M — Dinding UD sebagai Kolom"
            if lang == "ID" else
            "📈 P-M Interaction Diagram — UD Wall as Column"
        )
        fig_pm = _fig_pm_curve(dr.pm_curve)
        st.plotly_chart(fig_pm, use_container_width=True, key="pm_book")
        plt.close("all")


# =============================================================================
# TAHAP C — INLINE DIAGRAMS (rendered inside calculation flow)
# =============================================================================

def _render_inline_load_diagram(fr: "ForceResult", inp: dict, lang: str) -> None:
    """
    Inline load pressure diagram.
    Shows: triangular earth pressure + rectangular surcharge + (Boussinesq curve).
    Rendered with matplotlib inside the calculation flow.
    """
    st.markdown(
        f"**{'📊 Diagram Tekanan Lateral' if lang=='ID' else '📊 Lateral Pressure Diagram'}**"
    )
    try:
        fig = _fig_load_pressure(fr, inp, lang)
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
    except Exception as e:
        st.caption(f"Diagram tidak tersedia: {e}")


def _render_inline_force_diagram(fr: "ForceResult", lang: str) -> None:
    """
    Inline bending moment + shear force diagram along wall height.
    """
    st.markdown(
        f"**{'📊 Diagram Gaya Dalam' if lang=='ID' else '📊 Internal Forces Diagram'}**"
    )
    try:
        fig = _fig_wall_forces(fr, lang)
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
    except Exception as e:
        st.caption(f"Diagram tidak tersedia: {e}")


def _render_inline_stress_block(cap: "CapacityResult", inp: dict, lang: str) -> None:
    """
    Inline Whitney stress-block diagram.
    """
    st.markdown(
        f"**{'📊 Blok Tegangan Whitney' if lang=='ID' else '📊 Whitney Stress Block'}**"
    )
    try:
        fig = _fig_stress_block(cap, inp)
        st.pyplot(fig, use_container_width=True)
        plt.close(fig)
    except Exception as e:
        st.caption(f"Diagram tidak tersedia: {e}")


# =============================================================================
# TAHAP C — LOAD PRESSURE DIAGRAM  (matplotlib, engineering quality)
# =============================================================================

def _fig_load_pressure(fr: "ForceResult", inp: dict, lang: str) -> plt.Figure:
    """
    Professional lateral pressure diagram.
    Left panel : wall elevation with pressure distributions overlaid.
    Right panel: pressure vs height graph (numerical values).

    Distributions shown:
      • Triangular — earth pressure  (γs·Ka·z)
      • Rectangular — surcharge/timbunan
      • Boussinesq curve — vehicle point load (if selected)
    """
    from uditch.calc_engine import _compat, _rankine_Ka, _aashto_heq, _boussinesq_point_lateral
    inp_c = _compat(inp)

    H     = inp_c.get("ud_inner_height", 600) / 1000.0   # m
    Ka    = fr.lateral.Ka if fr.lateral else _rankine_Ka(inp_c.get("phi_soil", 30))
    gamma_s = inp_c.get("gamma_s", 18.0)
    c_soil  = inp_c.get("cohesion", 0.0)
    method  = inp_c.get("lat_method_idx", 0)
    P_wheel = inp_c.get("wheel_load", 112.5)
    x_m     = max(inp_c.get("wheel_dist", 0.25), 0.05)
    L_seg   = inp_c.get("ud_length", 1.2)
    q_bsd   = inp_c.get("udl_beside", 0.0)

    # ── Build pressure profiles ───────────────────────────────────────────────
    n_pts = 80
    z_arr = [i * H / n_pts for i in range(n_pts + 1)]

    def sigma_earth(z):
        s = Ka * gamma_s * z - 2 * c_soil * math.sqrt(Ka)
        return max(s, 0.0)

    sig_earth = [sigma_earth(z) for z in z_arr]

    if method == 0:
        # Surcharge: uniform = Ka*gamma_s*heq + Ka*q_beside
        heq = _aashto_heq(H)
        sig_sur = [Ka * gamma_s * heq + Ka * q_bsd] * len(z_arr)
        sig_bou = None
    else:
        # Boussinesq per-unit-height
        sig_bou = [_boussinesq_point_lateral(P_wheel, x_m, max(z, 0.001), L_seg) / L_seg
                   for z in z_arr]
        sig_sur = [Ka * q_bsd] * len(z_arr) if q_bsd > 0 else None

    # ── Figure ────────────────────────────────────────────────────────────────
    fig, (ax_wall, ax_plot) = plt.subplots(1, 2, figsize=(11, 5.5),
                                            facecolor=C["bg"],
                                            gridspec_kw={"width_ratios": [1, 2]})

    # ─── Left panel: wall elevation sketch with loads ─────────────────────────
    ax_wall.set_facecolor(C["bg"])
    ta_m = inp_c.get("ud_wall_thickness", 80) / 1000.0
    ts_m = inp_c.get("ud_base_thickness", 120) / 1000.0

    # Wall body
    wall_patch = patches.Rectangle((-ta_m, 0), ta_m, H,
                                    fc=C["concrete"], ec="#546E7A", lw=1.5)
    ax_wall.add_patch(wall_patch)
    # Base slab
    slab_patch = patches.Rectangle((-ta_m, -ts_m), ta_m * 3, ts_m,
                                    fc=C["concrete"], ec="#546E7A", lw=1.2)
    ax_wall.add_patch(slab_patch)

    # Soil hatch (right of wall)
    soil_patch = patches.Rectangle((0, 0), 0.4, H,
                                    fc="#A5D6A7", ec="none", alpha=0.35, hatch="///")
    ax_wall.add_patch(soil_patch)

    # Earth pressure (triangular fill)
    scale = 0.006    # kPa → display units
    xs_e = [s * scale for s in sig_earth]
    # Draw filled area from wall face (x=0) to pressure value
    ax_wall.fill_betweenx(z_arr, 0, xs_e, alpha=0.35, color="#7B1FA2", label="")
    ax_wall.plot(xs_e, z_arr, color="#7B1FA2", lw=2)

    if method == 0 and sig_sur:
        xs_s = [s * scale for s in sig_sur]
        ax_wall.fill_betweenx(z_arr, xs_e, [xe + xs for xe, xs in zip(xs_e, xs_s)],
                              alpha=0.25, color="#F57F17")
        ax_wall.plot([xe + xs for xe, xs in zip(xs_e, xs_s)], z_arr,
                     color="#F57F17", lw=1.5, ls="--")

    if method == 1 and sig_bou:
        xs_b = [s * scale for s in sig_bou]
        ax_wall.fill_betweenx(z_arr, xs_e, [xe + xb for xe, xb in zip(xs_e, xs_b)],
                              alpha=0.25, color="#D32F2F")
        ax_wall.plot([xe + xb for xe, xb in zip(xs_e, xs_b)], z_arr,
                     color="#D32F2F", lw=1.5, ls="--")

    # Wheel load arrow
    wx_arr = -ta_m - x_m
    ax_wall.annotate("", xy=(-ta_m, H - 0.05), xytext=(wx_arr, H + 0.12),
                     arrowprops=dict(arrowstyle="->", color=C["load"], lw=2))
    ax_wall.text(wx_arr, H + 0.18,
                 f"P={P_wheel:.0f}kN\nx={x_m:.2f}m",
                 ha="center", fontsize=8, color=C["load"])

    ax_wall.set_xlim(-ta_m - 0.4, max(xs_e) + 0.05 if xs_e else 0.3)
    ax_wall.set_ylim(-ts_m - 0.05, H + 0.35)
    ax_wall.set_xlabel("Tekanan lateral →", fontsize=8, color=C["neutral"])
    ax_wall.set_ylabel("z dari dasar dinding (m)", fontsize=8, color=C["neutral"])
    ax_wall.set_title("Elevasi Dinding", fontsize=9, color=C["neutral"])
    ax_wall.invert_yaxis()
    ax_wall.tick_params(labelsize=7)

    # ─── Right panel: pressure vs depth plot ──────────────────────────────────
    ax_plot.set_facecolor(C["bg"])
    ax_plot.set_xlabel("Tekanan lateral σh (kPa)", fontsize=9, color=C["neutral"])
    ax_plot.set_ylabel("Kedalaman z dari puncak dinding (m)", fontsize=9, color=C["neutral"])
    ax_plot.invert_yaxis()

    # Earth pressure
    ax_plot.fill_betweenx(z_arr, 0, sig_earth, alpha=0.35, color="#7B1FA2")
    ax_plot.plot(sig_earth, z_arr, color="#7B1FA2", lw=2,
                 label="Tanah aktif" if lang == "ID" else "Active earth")

    # Annotate at base
    sigma_base_val = sig_earth[-1]
    ax_plot.annotate(f"{sigma_base_val:.1f} kPa",
                     xy=(sigma_base_val, z_arr[-1]),
                     xytext=(sigma_base_val + 1, z_arr[-1] - 0.05),
                     fontsize=8, color="#7B1FA2",
                     arrowprops=dict(arrowstyle="-", color="#7B1FA2", lw=0.8))

    if method == 0 and sig_sur:
        total = [se + ss for se, ss in zip(sig_earth, sig_sur)]
        ax_plot.fill_betweenx(z_arr, sig_earth, total, alpha=0.30, color="#F57F17")
        ax_plot.plot(total, z_arr, color="#F57F17", lw=2, ls="--",
                     label=f"+ Surcharge (heq={_aashto_heq(H):.2f}m)" if lang=="ID"
                           else f"+ Surcharge (heq={_aashto_heq(H):.2f}m)")
        ax_plot.annotate(f"{total[-1]:.1f} kPa",
                         xy=(total[-1], z_arr[-1]),
                         xytext=(total[-1] + 1, z_arr[-1] - 0.05),
                         fontsize=8, color="#F57F17")

    elif method == 1 and sig_bou:
        total = [se + sb for se, sb in zip(sig_earth, sig_bou)]
        ax_plot.fill_betweenx(z_arr, sig_earth, total, alpha=0.30, color="#D32F2F")
        ax_plot.plot(total, z_arr, color="#D32F2F", lw=2, ls="--",
                     label=f"+ Boussinesq P={P_wheel:.0f}kN x={x_m:.2f}m")
        peak_idx = sig_bou.index(max(sig_bou))
        ax_plot.annotate(f"max {max(sig_bou):.1f} kPa",
                         xy=(total[peak_idx], z_arr[peak_idx]),
                         xytext=(total[peak_idx] + 1, z_arr[peak_idx]),
                         fontsize=8, color="#D32F2F")

    # Resultant force line
    F_total = fr.lateral.F_total if fr.lateral else 0
    arm     = (fr.lateral.arm_earth if fr.lateral else H / 3)
    # Mark resultant
    ax_plot.axhline(H - arm, color=C["axial"], lw=1, ls=":", alpha=0.7)
    ax_plot.text(0.5, H - arm - 0.02,
                 f"F={F_total:.2f}kN/m @ z={H-arm:.2f}m",
                 fontsize=8, color=C["axial"])

    ax_plot.legend(fontsize=8, loc="lower right")
    ax_plot.grid(color=C["grid"], lw=0.5)
    ax_plot.tick_params(labelsize=7)

    cond_title = (
        "Distribusi Tekanan Lateral pada Dinding UD"
        if lang == "ID" else
        "Lateral Pressure Distribution on UD Wall"
    )
    fig.suptitle(cond_title, fontsize=11, color=C["neutral"], y=1.01)
    fig.tight_layout()
    return fig


# =============================================================================
# TAHAP C — INTERNAL FORCES DIAGRAM  (matplotlib)
# =============================================================================

def _fig_wall_forces(fr: "ForceResult", lang: str) -> plt.Figure:
    """
    Bending moment and shear force diagrams along wall height.
    Engineering convention: wall as vertical column, z from base (0) to top (H).
    M diagram drawn to the right (tension on outer face = positive).
    V diagram drawn to the right.

    For cantilever: M(z) = ∫V·dz from top; V(z) = integral of distributed load.
    For propped: show both cantilever (dashed) and propped (solid).
    """
    wb = fr.wall_base
    wt = fr.wall_top
    if wb is None:
        return plt.figure(figsize=(8, 4))

    H = fr.inp_echo.get("ud_inner_height", 600) / 1000.0  # m
    n = 60
    z = [i * H / n for i in range(n + 1)]   # 0=base, H=top

    lat = fr.lateral
    if lat is None:
        return plt.figure(figsize=(8, 4))

    # Reconstruct distributed pressure profile p(z) from base:
    Ka      = lat.Ka
    gamma_s = fr.inp_echo.get("gamma_s", 18.0)
    c_soil  = fr.inp_echo.get("cohesion", 0.0)
    gamma_DL = fr.inp_echo.get("gamma_DL", 1.2)
    gamma_LL = fr.inp_echo.get("gamma_LL", 1.6)

    def sigma_e(zi):
        return max(Ka * gamma_s * (H - zi) - 2 * c_soil * math.sqrt(Ka), 0.0)

    sig_sur_per_m = lat.F_surcharge / H if H > 0 else 0.0
    # Factored total pressure at each z from base
    def p_factored(zi):
        return gamma_DL * sigma_e(zi) + gamma_LL * sig_sur_per_m

    # Shear at section z from base (cantilever: V(z) = ∫_z^H p(s)ds)
    def V_cant(zi):
        # Numeric integration from z to H
        dz = (H - zi) / 40
        return sum(p_factored(zi + (j + 0.5) * dz) * dz for j in range(40))

    # Moment at section z from base (cantilever: M(z) = ∫_z^H p(s)·(s-z)ds)
    def M_cant(zi):
        dz = (H - zi) / 40
        return sum(p_factored(zi + (j + 0.5) * dz) * ((j + 0.5) * dz) * dz
                   for j in range(40))

    V_arr = [V_cant(zi) for zi in z]
    M_arr = [M_cant(zi) for zi in z]

    # Propped cantilever correction (if gap closed)
    V_prop, M_prop = None, None
    if fr.gap_closed and wt is not None:
        H_strut = wt.Vu   # prop force at top (kN/m)
        V_prop  = [V_cant(zi) - H_strut * (1.0) for zi in z]  # subtract prop
        M_prop  = [M_cant(zi) - H_strut * (H - zi) for zi in z]

    # ── Figure ────────────────────────────────────────────────────────────────
    fig, (ax_V, ax_M) = plt.subplots(1, 2, figsize=(11, 5.5),
                                      facecolor=C["bg"])

    def draw_wall(ax):
        ax.set_facecolor(C["bg"])
        ax.axvline(0, color=C["neutral"], lw=2.5)
        ax.axhline(0, color=C["neutral"], lw=1, ls="--", alpha=0.5)   # base
        ax.axhline(H, color=C["neutral"], lw=0.8, ls="--", alpha=0.4)  # top
        ax.set_ylabel("z dari dasar (m)" if lang=="ID" else "z from base (m)",
                      fontsize=9, color=C["neutral"])
        ax.tick_params(labelsize=7)
        ax.grid(color=C["grid"], lw=0.5)
        ax.invert_yaxis()

    # ── Shear diagram ─────────────────────────────────────────────────────────
    draw_wall(ax_V)
    ax_V.fill_betweenx(z, 0, V_arr, alpha=0.30, color=C["shear"])
    ax_V.plot(V_arr, z, color=C["shear"], lw=2,
              label="Vu (kantilever)" if lang=="ID" else "Vu (cantilever)")
    if V_prop:
        ax_V.plot(V_prop, z, color=C["shear"], lw=2, ls="--",
                  label="Vu (propped)")
        ax_V.fill_betweenx(z, 0, V_prop, alpha=0.15, color=C["shear"])

    # Mark max and base values
    ax_V.annotate(f"Vu={wb.Vu:.2f}kN/m",
                  xy=(V_arr[0], 0), xytext=(V_arr[0] * 0.6 + 0.5, 0.05 * H),
                  fontsize=8, color=C["shear"],
                  arrowprops=dict(arrowstyle="->", color=C["shear"], lw=0.8))
    ax_V.set_xlabel("Vu (kN/m)", fontsize=9, color=C["shear"])
    ax_V.set_title("Diagram Geser / Shear", fontsize=10, color=C["neutral"])
    ax_V.legend(fontsize=8)

    # ── Moment diagram ────────────────────────────────────────────────────────
    draw_wall(ax_M)
    ax_M.fill_betweenx(z, 0, M_arr, alpha=0.30, color=C["moment"])
    ax_M.plot(M_arr, z, color=C["moment"], lw=2,
              label="Mu (kantilever)" if lang=="ID" else "Mu (cantilever)")
    if M_prop:
        ax_M.plot(M_prop, z, color=C["moment"], lw=2, ls="--",
                  label="Mu (propped)")
        ax_M.fill_betweenx(z, 0, M_prop, alpha=0.15, color=C["moment"])

    ax_M.annotate(f"Mu={wb.Mu:.2f}kN·m/m",
                  xy=(M_arr[0], 0), xytext=(M_arr[0] * 0.5 + 0.2, 0.08 * H),
                  fontsize=8, color=C["moment"],
                  arrowprops=dict(arrowstyle="->", color=C["moment"], lw=0.8))

    if fr.gap_closed and M_prop:
        Mmax_prop = max(M_prop, key=abs)
        z_Mmax    = z[M_prop.index(Mmax_prop)]
        ax_M.plot(Mmax_prop, z_Mmax, "o", color=C["moment"], ms=7)
        ax_M.annotate(f"Mu,max={Mmax_prop:.2f}",
                      xy=(Mmax_prop, z_Mmax),
                      xytext=(Mmax_prop + 0.3, z_Mmax - 0.05 * H),
                      fontsize=8, color=C["moment"])

    ax_M.set_xlabel("Mu (kN·m/m)", fontsize=9, color=C["moment"])
    ax_M.set_title("Diagram Momen / Moment", fontsize=10, color=C["neutral"])
    ax_M.legend(fontsize=8)
    ax_M.axvline(0, color=C["neutral"], lw=2)

    title_str = (f"Gaya Dalam Dinding UD — {fr.condition}" if lang == "ID"
                 else f"UD Wall Internal Forces — {fr.condition}")
    fig.suptitle(title_str, fontsize=11, color=C["neutral"])
    fig.tight_layout()
    return fig


def _section_calc_steps(steps: list, lang: str) -> None:
    """
    Render each CalcStep as a formatted expandable table row.
    Format: Symbol | Description | Formula → Substitution → Result | Unit | Code
    """
    st.markdown(
        f"### {'🧮 Detail Langkah Perhitungan' if lang=='ID' else '🧮 Step-by-Step Calculation'}"
    )

    if not steps:
        st.info("Tidak ada langkah perhitungan." if lang == "ID"
                else "No calculation steps available.")
        return

    # Group steps into collapsible sections by detecting separator steps
    # (steps whose symbol starts with "──" or whose description contains
    #  the big section headers)
    # We'll just stream them in a styled table with expandable groups

    # Build HTML table rows
    rows_html = []
    for i, s in enumerate(steps):
        # Determine row class
        ok_flag = ""
        if "✅" in s.substitution:
            ok_flag = "ok"
        elif "❌" in s.substitution:
            ok_flag = "ng"

        bg = ("#E8F5E9" if ok_flag == "ok" else
              "#FFEBEE" if ok_flag == "ng" else
              ("#F5F5F5" if i % 2 == 0 else "white"))

        # Escape for HTML
        def _e(x): return str(x).replace("<", "&lt;").replace(">", "&gt;")

        rows_html.append(
            f'<tr style="background:{bg};">'
            f'<td style="font-weight:600;color:#1A237E;padding:4px 8px;'
            f'white-space:nowrap;">{_e(s.symbol)}</td>'
            f'<td style="color:#455A64;padding:4px 8px;">{_e(s.description)}</td>'
            f'<td style="color:#1565C0;padding:4px 8px;font-family:monospace;">'
            f'{_e(f"{s.result:.4f}")} {_e(s.unit)}</td>'
            f'<td style="color:#6A1B9A;padding:4px 8px;font-size:11px;">'
            f'{_e(s.code_ref)}</td>'
            f'</tr>'
        )

    header = (
        '<tr style="background:#1A237E;color:white;font-weight:bold;">'
        '<td style="padding:6px 8px;">Simbol</td>'
        '<td style="padding:6px 8px;">Deskripsi</td>'
        '<td style="padding:6px 8px;">Hasil / Satuan</td>'
        '<td style="padding:6px 8px;">Ref. Code</td>'
        '</tr>'
    )

    table_html = (
        '<div style="overflow-x:auto;">'
        '<table style="width:100%;border-collapse:collapse;font-size:13px;">'
        + header
        + "".join(rows_html)
        + '</table></div>'
    )
    st.markdown(table_html, unsafe_allow_html=True)

    # ── LaTeX detail expander ─────────────────────────────────────────────────
    with st.expander("📐 " + ("Tampilkan formula LaTeX per langkah" if lang == "ID"
                              else "Show LaTeX formulas per step"), expanded=False):
        for s in steps:
            st.markdown(f"**{s.symbol}** — {s.description}")
            try:
                st.latex(s.formula)
            except Exception:
                st.code(s.formula)
            st.caption(
                f"Substitusi: `{s.substitution}` → "
                f"**{s.result:.4f} {s.unit}**  ·  _{s.code_ref}_"
            )
            st.divider()


# =============================================================================
# 8. EXPORT — WORD (.docx)
# =============================================================================

def export_docx(
    dr: DesignResult,
    fr: ForceResult,
    inp: dict,
    lang: str,
    fig_section: plt.Figure = None,
    fig_force:   go.Figure  = None,
    fig_stress:  plt.Figure = None,
    fig_pm:      go.Figure  = None,
) -> bytes:
    """
    Generate a Word report using python-docx.
    Returns bytes (for st.download_button).
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page setup (A4) ───────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Cm(2))

    # ── Styles ────────────────────────────────────────────────────────────────
    def _heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.runs[0].font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        return p

    def _para(text, bold=False, italic=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def _add_figure(fig_mpl=None, fig_plotly=None, width_cm=15):
        """Save figure to temp PNG and insert into doc."""
        buf = io.BytesIO()
        if fig_mpl is not None:
            fig_mpl.savefig(buf, format="png", dpi=_DPI,
                            bbox_inches="tight", facecolor=C["bg"])
        elif fig_plotly is not None:
            try:
                import plotly.io as pio
                buf_bytes = pio.to_image(fig_plotly, format="png",
                                         width=900, height=500, scale=1.5)
                buf = io.BytesIO(buf_bytes)
            except Exception:
                return   # kaleido not installed → skip figure
        buf.seek(0)
        doc.add_picture(buf, width=Cm(width_cm))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _table_row(table, cells_data: list, header=False):
        row = table.add_row()
        for i, (cell_txt, cell_width) in enumerate(cells_data):
            cell = row.cells[i]
            cell.text = str(cell_txt)
            if header:
                cell.paragraphs[0].runs[0].bold = True
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1A237E")
                shading.set(qn("w:color"), "FFFFFF")
                cell._tc.get_or_add_tcPr().append(shading)
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(
        "LAPORAN PERHITUNGAN STRUKTUR\n" if lang == "ID"
        else "STRUCTURAL CALCULATION REPORT\n"
    )
    run.bold = True; run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "Analisis Struktur U-Ditch & Cover Precast\n"
        + f"Kondisi: {fr.condition}\n"
        + f"Tanggal: {datetime.now().strftime('%d %B %Y')}"
    )
    sub_run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. INPUT SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    _heading("1. Data Input", 1)

    groups = {
        "Dimensi / Dimensions": [
            ("Lebar dalam UD — Wo",  f"{inp.get('ud_inner_width', '-')} mm"),
            ("Tinggi dalam UD — Ho", f"{inp.get('ud_inner_height', '-')} mm"),
            ("Tebal dinding atas — ta", f"{inp.get('ud_wall_thick_top', inp.get('ud_wall_thickness','-'))} mm"),
            ("Tebal dinding bawah — tb", f"{inp.get('ud_wall_thick_bot', '-')} mm"),
            ("Tebal slab dasar — ts", f"{inp.get('ud_slab_thick', inp.get('ud_base_thickness','-'))} mm"),
            ("Panjang segmen UD — L", f"{inp.get('ud_length', '-')} m"),
            ("Gap CU-UD", f"{inp.get('cu_gap', inp.get('gap_cu_ud','-'))} mm"),
            ("Tebal tengah CU — tcu", f"{inp.get('cu_thick_centre', inp.get('cu_thickness','-'))} mm"),
        ],
        "Material": [
            ("f'c (Beton / Concrete)", f"{inp.get('fc_prime','-')} MPa"),
            ("fy (Tulangan utama)", f"{inp.get('fy_main','-')} MPa"),
            ("fyt (Sengkang)", f"{inp.get('fy_shear',240)} MPa"),
            ("γc (Berat beton)", f"{inp.get('gamma_c','-')} kN/m³"),
        ],
        "Tanah / Soil": [
            ("γs", f"{inp.get('gamma_s','-')} kN/m³"),
            ("φ (Sudut gesek)", f"{inp.get('phi_soil','-')}°"),
            ("c (Kohesi)", f"{inp.get('cohesion',0)} kPa"),
            ("Tinggi timbunan samping Hf", f"{inp.get('soil_fill_beside',0)} m"),
        ],
        "Beban Kendaraan / Vehicle Load": [
            ("Beban gandar G", f"{inp.get('axle_load_G','-')} kN"),
            ("Beban satu roda P1 = G/2", f"{inp.get('axle_load_G',225)/2:.1f} kN"),
            ("Jarak roda x1", f"{inp.get('wheel_dist_x1','-')} m"),
            ("Jarak antar roda x2", f"{inp.get('wheel_spacing_x2',1.75)} m"),
            ("γDL / γLL", f"{inp.get('gamma_DL',1.2)} / {inp.get('gamma_LL',1.6)}"),
        ],
    }

    for group_name, rows in groups.items():
        _heading(group_name, 2)
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(9)
        tbl.columns[1].width = Cm(6)
        _table_row(tbl, [("Parameter", 9), ("Nilai / Value", 6)], header=True)
        for k, v in rows:
            r = tbl.add_row()
            r.cells[0].text = k
            r.cells[1].text = v
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. CROSS-SECTION FIGURE
    # ══════════════════════════════════════════════════════════════════════════
    _heading("2. Penampang Melintang / Cross-Section", 1)
    if fig_section:
        _add_figure(fig_mpl=fig_section, width_cm=14)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. FORCE DIAGRAMS
    # ══════════════════════════════════════════════════════════════════════════
    _heading("3. Diagram Gaya Dalam / Internal Forces", 1)
    wb = fr.wall_base
    if wb:
        tbl2 = doc.add_table(rows=1, cols=4)
        tbl2.style = "Table Grid"
        for col in tbl2.columns:
            col.width = Cm(3.7)
        _table_row(tbl2,
                   [("Penampang", 3.7), ("Mu (kN·m/m)", 3.7),
                    ("Vu (kN/m)", 3.7), ("Nu (kN/m)", 3.7)], header=True)
        r = tbl2.add_row()
        r.cells[0].text = "Dasar Dinding / Wall Base"
        r.cells[1].text = f"{wb.Mu:.3f}"
        r.cells[2].text = f"{wb.Vu:.3f}"
        r.cells[3].text = f"{wb.Nu:.3f}"
        if fr.cu_midspan:
            r2 = tbl2.add_row()
            r2.cells[0].text = "Tengah CU / CU Midspan"
            r2.cells[1].text = f"{fr.cu_midspan.Mu:.3f}"
            r2.cells[2].text = f"{fr.cu_midspan.Vu:.3f}"
            r2.cells[3].text = "—"
    if fig_force:
        _add_figure(fig_plotly=fig_force, width_cm=14)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. STEP-BY-STEP CALCULATIONS
    # ══════════════════════════════════════════════════════════════════════════
    _heading("4. Langkah Perhitungan / Calculation Steps", 1)

    steps_all = dr.all_steps
    tbl3 = doc.add_table(rows=1, cols=4)
    tbl3.style = "Table Grid"
    tbl3.columns[0].width = Cm(2)
    tbl3.columns[1].width = Cm(5)
    tbl3.columns[2].width = Cm(5)
    tbl3.columns[3].width = Cm(3)
    _table_row(tbl3,
               [("Simbol", 2), ("Deskripsi", 5), ("Hasil / Satuan", 5),
                ("Ref.", 3)], header=True)

    for s in steps_all:
        r = tbl3.add_row()
        r.cells[0].text = s.symbol
        r.cells[1].text = s.description
        r.cells[2].text = f"{s.result:.4f}  {s.unit}"
        r.cells[3].text = s.code_ref
        r.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. STRESS BLOCK FIGURE
    # ══════════════════════════════════════════════════════════════════════════
    _heading("5. Blok Tegangan / Stress Block", 1)
    if fig_stress:
        _add_figure(fig_mpl=fig_stress, width_cm=16)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. CAPACITY SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    _heading("6. Rangkuman Kapasitas / Capacity Summary", 1)

    cap_rows = []
    cap_checks = [
        (dr.wall_base_cap,  "Wall Base — Flexure"),
        (dr.wall_base_cap,  "Wall Base — Shear"),
        (dr.wall_base_dbl,  "Wall Base — 2-Rebar"),
        (dr.cu_midspan_cap, "CU Midspan"),
        (dr.cu_support_cap, "CU Support Shear"),
        (dr.base_slab_cap,  "Base Slab — Flexure"),
        (dr.base_slab_cap,  "Base Slab — Shear"),
    ]
    tbl4 = doc.add_table(rows=1, cols=5)
    tbl4.style = "Table Grid"
    for col in tbl4.columns:
        col.width = Cm(3)
    _table_row(tbl4,
               [("Penampang", 3), ("Demand", 3), ("φ Kapasitas", 3),
                ("Rasio", 3), ("Status", 3)], header=True)

    for cap, lbl in cap_checks:
        if cap is None:
            continue
        if "Shear" in lbl:
            dem = cap.Vu_kNm
            phi_c = cap.phi_Vn_kNm
            ok = cap.shear_ok
        else:
            dem = cap.Mu_kNmm
            phi_c = cap.phi_Mn_kNm
            ok = cap.flexure_ok
        if dem == 0 and phi_c == 0:
            continue
        ratio = phi_c / dem if dem > 0 else float("inf")
        r = tbl4.add_row()
        r.cells[0].text = lbl
        r.cells[1].text = f"{dem:.3f}"
        r.cells[2].text = f"{phi_c:.3f}"
        r.cells[3].text = f"{ratio:.3f}"
        r.cells[4].text = "✓ OK" if ok else "✗ NG"

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. P-M DIAGRAM (Condition 2)
    # ══════════════════════════════════════════════════════════════════════════
    if dr.pm_curve:
        _heading("7. Diagram Interaksi P-M / P-M Interaction Diagram", 1)
        pm = dr.pm_curve
        _para(f"Pn,max = {pm.Pn_max:.1f} kN/m  |  "
              f"Pb = {pm.Pb:.1f} kN/m  |  Mb = {pm.Mb:.2f} kN·m/m  |  "
              f"Mn,pure = {pm.Mn_pure:.2f} kN·m/m", bold=False)
        _para(f"Demand: Nu = {pm.Nu_demand:.2f} kN/m, "
              f"Mu = {pm.Mu_demand:.2f} kN·m/m → "
              + ("✓ Dalam kurva (AMAN)" if pm.inside_curve
                 else "✗ Di luar kurva (TIDAK AMAN)"), bold=True)
        if fig_pm:
            _add_figure(fig_plotly=fig_pm, width_cm=14)

    # ══════════════════════════════════════════════════════════════════════════
    # 8. FOOTER NOTE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    _para(
        "Catatan: Laporan ini dihasilkan secara otomatis. Selalu verifikasi dengan "
        "engineer yang bertanggung jawab.\n"
        "Note: This report is auto-generated. Always verify with a responsible engineer.",
        italic=True, size=9
    )
    _para(f"Kode Acuan / Code References: SNI 2847:2019 | SNI 1727:2020 | "
          f"AASHTO LRFD 9th Ed.", italic=True, size=9)
    _para(f"Dibuat oleh: U-Ditch Analysis App | {datetime.now().strftime('%Y-%m-%d %H:%M')}",
          italic=True, size=9)

    # ── Serialize to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# 9. EXPORT — PDF  (fpdf2)
# =============================================================================

def export_pdf(
    dr: DesignResult,
    fr: ForceResult,
    inp: dict,
    lang: str,
    fig_section: plt.Figure = None,
    fig_stress:  plt.Figure = None,
) -> bytes:
    """
    Generate a structured PDF report using fpdf2.
    Returns bytes for st.download_button.
    """
    from fpdf import FPDF, XPos, YPos

    class UDitchPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(26, 35, 126)
            self.cell(0, 6,
                      "U-Ditch & Cover Precast - Structural Analysis Report",
                      new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
            self.set_draw_color(26, 35, 126)
            self.line(self.l_margin, self.get_y(),
                      self.w - self.r_margin, self.get_y())
            self.ln(2)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(100, 100, 100)
            self.cell(0, 5,
                      f"SNI 2847:2019 | AASHTO LRFD 9th Ed.  |  "
                      f"Hal {self.page_no()}/{{nb}}  |  "
                      f"{datetime.now().strftime('%Y-%m-%d')}",
                      align="C")

    def _safe(text: str, maxlen: int = 200) -> str:
        """Sanitize text for fpdf2 Latin-1 font."""
        table = str.maketrans({
            "—": "-", "–": "-", "−": "-",
            "·": ".", "×": "x", "’": "'",
            "‘": "'", "“": '"', "”": '"',
            "≥": ">=", "≤": "<=", "°": " deg",
            "γ": "g", "φ": "phi", "α": "alpha",
            "β": "beta", "δ": "delta", "σ": "sigma",
            "ε": "eps", "π": "pi", "√": "sqrt",
            "•": "*", "→": "->", "←": "<-",
            "²": "2", "³": "3", "¹": "1",
            "✓": "OK", "✗": "NG",
            "✅": "[OK]", "❌": "[NG]",
            "é": "e", "è": "e", "ê": "e",
        })
        s = text.translate(table)
        # Strip any remaining non-latin-1 characters
        s = s.encode("latin-1", errors="replace").decode("latin-1")
        return s[:maxlen]

    pdf = UDitchPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_margins(20, 20, 20)

    def h1(text):
        pdf.set_font("Helvetica", "B", 13)
        pdf.set_text_color(26, 35, 126)
        pdf.set_fill_color(232, 240, 254)
        pdf.cell(0, 9, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT,
                 align="L", fill=True)
        pdf.ln(1)
        pdf.set_text_color(0, 0, 0)

    def h2(text):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(69, 90, 100)
        pdf.cell(0, 7, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="L")
        pdf.set_text_color(0, 0, 0)

    def body(text, bold=False, size=9, color=(0, 0, 0)):
        pdf.set_font("Helvetica", "B" if bold else "", size)
        pdf.set_text_color(*color)
        pdf.multi_cell(0, 5, text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    def divider():
        pdf.set_draw_color(200, 200, 200)
        pdf.line(pdf.l_margin, pdf.get_y(),
                 pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(2)

    def add_fig(fig_obj: plt.Figure, w_mm: float = 160):
        buf = io.BytesIO()
        fig_obj.savefig(buf, format="png", dpi=_DPI,
                        bbox_inches="tight", facecolor=C["bg"])
        buf.seek(0)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            tf.write(buf.read())
            tf_path = tf.name
        x_centre = (pdf.w - w_mm) / 2
        pdf.image(tf_path, x=x_centre, w=w_mm)
        os.unlink(tf_path)
        pdf.ln(3)

    def tbl_row(cells, widths, header=False):
        pdf.set_font("Helvetica", "B" if header else "", 8)
        if header:
            pdf.set_fill_color(26, 35, 126)
            pdf.set_text_color(255, 255, 255)
        else:
            pdf.set_fill_color(245, 245, 245)
            pdf.set_text_color(0, 0, 0)
        for txt, w in zip(cells, widths):
            pdf.cell(w, 6, _safe(str(txt)), border=1, fill=header,
                     new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)

    # ── Cover page ────────────────────────────────────────────────────────────
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(26, 35, 126)
    pdf.ln(20)
    pdf.cell(0, 12, "LAPORAN PERHITUNGAN STRUKTUR",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "U-Ditch & Cover Precast",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 10)
    pdf.ln(6)
    pdf.cell(0, 7, f"Kondisi Analisis: {fr.condition}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.cell(0, 7, f"Tanggal: {datetime.now().strftime('%d %B %Y')}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.add_page()

    # ── 1. Input Summary ─────────────────────────────────────────────────────
    h1("1. Data Input")
    rows_inp = [
        ("Wo (Lebar dalam UD)",   f"{inp.get('ud_inner_width','-')} mm"),
        ("Ho (Tinggi dalam UD)",  f"{inp.get('ud_inner_height','-')} mm"),
        ("ta (Tebal dinding atas)",f"{inp.get('ud_wall_thick_top',inp.get('ud_wall_thickness','-'))} mm"),
        ("tb (Tebal dinding bawah)",f"{inp.get('ud_wall_thick_bot','-')} mm"),
        ("ts (Tebal slab dasar)", f"{inp.get('ud_slab_thick',inp.get('ud_base_thickness','-'))} mm"),
        ("L (Panjang segmen UD)", f"{inp.get('ud_length','-')} m"),
        ("gap CU-UD",             f"{inp.get('cu_gap',inp.get('gap_cu_ud','-'))} mm"),
        ("tcu (Tebal tengah CU)", f"{inp.get('cu_thick_centre',inp.get('cu_thickness','-'))} mm"),
        ("f'c",                   f"{inp.get('fc_prime','-')} MPa"),
        ("fy",                    f"{inp.get('fy_main','-')} MPa"),
        ("gs (Tanah)",            f"{inp.get('gamma_s','-')} kN/m3"),
        ("phi (Sudut gesek)",     f"{inp.get('phi_soil','-')} deg"),
        ("G (Beban gandar)",      f"{inp.get('axle_load_G','-')} kN"),
        ("P1 = G/2 (Satu roda)", f"{inp.get('axle_load_G',225)/2:.1f} kN"),
        ("x1 (Jarak roda-UD)",   f"{inp.get('wheel_dist_x1','-')} m"),
        ("gDL / gLL",            f"{inp.get('gamma_DL',1.2)} / {inp.get('gamma_LL',1.6)}"),
    ]
    tbl_row(["Parameter", "Nilai"], [110, 60], header=True)
    for k, v in rows_inp:
        tbl_row([k, v], [110, 60])
    pdf.ln(4)

    # ── 2. Cross-section figure ───────────────────────────────────────────────
    h1("2. Penampang Melintang")
    if fig_section:
        add_fig(fig_section, w_mm=140)

    # ── 3. Internal Forces ───────────────────────────────────────────────────
    h1("3. Gaya Dalam / Internal Forces")
    wb = fr.wall_base
    if wb:
        tbl_row(["Penampang", "Mu (kN.m/m)", "Vu (kN/m)", "Nu (kN/m)"],
                [65, 40, 40, 40], header=True)
        tbl_row(["Wall Base", f"{wb.Mu:.3f}", f"{wb.Vu:.3f}", f"{wb.Nu:.3f}"],
                [65, 40, 40, 40])
        if fr.cu_midspan:
            tbl_row(["CU Midspan",
                     f"{fr.cu_midspan.Mu:.3f}", f"{fr.cu_midspan.Vu:.3f}", "-"],
                    [65, 40, 40, 40])
    pdf.ln(4)
    divider()

    # ── 4. Calculation Steps ─────────────────────────────────────────────────
    h1("4. Langkah Perhitungan / Calculation Steps")
    tbl_row(["Simbol", "Deskripsi", "Hasil", "Satuan", "Ref."],
            [28, 65, 28, 18, 46], header=True)
    for i, s in enumerate(dr.all_steps):
        fill = (i % 2 == 0)
        pdf.set_font("Helvetica", "B", 7.5)
        pdf.set_fill_color(248, 248, 252 if fill else 255)
        pdf.cell(28, 5.5, _safe(str(s.symbol))[:20], border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.set_font("Helvetica", "", 7)
        desc_short = _safe(str(s.description))[:60]
        pdf.cell(65, 5.5, desc_short, border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(28, 5.5, f"{s.result:.3f}", border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(18, 5.5, _safe(str(s.unit))[:10], border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        ref_short = _safe(str(s.code_ref))[:30]
        pdf.cell(46, 5.5, ref_short, border=1, fill=fill,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # ── 5. Stress block ──────────────────────────────────────────────────────
    if fig_stress:
        pdf.add_page()
        h1("5. Blok Tegangan / Stress Block")
        add_fig(fig_stress, w_mm=160)

    # ── 6. Capacity Summary ──────────────────────────────────────────────────
    pdf.add_page()
    h1("6. Rangkuman Kapasitas / Capacity Summary")
    tbl_row(["Penampang", "Demand", "phi-Cap.", "Rasio", "Status"],
            [60, 30, 30, 25, 25], header=True)

    cap_checks = [
        (dr.wall_base_cap,  "Wall Base Flex.",  False),
        (dr.wall_base_cap,  "Wall Base Shear",  True),
        (dr.wall_base_dbl,  "Wall (2-rebar)",   False),
        (dr.cu_midspan_cap, "CU Midspan",       False),
        (dr.cu_support_cap, "CU Support Shear", True),
        (dr.base_slab_cap,  "Slab Flex.",        False),
        (dr.base_slab_cap,  "Slab Shear",        True),
    ]
    for i, (cap, lbl, is_shear) in enumerate(cap_checks):
        if cap is None:
            continue
        dem   = cap.Vu_kNm if is_shear else cap.Mu_kNmm
        phi_c = cap.phi_Vn_kNm if is_shear else cap.phi_Mn_kNm
        ok    = cap.shear_ok if is_shear else cap.flexure_ok
        if dem == 0 and phi_c == 0:
            continue
        ratio = phi_c / dem if dem > 0 else 99.9
        fill  = (i % 2 == 0)
        ok_c  = (0, 100, 0) if ok else (180, 0, 0)
        pdf.set_font("Helvetica", "", 8)
        pdf.set_fill_color(248, 248, 252 if fill else 255)
        pdf.cell(60, 5.5, lbl, border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(30, 5.5, f"{dem:.3f}", border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(30, 5.5, f"{phi_c:.3f}", border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(25, 5.5, f"{ratio:.3f}", border=1, fill=fill,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.set_text_color(*ok_c)
        pdf.set_font("Helvetica", "B", 8)
        pdf.cell(25, 5.5, "OK" if ok else "NG", border=1, fill=fill,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)

    # ── 7. P-M summary ───────────────────────────────────────────────────────
    if dr.pm_curve:
        pdf.ln(4)
        h1("7. Diagram Interaksi P-M")
        pm = dr.pm_curve
        body(f"Pn,max={pm.Pn_max:.1f}kN/m | Pb={pm.Pb:.1f}kN/m | "
             f"Mb={pm.Mb:.2f}kN.m/m | Mn,pure={pm.Mn_pure:.2f}kN.m/m")
        status = "AMAN - Demand dalam kurva" if pm.inside_curve else \
                 "TIDAK AMAN - Demand di luar kurva"
        color  = (0, 100, 0) if pm.inside_curve else (180, 0, 0)
        body(f"Demand: Nu={pm.Nu_demand:.2f}kN/m, Mu={pm.Mu_demand:.2f}kN.m/m",
             bold=False)
        body(status, bold=True, color=color)

    # ── Footer note ───────────────────────────────────────────────────────────
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 7.5)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 4,
                   "Catatan: Laporan ini dihasilkan otomatis oleh U-Ditch Analysis App. "
                   "Selalu verifikasi hasil dengan engineer yang bertanggung jawab. | "
                   "Ref: SNI 2847:2019, SNI 1727:2020, AASHTO LRFD 9th Ed.")

    return bytes(pdf.output())


# =============================================================================
# 10. MASTER RENDER FUNCTION
# =============================================================================



def _render_calc_steps_streamlit(steps: list, lang: str, key_prefix: str = "out") -> None:
    """
    Public alias: renders calculation steps in book style.
    Called from 11_UDitch_CU.py tab_calc expander.
    """
    _render_calc_steps_book(steps, lang)

def render_output(
    dr:   DesignResult,
    fr:   ForceResult,
    inp:  dict,
    lang: str = "ID",
) -> None:
    """
    Master output renderer called from the Output tab in 11_UDitch_CU.py.

    Usage:
        from uditch.ui_output import render_output
        render_output(st.session_state["design_result"],
                      st.session_state["calc_results"],
                      st.session_state["input_data"], lang)
    """
    st.markdown(f"## {t('out_header', lang)}")

    # ── Summary metrics (top) ────────────────────────────────────────────────
    _section_summary_metrics(dr, fr, lang)
    st.divider()

    # ── Tab structure ─────────────────────────────────────────────────────────
    tab_labels = [
        "📐 " + t("out_section_sketch", lang),
        "🧮 " + t("calc_header", lang),
        "📊 " + ("P-M Diagram" if fr.condition == "Kondisi 2" else t("out_capacity_chart", lang)),
        "📄 " + (t("out_export_word", lang)[2:] + " / PDF"),
    ]
    tabs = st.tabs(tab_labels)

    # ── Tab 0: Cross-section sketch ───────────────────────────────────────────
    with tabs[0]:
        col_left, col_right = st.columns(2)
        with col_left:
            st.markdown(f"**{t('out_section_sketch', lang)}**")
            fig_sec = _fig_cross_section(inp, dr, fr)
            st.pyplot(fig_sec, use_container_width=True)
            plt.close(fig_sec)
        with col_right:
            st.markdown(
                f"**{'Ringkasan Gaya Dalam' if lang=='ID' else 'Internal Forces Summary'}**"
            )
            fig_cap = _fig_capacity_bars(dr)
            st.plotly_chart(fig_cap, use_container_width=True, key="plotly_cap_main")

    # ── Tab 1: Book-style calculation (THE MAIN OUTPUT) ───────────────────────
    with tabs[1]:
        st.markdown(
            f"### {'📋 Rincian Perhitungan Lengkap' if lang=='ID' else '📋 Complete Calculation Detail'}"
        )
        st.caption(
            "Urutan: A.Data → B.Tekanan (+ diagram) → C.Gaya Dalam (+ diagram) → "
            "D.Kapasitas (+ blok tegangan) → E.Kontrol"
            if lang == "ID" else
            "Order: A.Data → B.Pressure (+ diagram) → C.Forces (+ diagram) → "
            "D.Capacity (+ stress block) → E.Checks"
        )
        _render_calc_steps_book(dr.all_steps, lang, fr=fr, inp=inp, dr=dr)

    # ── Tab 2: P-M or capacity chart ─────────────────────────────────────────
    with tabs[2]:
        if fr.condition == "Kondisi 2" and dr.pm_curve:
            st.markdown(f"**{t('out_pm_curve', lang)}**")
            fig_pm = _fig_pm_curve(dr.pm_curve)
            st.plotly_chart(fig_pm, use_container_width=True, key="plotly_pm_tab")
            pm = dr.pm_curve
            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Pn,max (kN/m)", f"{pm.Pn_max:.1f}")
            c2.metric("Pb (kN/m)",     f"{pm.Pb:.1f}")
            c3.metric("Mb (kN·m/m)",   f"{pm.Mb:.2f}")
            c4.metric("Mn,pure",       f"{pm.Mn_pure:.2f}")
            inside_txt = ("✅ Aman — Titik beban dalam kurva" if pm.inside_curve
                          else "❌ Tidak Aman — Di luar kurva")
            st.markdown(
                f"**P-M Check:** Nu = {pm.Nu_demand:.2f} kN/m | "
                f"Mu = {pm.Mu_demand:.2f} kN·m/m → **{inside_txt}**"
            )
            plt.close("all")
        else:
            fig_cap2 = _fig_capacity_bars(dr)
            st.plotly_chart(fig_cap2, use_container_width=True, key="plotly_cap_tab2")

    # ── Tab 3: Export ─────────────────────────────────────────────────────────
    with tabs[3]:
        st.markdown(
            f"### {'📥 Ekspor Laporan Perhitungan' if lang=='ID' else '📥 Export Calculation Report'}"
        )
        st.caption(
            "Laporan berisi: Cover → Input → Diagram Tekanan → Gaya Dalam → "
            "Semua Langkah Perhitungan → Kapasitas → Kesimpulan."
            if lang == "ID" else
            "Report contains: Cover → Input → Pressure Diagram → Forces → "
            "All Calculation Steps → Capacity → Conclusions."
        )

        fig_sec_exp   = _fig_cross_section(inp, dr, fr)
        cap_for_exp   = dr.wall_base_cap or dr.wall_base_dbl
        fig_sb_exp    = _fig_stress_block(cap_for_exp, inp) if cap_for_exp else None
        fig_pres_exp  = _fig_pressure_full(fr, inp, lang)
        fig_force_exp = _fig_mvu_full(fr, inp, lang)
        fig_pm_exp    = _fig_pm_curve(dr.pm_curve) if dr.pm_curve else None

        col1, col2 = st.columns(2)
        with col1:
            st.markdown(f"#### {t('out_export_word', lang)}")
            try:
                docx_bytes = export_docx(
                    dr, fr, inp, lang,
                    fig_section=fig_sec_exp,
                    fig_force=fig_force_exp,
                    fig_stress=fig_sb_exp,
                    fig_pm=fig_pm_exp,
                )
                fname = (f"UD_Perhitungan_{fr.condition.replace(' ','_')}_"
                         f"{datetime.now().strftime('%Y%m%d')}.docx")
                st.download_button(
                    label=t("out_export_word", lang),
                    data=docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary",
                    use_container_width=True,
                    key="dl_docx",
                )
            except Exception as e:
                st.error(f"Word export error: {e}")

        with col2:
            st.markdown(f"#### {t('out_export_pdf', lang)}")
            try:
                pdf_bytes = export_pdf(
                    dr, fr, inp, lang,
                    fig_section=fig_sec_exp,
                    fig_stress=fig_sb_exp,
                )
                fname_pdf = (f"UD_Perhitungan_{fr.condition.replace(' ','_')}_"
                             f"{datetime.now().strftime('%Y%m%d')}.pdf")
                st.download_button(
                    label=t("out_export_pdf", lang),
                    data=pdf_bytes,
                    file_name=fname_pdf,
                    mime="application/pdf",
                    type="primary",
                    use_container_width=True,
                    key="dl_pdf",
                )
            except Exception as e:
                st.error(f"PDF export error: {e}")

        for _f in [fig_sec_exp, fig_sb_exp, fig_pres_exp, fig_force_exp]:
            if _f:
                try:
                    plt.close(_f)
                except Exception:
                    pass


# =============================================================================
# =============================================================================
#  TAHAP C — COMPLETE LOAD-TO-FORCE DERIVATION DISPLAY
#  "Dari Mana Gaya Dalam Berasal"
# =============================================================================
# =============================================================================
#
# This module explicitly shows the FULL CHAIN:
#   Geometri → Tekanan (σh per kedalaman) → Resultan (F, arm) → Mu, Vu, Nu
#
# Rendered as:
#   1. _render_load_derivation()   — show HOW pressures are built
#   2. _render_force_derivation()  — show HOW Mu/Vu/Nu come from pressures
#   3. _fig_pressure_full()        — engineering pressure diagram (2 panels)
#   4. _fig_mvu_full()             — Mu & Vu along wall height (2 panels)
#
# Called from _render_calc_steps_book() between Section B and Section C.
# Also available standalone for the Output tab.
# =============================================================================

import math as _math


def _render_load_derivation(fr: "ForceResult", inp: dict, lang: str) -> None:
    """
    Show step-by-step HOW the lateral pressure at each layer comes from
    the input data. Book-style: each line is Formula → Substitution → Result.
    """
    from uditch.calc_engine import _compat, _rankine_Ka, _aashto_heq
    inp_c  = _compat(inp)
    lat    = fr.lateral
    if lat is None:
        return

    H      = inp_c.get("ud_inner_height", 600) / 1000.0
    Ka     = lat.Ka
    gs     = inp_c.get("gamma_s", 18.0)
    c_soil = inp_c.get("cohesion", 0.0)
    method = inp_c.get("lat_method_idx", 0)
    P1     = inp_c.get("wheel_load", 112.5)
    x1     = inp_c.get("wheel_dist", 0.25)
    L_seg  = inp_c.get("ud_length", 1.2)
    Hf     = inp_c.get("soil_fill_beside", 0.0)
    gDL    = inp_c.get("gamma_DL", 1.2)
    gLL    = inp_c.get("gamma_LL", 1.6)

    def _box(title_id: str, title_en: str, content_fn):
        title = title_id if lang == "ID" else title_en
        st.markdown(
            f'<div style="border:1px solid #1A237E;border-radius:6px;'
            f'margin:10px 0;overflow:hidden;">'
            f'<div style="background:#1A237E;color:white;padding:6px 12px;'
            f'font-size:13px;font-weight:700;">{title}</div>'
            f'<div style="padding:10px 14px;">',
            unsafe_allow_html=True,
        )
        content_fn()
        st.markdown('</div></div>', unsafe_allow_html=True)

    def _step(symbol, formula_latex, sub_str, result_val, unit, ref):
        """Render one derivation step."""
        col1, col2 = st.columns([2, 3])
        with col1:
            st.markdown(
                f'<span style="font-weight:700;color:#1A237E;">{symbol}</span>',
                unsafe_allow_html=True,
            )
            try:
                st.latex(formula_latex)
            except Exception:
                st.code(formula_latex)
        with col2:
            result_fmt = f"{result_val:.4f}".rstrip("0").rstrip(".")
            st.markdown(
                f'<div style="background:#EDE7F6;border-radius:4px;padding:6px 10px;margin-top:8px;">'
                f'<code style="font-size:12px;">{sub_str}</code><br>'
                f'<span style="color:#1565C0;font-size:14px;font-weight:700;">= {result_fmt} {unit}</span>'
                f'<br><span style="color:#9C27B0;font-size:10px;font-style:italic;">{ref}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

    st.markdown(
        f'<div style="background:#E8EAF6;border-left:4px solid #1A237E;'
        f'padding:10px 16px;margin:12px 0;border-radius:4px;">'
        f'<b>{"📐 Derivasi Pembebanan Lateral — Dari Data ke Tekanan" if lang=="ID" else "📐 Lateral Load Derivation — From Data to Pressure"}</b><br>'
        f'<span style="font-size:12px;color:#455A64;">'
        f'{"Setiap baris menunjukkan: Rumus → Substitusi Angka → Hasil" if lang=="ID" else "Each line shows: Formula → Numeric Substitution → Result"}'
        f'</span></div>',
        unsafe_allow_html=True,
    )

    # ── Block 1: Rankine Ka ───────────────────────────────────────────────────
    st.markdown(f"**{'1️⃣  Koefisien Tekanan Tanah Aktif (Rankine)' if lang=='ID' else '1️⃣  Active Earth Pressure Coefficient (Rankine)'}**")
    phi_deg = inp_c.get("phi_soil", 30.0)
    _step("Ka",
          r"K_a = \tan^2\!\left(45° - \frac{\phi}{2}\right)",
          f"tan²(45° − {phi_deg}/2) = tan²({45-phi_deg/2}°)",
          Ka, "—", "Rankine (1857); SNI 8460:2017 §C.6.4")

    st.divider()

    # ── Block 2: Earth pressure profile ──────────────────────────────────────
    st.markdown(f"**{'2️⃣  Distribusi Tekanan Tanah Lateral' if lang=='ID' else '2️⃣  Lateral Earth Pressure Distribution'}**")
    st.caption(
        "σh(z) = Ka·γs·z − 2c·√Ka  (minimal 0)  |  "
        "z = kedalaman dari puncak dinding (m)"
        if lang == "ID" else
        "σh(z) = Ka·γs·z − 2c·√Ka  (minimum 0)  |  "
        "z = depth from wall top (m)"
    )
    sigma_top  = max(Ka * gs * 0    - 2*c_soil*_math.sqrt(Ka), 0.0)
    sigma_half = max(Ka * gs * H/2  - 2*c_soil*_math.sqrt(Ka), 0.0)
    sigma_base = max(Ka * gs * H    - 2*c_soil*_math.sqrt(Ka), 0.0)

    for z_val, sig_val, lbl in [
        (0,    sigma_top,  f"z=0 (puncak)"),
        (H/2,  sigma_half, f"z={H/2:.2f}m (tengah)"),
        (H,    sigma_base, f"z={H:.2f}m (dasar)"),
    ]:
        _step(f"σh  [{lbl}]",
              r"\sigma_h = K_a \cdot \gamma_s \cdot z - 2c\sqrt{K_a}",
              f"{Ka:.3f}×{gs}×{z_val:.3f} − 2×{c_soil}×√{Ka:.3f}",
              sig_val, "kPa",
              "SNI 8460:2017 §6.4.1; AASHTO LRFD §3.11.5.1")

    F_e   = lat.F_earth
    arm_e = lat.arm_earth
    _step("F_earth",
          r"F_{earth} = \tfrac{1}{2}(\sigma_{top}+\sigma_{base})\cdot H",
          f"½×({sigma_top:.2f}+{sigma_base:.2f})×{H:.3f}",
          F_e, "kN/m", "AASHTO LRFD §3.11.5.1")
    _step("y_earth",
          r"y_{earth} = \frac{H}{3}\cdot\frac{\sigma_{base}+2\sigma_{top}}{\sigma_{base}+\sigma_{top}}",
          f"({H:.3f}/3)×({sigma_base:.2f}+2×{sigma_top:.2f})/({sigma_base:.2f}+{sigma_top:.2f})",
          arm_e, "m dari dasar / from base", "Statics — centroid of trapezoid")

    st.divider()

    # ── Block 3: Surcharge/Boussinesq ─────────────────────────────────────────
    F_sur   = lat.F_surcharge
    arm_sur = lat.arm_surcharge

    if method == 0:
        st.markdown(f"**{'3️⃣  Surcharge Ekivalen (AASHTO Table 3.11.6.4-2)' if lang=='ID' else '3️⃣  Equivalent Surcharge (AASHTO Table 3.11.6.4-2)'}**")
        heq = _aashto_heq(H)
        sig_sur = Ka * gs * heq
        # Also soil fill beside if any
        q_beside = inp_c.get("udl_beside", 0.0)
        sig_q    = Ka * q_beside
        _step("heq",
              r"h_{eq} = f(H_{wall})",
              f"f({H:.2f}m) from AASHTO Table 3.11.6.4-2",
              heq, "m",
              "AASHTO LRFD 9th Ed. Table 3.11.6.4-2")
        _step("σ_sur",
              r"\sigma_{sur} = K_a\cdot\gamma_s\cdot h_{eq} + K_a\cdot q_{beside}",
              f"{Ka:.3f}×{gs}×{heq:.2f} + {Ka:.3f}×{q_beside:.1f}",
              sig_sur + sig_q, "kPa (uniform over H)",
              "AASHTO LRFD §3.11.6.4")
        _step("F_sur",
              r"F_{sur} = \sigma_{sur}\cdot H",
              f"{sig_sur+sig_q:.3f}×{H:.3f}",
              F_sur, "kN/m",
              "AASHTO LRFD §3.11.6.4")
        _step("y_sur",
              r"y_{sur} = H/2 \text{ (uniform)}",
              f"{H:.3f}/2",
              arm_sur, "m dari dasar", "Statics — uniform pressure midpoint")

    else:  # Boussinesq
        st.markdown(f"**{'3️⃣  Tekanan Lateral Boussinesq (Beban Terpusat P1)' if lang=='ID' else '3️⃣  Boussinesq Lateral Pressure (Point Load P1)'}**")
        st.caption(
            "Distribusi beban roda P1 sebagai beban terpusat 3D, diintegrasikan numerik "
            "sepanjang L segmen. Metode ini lebih realistis untuk segmen pendek (≤2.4m)."
            if lang == "ID" else
            "Wheel load P1 as 3D point load, numerically integrated over segment length L. "
            "More realistic for short precast segments (≤2.4m)."
        )
        _step("P1 = G/2",
              r"P_1 = G/2",
              f"{P1*2:.0f}/2",
              P1, "kN",
              "SNI 1725:2016 Tabel 1 — beban gandar G1/G2")
        _step("x (jarak roda ke dinding)",
              r"x = x_1 \text{ (tepi roda ke muka luar UD)}",
              f"x = {x1:.3f} m",
              x1, "m", "Input pengguna / user input")

        # Show formula for peak Boussinesq
        z_peak = x1 / _math.sqrt(2) if x1 > 0 else H / 2
        sigma_bou_peak = (3 * P1 * x1**2 * z_peak**3) / (
            2 * _math.pi * (x1**2 + z_peak**2)**2.5 * L_seg
        )
        _step("σh,max (Boussinesq)",
              r"\sigma_{h,max} \approx \frac{3P_1 x^2 z^3}{2\pi R^5 \cdot L}",
              f"z_peak≈{z_peak:.2f}m, R=√({x1:.2f}²+{z_peak:.2f}²)",
              sigma_bou_peak, "kPa",
              "Boussinesq (1885); NAVFAC DM7.02 Fig.3")
        _step("F_Boussinesq",
              r"F = \int_0^H \sigma_h(z)\,dz \text{ (numerik)}",
              f"P1={P1:.0f}kN, x={x1:.2f}m, H={H:.2f}m, L={L_seg}m",
              F_sur, "kN/m",
              "Integrasi numerik Boussinesq; NAVFAC DM7.02")
        _step("y_Boussinesq",
              r"y = M_{base}/F \text{ (numerik)}",
              f"≈ {arm_sur:.3f} m dari dasar / from base",
              arm_sur, "m", "Numerical integration")

    st.divider()

    # ── Block 4: Total & Factored ─────────────────────────────────────────────
    st.markdown(f"**{'4️⃣  Total Gaya Lateral & Faktor Beban' if lang=='ID' else '4️⃣  Total Lateral Force & Load Factors'}**")
    F_total = lat.F_total
    _step("F_total",
          r"F_{total} = F_{earth} + F_{sur}",
          f"{F_e:.3f} + {F_sur:.3f}",
          F_total, "kN/m", "Superposisi / Superposition")

    Mu_earth_serv = F_e   * arm_e
    Mu_sur_serv   = F_sur * arm_sur
    Mu_u = gDL * Mu_earth_serv + gLL * Mu_sur_serv
    Vu_u = gDL * F_e + gLL * F_sur

    _step("Mu,u (terfaktor / factored)",
          r"M_u = \gamma_{DL}\cdot F_{earth}\cdot y_{earth} + \gamma_{LL}\cdot F_{sur}\cdot y_{sur}",
          f"{gDL}×{F_e:.3f}×{arm_e:.3f} + {gLL}×{F_sur:.3f}×{arm_sur:.3f}",
          Mu_u, "kN·m/m", "SNI 2847:2019 Ps.5.3.1 (U=1.2D+1.6L)")
    _step("Vu,u (terfaktor / factored)",
          r"V_u = \gamma_{DL}\cdot F_{earth} + \gamma_{LL}\cdot F_{sur}",
          f"{gDL}×{F_e:.3f} + {gLL}×{F_sur:.3f}",
          Vu_u, "kN/m", "SNI 2847:2019 Ps.5.3.1")


def _render_force_derivation(fr: "ForceResult", inp: dict, lang: str) -> None:
    """
    Show HOW Mu, Vu, Nu at the critical sections are obtained.
    Includes gap check for Kondisi 1 and column analysis for Kondisi 2.
    """
    from uditch.calc_engine import _compat
    inp_c = _compat(inp)
    lat   = fr.lateral
    wb    = fr.wall_base
    if wb is None or lat is None:
        return

    gDL = inp_c.get("gamma_DL", 1.2)
    gLL = inp_c.get("gamma_LL", 1.6)
    H   = inp_c.get("ud_inner_height", 600) / 1000.0
    ta  = inp_c.get("ud_wall_thickness", 80) / 1000.0
    gam_c = inp_c.get("gamma_c", 24.0)
    cond  = fr.condition

    def _result_line(label, value, unit, note=""):
        st.markdown(
            f'<div style="background:#E3F2FD;border-left:4px solid #1565C0;'
            f'padding:6px 12px;margin:4px 0;border-radius:3px;">'
            f'<b style="color:#1565C0;">{label}</b> = '
            f'<span style="font-size:16px;font-weight:700;color:#0D47A1;">{value:.3f}</span> '
            f'<span style="color:#455A64;">{unit}</span>'
            + (f'  <span style="color:#7B1FA2;font-size:11px;">&nbsp;({note})</span>' if note else "")
            + f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown(
        f'<div style="background:#E8F5E9;border-left:4px solid #2E7D32;'
        f'padding:10px 16px;margin:12px 0;border-radius:4px;">'
        f'<b>{"⚡ Derivasi Gaya Dalam — Dari Tekanan ke Mu, Vu, Nu" if lang=="ID" else "⚡ Internal Force Derivation — From Pressure to Mu, Vu, Nu"}</b>'
        f'</div>',
        unsafe_allow_html=True,
    )

    # ── Wall self-weight & axial ──────────────────────────────────────────────
    st.markdown(f"**{'🧱 Berat Dinding & Gaya Aksial Nu' if lang=='ID' else '🧱 Wall Self-weight & Axial Nu'}**")
    W_wall = gam_c * H * ta
    Nu_wall = gDL * W_wall
    try:
        st.latex(r"W_{wall} = \gamma_c \cdot H \cdot t_a")
    except Exception:
        pass
    st.markdown(
        f'`{gam_c} × {H:.3f} × {ta:.3f}` → **W_wall = {W_wall:.3f} kN/m**<br>'
        f'`Nu = γDL × W_wall = {gDL} × {W_wall:.3f}` → **Nu = {wb.Nu:.3f} kN/m**',
        unsafe_allow_html=True,
    )
    _result_line("Nu (dasar dinding)", wb.Nu, "kN/m",
                 "tekan positif / compression positive")

    st.divider()

    # ── Cantilever moments ────────────────────────────────────────────────────
    st.markdown(f"**{'📐 Momen & Geser: Dinding sebagai Kantilever' if lang=='ID' else '📐 Moment & Shear: Wall as Cantilever'}**")
    st.caption(
        "Momen di dasar = jumlah (gaya × lengan) dari semua tekanan lateral."
        if lang == "ID" else
        "Moment at base = sum of (force × arm) for all lateral pressures."
    )

    Mu_e = gDL * lat.F_earth * lat.arm_earth
    Mu_s = gLL * lat.F_surcharge * lat.arm_surcharge
    Mu_cant = Mu_e + Mu_s
    Vu_cant = gDL * lat.F_earth + gLL * lat.F_surcharge

    try:
        st.latex(r"M_u = \gamma_{DL}\cdot F_{earth}\cdot y_{earth} + \gamma_{LL}\cdot F_{sur}\cdot y_{sur}")
    except Exception:
        pass
    st.markdown(
        f'`{gDL}×{lat.F_earth:.3f}×{lat.arm_earth:.3f}` + `{gLL}×{lat.F_surcharge:.3f}×{lat.arm_surcharge:.3f}`',
        unsafe_allow_html=True,
    )
    st.markdown(
        f'= `{Mu_e:.3f} + {Mu_s:.3f}`',
        unsafe_allow_html=True,
    )

    # ── Gap check (Kondisi 1) ─────────────────────────────────────────────────
    if cond == "Kondisi 1":
        st.divider()
        st.markdown(f"**{'🔍 Cek Gap — Apakah CU Aktif Sebagai Strut?' if lang=='ID' else '🔍 Gap Check — Does CU Activate as Strut?'}**")

        fc    = inp_c.get("fc_prime", 30.0)
        Ec    = 4700 * _math.sqrt(fc) * 1e3      # kN/m²
        I_w   = ta**3 / 12                         # m⁴/m
        EI    = Ec * I_w
        w_base= max(lat.Ka * inp_c.get("gamma_s",18) * H
                    - 2*inp_c.get("cohesion",0)*_math.sqrt(lat.Ka), 0.0)
        w_sur = lat.F_surcharge / H if H > 0 else 0
        d_e   = w_base * H**4 / (30 * EI)
        d_s   = w_sur  * H**4 / (8  * EI)
        d_tot = (d_e + d_s) * 1000  # mm
        gap_mm = inp_c.get("gap_cu_ud", 20.0)

        try:
            st.latex(r"\delta_{total} = \frac{w_{base}\cdot H^4}{30EI} + \frac{w_{sur}\cdot H^4}{8EI}")
        except Exception:
            pass
        st.markdown(
            f'Ec = 4700√{fc} = **{fc and 4700*_math.sqrt(fc):.0f} MPa** | '
            f'I = ta³/12 = {ta:.3f}³/12 = **{I_w:.6f} m⁴/m** | '
            f'EI = **{EI:.1f} kN·m²/m**',
            unsafe_allow_html=True,
        )
        st.markdown(
            f'δ_earth = `{w_base:.3f}×{H:.3f}⁴/(30×{EI:.1f})` = **{d_e*1000:.3f} mm**<br>'
            f'δ_sur   = `{w_sur:.3f}×{H:.3f}⁴/(8×{EI:.1f})` = **{d_s*1000:.3f} mm**<br>'
            f'δ_total = **{d_tot:.3f} mm**  vs  gap = **{gap_mm} mm**',
            unsafe_allow_html=True,
        )

        if d_tot >= gap_mm:
            st.success(
                f"✅ δ = {d_tot:.2f}mm ≥ gap = {gap_mm}mm → **CU aktif sebagai strut (propped cantilever)**"
                if lang == "ID" else
                f"✅ δ = {d_tot:.2f}mm ≥ gap = {gap_mm}mm → **CU active as strut (propped cantilever)**"
            )
            st.caption(
                "Pola momen berubah: bukan lagi kantilever murni. "
                "Momen di dasar BERKURANG karena CU menahan di puncak."
                if lang == "ID" else
                "Moment pattern changes: no longer pure cantilever. "
                "Base moment REDUCED because CU restrains at top."
            )
        else:
            st.info(
                f"ℹ️ δ = {d_tot:.2f}mm < gap = {gap_mm}mm → **Kantilever murni (CU belum aktif)**"
                if lang == "ID" else
                f"ℹ️ δ = {d_tot:.2f}mm < gap = {gap_mm}mm → **Pure cantilever (CU not yet active)**"
            )
            st.caption(
                "Dinding bekerja sebagai kantilever murni. "
                "Tulangan luar (tarik) didesain dari kondisi ini."
                if lang == "ID" else
                "Wall acts as pure cantilever. "
                "Outer (tension) reinforcement designed from this condition."
            )

    # ── Column analysis (Kondisi 2) ───────────────────────────────────────────
    elif cond == "Kondisi 2":
        st.divider()
        st.markdown(f"**{'🏛️ Kondisi 2: Dinding sebagai Kolom (N + M)' if lang=='ID' else '🏛️ Condition 2: Wall as Column (N + M)'}**")
        st.caption(
            "Beban dari CU → reaksi vertikal R ke puncak dinding. "
            "Dinding menerima: N (aksial) + M (dari eksentrisitas + lateral)."
            if lang == "ID" else
            "Load from CU → vertical reaction R to wall top. "
            "Wall carries: N (axial) + M (from eccentricity + lateral)."
        )

    # ── Final results box ────────────────────────────────────────────────────
    st.divider()
    st.markdown(
        f'<div style="background:#1A237E;color:white;padding:10px 16px;'
        f'border-radius:6px;margin:10px 0;">'
        f'<b>{"📋 Gaya Dalam Desain di Dasar Dinding UD" if lang=="ID" else "📋 Design Internal Forces at UD Wall Base"}</b>'
        f'</div>',
        unsafe_allow_html=True,
    )
    _result_line("Mu (Momen)", wb.Mu, "kN·m/m",
                 "momen lentur terfaktor di dasar")
    _result_line("Vu (Geser)",  wb.Vu, "kN/m",
                 "gaya geser terfaktor di dasar")
    _result_line("Nu (Aksial)", wb.Nu, "kN/m",
                 "gaya tekan aksial terfaktor")


def _fig_pressure_full(fr: "ForceResult", inp: dict, lang: str) -> "plt.Figure":
    """
    Two-panel professional pressure diagram.
    Left:  σh vs depth with areas coloured by load type
    Right: Force resultants & moment arms diagram
    Shows NUMBERS at key depths.
    """
    from uditch.calc_engine import _compat, _rankine_Ka, _aashto_heq, _boussinesq_point_lateral
    inp_c = _compat(inp)
    lat   = fr.lateral
    if lat is None:
        fig, ax = plt.subplots(figsize=(8, 4), facecolor=C["bg"])
        ax.text(0.5, 0.5, "No data", ha="center")
        return fig

    H      = inp_c.get("ud_inner_height", 600) / 1000.0
    Ka     = lat.Ka
    gs     = inp_c.get("gamma_s", 18.0)
    c_s    = inp_c.get("cohesion", 0.0)
    method = inp_c.get("lat_method_idx", 0)
    P1     = inp_c.get("wheel_load", 112.5)
    x1     = max(inp_c.get("wheel_dist", 0.25), 0.05)
    L_seg  = inp_c.get("ud_length", 1.2)
    q_side = inp_c.get("udl_beside", 0.0)

    n   = 100
    z   = [i * H / n for i in range(n + 1)]   # 0 = top, H = base

    # Earth pressure
    sig_e = [max(Ka * gs * zi - 2*c_s*_math.sqrt(Ka), 0.0) for zi in z]

    # Surcharge or Boussinesq
    if method == 0:
        heq    = _aashto_heq(H)
        sig_sur = [Ka * gs * heq + Ka * q_side] * len(z)
        sig_bou = None
        sur_label = f"Surcharge (heq={heq:.2f}m)"
    else:
        sig_bou  = [_boussinesq_point_lateral(P1, x1, max(zi, 0.001), L_seg) / L_seg
                    for zi in z]
        sig_sur  = [Ka * q_side] * len(z) if q_side > 0 else None
        sur_label = f"Boussinesq P1={P1:.0f}kN"

    fig, (ax_L, ax_R) = plt.subplots(1, 2, figsize=(12, 6),
                                      facecolor=C["bg"],
                                      gridspec_kw={"width_ratios": [3, 2]})

    # ────── Left panel: pressure vs depth ──────────────────────────────────
    ax_L.set_facecolor(C["bg"])
    ax_L.invert_yaxis()

    # Earth (fill left-to-right from 0)
    ax_L.fill_betweenx(z, 0, sig_e, alpha=0.40, color="#7B1FA2",
                        label=f"Tanah aktif / Active earth  (Ka={Ka:.3f})")
    ax_L.plot(sig_e, z, color="#7B1FA2", lw=2.5)

    # Annotate key values
    for zi, si, lbl in [(0, sig_e[0], "z=0"), (H/2, sig_e[n//2], f"z={H/2:.1f}m"),
                         (H, sig_e[-1], f"z={H:.2f}m (base)")]:
        ax_L.annotate(f"{si:.1f}kPa",
                      xy=(si, zi), xytext=(si+1.5, zi),
                      fontsize=8, color="#7B1FA2", va="center")

    # Surcharge / Boussinesq
    if method == 0 and sig_sur:
        tot = [se+ss for se, ss in zip(sig_e, sig_sur)]
        ax_L.fill_betweenx(z, sig_e, tot, alpha=0.30, color="#F57F17",
                            label=sur_label)
        ax_L.plot(tot, z, color="#F57F17", lw=2, ls="--")
        ax_L.annotate(f"{tot[-1]:.1f}kPa",
                      xy=(tot[-1], z[-1]), xytext=(tot[-1]+1.5, z[-1]),
                      fontsize=8, color="#F57F17", va="center")
    elif method == 1 and sig_bou:
        tot = [se+sb for se, sb in zip(sig_e, sig_bou)]
        ax_L.fill_betweenx(z, sig_e, tot, alpha=0.30, color="#D32F2F",
                            label=sur_label)
        ax_L.plot(tot, z, color="#D32F2F", lw=2, ls="--")
        # Mark peak
        peak_i = sig_bou.index(max(sig_bou))
        ax_L.plot(tot[peak_i], z[peak_i], "o", color="#D32F2F", ms=7, zorder=5)
        ax_L.annotate(f"Peak {tot[peak_i]:.1f}kPa\nz={z[peak_i]:.2f}m",
                      xy=(tot[peak_i], z[peak_i]),
                      xytext=(tot[peak_i]+1.5, z[peak_i]-0.05*H),
                      fontsize=8, color="#D32F2F")

    ax_L.axhline(H, color=C["neutral"], lw=0.8, ls=":", alpha=0.5)
    ax_L.text(0.5, H+0.02*H, "Dasar / Base", fontsize=8, color=C["neutral"])
    ax_L.set_xlabel("σh (kPa)", fontsize=10, color=C["neutral"])
    ax_L.set_ylabel("Kedalaman z dari puncak (m)" if lang=="ID"
                     else "Depth z from top (m)", fontsize=10, color=C["neutral"])
    ax_L.set_title(
        f"Distribusi Tekanan Lateral — {fr.condition}" if lang=="ID"
        else f"Lateral Pressure Distribution — {fr.condition}",
        fontsize=11, color=C["neutral"])
    ax_L.legend(fontsize=9, loc="lower right")
    ax_L.grid(color=C["grid"], lw=0.5)
    ax_L.tick_params(labelsize=8)
    ax_L.set_xlim(left=0)

    # ────── Right panel: resultant diagram ──────────────────────────────────
    ax_R.set_facecolor(C["bg"])
    ax_R.set_aspect("equal")
    ax_R.axis("off")

    # Wall outline
    wall_w = 0.15  # normalised
    ax_R.add_patch(patches.Rectangle((0, 0), wall_w, 1,
                                      fc=C["concrete"], ec=C["neutral"], lw=2))
    ax_R.text(wall_w/2, 0.5, "UD\nWall", ha="center", va="center",
              fontsize=9, color="#37474F")

    # Earth resultant arrow
    Fe_norm  = lat.F_earth / max(lat.F_total, 1)
    ye_norm  = 1 - lat.arm_earth / H
    ax_R.annotate("", xy=(wall_w, ye_norm),
                  xytext=(wall_w + Fe_norm * 0.6, ye_norm),
                  arrowprops=dict(arrowstyle="->", color="#7B1FA2", lw=2.5))
    ax_R.text(wall_w + 0.05, ye_norm - 0.06,
              f"F_earth={lat.F_earth:.2f}kN/m\ny={lat.arm_earth:.2f}m",
              fontsize=8, color="#7B1FA2")

    # Surcharge resultant arrow
    Fs_norm  = lat.F_surcharge / max(lat.F_total, 1)
    ys_norm  = 1 - lat.arm_surcharge / H
    ax_R.annotate("", xy=(wall_w, ys_norm),
                  xytext=(wall_w + Fs_norm * 0.6, ys_norm),
                  arrowprops=dict(arrowstyle="->",
                                  color="#D32F2F" if method==1 else "#F57F17",
                                  lw=2.5))
    ax_R.text(wall_w + 0.05, ys_norm + 0.03,
              f"F_sur={lat.F_surcharge:.2f}kN/m\ny={lat.arm_surcharge:.2f}m",
              fontsize=8, color="#D32F2F" if method==1 else "#F57F17")

    # Mu at base
    ax_R.text(wall_w/2, -0.08,
              f"Mu={fr.wall_base.Mu:.2f}kN·m/m",
              ha="center", fontsize=9, color=C["moment"], fontweight="bold")
    ax_R.text(wall_w/2, -0.14,
              f"Vu={fr.wall_base.Vu:.2f}kN/m",
              ha="center", fontsize=9, color=C["shear"], fontweight="bold")

    ax_R.set_xlim(-0.05, 1.1)
    ax_R.set_ylim(-0.2, 1.15)
    ax_R.set_title("Resultan Gaya\n& Momen Dasar" if lang=="ID"
                    else "Force Resultants\n& Base Moment",
                    fontsize=10, color=C["neutral"])

    fig.tight_layout()
    return fig


def _fig_mvu_full(fr: "ForceResult", inp: dict, lang: str) -> "plt.Figure":
    """
    Three-panel diagram: Vu, Mu, Nu along wall height.
    All computed from distributed pressure, shown with intermediate values.
    """
    from uditch.calc_engine import _compat
    inp_c = _compat(inp)
    lat   = fr.lateral
    wb    = fr.wall_base
    if lat is None or wb is None:
        return plt.figure(figsize=(10, 5))

    H    = inp_c.get("ud_inner_height", 600) / 1000.0
    ta   = inp_c.get("ud_wall_thickness", 80) / 1000.0
    gam_c = inp_c.get("gamma_c", 24.0)
    gDL  = inp_c.get("gamma_DL", 1.2)
    gLL  = inp_c.get("gamma_LL", 1.6)
    Ka   = lat.Ka
    gs   = inp_c.get("gamma_s", 18.0)
    c_s  = inp_c.get("cohesion", 0.0)

    n    = 80
    # z measured FROM BASE (0=base, H=top)
    z    = [i * H / n for i in range(n + 1)]

    def p_e(z_from_base):    # earth pressure (max at z=0 base)
        depth = H - z_from_base
        return max(Ka * gs * depth - 2*c_s*_math.sqrt(Ka), 0.0)

    p_sur = lat.F_surcharge / H    # uniform equivalent

    def p_total_factored(z_from_base):
        return gDL * p_e(z_from_base) + gLL * p_sur

    # V(z) = shear at height z from base = ∫_z^H p(s)ds (cantilever)
    def V_at(z_fb):
        dz = (H - z_fb) / 40
        return sum(p_total_factored(z_fb + (j+0.5)*dz) * dz for j in range(40))

    # M(z) = moment at height z from base = ∫_z^H p(s)·(s-z)ds
    def M_at(z_fb):
        dz = (H - z_fb) / 40
        return sum(p_total_factored(z_fb + (j+0.5)*dz) * ((j+0.5)*dz) * dz
                   for j in range(40))

    V_arr = [V_at(zi) for zi in z]
    M_arr = [M_at(zi) for zi in z]

    # Propped correction
    V_prop, M_prop = None, None
    if fr.gap_closed and fr.wall_top is not None:
        H_strut = fr.wall_top.Vu
        V_prop  = [V_at(zi) - H_strut for zi in z]
        M_prop  = [M_at(zi) - H_strut*(H-zi) for zi in z]

    # Nu along height (just self-weight, linearly increasing toward base)
    W_unit = gam_c * ta   # kN/m² per m height
    Nu_arr = [gDL * W_unit * (H - zi) for zi in z]  # nu increases toward base

    # ── Figure ────────────────────────────────────────────────────────────────
    fig, axes = plt.subplots(1, 3, figsize=(13, 6), facecolor=C["bg"],
                             sharey=True)

    def setup_ax(ax, title, xlabel, col):
        ax.set_facecolor(C["bg"])
        ax.set_title(title, fontsize=10, color=C["neutral"], pad=6)
        ax.set_xlabel(xlabel, fontsize=9, color=col)
        ax.set_ylabel("z dari dasar (m)" if lang=="ID" else "z from base (m)",
                       fontsize=9, color=C["neutral"])
        ax.axhline(0,  color=C["neutral"], lw=0.8, ls=":", alpha=0.5)
        ax.axhline(H,  color=C["neutral"], lw=0.8, ls=":", alpha=0.5)
        ax.axvline(0,  color=C["neutral"], lw=1.5)
        ax.grid(color=C["grid"], lw=0.5)
        ax.tick_params(labelsize=7)
        ax.text(-0.02, 0,   "Dasar" if lang=="ID" else "Base",
                fontsize=7, color=C["neutral"], ha="right")
        ax.text(-0.02, H,   "Puncak" if lang=="ID" else "Top",
                fontsize=7, color=C["neutral"], ha="right")

    # ── Shear ─────────────────────────────────────────────────────────────────
    ax0 = axes[0]
    ax0.fill_betweenx(z, 0, V_arr, alpha=0.25, color=C["shear"])
    ax0.plot(V_arr, z, color=C["shear"], lw=2,
             label="Vu (kantilever)" if lang=="ID" else "Vu (cantilever)")
    if V_prop:
        ax0.plot(V_prop, z, color=C["shear"], lw=2, ls="--",
                 label="Vu (propped)")
        ax0.fill_betweenx(z, 0, V_prop, alpha=0.12, color=C["shear"])
    ax0.annotate(f"Vu,max\n={wb.Vu:.2f}kN/m",
                 xy=(V_arr[0], 0), xytext=(V_arr[0]*0.5+0.5, H*0.12),
                 fontsize=8, color=C["shear"],
                 arrowprops=dict(arrowstyle="->", color=C["shear"], lw=0.8))
    ax0.legend(fontsize=8)
    setup_ax(ax0,
             "Diagram Geser Vu" if lang=="ID" else "Shear Diagram Vu",
             "Vu (kN/m)", C["shear"])

    # ── Moment ────────────────────────────────────────────────────────────────
    ax1 = axes[1]
    ax1.fill_betweenx(z, 0, M_arr, alpha=0.25, color=C["moment"])
    ax1.plot(M_arr, z, color=C["moment"], lw=2,
             label="Mu (kantilever)" if lang=="ID" else "Mu (cantilever)")
    if M_prop:
        ax1.plot(M_prop, z, color=C["moment"], lw=2, ls="--",
                 label="Mu (propped)")
        ax1.fill_betweenx(z, 0, M_prop, alpha=0.12, color=C["moment"])
        # Mark max propped moment
        M_max_p  = max(M_prop, key=abs)
        z_Mmaxp  = z[M_prop.index(M_max_p)]
        ax1.plot(M_max_p, z_Mmaxp, "o", color=C["moment"], ms=7, zorder=5)
        ax1.annotate(f"Mu,max={M_max_p:.2f}",
                     xy=(M_max_p, z_Mmaxp),
                     xytext=(M_max_p*0.5+0.3, z_Mmaxp+H*0.08),
                     fontsize=8, color=C["moment"])
    ax1.annotate(f"Mu,base\n={wb.Mu:.2f}kN·m",
                 xy=(M_arr[0], 0), xytext=(M_arr[0]*0.4+0.1, H*0.12),
                 fontsize=8, color=C["moment"],
                 arrowprops=dict(arrowstyle="->", color=C["moment"], lw=0.8))
    ax1.legend(fontsize=8)
    setup_ax(ax1,
             "Diagram Momen Mu" if lang=="ID" else "Moment Diagram Mu",
             "Mu (kN·m/m)", C["moment"])

    # ── Axial ─────────────────────────────────────────────────────────────────
    ax2 = axes[2]
    ax2.fill_betweenx(z, 0, Nu_arr, alpha=0.25, color=C["axial"])
    ax2.plot(Nu_arr, z, color=C["axial"], lw=2, label="Nu (berat sendiri)")
    ax2.annotate(f"Nu,base\n={wb.Nu:.2f}kN/m",
                 xy=(Nu_arr[0], 0), xytext=(Nu_arr[0]*0.5+0.1, H*0.12),
                 fontsize=8, color=C["axial"],
                 arrowprops=dict(arrowstyle="->", color=C["axial"], lw=0.8))
    ax2.legend(fontsize=8)
    setup_ax(ax2,
             "Diagram Aksial Nu" if lang=="ID" else "Axial Diagram Nu",
             "Nu (kN/m) — tekan +", C["axial"])

    title_str = (f"Gaya Dalam Dinding UD — {fr.condition}\n"
                 f"(Distribusi sepanjang tinggi dinding, dihitung dari tekanan lateral)"
                 if lang == "ID" else
                 f"UD Wall Internal Forces — {fr.condition}\n"
                 f"(Distribution along wall height, computed from lateral pressure)")
    fig.suptitle(title_str, fontsize=10, color=C["neutral"])
    fig.tight_layout()
    return fig
