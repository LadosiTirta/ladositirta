# =============================================================================
# uditch/laporan.py
# =============================================================================
# Laporan perhitungan profesional — Word (.docx) + PDF
#
# RANTAI LENGKAP di laporan (identik dengan layar Streamlit):
#   BAGIAN 1  Data & Geometri
#   BAGIAN 2  Tekanan Lateral per Kedalaman
#             • Tabel Ka, σh tanah, σh kendaraan
#             • GAMBAR distribusi tekanan
#   BAGIAN 3  Resultan Gaya & Lengan Momen
#   BAGIAN 4  Gaya Dalam: Mu, Vu, Nu
#             • Setiap baris: Rumus | Substitusi | Hasil
#   BAGIAN 5  GAMBAR Mu / Vu / Nu sepanjang dinding
#   BAGIAN 6  Kapasitas Penampang & Kontrol
#
# API publik:
#   build_laporan_docx(fr, dr, inp, lang) -> bytes   (.docx)
#   build_laporan_pdf (fr, dr, inp, lang) -> bytes   (.pdf)
#
# Keduanya bisa langsung dipakai di st.download_button()
# =============================================================================

from __future__ import annotations
import io, math, os, tempfile
from datetime import datetime
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# Import figure builders dari perhitungan.py
from uditch.perhitungan import _fig_distribusi, _fig_gaya_dalam
from uditch.calc_engine  import (
    _compat, _rankine_Ka, _aashto_heq, _boussinesq_point_lateral
)

def _fig_pm_curve_laporan(pm) -> plt.Figure:
    """Build P-M interaction diagram for the report (standalone, no plotly)."""
    # Nominal curve
    Mn_nom = [p.Mn for p in pm.points]
    Pn_nom = [p.Pn for p in pm.points]
    # Design (phi-reduced) curve
    Mn_des = [p.phi_Mn for p in pm.points]
    Pn_des = [p.phi_Pn for p in pm.points]

    fig, ax = plt.subplots(figsize=(9, 6.5), facecolor="#F8F9FA")
    ax.set_facecolor("#F8F9FA")

    # Nominal curve (grey dashed)
    ax.plot(Mn_nom, Pn_nom, color="#90A4AE", lw=1.5, ls="--",
            label="Nominal (Pn-Mn)")

    # Design curve (blue filled)
    ax.plot(Mn_des, Pn_des, color="#1565C0", lw=2.5,
            label="Desain phi(Pn-Mn)")
    ax.fill(Mn_des, Pn_des, alpha=0.08, color="#1565C0")

    # Control points
    ctrl = [
        (0,           pm.Pn_max,     "Pn,max\n(Pure Comp.)", "s", "#1565C0"),
        (pm.Mb,       pm.Pb,         "Balanced\n(Pb,Mb)",    "D", "#2E7D32"),
        (pm.Mn_pure,  0,             "Pure Flexure\n(Mn)",   "o", "#1565C0"),
        (0,           pm.Pn_tension, "Pure Tension\n(Pt)",   "v", "#E65100"),
    ]
    for mx, px, lbl, sym, col in ctrl:
        ax.plot(mx, px, sym, color=col, ms=9, zorder=5)
        ax.annotate(lbl, xy=(mx, px),
                    xytext=(mx + max(pm.Mn_pure*0.05, 0.3), px),
                    fontsize=8, color=col,
                    arrowprops=dict(arrowstyle="-", color=col, lw=0.6))

    # Demand point
    inside_col = "#2E7D32" if pm.inside_curve else "#C62828"
    inside_sym = "*" if pm.inside_curve else "X"
    ax.plot(pm.Mu_demand, pm.Nu_demand, inside_sym,
            color=inside_col, ms=14, zorder=6,
            label=f"Demand ({pm.Mu_demand:.2f}, {pm.Nu_demand:.2f})")
    ax.annotate(
        f"  Nu={pm.Nu_demand:.2f} kN/m\n  Mu={pm.Mu_demand:.2f} kN·m/m\n"
        + ("  [AMAN]" if pm.inside_curve else "  [TIDAK AMAN]"),
        xy=(pm.Mu_demand, pm.Nu_demand),
        xytext=(pm.Mu_demand + max(pm.Mn_pure*0.08, 0.5),
                pm.Nu_demand + pm.Pn_max * 0.05),
        fontsize=9, color=inside_col, fontweight="bold",
        arrowprops=dict(arrowstyle="->", color=inside_col, lw=1.2))

    ax.axhline(0, color="#455A64", lw=0.8, ls=":")
    ax.axvline(0, color="#455A64", lw=0.8, ls=":")
    ax.set_xlabel("Mn / φMn  (kN·m/m)", fontsize=10)
    ax.set_ylabel("Pn / φPn  (kN/m)   [+ = Tekan]", fontsize=10)
    ax.set_title(
        "Diagram Interaksi P-M — Dinding UD sebagai Kolom (Kondisi 2)\n"
        f"Pn,max={pm.Pn_max:.0f}  Pb={pm.Pb:.1f}  Mb={pm.Mb:.2f}  Mn,pure={pm.Mn_pure:.2f}",
        fontsize=10, pad=8)
    ax.legend(fontsize=9, loc="lower right")
    ax.grid(color="#ECEFF1", lw=0.6)
    ax.tick_params(labelsize=8)
    fig.tight_layout()
    return fig

_DPI   = 110   # resolusi gambar dalam laporan
_NPTS  = 60    # titik integrasi


# =============================================================================
# SHARED DATA BUILDER
# Mengumpulkan semua angka yang perlu ditampilkan di laporan
# =============================================================================

def _build_data(fr, dr, inp: dict) -> dict:
    """
    Kumpulkan semua nilai numerik dari fr/dr/inp ke dalam satu dict
    agar mudah dipakai oleh kedua format laporan.
    """
    c = _compat(inp)

    # Geometri
    H     = c.get("ud_inner_height", 600) / 1000
    ta    = c.get("ud_wall_thickness", 80) / 1000
    tb    = c.get("ud_wall_thick_bot", 100) / 1000
    ts    = c.get("ud_base_thickness", 120) / 1000
    Wo    = c.get("ud_inner_width", 600) / 1000
    L_seg = c.get("ud_length", 1.2)
    gap   = c.get("gap_cu_ud", 20.0)
    tcu   = c.get("cu_thickness", 100) / 1000

    # Material
    fc    = c.get("fc_prime", 30.0)
    fy    = c.get("fy_main",  420.0)
    fyt   = c.get("fy_shear", 240.0)
    gam_c = c.get("gamma_c",  24.0)

    # Tanah
    gs    = c.get("gamma_s",  18.0)
    phi   = c.get("phi_soil", 30.0)
    c_s   = c.get("cohesion", 0.0)
    Hf    = c.get("soil_fill_beside", 0.0)
    fill_t= c.get("fill_type_idx", 0)
    gam_fill = [gs, 22.0, 24.0][fill_t]
    q_beside = gam_fill * Hf

    # Kendaraan
    G     = c.get("axle_load_G", 225.0)
    P1    = G / 2.0
    x1    = c.get("wheel_dist", 0.25)
    method= c.get("lat_method_idx", 0)

    # Faktor
    gDL = c.get("gamma_DL", 1.2)
    gLL = c.get("gamma_LL", 1.6)
    phi_f = c.get("phi_flex", 0.90)
    phi_s = c.get("phi_shear_factor", c.get("phi_shear_f", 0.75))

    # Lateral hasil calc
    lat = fr.lateral
    Ka  = lat.Ka
    wb  = fr.wall_base
    cap = dr.wall_base_cap if dr else None

    # Pressure profile arrays (z = 0 top, H base)
    def sig_earth(z):
        return max(Ka * gs * z - 2 * c_s * math.sqrt(Ka), 0.0)

    z_top   = [i * H / _NPTS for i in range(_NPTS + 1)]
    sig_e   = [sig_earth(z) for z in z_top]

    if method == 0:
        heq   = _aashto_heq(H)
        sig_s_val = Ka * gs * heq + Ka * q_beside
        sig_s = [sig_s_val] * len(z_top)
        sig_b = None
    else:
        heq   = None
        sig_s_val = None
        sig_s = None
        sig_b = [
            _boussinesq_point_lateral(P1, max(x1, 0.05), max(z, 0.001), L_seg) / L_seg
            for z in z_top
        ]

    # Force arrays (z = 0 base, H top)
    z_fb = [i * H / _NPTS for i in range(_NPTS + 1)]

    def p_fact(z_from_base):
        depth = H - z_from_base
        pe = max(Ka * gs * depth - 2 * c_s * math.sqrt(Ka), 0.0)
        if method == 0:
            ps = Ka * gs * (heq or 0.0) + Ka * q_beside
        else:
            ps = (_boussinesq_point_lateral(P1, max(x1, 0.05),
                                            max(depth, 0.001), L_seg) / L_seg
                  if depth > 0 else 0.0)
        return gDL * pe + gLL * ps

    def V_at(z):
        dz = (H - z) / 40
        return sum(p_fact(z + (j + 0.5) * dz) * dz for j in range(40))

    def M_at(z):
        dz = (H - z) / 40
        return sum(p_fact(z + (j + 0.5) * dz) * ((j + 0.5) * dz) * dz
                   for j in range(40))

    V_arr  = [V_at(zi) for zi in z_fb]
    M_arr  = [M_at(zi) for zi in z_fb]
    Nu_arr = [gDL * gam_c * ta * (H - zi) for zi in z_fb]

    V_prop, M_prop = None, None
    if fr.gap_closed and fr.wall_top is not None:
        Hs = fr.wall_top.Vu
        V_prop = [V_at(zi) - Hs for zi in z_fb]
        M_prop = [M_at(zi) - Hs * (H - zi) for zi in z_fb]

    # Derived calc values
    W_wall = gam_c * H * ta
    Mu_e   = lat.F_earth      * lat.arm_earth
    Mu_s   = lat.F_surcharge  * lat.arm_surcharge

    Ec = 4700 * math.sqrt(fc) * 1e3
    I_w = ta**3 / 12
    EI  = Ec * I_w
    w_base_e = max(Ka * gs * H - 2 * c_s * math.sqrt(Ka), 0.0)
    w_sur_eq = lat.F_surcharge / H if H > 0 else 0.0
    d_e   = w_base_e * H**4 / (30 * EI)
    d_s   = w_sur_eq  * H**4 / (8  * EI)
    d_tot = (d_e + d_s) * 1000

    return dict(
        # Geometry
        H=H, ta=ta, tb=tb, ts=ts, Wo=Wo, L_seg=L_seg,
        gap=gap, tcu=tcu,
        # Material
        fc=fc, fy=fy, fyt=fyt, gam_c=gam_c,
        # Soil
        gs=gs, phi=phi, c_s=c_s, Hf=Hf, q_beside=q_beside,
        # Vehicle
        G=G, P1=P1, x1=x1, method=method, heq=heq,
        sig_s_val=sig_s_val,
        # Factors
        gDL=gDL, gLL=gLL, phi_f=phi_f, phi_s=phi_s,
        # Lateral
        Ka=Ka, lat=lat,
        # Forces
        wb=wb, cap=cap, pm=dr.pm_curve if dr else None,
        # Arrays
        z_top=z_top, sig_e=sig_e, sig_s=sig_s, sig_b=sig_b,
        z_fb=z_fb, V_arr=V_arr, M_arr=M_arr, Nu_arr=Nu_arr,
        V_prop=V_prop, M_prop=M_prop,
        # Derived
        W_wall=W_wall, Mu_e=Mu_e, Mu_s=Mu_s,
        Ec=Ec, I_w=I_w, EI=EI,
        w_base_e=w_base_e, w_sur_eq=w_sur_eq,
        d_e_mm=d_e*1000, d_s_mm=d_s*1000, d_tot=d_tot,
        # Condition
        cond=fr.condition, gap_closed=fr.gap_closed,
    )


def _save_fig_png(fig: plt.Figure, dpi: int = _DPI) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight",
                facecolor=fig.get_facecolor())
    buf.seek(0)
    return buf.read()


# =============================================================================
# WORD (.docx) EXPORT
# =============================================================================

def build_laporan_docx(fr, dr, inp: dict, lang: str) -> bytes:
    """
    Laporan Word berisi SEMUA langkah perhitungan identik dengan layar Streamlit.
    Returns bytes untuk st.download_button.
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    d   = _build_data(fr, dr, inp)
    doc = Document()

    # ── Page setup A4 ─────────────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Cm(2))

    # ── Style helpers ─────────────────────────────────────────────────────────
    def h1(text: str):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Blue background via paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1A237E")
        pPr.append(shd)
        return p

    def h2(text: str):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after  = Pt(2)

    def body(text: str, bold=False, italic=False, size=9.5, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)

    def step_row(label, formula, sub, result_val, unit, ref="", ok=None):
        """
        One calculation step: Label | Formula | sub → result  unit  [ref]
        Format:
          Label (bold blue)
          Formula = substitusi → hasil  unit  [ref]
        """
        # Label row
        p_lbl = doc.add_paragraph()
        p_lbl.paragraph_format.left_indent   = Cm(0.5)
        p_lbl.paragraph_format.space_before  = Pt(5)
        p_lbl.paragraph_format.space_after   = Pt(1)
        r_lbl = p_lbl.add_run(label)
        r_lbl.bold = True
        r_lbl.font.size = Pt(10)
        r_lbl.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

        # Formula line
        res_fmt = f"{result_val:.4f}".rstrip("0").rstrip(".")
        ok_txt = "  ✓ OK" if ok is True else ("  ✗ NG" if ok is False else "")
        ref_txt = f"  [{ref}]" if ref else ""

        p_eq = doc.add_paragraph()
        p_eq.paragraph_format.left_indent  = Cm(1.2)
        p_eq.paragraph_format.space_before = Pt(1)
        p_eq.paragraph_format.space_after  = Pt(4)

        r_form = p_eq.add_run(formula + "\n")
        r_form.font.size = Pt(9.5)
        r_form.font.color.rgb = RGBColor(0x45, 0x5A, 0x64)

        r_sub = p_eq.add_run(f"  =  {sub}  →  ")
        r_sub.font.size = Pt(9.5)

        r_res = p_eq.add_run(f"{res_fmt}  {unit}")
        r_res.bold = True
        r_res.font.size = Pt(10)
        r_res.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        if ok_txt:
            r_ok = p_eq.add_run(ok_txt)
            r_ok.bold = True
            r_ok.font.size = Pt(10)
            r_ok.font.color.rgb = (RGBColor(0x2E, 0x7D, 0x32) if ok
                                   else RGBColor(0xC6, 0x28, 0x28))
        if ref_txt:
            r_ref = p_eq.add_run(ref_txt)
            r_ref.italic = True
            r_ref.font.size = Pt(8)
            r_ref.font.color.rgb = RGBColor(0x7B, 0x1F, 0xA2)

    def add_fig_to_doc(fig_obj: plt.Figure, width_cm=15):
        png = _save_fig_png(fig_obj)
        buf = io.BytesIO(png)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(buf, width=Cm(width_cm))
        doc.add_paragraph()

    def two_col_table(rows: list[tuple], col_w=(9, 6)):
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        tbl.columns[0].width = Cm(col_w[0])
        tbl.columns[1].width = Cm(col_w[1])
        for k, v in rows:
            row = tbl.add_row()
            row.cells[0].text = k
            row.cells[1].text = str(v)
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)
        doc.add_paragraph()

    def three_col_table(headers, rows, col_ws=(6, 5, 5)):
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            # Header shading
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "1A237E")
            shd.set(qn("w:color"), "FFFFFF")
            tc_pr.append(shd)
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        for row_data in rows:
            row = tbl.add_row()
            for i, val in enumerate(row_data):
                row.cells[i].text = str(val)
                row.cells[i].paragraphs[0].runs[0].font.size = Pt(8.5)
        doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # COVER
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    cov = doc.add_paragraph()
    cov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cov.add_run("LAPORAN PERHITUNGAN STRUKTUR\nU-Ditch & Cover Precast\n")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        f"Kondisi Analisis: {d['cond']}\n"
        f"Berdasarkan SNI 2847:2019 & AASHTO LRFD 9th Ed.\n"
        f"Tanggal: {datetime.now().strftime('%d %B %Y')}"
    ).font.size = Pt(11)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 1 — DATA & GEOMETRI
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 1 — Data & Geometri Input")
    doc.add_paragraph()

    h2("U-Ditch (UD)")
    two_col_table([
        ("Lebar dalam UD — Wo",       f"{d['Wo']*1000:.0f} mm"),
        ("Tinggi dalam UD — Ho",      f"{d['H']*1000:.0f} mm"),
        ("Tebal dinding atas — ta",   f"{d['ta']*1000:.0f} mm"),
        ("Tebal dinding bawah — tb",  f"{d['tb']*1000:.0f} mm"),
        ("Tebal slab dasar — ts",     f"{d['ts']*1000:.0f} mm"),
        ("Panjang segmen — L_seg",    f"{d['L_seg']:.2f} m"),
        ("Gap CU–UD",                 f"{d['gap']:.0f} mm"),
        ("Tebal tengah CU — tcu",     f"{d['tcu']*1000:.0f} mm"),
    ])

    h2("Material & Tanah")
    two_col_table([
        ("f'c (beton)",          f"{d['fc']:.1f} MPa"),
        ("fy (tulangan utama)",  f"{d['fy']:.0f} MPa"),
        ("fyt (sengkang)",       f"{d['fyt']:.0f} MPa"),
        ("γc (beton)",           f"{d['gam_c']:.1f} kN/m³"),
        ("γs (tanah)",           f"{d['gs']:.1f} kN/m³"),
        ("φ (sudut gesek tanah)",f"{d['phi']:.1f}°"),
        ("c (kohesi)",           f"{d['c_s']:.1f} kPa"),
        ("Timbunan samping Hf",  f"{d['Hf']:.2f} m"),
    ])

    h2("Kendaraan & Faktor Beban")
    two_col_table([
        ("Beban gandar G",             f"{d['G']:.0f} kN"),
        ("Beban satu roda P1 = G/2",   f"{d['P1']:.1f} kN"),
        ("Jarak roda ke UD — x1",      f"{d['x1']:.3f} m"),
        ("Metode tekanan lateral",      "Surcharge Ekivalen" if d["method"]==0
                                         else "Boussinesq Terpusat"),
        ("Faktor beban DL — γDL",      f"{d['gDL']:.2f}"),
        ("Faktor beban LL — γLL",      f"{d['gLL']:.2f}"),
        ("Faktor reduksi lentur — φ",  f"{d['phi_f']:.2f}"),
        ("Faktor reduksi geser — φ",   f"{d['phi_s']:.2f}"),
    ])

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 2 — TEKANAN LATERAL PER KEDALAMAN
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 2 — Tekanan Lateral per Kedalaman")

    h2("2a. Koefisien Tekanan Tanah Aktif (Rankine)")
    step_row("Ka",
             "Ka = tan²(45° − φ/2)",
             f"tan²(45° − {d['phi']}/2) = tan²({45-d['phi']/2:.1f}°)",
             d["Ka"], "—",
             "Rankine (1857); SNI 8460:2017 §C.6.4")

    doc.add_paragraph()
    h2("2b. Distribusi Tekanan Tanah — SEGITIGA")
    body("σh(z) = Ka × γs × z − 2c√Ka  ≥ 0   (z = kedalaman dari puncak dinding)",
         italic=True)

    H  = d["H"]; Ka = d["Ka"]; gs = d["gs"]; c_s = d["c_s"]

    def se(z): return max(Ka * gs * z - 2 * c_s * math.sqrt(Ka), 0.0)

    step_row("σh  [z = 0  (puncak)]",
             "σh = Ka × γs × z − 2c√Ka",
             f"{Ka:.3f} × {gs} × 0 − 2×{c_s}×√{Ka:.3f}",
             0.0, "kPa", "SNI 8460:2017 §6.4.1")
    step_row(f"σh  [z = {H/2:.2f} m  (tengah)]",
             "σh = Ka × γs × z − 2c√Ka",
             f"{Ka:.3f} × {gs} × {H/2:.3f} − 2×{c_s}×√{Ka:.3f}",
             se(H/2), "kPa", "SNI 8460:2017 §6.4.1")
    step_row(f"σh  [z = {H:.2f} m  (DASAR)]",
             "σh = Ka × γs × z − 2c√Ka",
             f"{Ka:.3f} × {gs} × {H:.3f} − 2×{c_s}×√{Ka:.3f}",
             se(H), "kPa", "SNI 8460:2017 §6.4.1")

    # Pressure table
    body("Tabel tekanan tanah per kedalaman:", bold=True)
    step_z = max(H / 8, 0.05)
    tbl_rows_p = []
    z_v = 0.0
    while z_v <= H + 0.001:
        tbl_rows_p.append((f"{z_v:.2f} m", f"{se(z_v):.3f} kPa"))
        z_v += step_z
    two_col_table(tbl_rows_p, col_w=(6, 6))

    doc.add_paragraph()
    h2("2c. Tekanan Lateral Kendaraan")

    if d["method"] == 0:
        body("Metode: Surcharge Ekivalen — DISTRIBUSI KOTAK (uniform seluruh ketinggian)",
             italic=True)
        step_row("heq",
                 "heq = f(H_wall)  dari AASHTO Table 3.11.6.4-2",
                 f"H_wall = {H:.2f} m → interpolasi tabel",
                 d["heq"], "m",
                 "AASHTO LRFD 9th Ed. Table 3.11.6.4-2")
        step_row("σ_kend  (uniform sepanjang H)",
                 "σ_kend = Ka × γs × heq + Ka × q_timbunan",
                 f"{Ka:.3f} × {gs} × {d['heq']:.2f} + {Ka:.3f} × {d['q_beside']:.2f}",
                 d["sig_s_val"], "kPa  (konstan 0 s/d H)",
                 "AASHTO LRFD §3.11.6.4")
        body("→ Tekanan kotak/rectangular: nilai konstan dari puncak sampai dasar.", italic=True)
    else:
        body("Metode: Boussinesq 3D — DISTRIBUSI PARABOLA",
             italic=True)
        body(f"σh(z) = 3 × P1 × x1² × z³ / (2π × R⁵ × L)   dimana R = √(x1² + z²)",
             italic=True)
        step_row("P1 = G/2",
                 "P1 = G / 2",
                 f"{d['G']:.0f} / 2",
                 d["P1"], "kN", "SNI 1725:2016 — beban gandar")
        step_row("x1  (jarak roda ke muka dinding)",
                 "x1 — input pengguna",
                 f"{d['x1']:.3f} m",
                 d["x1"], "m", "—")
        body("Tabel σh Boussinesq per kedalaman:", bold=True)
        tbl_rows_b = []
        z_v = max(H / 10, 0.05)
        while z_v <= H + 0.001:
            R  = math.sqrt(d["x1"]**2 + z_v**2)
            sb = _boussinesq_point_lateral(d["P1"], max(d["x1"], 0.05),
                                           z_v, d["L_seg"]) / d["L_seg"]
            tbl_rows_b.append((f"{z_v:.2f} m", f"{R:.3f} m", f"{sb:.3f} kPa"))
            z_v += max(H / 10, 0.05)
        three_col_table(["z (m)", "R = √(x1²+z²) (m)", "σh Boussinesq (kPa)"],
                        tbl_rows_b)
        z_peak = d["x1"] / math.sqrt(2)
        body(f"→ Puncak tekanan terjadi sekitar z ≈ x1/√2 = {z_peak:.2f} m  "
             "(tekanan mengecil di atas dan di bawah titik ini)", italic=True)

    # ── Gambar distribusi tekanan ──────────────────────────────────────────────
    doc.add_paragraph()
    h2("Diagram Distribusi Tekanan Lateral")

    lat  = d["lat"]
    fig_dist = _fig_distribusi(
        d["sig_e"], d["sig_s"], d["sig_b"], d["z_top"],
        H, d["method"],
        lat.F_earth, lat.F_surcharge,
        lat.arm_earth, lat.arm_surcharge, lang,
    )
    add_fig_to_doc(fig_dist, width_cm=15)
    plt.close(fig_dist)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 3 — RESULTAN GAYA & LENGAN MOMEN
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 3 — Resultan Gaya & Lengan Momen")

    sig_top  = d["sig_e"][0]
    sig_base = d["sig_e"][-1]

    step_row("F_tanah  (resultante tekanan tanah — segitiga)",
             "F_tanah = ½ × (σ_top + σ_base) × H",
             f"½ × ({sig_top:.3f} + {sig_base:.3f}) × {H:.3f}",
             lat.F_earth, "kN/m",
             "AASHTO LRFD §3.11.5.1")

    denom = sig_base + sig_top
    arm_e_formula = (f"({H:.3f}/3) × ({sig_base:.3f} + 2×{sig_top:.3f})"
                     f" / ({sig_base:.3f} + {sig_top:.3f})"
                     if denom > 0 else f"{H:.3f}/3")
    step_row("y_tanah  (lengan dari dasar)",
             "y_tanah = (H/3) × (σ_base + 2σ_top) / (σ_base + σ_top)",
             arm_e_formula,
             lat.arm_earth, "m dari dasar",
             "Statics — centroid trapesium")

    if d["method"] == 0:
        step_row("F_kendaraan  (surcharge, kotak)",
                 "F_sur = σ_sur × H",
                 f"{d['sig_s_val']:.3f} × {H:.3f}",
                 lat.F_surcharge, "kN/m",
                 "AASHTO LRFD §3.11.6.4")
        step_row("y_kendaraan  (surcharge — tengah tinggi)",
                 "y_sur = H/2  (distribusi uniform)",
                 f"{H:.3f} / 2",
                 lat.arm_surcharge, "m dari dasar",
                 "Statics — centroid persegi panjang")
    else:
        step_row("F_Boussinesq  (integrasi numerik)",
                 "F_Bou = ∫₀ᴴ σh(z) dz   (numerik, N=60 titik)",
                 f"P1={d['P1']:.0f}kN, x1={d['x1']:.2f}m, H={H:.2f}m, L={d['L_seg']}m",
                 lat.F_surcharge, "kN/m",
                 "Boussinesq (1885); integrasi numerik")
        step_row("y_Boussinesq  (dari integrasi momen)",
                 "y_Bou = ∫σh(z)·z dz / F_Bou   (numerik)",
                 "integrasi numerik",
                 lat.arm_surcharge, "m dari dasar", "Numerik")

    doc.add_paragraph()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 4 — GAYA DALAM: Mu, Vu, Nu
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 4 — Gaya Dalam: Mu, Vu, Nu")

    wb  = d["wb"]
    gDL = d["gDL"]; gLL = d["gLL"]

    h2("4a. Gaya Aksial Nu — Berat Sendiri Dinding UD")
    step_row("W_wall  (berat dinding per m lebar)",
             "W_wall = γc × H × ta",
             f"{d['gam_c']} × {H:.3f} × {d['ta']:.3f}",
             d["W_wall"], "kN/m",
             "SNI 1727:2020 §3.1")
    step_row("Nu  (terfaktor, di dasar dinding)",
             "Nu = γDL × W_wall",
             f"{gDL} × {d['W_wall']:.4f}",
             wb.Nu, "kN/m",
             "SNI 2847:2019 Ps.5.3.1")

    doc.add_paragraph()
    h2("4b. Momen Lentur Mu — Kantilever di Dasar Dinding")
    body("M_dasar = Σ (Gaya × Lengan)  dari semua tekanan lateral", italic=True)

    step_row("M_tanah  [service, unfactored]",
             "M_tanah = F_tanah × y_tanah",
             f"{lat.F_earth:.4f} × {lat.arm_earth:.4f}",
             d["Mu_e"], "kN·m/m",
             "Statics — momen kantilever")
    step_row(f"M_kendaraan  [service]",
             "M_kend = F_kend × y_kend",
             f"{lat.F_surcharge:.4f} × {lat.arm_surcharge:.4f}",
             d["Mu_s"], "kN·m/m",
             "Statics — momen kantilever")
    step_row("Mu  [TERFAKTOR — di dasar dinding]",
             "Mu = γDL × M_tanah + γLL × M_kend",
             f"{gDL} × {d['Mu_e']:.4f} + {gLL} × {d['Mu_s']:.4f}",
             wb.Mu, "kN·m/m",
             "SNI 2847:2019 Ps.5.3.1  (U = 1.2D + 1.6L)")

    doc.add_paragraph()
    h2("4c. Gaya Geser Vu — di Dasar Dinding")
    step_row("Vu  [TERFAKTOR — di dasar dinding]",
             "Vu = γDL × F_tanah + γLL × F_kend",
             f"{gDL} × {lat.F_earth:.4f} + {gLL} × {lat.F_surcharge:.4f}",
             wb.Vu, "kN/m",
             "SNI 2847:2019 Ps.5.3.1")

    # Gap check (Kondisi 1)
    if d["cond"] == "Kondisi 1":
        doc.add_paragraph()
        h2("4d. Cek Gap — Kantilever Murni atau CU Aktif sebagai Strut?")
        step_row("Ec  (modulus elastisitas beton)",
                 "Ec = 4700√f'c",
                 f"4700 × √{d['fc']}",
                 d["Ec"] / 1e3, "MPa",
                 "SNI 2847:2019 Ps.19.2.2.1")
        step_row("I_dinding  (per m lebar)",
                 "I = ta³ / 12",
                 f"{d['ta']:.4f}³ / 12",
                 d["I_w"], "m⁴/m",
                 "Statics")
        step_row("EI  (kekakuan lentur dinding)",
                 "EI = Ec × I",
                 f"{d['Ec']:.1f} × {d['I_w']:.6f}",
                 d["EI"], "kN·m²/m", "—")
        step_row("δ_tanah  (defleksi dari tekanan tanah — segitiga)",
                 "δ_tanah = w_base × H⁴ / (30 × EI)",
                 f"{d['w_base_e']:.3f} × {H:.3f}⁴ / (30 × {d['EI']:.1f})",
                 d["d_e_mm"], "mm",
                 "Timoshenko — kantilever beban segitiga")
        step_row("δ_kend  (defleksi dari surcharge/Boussinesq ekivalen uniform)",
                 "δ_kend = w_eq × H⁴ / (8 × EI)",
                 f"{d['w_sur_eq']:.3f} × {H:.3f}⁴ / (8 × {d['EI']:.1f})",
                 d["d_s_mm"], "mm",
                 "Timoshenko — kantilever beban merata")
        step_row("δ_total",
                 "δ_total = δ_tanah + δ_kend",
                 f"{d['d_e_mm']:.3f} + {d['d_s_mm']:.3f}",
                 d["d_tot"], "mm",
                 "Superposisi")

        gap = d["gap"]
        if d["gap_closed"]:
            body(f"→ δ = {d['d_tot']:.3f} mm  ≥  gap = {gap:.0f} mm "
                 "→ CU aktif sebagai strut (PROPPED CANTILEVER)",
                 bold=True)
            body("  Momen di dasar BERKURANG. Tulangan dalam dicek untuk kondisi ini.")
        else:
            body(f"→ δ = {d['d_tot']:.3f} mm  <  gap = {gap:.0f} mm "
                 "→ Kantilever murni (CU belum aktif)",
                 bold=True)
            body("  Tulangan luar (tarik) didesain dari Mu kantilever di atas.")

    # Ringkasan gaya dalam
    doc.add_paragraph()
    body("RINGKASAN GAYA DALAM DI DASAR DINDING", bold=True)
    two_col_table([
        ("Mu  (momen lentur terfaktor)", f"{wb.Mu:.4f}  kN·m/m"),
        ("Vu  (gaya geser terfaktor)",   f"{wb.Vu:.4f}  kN/m"),
        ("Nu  (aksial tekan terfaktor)", f"{wb.Nu:.4f}  kN/m"),
    ])

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 5 — DIAGRAM GAYA DALAM
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 5 — Diagram Gaya Dalam Sepanjang Tinggi Dinding")
    body("Diagram dihitung dari integrasi distribusi tekanan lateral "
         "(bukan hanya nilai di dasar dinding).",
         italic=True)

    fig_gd = _fig_gaya_dalam(
        d["z_fb"], d["V_arr"], d["M_arr"], d["Nu_arr"],
        d["V_prop"], d["M_prop"],
        wb.Mu, wb.Vu, wb.Nu,
        H, d["gap_closed"], lang,
    )
    add_fig_to_doc(fig_gd, width_cm=16)
    plt.close(fig_gd)

    doc.add_page_break()

    # ─────────────────────────────────────────────────────────────────────────
    # BAGIAN 6 — KAPASITAS PENAMPANG & KONTROL
    # ─────────────────────────────────────────────────────────────────────────
    h1("BAGIAN 6 — Kapasitas Penampang & Kontrol")

    cap = d["cap"]
    if cap:
        h2("6a. Kapasitas Lentur")
        step_row("As,req  (tulangan tarik perlu dari Mu)",
                 "φ·As·fy·(d − a/2) = Mu   →  selesaikan As",
                 f"φ={d['phi_f']}, Mu={wb.Mu:.3f}kN·m/m",
                 cap.As_req_mm2, "mm²/m",
                 "SNI 2847:2019 Ps.22.3.2")
        step_row("As,min  (tulangan minimum dinding)",
                 "As,min = 0.0012 × b × h",
                 f"0.0012 × 1000 × {d['ta']*1000:.0f}",
                 cap.As_min_mm2, "mm²/m",
                 "SNI 2847:2019 Ps.11.6.1")
        body(f"As terpasang = {cap.As_prov_mm2:.0f} mm²/m")
        step_row("φMn  (kuat lentur rencana)",
                 "φMn = φ·As·fy·(d − a/2)",
                 f"φ={d['phi_f']}",
                 cap.phi_Mn_kNm, "kN·m/m",
                 "SNI 2847:2019 Ps.22.3.2",
                 ok=cap.flexure_ok)
        if cap.flexure_ok:
            body(f"✓ OK — φMn = {cap.phi_Mn_kNm:.3f} kN·m/m  ≥  Mu = {wb.Mu:.3f} kN·m/m",
                 bold=True)
        else:
            body(f"✗ NG — φMn = {cap.phi_Mn_kNm:.3f} kN·m/m  <  Mu = {wb.Mu:.3f} kN·m/m"
                 " — PERLU TAMBAH TULANGAN",
                 bold=True)

        doc.add_paragraph()
        h2("6b. Kapasitas Geser")
        step_row("φVc  (kapasitas geser beton)",
                 "φVc = φ × (0.17λ√f'c + Nu/(6Ag)) × b × d",
                 f"φ={d['phi_s']}",
                 cap.phi_Vc_kNm, "kN/m",
                 "SNI 2847:2019 Tabel 22.5.5.1 Pers.(a)")
        if cap.stirrup_needed:
            step_row("Vs,req",
                     "Vs = Vu/φ − Vc",
                     f"{wb.Vu:.3f}/{d['phi_s']} − {cap.phi_Vc_kNm/d['phi_s']:.3f}",
                     cap.Vs_req_kNm, "kN/m",
                     "SNI 2847:2019 Ps.22.5.1.1")
            body(f"→ PERLU SENGKANG: Av,req = {cap.Av_req_mm2:.0f} mm²/m", bold=True)
        else:
            body(f"✓ OK — φVc = {cap.phi_Vc_kNm:.3f} kN/m ≥ Vu = {wb.Vu:.3f} kN/m"
                 " → Beton cukup tanpa sengkang geser",
                 bold=True)

        step_row("φVn  (total kuat geser rencana)",
                 "φVn = φ(Vc + Vs)",
                 f"φ={d['phi_s']}",
                 cap.phi_Vn_kNm, "kN/m",
                 "SNI 2847:2019 Ps.22.5.1.1",
                 ok=cap.shear_ok)

    # P-M Kondisi 2
    pm = d.get("pm")
    if pm and d["cond"] == "Kondisi 2":
        doc.add_paragraph()
        h2("6c. Diagram Interaksi P-M — Dinding UD sebagai Kolom")

        # ── P-M calculation steps ─────────────────────────────────────────────
        body("Diagram interaksi dibangun dengan metode kompatibilitas regangan "
             "(strain compatibility), menyapu kedalaman sumbu netral c dari 0 s/d h.",
             italic=True)
        body("Titik kontrol kurva P-M:", bold=True)
        step_row("Pn,max  (tekan murni, c→∞)",
                 "Pn,max = 0.80 × [0.85f'c(Ag − Ast) + Ast×fy]",
                 f"0.80×[0.85×{d['fc']}×(Ag−Ast)+Ast×{d['fy']}]",
                 pm.Pn_max, "kN/m",
                 "SNI 2847:2019 Ps.22.4.2.1")
        step_row("cb  (kedalaman sumbu netral balanced)",
                 "cb = εu × d_in / (εu + εy)    εy = fy/Es",
                 f"0.003 × d / (0.003 + {d['fy']}/200000)",
                 abs(pm.Pb / pm.Pn_max * 100), "— (lihat tabel)",
                 "SNI 2847:2019 Ps.22.3.4")
        step_row("Pb  (aksial saat balanced)",
                 "Pb = f(cb)  dari kompatibilitas regangan",
                 f"cb={abs(pm.Pb):.1f}kN/m",
                 pm.Pb, "kN/m",
                 "SNI 2847:2019 Ps.22.3.4")
        step_row("Mb  (momen saat balanced)",
                 "Mb = f(cb)",
                 f"Mb dari titik balanced",
                 pm.Mb, "kN·m/m",
                 "SNI 2847:2019 Ps.22.3.4")
        step_row("Mn,pure  (lentur murni, Pn=0)",
                 "Pn=0 → iterasi c → Mn",
                 "biseksi hingga Pn=0",
                 pm.Mn_pure, "kN·m/m",
                 "SNI 2847:2019 Ps.22.3")
        step_row("Pn,tension  (tarik murni)",
                 "Pt = −(As,out + As,in) × fy",
                 f"−(As_out+As_in)×{d['fy']}",
                 pm.Pn_tension, "kN/m",
                 "SNI 2847:2019 Ps.22.4.3")

        # ── P-M Figure ────────────────────────────────────────────────────────
        doc.add_paragraph()
        body("Diagram Interaksi P-M (kurva nominal & desain φ-reduksi):", bold=True)
        fig_pm_doc = _fig_pm_curve_laporan(pm)
        add_fig_to_doc(fig_pm_doc, width_cm=14)
        plt.close(fig_pm_doc)

        doc.add_paragraph()
        body("Tabel titik kontrol P-M:", bold=True)
        two_col_table([
            ("Pn,max (tekan aksial maks)",  f"{pm.Pn_max:.2f} kN/m"),
            ("Pb (balanced axial)",          f"{pm.Pb:.2f} kN/m"),
            ("Mb (balanced moment)",         f"{pm.Mb:.3f} kN·m/m"),
            ("Mn,pure (lentur murni)",        f"{pm.Mn_pure:.3f} kN·m/m"),
            ("Nu demand",                    f"{pm.Nu_demand:.3f} kN/m"),
            ("Mu demand",                    f"{pm.Mu_demand:.3f} kN·m/m"),
        ])
        status = ("✓ AMAN — Titik beban dalam kurva P-M"
                  if pm.inside_curve else
                  "✗ TIDAK AMAN — Titik beban di luar kurva P-M")
        body(status, bold=True)

        # ── P-M Kurva Kedua (momen berbalik) ─────────────────────────────────
        pm_rev = getattr(dr, "pm_curve_reversed", None)
        if pm_rev is not None:
            doc.add_paragraph()
            h2("6d. Kurva P-M Kedua — Momen Berbalik (Zona Atas Dinding)")
            body(
                "Kurva ini berlaku untuk zona atas dinding saat CU aktif sebagai strut. "
                "Arah momen berbalik → tulangan DALAM menjadi tarik, tulangan LUAR menjadi tekan.",
                italic=True)
            body("Titik kontrol kurva P-M (tulangan dalam = tarik):", bold=True)
            step_row("Pn,max", "Pn,max = 0.80×[0.85f'c(Ag−Ast)+Ast×fy]",
                     f"sama dengan kurva 1", pm_rev.Pn_max, "kN/m",
                     "SNI 2847:2019 Ps.22.4.2.1")
            step_row("Pb  (balanced — tul.dalam sebagai tarik)",
                     "Pb = f(cb)", "cb dari kompatibilitas regangan",
                     pm_rev.Pb, "kN/m", "SNI 2847:2019 Ps.22.3.4")
            step_row("Mb", "Mb = f(cb)", "momen di titik balanced",
                     pm_rev.Mb, "kN·m/m", "SNI 2847:2019 Ps.22.3.4")
            step_row("Mn,pure", "Pn=0 → iterasi c", "biseksi",
                     pm_rev.Mn_pure, "kN·m/m", "SNI 2847:2019 Ps.22.3")
            step_row("Pn,tension", "Pt = −(As,out+As,in)×fy",
                     f"−Ast×{d['fy']}", pm_rev.Pn_tension, "kN/m",
                     "SNI 2847:2019 Ps.22.4.3")
            doc.add_paragraph()
            body("Diagram Interaksi P-M Kedua:", bold=True)
            fig_pm_rev_doc = _fig_pm_curve_laporan(pm_rev)
            add_fig_to_doc(fig_pm_rev_doc, width_cm=14)
            plt.close(fig_pm_rev_doc)
            doc.add_paragraph()
            two_col_table([
                ("Pn,max",             f"{pm_rev.Pn_max:.2f} kN/m"),
                ("Pb (balanced)",      f"{pm_rev.Pb:.2f} kN/m"),
                ("Mb (balanced)",      f"{pm_rev.Mb:.3f} kN·m/m"),
                ("Mn,pure",            f"{pm_rev.Mn_pure:.3f} kN·m/m"),
                ("Nu demand (zona atas)", f"{pm_rev.Nu_demand:.3f} kN/m"),
                ("Mu demand (zona atas)", f"{pm_rev.Mu_demand:.3f} kN·m/m"),
            ])
            status_rev = ("✓ AMAN — Zona atas dalam kurva P-M"
                          if pm_rev.inside_curve else
                          "✗ TIDAK AMAN — Zona atas DI LUAR kurva P-M → tambah tul.dalam")
            body(status_rev, bold=True)

    # ─────────────────────────────────────────────────────────────────────────
    # PENUTUP
    # ─────────────────────────────────────────────────────────────────────────
    doc.add_page_break()
    body("CATATAN / NOTES:", bold=True)
    body("1. Laporan ini dibuat otomatis oleh U-Ditch Analysis App.")
    body("2. Selalu verifikasi hasil dengan engineer yang bertanggung jawab.")
    body("3. Referensi: SNI 2847:2019, SNI 8460:2017, SNI 1727:2020, "
         "AASHTO LRFD Bridge Design Specifications 9th Ed. (2020).")
    body(f"   Dicetak: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# =============================================================================
# PDF EXPORT (fpdf2)
# =============================================================================

def build_laporan_pdf(fr, dr, inp: dict, lang: str) -> bytes:
    """
    Laporan PDF berisi SEMUA langkah perhitungan, identik dengan Word.
    """
    from fpdf import FPDF, XPos, YPos

    def _safe(text: str, maxlen=300) -> str:
        table = str.maketrans({
            "·": ".", "×": "x", "\u00b2": "2", "\u00b3": "3",
            "\u00b0": " deg", "\u03b3": "g", "\u03c6": "phi",
            "\u03b1": "a", "\u03b2": "b", "\u03b4": "d",
            "\u03c3": "s", "\u03b5": "e", "\u03c0": "pi",
            "\u221a": "sqrt", "\u2265": ">=", "\u2264": "<=",
            "\u2014": "-", "\u2013": "-", "\u2019": "'",
            "\u2018": "'", "\u201c": '"', "\u201d": '"',
            "\u2713": "OK", "\u2717": "NG",
            "\u2705": "[OK]", "\u274c": "[NG]",
            "✓": "OK", "✗": "NG",
        })
        s = text.translate(table)
        return s.encode("latin-1", errors="replace").decode("latin-1")[:maxlen]

    d = _build_data(fr, dr, inp)

    class ReportPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 8)
            self.set_text_color(26, 35, 126)
            self.cell(0, 5,
                      "Laporan Perhitungan Struktur U-Ditch & Cover Precast",
                      new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
            self.set_draw_color(26, 35, 126)
            self.line(self.l_margin, self.get_y(),
                      self.w - self.r_margin, self.get_y())
            self.ln(2)

        def footer(self):
            self.set_y(-12)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(120, 120, 120)
            self.cell(0, 5,
                      f"SNI 2847:2019 | AASHTO LRFD 9th Ed. | "
                      f"Hal {self.page_no()}/{{nb}} | "
                      f"{datetime.now().strftime('%Y-%m-%d')}",
                      align="C")

    pdf = ReportPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(18, 18, 18)
    pdf.add_page()

    # ── Helpers ───────────────────────────────────────────────────────────────
    def ph1(text: str):
        pdf.set_font("Helvetica", "B", 12)
        pdf.set_fill_color(26, 35, 126)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(0, 8, _safe(text),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT, fill=True)
        pdf.ln(2)
        pdf.set_text_color(0, 0, 0)

    def ph2(text: str):
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(26, 35, 126)
        pdf.cell(0, 6, _safe(text),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(1)

    def pbody(text: str, bold=False, indent=False):
        pdf.set_font("Helvetica", "B" if bold else "", 9)
        pdf.set_text_color(0, 0, 0)
        if indent:
            pdf.set_x(pdf.l_margin + 8)
        pdf.multi_cell(0, 5, _safe(text),
                       new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    def pstep(label, formula, sub, result_val, unit, ref="", ok=None):
        res_fmt = f"{result_val:.4f}".rstrip("0").rstrip(".")
        ok_str  = "  [OK]" if ok is True else ("  [NG]" if ok is False else "")
        ref_str = f"  ({_safe(ref, 50)})" if ref else ""

        # Label (blue bold)
        pdf.set_x(pdf.l_margin + 4)
        pdf.set_font("Helvetica", "B", 9.5)
        pdf.set_text_color(26, 35, 126)
        pdf.cell(0, 5, _safe(label),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        # Formula line
        pdf.set_x(pdf.l_margin + 10)
        pdf.set_font("Helvetica", "I", 8.5)
        pdf.set_text_color(80, 80, 80)
        pdf.cell(0, 4.5, _safe(formula),
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)

        # Substitution → Result
        pdf.set_x(pdf.l_margin + 10)
        pdf.set_font("Helvetica", "", 9)
        pdf.set_text_color(0, 0, 0)
        pdf.write(4.5, _safe(f"= {sub}  ->  "))
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(21, 101, 192)
        pdf.write(4.5, _safe(f"{res_fmt}  {unit}"))
        if ok_str:
            col = (46, 125, 50) if ok else (198, 40, 40)
            pdf.set_text_color(*col)
            pdf.set_font("Helvetica", "B", 9)
            pdf.write(4.5, ok_str)
        if ref_str:
            pdf.set_font("Helvetica", "I", 7.5)
            pdf.set_text_color(123, 31, 162)
            pdf.write(4.5, ref_str)
        pdf.ln(5.5)
        pdf.set_text_color(0, 0, 0)

    def ptbl_2col(rows, header1="Parameter", header2="Nilai", w1=95, w2=75):
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_fill_color(26, 35, 126)
        pdf.set_text_color(255, 255, 255)
        pdf.cell(w1, 6, _safe(header1), border=1, fill=True,
                 new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.cell(w2, 6, _safe(header2), border=1, fill=True,
                 new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        for i, (k, v) in enumerate(rows):
            pdf.set_fill_color(245, 245, 252 if i % 2 == 0 else 255)
            pdf.set_font("Helvetica", "", 8.5)
            pdf.cell(w1, 5.5, _safe(str(k)), border=1, fill=True,
                     new_x=XPos.RIGHT, new_y=YPos.LAST)
            pdf.cell(w2, 5.5, _safe(str(v)), border=1, fill=True,
                     new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)

    def ptbl_3col(hdr, rows, ws=(60, 50, 60)):
        pdf.set_font("Helvetica", "B", 8.5)
        pdf.set_fill_color(26, 35, 126)
        pdf.set_text_color(255, 255, 255)
        for h, w in zip(hdr, ws):
            pdf.cell(w, 6, _safe(h), border=1, fill=True,
                     new_x=XPos.RIGHT, new_y=YPos.LAST)
        pdf.ln()
        pdf.set_text_color(0, 0, 0)
        for i, row in enumerate(rows):
            pdf.set_fill_color(245, 245, 252 if i % 2 == 0 else 255)
            for v, w in zip(row, ws):
                pdf.set_font("Helvetica", "", 8)
                pdf.cell(w, 5.5, _safe(str(v)), border=1, fill=True,
                         new_x=XPos.RIGHT, new_y=YPos.LAST)
            pdf.ln()
        pdf.ln(3)

    def padd_fig(fig_obj: plt.Figure, w_mm=160):
        png = _save_fig_png(fig_obj, dpi=100)
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tf:
            tf.write(png)
            tf_path = tf.name
        x_c = (pdf.w - w_mm) / 2
        pdf.image(tf_path, x=x_c, w=w_mm)
        os.unlink(tf_path)
        pdf.ln(3)

    # ==========================================================================
    # COVER PAGE
    # ==========================================================================
    pdf.set_font("Helvetica", "B", 18)
    pdf.set_text_color(26, 35, 126)
    pdf.ln(20)
    pdf.cell(0, 12, "LAPORAN PERHITUNGAN STRUKTUR",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 9, "U-Ditch & Cover Precast",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.set_font("Helvetica", "", 10)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)
    pdf.cell(0, 7, f"Kondisi Analisis: {d['cond']}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.cell(0, 7, "Berdasarkan SNI 2847:2019 & AASHTO LRFD 9th Ed.",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.cell(0, 7, f"Tanggal: {datetime.now().strftime('%d %B %Y')}",
             new_x=XPos.LMARGIN, new_y=YPos.NEXT, align="C")
    pdf.add_page()

    # ==========================================================================
    # BAGIAN 1 — DATA
    # ==========================================================================
    ph1("BAGIAN 1 — Data & Geometri Input")
    ph2("U-Ditch (UD)")
    ptbl_2col([
        ("Lebar dalam UD - Wo",      f"{d['Wo']*1000:.0f} mm"),
        ("Tinggi dalam UD - Ho",     f"{d['H']*1000:.0f} mm"),
        ("Tebal dinding atas - ta",  f"{d['ta']*1000:.0f} mm"),
        ("Tebal dinding bawah - tb", f"{d['tb']*1000:.0f} mm"),
        ("Tebal slab dasar - ts",    f"{d['ts']*1000:.0f} mm"),
        ("Panjang segmen - L_seg",   f"{d['L_seg']:.2f} m"),
        ("Gap CU-UD",                f"{d['gap']:.0f} mm"),
    ])
    ph2("Material, Tanah & Kendaraan")
    ptbl_2col([
        ("f'c (beton)",             f"{d['fc']:.1f} MPa"),
        ("fy (tulangan)",           f"{d['fy']:.0f} MPa"),
        ("gc (berat beton)",        f"{d['gam_c']:.1f} kN/m3"),
        ("gs (berat tanah)",        f"{d['gs']:.1f} kN/m3"),
        ("phi (sudut gesek)",       f"{d['phi']:.1f} deg"),
        ("G (beban gandar)",        f"{d['G']:.0f} kN"),
        ("P1 = G/2",                f"{d['P1']:.1f} kN"),
        ("x1 (jarak roda ke UD)",   f"{d['x1']:.3f} m"),
        ("Metode tekanan lateral",  "Surcharge" if d["method"]==0 else "Boussinesq"),
        ("gDL / gLL",               f"{d['gDL']:.2f} / {d['gLL']:.2f}"),
    ])
    pdf.add_page()

    # ==========================================================================
    # BAGIAN 2 — TEKANAN
    # ==========================================================================
    ph1("BAGIAN 2 — Tekanan Lateral per Kedalaman")
    ph2("2a. Koefisien Tekanan Tanah Aktif (Rankine)")
    pstep("Ka",
          "Ka = tan^2(45 deg - phi/2)",
          f"tan2(45 - {d['phi']}/2) = tan2({45-d['phi']/2:.1f})",
          d["Ka"], "--",
          "Rankine (1857); SNI 8460:2017 C.6.4")

    ph2("2b. Tekanan Tanah - Distribusi SEGITIGA")
    pbody("sh(z) = Ka x gs x z - 2c x sqrt(Ka)  >= 0   (z = kedalaman dari puncak)")

    H = d["H"]; Ka = d["Ka"]; gs = d["gs"]; c_s = d["c_s"]
    def se(z): return max(Ka*gs*z - 2*c_s*math.sqrt(Ka), 0.0)

    pstep("sh [z=0 (puncak)]",
          "sh = Ka x gs x z - 2c x sqrt(Ka)",
          f"{Ka:.3f} x {gs} x 0",
          0.0, "kPa", "SNI 8460:2017 6.4.1")
    pstep(f"sh [z={H/2:.2f}m (tengah)]",
          "sh = Ka x gs x z - 2c x sqrt(Ka)",
          f"{Ka:.3f} x {gs} x {H/2:.3f}",
          se(H/2), "kPa", "SNI 8460:2017 6.4.1")
    pstep(f"sh [z={H:.2f}m (DASAR)]",
          "sh = Ka x gs x z - 2c x sqrt(Ka)",
          f"{Ka:.3f} x {gs} x {H:.3f}",
          se(H), "kPa", "SNI 8460:2017 6.4.1")

    pbody("Tabel sh tanah per kedalaman:", bold=True)
    step_z = max(H / 8, 0.05)
    tbl_rows_p = []
    z_v = 0.0
    while z_v <= H + 0.001:
        tbl_rows_p.append((f"{z_v:.2f} m", f"{se(z_v):.3f} kPa"))
        z_v += step_z
    ptbl_2col(tbl_rows_p, "z (m)", "sh tanah (kPa)", w1=85, w2=85)

    ph2("2c. Tekanan Lateral Kendaraan")
    if d["method"] == 0:
        pbody("Metode: Surcharge Ekivalen - Distribusi KOTAK (uniform seluruh ketinggian)")
        pstep("heq",
              "heq = f(H_wall) dari AASHTO Table 3.11.6.4-2",
              f"H_wall={H:.2f}m -> interpolasi tabel",
              d["heq"], "m", "AASHTO LRFD 9th Ed. Table 3.11.6.4-2")
        pstep("s_kend (uniform, kotak)",
              "s_kend = Ka x gs x heq + Ka x q_timbunan",
              f"{Ka:.3f} x {gs} x {d['heq']:.2f} + {Ka:.3f} x {d['q_beside']:.2f}",
              d["sig_s_val"], "kPa (konstan 0 s/d H)",
              "AASHTO LRFD 3.11.6.4")
    else:
        pbody("Metode: Boussinesq 3D - Distribusi PARABOLA/MELENGKUNG")
        pbody("sh(z) = 3 x P1 x x1^2 x z^3 / (2pi x R^5 x L)   R = sqrt(x1^2+z^2)")
        pstep("P1 = G/2", "P1 = G / 2", f"{d['G']:.0f}/2",
              d["P1"], "kN", "SNI 1725:2016")
        pbody("Tabel sh Boussinesq per kedalaman:", bold=True)
        tbl_bou = []
        z_v2 = max(H/10, 0.05)
        while z_v2 <= H + 0.001:
            R = math.sqrt(d["x1"]**2 + z_v2**2)
            sb = _boussinesq_point_lateral(d["P1"], max(d["x1"],0.05), z_v2, d["L_seg"])/d["L_seg"]
            tbl_bou.append((f"{z_v2:.2f}", f"{R:.3f}", f"{sb:.3f}"))
            z_v2 += max(H/10, 0.05)
        ptbl_3col(["z (m)", "R (m)", "sh Bou (kPa)"], tbl_bou, ws=(55, 55, 60))

    pbody("Diagram Distribusi Tekanan Lateral:", bold=True)
    lat = d["lat"]
    fig_dist = _fig_distribusi(
        d["sig_e"], d["sig_s"], d["sig_b"], d["z_top"],
        H, d["method"],
        lat.F_earth, lat.F_surcharge,
        lat.arm_earth, lat.arm_surcharge, lang,
    )
    padd_fig(fig_dist, w_mm=160)
    plt.close(fig_dist)
    pdf.add_page()

    # ==========================================================================
    # BAGIAN 3 — RESULTAN
    # ==========================================================================
    ph1("BAGIAN 3 — Resultan Gaya & Lengan Momen")

    sig_top  = d["sig_e"][0]
    sig_base = d["sig_e"][-1]

    pstep("F_tanah (resultante tekanan tanah)",
          "F_tanah = 0.5 x (sh_top + sh_base) x H",
          f"0.5 x ({sig_top:.3f} + {sig_base:.3f}) x {H:.3f}",
          lat.F_earth, "kN/m", "AASHTO LRFD 3.11.5.1")

    denom = sig_base + sig_top
    arm_sub = (f"({H:.3f}/3) x ({sig_base:.3f}+2x{sig_top:.3f}) / ({sig_base:.3f}+{sig_top:.3f})"
               if denom > 0 else f"{H:.3f}/3")
    pstep("y_tanah (lengan dari dasar)",
          "y = (H/3) x (sh_base+2sh_top) / (sh_base+sh_top)",
          arm_sub, lat.arm_earth, "m dari dasar",
          "Statics - centroid trapesium")

    if d["method"] == 0:
        pstep("F_kendaraan (surcharge)",
              "F_sur = s_sur x H",
              f"{d['sig_s_val']:.3f} x {H:.3f}",
              lat.F_surcharge, "kN/m", "AASHTO LRFD 3.11.6.4")
        pstep("y_kendaraan (uniform -> tengah H)",
              "y_sur = H/2",
              f"{H:.3f}/2",
              lat.arm_surcharge, "m dari dasar",
              "Statics - centroid persegi panjang")
    else:
        pstep("F_Boussinesq (numerik)",
              "F_Bou = integral sh(z) dz  (numerik N=60)",
              f"P1={d['P1']:.0f}kN, x1={d['x1']:.2f}m, L={d['L_seg']}m",
              lat.F_surcharge, "kN/m", "Boussinesq (1885)")
        pstep("y_Boussinesq (numerik)",
              "y_Bou = integral sh(z).z dz / F_Bou",
              "integrasi numerik",
              lat.arm_surcharge, "m dari dasar", "Numerik")

    # ==========================================================================
    # BAGIAN 4 — Mu, Vu, Nu
    # ==========================================================================
    ph1("BAGIAN 4 -- Gaya Dalam: Mu, Vu, Nu")
    wb = d["wb"]
    gDL = d["gDL"]; gLL = d["gLL"]

    ph2("4a. Gaya Aksial Nu")
    pstep("W_wall (berat dinding)",
          "W = gc x H x ta",
          f"{d['gam_c']} x {H:.3f} x {d['ta']:.3f}",
          d["W_wall"], "kN/m", "SNI 1727:2020 3.1")
    pstep("Nu (terfaktor)",
          "Nu = gDL x W_wall",
          f"{gDL} x {d['W_wall']:.4f}",
          wb.Nu, "kN/m", "SNI 2847:2019 Ps.5.3.1")

    ph2("4b. Momen Lentur Mu")
    pstep("M_tanah [service]",
          "M = F_tanah x y_tanah",
          f"{lat.F_earth:.4f} x {lat.arm_earth:.4f}",
          d["Mu_e"], "kN.m/m", "Statics")
    pstep("M_kend [service]",
          "M = F_kend x y_kend",
          f"{lat.F_surcharge:.4f} x {lat.arm_surcharge:.4f}",
          d["Mu_s"], "kN.m/m", "Statics")
    pstep("Mu [TERFAKTOR]",
          "Mu = gDL x M_tanah + gLL x M_kend",
          f"{gDL} x {d['Mu_e']:.4f} + {gLL} x {d['Mu_s']:.4f}",
          wb.Mu, "kN.m/m", "SNI 2847:2019 Ps.5.3.1 (U=1.2D+1.6L)")

    ph2("4c. Gaya Geser Vu")
    pstep("Vu [TERFAKTOR]",
          "Vu = gDL x F_tanah + gLL x F_kend",
          f"{gDL} x {lat.F_earth:.4f} + {gLL} x {lat.F_surcharge:.4f}",
          wb.Vu, "kN/m", "SNI 2847:2019 Ps.5.3.1")

    if d["cond"] == "Kondisi 1":
        ph2("4d. Cek Gap - Kantilever atau CU Aktif?")
        pstep("Ec", "Ec = 4700 x sqrt(f'c)",
              f"4700 x sqrt({d['fc']})", d["Ec"]/1e3, "MPa",
              "SNI 2847:2019 Ps.19.2.2.1")
        pstep("I_wall", "I = ta^3 / 12",
              f"{d['ta']:.4f}^3 / 12", d["I_w"], "m4/m", "Statics")
        pstep("delta_tanah", "d = w_base x H^4 / (30 x EI)",
              f"{d['w_base_e']:.3f} x {H:.3f}^4 / (30 x {d['EI']:.1f})",
              d["d_e_mm"], "mm", "Timoshenko - kantilever segitiga")
        pstep("delta_kend", "d = w_eq x H^4 / (8 x EI)",
              f"{d['w_sur_eq']:.3f} x {H:.3f}^4 / (8 x {d['EI']:.1f})",
              d["d_s_mm"], "mm", "Timoshenko - kantilever merata")
        pstep("delta_total", "d_total = d_tanah + d_kend",
              f"{d['d_e_mm']:.3f} + {d['d_s_mm']:.3f}",
              d["d_tot"], "mm", "Superposisi")

        if d["gap_closed"]:
            pbody(f"-> delta={d['d_tot']:.3f}mm >= gap={d['gap']:.0f}mm "
                  "-> CU AKTIF (Propped Cantilever)", bold=True)
        else:
            pbody(f"-> delta={d['d_tot']:.3f}mm < gap={d['gap']:.0f}mm "
                  "-> Kantilever Murni", bold=True)

    pbody("RINGKASAN GAYA DALAM DI DASAR DINDING:", bold=True)
    ptbl_2col([
        ("Mu (momen lentur terfaktor)", f"{wb.Mu:.4f} kN.m/m"),
        ("Vu (gaya geser terfaktor)",   f"{wb.Vu:.4f} kN/m"),
        ("Nu (aksial tekan terfaktor)", f"{wb.Nu:.4f} kN/m"),
    ])
    pdf.add_page()

    # ==========================================================================
    # BAGIAN 5 — DIAGRAM
    # ==========================================================================
    ph1("BAGIAN 5 -- Diagram Gaya Dalam Sepanjang Tinggi Dinding")
    pbody("Diagram dari integrasi distribusi tekanan lateral (bukan hanya nilai di dasar).")

    fig_gd = _fig_gaya_dalam(
        d["z_fb"], d["V_arr"], d["M_arr"], d["Nu_arr"],
        d["V_prop"], d["M_prop"],
        wb.Mu, wb.Vu, wb.Nu,
        H, d["gap_closed"], lang,
    )
    padd_fig(fig_gd, w_mm=170)
    plt.close(fig_gd)
    pdf.add_page()

    # ==========================================================================
    # BAGIAN 6 — KAPASITAS & KONTROL
    # ==========================================================================
    ph1("BAGIAN 6 -- Kapasitas Penampang & Kontrol")

    cap = d["cap"]
    if cap:
        ph2("6a. Kapasitas Lentur")
        pstep("As,req (dari Mu)",
              "phi x As x fy x (d - a/2) = Mu",
              f"phi={d['phi_f']}, Mu={wb.Mu:.3f}kN.m/m",
              cap.As_req_mm2, "mm2/m", "SNI 2847:2019 Ps.22.3.2")
        pstep("As,min", "As,min = 0.0012 x b x h",
              f"0.0012 x 1000 x {d['ta']*1000:.0f}",
              cap.As_min_mm2, "mm2/m", "SNI 2847:2019 Ps.11.6.1")
        pbody(f"As terpasang = {cap.As_prov_mm2:.0f} mm2/m")
        pstep("phi_Mn (kuat lentur rencana)",
              "phi_Mn = phi x As x fy x (d - a/2)",
              f"phi={d['phi_f']}",
              cap.phi_Mn_kNm, "kN.m/m",
              "SNI 2847:2019 Ps.22.3.2",
              ok=cap.flexure_ok)
        if cap.flexure_ok:
            pbody(f"OK: phi_Mn={cap.phi_Mn_kNm:.3f} >= Mu={wb.Mu:.3f} kN.m/m", bold=True)
        else:
            pbody(f"NG: phi_Mn={cap.phi_Mn_kNm:.3f} < Mu={wb.Mu:.3f} kN.m/m"
                  " -- PERLU TAMBAH TULANGAN", bold=True)

        ph2("6b. Kapasitas Geser")
        pstep("phi_Vc (kapasitas geser beton)",
              "phi_Vc = phi x (0.17 x sqrt(fc) + Nu/(6Ag)) x b x d",
              f"phi={d['phi_s']}",
              cap.phi_Vc_kNm, "kN/m",
              "SNI 2847:2019 Tbl 22.5.5.1 Pers.(a)")
        if not cap.stirrup_needed:
            pbody(f"OK: phi_Vc={cap.phi_Vc_kNm:.3f} >= Vu={wb.Vu:.3f} kN/m"
                  " -- beton cukup", bold=True)
        else:
            pbody(f"Perlu sengkang: Av,req={cap.Av_req_mm2:.0f} mm2/m", bold=True)
        pstep("phi_Vn (total)",
              "phi_Vn = phi x (Vc + Vs)",
              f"phi={d['phi_s']}",
              cap.phi_Vn_kNm, "kN/m",
              "SNI 2847:2019 Ps.22.5.1.1",
              ok=cap.shear_ok)

    pm = d.get("pm")
    if pm and d["cond"] == "Kondisi 2":
        ph2("6c. Diagram Interaksi P-M -- Dinding UD sebagai Kolom")
        pbody("Metode: kompatibilitas regangan (strain compatibility), c disapu 0 s/d h.",
              indent=True)

        pbody("Titik kontrol kurva P-M:", bold=True)
        pstep("Pn,max (tekan murni)",
              "Pn,max = 0.80[0.85fc(Ag-Ast)+Ast.fy]",
              f"0.80x[0.85x{pm.Pn_max:.0f}...]",
              pm.Pn_max, "kN/m", "SNI 2847:2019 Ps.22.4.2.1")
        pstep("Pb (balanced axial)",
              "Pb = f(cb) dari kompatibilitas regangan",
              "cb dari epsilon_u dan epsilon_y",
              pm.Pb, "kN/m", "SNI 2847:2019 Ps.22.3.4")
        pstep("Mb (balanced moment)",
              "Mb = f(cb)",
              "momen di titik balanced",
              pm.Mb, "kN.m/m", "SNI 2847:2019 Ps.22.3.4")
        pstep("Mn,pure (lentur murni, Pn=0)",
              "iterasi c hingga Pn=0",
              "biseksi",
              pm.Mn_pure, "kN.m/m", "SNI 2847:2019 Ps.22.3")
        pstep("Pn,tension (tarik murni)",
              "Pt = -(As_out+As_in) x fy",
              f"-(As_total)x{d['fy']}",
              pm.Pn_tension, "kN/m", "SNI 2847:2019 Ps.22.4.3")

        # P-M diagram figure
        pbody("Diagram Interaksi P-M:", bold=True)
        fig_pm_pdf = _fig_pm_curve_laporan(pm)
        padd_fig(fig_pm_pdf, w_mm=160)
        plt.close(fig_pm_pdf)

        # Summary table
        pbody("Ringkasan titik kontrol:", bold=True)
        ptbl_2col([
            ("Pn,max (tekan murni)",   f"{pm.Pn_max:.2f} kN/m"),
            ("Pb (balanced axial)",    f"{pm.Pb:.2f} kN/m"),
            ("Mb (balanced moment)",   f"{pm.Mb:.3f} kN.m/m"),
            ("Mn,pure (lentur murni)", f"{pm.Mn_pure:.3f} kN.m/m"),
            ("Pn,tension (tarik murni)", f"{pm.Pn_tension:.2f} kN/m"),
            ("Nu demand",              f"{pm.Nu_demand:.3f} kN/m"),
            ("Mu demand",              f"{pm.Mu_demand:.3f} kN.m/m"),
        ])
        st_txt = ("OK - Titik beban DALAM kurva P-M (Aman)"
                  if pm.inside_curve else
                  "NG - Titik beban DI LUAR kurva P-M (Tidak Aman)")
        pbody(st_txt, bold=True)

        # ── P-M Kurva Kedua (momen berbalik) ──────────────────────────────────
        pm_rev = getattr(dr, "pm_curve_reversed", None)
        if pm_rev is not None:
            pdf.add_page()
            ph2("6d. Kurva P-M Kedua -- Momen Berbalik (Zona Atas Dinding)")
            pbody("Berlaku saat CU aktif: momen berbalik, tul.DALAM menjadi tarik, tul.LUAR tekan.",
                  indent=True)
            pbody("Titik kontrol (tul.dalam = tarik):", bold=True)
            pstep("Pn,max", "Pn,max = 0.80[0.85fc(Ag-Ast)+Ast.fy]",
                  "sama kurva 1", pm_rev.Pn_max, "kN/m",
                  "SNI 2847:2019 Ps.22.4.2.1")
            pstep("Pb (balanced, tul.dalam tarik)",
                  "Pb = f(cb)", "cb dari kompatibilitas",
                  pm_rev.Pb, "kN/m", "SNI 2847:2019 Ps.22.3.4")
            pstep("Mb", "Mb = f(cb)", "balanced moment",
                  pm_rev.Mb, "kN.m/m", "SNI 2847:2019 Ps.22.3.4")
            pstep("Mn,pure", "Pn=0 -> iterasi", "biseksi",
                  pm_rev.Mn_pure, "kN.m/m", "SNI 2847:2019 Ps.22.3")
            pstep("Pn,tension", "Pt = -(As_total) x fy",
                  f"-(As_total)x{d['fy']}", pm_rev.Pn_tension,
                  "kN/m", "SNI 2847:2019 Ps.22.4.3")
            pbody("Diagram Interaksi P-M Kedua:", bold=True)
            fig_pm_rev_pdf = _fig_pm_curve_laporan(pm_rev)
            padd_fig(fig_pm_rev_pdf, w_mm=160)
            plt.close(fig_pm_rev_pdf)
            pbody("Ringkasan titik kontrol (kurva kedua):", bold=True)
            ptbl_2col([
                ("Pn,max",              f"{pm_rev.Pn_max:.2f} kN/m"),
                ("Pb (balanced)",       f"{pm_rev.Pb:.2f} kN/m"),
                ("Mb (balanced)",       f"{pm_rev.Mb:.3f} kN.m/m"),
                ("Mn,pure",             f"{pm_rev.Mn_pure:.3f} kN.m/m"),
                ("Nu demand zona atas", f"{pm_rev.Nu_demand:.3f} kN/m"),
                ("Mu demand zona atas", f"{pm_rev.Mu_demand:.3f} kN.m/m"),
            ])
            st_rev = ("OK - Zona atas AMAN" if pm_rev.inside_curve
                      else "NG - Zona atas TIDAK AMAN -> tambah tul.dalam")
            pbody(st_rev, bold=True)

    # Footer note
    pdf.ln(8)
    pdf.set_font("Helvetica", "I", 7.5)
    pdf.set_text_color(120, 120, 120)
    pdf.multi_cell(0, 4,
                   "Catatan: Laporan dibuat otomatis. Selalu verifikasi dengan engineer "
                   "yang bertanggung jawab. "
                   "Ref: SNI 2847:2019, SNI 8460:2017, SNI 1727:2020, "
                   "AASHTO LRFD 9th Ed.")

    return bytes(pdf.output())
