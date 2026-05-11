# modules/output_bc.py
# Output module: visualization, step-by-step math, and report export

import streamlit as st
import matplotlib.pyplot as plt
import matplotlib.patches as patches
import numpy as np
import io
import math
import os
import tempfile
from datetime import datetime
from boxculvert.lang_dict import lang_dict

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# ---------- Helper: multilingual text ----------
def t(key: str) -> str:
    return lang_dict[st.session_state.get("lang", "EN")].get(key, key)

# ---------- Utility: compute internal forces along a member for plotting ----------
def _compute_forces_along_member(L, N_i, V_i, M_i, loads, num_points=80):
    """
    Returns x (array from 0 to L), N, V, M arrays for plotting.
    Based on equilibrium from left end (node i) with applied loads.
    """
    x = np.linspace(0, L, num_points)
    N = np.zeros_like(x)
    V = np.zeros_like(x)
    M = np.zeros_like(x)
    for idx, xi in enumerate(x):
        # start with end forces
        Ni = N_i
        Vi = V_i
        Mi = M_i - V_i * xi   # moment from shear at left
        # subtract loads up to xi
        for load in loads:
            if load['type'] == 'udl':
                w_perp = load['w_perp']
                w_axial = load.get('w_axial', 0.0)
                if xi > 0:
                    l_eff = min(xi, L)
                    Mi -= w_perp * l_eff**2 / 2.0
                    Vi -= w_perp * l_eff
                    Ni -= w_axial * l_eff
            elif load['type'] == 'point':
                a = load['a']
                if a <= xi:
                    P_perp = load['P_perp']
                    P_axial = load.get('P_axial', 0.0)
                    Mi -= P_perp * (xi - a)
                    Vi -= P_perp
                    Ni -= P_axial
            elif load['type'] == 'trap':
                q1 = load['q1']; q2 = load['q2']
                # linear variation from q1 at i to q2 at j
                if xi > 0:
                    l_eff = min(xi, L)
                    # intensity at distance = q1 + (q2-q1)*(x/L)
                    # resultant of trapezoid up to xi: area = (q1 + q_xi)*xi/2, centroid at ...
                    # More simply integrate: V_load = integral_0^xi q(x) dx, M_load = integral_0^xi q(x)*(xi - x_bar) dx
                    # q(x) = q1 + (q2-q1)*x/L
                    # V_load = q1*xi + (q2-q1)*xi^2/(2L)
                    # M_load = q1*xi^2/2 + (q2-q1)*xi^3/(6L)
                    V_load = q1*xi + (q2-q1)*xi**2/(2*L)
                    M_load = q1*xi**2/2 + (q2-q1)*xi**3/(6*L)
                    Vi -= V_load
                    Mi -= M_load
        N[idx] = Ni
        V[idx] = Vi
        M[idx] = Mi
    return x, N, V, M

# ---------- Plot functions ----------
def plot_cross_section(inputs):
    """Return fig of BC cross‑section with dimensions."""
    # Convert mm to m for display
    clear_span = inputs["clear_span"]/1000.0
    clear_height = inputs["clear_height"]/1000.0
    t_top = inputs["t_top_slab"]/1000.0
    t_bot = inputs["t_bottom_slab"]/1000.0
    t_wall = inputs["t_wall"]/1000.0

    fig, ax = plt.subplots(figsize=(6, 5))
    # Outer box
    outer_width = clear_span + 2*t_wall
    outer_height = t_top + clear_height + t_bot
    # Draw outer rectangle
    rect_outer = patches.Rectangle((0,0), outer_width, outer_height, linewidth=2, edgecolor='black', facecolor='lightgrey')
    ax.add_patch(rect_outer)
    # Inner void
    rect_inner = patches.Rectangle((t_wall, t_bot), clear_span, clear_height, linewidth=2, edgecolor='white', facecolor='white')
    ax.add_patch(rect_inner)

    # Annotate dimensions
    ax.annotate(f"{clear_span:.2f}", xy=(t_wall + clear_span/2, t_bot/2), ha='center', va='center', color='blue')
    ax.annotate(f"{clear_height:.2f}", xy=(t_wall/2, t_bot + clear_height/2), ha='center', va='center', color='blue', rotation=90)
    ax.annotate(f"t_top={t_top:.2f}", xy=(outer_width/2, t_bot+clear_height + t_top/2), ha='center', va='center', color='red')
    ax.annotate(f"t_bot={t_bot:.2f}", xy=(outer_width/2, t_bot/2), ha='center', va='center', color='red')
    ax.annotate(f"t_wall={t_wall:.2f}", xy=(t_wall/2, outer_height/2), ha='center', va='center', color='red', rotation=90)
    ax.annotate(f"t_wall={t_wall:.2f}", xy=(outer_width - t_wall/2, outer_height/2), ha='center', va='center', color='red', rotation=90)

    ax.set_xlim(-0.5, outer_width+0.5)
    ax.set_ylim(-0.5, outer_height+0.5)
    ax.set_aspect('equal')
    ax.set_title(t("geometry_section"))
    ax.axis('off')
    return fig

def plot_frame_diagrams(analysis_res, inputs, case='full'):
    """
    Plot BMD, SFD, NFD over the box frame for a given load case.
    case: 'empty' or 'full'
    """
    crit = analysis_res['critical_forces']
    # Extract member end forces and member properties from analysis (we need to recompute from stored data).
    # We'll use the frame model from analysis results? We didn't store the element loads explicitly.
    # However, we can re-run the analysis for that case inside the output. That's heavy. Better: we stored the member end forces and element loads in analysis_res? Not yet.
    # We'll modify calc_engine to store the elements and member data needed for plotting. Since we already have the function, we can add a "plotting_data" key.
    # But to avoid altering calc_engine (maybe we can't), we can store it in session_state during calculation.
    # For this implementation, I'll assume we have a 'frame_model' stored in session_state from calculation page that contains nodes, elements list (with loads) and member_forces for each case.
    # That's what we need. I'll check for it. If not present, show a message.
    if 'frame_model_empty' not in st.session_state or 'frame_model_full' not in st.session_state:
        st.warning("Frame model data not found. Please run Analysis first.")
        return None
    model = st.session_state['frame_model_full'] if case=='full' else st.session_state['frame_model_empty']
    # model = {'nodes': dict, 'elements': list of dicts (i,j,E,A,I,loads), 'member_forces': list of dicts (i_node,j_node,N_i,V_i,M_i,N_j,V_j,M_j,L)}
    nodes = model['nodes']
    elems = model['elements']
    mem_forces = model['member_forces']

    # Determine global positions of nodes for plotting
    node_xy = {nid: np.array([x, y]) for nid, (x,y) in nodes.items()}

    # For each member, compute local x and global coordinates along its line
    diagrams = {'BMD': [], 'SFD': [], 'NFD': []}  # list of (x_global, y_global, value)
    for mem, el in zip(mem_forces, elems):
        i_node = mem['i_node']; j_node = mem['j_node']
        L = mem['L']
        xi, yi = node_xy[i_node]
        xj, yj = node_xy[j_node]
        # Unit vector
        dx = xj - xi; dy = yj - yi
        # parameter t in [0, L]
        t_arr = np.linspace(0, L, 100)
        # local coordinates along member
        x_local, N_local, V_local, M_local = _compute_forces_along_member(
            L, mem['N_i'], mem['V_i'], mem['M_i'], el['loads'], num_points=100)
        # Map to global coordinates: (xi + dx/L * x_local, yi + dy/L * x_local)
        glob_x = xi + dx/L * x_local
        glob_y = yi + dy/L * x_local
        # store for diagrams
        diagrams['BMD'].append((glob_x, glob_y, M_local))
        diagrams['SFD'].append((glob_x, glob_y, V_local))
        diagrams['NFD'].append((glob_x, glob_y, N_local))

    fig, axes = plt.subplots(1, 3, figsize=(18, 5))
    titles = ['Bending Moment (kN·m/m)', 'Shear Force (kN/m)', 'Axial Force (kN/m)']
    for ax, key, title in zip(axes, ['BMD','SFD','NFD'], titles):
        for glob_x, glob_y, val in diagrams[key]:
            # Plot line colour by value
            points = np.array([glob_x, glob_y]).T
            # For line segments, we can use LineCollection but simpler: plot as scatter with color.
            ax.scatter(glob_x, glob_y, c=val, cmap='jet', s=10)
        ax.set_aspect('equal')
        ax.set_title(title)
        ax.axis('equal')
        # Set limits based on node positions
    # add node points
    for ax in axes:
        for nid, (x,y) in nodes.items():
            ax.plot(x, y, 'ko', markersize=4)
    fig.tight_layout()
    return fig

def plot_pm_curve(pm_points, wall_check=None):
    fig, ax = plt.subplots()
    P_values = [p[0] for p in pm_points]
    M_values = [p[1] for p in pm_points]
    ax.plot(M_values, P_values, 'b-')
    ax.set_xlabel('Moment (kN·m/m)')
    ax.set_ylabel('Axial Load (kN/m)')
    ax.set_title('P‑M Interaction Diagram (Wall)')
    ax.grid(True)
    if wall_check:
        Pu = wall_check.get('Pu_applied', None)
        Mu = wall_check.get('Mu_applied', None)
        if Pu is not None and Mu is not None:
            ax.plot(Mu, Pu, 'ro', label='Applied')
            ax.legend()
    return fig

# ---------- Step-by-step display ----------
def display_step_by_step(design_res):
    """Show all calculation steps from design results."""
    sections = {
        "Top Slab Flexure": design_res["top_slab"]["flexural_design"]["steps"],
        "Top Slab Shear": design_res["top_slab"]["shear_check"]["steps"],
        "Bottom Slab Flexure": design_res["bottom_slab"]["flexural_design"]["steps"],
        "Bottom Slab Shear": design_res["bottom_slab"]["shear_check"]["steps"],
        "Wall P‑M Interaction": design_res["walls"]["pm_steps"],
        "Wall Capacity Check": design_res["walls"]["capacity_check"]["steps"]
    }
    for title, steps in sections.items():
        with st.expander(title):
            for line in steps:
                if line.startswith("==="):
                    st.markdown(f"**{line.strip('= ')}**")
                else:
                    st.write(line)

# ---------- Report generation ----------
def generate_word_report(inputs, analysis_res, design_res, figs):
    """Generate Word document and return bytes."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Title
    doc.add_heading('Precast Box Culvert Design Report', 0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M")}')
    doc.add_paragraph()

    # 1. Inputs
    doc.add_heading('1. Input Data', level=1)
    doc.add_heading('Geometry', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells; hdr[0].text = 'Parameter'; hdr[1].text = 'Value'
    geom_items = [
        ('Clear Span', f"{inputs['clear_span']} mm"),
        ('Clear Height', f"{inputs['clear_height']} mm"),
        ('Top Slab Thickness', f"{inputs['t_top_slab']} mm"),
        ('Bottom Slab Thickness', f"{inputs['t_bottom_slab']} mm"),
        ('Wall Thickness', f"{inputs['t_wall']} mm")
    ]
    for i, (par, val) in enumerate(geom_items):
        row = table.rows[i+1].cells
        row[0].text = par; row[1].text = val

    doc.add_heading('Materials', level=2)
    doc.add_paragraph(f"Concrete f'c = {inputs['fc']} MPa")
    doc.add_paragraph(f"Rebar fy = {inputs['fy']} MPa")
    doc.add_paragraph(f"Soil density γ = {inputs['soil_density']} kN/m³")
    doc.add_paragraph(f"Friction angle φ = {inputs['friction_angle']}°")
    doc.add_paragraph(f"Cohesion c = {inputs['cohesion']} kPa")

    doc.add_heading('Loading & Arching', level=2)
    doc.add_paragraph(f"Installation: {inputs['install_condition']}")
    doc.add_paragraph(f"Trench width Bd = {inputs['Bd']} mm, fill height H = {inputs['H']} mm")
    arch = analysis_res['marston_arch']
    doc.add_paragraph(arch['formula_used'])

    # 2. Analysis results (force summary)
    doc.add_heading('2. Internal Forces', level=1)
    crit = analysis_res['critical_forces']
    for comp in ['top_slab','bottom_slab','left_wall','right_wall']:
        doc.add_heading(comp.replace('_',' ').title(), level=2)
        for case in ['empty','full']:
            if case in crit[comp] and crit[comp][case]:
                doc.add_paragraph(f"Case: {case}")
                sec_forces = crit[comp][case]
                for sec, vals in sec_forces.items():
                    doc.add_paragraph(f"  {sec}: N={vals.get('N',0):.2f} kN/m, V={vals.get('V',0):.2f} kN/m, M={vals.get('M',0):.2f} kN·m/m")

    # 3. Design step-by-step
    doc.add_heading('3. Reinforcement Design', level=1)
    for slab in ['top_slab','bottom_slab']:
        doc.add_heading(slab.replace('_',' ').title(), level=2)
        flex_steps = design_res[slab]['flexural_design']['steps']
        for line in flex_steps:
            doc.add_paragraph(line, style='List Bullet')
        shear_steps = design_res[slab]['shear_check']['steps']
        for line in shear_steps:
            doc.add_paragraph(line, style='List Bullet')

    doc.add_heading('Wall Design', level=2)
    for line in design_res['walls']['pm_steps']:
        doc.add_paragraph(line, style='List Bullet')
    for line in design_res['walls']['capacity_check']['steps']:
        doc.add_paragraph(line, style='List Bullet')

    # 4. Figures
    doc.add_heading('4. Diagrams', level=1)
    for title, fig in figs.items():
        doc.add_heading(title, level=2)
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', dpi=150)
        img_stream.seek(0)
        doc.add_picture(img_stream, width=Inches(5.5))
        doc.add_paragraph()

    # Save to bytes
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.getvalue()

class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Box Culvert Design Report', 0, 1, 'C')
    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def generate_pdf_report(inputs, analysis_res, design_res, figs):
    pdf = PDF()
    pdf.add_page()
    pdf.set_font('Arial', '', 11)

    # Title
    pdf.cell(0, 10, f'Date: {datetime.now().strftime("%Y-%m-%d %H:%M")}', 0, 1)
    pdf.ln(5)

    # Inputs
    pdf.cell(0, 10, '1. Input Data', 0, 1, 'L')
    pdf.set_font('Arial', 'B', 10)
    pdf.cell(0, 8, 'Geometry', 0, 1)
    pdf.set_font('Arial', '', 10)
    geom_str = f"Clear span: {inputs['clear_span']} mm, Height: {inputs['clear_height']} mm, t_top={inputs['t_top_slab']}, t_bot={inputs['t_bottom_slab']}, t_wall={inputs['t_wall']}"
    pdf.multi_cell(0, 6, geom_str)
    pdf.ln(2)
    pdf.cell(0, 8, 'Materials', 0, 1, 'L')
    pdf.set_font('Arial', '', 10)
    pdf.cell(0, 6, f"f'c = {inputs['fc']} MPa, fy = {inputs['fy']} MPa", 0, 1)
    pdf.cell(0, 6, f"Soil: γ={inputs['soil_density']} kN/m^3, φ={inputs['friction_angle']}°, c={inputs['cohesion']} kPa", 0, 1)
    pdf.ln(2)
    pdf.cell(0, 8, 'Arching', 0, 1, 'L')
    arch = analysis_res['marston_arch']
    for line in arch['formula_used'].split('\n'):
        pdf.cell(0, 6, line, 0, 1)

    # Internal Forces
    pdf.add_page()
    pdf.cell(0, 10, '2. Internal Forces', 0, 1, 'L')
    crit = analysis_res['critical_forces']
    for comp in ['top_slab','bottom_slab','left_wall','right_wall']:
        pdf.set_font('Arial', 'B', 10)
        pdf.cell(0, 8, comp.replace('_',' ').title(), 0, 1)
        pdf.set_font('Arial', '', 9)
        for case in ['empty','full']:
            if case in crit[comp] and crit[comp][case]:
                pdf.cell(0, 6, f"Case {case}:", 0, 1)
                for sec, vals in crit[comp][case].items():
                    txt = f"  {sec}: N={vals.get('N',0):.2f}, V={vals.get('V',0):.2f}, M={vals.get('M',0):.2f}"
                    pdf.cell(0, 6, txt, 0, 1)
        pdf.ln(2)

    # Design steps
    pdf.add_page()
    pdf.cell(0, 10, '3. Reinforcement Design Steps', 0, 1, 'L')
    design_steps_text = []
    for slab in ['top_slab','bottom_slab']:
        design_steps_text += design_res[slab]['flexural_design']['steps'] + design_res[slab]['shear_check']['steps']
    design_steps_text += design_res['walls']['pm_steps'] + design_res['walls']['capacity_check']['steps']
    for line in design_steps_text:
        pdf.set_font('Arial', '', 8)
        if line.startswith("==="):
            pdf.set_font('Arial', 'B', 9)
            pdf.cell(0, 6, line.strip('= '), 0, 1)
        else:
            pdf.multi_cell(0, 5, line)
    pdf.ln(5)

    # Figures
    for title, fig in figs.items():
        pdf.add_page()
        pdf.cell(0, 10, title, 0, 1, 'L')
        img_stream = io.BytesIO()
        fig.savefig(img_stream, format='png', dpi=100)
        img_stream.seek(0)
        # Save temp file for FPDF
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp:
            tmp.write(img_stream.read())
            tmp_path = tmp.name
        pdf.image(tmp_path, x=10, w=190)
        os.unlink(tmp_path)

    pdf_output = pdf.output(dest='S').encode('latin-1')
    return pdf_output

# ---------- Main output page ----------
def output_page():
    st.header(t("output_title"))

    # Check that analysis and design results exist
    if 'analysis_results' not in st.session_state or 'design_results' not in st.session_state:
        st.error("Please run the Analysis and Design first (go to 'Analysis & Design' page).")
        return

    inputs = st.session_state.bc_inputs
    analysis_res = st.session_state.analysis_results
    design_res = st.session_state.design_results

    # Tabs
    tab1, tab2, tab3 = st.tabs(["📊 Visuals", "📝 Calculation Steps", "📥 Export Reports"])

    with tab1:
        st.subheader("Cross‑Section")
        fig_sec = plot_cross_section(inputs)
        st.pyplot(fig_sec)

        st.subheader("Frame Analysis Diagrams")
        case_choice = st.radio("Load case for diagrams", ['empty', 'full'],
                               index=1 if inputs['water_level_fraction'] > 0 else 0,
                               key='diag_case')
        # Need frame model stored; if not, recompute? We'll assume stored during calculation.
        if 'frame_model_empty' not in st.session_state:
            st.warning("Frame model data missing. Diagrams unavailable.")
        else:
            fig_frame = plot_frame_diagrams(analysis_res, inputs, case=case_choice)
            if fig_frame:
                st.pyplot(fig_frame)

        st.subheader("Wall P‑M Interaction Curve")
        pm_points = design_res['walls']['pm_curve']
        Pu_wall = design_res['walls']['Pu_max']
        Mu_wall = design_res['walls']['Mu_max']
        fig_pm = plot_pm_curve(pm_points, wall_check={'Pu_applied': Pu_wall, 'Mu_applied': Mu_wall})
        st.pyplot(fig_pm)

    with tab2:
        st.subheader("Step‑by‑Step Design Calculations")
        display_step_by_step(design_res)

    with tab3:
        st.subheader("Export Reports")
        st.write("Generate a Word (.docx) or PDF report containing all inputs, calculations, and diagrams.")

        if st.button("Generate Word Report", key='word'):
            with st.spinner("Creating Word report..."):
                # Prepare figures for report (we'll reuse ones already generated but must avoid re-displaying)
                figs = {
                    "Cross Section": plot_cross_section(inputs),
                    "Frame Diagrams (Empty)": plot_frame_diagrams(analysis_res, inputs, 'empty') if 'frame_model_empty' in st.session_state else None,
                    "Frame Diagrams (Full)": plot_frame_diagrams(analysis_res, inputs, 'full') if 'frame_model_full' in st.session_state else None,
                    "P‑M Interaction": plot_pm_curve(pm_points, {'Pu_applied': Pu_wall, 'Mu_applied': Mu_wall})
                }
                # Remove None figures
                figs = {k: v for k, v in figs.items() if v is not None}
                doc_bytes = generate_word_report(inputs, analysis_res, design_res, figs)
            st.download_button(
                label="Download Word Report",
                data=doc_bytes,
                file_name="box_culvert_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        if st.button("Generate PDF Report", key='pdf'):
            with st.spinner("Creating PDF report..."):
                figs = {
                    "Cross Section": plot_cross_section(inputs),
                    "Frame Diagrams (Empty)": plot_frame_diagrams(analysis_res, inputs, 'empty') if 'frame_model_empty' in st.session_state else None,
                    "Frame Diagrams (Full)": plot_frame_diagrams(analysis_res, inputs, 'full') if 'frame_model_full' in st.session_state else None,
                    "P‑M Interaction": plot_pm_curve(pm_points, {'Pu_applied': Pu_wall, 'Mu_applied': Mu_wall})
                }
                figs = {k: v for k, v in figs.items() if v is not None}
                pdf_bytes = generate_pdf_report(inputs, analysis_res, design_res, figs)
            st.download_button(
                label="Download PDF Report",
                data=pdf_bytes,
                file_name="box_culvert_report.pdf",
                mime="application/pdf"
            )
