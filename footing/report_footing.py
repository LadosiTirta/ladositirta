"""
footing/report_footing.py
Professional report generator — Word (.docx) and PDF.
Both outputs are IDENTICAL in content.

Step-by-step format: formula → substitution → result → unit → code clause.
"""
import io
import numpy as np
from datetime import datetime

# ─── Word ────────────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

# ─── PDF ─────────────────────────────────────────────────────────────────────
try:
    from fpdf import FPDF
    PDF_OK = True
except ImportError:
    PDF_OK = False


# ─────────────────────────────────────────────────────────────────────────────
# SHARED FORMATTING HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def _ok(flag): return "✓ OK" if flag else "✗ FAIL"
def _fmt(v, dec=2): return f"{v:,.{dec}f}"


def _section_title(n, title):
    return f"\n{'─'*60}\n{n}. {title.upper()}\n{'─'*60}"


def build_report_lines(proj, mat, geo, columns_data,
                        soil_result, struct_result, settle_result,
                        method_used: str = "sni"):
    """
    Build ordered list of report lines (strings).
    Each line is a text string; blank lines = paragraph breaks.

    Parameters:
        proj           : dict — project info
        mat            : dict — material properties
        geo            : dict — footing geometry
        columns_data   : list of dicts — column inputs & results
        soil_result    : dict — bearing capacity check results
        struct_result  : dict — structural checks (punching, shear, flexure)
        settle_result  : dict — settlement results
        method_used    : "meyerhof" | "bowles" | "sni" | "sondir"
    """
    lines = []
    A = lines.append

    # ── COVER ────────────────────────────────────────────────────────────────
    A("ISOLATED SPREAD FOOTING DESIGN")
    A("Structural Calculation Report")
    A("")
    A(f"Project     : {proj.get('name', '-')}")
    A(f"Location    : {proj.get('location', '-')}")
    A(f"Engineer    : {proj.get('engineer', '-')}")
    A(f"Date        : {proj.get('date', datetime.today().strftime('%d %B %Y'))}")
    A(f"Document No : {proj.get('doc_no', '-')}")
    A("")

    # ── DESIGN BASIS ─────────────────────────────────────────────────────────
    A(_section_title("1", "Design Basis & References"))
    A("The following codes and references are used in this calculation:")
    A("")
    A("  [1]  SNI 2847:2019  — Persyaratan Beton Struktural untuk Bangunan Gedung")
    A("       (Structural Concrete Requirements for Buildings)")
    A("       Equivalent: ACI 318-19")
    A("")
    A("  [2]  SNI 8460:2017  — Persyaratan Perancangan Geoteknik")
    A("       (Geotechnical Design Requirements)")
    A("")
    A("  [3]  SNI 1726:2019  — Tata Cara Perencanaan Ketahanan Gempa untuk Struktur")
    A("       (Seismic Design Procedure)")
    A("")
    A("  [4]  SNI 1727:2020  — Beban Minimum untuk Perancangan Bangunan")
    A("       (Minimum Design Loads)")
    A("")
    A("  [5]  Meyerhof, G.G. (1963). Some recent research on bearing capacity.")
    A("       Canadian Geotechnical Journal.")
    A("")
    A("  [6]  Bowles, J.E. (1996). Foundation Analysis and Design, 5th Ed.")
    A("")
    A("  [7]  ACI 336.2R — Bearing and Settlement of Shallow Foundations")
    A("")

    # ── INPUT SUMMARY ─────────────────────────────────────────────────────────
    A(_section_title("2", "Input Summary"))
    A("")
    A("2.1 Material Properties")
    A(f"    Concrete strength    f'c = {mat['fc']} MPa")
    A(f"    Main steel (fy)           = {mat['fy']} MPa")
    A(f"    Shrinkage steel (fy_s)    = {mat.get('fy_s', mat['fy'])} MPa")
    A(f"    Unit weight concrete γ_c  = {mat['gamma_c']} kN/m³")
    A(f"    Unit weight soil    γ_s   = {mat['gamma_s']} kN/m³")
    A(f"    λ (concrete density)      = {mat.get('lambda', 1.0)}")
    A("")
    A("2.2 Footing Geometry")
    A(f"    Shape           = {geo['shape']}")
    A(f"    Width  B        = {geo['B']} m")
    A(f"    Length L        = {geo['L']} m")
    A(f"    Thickness t     = {geo['t']} m  ({geo['t']*1000:.0f} mm)")
    A(f"    Found. depth Df = {geo['Df']} m")
    A(f"    Soil above top  = {geo['h_soil']} m")
    A(f"    Cover           = {geo['cover']} mm")
    A(f"    Sloof position  = {geo.get('sloof_note', 'On top of footing')}")
    A("")
    A("2.3 Column Data")
    for i, col in enumerate(columns_data):
        A(f"    Column {col.get('label', f'C{i+1}')}")
        A(f"      Dimensions     : {col['bc']*1000:.0f} mm × {col['hc']*1000:.0f} mm")
        A(f"      Position (x,y) : ({col['x']:.3f} m, {col['y']:.3f} m) from footing centroid")
        A(f"      Ultimate loads : Nu = {col['Nu']:.1f} kN, Mux = {col['Mux']:.1f} kN·m, Muy = {col['Muy']:.1f} kN·m")
        A(f"                       Vux = {col['Vux']:.1f} kN, Vuy = {col['Vuy']:.1f} kN")
        A(f"      Service loads  : Ns = {col['Ns']:.1f} kN, Msx = {col['Msx']:.1f} kN·m, Msy = {col['Msy']:.1f} kN·m")
        A("")

    # ── BEARING CAPACITY ──────────────────────────────────────────────────────
    A(_section_title("3", "Bearing Capacity & Soil Pressure Check (Service Loads)"))
    A("")
    A(f"  Method Used: {soil_result.get('method_name', method_used)}")
    A(f"  Code Reference: {soil_result.get('code_ref', 'SNI 8460:2017')}")
    A("")

    A("3.1 Self-Weight & Overburden")
    gd = soil_result.get("geo_detail", {})
    A(f"    Self-weight of footing (SNI 8460:2017 Pasal 4.2):")
    A(f"    W_foot = γ_c × B × L × t")
    A(f"           = {gd.get('gamma_c','?')} × {gd.get('B','?')} × {gd.get('L','?')} × {gd.get('t','?')}")
    A(f"           = {gd.get('W_foot_kN', 0):.2f} kN")
    A("")
    A(f"    Weight of soil above footing:")
    A(f"    W_soil = γ_s × B × L × h_soil_eff")
    A(f"           = {gd.get('gamma_s','?')} × {gd.get('B','?')} × {gd.get('L','?')} × {gd.get('h_soil_eff_m','?')}")
    A(f"           = {gd.get('W_soil_kN', 0):.2f} kN")
    A("")

    sc = soil_result.get("service_check", {})
    A("3.2 Resultant Service Loads")
    A(f"    ΣNs_total = ΣNs_columns + W_foot + W_soil")
    A(f"              = {sc.get('sum_Ns_col', 0):.2f} + {gd.get('W_foot_kN',0):.2f} + {gd.get('W_soil_kN',0):.2f}")
    A(f"              = {sc.get('sum_Ns_total', 0):.2f} kN")
    A("")
    A(f"    ΣMsx_total = ΣMsx_col + ΣVsy × (t + h_soil)")
    A(f"               = {sc.get('sum_Msx_col',0):.2f} + {sc.get('Msx_from_V',0):.2f}")
    A(f"               = {sc.get('sum_Msx_total',0):.2f} kN·m")
    A("")
    A(f"    ΣMsy_total = ΣMsy_col + ΣVsx × (t + h_soil)")
    A(f"               = {sc.get('sum_Msy_col',0):.2f} + {sc.get('Msy_from_V',0):.2f}")
    A(f"               = {sc.get('sum_Msy_total',0):.2f} kN·m")
    A("")

    A("3.3 Section Properties of Footing Base")
    A(f"    A  = B × L = {geo['B']} × {geo['L']} = {geo['B']*geo['L']:.3f} m²")
    A(f"    Wx = L × B²/6 = {geo['L']} × {geo['B']}²/6 = {geo['L']*geo['B']**2/6:.4f} m³")
    A(f"    Wy = B × L²/6 = {geo['B']} × {geo['L']}²/6 = {geo['B']*geo['L']**2/6:.4f} m³")
    A("")

    A("3.4 Eccentricity Check (SNI 8460:2017 Pasal 5.3)")
    cc = soil_result.get("pressure_check", {})
    A(f"    ex = |ΣMsy_total| / ΣNs_total")
    A(f"       = {abs(sc.get('sum_Msy_total',0)):.2f} / {sc.get('sum_Ns_total',1):.2f}")
    A(f"       = {cc.get('ex_m',0):.4f} m")
    A(f"    Limit B/6 = {geo['B']}/6 = {geo['B']/6:.4f} m")
    A(f"    ex = {cc.get('ex_m',0):.4f} m {'≤' if cc.get('ok_ex') else '>'} B/6 = {geo['B']/6:.4f} m → {_ok(cc.get('ok_ex'))}")
    A("")
    A(f"    ey = |ΣMsx_total| / ΣNs_total")
    A(f"       = {abs(sc.get('sum_Msx_total',0)):.2f} / {sc.get('sum_Ns_total',1):.2f}")
    A(f"       = {cc.get('ey_m',0):.4f} m")
    A(f"    Limit L/6 = {geo['L']}/6 = {geo['L']/6:.4f} m")
    A(f"    ey = {cc.get('ey_m',0):.4f} m {'≤' if cc.get('ok_ey') else '>'} L/6 = {geo['L']/6:.4f} m → {_ok(cc.get('ok_ey'))}")
    A("")

    A("3.5 Soil Pressure at 4 Corners")
    A(f"    q = ΣNs/A ± ΣMsy/Wx ± ΣMsx/Wy")
    A(f"      = {sc.get('sum_Ns_total',0):.2f}/{geo['B']*geo['L']:.3f} ± "
      f"{sc.get('sum_Msy_total',0):.2f}/{geo['L']*geo['B']**2/6:.4f} ± "
      f"{sc.get('sum_Msx_total',0):.2f}/{geo['B']*geo['L']**2/6:.4f}")
    A(f"    q₁ (+x,+y) = {cc.get('q1', 0):.2f} kN/m²")
    A(f"    q₂ (-x,+y) = {cc.get('q2', 0):.2f} kN/m²")
    A(f"    q₃ (-x,-y) = {cc.get('q3', 0):.2f} kN/m²")
    A(f"    q₄ (+x,-y) = {cc.get('q4', 0):.2f} kN/m²")
    A("")
    qa = soil_result.get("qa_kPa", 0)
    A(f"    q_max = {cc.get('q_max_kPa',0):.2f} kN/m² {'≤' if cc.get('ok_bearing') else '>'} "
      f"qa = {qa:.2f} kN/m² → {_ok(cc.get('ok_bearing'))}")
    A(f"    q_min = {cc.get('q_min_kPa',0):.2f} kN/m² {'≥' if cc.get('ok_no_uplift') else '<'} "
      f"0 → {_ok(cc.get('ok_no_uplift'))} (no uplift)")
    A("")

    # ── STRUCTURAL DESIGN ─────────────────────────────────────────────────────
    A(_section_title("4", "Structural Design (Factored/Ultimate Loads)"))
    A("")
    qu_avg = struct_result.get("qu_avg_kPa", 0)
    A(f"    Average factored soil pressure:")
    A(f"    qu_avg = ΣNu_total / A")
    A(f"           = {struct_result.get('sum_Nu_total',0):.2f} / {geo['B']*geo['L']:.3f}")
    A(f"           = {qu_avg:.2f} kN/m²")
    A("")

    A("4.1 Effective Depth")
    A("    REINFORCEMENT DIRECTION CONVENTION (SNI 2847:2019):")
    A("    ┌──────────────────────────────────────────────────────────┐")
    A("    │ X-Direction Bars: bars that RUN ALONG the X-axis        │")
    A("    │   (bar length ≈ footing dimension in X = B)             │")
    A("    │   Bars are SPACED at distance 's' in the Y-direction    │")
    A("    │                                                          │")
    A("    │ Y-Direction Bars: bars that RUN ALONG the Y-axis        │")
    A("    │   (bar length ≈ footing dimension in Y = L)             │")
    A("    │   Bars are SPACED at distance 's' in the X-direction    │")
    A("    │                                                          │")
    A("    │ Bottom X-bars placed LOWER (larger d = dx)              │")
    A("    │ Bottom Y-bars ON TOP of X-bars (smaller d = dy)         │")
    A("    └──────────────────────────────────────────────────────────┘")
    A("")
    ed = struct_result.get("eff_depth", {})
    A(f"    Cover = {geo['cover']} mm")
    A(f"    Bar X diameter (bottom) = {ed.get('bar_x_mm', 0):.0f} mm")
    A(f"    Bar Y diameter (bottom) = {ed.get('bar_y_mm', 0):.0f} mm")
    A("")
    A(f"    dx = t - cover - Ø_x/2")
    A(f"       = {geo['t']*1000:.0f} - {geo['cover']:.0f} - {ed.get('bar_x_mm',0)/2:.1f}")
    A(f"       = {ed.get('dx', 0):.1f} mm")
    A("")
    A(f"    dy = t - cover - Ø_x - Ø_y/2")
    A(f"       = {geo['t']*1000:.0f} - {geo['cover']:.0f} - {ed.get('bar_x_mm',0):.0f} - {ed.get('bar_y_mm',0)/2:.1f}")
    A(f"       = {ed.get('dy', 0):.1f} mm")
    A("")

    # Punching shear per column
    A("4.2 Two-Way (Punching) Shear — SNI 2847:2019 Pasal 22.6")
    for i, col in enumerate(columns_data):
        label = col.get('label', f'C{i+1}')
        ps = struct_result.get("punching", {}).get(label, {})
        cap = ps.get("capacity", {})
        dem = ps.get("demand", {})
        A(f"")
        A(f"    Column {label} (bc={col['bc']*1000:.0f}mm × hc={col['hc']*1000:.0f}mm):")
        A(f"    Critical perimeter at d/2 from column face:")
        A(f"    b_crit = bc + d = {col['bc']*1000:.0f} + {ed.get('dx',0):.1f} = {dem.get('b_crit_m',0)*1000:.1f} mm")
        A(f"    h_crit = hc + d = {col['hc']*1000:.0f} + {ed.get('dx',0):.1f} = {dem.get('h_crit_m',0)*1000:.1f} mm")
        A(f"    bo = 2 × (b_crit + h_crit) = {cap.get('bo_mm',0):.1f} mm")
        A(f"    β  = hc/bc = {col['hc']/col['bc']:.3f}")
        A(f"    αs = 40 (interior column)")
        A(f"")
        A(f"    Capacity — SNI 2847:2019 Pasal 22.6.5.2:")
        A(f"    Vc1 = 0.33 × λ√f'c × bo × d                    [Butir (a)]")
        A(f"        = 0.33 × {mat.get('lambda',1.0)} × √{mat['fc']} × {cap.get('bo_mm',0):.1f} × {ed.get('dx',0):.1f}")
        A(f"        = {cap.get('Vc1_N',0)/1000:.2f} kN")
        A(f"    Vc2 = (0.17 + 0.33/β) × λ√f'c × bo × d         [Butir (b)]")
        A(f"        = (0.17 + 0.33/{col['hc']/col['bc']:.3f}) × {cap.get('sqrt_fc',0):.3f} × {cap.get('bo_mm',0):.1f} × {ed.get('dx',0):.1f}")
        A(f"        = {cap.get('Vc2_N',0)/1000:.2f} kN")
        A(f"    Vc3 = 0.083(αs×d/bo+2) × λ√f'c × bo × d        [Butir (c)]")
        A(f"        = {cap.get('Vc3_N',0)/1000:.2f} kN")
        A(f"    φVc = 0.75 × min(Vc1,Vc2,Vc3)  [governing: {cap.get('governing_eq','?')}]")
        A(f"        = 0.75 × {cap.get('Vc_governing_N',0)/1000:.2f}")
        A(f"        = {cap.get('phiVc_kN',0):.2f} kN")
        A(f"")
        A(f"    Demand:")
        A(f"    Vu = Nu_col - qu_avg × b_crit × h_crit")
        A(f"       = {col['Nu']:.2f} - {qu_avg:.2f} × {dem.get('b_crit_m',0):.4f} × {dem.get('h_crit_m',0):.4f}")
        A(f"       = {ps.get('Vu_kN',0):.2f} kN")
        A(f"")
        A(f"    Vu = {ps.get('Vu_kN',0):.2f} kN {'≤' if ps.get('ok') else '>'} φVc = {cap.get('phiVc_kN',0):.2f} kN  → {_ok(ps.get('ok'))}")
        if not ps.get('ok'):
            A(f"    *** {ps.get('recommendation', '')} ***")

    # One-way shear
    A("")
    A("4.3 One-Way (Beam) Shear — SNI 2847:2019 Pasal 22.5")
    ow = struct_result.get("one_way", {})
    for dirk, dname in [("x_dir", "X-Direction"), ("y_dir", "Y-Direction")]:
        od = ow.get(dirk, {})
        cap_ow = od.get("capacity", {})
        dem_ow = od.get("demand", {})
        bw_used = od.get("bw_mm", 0)
        arm_used = od.get("arm_m", 0)
        A(f"")
        A(f"    {dname} (critical section at d from column face):")
        A(f"    Vc = 0.17 × λ × √f'c × bw × d  [SNI 2847:2019 Pasal 22.5.5.1]")
        A(f"       = 0.17 × {mat.get('lambda',1.0)} × √{mat['fc']} × {bw_used:.1f} × {ed.get('dx',0):.1f}")
        A(f"       = {cap_ow.get('Vc_kN',0):.2f} kN")
        A(f"    φVc = 0.75 × {cap_ow.get('Vc_kN',0):.2f} = {cap_ow.get('phiVc_kN',0):.2f} kN")
        A(f"")
        A(f"    Vu = qu_avg × bw × arm")
        A(f"       = {qu_avg:.2f} × {bw_used/1000:.3f} × {arm_used:.4f}")
        A(f"       = {od.get('Vu_kN',0):.2f} kN")
        A(f"    Vu = {od.get('Vu_kN',0):.2f} kN {'≤' if od.get('ok') else '>'} φVc = {cap_ow.get('phiVc_kN',0):.2f} kN  → {_ok(od.get('ok'))}")

    # Flexure
    A("")
    A("4.4 Flexural Design — SNI 2847:2019 Pasal 22.2")
    A("    (Critical section at face of column — SNI 2847:2019 Pasal 13.2.7.1)")
    flex_results = struct_result.get("flexure", [])
    for fr in flex_results:
        A(f"")
        A(f"    {fr['location']} Bars — {fr['direction']}-Direction:")
        A(f"    Design moment at column face:")
        fd = fr.get("flex_detail", {})
        A(f"    Mu = {fr['Mu_kNm']:.2f} kN·m")
        A(f"")
        A(f"    Rn = Mu / (φ × b × d²)")
        A(f"       = {fr['Mu_kNm']*1e6:.0f} / ({fd.get('phi',0.9)} × {fd.get('b_mm',0):.0f} × {fd.get('d_mm',0):.1f}²)")
        A(f"       = {fd.get('Rn_MPa',0):.4f} MPa")
        A(f"")
        A(f"    ρ = (0.85f'c/fy) × [1 - √(1 - 2Rn/0.85f'c)]    [SNI 2847:2019 Pasal 22.2]")
        A(f"      = (0.85×{mat['fc']}/{mat['fy']}) × [1 - √(1 - 2×{fd.get('Rn_MPa',0):.4f}/(0.85×{mat['fc']}))]")
        A(f"      = {fd.get('rho',0):.6f}")
        A(f"")
        A(f"    As_req = ρ × b × d = {fd.get('rho',0):.6f} × {fd.get('b_mm',0):.0f} × {fd.get('d_mm',0):.1f}")
        A(f"           = {fr['As_required_mm2']:.2f} mm²")
        A(f"")
        md = fr.get("min_detail", {})
        A(f"    As_min (SNI 2847:2019 Pasal 9.6.1.2):")
        A(f"    As_min1 = 0.25√f'c/fy × b×d = {md.get('As_min1_mm2',0):.2f} mm²")
        A(f"    As_min2 = 1.4/fy × b×d      = {md.get('As_min2_mm2',0):.2f} mm²")
        A(f"    As_min  = max(As_min1, As_min2) = {fr['As_min_mm2']:.2f} mm²")
        A(f"")
        A(f"    As_shrink = 0.0018 × b × h  [SNI 2847:2019 Pasal 24.4.3]")
        A(f"              = 0.0018 × {fd.get('b_mm',0):.0f} × {geo['t']*1000:.0f} = {md.get('As_shrink_mm2',0):.2f} mm²")
        A(f"")
        A(f"    As_design = max(As_req, {fr['min_basis'].split('(')[0].strip()}) = {fr['As_design_mm2']:.2f} mm²")
        A(f"")
        bar = fr.get("bar", {})
        A(f"    SELECTED REINFORCEMENT:")
        A(f"    {fr['location']} {fr['direction']}-Direction Bars: D{bar.get('dia_mm',0)} @ {bar.get('spacing_mm',0)} mm c/c")
        A(f"    {fr.get('bar_description','')}")
        A(f"    As_provided = {bar.get('As_provided_mm2',0):.2f} mm² {'≥' if bar.get('As_provided_mm2',0) >= fr['As_design_mm2'] else '<'} As_design = {fr['As_design_mm2']:.2f} mm²  → {_ok(bar.get('As_provided_mm2',0) >= fr['As_design_mm2'])}")
        A(f"")
        A(f"    Capacity check: φMn = φ × As_prov × fy × (d - a/2)")
        A(f"    a = As_prov×fy/(0.85×f'c×b) = {fr.get('a_mm',0):.2f} mm")
        A(f"    φMn = 0.90 × {bar.get('As_provided_mm2',0):.2f} × {mat['fy']} × ({fd.get('d_mm',0):.1f} - {fr.get('a_mm',0):.2f}/2) × 10⁻⁶")
        A(f"        = {fr['phi_Mn_kNm']:.2f} kN·m {'≥' if fr['ok_strength'] else '<'} Mu = {fr['Mu_kNm']:.2f} kN·m  → {_ok(fr['ok_strength'])}")

    # Settlement
    A("")
    A(_section_title("5", "Settlement Check (SNI 8460:2017 Pasal 9)"))
    st = settle_result
    if st.get("immediate"):
        si = st["immediate"]
        A("")
        A("5.1 Immediate (Elastic) Settlement — Bowles (1996) / SNI 8460:2017 Lampiran D")
        A(f"    {si.get('formula','')}")
        A(f"    {si.get('substitution','')}")
        A(f"    δi = {si.get('delta_mm',0):.2f} mm")
    if st.get("consolidation"):
        sc2 = st["consolidation"]
        A("")
        A("5.2 Consolidation Settlement — Terzaghi / SNI 8460:2017 Pasal 9.3")
        A(f"    Case: {sc2.get('case','')}")
        A(f"    {sc2.get('formula','')}")
        A(f"    {sc2.get('substitution','')}")
        A(f"    Sc = {sc2.get('Sc_mm',0):.2f} mm")
    if st.get("check"):
        stc = st["check"]
        A("")
        A("5.3 Total Settlement & Allowable (SNI 8460:2017 Pasal 9.3.1)")
        A(f"    δ_total = δi + Sc = {stc.get('delta_i_mm',0):.2f} + {stc.get('delta_c_mm',0):.2f}")
        A(f"            = {stc.get('delta_total_mm',0):.2f} mm")
        A(f"    Allowable total = {stc.get('allow_total_mm',25)} mm")
        A(f"    {stc.get('delta_total_mm',0):.2f} mm {'≤' if stc.get('ok_total') else '>'} {stc.get('allow_total_mm',25)} mm → {_ok(stc.get('ok_total'))}")

    # Summary table
    A("")
    A(_section_title("6", "Summary of All Checks"))
    A("")
    A(f"  {'CHECK':<35} {'DEMAND':>12} {'CAPACITY':>12} {'STATUS':>8}")
    A(f"  {'─'*35} {'─'*12} {'─'*12} {'─'*8}")

    sum_checks = struct_result.get("summary_checks", [])
    for row in sum_checks:
        A(f"  {row[0]:<35} {row[1]:>12} {row[2]:>12} {row[3]:>8}")

    A("")
    A(f"  Code: SNI 2847:2019, SNI 8460:2017, SNI 1726:2019, SNI 1727:2020")
    A("")
    A("━" * 60)
    A(f"  Prepared by : {proj.get('engineer','-')}")
    A(f"  Date        : {proj.get('date', datetime.today().strftime('%d %B %Y'))}")
    A("━" * 60)

    return lines


# ─────────────────────────────────────────────────────────────────────────────
# WORD GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def generate_word(lines: list, fig_bytes_list: list = None) -> bytes:
    """Generate Word document from report lines. Returns bytes."""
    if not DOCX_OK:
        raise ImportError("python-docx not installed")

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)

    # Styles
    style = doc.styles['Normal']
    style.font.name = 'Courier New'
    style.font.size = Pt(9)

    for line in lines:
        if line.startswith("ISOLATED SPREAD") or line.startswith("Structural Calc"):
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(14)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("─" * 10) or line.startswith("━" * 10):
            p = doc.add_paragraph(line)
            p.runs[0].font.size = Pt(8)
        elif line and line[0].isdigit() and "." in line[:3]:
            p = doc.add_paragraph(line)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(line)

    # Embed figures if provided
    if fig_bytes_list:
        doc.add_paragraph("")
        doc.add_paragraph("FOOTING SKETCHES").runs[0].bold = True
        for fig_bytes in fig_bytes_list:
            if fig_bytes:
                doc.add_picture(io.BytesIO(fig_bytes), width=Cm(15))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATOR
# ─────────────────────────────────────────────────────────────────────────────

def _safe_str(text: str) -> str:
    """
    Replace common Unicode symbols with ASCII equivalents
    so they render safely on any PDF font.
    """
    replacements = {
        "\u2014": "--",     # em dash —
        "\u2013": "-",      # en dash –
        "\u00d7": "x",      # × multiplication
        "\u00b2": "2",      # superscript 2
        "\u00b3": "3",      # superscript 3
        "\u2019": "'",      # right single quote
        "\u2018": "'",      # left single quote
        "\u201c": '"',      # left double quote
        "\u201d": '"',      # right double quote
        "\u2265": ">=",     # ≥
        "\u2264": "<=",     # ≤
        "\u2260": "!=",     # ≠
        "\u221a": "sqrt",   # √
        "\u03c6": "phi",    # φ
        "\u03b1": "alpha",  # α
        "\u03b2": "beta",   # β
        "\u03b3": "gamma",  # γ
        "\u03bb": "lambda", # λ
        "\u03bd": "nu",     # ν
        "\u03c3": "sigma",  # σ
        "\u03a3": "Sum",    # Σ
        "\u2211": "Sum",    # ∑
        "\u00b7": ".",      # middle dot ·
        "\u2022": "-",      # bullet •
        "\u25b6": ">",      # triangle ▶
        "\u2713": "OK",     # ✓
        "\u2717": "FAIL",   # ✗
        "\u2714": "OK",     # ✔
        "\u2716": "FAIL",   # ✖
        # Box drawing characters → ASCII
        "\u2500": "-",      # ─
        "\u2501": "=",      # ━
        "\u2502": "|",      # │
        "\u250c": "+",      # ┌
        "\u2510": "+",      # ┐
        "\u2514": "+",      # └
        "\u2518": "+",      # ┘
        "\u251c": "+",      # ├
        "\u2524": "+",      # ┤
        "\u252c": "+",      # ┬
        "\u2534": "+",      # ┴
        "\u253c": "+",      # ┼
        "\u2588": "#",      # █
        "\u00e9": "e",      # é
        "\u00e8": "e",      # è
        "\u00b0": "deg",    # °
        "\u2248": "~=",     # ≈
        "\u00b1": "+/-",    # ±
        "\u00f8": "o",      # ø (diameter symbol sometimes)
        "\u00d8": "O",      # Ø
    }
    for uni, asc in replacements.items():
        text = text.replace(uni, asc)
    # Final fallback: encode to latin-1, replace unknowns
    return text.encode("latin-1", errors="replace").decode("latin-1")


class FootingPDF(FPDF if PDF_OK else object):
    def header(self):
        self.set_font("Courier", "B", 9)
        self.cell(0, 5,
                  "ISOLATED SPREAD FOOTING DESIGN -- STRUCTURAL CALCULATION",
                  align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Courier", "", 7)
        self.cell(0, 4,
                  "Ref: SNI 2847:2019 | SNI 8460:2017 | SNI 1726:2019",
                  align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self):
        self.set_y(-12)
        self.set_font("Courier", "I", 7)
        self.cell(0, 5, f"Page {self.page_no()}", align="C")


def generate_pdf(lines: list, fig_bytes_list: list = None) -> bytes:
    """
    Generate PDF from report lines. Returns bytes.
    All Unicode special characters are converted to ASCII equivalents
    to ensure compatibility with the built-in Courier font in fpdf2.
    """
    if not PDF_OK:
        raise ImportError("fpdf2 not installed")

    pdf = FootingPDF(orientation="P", unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_left_margin(15)
    pdf.set_right_margin(12)

    # Flatten: split any embedded newlines into separate lines
    flat_lines = []
    for raw_line in lines:
        for sub in raw_line.split("\n"):
            flat_lines.append(sub)

    # Usable page width (explicit — avoids fpdf2 multi_cell(0,...) bug)
    page_w = pdf.w - pdf.l_margin - pdf.r_margin

    for raw_line in flat_lines:
        # Convert all Unicode → ASCII-safe for latin-1 font
        line = _safe_str(raw_line)

        is_title   = raw_line.startswith("ISOLATED SPREAD") or raw_line.startswith("Structural Calc")
        is_divider = raw_line.startswith("─" * 5) or raw_line.startswith("━" * 5)
        is_section = (bool(raw_line) and raw_line[0].isdigit()
                      and len(raw_line) > 2 and raw_line[1] == ".")

        if is_title:
            pdf.set_font("Courier", "B", 12)
            pdf.multi_cell(page_w, 7, line, align="C")
            pdf.set_font("Courier", "", 8)
        elif is_divider:
            pdf.set_font("Courier", "", 7)
            pdf.multi_cell(page_w, 4, line[:90])
            pdf.set_font("Courier", "", 8)
        elif is_section:
            pdf.set_font("Courier", "B", 9)
            pdf.multi_cell(page_w, 5, line)
            pdf.set_font("Courier", "", 8)
        else:
            pdf.set_font("Courier", "", 8)
            pdf.multi_cell(page_w, 4.5, line)

    if fig_bytes_list:
        for fig_bytes in fig_bytes_list:
            if fig_bytes:
                pdf.add_page()
                pdf.set_font("Courier", "B", 9)
                pdf.cell(0, 6, "FOOTING SKETCH", align="C",
                         new_x="LMARGIN", new_y="NEXT")
                with io.BytesIO(fig_bytes) as f:
                    pdf.image(f, x=15, w=180)

    return bytes(pdf.output())
