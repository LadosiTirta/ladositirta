# utils/report_generator.py
# Generator laporan Word (.docx) dan PDF yang identik
#
# Strategi:
#   1. Buat laporan Word lengkap via python-docx
#   2. Konversi ke PDF via LibreOffice headless (hasil identik dengan Word)
#   3. Fallback PDF: fpdf2 jika LibreOffice tidak tersedia
#
# Library: python-docx, matplotlib (embed grafik), subprocess (LibreOffice)

import io
import os
import subprocess
import tempfile
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==============================================================
# KONSTANTA DESAIN LAPORAN
# ==============================================================
WARNA_HEADER      = RGBColor(0x1A, 0x53, 0x76)   # biru gelap
WARNA_BARIS_GENAP = RGBColor(0xEB, 0xF5, 0xFB)   # biru sangat muda
WARNA_BARIS_GANJIL= RGBColor(0xFF, 0xFF, 0xFF)   # putih
FONT_UTAMA        = "Times New Roman"
FONT_HEADING      = "Arial"
UKURAN_BODY       = Pt(11)
UKURAN_HEADING1   = Pt(14)
UKURAN_HEADING2   = Pt(12)
MARGIN_CM         = Cm(2.5)


# ==============================================================
# HELPER FUNGSI DOCX
# ==============================================================

def _set_warna_sel(sel, rgb: RGBColor):
    """Mengatur warna latar belakang sel tabel."""
    tc   = sel._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    r, g, b = rgb[0], rgb[1], rgb[2]
    shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


def _set_border_sel(sel, sisi: list[str] = None, ukuran: int = 6, warna: str = "AAAAAA"):
    """Mengatur border sel tabel."""
    if sisi is None:
        sisi = ["top", "left", "bottom", "right"]
    tc   = sel._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for s in sisi:
        border = OxmlElement(f"w:{s}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(ukuran))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), warna)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _tambah_heading(doc: Document, teks: str, level: int = 1) -> None:
    """Menambahkan heading dengan format konsisten."""
    p = doc.add_heading(teks, level=level)
    run = p.runs[0] if p.runs else p.add_run(teks)
    run.font.name  = FONT_HEADING
    run.font.size  = UKURAN_HEADING1 if level == 1 else UKURAN_HEADING2
    run.font.color.rgb = WARNA_HEADER
    run.font.bold  = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(6)


def _tambah_paragraf(doc: Document, teks: str, bold: bool = False,
                     ukuran: Pt = None, spasi_sesudah: Pt = Pt(4)) -> None:
    """Menambahkan paragraf teks biasa."""
    p   = doc.add_paragraph()
    run = p.add_run(teks)
    run.font.name  = FONT_UTAMA
    run.font.size  = ukuran or UKURAN_BODY
    run.font.bold  = bold
    p.paragraph_format.space_after = spasi_sesudah
    p.paragraph_format.space_before = Pt(0)
    return p


def _tambah_kode(doc: Document, baris_kode: list[str]) -> None:
    """Menambahkan blok teks langkah perhitungan dalam font monospace."""
    for baris in baris_kode:
        p   = doc.add_paragraph()
        run = p.add_run(baris if baris else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)


def _buat_tabel_header(doc: Document, kolom: list[str],
                        lebar_kolom: list[float] = None) -> object:
    """
    Membuat tabel dengan baris header berwarna.

    Parameter:
        kolom        : list nama kolom header
        lebar_kolom  : list lebar kolom dalam cm (opsional)

    Mengembalikan objek tabel docx.
    """
    tabel = doc.add_table(rows=1, cols=len(kolom))
    tabel.style = "Table Grid"
    tabel.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Atur lebar kolom
    if lebar_kolom:
        for i, (sel, lebar) in enumerate(zip(tabel.columns, lebar_kolom)):
            for cell in sel.cells:
                cell.width = Cm(lebar)

    # Header row
    baris_h = tabel.rows[0]
    baris_h.height = Cm(0.8)
    for i, (sel, nama) in enumerate(zip(baris_h.cells, kolom)):
        _set_warna_sel(sel, WARNA_HEADER)
        _set_border_sel(sel, warna="1A5376")
        p   = sel.paragraphs[0]
        run = p.add_run(nama)
        run.font.name   = FONT_HEADING
        run.font.size   = Pt(10)
        run.font.bold   = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return tabel


def _tambah_baris_tabel(tabel, data: list[str], rata_kanan: list[int] = None,
                         nomor_baris: int = 0) -> None:
    """
    Menambahkan baris data ke tabel dengan warna alternating.

    Parameter:
        data        : list nilai per kolom
        rata_kanan  : list indeks kolom yang rata kanan (angka)
        nomor_baris : indeks baris (untuk warna alternating)
    """
    if rata_kanan is None:
        rata_kanan = []

    baris = tabel.add_row()
    warna = WARNA_BARIS_GENAP if nomor_baris % 2 == 0 else WARNA_BARIS_GANJIL

    for i, (sel, nilai) in enumerate(zip(baris.cells, data)):
        _set_warna_sel(sel, warna)
        _set_border_sel(sel, warna="CCCCCC", ukuran=4)
        p   = sel.paragraphs[0]
        run = p.add_run(str(nilai))
        run.font.name = FONT_UTAMA
        run.font.size = Pt(10)
        p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if i in rata_kanan
                       else WD_ALIGN_PARAGRAPH.LEFT)
        sel.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _embed_grafik(doc: Document, fig: plt.Figure,
                  lebar_inch: float = 6.0, dpi: int = 150) -> None:
    """Mengubah Figure matplotlib menjadi gambar dan menyisipkan ke dokumen."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=dpi, bbox_inches="tight")
    buf.seek(0)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(lebar_inch))
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    plt.close(fig)


def _tambah_garis_pemisah(doc: Document) -> None:
    """Menambahkan garis horizontal pemisah."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A5376")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def _setup_halaman(doc: Document) -> None:
    """Mengatur ukuran halaman, margin, header, dan footer."""
    section = doc.sections[0]
    # A4
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin    = MARGIN_CM
    section.bottom_margin = MARGIN_CM
    section.left_margin   = MARGIN_CM
    section.right_margin  = MARGIN_CM

    # Header
    header = section.header
    p_hdr  = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p_hdr.clear()
    run_hdr = p_hdr.add_run("LAPORAN PERHITUNGAN KAPASITAS PONDASI TIANG DALAM")
    run_hdr.font.name  = FONT_HEADING
    run_hdr.font.size  = Pt(9)
    run_hdr.font.bold  = True
    run_hdr.font.color.rgb = WARNA_HEADER
    p_hdr.alignment    = WD_ALIGN_PARAGRAPH.CENTER

    # Garis bawah header
    pPr  = p_hdr._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1A5376")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Footer dengan nomor halaman
    footer = section.footer
    p_ftr  = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_ftr.clear()
    p_ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_ftr = p_ftr.add_run("Halaman ")
    run_ftr.font.name = FONT_UTAMA
    run_ftr.font.size = Pt(9)

    # Field nomor halaman
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run_pg = p_ftr.add_run()
    run_pg._r.append(fldChar1)
    run_pg._r.append(instrText)
    run_pg._r.append(fldChar2)

    run_ftr2 = p_ftr.add_run(f"  |  Dibuat: {datetime.now().strftime('%d %B %Y')}")
    run_ftr2.font.name = FONT_UTAMA
    run_ftr2.font.size = Pt(9)
    run_ftr2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ==============================================================
# GENERATOR LAPORAN UTAMA
# ==============================================================

def buat_laporan_word(
    param_tiang: dict,
    df_tanah,
    hasil_tekan: dict,
    hasil_lateral: dict | None = None,
    metode_lateral: str = "Broms",
    nama_proyek: str = "Proyek Pondasi",
    nama_konsultan: str = "",
    nomor_laporan: str = "LAP-001",
) -> io.BytesIO:
    """
    Membuat laporan Word lengkap berisi:
    1. Cover / Header proyek
    2. Data tiang dan kondisi tanah
    3. Profil data SPT per lapisan
    4. Langkah perhitungan daya dukung tekan & tarik (step-by-step)
    5. Tabel distribusi skin friction per lapisan
    6. Ringkasan kapasitas tekan & tarik (dengan SF)
    7. Langkah perhitungan gaya lateral
    8. Grafik profil tanah, distribusi, variasi kedalaman
    9. Grafik gaya lateral (jika ada)
    10. Footer disclaimer & referensi

    Mengembalikan BytesIO siap di-download.
    """
    from utils.grapher import (
        buat_grafik_profil,
        buat_grafik_distribusi_skin,
        buat_grafik_variasi_kedalaman,
    )
    from utils.grapher_lateral import buat_grafik_broms, buat_grafik_py
    from calculations.bearing_capacity import hitung_variasi_kedalaman
    from utils.input_handler import KOLOM_TANAH

    doc = Document()
    _setup_halaman(doc)

    # ===========================================================
    # BAGIAN 1: COVER / INFO PROYEK
    # ===========================================================
    p_judul = doc.add_paragraph()
    p_judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_judul.add_run("LAPORAN PERHITUNGAN")
    r.font.name = FONT_HEADING; r.font.size = Pt(18); r.font.bold = True
    r.font.color.rgb = WARNA_HEADER

    p_judul2 = doc.add_paragraph()
    p_judul2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_judul2.add_run("KAPASITAS PONDASI TIANG DALAM")
    r2.font.name = FONT_HEADING; r2.font.size = Pt(16); r2.font.bold = True
    r2.font.color.rgb = WARNA_HEADER

    _tambah_garis_pemisah(doc)

    # Tabel info proyek
    tbl_info = doc.add_table(rows=5, cols=2)
    tbl_info.style = "Table Grid"
    info_data = [
        ("Nama Proyek",     nama_proyek),
        ("No. Laporan",     nomor_laporan),
        ("Tanggal Hitung",  datetime.now().strftime("%d %B %Y")),
        ("Konsultan/Dibuat",nama_konsultan or "—"),
        ("Acuan Standar",   "SNI 8460:2017 · Meyerhof (1976) · Tomlinson (2008) · Broms (1964) · API RP 2GEO"),
    ]
    for i, (label, nilai) in enumerate(info_data):
        sel_l = tbl_info.rows[i].cells[0]
        sel_r = tbl_info.rows[i].cells[1]
        sel_l.width = Cm(5)
        sel_r.width = Cm(13)
        _set_warna_sel(sel_l, RGBColor(0xEB, 0xF5, 0xFB))
        _set_border_sel(sel_l); _set_border_sel(sel_r)
        p_l = sel_l.paragraphs[0]; r_l = p_l.add_run(label)
        r_l.font.name = FONT_HEADING; r_l.font.size = Pt(10); r_l.font.bold = True
        p_r = sel_r.paragraphs[0]; r_r = p_r.add_run(nilai)
        r_r.font.name = FONT_UTAMA; r_r.font.size = Pt(10)

    doc.add_paragraph()

    # ===========================================================
    # BAGIAN 2: DATA TIANG
    # ===========================================================
    _tambah_heading(doc, "1. Data Tiang", level=1)
    _tambah_garis_pemisah(doc)

    tbl_tiang = _buat_tabel_header(doc,
        ["Parameter", "Nilai", "Satuan"],
        lebar_kolom=[6, 6, 4]
    )
    data_tiang_rows = [
        ("Tipe tiang",                param_tiang["tipe"],                "—"),
        ("Dimensi",                   param_tiang["dim_label"],           "—"),
        ("Kedalaman tiang (L)",       f"{param_tiang['kedalaman']:.2f}", "m"),
        ("Luas ujung (Ab)",           f"{param_tiang['area_ujung']:.4f}","m²"),
        ("Keliling (p)",              f"{param_tiang['keliling']:.4f}",  "m"),
        ("Muka air tanah (MAT)",      f"{param_tiang['muka_air']:.2f}",  "m"),
        ("Mutu material",             param_tiang["material"],            "—"),
        ("Safety factor tekan (SF)",  f"{param_tiang['sf_tekan']:.1f}",  "—"),
        ("Safety factor tarik (SF)",  f"{param_tiang['sf_tarik']:.1f}",  "—"),
    ]
    for i, baris in enumerate(data_tiang_rows):
        _tambah_baris_tabel(tbl_tiang, list(baris), rata_kanan=[1], nomor_baris=i)

    doc.add_paragraph()

    # ===========================================================
    # BAGIAN 3: DATA SPT PER LAPISAN
    # ===========================================================
    _tambah_heading(doc, "2. Profil Data Tanah (SPT)", level=1)
    _tambah_garis_pemisah(doc)

    tbl_spt = _buat_tabel_header(doc,
        ["No.", "Jenis Tanah", "z atas\n(m)", "z bawah\n(m)",
         "Tebal\n(m)", "SPT-N", "Cu\n(kPa)", "φ\n(°)", "γ\n(kN/m³)"],
        lebar_kolom=[1.0, 4.5, 1.5, 1.5, 1.5, 1.2, 1.5, 1.2, 1.8]
    )
    for i, (_, baris) in enumerate(df_tanah.iterrows()):
        _tambah_baris_tabel(tbl_spt, [
            str(i + 1),
            str(baris[KOLOM_TANAH["jenis"]]),
            f"{baris[KOLOM_TANAH['z_atas']]:.2f}",
            f"{baris[KOLOM_TANAH['z_bawah']]:.2f}",
            f"{baris[KOLOM_TANAH['z_bawah']] - baris[KOLOM_TANAH['z_atas']]:.2f}",
            f"{int(baris[KOLOM_TANAH['spt']])}",
            f"{baris[KOLOM_TANAH['cu']]:.1f}",
            f"{baris[KOLOM_TANAH['phi']]:.1f}",
            f"{baris[KOLOM_TANAH['gamma']]:.1f}",
        ], rata_kanan=[0, 2, 3, 4, 5, 6, 7, 8], nomor_baris=i)

    doc.add_paragraph()

    # ===========================================================
    # BAGIAN 4: LANGKAH PERHITUNGAN DAYA DUKUNG TEKAN
    # ===========================================================
    _tambah_heading(doc, "3. Perhitungan Daya Dukung Tekan", level=1)
    _tambah_garis_pemisah(doc)

    _tambah_heading(doc, "3.1 End Bearing (Qpoint)", level=2)
    _tambah_kode(doc, hasil_tekan["langkah_qpoint"])
    doc.add_paragraph()

    _tambah_heading(doc, "3.2 Skin Friction per Lapisan (Qskin)", level=2)
    _tambah_kode(doc, hasil_tekan["langkah_qskin"])
    doc.add_paragraph()

    _tambah_heading(doc, "3.3 Kapasitas Total Tekan & Tarik", level=2)
    _tambah_kode(doc, hasil_tekan["langkah_total"])
    doc.add_paragraph()

    _tambah_heading(doc, "3.4 Kapasitas Struktur Tiang", level=2)
    _tambah_kode(doc, hasil_tekan["langkah_struktur"])
    doc.add_paragraph()

    # ===========================================================
    # BAGIAN 5: TABEL DISTRIBUSI SKIN FRICTION
    # ===========================================================
    _tambah_heading(doc, "4. Distribusi Daya Dukung Selimut per Lapisan", level=1)
    _tambah_garis_pemisah(doc)

    tbl_skin = _buat_tabel_header(doc,
        ["No.", "Jenis Tanah", "z\natas(m)", "z\nbawah(m)",
         "Tebal\n(m)", "Metode", "σ'v\n(kPa)", "α / β", "fs\n(kPa)", "Qs\n(kN)"],
        lebar_kolom=[0.8, 3.5, 1.2, 1.2, 1.2, 2.8, 1.3, 1.0, 1.3, 1.5]
    )
    for i, lap in enumerate(hasil_tekan["detail_lapisan"]):
        alfa_beta = (f"{lap['alpha']:.2f}" if lap["kategori"] == "lempung"
                     else f"β={lap['beta']:.4f}")
        _tambah_baris_tabel(tbl_skin, [
            str(lap["no"]),
            lap["jenis"][:22],
            f"{lap['z_atas']:.2f}",
            f"{lap['z_bawah']:.2f}",
            f"{lap['tebal']:.2f}",
            lap["metode"].split("(")[0].strip(),
            f"{lap['sigma_v_eff']:.2f}",
            alfa_beta,
            f"{lap['fs']:.3f}",
            f"{lap['Qs']:.2f}",
        ], rata_kanan=[0, 2, 3, 4, 6, 7, 8, 9], nomor_baris=i)

    doc.add_paragraph()
    _tambah_paragraf(doc,
        f"Total ΣQskin = {hasil_tekan['Qskin']:.2f} kN  |  "
        f"Qpoint = {hasil_tekan['Qpoint']:.2f} kN  |  "
        f"Qultimate tekan = {hasil_tekan['Qult_tekan']:.2f} kN",
        bold=True
    )

    # ===========================================================
    # BAGIAN 6: RINGKASAN KAPASITAS
    # ===========================================================
    _tambah_heading(doc, "5. Ringkasan Daya Dukung", level=1)
    _tambah_garis_pemisah(doc)

    tbl_resume = _buat_tabel_header(doc,
        ["Komponen", "Ultimit (kN)", "SF", "Ijin (kN)", "Keterangan"],
        lebar_kolom=[5.0, 3.0, 1.5, 3.0, 5.0]
    )
    resume_rows = [
        ("End Bearing (Qpoint)",
         f"{hasil_tekan['Qpoint']:.2f}", "—", "—", "Meyerhof (1976)"),
        ("Skin Friction (ΣQskin)",
         f"{hasil_tekan['Qskin']:.2f}", "—", "—", "α-method / β-method"),
        ("Daya Dukung Tekan",
         f"{hasil_tekan['Qult_tekan']:.2f}",
         f"{hasil_tekan['sf_tekan']:.1f}",
         f"{hasil_tekan['Qijin_tekan']:.2f}",
         "Qijin = Qult / SF"),
        ("Daya Dukung Tarik",
         f"{hasil_tekan['Qult_tarik']:.2f}",
         f"{hasil_tekan['sf_tarik']:.1f}",
         f"{hasil_tekan['Qijin_tarik']:.2f}",
         "Qijin_tarik = ΣQskin×fr / SF"),
    ]
    if hasil_tekan.get("Pn_struktur"):
        resume_rows.append((
            "Kapasitas Struktur (Pn)",
            f"{hasil_tekan['Pn_struktur']:.2f}",
            "—", "—",
            "SNI 2847:2019 (φ×0.85×fc'×Ag)"
        ))
    for i, baris in enumerate(resume_rows):
        _tambah_baris_tabel(tbl_resume, list(baris),
                             rata_kanan=[1, 2, 3], nomor_baris=i)

    doc.add_paragraph()

    # ===========================================================
    # BAGIAN 7: GAYA LATERAL (jika tersedia)
    # ===========================================================
    if hasil_lateral:
        _tambah_heading(doc, "6. Perhitungan Gaya Lateral", level=1)
        _tambah_garis_pemisah(doc)
        _tambah_paragraf(doc, f"Metode: {metode_lateral}", bold=True)
        doc.add_paragraph()

        if metode_lateral == "Broms (1964) — solusi cepat":
            _tambah_kode(doc, hasil_lateral.get("langkah", []))
            doc.add_paragraph()

            # Tabel ringkasan Broms
            tbl_lat = _buat_tabel_header(doc,
                ["Parameter", "Nilai", "Satuan"],
                lebar_kolom=[7, 5, 4]
            )
            lat_rows = [
                ("Kapasitas lateral ultimit (Hu)", f"{hasil_lateral['Hu']:.2f}", "kN"),
                ("Kapasitas lateral ijin (Hijin)", f"{hasil_lateral['Hijin']:.2f}", "kN"),
                ("Momen maksimum (Mmax)",          f"{hasil_lateral['Mmax']:.2f}", "kN·m"),
                ("Defleksi kepala tiang (y₀)",      f"{hasil_lateral['defleksi_mm']:.2f}", "mm"),
                ("Kontrol (H ≤ Hijin)",
                 "OK ✓" if hasil_lateral.get("kontrol_ok") else "TIDAK OK ✗", "—"),
            ]
            for i, r in enumerate(lat_rows):
                _tambah_baris_tabel(tbl_lat, list(r), rata_kanan=[1], nomor_baris=i)

        else:  # P-Y Curve
            _tambah_kode(doc, hasil_lateral.get("langkah", []))
            doc.add_paragraph()

            tbl_lat = _buat_tabel_header(doc,
                ["Parameter", "Nilai", "Satuan"],
                lebar_kolom=[7, 5, 4]
            )
            lat_rows = [
                ("Defleksi kepala tiang (y₀)",  f"{hasil_lateral['y0_mm']:.2f}", "mm"),
                ("Momen maksimum (Mmax)",        f"{hasil_lateral['Mmax']:.2f}", "kN·m"),
                ("z titik Mmax",                 f"{hasil_lateral['z_Mmax']:.2f}", "m"),
                ("Gaya geser maks (Vmax)",       f"{hasil_lateral['Vmax']:.2f}", "kN"),
                ("Status konvergensi",
                 f"Konvergen ({hasil_lateral['iterasi']} iterasi)" if hasil_lateral["konvergen"]
                 else "Belum konvergen", "—"),
            ]
            for i, r in enumerate(lat_rows):
                _tambah_baris_tabel(tbl_lat, list(r), rata_kanan=[1], nomor_baris=i)

        doc.add_paragraph()

    # ===========================================================
    # BAGIAN 8: GRAFIK
    # ===========================================================
    _tambah_heading(doc, "7. Grafik dan Visualisasi", level=1)
    _tambah_garis_pemisah(doc)

    # Grafik profil tanah
    _tambah_heading(doc, "7.1 Profil Tanah", level=2)
    fig_profil = buat_grafik_profil(
        hasil_tekan["semua_lapisan"],
        param_tiang["kedalaman"],
        param_tiang["muka_air"],
        param_tiang["diameter"]
    )
    _embed_grafik(doc, fig_profil, lebar_inch=5.0)

    # Grafik distribusi skin friction
    _tambah_heading(doc, "7.2 Distribusi Daya Dukung Selimut", level=2)
    fig_dist = buat_grafik_distribusi_skin(
        hasil_tekan["detail_lapisan"],
        hasil_tekan["Qpoint"],
        param_tiang["kedalaman"]
    )
    _embed_grafik(doc, fig_dist, lebar_inch=6.0)

    # Grafik variasi kedalaman
    _tambah_heading(doc, "7.3 Kapasitas Tiang vs Variasi Kedalaman", level=2)
    hasil_var = hitung_variasi_kedalaman(
        df_tanah, param_tiang,
        z_min=max(param_tiang["kedalaman"] * 0.3, 3.0),
        z_max=param_tiang["kedalaman"]
    )
    if hasil_var:
        fig_var = buat_grafik_variasi_kedalaman(hasil_var)
        _embed_grafik(doc, fig_var, lebar_inch=6.0)

    # Grafik lateral
    if hasil_lateral:
        _tambah_heading(doc, "7.4 Gaya Lateral", level=2)
        if "Broms" in metode_lateral:
            hasil_lateral["H_input"] = hasil_lateral.get("H_input", 0)
            hasil_lateral["L"] = param_tiang["kedalaman"]
            hasil_lateral["D"] = param_tiang["diameter"]
            fig_lat = buat_grafik_broms(hasil_lateral)
            _embed_grafik(doc, fig_lat, lebar_inch=6.0)
        else:
            from calculations.bearing_capacity import hitung_variasi_kedalaman as _
            fig_py  = buat_grafik_py(hasil_lateral, param_tiang, df_tanah)
            _embed_grafik(doc, fig_py, lebar_inch=6.5)

    # ===========================================================
    # BAGIAN 9: DISCLAIMER & REFERENSI
    # ===========================================================
    doc.add_page_break()
    _tambah_heading(doc, "8. Referensi dan Catatan", level=1)
    _tambah_garis_pemisah(doc)

    refs = [
        "SNI 8460:2017 — Persyaratan Perancangan Geoteknik.",
        "Meyerhof, G.G. (1976). Bearing capacity and settlement of pile foundations. "
        "ASCE J. Geotech. Eng. Div., 102(3), 197–228.",
        "Tomlinson, M.J. (2008). Pile Design and Construction Practice, 5th Ed. Taylor & Francis.",
        "Broms, B.B. (1964). Lateral resistance of piles in cohesive soils. "
        "ASCE J. Soil Mech. Found. Div., 90(SM2), 27–63.",
        "Broms, B.B. (1964). Lateral resistance of piles in cohesionless soils. "
        "ASCE J. Soil Mech. Found. Div., 90(SM3), 123–156.",
        "Matlock, H. (1970). Correlations for design of laterally loaded piles in soft clay. "
        "OTC Paper 1204, Houston.",
        "Reese, L.C., Cox, W.R., Koop, F.D. (1974). Analysis of laterally loaded piles in sand. "
        "OTC Paper 2080, Houston.",
        "API RP 2GEO (2011). Geotechnical and Foundation Design Considerations.",
    ]
    for ref in refs:
        p_ref = doc.add_paragraph(style="List Bullet")
        run_ref = p_ref.add_run(ref)
        run_ref.font.name = FONT_UTAMA
        run_ref.font.size = Pt(10)

    doc.add_paragraph()
    _tambah_garis_pemisah(doc)
    _tambah_paragraf(doc,
        "DISCLAIMER: Laporan ini dibuat berdasarkan data yang diinput pengguna dan metode "
        "empiris yang berlaku. Hasil perhitungan bersifat referensi dan harus diverifikasi "
        "oleh insinyur geoteknik yang bertanggung jawab sebelum digunakan untuk konstruksi.",
        ukuran=Pt(9)
    )

    # Simpan ke BytesIO
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ==============================================================
# KONVERSI WORD → PDF via LibreOffice
# ==============================================================

def konversi_ke_pdf(docx_bytes: io.BytesIO) -> io.BytesIO | None:
    """
    Mengkonversi dokumen Word (BytesIO) ke PDF menggunakan LibreOffice headless.
    Hasil PDF identik dengan Word karena menggunakan renderer yang sama.

    Mengembalikan BytesIO PDF, atau None jika gagal.
    """
    with tempfile.TemporaryDirectory() as tmpdir:
        # Tulis DOCX sementara
        path_docx = os.path.join(tmpdir, "laporan_tiang.docx")
        with open(path_docx, "wb") as f:
            f.write(docx_bytes.getvalue())

        # Jalankan LibreOffice headless
        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    "--outdir", tmpdir,
                    path_docx,
                ],
                capture_output=True, text=True, timeout=60
            )
        except (subprocess.TimeoutExpired, FileNotFoundError):
            return None

        if result.returncode != 0:
            return None

        # Baca PDF hasil
        path_pdf = os.path.join(tmpdir, "laporan_tiang.pdf")
        if not os.path.exists(path_pdf):
            return None

        with open(path_pdf, "rb") as f:
            pdf_bytes = f.read()

    buf_pdf = io.BytesIO(pdf_bytes)
    buf_pdf.seek(0)
    return buf_pdf
