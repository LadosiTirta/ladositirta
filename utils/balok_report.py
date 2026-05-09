"""
=============================================================
utils/balok_report.py
Generator laporan Word & PDF untuk evaluasi balok beton bertulang.
Referensi : SNI 2847:2019 (ACI 318-14)

TIDAK mengandung: streamlit
Import dari    : utils/balok_calc (sp, konstanta)
Dipanggil oleh : pages/1_Lentur_Balok.py
=============================================================
"""

import io
import datetime
import math

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from utils.balok_calc import sp, luas_batang


# ============================================================
# KONSTANTA WARNA PDF
# ============================================================
WATERMARK_TEXT = "DIHASILKAN OLEH: LADOSI ENGINEERING"
BRAND_COLOR    = (26, 60, 94)
OK_COLOR       = (27, 94, 32)
FAIL_COLOR     = (183, 28, 28)
GRAY           = (120, 120, 120)


# ============================================================
# DICTIONARY BAHASA (ID & EN)
# ============================================================
TEXT = {
    "ID": {
        "lang_select": "Pilih Bahasa",
        "title": "Evaluasi Kapasitas Lentur, Geser & Torsi Balok Beton Bertulang",
        "subtitle": "Tulangan Rangkap | Strain Compatibility | Referensi: SNI 2847:2019 (setara ACI 318-14)",
        "data_input": "Data Input",
        # --- Project Info ---
        "proj_info":    "Informasi Proyek",
        "proj_name":    "Nama Proyek",
        "proj_val":     "Laporan Analisa Struktur",
        "proj_lokasi":  "Lokasi Proyek",
        "proj_eng":     "Nama Engineer",
        "proj_nodok":   "No. Dokumen",
        "proj_tgl":     "Tanggal",
        "proj_catatan": "Catatan (opsional)",
        # --- Sistem Seismik ---
        "sistem_rangka": "Sistem Rangka",
        "sistem_help":   "Pilih sesuai kategori desain seismik proyek (SNI 1726:2019)",
        "sistem_list":   [
            "Biasa (Non-Seismik)",
            "SRPMB — Sistem Rangka Pemikul Momen Biasa",
            "SRPMM — Sistem Rangka Pemikul Momen Menengah  [SNI 2847 Ps. 18.4]",
            "SRPMK — Sistem Rangka Pemikul Momen Khusus   [SNI 2847 Ps. 18.6]",
        ],
        # --- Gaya Dalam ---
        "gaya_dalam": "Gaya Dalam",
        "mu": "Mu - Momen rencana (kN.m)",
        "vu": "Vu - Gaya geser ultimit (kN)",
        "tu": "Tu - Momen torsi ultimit (kN.m)",
        "tu_help": "Isi 0 jika tidak ada torsi. Evaluasi otomatis berjalan.",
        # --- Material ---
        "material": "Material",
        "fc": "f'c - Kuat tekan beton (MPa)",
        "fy": "fy - Kuat leleh tul. lentur (MPa)",
        "fyt": "fyt - Kuat leleh sengkang (MPa)",
        "fyt_help": "Umumnya 240 MPa untuk BJTP, atau samakan dengan fy bila pakai BJTS",
        # --- Geometri ---
        "geometri": "Geometri",
        "b": "b (mm) - Lebar",
        "h": "h (mm) - Tinggi",
        "cc": "cc (mm) - Selimut bersih",
        "ds": "ds (mm) - Dia. sengkang",
        # --- Sengkang ---
        "sengkang": "Sengkang",
        "s": "s - Jarak sengkang (mm)",
        "nkaki": "Jumlah kaki (n_kaki)",
        "nkaki_help": "Default 2 kaki. Untuk balok lebar bisa 4 kaki.",
        # --- Torsi ---
        "torsi_add": "Input Torsi Tambahan",
        "tipe_torsi": "Tipe Torsi",
        "tipe_torsi_help": "Equilibrium: tidak boleh redistribusi. Compatibility: boleh direduksi.",
        "db_long": "Dia. tul. long. torsi (mm)",
        "db_long_help": "Diameter tulangan longitudinal tambahan untuk menahan torsi.",
        # --- Tulangan ---
        "tarik": "Tulangan TARIK (bawah)",
        "tarik_cap": "Lapis 1 = paling bawah, Lapis 2 = di atas Lapis 1",
        "tekan": "Tulangan TEKAN (atas)",
        "tekan_cap": "Lapis 1 = paling atas, Lapis 2 = di bawah Lapis 1",
        "jml": "Jumlah",
        "dia": "Diameter (mm)",
        "lapis": "Lapis",
        # --- Tombol & pesan ---
        "btn_hitung": "HITUNG EVALUASI LENTUR, GESER & TORSI",
        "err_tarik": "Harus ada minimal 1 lapis tulangan tarik!",
        "err_mu": "Mu harus lebih besar dari 0!",
        "err_vu": "Vu tidak boleh negatif!",
        "err_tu": "Tu tidak boleh negatif!",
        "err_d": "d-aktual ({d:.1f} mm) tidak valid. Cek selimut, sengkang, dan diameter tulangan.",
        "warn_ubah": "Perhatian: Data input telah diubah. Hasil di bawah masih menggunakan data sebelumnya. Klik HITUNG kembali.",
        # --- Hasil ---
        "res_lentur": "Hasil Utama - Lentur",
        "res_geser":  "Hasil Utama - Geser",
        "res_torsi":  "Hasil Utama - Torsi",
        "ok_lentur":   "[OK] LENTUR AMAN - D/C = {dc:.3f} <= 1.0 | Tension-controlled",
        "warn_lentur": "[!] LENTUR AMAN secara D/C ({dc:.3f}) namun perlu perhatian: {note}",
        "fail_lentur": "[X] LENTUR TIDAK AMAN - D/C = {dc:.3f} > 1.0",
        "ok_geser":   "[OK] GESER AMAN - D/C = {dc:.3f} <= 1.0",
        "warn_geser": "[!] GESER kapasitas terpenuhi (D/C = {dc:.3f}) namun: {note}",
        "fail_geser": "[X] GESER TIDAK AMAN - D/C = {dc:.3f} > 1.0",
        "ok_torsi":   "[OK] TORSI AMAN - D/C = {dc:.3f} <= 1.0",
        "ign_torsi":  "[DIABAIKAN] TORSI - Tu <= Phi_t x Tth (Efek torsi diabaikan)",
        "fail_torsi": "[X] TORSI TIDAK AMAN - {note}",
        "stat_gabungan": "STATUS GABUNGAN",
        "aman":      "AMAN",
        "tidak_aman":"TIDAK AMAN",
        "dl_word": "Download Laporan Word (.docx)",
        "dl_pdf":  "Download Laporan PDF (.pdf)",
        # --- Label laporan (Word & PDF) ---
        "lap_judul":    "LAPORAN PERHITUNGAN STRUKTUR",
        "lap_subjudul": "Evaluasi Kapasitas Lentur, Geser & Torsi Balok Beton Bertulang",
        "lap_ref":      "Referensi: SNI 2847:2019 (ACI 318-14)",
        "lap_A": "A.  DATA INPUT PENAMPANG",
        "lap_B": "B.  ANALISA PERHITUNGAN - LENTUR",
        "lap_C": "C.  RANGKUMAN HASIL",
        "lap_D": "D.  KESIMPULAN",
        "lap_E": "E.  ANALISA PERHITUNGAN - GESER",
        "lap_F": "F.  ANALISA PERHITUNGAN - TORSI",
        "lap_material":  "Material:",
        "lap_geometri":  "Geometri:",
        "lap_gayadalm":  "Gaya Dalam:",
        "lap_tul_tarik": "Tulangan TARIK (bawah):",
        "lap_tul_tekan": "Tulangan TEKAN (atas):",
        "lap_visualisasi": "Visualisasi Penampang :",
        "lap_tidak_ada_tekan": "(Tidak ada tulangan tekan)",
        "lap_kes_lentur": "KESIMPULAN LENTUR",
        "lap_kes_geser":  "KESIMPULAN GESER",
        "lap_kes_torsi":  "KESIMPULAN TORSI",
    },
    "EN": {
        "lang_select": "Select Language",
        "title": "Flexural, Shear & Torsion Capacity Evaluation for Concrete Beams",
        "subtitle": "Double Reinforcement | Strain Compatibility | Ref: ACI 318-14 (SNI 2847:2019)",
        "data_input": "Input Data",
        # --- Project Info ---
        "proj_info":    "Project Information",
        "proj_name":    "Project Name",
        "proj_val":     "Structural Analysis Report",
        "proj_lokasi":  "Project Location",
        "proj_eng":     "Engineer Name",
        "proj_nodok":   "Document No.",
        "proj_tgl":     "Date",
        "proj_catatan": "Notes (optional)",
        # --- Seismic System ---
        "sistem_rangka": "Frame System",
        "sistem_help":   "Select based on project seismic design category (SNI 1726:2019 / ASCE 7)",
        "sistem_list":   [
            "Ordinary (Non-Seismic)",
            "OMRF — Ordinary Moment Resisting Frame",
            "IMRF — Intermediate Moment Resisting Frame  [ACI 318 Sec. 18.4]",
            "SMRF — Special Moment Resisting Frame        [ACI 318 Sec. 18.6]",
        ],
        # --- Internal Forces ---
        "gaya_dalam": "Internal Forces",
        "mu": "Mu - Design Moment (kN.m)",
        "vu": "Vu - Ultimate Shear (kN)",
        "tu": "Tu - Ultimate Torsion (kN.m)",
        "tu_help": "Enter 0 if no torsion. Evaluation will run automatically.",
        # --- Material ---
        "material": "Material",
        "fc": "f'c - Concrete Compressive Strength (MPa)",
        "fy": "fy - Flexural Rebar Yield Strength (MPa)",
        "fyt": "fyt - Stirrup Yield Strength (MPa)",
        "fyt_help": "Usually 240 MPa for plain bars, or match fy for deformed bars.",
        # --- Geometry ---
        "geometri": "Geometry",
        "b": "b (mm) - Width",
        "h": "h (mm) - Height",
        "cc": "cc (mm) - Clear Cover",
        "ds": "ds (mm) - Stirrup Diameter",
        # --- Stirrups ---
        "sengkang": "Stirrups",
        "s": "s - Stirrup Spacing (mm)",
        "nkaki": "Number of Legs (n_legs)",
        "nkaki_help": "Default 2 legs. Use 4 legs for wider beams.",
        # --- Torsion ---
        "torsi_add": "Additional Torsion Input",
        "tipe_torsi": "Torsion Type",
        "tipe_torsi_help": "Equilibrium: no redistribution allowed. Compatibility: can be reduced.",
        "db_long": "Longitudinal Torsion Rebar Dia. (mm)",
        "db_long_help": "Additional longitudinal rebar diameter to resist torsion.",
        # --- Reinforcement ---
        "tarik": "TENSION Reinforcement (Bottom)",
        "tarik_cap": "Layer 1 = Bottommost, Layer 2 = Above Layer 1",
        "tekan": "COMPRESSION Reinforcement (Top)",
        "tekan_cap": "Layer 1 = Topmost, Layer 2 = Below Layer 1",
        "jml": "Quantity",
        "dia": "Diameter (mm)",
        "lapis": "Layer",
        # --- Buttons & messages ---
        "btn_hitung": "CALCULATE FLEXURE, SHEAR & TORSION",
        "err_tarik": "At least 1 layer of tension reinforcement is required!",
        "err_mu": "Mu must be greater than 0!",
        "err_vu": "Vu cannot be negative!",
        "err_tu": "Tu cannot be negative!",
        "err_d": "Effective depth d ({d:.1f} mm) is invalid. Check cover, stirrups, and rebar dia.",
        "warn_ubah": "Warning: Input data changed. Results below are from previous data. Click CALCULATE again.",
        # --- Results ---
        "res_lentur": "Main Results - Flexure",
        "res_geser":  "Main Results - Shear",
        "res_torsi":  "Main Results - Torsion",
        "ok_lentur":   "[OK] FLEXURE SAFE - D/C = {dc:.3f} <= 1.0 | Tension-controlled",
        "warn_lentur": "[!] FLEXURE SAFE by D/C ({dc:.3f}) but needs attention: {note}",
        "fail_lentur": "[X] FLEXURE UNSAFE - D/C = {dc:.3f} > 1.0",
        "ok_geser":   "[OK] SHEAR SAFE - D/C = {dc:.3f} <= 1.0",
        "warn_geser": "[!] SHEAR capacity met (D/C = {dc:.3f}) but: {note}",
        "fail_geser": "[X] SHEAR UNSAFE - D/C = {dc:.3f} > 1.0",
        "ok_torsi":   "[OK] TORSION SAFE - D/C = {dc:.3f} <= 1.0",
        "ign_torsi":  "[IGNORED] TORSION - Tu <= Phi_t x Tth (Torsion effects ignored)",
        "fail_torsi": "[X] TORSION UNSAFE - {note}",
        "stat_gabungan": "COMBINED STATUS",
        "aman":      "SAFE",
        "tidak_aman":"UNSAFE",
        "dl_word": "Download Word Report (.docx)",
        "dl_pdf":  "Download PDF Report (.pdf)",
        # --- Report labels (Word & PDF) ---
        "lap_judul":    "STRUCTURAL CALCULATION REPORT",
        "lap_subjudul": "Flexural, Shear & Torsion Capacity — Reinforced Concrete Beam",
        "lap_ref":      "Reference: ACI 318-14 (SNI 2847:2019)",
        "lap_A": "A.  INPUT DATA",
        "lap_B": "B.  FLEXURE ANALYSIS",
        "lap_C": "C.  SUMMARY",
        "lap_D": "D.  CONCLUSION",
        "lap_E": "E.  SHEAR ANALYSIS",
        "lap_F": "F.  TORSION ANALYSIS",
        "lap_material":  "Material:",
        "lap_geometri":  "Geometry:",
        "lap_gayadalm":  "Internal Forces:",
        "lap_tul_tarik": "TENSION Reinforcement (Bottom):",
        "lap_tul_tekan": "COMPRESSION Reinforcement (Top):",
        "lap_visualisasi": "Cross-section Visualization:",
        "lap_tidak_ada_tekan": "(No compression reinforcement)",
        "lap_kes_lentur": "FLEXURE CONCLUSION",
        "lap_kes_geser":  "SHEAR CONCLUSION",
        "lap_kes_torsi":  "TORSION CONCLUSION",
    }
}


# ============================================================
# CLASS PDF
# ============================================================
class LaporanBalokPDF(FPDF):
    def __init__(self, nama_proyek):
        super().__init__()
        self.nama_proyek = sp(nama_proyek)
        self.set_margins(25, 25, 20)
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.8)
        self.line(25, 15, 190, 15)
        self.set_font("Helvetica", "B", 9); self.set_text_color(*BRAND_COLOR)
        self.set_xy(25, 17)
        self.cell(0, 5, sp("LAPORAN EVALUASI LENTUR, GESER & TORSI BALOK  |  SNI 2847:2019"),
                  ln=False, align="L")
        self.set_font("Helvetica", "", 8); self.set_text_color(*GRAY)
        self.set_xy(25, 17)
        self.cell(0, 5, sp(f"Proyek: {self.nama_proyek}"), ln=False, align="R")
        self.ln(10)

    def footer(self):
        self.set_y(-18)
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.4)
        self.line(25, self.get_y(), 190, self.get_y())
        self.set_font("Helvetica", "I", 7.5); self.set_text_color(*GRAY)
        self.cell(0, 6,
            sp("Referensi: SNI 2847:2019 | ACI 318-14 - "
               "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan."),
            align="C")
        self.set_xy(25, self.get_y())
        self.set_font("Helvetica", "", 7.5)
        self.cell(0, 6, sp(f"Halaman {self.page_no()}"), align="R")

    def watermark(self):
        self.set_font("Helvetica", "B", 28); self.set_text_color(210, 215, 220)
        xc, yc = self.w / 2, self.h / 2
        with self.rotation(40, xc, yc):
            self.set_xy(xc - 65, yc - 6)
            self.cell(130, 12, sp(WATERMARK_TEXT), align="C")
        self.set_text_color(0, 0, 0)

    def section_title(self, teks):
        self.set_font("Helvetica", "B", 11); self.set_text_color(*BRAND_COLOR)
        self.ln(4); self.cell(0, 7, sp(teks), ln=True)
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.4)
        self.line(self.get_x(), self.get_y(), 190, self.get_y())
        self.ln(3); self.set_text_color(0, 0, 0)

    def mono_line(self, teks, bold=False, color=None):
        self.set_font("Courier", "B" if bold else "", 9)
        if color: self.set_text_color(*color)
        self.set_x(28); self.multi_cell(0, 4.5, sp(teks))
        self.set_text_color(0, 0, 0)


def _tulis_steps_pdf(pdf, steps_, ok_color=OK_COLOR, fail_color=FAIL_COLOR):
    for s in steps_:
        if pdf.get_y() > 240:
            pdf.add_page(); pdf.watermark()
        pdf.set_font("Helvetica", "B", 10); pdf.set_text_color(*BRAND_COLOR); pdf.set_x(25)
        pdf.cell(0, 6, sp(f"{s['no']}  {s['judul']}"), ln=True)
        pdf.set_font("Helvetica", "I", 8); pdf.set_text_color(*GRAY); pdf.set_x(28)
        pdf.cell(0, 4, sp(f"[{s['ref']}]"), ln=True); pdf.set_text_color(0, 0, 0)
        for baris in s["isi"].split("\n"):
            is_result = (baris.strip().startswith("-->") or
                         "[OK]" in baris or "[TIDAK OK]" in baris or
                         "AMAN" in baris or "[DIABAIKAN]" in baris)
            pdf.mono_line(baris if baris.strip() else "",
                          bold=is_result,
                          color=(ok_color if s["ok"] else fail_color) if is_result else None)
        pdf.ln(2)


# ============================================================
# HELPER: RANGKUMAN (sumber tunggal untuk Word & PDF)
# ============================================================
def _build_rangkuman(R, G, Tor):
    """Kembalikan list (simbol, nilai, keterangan) -- dipakai Word & PDF agar identik."""
    rg = [
        ("Beta-1",     f"{R['beta1']:.4f}",        ""),
        ("c",          f"{R['c']:.2f} mm",          "Kedalaman sumbu netral"),
        ("a",          f"{R['a']:.2f} mm",          "Kedalaman blok tegangan"),
        ("Cc",         f"{R['Cc']/1000:.2f} kN",    "Gaya tekan beton"),
        ("Cs",         f"{R['Cs']/1000:.2f} kN",    "Gaya tekan baja total"),
        ("T",          f"{R['T']/1000:.2f} kN",     "Gaya tarik baja total"),
        ("et",         f"{R['et']:.5f}",            "Regangan tarik terjauh"),
        ("Phi",        f"{R['phi']:.4f}",           "Faktor reduksi lentur"),
        ("Rho",        f"{R['rho']*100:.4f}%",      "Rasio tulangan tarik"),
        ("Rho-min",    f"{R['rho_min']*100:.4f}%",  "Batas minimum"),
        ("Rho-max",    f"{R['rho_max']*100:.4f}%",  "Batas maksimum"),
        ("Mn",         f"{R['Mn']:.3f} kN.m",       "Momen nominal"),
        ("Phi.Mn",     f"{R['phiMn']:.3f} kN.m",    "Momen rencana kapasitas"),
        ("Mu",         f"{R['Mu']:.3f} kN.m",       "Momen rencana ultimit"),
        ("D/C-Lentur", f"{R['DC']:.3f}",            "Demand-to-Capacity Lentur"),
    ]
    if G is not None:
        rg.extend([
            ("Phi_v",      f"{G['Phi_v']:.2f}",             "Faktor reduksi geser"),
            ("Vc",         f"{G['Vc']:.2f} kN",             "Kapasitas geser beton"),
            ("Vs",         f"{G['Vs_efektif']:.2f} kN",     "Kapasitas geser sengkang"),
            ("Vn",         f"{G['Vn_efektif']:.2f} kN",     "Kapasitas geser nominal"),
            ("Phi.Vn",     f"{G['PhiVn_efektif']:.2f} kN",  "Kapasitas geser rencana"),
            ("Vu",         f"{G['Vu']:.2f} kN",             "Geser ultimit"),
            ("Av_pasang",  f"{G['Av_pasang']:.1f} mm2",     "Av sengkang terpasang"),
            ("Av_min",     f"{G['Av_min']:.1f} mm2",        "Av sengkang minimum"),
            ("D/C-Geser",  f"{G['DC_geser']:.3f}",          "Demand-to-Capacity Geser"),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        rg.extend([
            ("Phi_t",      f"{Tor['Phi_t']:.2f}",               "Faktor reduksi torsi"),
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",            "Batas ambang torsi"),
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",             "Torsi ultimit"),
            ("Tu_desain",  f"{Tor['Tu_desain']:.4f} kN.m",      "Tu desain (setelah reduksi)"),
            ("Phi.Tn",     f"{Tor['PhiTn_cap']:.4f} kN.m",      "Kapasitas torsi rencana"),
            ("D/C-Torsi",  f"{Tor['DC_torsi']:.3f}",            "Demand-to-Capacity Torsi"),
            ("At/s",       f"{Tor['At_per_s']:.6f} mm2/mm",     "Kebutuhan sengkang torsi"),
            ("Al_pakai",   f"{Tor['Al_pakai']:.2f} mm2",        "Tul. longitudinal torsi"),
            ("n_batang",   f"{Tor['n_batang']} D{int(Tor['db_long_torsi'])}", "Batang longitudinal torsi"),
            ("s_max_torsi",f"{Tor['s_max_torsi']:.1f} mm",      "Spasi maks sengkang torsi"),
        ])
    elif Tor is not None and Tor["abaikan_torsi"]:
        rg.extend([
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",    "Torsi ultimit"),
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",   "Batas ambang torsi"),
            ("D/C-Torsi",  "DIABAIKAN",                 "Tu < Phi_t x Tth"),
        ])
    return rg


def _build_cek_list(R, G, Tor, geser_inputs):
    """Kembalikan list (teks, ok_bool) untuk seksi Kesimpulan -- identik Word & PDF."""
    cl = [
        (f"Rho-min = {R['rho_min']*100:.4f}%  <=  Rho = {R['rho']*100:.4f}%", R["ok_rho_min"]),
        (f"Rho-max = {R['rho_max']*100:.4f}%  >=  Rho = {R['rho']*100:.4f}%", R["ok_rho_max"]),
        (f"et      = {R['et']:.5f}  >=  0.004",                                R["ok_et"]),
        (f"D/C-Lentur = {R['DC']:.3f}  <=  1.000",                             R["ok_dc"]),
    ]
    if G is not None:
        cl.extend([
            (f"D/C-Geser  = {G['DC_geser']:.3f}  <=  1.000",                G["ok_dc"]),
            (f"Vs_perlu   = {G['Vs_perlu']:.2f}  <=  Vs_max = {G['Vs_max']:.2f} kN", G["ok_Vs_max"]),
            (f"Av_pasang  = {G['Av_pasang']:.1f}  >=  Av_min = {G['Av_min']:.1f} mm2", G["ok_Av"]),
            (f"s_pasang   = {geser_inputs['s_seng']:.0f}  <=  s_max = {G['s_max']:.1f} mm", G["ok_spasi"]),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        cl.extend([
            (f"D/C-Torsi  = {Tor['DC_torsi']:.3f}  <=  1.000",              Tor["ok_DC_torsi"]),
            (f"Dimensi penampang (cek geser+torsi)",                          Tor["ok_dimensi"]),
            (f"Av+2At terpasang  >= Av+2At minimum",                         Tor["ok_Avt_min"]),
            (f"s_pasang = {geser_inputs['s_seng']:.0f} <= s_max_torsi = {Tor['s_max_torsi']:.1f} mm",
             Tor["ok_spasi_torsi"]),
        ])
    return cl


def _status_strings(R, G, Tor):
    """Kembalikan (status_lentur, status_geser, status_torsi, dc_torsi_str, aman_torsi, aman_total)."""
    status_lentur = "AMAN" if R["ok_dc"] else "TIDAK AMAN"
    status_geser  = ("AMAN" if (G is not None and G["ok_total"]) else
                     ("TIDAK AMAN" if G is not None else "belum dievaluasi"))
    if Tor is None:
        status_torsi = "belum dievaluasi"; dc_torsi_str = "-"; aman_torsi = True
    elif Tor["abaikan_torsi"]:
        status_torsi = "DIABAIKAN"
        dc_torsi_str = f"Tu={Tor['Tu']:.3f} < phi_t x Tth={Tor['phi_Tth']:.3f} kN.m"
        aman_torsi = True
    else:
        status_torsi = "AMAN" if Tor["ok_torsi_total"] else "TIDAK AMAN"
        dc_torsi_str = f"D/C = {Tor['DC_torsi']:.3f}"
        aman_torsi = Tor["ok_torsi_total"]
    aman_total = R["ok_dc"] and (G is None or G["ok_total"]) and aman_torsi
    return status_lentur, status_geser, status_torsi, dc_torsi_str, aman_torsi, aman_total


# ============================================================
# GENERATOR WORD
# ============================================================
def create_word_balok(fc, fy, b, h, cc_sel, ds, Mu, lapis_tarik, lapis_tekan,
                      R, steps, proj_info, png_buf,
                      lang="ID",
                      G=None, steps_geser=None, geser_inputs=None,
                      Tor=None, steps_torsi=None, torsi_inputs=None,
                      timestamp_str=None):
    LT = TEXT[lang]
    doc = Document()
    for sect in doc.sections:
        sect.top_margin    = Cm(2.5)
        sect.bottom_margin = Cm(2.5)
        sect.left_margin   = Cm(3.0)
        sect.right_margin  = Cm(2.5)

    def par(teks="", bold=False, italic=False, size=11, indent=0.0,
            align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(teks)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = "Calibri"
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    def judul_utama(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(teks); r.bold = True; r.font.size = Pt(14)
        r.font.name = "Calibri"; r.font.color.rgb = RGBColor(26, 60, 94)

    def subjudul(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(teks); r.bold = True; r.font.size = Pt(11.5)
        r.font.name = "Calibri"; r.font.color.rgb = RGBColor(26, 60, 94)

    def garis():
        par("", space_after=2)
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)

    def tulis_steps(steps_list):
        for s in steps_list:
            ok_col = (0x1B, 0x5E, 0x20) if s["ok"] else (0xB7, 0x1C, 0x1C)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(f"{s['no']} {s['judul']}")
            r1.bold = True; r1.font.size = Pt(10.5); r1.font.name = "Calibri"
            r1.font.color.rgb = RGBColor(26, 60, 94)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.left_indent = Cm(0.4)
            r2 = p2.add_run(f"[{s['ref']}]")
            r2.italic = True; r2.font.size = Pt(9); r2.font.name = "Calibri"
            r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
            for baris in s["isi"].split("\n"):
                is_result = (baris.strip().startswith("-->") or
                             "[OK]" in baris or "[TIDAK OK]" in baris or
                             "AMAN" in baris or "[DIABAIKAN]" in baris)
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_before = Pt(0)
                p3.paragraph_format.space_after  = Pt(0)
                p3.paragraph_format.left_indent  = Cm(0.5)
                r3 = p3.add_run(baris if baris.strip() else "")
                r3.font.name = "Courier New"; r3.font.size = Pt(9.5)
                if is_result:
                    r3.bold = True
                    r3.font.color.rgb = RGBColor(*ok_col)
            par(space_after=4)

    if timestamp_str is None:
        timestamp_str = datetime.datetime.now().strftime('%d %B %Y  %H:%M')

    # Header
    judul_utama(LT["lap_judul"])
    judul_utama(LT["lap_subjudul"])
    par(TEXT[lang]["subtitle"],
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55))
    par(LT["lap_ref"],
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55))
    par(f"Proyek: {proj_info['nama']}   |   Tanggal: {timestamp_str}",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x33, 0x33, 0x33), space_after=8)
    par(space_after=4)

    # A. DATA INPUT
    subjudul(LT["lap_A"])
    par(f"Proyek    : {proj_info['nama']}", size=10, indent=0.5, space_after=2)
    if proj_info.get("lokasi"):
        par(f"Lokasi    : {proj_info['lokasi']}", size=10, indent=0.5, space_after=2)
    if proj_info.get("engineer"):
        par(f"Engineer  : {proj_info['engineer']}", size=10, indent=0.5, space_after=2)
    if proj_info.get("nodok"):
        par(f"No. Dok   : {proj_info['nodok']}", size=10, indent=0.5, space_after=2)
    par(f"Tanggal   : {timestamp_str}", size=10, indent=0.5, space_after=6)

    par(LT["lap_material"], bold=True, size=10, indent=0.5, space_after=2)
    for simb, nilai in [("f'c", f"{fc:.1f} MPa"), ("fy", f"{fy:.0f} MPa"),
                        ("fyt", f"{geser_inputs['fyt']:.0f} MPa" if geser_inputs else "-")]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.add_run(f"{simb:<6} = {nilai}").font.size = Pt(10)
        p.runs[0].font.name = "Courier New"

    par(LT["lap_geometri"], bold=True, size=10, indent=0.5, space_after=2)
    for simb, nilai in [("b", f"{b:.0f} mm"), ("h", f"{h:.0f} mm"),
                        ("cc", f"{cc_sel:.0f} mm"), ("ds", f"{ds:.0f} mm")]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.add_run(f"{simb:<6} = {nilai}").font.size = Pt(10)
        p.runs[0].font.name = "Courier New"

    par(LT["lap_gayadalm"], bold=True, size=10, indent=0.5, space_after=2)
    gd_items = [("Mu", f"{R['Mu']:.2f} kN.m")]
    if G is not None:  gd_items.append(("Vu", f"{G['Vu']:.2f} kN"))
    if Tor is not None: gd_items.append(("Tu", f"{Tor['Tu']:.2f} kN.m"))
    for simb, nilai in gd_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.add_run(f"{simb:<6} = {nilai}").font.size = Pt(10)
        p.runs[0].font.name = "Courier New"

    par(LT["lap_tul_tarik"], bold=True, size=10, indent=0.5, space_after=2)
    for L in lapis_tarik:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        p.add_run(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                  f"As = {L['As']:.1f} mm2   y = {L['y']:.1f} mm").font.size = Pt(10)
        p.runs[0].font.name = "Courier New"
    par(f"  As-total = {R['As_tarik']:.1f} mm2  |  d-aktual = {R['d_tarik']:.2f} mm",
        bold=True, size=10, indent=1.0, space_after=4)

    par(LT["lap_tul_tekan"], bold=True, size=10, indent=0.5, space_after=2)
    if lapis_tekan:
        for L in lapis_tekan:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent  = Cm(1.0)
            p.add_run(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                      f"As' = {L['As']:.1f} mm2   y = {L['y']:.1f} mm").font.size = Pt(10)
            p.runs[0].font.name = "Courier New"
        par(f"  As'-total = {R['As_tekan']:.1f} mm2  |  d'-aktual = {R['d_tekan']:.2f} mm",
            bold=True, size=10, indent=1.0, space_after=4)
    else:
        par(LT["lap_tidak_ada_tekan"], size=10, indent=1.0,
            color=(0x55, 0x55, 0x55), space_after=4)

    par(space_after=4)
    par(LT["lap_visualisasi"], bold=True, size=10, indent=0.5, space_after=4)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(png_buf, width=Cm(9))
    par(space_after=6)

    # B. LENTUR
    subjudul(LT["lap_B"])
    par("Urutan perhitungan lentur mengacu pada SNI 2847:2019.",
        size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
    tulis_steps(steps)
    par(space_after=6)

    # C. RANGKUMAN (dari helper)
    subjudul(LT["lap_C"])
    for simb, nilai, ket in _build_rangkuman(R, G, Tor):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(f"{simb:<14}"); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = "Courier New"
        r2 = p.add_run(f"=  {nilai:<24}  {ket}")
        r2.font.size = Pt(10); r2.font.name = "Courier New"
    par(space_after=6)

    # D. KESIMPULAN (dari helper)
    subjudul(LT["lap_D"])
    for teks_k, ok_k in _build_cek_list(R, G, Tor, geser_inputs or {}):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.5)
        r_k = p.add_run(f"{teks_k}   --> {'[OK]' if ok_k else '[TIDAK OK]'}")
        r_k.font.name = "Courier New"; r_k.font.size = Pt(10); r_k.bold = True
        r_k.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if ok_k else RGBColor(0xB7, 0x1C, 0x1C))
    par(space_after=4)

    sl, sg, st, dc_str, _, aman_total = _status_strings(R, G, Tor)
    kes_lines = [
        f"Status Lentur  : {sl}  (D/C = {R['DC']:.3f})",
        f"Status Geser   : {sg}  (D/C = {G['DC_geser']:.3f})" if G else "Status Geser   : belum dievaluasi",
        f"Status Torsi   : {st}  ({dc_str})",
        "",
        "KESIMPULAN AKHIR: Penampang " + ("AMAN" if aman_total else "TIDAK AMAN") + " secara LENTUR + GESER + TORSI",
    ]
    for ln in kes_lines:
        p_kes = doc.add_paragraph()
        p_kes.paragraph_format.space_before = Pt(2); p_kes.paragraph_format.space_after = Pt(2)
        r_kes = p_kes.add_run(ln)
        r_kes.bold = True; r_kes.font.size = Pt(10.5); r_kes.font.name = "Calibri"
        if "TIDAK AMAN" in ln: r_kes.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
        elif "AMAN" in ln:     r_kes.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    par(space_after=6)

    # E. GESER
    if G is not None and steps_geser:
        subjudul(LT["lap_E"])
        par("Urutan perhitungan geser mengacu pada SNI 2847:2019.",
            size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
        tulis_steps(steps_geser)
        par(space_after=6)
        kes_g = (f"KESIMPULAN GESER : AMAN  |  D/C-Geser = {G['DC_geser']:.3f} <= 1.0"
                 if G["ok_total"] else
                 f"KESIMPULAN GESER : Kapasitas terpenuhi (D/C = {G['DC_geser']:.3f}), namun ada detailing yang tidak terpenuhi."
                 if G["ok_dc"] else
                 f"KESIMPULAN GESER : TIDAK AMAN  |  D/C-Geser = {G['DC_geser']:.3f} > 1.0")
        ok_g = G["ok_total"]
        p_kg = doc.add_paragraph()
        p_kg.paragraph_format.space_before = Pt(6); p_kg.paragraph_format.space_after = Pt(4)
        r_kg = p_kg.add_run(kes_g); r_kg.bold = True; r_kg.font.size = Pt(10.5)
        r_kg.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20) if ok_g else RGBColor(0xB7, 0x1C, 0x1C)

    # F. TORSI
    if Tor is not None and steps_torsi:
        subjudul(LT["lap_F"])
        par("Urutan perhitungan torsi mengacu pada SNI 2847:2019 Pasal 22.7.",
            size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
        tulis_steps(steps_torsi)
        par(space_after=6)
        if Tor["abaikan_torsi"]:
            kes_t = f"KESIMPULAN TORSI : DIABAIKAN  |  Tu = {Tor['Tu']:.3f} kN.m <= Phi_t x Tth = {Tor['phi_Tth']:.3f} kN.m"
            ok_t = True
        elif Tor["ok_torsi_total"]:
            kes_t = f"KESIMPULAN TORSI : AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} <= 1.0"
            ok_t = True
        else:
            kes_t = f"KESIMPULAN TORSI : TIDAK AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} > 1.0"
            ok_t = False
        p_kt = doc.add_paragraph()
        p_kt.paragraph_format.space_before = Pt(6); p_kt.paragraph_format.space_after = Pt(4)
        r_kt = p_kt.add_run(kes_t); r_kt.bold = True; r_kt.font.size = Pt(10.5)
        r_kt.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20) if ok_t else RGBColor(0xB7, 0x1C, 0x1C)

    par(space_after=6)
    garis()
    par("Referensi: SNI 2847:2019 | ACI 318-14  --  "
        "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan.",
        size=8, italic=True, color=(0x99, 0x99, 0x99),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ============================================================
# GENERATOR PDF
# ============================================================
def create_pdf_balok(fc, fy, b, h, cc_sel, ds, Mu, lapis_tarik, lapis_tekan,
                     R, steps, proj_info, png_buf,
                     lang="ID",
                     G=None, steps_geser=None, geser_inputs=None,
                     Tor=None, steps_torsi=None, torsi_inputs=None,
                     timestamp_str=None):
    LT  = TEXT[lang]
    pdf = LaporanBalokPDF(proj_info["nama"])
    pdf.add_page(); pdf.watermark()

    if timestamp_str is None:
        timestamp_str = datetime.datetime.now().strftime('%d %B %Y  %H:%M')

    # Header
    pdf.set_font("Helvetica", "B", 15); pdf.set_text_color(*BRAND_COLOR); pdf.ln(2)
    pdf.cell(0, 9, sp(LT["lap_judul"]), ln=True, align="C")
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, sp(LT["lap_subjudul"]), ln=True, align="C")
    pdf.set_font("Helvetica", "I", 9); pdf.set_text_color(*GRAY)
    pdf.cell(0, 5, sp(TEXT[lang]["subtitle"]), ln=True, align="C")
    pdf.cell(0, 5, sp(LT["lap_ref"]), ln=True, align="C")
    pdf.cell(0, 5, sp(f"Proyek: {proj_info['nama']}   |   Tanggal: {timestamp_str}"), ln=True, align="C")
    if proj_info.get("engineer"):
        pdf.cell(0, 5, sp(f"Engineer: {proj_info['engineer']}"), ln=True, align="C")
    pdf.ln(6); pdf.set_draw_color(*BRAND_COLOR); pdf.set_line_width(0.6)
    pdf.line(25, pdf.get_y(), 190, pdf.get_y()); pdf.ln(6); pdf.set_text_color(0, 0, 0)

    # A. DATA INPUT
    pdf.section_title(LT["lap_A"])
    rows = [
        ("Mu",   f"{Mu:.2f} kN.m",    "Momen rencana"),
        ("fc",   f"{fc:.1f} MPa",     "Kuat tekan beton"),
        ("fy",   f"{fy:.0f} MPa",     "Kuat leleh tulangan lentur"),
        ("b",    f"{b:.0f} mm",       "Lebar balok"),
        ("h",    f"{h:.0f} mm",       "Tinggi total balok"),
        ("cc",   f"{cc_sel:.0f} mm",  "Tebal selimut bersih"),
        ("ds",   f"{ds:.0f} mm",      "Diameter sengkang"),
    ]
    if geser_inputs is not None:
        rows.extend([
            ("Vu",     f"{geser_inputs['Vu']:.2f} kN",    "Gaya geser ultimit"),
            ("fyt",    f"{geser_inputs['fyt']:.0f} MPa",  "Kuat leleh sengkang"),
            ("s",      f"{geser_inputs['s_seng']:.0f} mm","Jarak sengkang"),
            ("n_kaki", f"{geser_inputs['n_kaki']:d}",     "Jumlah kaki sengkang"),
        ])
    if torsi_inputs is not None:
        rows.extend([
            ("Tu",            f"{torsi_inputs['Tu']:.2f} kN.m",         "Momen torsi ultimit"),
            ("tipe_torsi",    f"{torsi_inputs['tipe_torsi']}",           "Tipe torsi"),
            ("db_long_torsi", f"D{int(torsi_inputs['db_long_torsi'])} mm","Dia. tul. long. torsi"),
        ])
    for simb, nilai, ket in rows:
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.cell(22, 5, sp(f"{simb:<10}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(36, 5, sp(f"=  {nilai}"), ln=False)
        pdf.set_font("Helvetica", "I", 8.5); pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, sp(f"({ket})"), ln=True); pdf.set_text_color(0, 0, 0)

    pdf.ln(2)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp(LT["lap_tul_tarik"]), ln=True); pdf.set_text_color(0, 0, 0)
    for L in lapis_tarik:
        pdf.set_x(30); pdf.set_font("Courier", "", 9)
        pdf.cell(0, 4.5, sp(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                            f"As = {L['As']:.1f} mm2   y = {L['y']:.1f} mm"), ln=True)
    pdf.set_x(30); pdf.set_font("Courier", "B", 9)
    pdf.cell(0, 4.5, sp(f"As-total = {R['As_tarik']:.1f} mm2  |  d-aktual = {R['d_tarik']:.2f} mm"), ln=True)

    pdf.ln(1)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp(LT["lap_tul_tekan"]), ln=True); pdf.set_text_color(0, 0, 0)
    if lapis_tekan:
        for L in lapis_tekan:
            pdf.set_x(30); pdf.set_font("Courier", "", 9)
            pdf.cell(0, 4.5, sp(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                                f"As' = {L['As']:.1f} mm2   y = {L['y']:.1f} mm"), ln=True)
        pdf.set_x(30); pdf.set_font("Courier", "B", 9)
        pdf.cell(0, 4.5, sp(f"As'-total = {R['As_tekan']:.1f} mm2  |  d'-aktual = {R['d_tekan']:.2f} mm"), ln=True)
    else:
        pdf.set_x(30); pdf.set_font("Helvetica", "I", 9); pdf.set_text_color(*GRAY)
        pdf.cell(0, 4.5, sp(LT["lap_tidak_ada_tekan"]), ln=True); pdf.set_text_color(0, 0, 0)

    pdf.ln(3)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp(LT["lap_visualisasi"]), ln=True); pdf.set_text_color(0, 0, 0)
    img_path = "/tmp/_penampang_balok.png"
    with open(img_path, "wb") as fimg:
        fimg.write(png_buf.getvalue())
    pdf.image(img_path, x=70, y=pdf.get_y() + 1, w=70)
    pdf.ln(82)

    # B. LENTUR
    if pdf.get_y() > 220: pdf.add_page(); pdf.watermark()
    pdf.section_title(LT["lap_B"])
    _tulis_steps_pdf(pdf, steps)

    # C. RANGKUMAN (dari helper)
    if pdf.get_y() > 200: pdf.add_page(); pdf.watermark()
    pdf.section_title(LT["lap_C"])
    for simb, nilai, ket in _build_rangkuman(R, G, Tor):
        if pdf.get_y() > 260: pdf.add_page(); pdf.watermark()
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.cell(28, 5, sp(f"{simb:<13}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(38, 5, sp(f"=  {nilai}"), ln=False)
        if ket:
            pdf.set_font("Helvetica", "I", 8.5); pdf.set_text_color(*GRAY)
            pdf.cell(0, 5, sp(f"({ket})"), ln=True); pdf.set_text_color(0, 0, 0)
        else:
            pdf.ln()
    pdf.ln(4)

    # D. KESIMPULAN (dari helper)
    if pdf.get_y() > 220: pdf.add_page(); pdf.watermark()
    pdf.section_title(LT["lap_D"])
    for teks_k, ok_k in _build_cek_list(R, G, Tor, geser_inputs or {}):
        if pdf.get_y() > 260: pdf.add_page(); pdf.watermark()
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.set_text_color(*(OK_COLOR if ok_k else FAIL_COLOR))
        pdf.cell(0, 5.5, sp(f"{teks_k}   --> {'[OK]' if ok_k else '[TIDAK OK]'}"), ln=True)
        pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    sl, sg, st, dc_str, _, aman_total = _status_strings(R, G, Tor)
    kes_lines = [
        f"Status Lentur  : {sl}  (D/C = {R['DC']:.3f})",
        f"Status Geser   : {sg}  (D/C = {G['DC_geser']:.3f})" if G else "Status Geser   : belum dievaluasi",
        f"Status Torsi   : {st}  ({dc_str})",
        "",
        "KESIMPULAN AKHIR: Penampang " + ("AMAN" if aman_total else "TIDAK AMAN") + " secara LENTUR + GESER + TORSI",
    ]
    for ln_t in kes_lines:
        if not ln_t: pdf.ln(2); continue
        ok_line   = ("AMAN" in ln_t and "TIDAK" not in ln_t)
        fail_line = "TIDAK AMAN" in ln_t
        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        if fail_line:   pdf.set_text_color(*FAIL_COLOR)
        elif ok_line:   pdf.set_text_color(*OK_COLOR)
        else:           pdf.set_text_color(*GRAY)
        pdf.multi_cell(0, 6, sp(ln_t))
        pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    # E. GESER
    if G is not None and steps_geser:
        if pdf.get_y() > 220: pdf.add_page(); pdf.watermark()
        pdf.section_title(LT["lap_E"])
        _tulis_steps_pdf(pdf, steps_geser)
        kes_g = (f"KESIMPULAN GESER : AMAN  |  D/C-Geser = {G['DC_geser']:.3f} <= 1.0"
                 if G["ok_total"] else
                 f"KESIMPULAN GESER : Kapasitas terpenuhi (D/C = {G['DC_geser']:.3f}), namun ada detailing yang tidak terpenuhi."
                 if G["ok_dc"] else
                 f"KESIMPULAN GESER : TIDAK AMAN  |  D/C-Geser = {G['DC_geser']:.3f} > 1.0")
        ok_g = G["ok_total"]
        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*(OK_COLOR if ok_g else FAIL_COLOR))
        pdf.multi_cell(0, 6, sp(kes_g)); pdf.set_text_color(0, 0, 0)

    # F. TORSI
    if Tor is not None and steps_torsi:
        if pdf.get_y() > 220: pdf.add_page(); pdf.watermark()
        pdf.section_title(LT["lap_F"])
        _tulis_steps_pdf(pdf, steps_torsi)
        if Tor["abaikan_torsi"]:
            kes_t = f"KESIMPULAN TORSI : DIABAIKAN  |  Tu = {Tor['Tu']:.3f} kN.m <= Phi_t x Tth = {Tor['phi_Tth']:.3f} kN.m"
            ok_t = True
        elif Tor["ok_torsi_total"]:
            kes_t = f"KESIMPULAN TORSI : AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} <= 1.0"
            ok_t = True
        else:
            kes_t = f"KESIMPULAN TORSI : TIDAK AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} > 1.0"
            ok_t = False
        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*(OK_COLOR if ok_t else FAIL_COLOR))
        pdf.multi_cell(0, 6, sp(kes_t)); pdf.set_text_color(0, 0, 0)

    buf = io.BytesIO(); pdf.output(buf); buf.seek(0)
    return buf
