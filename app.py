"""
=============================================================
APLIKASI WEB — KAPASITAS LENTUR BALOK BETON BERTULANG
Referensi : SNI 2847:2019 (ACI 318-14)
Framework  : Streamlit
Library    : python-docx (laporan Word)
Jalankan   : streamlit run app.py
=============================================================
"""

import math
import io
import datetime
import streamlit as st
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── KONFIGURASI HALAMAN ─────────────────────────────────────
st.set_page_config(
    page_title="Lentur Balok Beton | SNI 2847:2019",
    page_icon="🏗️",
    layout="wide",
)

# ─── CSS KUSTOM ──────────────────────────────────────────────
st.markdown("""
<style>
    .main-title { font-size:1.6rem; font-weight:600; color:#1a3c5e; margin-bottom:0; }
    .sub-title   { font-size:0.9rem; color:#666; margin-bottom:1.5rem; }
    .step-box    { background:#f8f9fa; border-left:4px solid #1a3c5e; border-radius:0 8px 8px 0;
                   padding:12px 16px; margin-bottom:10px; font-family:monospace; font-size:0.85rem; }
    .step-header { font-weight:700; color:#1a3c5e; font-size:0.8rem;
                   text-transform:uppercase; letter-spacing:0.5px; margin-bottom:4px; }
    .result-ok   { background:#e8f5e9; border-color:#2e7d32; color:#1b5e20;
                   padding:12px 16px; border-radius:8px; border-left:4px solid; font-weight:600; }
    .result-warn { background:#fff8e1; border-color:#f9a825; color:#5d4037;
                   padding:12px 16px; border-radius:8px; border-left:4px solid; font-weight:600; }
    .result-fail { background:#ffebee; border-color:#c62828; color:#b71c1c;
                   padding:12px 16px; border-radius:8px; border-left:4px solid; font-weight:600; }
    .metric-card  { background:white; border:1px solid #e0e0e0; border-radius:10px;
                    padding:14px; text-align:center; }
    .metric-label { font-size:0.75rem; color:#888; margin-bottom:4px; }
    .metric-value { font-size:1.5rem; font-weight:700; color:#1a3c5e; }
    .metric-unit  { font-size:0.75rem; color:#aaa; }
    .section-divider { border:none; border-top:2px solid #e0e0e0; margin:1.5rem 0; }
    .ref-badge    { display:inline-block; background:#e3f2fd; color:#1565c0;
                    font-size:0.7rem; padding:2px 8px; border-radius:20px; margin-bottom:6px; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════
# FUNGSI PERHITUNGAN
# ═══════════════════════════════════════════════════════════
def hitung_lentur_balok(fc, fy, b, h, d, As, Asp=0.0):
    R = {}

    # Langkah 1 — beta1
    if fc <= 28:
        beta1 = 0.85
        R['beta1_cara'] = f"f'c = {fc} MPa ≤ 28 MPa  →  β₁ = 0.85"
    elif fc >= 56:
        beta1 = 0.65
        R['beta1_cara'] = f"f'c = {fc} MPa ≥ 56 MPa  →  β₁ = 0.65"
    else:
        beta1 = 0.85 - 0.05 * (fc - 28) / 7
        beta1 = max(0.65, min(0.85, beta1))
        R['beta1_cara'] = (f"β₁ = 0.85 - 0.05 × (f'c - 28) / 7\n"
                           f"   = 0.85 - 0.05 × ({fc} - 28) / 7")
    R['beta1'] = beta1

    As_net   = As - Asp
    T        = As_net * fy
    a        = (As_net * fy) / (0.85 * fc * b)
    c        = a / beta1
    rho      = As / (b * d)
    rho_minA = 0.25 * math.sqrt(fc) / fy
    rho_minB = 1.4 / fy
    rho_min  = max(rho_minA, rho_minB)
    rho_bal  = (0.85 * beta1 * fc / fy) * (600 / (600 + fy))
    rho_max  = 0.75 * rho_bal
    et       = 0.003 * (d - c) / c

    if et >= 0.005:
        phi = 0.90
        R['phi_cara'] = "εt ≥ 0.005 → Tension-controlled → φ = 0.90"
    elif et <= 0.002:
        phi = 0.65
        R['phi_cara'] = "εt ≤ 0.002 → Compression-controlled → φ = 0.65"
    else:
        phi = 0.65 + (et - 0.002) * (250 / 3)
        R['phi_cara'] = (f"φ = 0.65 + (εt - 0.002) × 250/3\n"
                         f"  = 0.65 + ({et:.5f} - 0.002) × 83.333")

    Mn    = As_net * fy * (d - a / 2) / 1_000_000
    phiMn = phi * Mn

    R.update({
        'As_net': As_net, 'T': T, 'a': a, 'c': c,
        'rho': rho, 'rho_minA': rho_minA, 'rho_minB': rho_minB,
        'rho_min': rho_min, 'rho_bal': rho_bal, 'rho_max': rho_max,
        'et': et, 'phi': phi, 'Mn': Mn, 'phiMn': phiMn,
        'ok_rho_min': rho >= rho_min,
        'ok_rho_max': rho <= rho_max,
        'ok_et':      et >= 0.004,
    })
    return R


def buat_steps(fc, fy, b, h, d, As, Asp, R):
    """Kembalikan list dict step — dipakai baik UI maupun Word."""
    return [
        {
            "no": "Langkah 1", "ref": "SNI 2847:2019 Pasal 22.2.2.4.3",
            "judul": "Faktor blok tegangan ekivalen β₁",
            "isi": f"{R['beta1_cara']}\n► β₁ = {R['beta1']:.4f}",
            "ok": True,
        },
        {
            "no": "Langkah 2", "ref": "SNI 2847:2019 Pasal 22.2.1",
            "judul": "Gaya tarik tulangan (T)",
            "isi": (f"T = As_net × fy\n"
                    f"  = ({As:.0f} - {Asp:.0f}) × {fy:.0f}\n"
                    f"► T = {R['T']:,.0f} N  =  {R['T']/1000:.1f} kN"),
            "ok": True,
        },
        {
            "no": "Langkah 3", "ref": "SNI 2847:2019 Pasal 22.2.2.4",
            "judul": "Kedalaman blok tegangan ekivalen (a)",
            "isi": (f"a = As_net × fy / (0.85 × f'c × b)\n"
                    f"  = {R['As_net']:.0f} × {fy:.0f} / (0.85 × {fc} × {b:.0f})\n"
                    f"► a = {R['a']:.2f} mm"),
            "ok": True,
        },
        {
            "no": "Langkah 4", "ref": "SNI 2847:2019 Pasal 22.2.2.4.1",
            "judul": "Kedalaman sumbu netral (c)",
            "isi": (f"c = a / β₁  =  {R['a']:.2f} / {R['beta1']:.4f}\n"
                    f"► c = {R['c']:.2f} mm"),
            "ok": True,
        },
        {
            "no": "Langkah 5", "ref": "SNI 2847:2019 Pasal 9.6.1",
            "judul": "Rasio tulangan aktual (ρ)",
            "isi": (f"ρ = As / (b × d)  =  {As:.0f} / ({b:.0f} × {d:.0f})\n"
                    f"► ρ = {R['rho']:.6f}  =  {R['rho']*100:.4f}%"),
            "ok": R['ok_rho_min'] and R['ok_rho_max'],
        },
        {
            "no": "Langkah 6", "ref": "SNI 2847:2019 Pasal 9.6.1.2",
            "judul": "Rasio tulangan minimum (ρ_min)",
            "isi": (f"ρ_min = max( 0.25√f'c/fy  ,  1.4/fy )\n"
                    f"      = max( {R['rho_minA']:.6f}  ,  {R['rho_minB']:.6f} )\n"
                    f"► ρ_min = {R['rho_min']:.6f}  =  {R['rho_min']*100:.4f}%\n"
                    f"  Kontrol: ρ {'≥' if R['ok_rho_min'] else '<'} ρ_min  "
                    f"{'[OK]' if R['ok_rho_min'] else '[TIDAK OK]'}"),
            "ok": R['ok_rho_min'],
        },
        {
            "no": "Langkah 7", "ref": "SNI 2847:2019 Pasal 21.2.2",
            "judul": "Rasio tulangan maksimum (ρ_max)",
            "isi": (f"ρ_bal = 0.85 × β₁ × f'c/fy × 600/(600+fy)\n"
                    f"      = 0.85 × {R['beta1']:.4f} × {fc}/{fy:.0f} × 600/{600+fy:.0f}\n"
                    f"ρ_bal = {R['rho_bal']:.6f}  =  {R['rho_bal']*100:.4f}%\n"
                    f"ρ_max = 0.75 × ρ_bal = {R['rho_max']:.6f}  =  {R['rho_max']*100:.4f}%\n"
                    f"  Kontrol: ρ {'≤' if R['ok_rho_max'] else '>'} ρ_max  "
                    f"{'[OK]' if R['ok_rho_max'] else '[TIDAK OK]'}"),
            "ok": R['ok_rho_max'],
        },
        {
            "no": "Langkah 8", "ref": "SNI 2847:2019 Pasal 21.2.2",
            "judul": "Regangan tarik tulangan (εt)",
            "isi": (f"εt = 0.003 × (d - c) / c\n"
                    f"   = 0.003 × ({d:.0f} - {R['c']:.2f}) / {R['c']:.2f}\n"
                    f"► εt = {R['et']:.5f}\n"
                    f"  {'εt >= 0.005 → Tension-controlled [OK]' if R['et'] >= 0.005 else '0.004 <= εt < 0.005 → Zona transisi [PERLU TINJAUAN]' if R['et'] >= 0.004 else 'εt < 0.004 → Tidak memenuhi syarat [TIDAK OK]'}"),
            "ok": R['ok_et'],
        },
        {
            "no": "Langkah 9", "ref": "SNI 2847:2019 Tabel 21.2.2",
            "judul": "Faktor reduksi kekuatan (φ)",
            "isi": f"{R['phi_cara']}\n► φ = {R['phi']:.4f}",
            "ok": R['et'] >= 0.004,
        },
        {
            "no": "Langkah 10", "ref": "SNI 2847:2019 Pasal 22.3.2",
            "judul": "Momen nominal dan momen rencana",
            "isi": (f"Mn = As_net × fy × (d - a/2)\n"
                    f"   = {R['As_net']:.0f} × {fy:.0f} × ({d:.0f} - {R['a']/2:.2f})\n"
                    f"   = {R['As_net'] * fy * (d - R['a']/2):,.0f} N.mm\n"
                    f"► Mn   = {R['Mn']:.3f} kN.m\n\n"
                    f"φMn = φ × Mn = {R['phi']:.4f} × {R['Mn']:.3f}\n"
                    f"► φMn = {R['phiMn']:.3f} kN.m"),
            "ok": True,
        },
    ]


# ═══════════════════════════════════════════════════════════
# FUNGSI GENERATE WORD (.docx)
# ═══════════════════════════════════════════════════════════
def set_cell_bg(cell, hex_color):
    """Warnai background cell tabel."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """Atur border cell tabel."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'),   kwargs.get('val',   'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',    '4'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color', '2E75B6'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def generate_word(fc, fy, b, h, d, As, Asp, R, steps):
    doc = Document()

    # ── Margin A4 ──
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2.5)
        section.bottom_margin= Cm(2.0)

    # ── Garis header dekoratif ──
    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after  = Pt(4)
    run_line = p_line.add_run("─" * 80)
    run_line.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run_line.font.size = Pt(7)

    # ── Judul Utama ──
    judul = doc.add_paragraph()
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    judul.paragraph_format.space_after = Pt(2)
    r = judul.add_run("PERHITUNGAN KAPASITAS LENTUR BALOK BETON BERTULANG")
    r.font.bold  = True
    r.font.size  = Pt(14)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    subjudul = doc.add_paragraph()
    subjudul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subjudul.paragraph_format.space_after = Pt(2)
    r2 = subjudul.add_run("Referensi: SNI 2847:2019 (setara ACI 318-14)")
    r2.font.size  = Pt(10)
    r2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    tgl = doc.add_paragraph()
    tgl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tgl.paragraph_format.space_after = Pt(6)
    r3 = tgl.add_run(f"Tanggal: {datetime.datetime.now().strftime('%d %B %Y')}")
    r3.font.size  = Pt(9)
    r3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    p_line2 = doc.add_paragraph()
    p_line2.paragraph_format.space_before = Pt(0)
    p_line2.paragraph_format.space_after  = Pt(10)
    run_line2 = p_line2.add_run("─" * 80)
    run_line2.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
    run_line2.font.size = Pt(7)

    def heading2(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(teks.upper())
        r.font.bold  = True
        r.font.size  = Pt(11)
        r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
        # Garis bawah tipis
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '2E75B6')
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # ════════════════════════════════════
    # BAGIAN 1 — DATA INPUT
    # ════════════════════════════════════
    heading2("1. Data Input Penampang")

    tbl_input = doc.add_table(rows=1, cols=4)
    tbl_input.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_input.style = 'Table Grid'

    # Header
    hdrs = ["Parameter", "Simbol", "Nilai", "Satuan"]
    warna_hdr = "1A3C5E"
    for i, h_teks in enumerate(hdrs):
        cell = tbl_input.rows[0].cells[i]
        set_cell_bg(cell, warna_hdr)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_teks)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)

    # Data baris
    data_input = [
        ("Kuat tekan beton",        "f'c",  f"{fc:.1f}",   "MPa"),
        ("Kuat leleh tulangan",     "fy",   f"{fy:.0f}",   "MPa"),
        ("Lebar balok",             "b",    f"{b:.0f}",    "mm"),
        ("Tinggi total balok",      "h",    f"{h:.0f}",    "mm"),
        ("Tinggi efektif",          "d",    f"{d:.0f}",    "mm"),
        ("Luas tulangan tarik",     "As",   f"{As:.0f}",   "mm²"),
        ("Luas tulangan tekan",     "As'",  f"{Asp:.0f}",  "mm²"),
    ]
    for i, (nama, simb, nilai, sat) in enumerate(data_input):
        row = tbl_input.add_row()
        bg  = "EBF2FA" if i % 2 == 0 else "FFFFFF"
        vals = [nama, simb, nilai, sat]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER]
        for j, (v, al) in enumerate(zip(vals, aligns)):
            c = row.cells[j]
            set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = al
            run = p.add_run(v)
            run.font.size = Pt(10)
            if j == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

    # Set lebar kolom
    for i, w in enumerate([Cm(6), Cm(2.5), Cm(3), Cm(2.5)]):
        for row in tbl_input.rows:
            row.cells[i].width = w

    doc.add_paragraph()

    # ════════════════════════════════════
    # BAGIAN 2 — RINGKASAN HASIL
    # ════════════════════════════════════
    heading2("2. Ringkasan Hasil Utama")

    tbl_hasil = doc.add_table(rows=1, cols=5)
    tbl_hasil.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_hasil.style = 'Table Grid'

    hdrs2 = ["Parameter", "Simbol", "Nilai", "Satuan", "Status"]
    for i, h_teks in enumerate(hdrs2):
        cell = tbl_hasil.rows[0].cells[i]
        set_cell_bg(cell, "1A3C5E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h_teks)
        run.font.bold  = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size  = Pt(10)

    all_ok = R['ok_rho_min'] and R['ok_rho_max'] and R['ok_et']
    status_global = "OK" if all_ok else "PERIKSA"

    et_status = "OK" if R['et'] >= 0.005 else ("TRANSISI" if R['et'] >= 0.004 else "TIDAK OK")
    rho_status = "OK" if (R['ok_rho_min'] and R['ok_rho_max']) else "TIDAK OK"

    def status_color(s):
        if "OK" in s and "TIDAK" not in s and "TRANSISI" not in s:
            return "D6F0DA", "1B5E20"   # hijau
        elif "TRANSISI" in s or "PERIKSA" in s:
            return "FFF8DC", "5D4037"   # kuning
        else:
            return "FDECEA", "B71C1C"   # merah

    data_hasil = [
        ("Faktor β₁",                   "β₁",    f"{R['beta1']:.4f}",   "—",     "—"),
        ("Kedalaman blok tegangan",      "a",     f"{R['a']:.2f}",       "mm",    "—"),
        ("Kedalaman sumbu netral",       "c",     f"{R['c']:.2f}",       "mm",    "—"),
        ("Faktor reduksi kekuatan",      "φ",     f"{R['phi']:.4f}",     "—",     "—"),
        ("Regangan tarik tulangan",      "εt",    f"{R['et']:.5f}",      "—",     et_status),
        ("Rasio tulangan",               "ρ",     f"{R['rho']*100:.4f}", "%",     rho_status),
        ("Rasio tulangan minimum",       "ρ_min", f"{R['rho_min']*100:.4f}", "%", "—"),
        ("Rasio tulangan maksimum",      "ρ_max", f"{R['rho_max']*100:.4f}", "%", "—"),
        ("Momen nominal",                "Mn",    f"{R['Mn']:.3f}",      "kN·m", "—"),
        ("Momen rencana",                "φMn",   f"{R['phiMn']:.3f}",   "kN·m", status_global),
    ]
    for i, (nama, simb, nilai, sat, sts) in enumerate(data_hasil):
        row = tbl_hasil.add_row()
        bg  = "EBF2FA" if i % 2 == 0 else "FFFFFF"
        vals   = [nama, simb, nilai, sat, sts]
        aligns = [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                  WD_ALIGN_PARAGRAPH.CENTER]
        for j, (v, al) in enumerate(zip(vals, aligns)):
            c = row.cells[j]
            if j == 4 and sts not in ("—", ""):
                sbg, sfg = status_color(sts)
                set_cell_bg(c, sbg)
            else:
                set_cell_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = al
            run = p.add_run(v)
            run.font.size = Pt(10)
            if j == 1:
                run.font.bold  = True
                run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
            if j == 4 and sts not in ("—", ""):
                _, sfg = status_color(sts)
                r_val = int(sfg[0:2], 16)
                g_val = int(sfg[2:4], 16)
                b_val = int(sfg[4:6], 16)
                run.font.bold  = True
                run.font.color.rgb = RGBColor(r_val, g_val, b_val)

    for i, w in enumerate([Cm(5.5), Cm(2), Cm(3), Cm(2), Cm(2.5)]):
        for row in tbl_hasil.rows:
            row.cells[i].width = w

    doc.add_paragraph()

    # ════════════════════════════════════
    # BAGIAN 3 — STATUS KESELURUHAN
    # ════════════════════════════════════
    heading2("3. Status Penampang")

    all_ok = R['ok_rho_min'] and R['ok_rho_max'] and R['ok_et']
    if all_ok and R['et'] >= 0.005:
        status_teks = "PENAMPANG OK — Tension-controlled, memenuhi seluruh syarat SNI 2847:2019"
        bg_status, fg_status = "D6F0DA", "1B5E20"
    elif all_ok:
        status_teks = "PERLU TINJAUAN — Zona transisi (0.004 <= εt < 0.005)"
        bg_status, fg_status = "FFF8DC", "5D4037"
    else:
        masalah = []
        if not R['ok_rho_min']: masalah.append("ρ < ρ_min (tambah tulangan)")
        if not R['ok_rho_max']: masalah.append("ρ > ρ_max (kurangi tulangan)")
        if not R['ok_et']:      masalah.append("εt < 0.004 (tidak memenuhi syarat)")
        status_teks = "TIDAK OK — " + " | ".join(masalah)
        bg_status, fg_status = "FDECEA", "B71C1C"

    tbl_status = doc.add_table(rows=1, cols=1)
    tbl_status.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_status.style = 'Table Grid'
    c_status = tbl_status.rows[0].cells[0]
    set_cell_bg(c_status, bg_status)
    p_status = c_status.paragraphs[0]
    p_status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_status = p_status.add_run(status_teks)
    r_status.font.bold = True
    r_status.font.size = Pt(11)
    r_val = int(fg_status[0:2], 16)
    g_val = int(fg_status[2:4], 16)
    b_val = int(fg_status[4:6], 16)
    r_status.font.color.rgb = RGBColor(r_val, g_val, b_val)
    c_status.width = Cm(15)

    doc.add_paragraph()

    # ════════════════════════════════════
    # BAGIAN 4 — URUTAN PERHITUNGAN
    # ════════════════════════════════════
    heading2("4. Urutan Perhitungan Lengkap")

    for step in steps:
        # Header step
        p_step_hdr = doc.add_paragraph()
        p_step_hdr.paragraph_format.space_before = Pt(8)
        p_step_hdr.paragraph_format.space_after  = Pt(0)

        r_no  = p_step_hdr.add_run(f"{step['no']}  ")
        r_no.font.bold  = True
        r_no.font.size  = Pt(10)
        r_no.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

        r_jdl = p_step_hdr.add_run(step['judul'])
        r_jdl.font.bold  = True
        r_jdl.font.size  = Pt(10)
        r_jdl.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)

        # Badge referensi
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(0)
        p_ref.paragraph_format.space_after  = Pt(2)
        p_ref.paragraph_format.left_indent  = Cm(0.3)
        r_ref = p_ref.add_run(step['ref'])
        r_ref.font.italic = True
        r_ref.font.size   = Pt(8)
        r_ref.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        # Isi perhitungan (monospace)
        tbl_step = doc.add_table(rows=1, cols=1)
        tbl_step.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl_step.style = 'Table Grid'
        c_step = tbl_step.rows[0].cells[0]
        bg_step = "F0F7F0" if step['ok'] else "FFF3F3"
        set_cell_bg(c_step, bg_step)

        for baris in step['isi'].split('\n'):
            p_baris = c_step.add_paragraph() if c_step.paragraphs[0].text else c_step.paragraphs[0]
            if c_step.paragraphs[0].text and baris:
                p_baris = c_step.add_paragraph()
            p_baris.paragraph_format.space_before = Pt(0)
            p_baris.paragraph_format.space_after  = Pt(0)
            run_b = p_baris.add_run(baris)
            run_b.font.name = "Courier New"
            run_b.font.size = Pt(9)
            if baris.startswith("►") or baris.startswith(">>"):
                run_b.font.bold = True
                run_b.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20)
                                        if step['ok'] else RGBColor(0xB7, 0x1C, 0x1C))

        c_step.width = Cm(15)

    # ── Footer ──
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(6)
    run_foot = p_foot.add_run(
        "Referensi: SNI 2847:2019 | ACI 318-14  —  "
        "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan."
    )
    run_foot.font.size  = Pt(8)
    run_foot.font.italic = True
    run_foot.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Simpan ke buffer memory
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════════════════════
# UI — HEADER
# ═══════════════════════════════════════════════════════════
st.markdown('<p class="main-title">🏗️ Kapasitas Lentur Balok Beton Bertulang</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">Referensi: SNI 2847:2019 (setara ACI 318-14) | Urutan perhitungan sesuai urutan pasal</p>', unsafe_allow_html=True)
st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════
# UI — LAYOUT DUA KOLOM
# ═══════════════════════════════════════════════════════════
col_input, col_hasil = st.columns([1, 2], gap="large")

with col_input:
    st.markdown("### 📋 Data Input Penampang")

    st.markdown("**Material**")
    fc  = st.number_input("f'c — Kuat tekan beton (MPa)",   min_value=17.0, max_value=100.0, value=30.0,   step=1.0,  format="%.1f")
    fy  = st.number_input("fy — Kuat leleh tulangan (MPa)", min_value=240.0, max_value=600.0, value=400.0, step=10.0, format="%.0f")

    st.markdown("**Geometri Penampang**")
    col_b, col_h = st.columns(2)
    with col_b:
        b = st.number_input("b (mm)\nLebar", min_value=100.0, max_value=2000.0, value=300.0, step=10.0, format="%.0f")
    with col_h:
        h = st.number_input("h (mm)\nTinggi total", min_value=100.0, max_value=5000.0, value=500.0, step=10.0, format="%.0f")
    d = st.number_input("d (mm) — Tinggi efektif (h − selimut − ½D)",
                        min_value=50.0, max_value=4900.0, value=440.0, step=5.0, format="%.0f")

    st.markdown("**Tulangan**")
    As  = st.number_input("As (mm²) — Tulangan tarik",  min_value=0.0, value=1520.0, step=10.0, format="%.0f",
                          help="Contoh: 4D22 = 4 × 380.1 = 1520 mm²")
    Asp = st.number_input("As' (mm²) — Tulangan tekan", min_value=0.0, value=0.0,    step=10.0, format="%.0f",
                          help="Isi 0 jika tidak ada tulangan tekan")

    st.markdown("")
    tombol = st.button("⚡ HITUNG KAPASITAS LENTUR", use_container_width=True, type="primary")

    with st.expander("📌 Referensi luas tulangan (mm²)"):
        st.markdown("""
        | Ø | 1 btg | 2 | 3 | 4 | 5 | 6 |
        |---|---|---|---|---|---|---|
        | D10 | 78.5 | 157 | 236 | 314 | 393 | 471 |
        | D13 | 132.7 | 265 | 398 | 531 | 663 | 796 |
        | D16 | 201.1 | 402 | 603 | 804 | 1005 | 1206 |
        | D19 | 283.5 | 567 | 851 | 1134 | 1418 | 1701 |
        | D22 | 380.1 | 760 | 1140 | 1520 | 1901 | 2281 |
        | D25 | 490.9 | 982 | 1473 | 1964 | 2454 | 2945 |
        | D29 | 660.5 | 1321 | 1981 | 2642 | 3302 | 3963 |
        | D32 | 804.2 | 1608 | 2413 | 3217 | 4021 | 4825 |
        """)

# ═══════════════════════════════════════════════════════════
# UI — HASIL
# ═══════════════════════════════════════════════════════════
with col_hasil:
    if tombol:
        if d >= h:
            st.error("⚠️ Tinggi efektif d harus lebih kecil dari tinggi total h!")
        elif As <= 0:
            st.error("⚠️ Luas tulangan tarik As harus lebih besar dari 0!")
        else:
            R     = hitung_lentur_balok(fc, fy, b, h, d, As, Asp)
            steps = buat_steps(fc, fy, b, h, d, As, Asp, R)

            # ── Metrik Utama ────────────────────────────────
            st.markdown("### 📊 Hasil Utama")
            m1, m2, m3, m4 = st.columns(4)
            for col, label, val, unit in [
                (m1, "Momen Nominal Mn",   f"{R['Mn']:.2f}",    "kN·m"),
                (m2, "φMn (Momen Rencana)",f"{R['phiMn']:.2f}", "kN·m"),
                (m3, "Faktor Reduksi φ",   f"{R['phi']:.3f}",   "—"),
                (m4, "Regangan εt",        f"{R['et']:.4f}",    "—"),
            ]:
                with col:
                    st.markdown(f"""<div class="metric-card">
                        <div class="metric-label">{label}</div>
                        <div class="metric-value">{val}</div>
                        <div class="metric-unit">{unit}</div></div>""",
                        unsafe_allow_html=True)
            st.markdown("")

            # ── Status ──────────────────────────────────────
            all_ok = R['ok_rho_min'] and R['ok_rho_max'] and R['ok_et']
            if all_ok and R['et'] >= 0.005:
                st.markdown('<div class="result-ok">✅ PENAMPANG OK — Tension-controlled, memenuhi seluruh syarat SNI 2847:2019</div>', unsafe_allow_html=True)
            elif all_ok:
                st.markdown('<div class="result-warn">⚠️ PERLU TINJAUAN — Zona transisi (0.004 ≤ εt &lt; 0.005)</div>', unsafe_allow_html=True)
            else:
                masalah = []
                if not R['ok_rho_min']: masalah.append("ρ &lt; ρ_min → tambah tulangan")
                if not R['ok_rho_max']: masalah.append("ρ > ρ_max → kurangi tulangan")
                if not R['ok_et']:      masalah.append("εt &lt; 0.004 → tidak memenuhi syarat")
                st.markdown(f'<div class="result-fail">❌ TIDAK OK — {" | ".join(masalah)}</div>', unsafe_allow_html=True)

            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)

            # ── Langkah-langkah ─────────────────────────────
            st.markdown("### 📐 Urutan Perhitungan")
            for s in steps:
                warna = "#2e7d32" if s['ok'] else "#c62828"
                tanda = "✓" if s['ok'] else "✗"
                st.markdown(f"""
                <div class="step-box" style="border-left-color:{warna}">
                    <div class="ref-badge">{s['ref']}</div><br>
                    <div class="step-header">{s['no']} — {s['judul']} &nbsp; {tanda}</div>
                    <pre style="margin:0;font-size:0.82rem;white-space:pre-wrap;font-family:monospace;">{s['isi']}</pre>
                </div>""", unsafe_allow_html=True)

            # ── Tabel Rangkuman ──────────────────────────────
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
            st.markdown("### 📋 Tabel Rangkuman")
            import pandas as pd
            et_sts = "✅ OK" if R['et'] >= 0.005 else ("⚠️ Transisi" if R['et'] >= 0.004 else "❌ Tidak OK")
            rh_sts = "✅ OK" if (R['ok_rho_min'] and R['ok_rho_max']) else "❌ Tidak OK"
            df = pd.DataFrame({
                "Parameter": ["β₁","a (mm)","c (mm)","εt","φ","ρ (%)","ρ_min (%)","ρ_max (%)","Mn (kN·m)","φMn (kN·m)"],
                "Nilai": [
                    f"{R['beta1']:.4f}", f"{R['a']:.2f}", f"{R['c']:.2f}",
                    f"{R['et']:.5f}", f"{R['phi']:.4f}", f"{R['rho']*100:.4f}",
                    f"{R['rho_min']*100:.4f}", f"{R['rho_max']*100:.4f}",
                    f"{R['Mn']:.3f}", f"{R['phiMn']:.3f}",
                ],
                "Status": ["—","—","—", et_sts,"—", rh_sts,"—","—","—","—"],
            })
            st.dataframe(df, use_container_width=True, hide_index=True)

            # ════════════════════════════════════════════════
            # TOMBOL DOWNLOAD WORD
            # ════════════════════════════════════════════════
            st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
            st.markdown("### 📄 Download Laporan")

            word_buf = generate_word(fc, fy, b, h, d, As, Asp, R, steps)
            nama_file = f"Laporan_Lentur_Balok_fc{int(fc)}_fy{int(fy)}_b{int(b)}x{int(h)}.docx"

            st.download_button(
                label="⬇️ Download Laporan Word (.docx)",
                data=word_buf,
                file_name=nama_file,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Laporan berisi data input, ringkasan hasil, dan 10 langkah perhitungan lengkap.",
            )
            st.caption("File siap dibuka di Microsoft Word atau LibreOffice Writer.")

    else:
        st.info("👈  Isi data di panel kiri, lalu klik **HITUNG KAPASITAS LENTUR**")
        st.markdown("""
        **Yang akan dihitung secara runtut:**
        1. Faktor β₁ (blok tegangan Whitney)
        2. Gaya tarik tulangan T
        3. Kedalaman blok tegangan **a**
        4. Kedalaman sumbu netral **c**
        5. Rasio tulangan aktual ρ
        6. Batas ρ_min (SNI Pasal 9.6.1.2)
        7. Batas ρ_max (0.75 ρ_bal)
        8. Regangan tarik εt
        9. Faktor reduksi φ
        10. Momen nominal **Mn** dan momen rencana **φMn**
        """)

# ── Footer ────────────────────────────────────────────────
st.markdown('<hr class="section-divider">', unsafe_allow_html=True)
st.markdown(
    "<p style='text-align:center;font-size:0.75rem;color:#aaa;'>"
    "Referensi: SNI 2847:2019 | ACI 318-14 | "
    "Untuk keperluan profesional — verifikasi mandiri tetap diperlukan"
    "</p>",
    unsafe_allow_html=True,
)
