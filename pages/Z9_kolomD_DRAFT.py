# =============================================================================
# FILE: app.py
# APPLICATION: Perhitungan Kapasitas Kolom Beton Bertulang - Diagram Interaksi P-M
# STANDAR: SNI 2847:2019
# ENGINEER DEFAULT: Ladosi
# =============================================================================
import streamlit as st
import pandas as pd
import numpy as np
import math
import matplotlib.pyplot as plt
import plotly.graph_objects as go
from io import BytesIO
import tempfile
import os

# -------------------- Report libraries --------------------
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
# -----------------------------------------------------------

# =============================================================================
# BAGIAN 1: LOGIKA KALKULASI (FUNGSI-FUNGSI TEKNIK)
# =============================================================================

# ---------- 1.1 Properti Penampang ----------
def hitung_properti(b, h, c, dia_s, D, n_b, n_h, fc):
    """
    Menghitung Ag, Ig, Ec, d, d', Ast, rasio tulangan
    SNI 2847:2019 Pasal 19.2.2, Pasal 10.6.1
    """
    Ag = b * h
    Ig = (1/12) * b * h**3
    Ec = 4700 * math.sqrt(fc)          # Pasal 19.2.2
    d_eff = h - c - dia_s - D/2         # tinggi efektif tarik
    d_prima = c + dia_s + D/2           # jarak tekan terluar
    # Susunan tulangan berdasarkan n_b dan n_h (sisi lebar & tinggi)
    n_total = 2 * n_b + 2 * (n_h - 2)
    A_per_bar = 0.25 * math.pi * D**2
    Ast = n_total * A_per_bar
    rho = (Ast / Ag) * 100              # persen
    return Ag, Ig, Ec, d_eff, d_prima, Ast, n_total, rho

def cek_rasio_tulangan(rho):
    """SNI 2847:2019 Pasal 10.6.1 : ρmin=1%, ρmax=8%"""
    ok = 1.0 <= rho <= 8.0
    return ok

# ---------- 1.2 Kekakuan Elemen & Faktor Ψ ----------
def kekakuan_EIL(b, h, L, Ec):
    """Menghitung EI/L (N.mm) untuk elemen balok/kolom"""
    if L == 0:
        return 0.0
    I_val = (1/12) * b * h**3
    return Ec * I_val / L

def psi_joint(EIL_kolom, EIL_balok):
    """
    Ψ = Σ(EI/L kolom) / Σ(EI/L balok)
    SNI 2847:2019 Pasal 6.6.4
    """
    sum_balok = sum(EIL_balok) if sum(EIL_balok) > 1e-12 else 1e-12
    return sum(EIL_kolom) / sum_balok

# ---------- 1.3 Faktor panjang efektif k ----------
def hitung_k_braced(psiA, psiB):
    """SNI 2847:2019 Pasal 6.6.4 : braced frame"""
    k1 = 0.7 + 0.05 * (psiA + psiB)
    k2 = 0.85 + 0.05 * min(psiA, psiB)
    k = min(k1, k2, 1.0)
    return k

def hitung_k_unbraced(psiA, psiB):
    """SNI 2847:2019 Pasal 6.6.4 : unbraced frame"""
    psi_m = 0.5 * (psiA + psiB)
    if psi_m < 2:
        k = ((20 - psi_m) / 20) * math.sqrt(1 + psi_m)
    else:
        k = 0.9 * math.sqrt(1 + psi_m)
    return k

# ---------- 1.4 Kelangsingan Kolom ----------
def cek_kelangsingan(b, h, Lu, k, M1, M2, frame_type):
    """
    SNI 2847:2019 Pasal 6.2.5
    Klasifikasi Short / Slender
    M1 dan M2 sudah bertanda (single/double curvature)
    """
    r = 0.3 * h                           # radius girasi persegi
    kLu = k * Lu
    rasio = kLu / r
    if frame_type == 'Braced':
        batas = 34 - 12 * (M1 / M2)        # M1/M2 dengan tanda
    else:  # Unbraced
        batas = 22
    klas = "Short Column" if rasio <= batas else "Slender Column"
    return r, kLu, rasio, batas, klas

# ---------- 1.5 Pembesaran Momen (Braced Frame) ----------
def pembesaran_momen_braced(b, h, Lu, k, Ec, Ig, Pu, M1, M2, beta_dns):
    """
    SNI 2847:2019 Pasal 6.6.4.4 & 6.6.4.5
    Menghitung δns dan momen desain Mc (kN.m)
    """
    # Cm = 0.6 + 0.4*(M1/M2) ≥ 0.4
    Cm = 0.6 + 0.4 * (M1 / M2)
    if Cm < 0.4:
        Cm = 0.4
    # EI efektif (Pers. 6.6.4.4.4)
    EI_eff = 0.4 * Ec * Ig / (1 + beta_dns)       # N.mm²
    kLu = k * Lu
    Pc = (math.pi**2 * EI_eff) / (kLu**2)         # N → kN
    Pc_kN = Pc / 1000.0
    # δns = Cm / (1 - Pu/(0.75*Pc)) ≥ 1.0
    denominator = 1.0 - Pu / (0.75 * Pc_kN)
    if denominator <= 0:
        delta_ns = 1e6   # tak hingga (tanda kolom sangat langsing)
    else:
        delta_ns = Cm / denominator
        if delta_ns < 1.0:
            delta_ns = 1.0
    Mc = delta_ns * M2
    return Cm, EI_eff, Pc_kN, delta_ns, Mc

# ---------- 1.6 Diagram Interaksi P-M ----------
def generate_interaction(b, h, fc, fy, Es, beta1, cover, dia_s, D, n_b, n_h):
    """
    Menghasilkan titik-titik diagram interaksi φPn - φMn (kN, kN.m)
    50 titik c/h + tekan murni + tarik murni
    SNI 2847:2019 Pasal 22.2.2
    """
    Ag = b * h
    d_prime = cover + dia_s + D/2
    d_eff = h - d_prime
    # susunan layer
    layers = [
        {"y": d_prime, "n": n_b},                               # tekan
        {"y": h/2, "n": 2 * (n_h - 2)},                         # tengah
        {"y": h - d_prime, "n": n_b}                            # tarik
    ]
    A_bar = 0.25 * math.pi * D**2
    for lyr in layers:
        lyr["As"] = lyr["n"] * A_bar

    # Titik-titik c/h (50 titik dari 0.02 sd 1.20)
    titik_ch = np.linspace(0.02, 1.20, 50)
    points = []

    for ch in titik_ch:
        c_val = ch * h
        a = beta1 * c_val
        if a > h:
            a = h                                                # batas fisik
        # Gaya beton
        Cc = 0.85 * fc * a * b
        # Resultan baja
        Fs_list = []
        M_list = []
        for lyr in layers:
            yi = lyr["y"]
            eps_si = 0.003 * (c_val - yi) / c_val
            fs_i = Es * eps_si
            if fs_i > fy:
                fs_i = fy
            elif fs_i < -fy:
                fs_i = -fy
            Fi = lyr["As"] * fs_i
            Fs_list.append(Fi)
            M_list.append(Fi * (h/2 - yi))
        # Momen dari beton
        M_beton = Cc * (h/2 - a/2)
        Pn = Cc + sum(Fs_list)
        Mn = M_beton + sum(M_list)           # N.mm
        # Strain baja ekstrim tarik (layer bawah)
        eps_t = 0.003 * (c_val - layers[-1]["y"]) / c_val
        # Faktor reduksi φ
        if eps_t >= 0.005:
            phi = 0.90
        elif eps_t <= 0.002:
            phi = 0.65
        else:
            phi = 0.65 + 83.333 * (eps_t - 0.002)   # interpolasi linier (0.25/0.003)
        if phi < 0.65:
            phi = 0.65
        if phi > 0.90:
            phi = 0.90
        points.append({
            "c/h": ch, "c (mm)": c_val, "a (mm)": a,
            "εt": eps_t, "Pn (kN)": Pn/1000, "Mn (kN.m)": Mn/1e6,
            "φ": phi, "φPn (kN)": phi*Pn/1000, "φMn (kN.m)": phi*Mn/1e6
        })

    # Titik tekan murni (Po)
    Po = 0.85 * fc * (Ag - sum(lyr["As"] for lyr in layers)) + fy * sum(lyr["As"] for lyr in layers)
    phi_po = 0.65
    points.append({
        "c/h": "∞", "c (mm)": "∞", "a (mm)": h,
        "εt": 0.0, "Pn (kN)": Po/1000, "Mn (kN.m)": 0.0,
        "φ": phi_po, "φPn (kN)": phi_po*Po/1000, "φMn (kN.m)": 0.0
    })
    # Titik tarik murni (To)
    To = -fy * sum(lyr["As"] for lyr in layers)
    phi_to = 0.90
    points.append({
        "c/h": "0", "c (mm)": 0, "a (mm)": 0,
        "εt": 0.0, "Pn (kN)": To/1000, "Mn (kN.m)": 0.0,
        "φ": phi_to, "φPn (kN)": phi_to*To/1000, "φMn (kN.m)": 0.0
    })

    return points, layers

def cek_kapasitas(points, Pu, Mu):
    """
    Mengecek apakah (Pu (kN), Mu (kN.m)) berada di dalam kurva φPn-φMn
    Menggunakan interpolasi pada φPn untuk mendapatkan φMn kapasitas
    Kembalikan (status, φMn_cap, ratio)
    """
    # Filter titik non-spesial (c/h numerik) untuk interpolasi
    data = [p for p in points if isinstance(p["c/h"], float)]
    data_sorted = sorted(data, key=lambda x: x["φPn (kN)"])  # ascending
    P_vals = np.array([d["φPn (kN)"] for d in data_sorted])
    M_vals = np.array([d["φMn (kN.m)"] for d in data_sorted])
    # Jika Pu di luar rentang
    if Pu < P_vals[0] or Pu > P_vals[-1]:
        return "OVER - beban di luar kurva", 0.0, float('inf')
    # Interpolasi linier
    idx = np.searchsorted(P_vals, Pu)
    if idx == 0:
        Mcap = M_vals[0]
    elif idx == len(P_vals):
        Mcap = M_vals[-1]
    else:
        x0, x1 = P_vals[idx-1], P_vals[idx]
        y0, y1 = M_vals[idx-1], M_vals[idx]
        Mcap = y0 + (y1 - y0) * (Pu - x0) / (x1 - x0)
    ratio = abs(Mu) / Mcap if Mcap != 0 else 1e12
    status = "OK - AMAN" if ratio <= 1.0 else "NOT OK - OVER"
    return status, Mcap, ratio

# =============================================================================
# BAGIAN 2: UI STREAMLIT
# =============================================================================
st.set_page_config(page_title="Kapasitas Kolom Beton - SNI 2847:2019", layout="wide")
st.title("🔧 Kalkulator Kapasitas Kolom Beton Bertulang (Diagram Interaksi P-M)")
st.markdown("**Standar : SNI 2847:2019 | Engineer : Ladosi**")

# Initialize session state for results
if "hasil" not in st.session_state:
    st.session_state.hasil = None

# ---------- SIDEBAR INPUT ----------
st.sidebar.header("📐 DATA INPUT")

with st.sidebar.expander("1. Material", expanded=False):
    fc = st.number_input("Mutu Beton fc' (MPa)", value=30, min_value=15, max_value=55, step=5)
    fy = st.number_input("Mutu Baja fy (MPa)", value=400, min_value=240, max_value=600, step=10)
    Es = st.number_input("Modulus Elastisitas Baja Es (MPa)", value=200000, step=1000)

with st.sidebar.expander("2. Dimensi Kolom", expanded=False):
    b = st.number_input("Lebar b (mm)", value=400, min_value=200, step=50)
    h = st.number_input("Tinggi h (mm)", value=500, min_value=200, step=50)
    c = st.number_input("Selimut Beton c (mm)", value=40, min_value=20, step=5)
    dia_s = st.number_input("Diameter Sengkang Øs (mm)", value=10, min_value=6, step=2)

with st.sidebar.expander("3. Tulangan Longitudinal", expanded=False):
    D = st.number_input("Diameter Tulangan D (mm)", value=22, min_value=10, step=2)
    n_b = st.number_input("Jumlah Tulangan Sisi Lebar n_b", value=4, min_value=2, step=1)
    n_h = st.number_input("Jumlah Tulangan Sisi Tinggi n_h", value=3, min_value=2, step=1)

with st.sidebar.expander("4. Panjang & Rangka", expanded=False):
    Lu = st.number_input("Panjang Tak Tertahan Lu (mm)", value=6000, min_value=1000, step=500)
    frame_type = st.radio("Kondisi Rangka", ["Braced", "Unbraced"], index=0)
    curv_type = st.radio("Kelengkungan", ["Single", "Double"], index=0)
    beta_dns = st.number_input("Rasio βdns (beban tetap/total)", value=0.6, min_value=0.0, max_value=1.0, step=0.05)

with st.sidebar.expander("5. Kekakuan Elemen Penghubung", expanded=False):
    st.markdown("**Balok Atas**")
    balok_atas_kiri = st.checkbox("Balok Atas Kiri", value=True)
    if balok_atas_kiri:
        b_bal_atas = st.number_input("Lebar Balok Atas (mm)", value=300, step=10)
        h_bal_atas = st.number_input("Tinggi Balok Atas (mm)", value=500, step=10)
        L_bal_atas = st.number_input("Panjang Balok Atas (mm)", value=6000, step=500)
    else:
        b_bal_atas, h_bal_atas, L_bal_atas = 0,0,0
    balok_atas_kanan = st.checkbox("Balok Atas Kanan", value=True)
    if balok_atas_kanan:
        b_bal_atas_k = st.number_input("Lebar Balok Atas Kanan (mm)", value=300, step=10)
        h_bal_atas_k = st.number_input("Tinggi Balok Atas Kanan (mm)", value=500, step=10)
        L_bal_atas_k = st.number_input("Panjang Balok Atas Kanan (mm)", value=6000, step=500)
    else:
        b_bal_atas_k, h_bal_atas_k, L_bal_atas_k = 0,0,0

    st.markdown("**Balok Bawah**")
    balok_bawah_kiri = st.checkbox("Balok Bawah Kiri", value=True)
    if balok_bawah_kiri:
        b_bal_bawah = st.number_input("Lebar Balok Bawah (mm)", value=300, step=10)
        h_bal_bawah = st.number_input("Tinggi Balok Bawah (mm)", value=500, step=10)
        L_bal_bawah = st.number_input("Panjang Balok Bawah (mm)", value=6000, step=500)
    else:
        b_bal_bawah, h_bal_bawah, L_bal_bawah = 0,0,0
    balok_bawah_kanan = st.checkbox("Balok Bawah Kanan", value=True)
    if balok_bawah_kanan:
        b_bal_bawah_k = st.number_input("Lebar Balok Bawah Kanan (mm)", value=300, step=10)
        h_bal_bawah_k = st.number_input("Tinggi Balok Bawah Kanan (mm)", value=500, step=10)
        L_bal_bawah_k = st.number_input("Panjang Balok Bawah Kanan (mm)", value=6000, step=500)
    else:
        b_bal_bawah_k, h_bal_bawah_k, L_bal_bawah_k = 0,0,0

    st.markdown("**Kolom Atas & Bawah**")
    kolom_atas_ada = st.checkbox("Ada Kolom Atas", value=True)
    if kolom_atas_ada:
        b_kol_atas = st.number_input("Lebar Kolom Atas (mm)", value=400, step=10)
        h_kol_atas = st.number_input("Tinggi Kolom Atas (mm)", value=500, step=10)
        L_kol_atas = st.number_input("Panjang Kolom Atas (mm)", value=3500, step=500)
    else:
        b_kol_atas, h_kol_atas, L_kol_atas = 0,0,0
    # Kondisi bawah
    jenis_bawah = st.radio("Kondisi Ujung Bawah", ["Kolom Bawah", "Pondasi (Fixed)", "Bebas (Hinge)"], index=0)
    if jenis_bawah == "Kolom Bawah":
        b_kol_bawah = st.number_input("Lebar Kolom Bawah (mm)", value=400, step=10)
        h_kol_bawah = st.number_input("Tinggi Kolom Bawah (mm)", value=500, step=10)
        L_kol_bawah = st.number_input("Panjang Kolom Bawah (mm)", value=3500, step=500)
    else:
        b_kol_bawah = h_kol_bawah = L_kol_bawah = 0

with st.sidebar.expander("6. Beban Terfaktor", expanded=False):
    Pu = st.number_input("Gaya Aksial Pu (kN)", value=1500.0, step=100.0)
    M1_abs = st.number_input("Momen Ujung M1 (kN.m) (lebih kecil, absolut)", value=80.0, step=10.0)
    M2_abs = st.number_input("Momen Ujung M2 (kN.m) (lebih besar, absolut)", value=150.0, step=10.0)

# Tombol Kalkulasi
if st.sidebar.button("⚡ HITUNG KAPASITAS", use_container_width=True):
    # ---------- PROSES KALKULASI ----------
    try:
        # 1. Properti penampang
        Ag, Ig, Ec, d_eff, d_prima, Ast, n_total, rho = hitung_properti(b, h, c, dia_s, D, n_b, n_h, fc)
        ok_rasio = cek_rasio_tulangan(rho)

        # 2. Kekakuan EIL
        EIL_kolom_design = kekakuan_EIL(b, h, Lu, Ec)
        EIL_kol_atas = kekakuan_EIL(b_kol_atas, h_kol_atas, L_kol_atas, Ec) if kolom_atas_ada else 0.0
        EIL_kol_bawah = 0.0
        if jenis_bawah == "Kolom Bawah":
            EIL_kol_bawah = kekakuan_EIL(b_kol_bawah, h_kol_bawah, L_kol_bawah, Ec)
        # Balok atas
        EIL_bal_atas_kiri = kekakuan_EIL(b_bal_atas, h_bal_atas, L_bal_atas, Ec) if balok_atas_kiri else 0.0
        EIL_bal_atas_kanan = kekakuan_EIL(b_bal_atas_k, h_bal_atas_k, L_bal_atas_k, Ec) if balok_atas_kanan else 0.0
        EIL_bal_bawah_kiri = kekakuan_EIL(b_bal_bawah, h_bal_bawah, L_bal_bawah, Ec) if balok_bawah_kiri else 0.0
        EIL_bal_bawah_kanan = kekakuan_EIL(b_bal_bawah_k, h_bal_bawah_k, L_bal_bawah_k, Ec) if balok_bawah_kanan else 0.0

        # Ψ atas
        if kolom_atas_ada:
            EIL_kol_atas_sum = EIL_kol_atas + EIL_kolom_design
        else:
            EIL_kol_atas_sum = EIL_kolom_design   # hanya kolom yang ditinjau
        psiA = psi_joint([EIL_kol_atas_sum], [EIL_bal_atas_kiri + EIL_bal_atas_kanan])

        # Ψ bawah
        if jenis_bawah == "Pondasi (Fixed)":
            psiB = 0.0
        elif jenis_bawah == "Bebas (Hinge)":
            psiB = 10.0
        else:  # Kolom Bawah
            EIL_kol_bawah_sum = EIL_kol_bawah + EIL_kolom_design
            psiB = psi_joint([EIL_kol_bawah_sum], [EIL_bal_bawah_kiri + EIL_bal_bawah_kanan])

        # Faktor k
        if frame_type == "Braced":
            k = hitung_k_braced(psiA, psiB)
        else:
            k = hitung_k_unbraced(psiA, psiB)

        # Momen bertanda sesuai kelengkungan
        sign = 1.0 if curv_type == "Single" else -1.0
        M1 = M1_abs * sign
        M2 = M2_abs * sign

        # Kelangsingan
        r, kLu, rasio_kel, batas, klas = cek_kelangsingan(b, h, Lu, k, M1, M2, frame_type)

        # Pembesaran momen (hanya jika braced dan slender)
        if frame_type == "Braced" and klas == "Slender Column":
            Cm, EI_eff, Pc, delta_ns, Mc = pembesaran_momen_braced(b, h, Lu, k, Ec, Ig, Pu, M1, M2, beta_dns)
        else:
            # Tidak ada pembesaran (short atau unbraced tidak dibahas di sini)
            Cm, EI_eff, Pc, delta_ns, Mc = None, None, None, 1.0, M2

        # Diagram Interaksi
        beta1 = 0.85 - max(0, (fc - 28) / 7) * 0.05
        titik_interaksi, layers = generate_interaction(b, h, fc, fy, Es, beta1, c, dia_s, D, n_b, n_h)

        # Cek kapasitas
        status_kap, Mcap, ratio_cap = cek_kapasitas(titik_interaksi, Pu, Mc if Mc is not None else M2)

        # Simpan di session
        hasil = {
            "Ag": Ag, "Ig": Ig, "Ec": Ec, "d": d_eff, "d'": d_prima,
            "Ast": Ast, "n_total": n_total, "rho": rho, "ok_rasio": ok_rasio,
            "EIL_kol_design": EIL_kolom_design,
            "psiA": psiA, "psiB": psiB, "k": k,
            "r": r, "kLu": kLu, "rasio_kel": rasio_kel, "batas": batas, "klas": klas,
            "Cm": Cm, "EI_eff": EI_eff, "Pc": Pc, "delta_ns": delta_ns, "Mc": Mc,
            "titik_interaksi": titik_interaksi,
            "status_kap": status_kap, "Mcap": Mcap, "ratio_cap": ratio_cap,
            "Pu": Pu, "Mu_desain": Mc if Mc is not None else M2,
            "fc": fc, "fy": fy, "Es": Es, "b": b, "h": h, "c": c, "dia_s": dia_s,
            "D": D, "n_b": n_b, "n_h": n_h, "Lu": Lu, "frame_type": frame_type,
            "curv_type": curv_type, "beta_dns": beta_dns,
            "M1_abs": M1_abs, "M2_abs": M2_abs,
        }
        st.session_state.hasil = hasil
        st.success("✅ Perhitungan selesai! Lihat hasil di bawah.")
    except Exception as e:
        st.error(f"Terjadi kesalahan: {str(e)}")

# Tampilkan hasil jika sudah dihitung
hasil = st.session_state.hasil
if hasil is not None:
    # ---------- TABEL HASIL ----------
    st.header("📊 HASIL PERHITUNGAN STEP‑BY‑STEP (SNI 2847:2019)")

    with st.expander("1. PROPERTIES PENAMPANG", expanded=True):
        st.markdown(f"""
        - Luas Bruto Ag = b × h = {hasil['b']} × {hasil['h']} = **{hasil['Ag']:,.2f} mm²**
        - Momen Inersia Ig = (1/12)×b×h³ = (1/12)×{hasil['b']}×{hasil['h']}³ = **{hasil['Ig']:,.2f} mm⁴**
        - Modulus Elastisitas Beton Ec = 4700×√fc' = 4700×√{hasil['fc']} = **{hasil['Ec']:,.2f} MPa** (SNI 2847:2019 Pasal 19.2.2)
        - Tinggi Efektif d = h - c - Øs - D/2 = {hasil['h']} - {hasil['c']} - {hasil['dia_s']} - {hasil['D']}/2 = **{hasil['d']:.2f} mm**
        - Jarak Tulangan Tekan Terluar d' = c + Øs + D/2 = {hasil['c']} + {hasil['dia_s']} + {hasil['D']}/2 = **{hasil["d'"]:.2f} mm**
        - Total Tulangan = 2×n_b + 2×(n_h‑2) = 2×{hasil['n_b']} + 2×({hasil['n_h']}‑2) = **{hasil['n_total']} buah**
        - Luas Tulangan Ast = n_total × π×D²/4 = {hasil['n_total']} × π×{hasil['D']}²/4 = **{hasil['Ast']:,.2f} mm²**
        - Rasio Tulangan ρg = Ast/Ag × 100 = **{hasil['rho']:.2f} %** (Syarat 1%‑8% : {'✅ OK' if hasil['ok_rasio'] else '❌ NOT OK'})
        """)

    with st.expander("2. FAKTOR KEKAUKAN Ψ & PANJANG EFEKTIF k", expanded=True):
        st.markdown(f"""
        - EI/L Kolom Ditinjau = {hasil['EIL_kol_design']:,.2f} N.mm  
        - Ψ Atas = {hasil['psiA']:.2f}, Ψ Bawah = {hasil['psiB']:.2f}  
        - Faktor k ({hasil['frame_type']}) = **{hasil['k']:.4f}** (SNI 2847:2019 Pasal 6.6.4)
        """)

    with st.expander("3. KELANGSINGAN KOLOM", expanded=True):
        st.markdown(f"""
        - Radius Girasi r = 0.3 × h = 0.3 × {hasil['h']} = **{hasil['r']:.2f} mm**
        - Panjang Efektif k×Lu = {hasil['kLu']:.2f} mm  
        - Rasio Kelangsingan k×Lu/r = **{hasil['rasio_kel']:.2f}**  
        - Batas ({hasil['frame_type']}) = {hasil['batas']:.2f}  
        - Klasifikasi : **{hasil['klas']}**
        """)

    if hasil['klas'] == "Slender Column" and hasil['frame_type'] == "Braced":
        with st.expander("4. PEMBESARAN MOMEN (Braced Slender)", expanded=True):
            st.markdown(f"""
            - Cm = 0.6 + 0.4×(M1/M2) = 0.6 + 0.4×({hasil['M1_abs']*sign:.2f}/{hasil['M2_abs']:.2f}) = **{hasil['Cm']:.2f}**  
            - EI efektif = 0.4×Ec×Ig/(1+βdns) = **{hasil['EI_eff']:,.2f} N.mm²**  
            - Beban Kritis Euler Pc = π²×EI/(kLu)² = **{hasil['Pc']:,.2f} kN**  
            - δns = Cm / (1 - Pu/(0.75×Pc)) = **{hasil['delta_ns']:.2f}**  
            - Momen Desain Mc = δns × M2 = {hasil['delta_ns']:.2f} × {hasil['M2_abs']} = **{hasil['Mc']:.2f} kN.m**
            """)

    with st.expander("5. DIAGRAM INTERAKSI P‑M (52 TITIK)", expanded=True):
        df = pd.DataFrame(hasil['titik_interaksi'])
        st.dataframe(df.style.format(precision=2), use_container_width=True)

    # Plot interaktif dengan Plotly
    st.subheader("📈 Kurva Interaksi φPn – φMn")
    points = hasil['titik_interaksi']
    P_plot = [p["φPn (kN)"] for p in points]
    M_plot = [p["φMn (kN.m)"] for p in points]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=M_plot, y=P_plot,
        mode='lines+markers', name='Kurva φPn-φMn',
        line=dict(color='blue')
    ))
    fig.add_trace(go.Scatter(
        x=[hasil['Mu_desain']], y=[hasil['Pu']],
        mode='markers', marker=dict(color='red', size=12),
        name='Beban (Pu, Mu)'
    ))
    fig.update_layout(
        title="Diagram Interaksi Kolom",
        xaxis_title="φMn (kN.m)",
        yaxis_title="φPn (kN)",
        legend=dict(x=0.01, y=0.99),
    )
    st.plotly_chart(fig, use_container_width=True)

    # Status kapasitas
    st.subheader("✅ CEK KAPASITAS")
    st.markdown(f"""
    - Status : **{hasil['status_kap']}**
    - Rasio Mu/φMn = {hasil['Mu_desain']:.2f} / {hasil['Mcap']:.2f} = **{hasil['ratio_cap']:.2f}**
    """)

    # Tombol unduh laporan
    col1, col2 = st.columns(2)
    with col1:
        if st.button("📄 Unduh Laporan Word (.docx)"):
            try:
                doc_bytes = create_word_report(hasil)
                st.download_button(
                    label="Download Laporan Word",
                    data=doc_bytes,
                    file_name="Laporan_Kolom_Beton.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Gagal membuat laporan Word: {e}")
    with col2:
        if st.button("📄 Unduh Laporan PDF (.pdf)"):
            try:
                pdf_bytes = create_pdf_report(hasil)
                st.download_button(
                    label="Download Laporan PDF",
                    data=pdf_bytes,
                    file_name="Laporan_Kolom_Beton.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error(f"Gagal membuat laporan PDF: {e}")

# =============================================================================
# BAGIAN 3: GENERATOR LAPORAN (WORD & PDF)
# =============================================================================

@st.cache_data
def create_word_report(data):
    """Membuat laporan Word (.docx) terformat sebagai Engineering Report."""
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    # Header
    doc.add_heading('LAPORAN PERHITUNGAN KAPASITAS KOLOM BETON BERTULANG', 0)
    doc.add_paragraph(f"Standar: SNI 2847:2019\nEngineer: Ladosi\nTanggal: {pd.Timestamp.now().strftime('%d/%m/%Y')}")

    # BAB 1 - Data Input
    doc.add_heading('BAB 1: DATA INPUT PARAMETER', level=1)
    inputs = [
        ("Material", [
            ("Mutu Beton fc'", f"{data['fc']} MPa"),
            ("Mutu Baja fy", f"{data['fy']} MPa"),
            ("Modulus Elastisitas Baja Es", f"{data['Es']} MPa"),
        ]),
        ("Dimensi Kolom", [
            ("Lebar b", f"{data['b']} mm"),
            ("Tinggi h", f"{data['h']} mm"),
            ("Selimut Beton c", f"{data['c']} mm"),
            ("Diameter Sengkang Øs", f"{data['dia_s']} mm"),
        ]),
        ("Tulangan Longitudinal", [
            ("Diameter D", f"{data['D']} mm"),
            ("Jumlah Sisi Lebar n_b", str(data['n_b'])),
            ("Jumlah Sisi Tinggi n_h", str(data['n_h'])),
        ]),
        ("Panjang & Rangka", [
            ("Panjang Lu", f"{data['Lu']} mm"),
            ("Tipe Rangka", data['frame_type']),
            ("Kelengkungan", data['curv_type']),
            ("βdns", str(data['beta_dns'])),
        ]),
        ("Beban", [
            ("Pu", f"{data['Pu']} kN"),
            ("M1 (abs)", f"{data['M1_abs']} kN.m"),
            ("M2 (abs)", f"{data['M2_abs']} kN.m"),
        ]),
    ]
    for heading, items in inputs:
        doc.add_heading(heading, level=2)
        table = doc.add_table(rows=len(items), cols=2, style='Light Grid Accent 1')
        for i, (param, val) in enumerate(items):
            table.rows[i].cells[0].text = param
            table.rows[i].cells[1].text = val

    # BAB 2 - Step-by-step
    doc.add_heading('BAB 2: PROSES PERHITUNGAN (STEP‑BY‑STEP)', level=1)

    # Properties
    doc.add_heading('2.1 Properti Penampang', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Ag = b × h = {data['b']} × {data['h']} = {data['Ag']:,.2f} mm²\n")
    p.add_run(f"Ig = (1/12)×b×h³ = (1/12)×{data['b']}×{data['h']}³ = {data['Ig']:,.2f} mm⁴\n")
    p.add_run(f"Ec = 4700×√fc' = 4700×√{data['fc']} = {data['Ec']:,.2f} MPa\n")
    p.add_run(f"d = {data['d']:.2f} mm\n")
    p.add_run(f"d' = {data['d_prime']:.2f} mm\n")
    p.add_run(f"Ast = {data['Ast']:,.2f} mm²\n")
    p.add_run(f"ρg = {data['rho']:.2f}% (OK)" if data['ok_rasio'] else f"ρg = {data['rho']:.2f}% (NOT OK)")

    # Kelangsingan & momen
    doc.add_heading('2.2 Kelangsingan & Pembesaran Momen', level=2)
    p = doc.add_paragraph()
    p.add_run(f"Ψ atas = {data['psiA']:.2f}, Ψ bawah = {data['psiB']:.2f}\n")
    p.add_run(f"k = {data['k']:.4f}\n")
    p.add_run(f"r = 0.3×h = {data['r']:.2f} mm\n")
    p.add_run(f"kLu/r = {data['rasio_kel']:.2f} (Batas {data['batas']:.2f}) → {data['klas']}\n")
    if data['Cm'] is not None:
        p.add_run(f"Cm = {data['Cm']:.2f}\n")
        p.add_run(f"Pc = {data['Pc']:.2f} kN\n")
        p.add_run(f"δns = {data['delta_ns']:.2f}\n")
        p.add_run(f"Mc = δns×M2 = {data['Mc']:.2f} kN.m\n")
    else:
        p.add_run("Tidak ada pembesaran momen.\n")

    # Diagram Interaksi
    doc.add_heading('2.3 Diagram Interaksi P-M', level=2)
    p = doc.add_paragraph(f"β1 = {0.85 - max(0, (data['fc'] - 28) / 7) * 0.05:.4f}\n")
    p.add_run("Iterasi 52 titik dilakukan. (Lihat lampiran tabel)\n")

    # BAB 3 - Kesimpulan
    doc.add_heading('BAB 3: KESIMPULAN & HASIL AKHIR', level=1)
    p = doc.add_paragraph()
    p.add_run(f"Pu = {data['Pu']} kN\n")
    p.add_run(f"Momen desain = {data['Mu_desain']:.2f} kN.m\n")
    p.add_run(f"Kapasitas φMn pada Pu = {data['Mcap']:.2f} kN.m\n")
    p.add_run(f"Rasio = {data['ratio_cap']:.2f} → {data['status_kap']}\n")

    # Simpan ke BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# Sanitasi karakter khusus untuk PDF
def sanitize_pdf(text):
    """Mengganti karakter unicode yang tidak didukung latin-1."""
    replacements = {
        'φ': 'phi', '²': '^2', '³': '^3', '√': 'sqrt',
        '·': '*', 'Ø': 'Phi', '≤': '<=', '≥': '>=',
        '°': 'deg', '′': "'", '″': "''",
        'ρ': 'rho', 'δ': 'delta', 'ε': 'epsilon',
        'β': 'beta', 'ψ': 'psi', '∞': 'inf',
    }
    for uni, latin in replacements.items():
        text = text.replace(uni, latin)
    return text

@st.cache_data
def create_pdf_report(data):
    """Membuat laporan PDF dengan FPDF2."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Times", size=12)
    # Judul
    pdf.set_font("Times", 'B', 16)
    pdf.cell(0, 10, "LAPORAN PERHITUNGAN KAPASITAS KOLOM BETON BERTULANG", ln=True, align='C')
    pdf.set_font("Times", size=12)
    pdf.cell(0, 10, "Standar: SNI 2847:2019 | Engineer: Ladosi", ln=True, align='C')
    pdf.ln(5)

    # BAB 1
    pdf.set_font("Times", 'B', 14)
    pdf.cell(0, 10, "BAB 1: DATA INPUT PARAMETER", ln=True)
    pdf.set_font("Times", size=12)
    inputs = [
        f"fc' = {data['fc']} MPa",
        f"fy = {data['fy']} MPa",
        f"Es = {data['Es']} MPa",
        f"b = {data['b']} mm, h = {data['h']} mm",
        f"c = {data['c']} mm, D sengkang = {data['dia_s']} mm",
        f"D = {data['D']} mm, n_b = {data['n_b']}, n_h = {data['n_h']}",
        f"Lu = {data['Lu']} mm, Rangka = {data['frame_type']}",
        f"βdns = {data['beta_dns']}",
        f"Pu = {data['Pu']} kN, M1 = {data['M1_abs']} kN.m, M2 = {data['M2_abs']} kN.m",
    ]
    for line in inputs:
        pdf.cell(0, 8, sanitize_pdf("- " + line), ln=True)

    pdf.ln(5)
    pdf.set_font("Times", 'B', 14)
    pdf.cell(0, 10, "BAB 2: PROSES PERHITUNGAN", ln=True)
    pdf.set_font("Times", size=12)

    step_text = [
        f"Ag = {data['Ag']:,.2f} mm^2",
        f"Ig = {data['Ig']:,.2f} mm^4",
        f"Ec = {data['Ec']:,.2f} MPa",
        f"d = {data['d']:.2f} mm, d' = {data['d_prime']:.2f} mm",
        f"Ast = {data['Ast']:,.2f} mm^2, rho = {data['rho']:.2f}%",
        f"psiA = {data['psiA']:.2f}, psiB = {data['psiB']:.2f}",
        f"k = {data['k']:.4f}",
        f"r = {data['r']:.2f} mm",
        f"kLu/r = {data['rasio_kel']:.2f} (batas {data['batas']:.2f}) -> {data['klas']}",
    ]
    if data['Cm'] is not None:
        step_text += [
            f"Cm = {data['Cm']:.2f}",
            f"Pc = {data['Pc']:.2f} kN",
            f"delta_ns = {data['delta_ns']:.2f}",
            f"Mc = {data['Mc']:.2f} kN.m",
        ]
    for t in step_text:
        pdf.cell(0, 8, sanitize_pdf(t), ln=True)

    # Kurva interaksi
    pdf.ln(5)
    pdf.set_font("Times", 'B', 14)
    pdf.cell(0, 10, "BAB 3: KESIMPULAN", ln=True)
    pdf.set_font("Times", size=12)
    pdf.cell(0, 8, sanitize_pdf(f"Pu = {data['Pu']} kN, Mu desain = {data['Mu_desain']:.2f} kN.m"), ln=True)
    pdf.cell(0, 8, sanitize_pdf(f"Kapasitas: phiMn = {data['Mcap']:.2f} kN.m, Rasio = {data['ratio_cap']:.2f}"), ln=True)
    pdf.cell(0, 8, sanitize_pdf(f"Status: {data['status_kap']}"), ln=True)

    # Simpan sebagai bytes
    pdf_out = pdf.output(dest='S')  # string bytes
    return pdf_out if isinstance(pdf_out, bytes) else pdf_out.encode('latin-1')
