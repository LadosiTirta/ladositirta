"""
================================================================================
app.py  —  Streamlit UI  —  Program Perhitungan Turap / Sheet Pile
================================================================================
Antarmuka web untuk perhitungan turap (sheet pile) secara lengkap.

Modul yang dipanggil (semua dari folder yang sama):
    earth_pressure.py    — tekanan tanah lateral
    stability.py         — analisis stabilitas, SF heave, piping
    internal_forces.py   — gaya dalam V dan M
    section_design.py    — desain penampang baja / beton pracetak
    anchor_design.py     — desain angkur, waling, deadman, strut

Cara menjalankan:
    streamlit run app.py

Referensi standar: lihat masing-masing modul perhitungan.
================================================================================
"""

from __future__ import annotations

import io
import os
import sys
import math
import datetime
import traceback

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import streamlit as st

# Tambahkan folder script ke path
_DIR = os.path.dirname(os.path.abspath(__file__))
if _DIR not in sys.path:
    sys.path.insert(0, _DIR)

# Impor modul perhitungan
from earth_pressure import (
    hitung_Ka_rankine, hitung_Kp_rankine, hitung_Ko,
    hitung_distribusi_tekanan, plot_diagram_tekanan,
    format_langkah as ep_fmt,
)
from stability import (
    free_earth_support, fixed_earth_support_cantilever,
    hitung_SF_heave_terzaghi, hitung_SF_heave_bjerrum_eide,
    hitung_SF_piping, rangkuman_stabilitas,
    format_langkah as st_fmt,
    SF_GULING_MIN, SF_HEAVE_MIN, SF_PIPING_MIN,
)
from internal_forces import (
    hitung_gaya_dalam, hitung_gaya_dalam_kantilever,
    plot_V_M_diagram, format_langkah as if_fmt,
)
from section_design import (
    desain_baja, desain_beton_pracetak,
    format_langkah as sd_fmt,
    MUTU_BAJA_DEFAULT,
)
from anchor_design import (
    desain_tie_rod, desain_waling, desain_deadman, desain_strut,
    rangkuman_angkur, format_langkah as ad_fmt,
    MUTU_BAJA, PROFIL_WF,
)


# ─────────────────────────────────────────────────────────────────────────────
# KONFIGURASI HALAMAN
# ─────────────────────────────────────────────────────────────────────────────

st.set_page_config(
    page_title = "Perhitungan Turap / Sheet Pile",
    page_icon  = "🏗️",
    layout     = "wide",
    initial_sidebar_state = "expanded",
)

# CSS kustom
st.markdown("""
<style>
/* Header utama */
.main-title {
    font-size: 1.6rem; font-weight: 700;
    color: #1a3a5c; border-bottom: 3px solid #2196F3;
    padding-bottom: 6px; margin-bottom: 4px;
}
/* Judul tab */
.tab-title {
    font-size: 1.15rem; font-weight: 600;
    color: #1565C0; margin-top: 4px;
}
/* Kotak status AMAN */
.status-aman {
    background: #E8F5E9; border-left: 5px solid #2E7D32;
    padding: 10px 14px; border-radius: 5px;
    font-weight: 700; color: #1B5E20; margin: 6px 0;
}
/* Kotak status TIDAK AMAN */
.status-tidak {
    background: #FFEBEE; border-left: 5px solid #C62828;
    padding: 10px 14px; border-radius: 5px;
    font-weight: 700; color: #B71C1C; margin: 6px 0;
}
/* Kotak info referensi */
.ref-box {
    background: #E3F2FD; border-left: 4px solid #1565C0;
    padding: 7px 12px; border-radius: 4px;
    font-size: 0.82rem; color: #1565C0; margin: 4px 0;
}
/* Label sidebar section */
.sidebar-section {
    font-weight: 700; color: #1a3a5c;
    border-bottom: 1px solid #ccc; padding-bottom: 3px;
    margin-top: 10px; margin-bottom: 6px;
}
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# FUNGSI PEMBANTU UI
# ─────────────────────────────────────────────────────────────────────────────

def _status_box(teks: str, aman: bool) -> None:
    css_kelas = "status-aman" if aman else "status-tidak"
    ikon      = "✅" if aman else "❌"
    st.markdown(f'<div class="{css_kelas}">{ikon}  {teks}</div>', unsafe_allow_html=True)


def _ref_box(*refs: str) -> None:
    teks = "<br>".join([f"📋 {r}" for r in refs])
    st.markdown(f'<div class="ref-box">{teks}</div>', unsafe_allow_html=True)


def _tab_header(nomor: str, judul: str) -> None:
    st.markdown(f'<div class="tab-title">Tab {nomor} — {judul}</div>', unsafe_allow_html=True)


def _fig_to_bytes(fig: plt.Figure) -> bytes:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight")
    buf.seek(0)
    return buf.read()


def _build_lapisan(sesi_lapisan: list[dict]) -> list[dict]:
    """Konversi format lapisan dari session_state ke format modul."""
    return [
        {
            "nama"     : lyr["nama"],
            "tebal"    : lyr["tebal"],
            "gamma"    : lyr["gamma"],
            "gamma_sat": lyr["gamma_sat"],
            "phi"      : lyr["phi"],
            "cohesion" : lyr["cohesion"],
        }
        for lyr in sesi_lapisan
    ]


# ─────────────────────────────────────────────────────────────────────────────
# INISIALISASI SESSION STATE
# ─────────────────────────────────────────────────────────────────────────────

_DEFAULT_LAPISAN = [
    {"nama": "Lempung lunak",  "tebal": 3.0, "gamma": 16.0, "gamma_sat": 17.0, "phi": 10.0, "cohesion": 10.0},
    {"nama": "Lempung medium", "tebal": 5.0, "gamma": 17.0, "gamma_sat": 18.5, "phi": 20.0, "cohesion": 20.0},
]

def _init_state():
    defaults = {
        "lapisan"         : _DEFAULT_LAPISAN.copy(),
        "hasil_ep"        : None,
        "hasil_stability" : None,
        "hasil_gd"        : None,
        "hasil_section"   : None,
        "hasil_anchor"    : None,
        "sudah_hitung"    : False,
        "error_msg"       : "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init_state()


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────

with st.sidebar:
    st.markdown('<div class="main-title">🏗️ TURAP / SHEET PILE</div>', unsafe_allow_html=True)
    st.caption("Program Perhitungan Geoteknik")

    # ── Data Proyek ───────────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">📁 Data Proyek</div>', unsafe_allow_html=True)
    nama_proyek   = st.text_input("Nama Proyek",    value="Galian Basement", key="nama_proyek")
    lokasi        = st.text_input("Lokasi",          value="Jakarta Selatan", key="lokasi")
    tgl           = st.date_input("Tanggal",         value=datetime.date.today(), key="tgl")
    dibuat_oleh   = st.text_input("Dibuat oleh",     value="Engineer", key="dibuat_oleh")
    diperiksa_oleh= st.text_input("Diperiksa oleh",  value="Senior Engineer", key="diperiksa_oleh")

    # ── Tipe Turap ────────────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">⚙️ Tipe Turap</div>', unsafe_allow_html=True)
    tipe_turap    = st.radio("Sistem Penahan",
                             ["Kantilever", "Dengan Angkur", "Dengan Strut"],
                             key="tipe_turap")
    tipe_material = st.radio("Material",
                             ["Baja", "Beton Pracetak"],
                             key="tipe_material")

    # ── Data Geometri ─────────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">📐 Geometri</div>', unsafe_allow_html=True)
    H           = st.number_input("H = Tinggi galian (m)",
                                  min_value=0.5, max_value=30.0, value=4.0, step=0.5, key="H")
    B_galian    = st.number_input("B = Lebar galian (m)",
                                  min_value=1.0, max_value=100.0, value=6.0, step=0.5, key="B_galian")
    surcharge   = st.number_input("Surcharge (kPa)",
                                  min_value=0.0, max_value=200.0, value=10.0, step=2.5, key="surcharge")

    if tipe_turap in ("Dengan Angkur", "Dengan Strut"):
        tinggi_angkur = st.number_input("Tinggi angkur dari permukaan (m)",
                                        min_value=0.0, max_value=H - 0.5,
                                        value=min(0.5, H - 0.5), step=0.25, key="tinggi_angkur")
        spasi_angkur  = st.number_input("Spasi angkur (m)",
                                        min_value=0.5, max_value=10.0, value=2.5, step=0.25, key="spasi_angkur")
    else:
        tinggi_angkur = 0.0
        spasi_angkur  = 1.0

    # ── Muka Air Tanah ────────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">💧 Muka Air Tanah</div>', unsafe_allow_html=True)
    MAT         = st.number_input("MAT sisi aktif (m dari permukaan)",
                                  min_value=0.0, max_value=30.0, value=1.5, step=0.25, key="MAT")
    MAT_pasif   = st.number_input("MAT sisi pasif (m dari dasar galian)",
                                  min_value=0.0, max_value=10.0, value=0.0, step=0.25, key="MAT_pasif")

    # ── Data Lapisan Tanah ────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">🪨 Data Lapisan Tanah</div>', unsafe_allow_html=True)

    # Tombol tambah / hapus lapisan
    col_btn1, col_btn2 = st.columns(2)
    with col_btn1:
        if st.button("➕ Tambah Lapisan", use_container_width=True):
            st.session_state["lapisan"].append(
                {"nama": f"Lapisan {len(st.session_state['lapisan'])+1}",
                 "tebal": 3.0, "gamma": 17.0, "gamma_sat": 18.5,
                 "phi": 15.0, "cohesion": 15.0}
            )
    with col_btn2:
        if st.button("➖ Hapus Terakhir", use_container_width=True):
            if len(st.session_state["lapisan"]) > 1:
                st.session_state["lapisan"].pop()

    # Input setiap lapisan
    for i, lyr in enumerate(st.session_state["lapisan"]):
        with st.expander(f"Lapisan {i+1}: {lyr['nama']}", expanded=(i == 0)):
            lyr["nama"]      = st.text_input("Nama",        value=lyr["nama"],      key=f"lyr_nama_{i}")
            lyr["tebal"]     = st.number_input("Tebal (m)", value=lyr["tebal"],     min_value=0.1, max_value=30.0, step=0.5, key=f"lyr_tebal_{i}")
            c1, c2 = st.columns(2)
            with c1:
                lyr["gamma"]     = st.number_input("gamma (kN/m3)",     value=lyr["gamma"],     min_value=10.0, max_value=25.0, step=0.5, key=f"lyr_gamma_{i}")
                lyr["phi"]       = st.number_input("phi (derajat)",      value=lyr["phi"],       min_value=0.0,  max_value=45.0, step=1.0, key=f"lyr_phi_{i}")
            with c2:
                lyr["gamma_sat"] = st.number_input("gamma_sat (kN/m3)", value=lyr["gamma_sat"], min_value=10.0, max_value=25.0, step=0.5, key=f"lyr_gsat_{i}")
                lyr["cohesion"]  = st.number_input("cohesion (kPa)",    value=lyr["cohesion"],  min_value=0.0,  max_value=500.0, step=5.0, key=f"lyr_c_{i}")

    # ── Material ──────────────────────────────────────────────────────────────
    st.markdown('<div class="sidebar-section">🔩 Material Turap</div>', unsafe_allow_html=True)

    if tipe_material == "Baja":
        mutu_options  = list(MUTU_BAJA.keys())
        mutu_pilih    = st.selectbox("Mutu Baja", mutu_options,
                                     index=mutu_options.index("BJ41") if "BJ41" in mutu_options else 0,
                                     key="mutu_baja")
        fy_baja       = st.number_input("fy (MPa)", value=float(MUTU_BAJA[mutu_pilih]["fy"]),
                                        min_value=200.0, max_value=900.0, step=10.0, key="fy_baja")
        fu_baja       = st.number_input("fu (MPa)", value=float(MUTU_BAJA[mutu_pilih]["fu"]),
                                        min_value=300.0, max_value=1100.0, step=10.0, key="fu_baja")
        seri_profil   = st.selectbox("Seri profil baja", ["Semua", "PZ", "AU", "PSA", "WF", "HZ"],
                                     key="seri_profil")
        seri_profil   = None if seri_profil == "Semua" else seri_profil
    else:
        fc_beton      = st.number_input("fc (MPa) — kuat tekan beton", value=25.0,
                                        min_value=17.0, max_value=60.0, step=1.0, key="fc_beton")
        fy_tulangan   = st.number_input("fy tulangan (MPa)", value=390.0,
                                        min_value=240.0, max_value=600.0, step=10.0, key="fy_tul")
        b_beton       = st.number_input("b = lebar penampang (mm)", value=1000.0,
                                        min_value=200.0, max_value=2000.0, step=50.0, key="b_beton")
        d_beton       = st.number_input("d = tinggi efektif (mm)", value=350.0,
                                        min_value=100.0, max_value=1000.0, step=25.0, key="d_beton")
        cover_beton   = st.number_input("Selimut beton (mm)", value=60.0,
                                        min_value=20.0, max_value=150.0, step=5.0, key="cover_beton")

    # ── Tombol Hitung ─────────────────────────────────────────────────────────
    st.markdown("---")
    tombol_hitung = st.button("🔢  HITUNG", type="primary", use_container_width=True)


# ─────────────────────────────────────────────────────────────────────────────
# PROSES PERHITUNGAN
# ─────────────────────────────────────────────────────────────────────────────

if tombol_hitung:
    st.session_state["error_msg"]       = ""
    st.session_state["sudah_hitung"]    = False
    st.session_state["hasil_ep"]        = None
    st.session_state["hasil_stability"] = None
    st.session_state["hasil_gd"]        = None
    st.session_state["hasil_section"]   = None
    st.session_state["hasil_anchor"]    = None

    lapisan = _build_lapisan(st.session_state["lapisan"])
    prog    = st.progress(0, text="Memulai perhitungan...")

    try:
        # ── TAHAP 1: Tekanan Tanah ─────────────────────────────────────────
        prog.progress(10, text="Tahap 1: Menghitung tekanan tanah lateral...")

        phi_rata = sum(l["phi"] * l["tebal"] for l in lapisan) / sum(l["tebal"] for l in lapisan)
        r_ka     = hitung_Ka_rankine(phi_rata)
        r_kp     = hitung_Kp_rankine(phi_rata)
        r_ko     = hitung_Ko(phi_rata)

        D_estimasi = max(H * 0.8, 2.0)
        r_dist  = hitung_distribusi_tekanan(
            lapisan_tanah = lapisan,
            H             = H,
            D             = D_estimasi,
            muka_air      = MAT,
            surcharge     = surcharge,
            dz            = 0.1,
        )

        fig_tekanan = plot_diagram_tekanan(
            z_array       = r_dist["nilai"]["z_array"],
            tekanan_aktif = r_dist["nilai"]["tekanan_aktif"],
            tekanan_pasif = r_dist["nilai"]["tekanan_pasif"],
            tekanan_neto  = r_dist["nilai"]["tekanan_neto"],
            H             = H,
            D             = D_estimasi,
            muka_air      = MAT,
            batas_lapisan = r_dist["nilai"]["batas_lapisan"],
            judul         = f"Distribusi Tekanan Tanah — {nama_proyek}",
        )

        st.session_state["hasil_ep"] = {
            "Ka"         : r_ka,
            "Kp"         : r_kp,
            "Ko"         : r_ko,
            "distribusi" : r_dist,
            "fig"        : fig_tekanan,
            "phi_rata"   : phi_rata,
        }

        # ── TAHAP 2: Stabilitas & Penetrasi ───────────────────────────────
        prog.progress(30, text="Tahap 2: Menghitung stabilitas dan kedalaman penetrasi...")

        if tipe_turap == "Kantilever":
            r_stab = fixed_earth_support_cantilever(
                H=H, lapisan_tanah=lapisan, surcharge=surcharge, muka_air=MAT, dz=0.1,
            )
            Ra_val      = 0.0
            z_angkur_val= 0.0
        else:
            r_stab = free_earth_support(
                H=H, lapisan_tanah=lapisan, surcharge=surcharge,
                muka_air=MAT, tinggi_angkur=tinggi_angkur, dz=0.1,
            )
            Ra_val       = r_stab["nilai"]["Ra"]
            z_angkur_val = tinggi_angkur

        cohesion_dasar = lapisan[min(1, len(lapisan)-1)]["cohesion"]
        r_heave_t = hitung_SF_heave_terzaghi(
            H=H, B_galian=B_galian, cohesion_dasar=cohesion_dasar,
            gamma=lapisan[0]["gamma"], surcharge=surcharge,
        )
        r_heave_b = hitung_SF_heave_bjerrum_eide(
            H=H, B_galian=B_galian, L_galian=max(B_galian * 3, 20.0),
            cohesion_dasar=cohesion_dasar, gamma=lapisan[0]["gamma"],
            surcharge=surcharge,
        )

        D_design = r_stab["nilai"]["D_design"]
        r_piping = hitung_SF_piping(
            H_total     = H + D_design,
            D           = D_design,
            delta_H     = max(MAT, 0.5),
            jenis_tanah = "lempung lunak s/d medium",
        )

        r_rangkuman_stab = rangkuman_stabilitas(
            hasil_analisis = {
                ("free_earth" if tipe_turap != "Kantilever" else "fixed_earth"): r_stab,
                "heave_terzaghi": r_heave_t,
                "heave_bjerrum" : r_heave_b,
                "piping"        : r_piping,
            },
            nama_proyek = nama_proyek,
        )

        st.session_state["hasil_stability"] = {
            "stab"          : r_stab,
            "heave_t"       : r_heave_t,
            "heave_b"       : r_heave_b,
            "piping"        : r_piping,
            "rangkuman"     : r_rangkuman_stab,
            "Ra"            : Ra_val,
            "z_angkur"      : z_angkur_val,
        }

        # ── TAHAP 3: Gaya Dalam ────────────────────────────────────────────
        prog.progress(55, text="Tahap 3: Menghitung diagram gaya geser dan momen...")

        z_arr   = r_stab["nilai"]["z_array"]
        p_neto  = r_stab["nilai"]["p_neto"]
        p_aktif = r_stab["nilai"]["p_aktif"]
        p_pasif = r_stab["nilai"]["p_pasif"]

        if tipe_turap == "Kantilever":
            r_gd = hitung_gaya_dalam_kantilever(
                z_array=z_arr, tekanan_array=p_neto,
                H=H, p_aktif=p_aktif, p_pasif=p_pasif,
            )
            z_angkur_plot = None
        else:
            r_gd = hitung_gaya_dalam(
                z_array=z_arr, tekanan_array=p_neto,
                Ra=Ra_val, z_angkur=z_angkur_val,
                p_aktif=p_aktif, p_pasif=p_pasif,
            )
            z_angkur_plot = z_angkur_val

        fig_vm = plot_V_M_diagram(
            z_array      = z_arr,
            V_array      = r_gd["nilai"]["V_array"],
            M_array      = r_gd["nilai"]["M_array"],
            M_max        = r_gd["nilai"]["M_max"],
            z_Mmax       = r_gd["nilai"]["z_Mmax"],
            H            = H,
            z_angkur     = z_angkur_plot,
            z_V0         = r_gd["nilai"]["z_V0"],
            D_design     = D_design,
            tekanan_neto = p_neto,
            judul        = f"Diagram V & M — {nama_proyek}",
        )

        st.session_state["hasil_gd"] = {
            "gd"   : r_gd,
            "fig"  : fig_vm,
        }

        # ── TAHAP 4: Desain Penampang ──────────────────────────────────────
        prog.progress(72, text="Tahap 4: Mendesain penampang turap...")

        M_max_val = r_gd["nilai"]["M_max"]
        V_max_val = float(np.max(np.abs(r_gd["nilai"]["V_array"])))

        if tipe_material == "Baja":
            csv_path = os.path.join(_DIR, "steel_sections.csv")
            r_sec = desain_baja(
                M_max         = M_max_val,
                fy            = fy_baja,
                fu            = fu_baja,
                path_csv      = csv_path if os.path.exists(csv_path) else None,
                nama_material = mutu_pilih,
                seri_profil   = seri_profil,
            )
        else:
            r_sec = desain_beton_pracetak(
                M_max  = M_max_val,
                V_max  = V_max_val,
                fc     = fc_beton,
                fy     = fy_tulangan,
                b      = b_beton,
                d      = d_beton,
                cover  = cover_beton,
                L_span = H,
            )

        st.session_state["hasil_section"] = {"section": r_sec}

        # ── TAHAP 5: Angkur / Strut ────────────────────────────────────────
        prog.progress(87, text="Tahap 5: Mendesain sistem angkur / strut...")

        r_anchor_dict: dict = {}

        if tipe_turap == "Dengan Angkur":
            if Ra_val > 0:
                r_tr = desain_tie_rod(
                    Ra=Ra_val, spasi_angkur=spasi_angkur,
                    fy_batang=fy_baja if tipe_material == "Baja" else 250.0,
                    fu_batang=fu_baja if tipe_material == "Baja" else 400.0,
                    nama_material=mutu_pilih if tipe_material == "Baja" else "BJ41",
                )
                r_wal = desain_waling(
                    Ra=Ra_val, spasi_angkur=spasi_angkur,
                    fy_waling=fy_baja if tipe_material == "Baja" else 240.0,
                )
                phi_dasar = lapisan[-1]["phi"]
                gamma_dasar = lapisan[-1]["gamma"]
                r_dm = desain_deadman(
                    Ra=Ra_val, gamma_tanah=gamma_dasar, phi_tanah=phi_dasar,
                    kedalaman_deadman=max(H * 0.5, 1.5),
                    cohesion_tanah=lapisan[-1]["cohesion"],
                    spasi_angkur=spasi_angkur,
                )
                r_anchor_dict = {
                    "tie_rod": r_tr,
                    "waling" : r_wal,
                    "deadman": r_dm,
                }

        elif tipe_turap == "Dengan Strut":
            Ra_strut = abs(Ra_val) if Ra_val != 0 else H * 15.0
            r_st = desain_strut(
                Ra=Ra_strut, spasi_strut=spasi_angkur,
                panjang_strut=B_galian,
                fy_strut=fy_baja if tipe_material == "Baja" else 250.0,
            )
            r_wal2 = desain_waling(
                Ra=Ra_strut, spasi_angkur=spasi_angkur,
                fy_waling=fy_baja if tipe_material == "Baja" else 240.0,
            )
            r_anchor_dict = {"strut": r_st, "waling": r_wal2}

        if r_anchor_dict:
            r_rk_anchor = rangkuman_angkur(
                hasil_tie_rod = r_anchor_dict.get("tie_rod"),
                hasil_waling  = r_anchor_dict.get("waling"),
                hasil_deadman = r_anchor_dict.get("deadman"),
                hasil_strut   = r_anchor_dict.get("strut"),
                nama_proyek   = nama_proyek,
            )
            r_anchor_dict["rangkuman"] = r_rk_anchor

        st.session_state["hasil_anchor"] = r_anchor_dict

        # ── Selesai ────────────────────────────────────────────────────────
        prog.progress(100, text="Perhitungan selesai!")
        st.session_state["sudah_hitung"] = True

    except Exception as exc:
        st.session_state["error_msg"]  = str(exc)
        st.session_state["sudah_hitung"] = False
        prog.empty()
        st.error(f"❌ **Error saat perhitungan:**\n\n`{exc}`\n\n"
                 f"```\n{traceback.format_exc()}\n```")


# ─────────────────────────────────────────────────────────────────────────────
# HEADER UTAMA
# ─────────────────────────────────────────────────────────────────────────────

st.markdown('<div class="main-title">🏗️ Perhitungan Turap / Sheet Pile</div>', unsafe_allow_html=True)

if st.session_state["sudah_hitung"]:
    info_col1, info_col2, info_col3 = st.columns(3)
    with info_col1:
        st.info(f"**Proyek:** {nama_proyek} | **Lokasi:** {lokasi}")
    with info_col2:
        D_show = st.session_state["hasil_stability"]["stab"]["nilai"]["D_design"]
        M_show = st.session_state["hasil_gd"]["gd"]["nilai"]["M_max"]
        st.info(f"**D_design:** {D_show:.2f} m | **M_max:** {M_show:.1f} kN.m/m")
    with info_col3:
        status_all = st.session_state["hasil_stability"]["rangkuman"]["status"]
        aman_all   = "SEMUA AMAN" in status_all
        _status_box(status_all, aman_all)
else:
    st.info("👈 Isi data di sidebar lalu tekan **HITUNG** untuk memulai perhitungan.")

st.markdown("---")


# ─────────────────────────────────────────────────────────────────────────────
# TABS
# ─────────────────────────────────────────────────────────────────────────────

tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "1️⃣ Tekanan Tanah",
    "2️⃣ Stabilitas & Penetrasi",
    "3️⃣ Gaya Dalam (V & M)",
    "4️⃣ Desain Penampang",
    "5️⃣ Angkur / Strut",
    "6️⃣ Ringkasan & Download",
])


# ═════════════════════════════════════════════════════════════════════════════
# TAB 1 — TEKANAN TANAH
# ═════════════════════════════════════════════════════════════════════════════

with tab1:
    _tab_header("1", "Tekanan Tanah Lateral")
    _ref_box(
        "SNI 8460:2017, Pasal 5.2–5.4 (tekanan tanah lateral)",
        "NAVFAC DM-7.01, Chapter 4 (Lateral Earth Pressure)",
        "USS Sheet Pile Design Manual (1975), Chapter 2",
    )

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat hasil.")
    else:
        ep = st.session_state["hasil_ep"]

        # ── Koefisien ─────────────────────────────────────────────────────────
        st.subheader("Koefisien Tekanan Tanah")
        k_col1, k_col2, k_col3, k_col4 = st.columns(4)
        with k_col1:
            st.metric("phi rata-rata", f"{ep['phi_rata']:.2f} °")
        with k_col2:
            st.metric("Ka (aktif)", f"{ep['Ka']['nilai']:.4f}")
        with k_col3:
            st.metric("Kp (pasif)", f"{ep['Kp']['nilai']:.4f}")
        with k_col4:
            st.metric("Ko (at-rest)", f"{ep['Ko']['nilai']:.4f}")

        # ── Diagram ───────────────────────────────────────────────────────────
        st.subheader("Diagram Distribusi Tekanan")
        col_fig, col_info = st.columns([3, 1])
        with col_fig:
            st.pyplot(ep["fig"], use_container_width=True)
        with col_info:
            st.markdown("**Keterangan:**")
            st.markdown("- **Pa** = tekanan aktif (kPa)")
            st.markdown("- **Pp** = tekanan pasif (kPa)")
            st.markdown("- **Pneto** = Pa − Pp")
            st.markdown("- Arsir merah = aktif mendominasi")
            st.markdown("- Arsir biru = pasif mendominasi")
            st.markdown(f"- **MAT** = {MAT:.2f} m")

        # ── Langkah Ka ────────────────────────────────────────────────────────
        with st.expander("📋 Langkah perhitungan Ka Rankine"):
            st.code(ep_fmt(ep["Ka"]["langkah"]), language="")

        with st.expander("📋 Langkah perhitungan Ko (Jaky)"):
            st.code(ep_fmt(ep["Ko"]["langkah"]), language="")

        with st.expander("📋 Tabel distribusi tekanan"):
            dist = ep["distribusi"]["nilai"]
            z_a = dist["z_array"]
            pa  = dist["tekanan_aktif"]
            pp  = dist["tekanan_pasif"]
            pn  = dist["tekanan_neto"]
            u   = dist["tekanan_air_pori"]
            import pandas as pd
            mask = np.arange(0, len(z_a), max(1, len(z_a)//40))
            df = pd.DataFrame({
                "z (m)"         : np.round(z_a[mask], 2),
                "Pa (kPa)"      : np.round(pa[mask], 3),
                "Pp (kPa)"      : np.round(pp[mask], 3),
                "Pneto (kPa)"   : np.round(pn[mask], 3),
                "u (kPa)"       : np.round(u[mask],  3),
            })
            st.dataframe(df, use_container_width=True, height=300)


# ═════════════════════════════════════════════════════════════════════════════
# TAB 2 — STABILITAS & PENETRASI
# ═════════════════════════════════════════════════════════════════════════════

with tab2:
    _tab_header("2", "Stabilitas dan Kedalaman Penetrasi")
    _ref_box(
        "SNI 8460:2017, Pasal 9.6.2 (stabilitas turap), 10.6.3 (heave)",
        "NAVFAC DM-7.02, Chapter 3 (sheet pile) dan Chapter 7 (heave)",
        "USS Sheet Pile Design Manual (1975), Chapter 3",
        "Terzaghi (1943) | Bjerrum & Eide (1956) | Lane (1935)",
    )

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat hasil.")
    else:
        stab_data = st.session_state["hasil_stability"]
        stab      = stab_data["stab"]["nilai"]
        heave_t   = stab_data["heave_t"]["nilai"]
        heave_b   = stab_data["heave_b"]["nilai"]
        piping    = stab_data["piping"]["nilai"]

        # ── Metrik utama ──────────────────────────────────────────────────────
        st.subheader(f"Penetrasi Turap ({tipe_turap})")
        m1, m2, m3, m4 = st.columns(4)
        with m1:
            st.metric("D_min",    f"{stab['D_min']:.3f} m")
        with m2:
            st.metric("D_design", f"{stab['D_design']:.3f} m",
                      delta=f"x{stab['D_design']/stab['D_min']:.2f} D_min")
        with m3:
            sf_key = "SF_momen"
            sf_val = stab.get(sf_key, 0.0)
            st.metric("SF Momen", f"{sf_val:.3f}",
                      delta="AMAN" if sf_val >= SF_GULING_MIN else "TIDAK AMAN",
                      delta_color="normal" if sf_val >= SF_GULING_MIN else "inverse")
        with m4:
            if tipe_turap != "Kantilever":
                st.metric("Ra (gaya angkur)", f"{stab.get('Ra', 0):.2f} kN/m")
            else:
                st.metric("z0 (jepit)", f"{stab.get('z0', 0):.3f} m")

        # ── Tabel SF ─────────────────────────────────────────────────────────
        st.subheader("Faktor Keamanan")
        sf_items = [
            ("SF Momen / Guling", sf_val,              SF_GULING_MIN),
            ("SF Heave (Terzaghi)", heave_t["SF_heave"], SF_HEAVE_MIN),
            ("SF Heave (Bjerrum-Eide)", heave_b["SF_heave"], SF_HEAVE_MIN),
            ("SF Piping (Lane)", piping["SF_piping"],   SF_PIPING_MIN),
        ]
        import pandas as pd
        df_sf = pd.DataFrame([
            {
                "Parameter"  : nm,
                "SF Hitung"  : f"{val:.3f}",
                "SF Min"     : f"{mn:.2f}",
                "Status"     : "✅ AMAN" if val >= mn else "❌ TIDAK AMAN",
            }
            for nm, val, mn in sf_items
        ])
        st.dataframe(df_sf, use_container_width=True, hide_index=True)

        # ── Status keseluruhan ────────────────────────────────────────────────
        semua_aman_stab = all(v >= m for _, v, m in sf_items)
        _status_box(
            "Semua faktor keamanan stabilitas memenuhi syarat." if semua_aman_stab
            else "Ada faktor keamanan yang TIDAK memenuhi syarat — perlu revisi.",
            semua_aman_stab,
        )

        # ── Detail langkah ────────────────────────────────────────────────────
        with st.expander(f"📋 Langkah {'Free' if tipe_turap != 'Kantilever' else 'Fixed'} Earth Support"):
            st.code(st_fmt(stab_data["stab"]["langkah"]), language="")

        col_h1, col_h2 = st.columns(2)
        with col_h1:
            with st.expander("📋 Langkah SF Heave — Terzaghi"):
                st.code(st_fmt(stab_data["heave_t"]["langkah"]), language="")
        with col_h2:
            with st.expander("📋 Langkah SF Heave — Bjerrum-Eide"):
                st.code(st_fmt(stab_data["heave_b"]["langkah"]), language="")

        with st.expander("📋 Langkah SF Piping — Lane"):
            st.code(st_fmt(stab_data["piping"]["langkah"]), language="")

        with st.expander("📋 Rangkuman stabilitas (semua SF)"):
            st.code(st_fmt(stab_data["rangkuman"]["langkah"]), language="")


# ═════════════════════════════════════════════════════════════════════════════
# TAB 3 — GAYA DALAM
# ═════════════════════════════════════════════════════════════════════════════

with tab3:
    _tab_header("3", "Gaya Dalam — Diagram V dan M")
    _ref_box(
        "NAVFAC DM-7.02, Section 3.2.5 (diagram gaya dalam turap)",
        "USS Sheet Pile Design Manual (1975), Chapter 4",
        "Das, B.M., Principles of Foundation Engineering, 8th Ed., Ch. 9.5",
    )

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat hasil.")
    else:
        gd_data = st.session_state["hasil_gd"]
        gd      = gd_data["gd"]["nilai"]

        # ── Metrik ────────────────────────────────────────────────────────────
        st.subheader("Gaya Dalam Kritis")
        g1, g2, g3, g4 = st.columns(4)
        with g1:
            st.metric("M_max", f"{gd['M_max']:.2f} kN.m/m")
        with g2:
            st.metric("z_Mmax", f"{gd['z_Mmax']:.3f} m")
        with g3:
            st.metric("V_max", f"{float(np.max(np.abs(gd['V_array']))):.2f} kN/m")
        with g4:
            st.metric("z_V0 (V=0)", f"{gd['z_V0']:.3f} m")

        # ── Diagram V & M ─────────────────────────────────────────────────────
        st.subheader("Diagram Gaya Geser dan Momen Lentur")
        st.pyplot(gd_data["fig"], use_container_width=True)

        # ── Tabel ringkasan ───────────────────────────────────────────────────
        with st.expander("📋 Tabel V dan M sepanjang kedalaman"):
            import pandas as pd
            z_a = gd["z_array"]
            V_a = gd["V_array"]
            M_a = gd["M_array"]
            p_a = gd["tekanan_array"]
            mask = np.arange(0, len(z_a), max(1, len(z_a)//50))
            df_gd = pd.DataFrame({
                "z (m)"      : np.round(z_a[mask], 2),
                "Pnet (kPa)" : np.round(p_a[mask], 3),
                "V (kN/m)"   : np.round(V_a[mask], 3),
                "M (kN.m/m)" : np.round(M_a[mask], 3),
            })
            st.dataframe(df_gd, use_container_width=True, height=320)

        with st.expander("📋 Langkah perhitungan gaya dalam"):
            st.code(if_fmt(gd_data["gd"]["langkah"]), language="")


# ═════════════════════════════════════════════════════════════════════════════
# TAB 4 — DESAIN PENAMPANG
# ═════════════════════════════════════════════════════════════════════════════

with tab4:
    _tab_header("4", f"Desain Penampang — {tipe_material}")
    if tipe_material == "Baja":
        _ref_box(
            "AISC Steel Construction Manual, 16th Ed., Section F11.1",
            "USS Steel Sheet Pile Design Manual (1975), Chapter 4",
            "SNI 1729:2020, Pasal F11",
        )
    else:
        _ref_box(
            "SNI 2847:2019, Pasal 9.6.1.2, 21.2.1, 21.2.2, 22.5.5.1",
            "ACI 318-19, Section 9.6.1.2 dan 22.5.5.1",
            "PCI Design Handbook, 8th Ed., Chapter 5",
        )

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat hasil.")
    else:
        sec_data = st.session_state["hasil_section"]["section"]
        sec_val  = sec_data["nilai"]

        # ── Metrik ────────────────────────────────────────────────────────────
        st.subheader("Hasil Desain Penampang")
        _status_box(sec_data["status"], "MEMENUHI" in sec_data["status"])

        if tipe_material == "Baja":
            s1, s2, s3, s4 = st.columns(4)
            with s1:
                st.metric("sigma_allow", f"{sec_val['sigma_allow']:.1f} MPa")
            with s2:
                st.metric("S_req", f"{sec_val['S_req']:.2f} cm3/m")
            with s3:
                pt = sec_val.get("profil_terpilih")
                st.metric("Profil", pt["tipe"] if pt else "—")
            with s4:
                st.metric("S_pakai", f"{pt['S']:.1f} cm3/m" if pt else "—")

            if pt:
                st.subheader("Profil Terpilih")
                import pandas as pd
                df_profil = pd.DataFrame([{
                    "Tipe"          : pt["tipe"],
                    "S (cm3/m)"     : pt["S"],
                    "Berat (kg/m2)" : pt["berat"],
                    "Ix (cm4/m)"    : pt["Ix"],
                    "Material"      : pt["material"],
                    "Keterangan"    : pt["keterangan"],
                }])
                st.dataframe(df_profil, use_container_width=True, hide_index=True)

            if sec_val.get("kandidat_profil"):
                with st.expander("📋 Semua kandidat profil yang memenuhi"):
                    import pandas as pd
                    df_kand = pd.DataFrame([
                        {
                            "Tipe"     : p["tipe"],
                            "S(cm3/m)" : p["S"],
                            "Berat(kg/m2)": p["berat"],
                            "Ix(cm4/m)": p["Ix"],
                            "Material" : p["material"],
                        }
                        for p in sec_val["kandidat_profil"]
                    ])
                    st.dataframe(df_kand, use_container_width=True, hide_index=True)

        else:  # Beton Pracetak
            b1, b2, b3 = st.columns(3)
            with b1:
                st.metric("Mu", f"{sec_val['Mu']:.2f} kN.m/m")
                st.metric("Rn", f"{sec_val['Rn']:.4f} MPa")
            with b2:
                st.metric("As_req", f"{sec_val['As_req']:.2f} mm2/m")
                st.metric("rho_pakai", f"{sec_val['rho_pakai']:.5f}")
            with b3:
                st.metric("phi_Vc", f"{sec_val['phi_Vc']:.2f} kN/m")
                aman_g = sec_val["aman_geser"]
                _status_box(f"Geser: phi_Vc = {sec_val['phi_Vc']:.2f} >= Vu = {sec_val['Vu']:.2f} kN/m", aman_g)

            # Tabel kontrol tulangan
            import pandas as pd
            df_tul = pd.DataFrame([{
                "Parameter"   : n,
                "Nilai"       : v,
            } for n, v in [
                ("rho_tulangan", f"{sec_val['rho_tulangan']:.6f}"),
                ("rho_min",      f"{sec_val['rho_min']:.6f}"),
                ("rho_max",      f"{sec_val['rho_max']:.6f}"),
                ("rho_pakai",    f"{sec_val['rho_pakai']:.6f}"),
                ("As_req (mm2/m)", f"{sec_val['As_req']:.2f}"),
                ("As_min (mm2/m)", f"{sec_val['As_min']:.2f}"),
                ("As_max (mm2/m)", f"{sec_val['As_maks']:.2f}"),
                ("Status lentur",  "OK" if sec_val["aman_lentur"] else "REVISI"),
                ("Status geser",   "OK" if sec_val["aman_geser"]  else "REVISI"),
            ]])
            st.dataframe(df_tul, use_container_width=True, hide_index=True)

        with st.expander("📋 Langkah perhitungan desain penampang"):
            st.code(sd_fmt(sec_data["langkah"]), language="")


# ═════════════════════════════════════════════════════════════════════════════
# TAB 5 — ANGKUR / STRUT
# ═════════════════════════════════════════════════════════════════════════════

with tab5:
    _tab_header("5", f"Sistem {tipe_turap}")
    _ref_box(
        "NAVFAC DM-7.02, Section 3.3 (angkur) dan 3.4 (deadman)",
        "USS Sheet Pile Design Manual (1975), Chapter 5",
        "AISC 16th Ed., Ch. E (tekan) dan J (sambungan)  |  SNI 1729:2020",
    )

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat hasil.")
    elif tipe_turap == "Kantilever":
        st.info("Turap kantilever tidak memiliki angkur atau strut.")
    elif not st.session_state["hasil_anchor"]:
        st.warning("Data angkur belum tersedia. Periksa hasil Tahap 2.")
    else:
        anc = st.session_state["hasil_anchor"]

        # ── Tie Rod ───────────────────────────────────────────────────────────
        if "tie_rod" in anc:
            st.subheader("Tie Rod (Batang Angkur Tarik)")
            tr = anc["tie_rod"]["nilai"]
            t1, t2, t3, t4 = st.columns(4)
            with t1: st.metric("T_angkur", f"{tr['T_angkur']:.2f} kN")
            with t2: st.metric("A_req", f"{tr['A_req']:.2f} mm2")
            with t3: st.metric("Diameter pakai", f"{tr['diameter_pakai']:.0f} mm")
            with t4: st.metric("SF", f"{tr['SF']:.3f}")
            _status_box(f"Tie Rod D{tr['diameter_pakai']:.0f}mm — SF={tr['SF']:.3f}", tr["aman"])

            with st.expander("📋 Langkah desain tie rod"):
                st.code(ad_fmt(anc["tie_rod"]["langkah"]), language="")

        # ── Waling ────────────────────────────────────────────────────────────
        if "waling" in anc:
            st.subheader("Waling (Balok Distribusi)")
            wal = anc["waling"]["nilai"]
            w1, w2, w3 = st.columns(3)
            with w1: st.metric("M_waling", f"{wal['M_waling']:.3f} kN.m")
            with w2: st.metric("S_req", f"{wal['S_req']:.2f} cm3")
            pt_wal = wal.get("profil_terpilih")
            with w3: st.metric("Profil", pt_wal[0] if pt_wal else "—")
            _status_box(
                f"Waling {pt_wal[0] if pt_wal else '—'} — rasio_S={wal.get('rasio_S', 0):.3f}",
                wal.get("aman", False),
            )

            with st.expander("📋 Langkah desain waling"):
                st.code(ad_fmt(anc["waling"]["langkah"]), language="")

        # ── Deadman ───────────────────────────────────────────────────────────
        if "deadman" in anc:
            st.subheader("Deadman Anchor (Angkur Pelat)")
            dm = anc["deadman"]["nilai"]
            d1, d2, d3 = st.columns(3)
            with d1: st.metric("Pp_per_m", f"{dm['Pp_per_m']:.3f} kN/m")
            with d2: st.metric("L_deadman (min)", f"{dm['L_deadman']:.3f} m")
            with d3: st.metric("SF Deadman", f"{dm['SF_actual']:.3f}")
            _status_box(f"Deadman — SF={dm['SF_actual']:.3f}", dm["aman"])

            with st.expander("📋 Langkah desain deadman"):
                st.code(ad_fmt(anc["deadman"]["langkah"]), language="")

        # ── Strut ─────────────────────────────────────────────────────────────
        if "strut" in anc:
            st.subheader("Strut / Bracing (Batang Tekan)")
            st_val  = anc["strut"]["nilai"]
            pt_st   = st_val.get("profil_terpilih")
            s1, s2, s3 = st.columns(3)
            with s1: st.metric("P_strut", f"{st_val['P_strut']:.2f} kN")
            with s2: st.metric("Profil", pt_st["nama"] if pt_st else "—")
            with s3:
                if pt_st:
                    st.metric("SF Strut", f"{pt_st['SF']:.3f}")
                    st.metric("lambda_sr", f"{pt_st['lambda_sr']:.1f}")
            _status_box(
                f"Strut {pt_st['nama'] if pt_st else '—'} — SF={pt_st['SF']:.3f if pt_st else 0}",
                st_val.get("aman", False),
            )

            with st.expander("📋 Langkah desain strut"):
                st.code(ad_fmt(anc["strut"]["langkah"]), language="")

        # ── Rangkuman angkur ──────────────────────────────────────────────────
        if "rangkuman" in anc:
            with st.expander("📋 Rangkuman semua komponen angkur"):
                st.code(ad_fmt(anc["rangkuman"]["langkah"]), language="")


# ═════════════════════════════════════════════════════════════════════════════
# TAB 6 — RINGKASAN & DOWNLOAD
# ═════════════════════════════════════════════════════════════════════════════

with tab6:
    _tab_header("6", "Ringkasan Hasil dan Download")

    if not st.session_state["sudah_hitung"]:
        st.warning("Tekan HITUNG untuk melihat ringkasan.")
    else:
        stab_data = st.session_state["hasil_stability"]
        gd_data   = st.session_state["hasil_gd"]
        sec_data  = st.session_state["hasil_section"]["section"]
        anc_data  = st.session_state["hasil_anchor"]

        stab_val = stab_data["stab"]["nilai"]
        gd_val   = gd_data["gd"]["nilai"]
        sec_val  = sec_data["nilai"]

        # ── Info proyek ───────────────────────────────────────────────────────
        st.subheader("Informasi Proyek")
        import pandas as pd
        df_info = pd.DataFrame([
            {"Parameter": k, "Nilai": v}
            for k, v in [
                ("Nama Proyek",     nama_proyek),
                ("Lokasi",          lokasi),
                ("Tanggal",         str(tgl)),
                ("Dibuat oleh",     dibuat_oleh),
                ("Diperiksa oleh",  diperiksa_oleh),
                ("Tipe Turap",      tipe_turap),
                ("Material",        tipe_material),
                ("H Galian",        f"{H:.2f} m"),
                ("B Galian",        f"{B_galian:.2f} m"),
                ("Surcharge",       f"{surcharge:.2f} kPa"),
                ("MAT",             f"{MAT:.2f} m dari permukaan"),
            ]
        ])
        st.dataframe(df_info, use_container_width=True, hide_index=True)

        # ── Ringkasan parameter desain ────────────────────────────────────────
        st.subheader("Parameter Desain Final")

        import pandas as pd
        baris_ringkasan = [
            ("D_min",     f"{stab_val['D_min']:.4f} m",    "—"),
            ("D_design",  f"{stab_val['D_design']:.4f} m", "—"),
        ]
        if tipe_turap != "Kantilever":
            baris_ringkasan.append(("Ra (gaya angkur)", f"{stab_val.get('Ra', 0):.3f} kN/m", "—"))
        else:
            baris_ringkasan.append(("z0 (titik jepit)", f"{stab_val.get('z0', 0):.4f} m", "—"))

        baris_ringkasan += [
            ("M_max",     f"{gd_val['M_max']:.3f} kN.m/m", "—"),
            ("z_Mmax",    f"{gd_val['z_Mmax']:.3f} m",     "—"),
            ("V_max",     f"{float(np.max(np.abs(gd_val['V_array']))):.3f} kN/m", "—"),
        ]

        # Material
        if tipe_material == "Baja":
            pt = sec_val.get("profil_terpilih")
            baris_ringkasan += [
                ("sigma_allow",  f"{sec_val['sigma_allow']:.2f} MPa",  "—"),
                ("S_req",        f"{sec_val['S_req']:.2f} cm3/m",     "—"),
                ("Profil baja",  pt["tipe"] if pt else "—",            "—"),
                ("S_pakai",      f"{pt['S']:.1f} cm3/m" if pt else "—","—"),
                ("Berat turap",  f"{pt['berat']:.1f} kg/m2" if pt else "—", "—"),
            ]
        else:
            baris_ringkasan += [
                ("Mu",           f"{sec_val['Mu']:.3f} kN.m/m", "—"),
                ("As_req",       f"{sec_val['As_req']:.2f} mm2/m", "—"),
                ("rho_pakai",    f"{sec_val['rho_pakai']:.6f}", "—"),
                ("phi_Vc",       f"{sec_val['phi_Vc']:.2f} kN/m", "—"),
            ]

        df_final = pd.DataFrame([
            {"Parameter": p, "Nilai": v, "Catatan": c}
            for p, v, c in baris_ringkasan
        ])
        st.dataframe(df_final, use_container_width=True, hide_index=True)

        # ── Tabel SF ─────────────────────────────────────────────────────────
        st.subheader("Faktor Keamanan — Semua Komponen")
        sf_final = [
            ("SF Momen",              stab_val.get("SF_momen", 0),         SF_GULING_MIN),
            ("SF Heave (Terzaghi)",   stab_data["heave_t"]["nilai"]["SF_heave"], SF_HEAVE_MIN),
            ("SF Heave (Bjerrum)",    stab_data["heave_b"]["nilai"]["SF_heave"], SF_HEAVE_MIN),
            ("SF Piping (Lane)",      stab_data["piping"]["nilai"]["SF_piping"], SF_PIPING_MIN),
        ]
        if anc_data and "tie_rod" in anc_data:
            sf_final.append(("SF Tie Rod", anc_data["tie_rod"]["nilai"]["SF"], 2.0))
        if anc_data and "deadman" in anc_data:
            sf_final.append(("SF Deadman", anc_data["deadman"]["nilai"]["SF_actual"], 2.0))
        if anc_data and "strut" in anc_data:
            pt_st = anc_data["strut"]["nilai"].get("profil_terpilih")
            if pt_st:
                sf_final.append(("SF Strut", pt_st["SF"], 2.0))

        df_sf_final = pd.DataFrame([
            {
                "Komponen"   : nm,
                "SF Hitung"  : f"{val:.3f}",
                "SF Min"     : f"{mn:.2f}",
                "Status"     : "✅ AMAN" if val >= mn else "❌ TIDAK AMAN",
            }
            for nm, val, mn in sf_final
        ])
        st.dataframe(df_sf_final, use_container_width=True, hide_index=True)

        semua_ok = all(v >= m for _, v, m in sf_final)
        _status_box(
            "SEMUA FAKTOR KEAMANAN MEMENUHI — DESAIN DAPAT DITERIMA" if semua_ok
            else "ADA FAKTOR KEAMANAN TIDAK MEMENUHI — REVISI DIPERLUKAN",
            semua_ok,
        )

        # ── Download ──────────────────────────────────────────────────────────
        st.markdown("---")
        st.subheader("Download Laporan")

        def _buat_teks_laporan() -> str:
            """Susun teks laporan lengkap dari semua langkah."""
            garis = "=" * 70
            baris: list[str] = [
                garis,
                f"LAPORAN PERHITUNGAN TURAP / SHEET PILE",
                garis,
                f"Proyek    : {nama_proyek}",
                f"Lokasi    : {lokasi}",
                f"Tanggal   : {tgl}",
                f"Dibuat    : {dibuat_oleh}",
                f"Diperiksa : {diperiksa_oleh}",
                "",
                f"H galian   = {H} m  |  B galian = {B_galian} m",
                f"Surcharge  = {surcharge} kPa  |  MAT = {MAT} m",
                f"Tipe turap = {tipe_turap}  |  Material = {tipe_material}",
                garis,
            ]

            sesi = st.session_state
            for judul, kunci_luar, kunci_dalam in [
                ("TEKANAN TANAH — Ka Rankine", "hasil_ep",        "Ka"),
                ("TEKANAN TANAH — Ko",          "hasil_ep",        "Ko"),
            ]:
                obj = sesi.get(kunci_luar)
                if obj and kunci_dalam in obj:
                    baris += ["", judul, "-" * 50]
                    baris += obj[kunci_dalam].get("langkah", [])

            # Stabilitas
            hs = sesi.get("hasil_stability")
            if hs:
                baris += ["", "ANALISIS STABILITAS", "-" * 50]
                baris += hs["stab"].get("langkah", [])
                baris += ["", "SF HEAVE — Terzaghi", "-" * 50]
                baris += hs["heave_t"].get("langkah", [])
                baris += ["", "SF HEAVE — Bjerrum-Eide", "-" * 50]
                baris += hs["heave_b"].get("langkah", [])
                baris += ["", "SF PIPING — Lane", "-" * 50]
                baris += hs["piping"].get("langkah", [])
                baris += ["", "RANGKUMAN STABILITAS", "-" * 50]
                baris += hs["rangkuman"].get("langkah", [])

            # Gaya dalam
            hg = sesi.get("hasil_gd")
            if hg:
                baris += ["", "GAYA DALAM — V DAN M", "-" * 50]
                baris += hg["gd"].get("langkah", [])

            # Desain penampang
            hss = sesi.get("hasil_section")
            if hss:
                baris += ["", f"DESAIN PENAMPANG — {tipe_material}", "-" * 50]
                baris += hss["section"].get("langkah", [])

            # Angkur
            ha = sesi.get("hasil_anchor")
            if ha:
                for k, judul_k in [
                    ("tie_rod", "TIE ROD"),
                    ("waling",  "WALING"),
                    ("deadman", "DEADMAN ANCHOR"),
                    ("strut",   "STRUT / BRACING"),
                ]:
                    if k in ha:
                        baris += ["", f"DESAIN {judul_k}", "-" * 50]
                        baris += ha[k].get("langkah", [])

            return "\n".join(baris)

        def _buat_docx() -> bytes:
            """Ekspor laporan ke format .docx."""
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style heading
            judul_para = doc.add_heading("LAPORAN PERHITUNGAN TURAP / SHEET PILE", level=1)
            judul_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Info proyek
            doc.add_heading("Data Proyek", level=2)
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Parameter"; hdr[1].text = "Nilai"
            for k, v in [
                ("Nama Proyek", nama_proyek), ("Lokasi", lokasi),
                ("Tanggal", str(tgl)), ("Dibuat oleh", dibuat_oleh),
                ("Diperiksa oleh", diperiksa_oleh),
                ("H Galian", f"{H:.2f} m"), ("D_design", f"{stab_val['D_design']:.4f} m"),
                ("M_max", f"{gd_val['M_max']:.3f} kN.m/m"),
            ]:
                row = tbl.add_row().cells
                row[0].text = k; row[1].text = v

            # Langkah perhitungan tiap bagian
            for judul_bab, teks_langkah in [
                ("Koefisien Ka Rankine",
                 ep_fmt(st.session_state["hasil_ep"]["Ka"]["langkah"])),
                ("Analisis Stabilitas",
                 st_fmt(st.session_state["hasil_stability"]["stab"]["langkah"])),
                ("Gaya Dalam V dan M",
                 if_fmt(st.session_state["hasil_gd"]["gd"]["langkah"])),
                (f"Desain Penampang {tipe_material}",
                 sd_fmt(st.session_state["hasil_section"]["section"]["langkah"])),
            ]:
                doc.add_heading(judul_bab, level=2)
                doc.add_paragraph(teks_langkah, style="No Spacing")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return buf.read()

        def _buat_pdf() -> bytes:
            """Ekspor laporan ke format .pdf."""
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import mm
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Preformatted
            from reportlab.lib import colors

            buf = io.BytesIO()
            doc_pdf = SimpleDocTemplate(buf, pagesize=A4,
                                        leftMargin=20*mm, rightMargin=20*mm,
                                        topMargin=20*mm, bottomMargin=20*mm)
            styles = getSampleStyleSheet()
            kode_style = ParagraphStyle(
                "Kode", fontName="Courier", fontSize=6.5,
                leading=9, parent=styles["Normal"],
            )
            story = []
            story.append(Paragraph("LAPORAN PERHITUNGAN TURAP / SHEET PILE", styles["Title"]))
            story.append(Spacer(1, 6*mm))
            story.append(Paragraph(f"Proyek: {nama_proyek}  |  Lokasi: {lokasi}  |  Tgl: {tgl}", styles["Normal"]))
            story.append(Spacer(1, 4*mm))

            teks_penuh = _buat_teks_laporan()
            for baris in teks_penuh.split("\n"):
                story.append(Preformatted(baris, kode_style))

            doc_pdf.build(story)
            buf.seek(0)
            return buf.read()

        # ── Tombol download ───────────────────────────────────────────────────
        dl1, dl2, dl3 = st.columns(3)

        with dl1:
            teks_lap = _buat_teks_laporan()
            st.download_button(
                label      = "⬇️ Download TXT",
                data       = teks_lap.encode("utf-8"),
                file_name  = f"Turap_{nama_proyek.replace(' ','_')}_{tgl}.txt",
                mime       = "text/plain",
                use_container_width=True,
            )

        with dl2:
            try:
                docx_bytes = _buat_docx()
                st.download_button(
                    label      = "⬇️ Download Word (.docx)",
                    data       = docx_bytes,
                    file_name  = f"Turap_{nama_proyek.replace(' ','_')}_{tgl}.docx",
                    mime       = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"Word tidak tersedia: {e}")

        with dl3:
            try:
                pdf_bytes = _buat_pdf()
                st.download_button(
                    label      = "⬇️ Download PDF",
                    data       = pdf_bytes,
                    file_name  = f"Turap_{nama_proyek.replace(' ','_')}_{tgl}.pdf",
                    mime       = "application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"PDF tidak tersedia: {e}")

        # ── Gambar untuk download ─────────────────────────────────────────────
        st.markdown("---")
        st.subheader("Download Diagram")
        img1, img2 = st.columns(2)
        with img1:
            fig_bytes_t = _fig_to_bytes(st.session_state["hasil_ep"]["fig"])
            st.download_button(
                "⬇️ Diagram Tekanan Tanah (PNG)",
                data       = fig_bytes_t,
                file_name  = f"tekanan_{nama_proyek}_{tgl}.png",
                mime       = "image/png",
                use_container_width=True,
            )
        with img2:
            fig_bytes_vm = _fig_to_bytes(st.session_state["hasil_gd"]["fig"])
            st.download_button(
                "⬇️ Diagram V & M (PNG)",
                data       = fig_bytes_vm,
                file_name  = f"V_M_{nama_proyek}_{tgl}.png",
                mime       = "image/png",
                use_container_width=True,
            )

        # ── Footer ────────────────────────────────────────────────────────────
        st.markdown("---")
        st.caption(
            "Perhitungan mengacu pada: SNI 8460:2017 | SNI 2847:2019 | SNI 1729:2020 | "
            "NAVFAC DM-7.01 & DM-7.02 | USS Sheet Pile Design Manual (1975) | "
            "AISC 16th Ed. | ACI 318-19 | Terzaghi (1943) | Lane (1935)"
        )

