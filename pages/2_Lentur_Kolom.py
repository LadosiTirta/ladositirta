"""
=============================================================
HALAMAN 2 - KAPASITAS LENTUR & AKSIAL KOLOM BETON BERTULANG
Referensi : SNI 2847:2019 (ACI 318-14)
Asumsi    : Kolom Pendek (Short Column) - tanpa kelangsingan
Framework  : Streamlit (multipage)
Output     : Word (.docx) & PDF (.pdf) + Watermark
=============================================================
"""

import math
import io
import datetime
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# ============================================================
# KONFIGURASI HALAMAN
# ============================================================
st.set_page_config(
    page_title="Lentur Kolom Beton | SNI 2847:2019",
    page_icon="",
    layout="wide",
)

st.markdown("""
<style>
  .main-title{font-size:1.6rem;font-weight:600;color:#1a3c5e;margin-bottom:0}
  .sub-title{font-size:.9rem;color:#666;margin-bottom:1.5rem}
  .step-box{background:#f8f9fa;border-left:4px solid #1a3c5e;border-radius:0 8px 8px 0;
            padding:12px 16px;margin-bottom:10px;font-size:.85rem}
  .step-hdr{font-weight:700;color:#1a3c5e;font-size:.8rem;
            text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}
  .result-ok  {background:#e8f5e9;border-left:4px solid #2e7d32;color:#1b5e20;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .result-warn{background:#fff8e1;border-left:4px solid #f9a825;color:#5d4037;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .result-fail{background:#ffebee;border-left:4px solid #c62828;color:#b71c1c;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .metric-card{background:white;border:1px solid #e0e0e0;border-radius:10px;
               padding:14px;text-align:center;margin-bottom:6px}
  .metric-lbl{font-size:.75rem;color:#888;margin-bottom:4px}
  .metric-val{font-size:1.4rem;font-weight:700;color:#1a3c5e}
  .metric-unt{font-size:.75rem;color:#aaa}
  .ref-badge{display:inline-block;background:#e3f2fd;color:#1565c0;
             font-size:.7rem;padding:2px 8px;border-radius:20px;margin-bottom:6px}
  hr.divider{border:none;border-top:2px solid #e0e0e0;margin:1.5rem 0}
  .asumsi-box{background:#f0f4ff;border-left:4px solid #3f51b5;border-radius:0 8px 8px 0;
              padding:10px 14px;margin-bottom:1rem;font-size:.85rem;color:#283593}
</style>
""", unsafe_allow_html=True)


# ============================================================
# SANITASI STRING UNTUK PDF (Latin-1 only)
# ============================================================
_UNICODE_MAP = {
    "\u2014": "-",    "\u2013": "-",    "\u2019": "'",   "\u2018": "'",
    "\u201c": '"',    "\u201d": '"',    "\u00b2": "2",   "\u00b3": "3",
    "\u00b0": " deg", "\u00d7": "x",   "\u2265": ">=",  "\u2264": "<=",
    "\u2260": "!=",   "\u221a": "sqrt", "\u03c6": "Phi", "\u03b2": "Beta",
    "\u03b5": "et",   "\u03c1": "Rho",  "\u03bc": "mu",  "\u2022": "-",
    "\u2192": "->",   "\u00b7": ".",    "\u00e9": "e",   "\u00e8": "e",
    "\u00e0": "a",    "\u2081": "1",    "\u2082": "2",
}

def sp(teks: str) -> str:
    """Sanitasi string agar aman untuk fpdf2 (Latin-1 only)."""
    if not isinstance(teks, str):
        teks = str(teks)
    for ch, repl in _UNICODE_MAP.items():
        teks = teks.replace(ch, repl)
    return teks.encode("latin-1", errors="replace").decode("latin-1")


# ============================================================
# FUNGSI BANTU - LUAS TULANGAN
# ============================================================
# Diameter standar (mm) dan luasnya (mm2)
DIAM_OPTS = {
    "D10 (78.5 mm2)": (10, 78.54),
    "D13 (132.7 mm2)": (13, 132.73),
    "D16 (201.1 mm2)": (16, 201.06),
    "D19 (283.5 mm2)": (19, 283.53),
    "D22 (380.1 mm2)": (22, 380.13),
    "D25 (490.9 mm2)": (25, 490.87),
    "D29 (660.5 mm2)": (29, 660.52),
    "D32 (804.2 mm2)": (32, 804.25),
    "D36 (1017.9 mm2)": (36, 1017.88),
}


# ============================================================
# FUNGSI PERHITUNGAN KOLOM
# ============================================================
def hitung_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h):
    """
    Menghitung kapasitas aksial-lentur kolom persegi pendek simetris.

    Parameter:
      fc      : kuat tekan beton (MPa)
      fy      : kuat leleh tulangan (MPa)
      b       : lebar kolom (mm)
      h       : tinggi kolom (mm)
      cover   : tebal selimut beton ke pusat tulangan (mm)
      D_tul   : diameter tulangan longitudinal (mm)
      n_total : jumlah total tulangan longitudinal
      n_b     : jumlah tulangan sisi lebar (termasuk sudut)
      n_h     : jumlah tulangan sisi tinggi (termasuk sudut)

    Return: dict hasil perhitungan
    """
    R = {}

    # -- Luas tulangan per batang
    Ab = math.pi * D_tul**2 / 4
    Ast = n_total * Ab

    # -- Dimensi efektif
    d_prime = cover                        # jarak selimut ke pusat tul.
    d       = h - cover                    # tinggi efektif sisi tarik
    Ag      = b * h                        # luas penampang bruto

    # -- Rasio tulangan
    rho_g = Ast / Ag

    # -- Beta1
    if fc <= 28:
        beta1 = 0.85
    elif fc >= 56:
        beta1 = 0.65
    else:
        beta1 = max(0.65, min(0.85, 0.85 - 0.05 * (fc - 28) / 7))

    # ======================================================
    # KAPASITAS AKSIAL MURNI (Pn0)
    # SNI 2847:2019 Pasal 22.4.2.2
    # Pn0 = 0.85 fc (Ag - Ast) + fy Ast
    # ======================================================
    Pn0   = 0.85 * fc * (Ag - Ast) + fy * Ast          # N
    phi_c = 0.65                                         # tied column
    # Batas maksimum aksial SNI (untuk kolom ikat / tied):
    # phi.Pn,max = phi x 0.80 x Pn0
    phiPn_max = phi_c * 0.80 * Pn0                      # N

    # ======================================================
    # KAPASITAS LENTUR MURNI (Mn0) - kondisi beban aksial = 0
    # Gunakan pendekatan: tulangan simetris, hitung Mn seperti balok
    # dengan semua tulangan tarik pada satu sisi
    # ======================================================
    # Tulangan sisi tarik = n_b batang (sisi bawah/satu sisi lebar)
    # Tulangan sisi tekan = n_b batang (sisi atas/satu sisi lebar)
    # Tulangan sisi kiri & kanan = n_h - 2 batang tiap sisi (sudut sudah terhitung)
    n_sisi_tarik  = n_b                           # tulangan sisi tarik (bawah)
    n_sisi_tekan  = n_b                           # tulangan sisi tekan (atas)

    As_tarik  = n_sisi_tarik * Ab                 # mm2
    As_tekan  = n_sisi_tekan * Ab                 # mm2

    # Tinggi blok tegangan - pendekatan As_net
    As_net = As_tarik - As_tekan
    if As_net <= 0:
        As_net = As_tarik   # fallback: abaikan tulangan tekan

    a_flex  = (As_tarik * fy) / (0.85 * fc * b)
    c_flex  = a_flex / beta1
    et_flex = 0.003 * (d - c_flex) / c_flex if c_flex > 0 else 0.010
    phi_f   = 0.90 if et_flex >= 0.005 else max(0.65, 0.65 + (et_flex - 0.002) * (250/3))
    Mn0     = As_tarik * fy * (d - a_flex / 2) / 1_000_000   # kN.m
    phiMn0  = phi_f * Mn0

    # ======================================================
    # TITIK SEIMBANG (BALANCED CONDITION)
    # SNI 2847:2019 Pasal 22.2.2
    # ======================================================
    cb   = (600 / (600 + fy)) * d                          # mm
    ab   = beta1 * cb                                      # mm
    # Regangan baja tekan pada balanced
    eps_sp = 0.003 * (cb - d_prime) / cb
    fsp    = min(fy, eps_sp * 200_000)                     # MPa

    Cc_b   = 0.85 * fc * ab * b / 1000                    # kN
    Cs_b   = As_tekan * fsp / 1000                        # kN
    T_b    = As_tarik * fy / 1000                         # kN

    Pb     = (Cc_b + Cs_b - T_b)                          # kN
    Mb_arm = (h / 2 - ab / 2) / 1000                      # m (dari pusat penampang)
    # Momen balanced terhadap pusat penampang
    e_Cc   = h / 2 - ab / 2                               # mm
    e_Cs   = h / 2 - d_prime                              # mm
    e_T    = d - h / 2                                    # mm
    Mb     = (Cc_b * e_Cc + Cs_b * e_Cs - T_b * (-e_T)) / 1000  # kN.m - salah tanda
    Mb     = abs((Cc_b * e_Cc / 1000) + (Cs_b * e_Cs / 1000) + (T_b * e_T / 1000))

    phi_b  = 0.65
    phiPb  = phi_b * Pb
    phiMb  = phi_b * Mb

    # ======================================================
    # KAPASITAS NOMINAL INTERAKSI - TIGA TITIK KONTROL
    # ======================================================
    # Titik A : Aksial murni    (Mn=0, Pn=Pn0)
    # Titik B : Balanced        (Pn=Pb, Mn=Mb)
    # Titik C : Lentur murni    (Pn=0,  Mn=Mn0)

    # Rasio tulangan - kontrol
    rho_min_col = 0.01
    rho_max_col = 0.08
    ok_rho      = rho_min_col <= rho_g <= rho_max_col

    R.update({
        # Geometri
        "Ab": Ab, "Ast": Ast, "Ag": Ag,
        "d": d, "d_prime": d_prime,
        "beta1": beta1,
        # Rasio tulangan
        "rho_g": rho_g,
        "rho_min": rho_min_col, "rho_max": rho_max_col,
        "ok_rho": ok_rho,
        # Aksial murni
        "Pn0": Pn0 / 1000,              # kN
        "phiPn_max": phiPn_max / 1000,  # kN
        "phi_c": phi_c,
        # Lentur murni (Pn=0)
        "a_flex": a_flex, "c_flex": c_flex,
        "et_flex": et_flex, "phi_f": phi_f,
        "Mn0": Mn0, "phiMn0": phiMn0,
        "n_sisi_tarik": n_sisi_tarik,
        "As_tarik": As_tarik,
        # Balanced
        "cb": cb, "ab": ab,
        "eps_sp": eps_sp, "fsp": fsp,
        "Cc_b": Cc_b, "Cs_b": Cs_b, "T_b": T_b,
        "Pb": Pb, "Mb": Mb,
        "phiPb": phiPb, "phiMb": phiMb,
        "phi_b": phi_b,
    })
    return R


def buat_steps_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h, R):
    """Susun urutan langkah perhitungan kolom sebagai list dict."""
    Ab   = R["Ab"]
    Ast  = R["Ast"]
    Ag   = R["Ag"]
    d    = R["d"]
    dp   = R["d_prime"]

    steps = [
        dict(
            no="Langkah 1", ref="SNI 2847:2019 Pasal 22.2.2.4.3",
            judul="Faktor blok tegangan ekivalen (Beta-1)",
            isi=(
                f"fc = {fc} MPa\n"
                + (f"fc <= 28 MPa  -->  Beta-1 = 0.85" if fc <= 28
                   else f"fc >= 56 MPa  -->  Beta-1 = 0.65" if fc >= 56
                   else f"Beta-1 = 0.85 - 0.05 x (fc - 28) / 7\n"
                        f"       = 0.85 - 0.05 x ({fc} - 28) / 7")
                + f"\n  -->  Beta-1 = {R['beta1']:.4f}"
            ),
            ok=True,
        ),
        dict(
            no="Langkah 2", ref="SNI 2847:2019 Pasal 10.6.1.1",
            judul="Data tulangan longitudinal",
            isi=(
                f"Diameter tulangan  : D{D_tul:.0f}\n"
                f"Luas per batang    : Ab = pi/4 x {D_tul:.0f}^2 = {Ab:.2f} mm2\n"
                f"Jumlah total       : n  = {n_total} batang "
                f"({n_b} sisi lebar + {n_h} sisi tinggi, simetris 4 sisi)\n"
                f"  Sisi lebar (atas/bawah) : {n_b} batang x 2 = {n_b*2} batang\n"
                f"  Sisi tinggi (kiri/kanan): {n_h} batang x 2 = {n_h*2} batang\n"
                f"  Total tulangan sisi      = {n_b*2 + n_h*2} btg "
                f"(sudut terhitung di sisi lebar)\n"
                f"Luas total tulangan: Ast = {n_total} x {Ab:.2f}\n"
                f"  -->  Ast = {Ast:.2f} mm2"
            ),
            ok=True,
        ),
        dict(
            no="Langkah 3", ref="SNI 2847:2019 Pasal 10.6.1.1",
            judul="Rasio tulangan longitudinal (Rho-g)",
            isi=(
                f"Ag    = b x h = {b:.0f} x {h:.0f} = {Ag:.0f} mm2\n"
                f"Rho-g = Ast / Ag\n"
                f"      = {Ast:.2f} / {Ag:.0f}\n"
                f"      = {R['rho_g']:.6f}  =  {R['rho_g']*100:.4f}%\n"
                f"Syarat: 0.01 <= Rho-g <= 0.08\n"
                f"        {R['rho_min']*100:.0f}% <= {R['rho_g']*100:.4f}% <= {R['rho_max']*100:.0f}%  "
                + ("[OK]" if R['ok_rho'] else "[TIDAK OK]")
            ),
            ok=R["ok_rho"],
        ),
        dict(
            no="Langkah 4", ref="SNI 2847:2019 Pasal 22.4.2.2",
            judul="Kapasitas aksial nominal murni (Pn0) - Beban sentris",
            isi=(
                f"Pn0 = 0.85 x fc x (Ag - Ast) + fy x Ast\n"
                f"    = 0.85 x {fc} x ({Ag:.0f} - {Ast:.2f}) + {fy:.0f} x {Ast:.2f}\n"
                f"    = {0.85*fc*(Ag-Ast):,.0f} + {fy*Ast:,.0f}\n"
                f"  -->  Pn0 = {R['Pn0']:,.2f} kN\n"
                f"\nFaktor reduksi kolom ikat (tied): Phi-c = {R['phi_c']}\n"
                f"Batas maks SNI (Phi x 0.80 x Pn0):\n"
                f"  Phi.Pn,max = {R['phi_c']} x 0.80 x {R['Pn0']:,.2f}\n"
                f"  -->  Phi.Pn,max = {R['phiPn_max']:,.2f} kN"
            ),
            ok=True,
        ),
        dict(
            no="Langkah 5", ref="SNI 2847:2019 Pasal 22.2",
            judul="Kapasitas lentur murni (Mn0) - Kondisi Pn = 0",
            isi=(
                f"Tulangan sisi tarik : {R['n_sisi_tarik']} batang (sisi bawah)\n"
                f"As-tarik = {R['n_sisi_tarik']} x {Ab:.2f} = {R['As_tarik']:.2f} mm2\n"
                f"\nBlok tegangan lentur:\n"
                f"a  = As-tarik x fy / (0.85 x fc x b)\n"
                f"   = {R['As_tarik']:.2f} x {fy:.0f} / (0.85 x {fc} x {b:.0f})\n"
                f"   = {R['a_flex']:.2f} mm\n"
                f"c  = a / Beta-1 = {R['a_flex']:.2f} / {R['beta1']:.4f}\n"
                f"   = {R['c_flex']:.2f} mm\n"
                f"et = 0.003 x (d - c) / c\n"
                f"   = 0.003 x ({d:.0f} - {R['c_flex']:.2f}) / {R['c_flex']:.2f}\n"
                f"   = {R['et_flex']:.5f}  -->  Phi = {R['phi_f']:.4f}\n"
                f"\nMn0 = As-tarik x fy x (d - a/2)\n"
                f"    = {R['As_tarik']:.2f} x {fy:.0f} x ({d:.0f} - {R['a_flex']/2:.2f})\n"
                f"  -->  Mn0   = {R['Mn0']:.3f} kN.m\n"
                f"  -->  Phi.Mn0 = {R['phi_f']:.4f} x {R['Mn0']:.3f} = {R['phiMn0']:.3f} kN.m"
            ),
            ok=True,
        ),
        dict(
            no="Langkah 6", ref="SNI 2847:2019 Pasal 22.2.2",
            judul="Kondisi seimbang (Balanced Condition)",
            isi=(
                f"Kedalaman sumbu netral balanced:\n"
                f"cb = [600 / (600 + fy)] x d\n"
                f"   = [600 / (600 + {fy:.0f})] x {d:.0f}\n"
                f"   = {R['cb']:.2f} mm\n"
                f"ab = Beta-1 x cb = {R['beta1']:.4f} x {R['cb']:.2f}\n"
                f"   = {R['ab']:.2f} mm\n"
                f"\nRegangan baja tulangan tekan:\n"
                f"eps-s' = 0.003 x (cb - d') / cb\n"
                f"       = 0.003 x ({R['cb']:.2f} - {dp:.0f}) / {R['cb']:.2f}\n"
                f"       = {R['eps_sp']:.5f}\n"
                f"fs'    = min(fy, eps-s' x Es) = min({fy:.0f}, {R['eps_sp']:.5f} x 200000)\n"
                f"       = {R['fsp']:.2f} MPa\n"
                f"\nGaya-gaya dalam pada kondisi balanced:\n"
                f"Cc = 0.85 x fc x ab x b\n"
                f"   = 0.85 x {fc} x {R['ab']:.2f} x {b:.0f} = {R['Cc_b']:.2f} kN\n"
                f"Cs = As-tekan x fs'\n"
                f"   = {R['As_tarik']:.2f} x {R['fsp']:.2f} = {R['Cs_b']:.2f} kN\n"
                f"T  = As-tarik x fy\n"
                f"   = {R['As_tarik']:.2f} x {fy:.0f} = {R['T_b']:.2f} kN\n"
                f"\nPb = Cc + Cs - T\n"
                f"   = {R['Cc_b']:.2f} + {R['Cs_b']:.2f} - {R['T_b']:.2f}\n"
                f"   = {R['Pb']:.2f} kN\n"
                f"\nMomen balanced (terhadap pusat penampang):\n"
                f"Mb = Cc x (h/2 - ab/2) + Cs x (h/2 - d') + T x (d - h/2)\n"
                f"   = {R['Cc_b']:.2f} x {(h/2-R['ab']/2)/1000:.4f} + "
                f"{R['Cs_b']:.2f} x {(h/2-dp)/1000:.4f} + "
                f"{R['T_b']:.2f} x {(d-h/2)/1000:.4f}\n"
                f"   = {R['Mb']:.3f} kN.m\n"
                f"\nPhi = {R['phi_b']} (kolom ikat, zona tekan)\n"
                f"  -->  Phi.Pb = {R['phiPb']:.2f} kN\n"
                f"  -->  Phi.Mb = {R['phiMb']:.3f} kN.m"
            ),
            ok=True,
        ),
        dict(
            no="Langkah 7", ref="SNI 2847:2019 Pasal 22.4 & 21.2",
            judul="Diagram Interaksi - Tiga Titik Kontrol",
            isi=(
                f"Titik A (Aksial Murni):\n"
                f"  Pn  = {R['Pn0']:.2f} kN\n"
                f"  Mn  = 0 kN.m\n"
                f"  Phi.Pn,max = {R['phiPn_max']:.2f} kN  (Phi x 0.80 x Pn0)\n"
                f"\nTitik B (Balanced):\n"
                f"  Pn  = {R['Pb']:.2f} kN\n"
                f"  Mn  = {R['Mb']:.3f} kN.m\n"
                f"  Phi.Pn = {R['phiPb']:.2f} kN\n"
                f"  Phi.Mn = {R['phiMb']:.3f} kN.m\n"
                f"\nTitik C (Lentur Murni, Pn = 0):\n"
                f"  Pn  = 0 kN\n"
                f"  Mn  = {R['Mn0']:.3f} kN.m\n"
                f"  Phi.Mn0 = {R['phiMn0']:.3f} kN.m\n"
                f"\nKolom aman selama titik (Pu, Mu) berada\n"
                f"DI DALAM kurva diagram interaksi di atas."
            ),
            ok=True,
        ),
    ]
    return steps


# ============================================================
# GENERATOR WORD (.docx) - Format natural ketikan
# ============================================================
def create_word_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h,
                      R, steps, nama_proyek):
    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Cm(21);  sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(3);   sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(3);   sec.bottom_margin = Cm(2.5)

    def par(teks="", bold=False, size=11, indent=0, space_after=4,
            align=WD_ALIGN_PARAGRAPH.LEFT, color=None, italic=False, mono=False):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        if teks:
            r = p.add_run(teks)
            r.bold = bold; r.italic = italic
            r.font.size = Pt(size)
            if color:   r.font.color.rgb = RGBColor(*color)
            if mono:    r.font.name = "Courier New"
        return p

    def garis():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        run = p.add_run("_" * 72)
        run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    def subjudul(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(teks); r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
        garis()

    # Judul
    par("LAPORAN PERHITUNGAN STRUKTUR", bold=True, size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, color=(0x1A, 0x3C, 0x5E))
    par("Kapasitas Aksial dan Lentur Kolom Beton Bertulang", size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    par("Asumsi: Kolom Pendek (Short Column) - Tanpa Kelangsingan", size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, italic=True, color=(0x44, 0x44, 0x44))
    par(f"Referensi  :  SNI 2847:2019 (ACI 318-14)", size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, italic=True, color=(0x55, 0x55, 0x55))
    par(f"Proyek     :  {nama_proyek}", size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, italic=True, color=(0x55, 0x55, 0x55))
    par(f"Tanggal    :  {datetime.datetime.now().strftime('%d %B %Y')}", size=10,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, italic=True, color=(0x55, 0x55, 0x55))
    garis(); par(space_after=6)

    # 1. Data Input
    subjudul("1.  DATA INPUT PENAMPANG KOLOM")
    for simb, nilai in [
        ("fc",      f"{fc:.1f} MPa"),
        ("fy",      f"{fy:.0f} MPa"),
        ("b",       f"{b:.0f} mm"),
        ("h",       f"{h:.0f} mm"),
        ("selimut", f"{cover:.0f} mm  (ke pusat tulangan)"),
        ("D-tul",   f"D{D_tul:.0f}"),
        ("n-total", f"{n_total} batang  ({n_b} sisi lebar + {n_h} sisi tinggi)"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(f"{simb:<12}"); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = "Courier New"
        r2 = p.add_run(f"=  {nilai}"); r2.font.size = Pt(10); r2.font.name = "Courier New"
    par(space_after=6)

    # 2. Analisa Perhitungan
    subjudul("2.  ANALISA PERHITUNGAN")
    par("Urutan perhitungan mengacu pada SNI 2847:2019.", size=10, italic=True,
        color=(0x55, 0x55, 0x55), space_after=8)

    for s in steps:
        p_hdr = doc.add_paragraph()
        p_hdr.paragraph_format.space_before = Pt(8); p_hdr.paragraph_format.space_after = Pt(1)
        r_no  = p_hdr.add_run(f"{s['no']}  "); r_no.bold = True
        r_no.font.size = Pt(10); r_no.font.color.rgb = RGBColor(0x1A, 0x3C, 0x5E)
        r_jdl = p_hdr.add_run(s["judul"]); r_jdl.bold = True; r_jdl.font.size = Pt(10)

        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(0); p_ref.paragraph_format.space_after = Pt(2)
        p_ref.paragraph_format.left_indent  = Cm(0.5)
        r_ref = p_ref.add_run(f"[{s['ref']}]"); r_ref.italic = True
        r_ref.font.size = Pt(8.5); r_ref.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

        for baris in s["isi"].split("\n"):
            p_b = doc.add_paragraph()
            p_b.paragraph_format.space_before = Pt(0); p_b.paragraph_format.space_after = Pt(1)
            p_b.paragraph_format.left_indent  = Cm(0.5)
            r_b = p_b.add_run(baris if baris.strip() else " ")
            r_b.font.name = "Courier New"; r_b.font.size = Pt(9.5)
            if baris.strip().startswith("-->") or "[OK]" in baris or "[TIDAK OK]" in baris:
                r_b.bold = True
                r_b.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if s["ok"]
                                      else RGBColor(0xB7, 0x1C, 0x1C))

    par(space_after=6)

    # 3. Rangkuman
    subjudul("3.  RANGKUMAN HASIL")
    rangkuman = [
        ("Beta-1",       f"{R['beta1']:.4f}",              ""),
        ("Ag",           f"{R['Ag']:.0f} mm2",             "Luas penampang bruto"),
        ("Ast",          f"{R['Ast']:.2f} mm2",            "Luas tulangan total"),
        ("Rho-g",        f"{R['rho_g']*100:.4f}%",         "Rasio tulangan"),
        ("Pn0",          f"{R['Pn0']:.2f} kN",             "Aksial murni nominal"),
        ("Phi.Pn,max",   f"{R['phiPn_max']:.2f} kN",       "Aksial maks rencana"),
        ("Mn0",          f"{R['Mn0']:.3f} kN.m",           "Lentur murni nominal"),
        ("Phi.Mn0",      f"{R['phiMn0']:.3f} kN.m",        "Lentur murni rencana"),
        ("Pb (balanced)",f"{R['Pb']:.2f} kN",              "Aksial balanced"),
        ("Mb (balanced)",f"{R['Mb']:.3f} kN.m",            "Lentur balanced"),
        ("Phi.Pb",       f"{R['phiPb']:.2f} kN",           ""),
        ("Phi.Mb",       f"{R['phiMb']:.3f} kN.m",         ""),
    ]
    for simb, nilai, ket in rangkuman:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(f"{simb:<18}"); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = "Courier New"
        r2 = p.add_run(f"=  {nilai:<22}  {ket}")
        r2.font.size = Pt(10); r2.font.name = "Courier New"
    par(space_after=6)

    # 4. Kontrol
    subjudul("4.  KONTROL PENAMPANG")
    kontrol = [
        (f"Rho-g = {R['rho_g']*100:.4f}%  dalam batas 1% s/d 8%", R["ok_rho"]),
    ]
    for tk, ok_k in kontrol:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.5)
        r_k = p.add_run(f"{tk}   --> {'[OK]' if ok_k else '[TIDAK OK]'}")
        r_k.font.name = "Courier New"; r_k.font.size = Pt(10); r_k.bold = True
        r_k.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if ok_k
                              else RGBColor(0xB7, 0x1C, 0x1C))
    par(space_after=4)

    kes_ok  = R["ok_rho"]
    kes_txt = (
        "KESIMPULAN : Penampang kolom OK - Rasio tulangan memenuhi syarat SNI 2847:2019.\n"
        "             Cek titik (Pu, Mu) terhadap diagram interaksi di atas."
        if kes_ok else
        "KESIMPULAN : Penampang kolom PERLU DIREVISI - Rasio tulangan di luar batas."
    )
    p_kes = doc.add_paragraph()
    p_kes.paragraph_format.space_before = Pt(6); p_kes.paragraph_format.space_after = Pt(4)
    for baris_kes in kes_txt.split("\n"):
        if baris_kes != kes_txt.split("\n")[0]:
            p_kes = doc.add_paragraph()
            p_kes.paragraph_format.space_before = Pt(0)
            p_kes.paragraph_format.space_after  = Pt(4)
        r_kes = p_kes.add_run(baris_kes); r_kes.bold = True; r_kes.font.size = Pt(10.5)
        r_kes.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if kes_ok
                                else RGBColor(0xB7, 0x1C, 0x1C))

    garis()
    par("Referensi: SNI 2847:2019 | ACI 318-14  --  "
        "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan.",
        size=8, italic=True, color=(0x99, 0x99, 0x99),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf


# ============================================================
# GENERATOR PDF (.pdf) - Formal + Watermark LADOSI ENGINEERING
# ============================================================
WATERMARK_TEXT = "DIHASILKAN OLEH: LADOSI ENGINEERING"
BRAND_COLOR    = (26, 60, 94)
OK_COLOR       = (27, 94, 32)
FAIL_COLOR     = (183, 28, 28)
GRAY           = (120, 120, 120)


class LaporanKolomPDF(FPDF):
    def __init__(self, nama_proyek):
        super().__init__()
        self.nama_proyek = sp(nama_proyek)
        self.set_margins(25, 25, 20)
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        self.set_draw_color(*BRAND_COLOR)
        self.set_line_width(0.8)
        self.line(25, 15, 190, 15)
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(*BRAND_COLOR)
        self.set_xy(25, 17)
        self.cell(0, 5, sp("LAPORAN PERHITUNGAN KOLOM BETON  |  SNI 2847:2019"),
                  ln=False, align="L")
        self.set_font("Helvetica", "", 8)
        self.set_text_color(*GRAY)
        self.set_xy(25, 17)
        self.cell(0, 5, sp(f"Proyek: {self.nama_proyek}"), ln=False, align="R")
        self.ln(10)

    def footer(self):
        self.set_y(-18)
        self.set_draw_color(*BRAND_COLOR)
        self.set_line_width(0.4)
        self.line(25, self.get_y(), 190, self.get_y())
        self.set_font("Helvetica", "I", 7.5)
        self.set_text_color(*GRAY)
        self.cell(0, 6,
            sp("Referensi: SNI 2847:2019 | ACI 318-14 - "
               "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan."),
            align="C")
        self.set_xy(25, self.get_y())
        self.set_font("Helvetica", "", 7.5)
        self.cell(0, 6, sp(f"Halaman {self.page_no()}"), align="R")

    def watermark(self):
        self.set_font("Helvetica", "B", 28)
        self.set_text_color(210, 215, 220)
        xc, yc = self.w / 2, self.h / 2
        with self.rotation(40, xc, yc):
            self.set_xy(xc - 65, yc - 6)
            self.cell(130, 12, sp(WATERMARK_TEXT), align="C")
        self.set_text_color(0, 0, 0)

    def section_title(self, teks):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(*BRAND_COLOR)
        self.ln(4)
        self.cell(0, 7, sp(teks), ln=True)
        self.set_draw_color(*BRAND_COLOR)
        self.set_line_width(0.4)
        self.line(self.get_x(), self.get_y(), 190, self.get_y())
        self.ln(3)
        self.set_text_color(0, 0, 0)

    def mono_line(self, teks, bold=False, color=None):
        self.set_font("Courier", "B" if bold else "", 9)
        if color:
            self.set_text_color(*color)
        self.set_x(28)
        self.multi_cell(0, 4.5, sp(teks))
        self.set_text_color(0, 0, 0)


def create_pdf_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h,
                     R, steps, nama_proyek):
    pdf = LaporanKolomPDF(nama_proyek)
    pdf.add_page()
    pdf.watermark()

    # Judul
    pdf.set_font("Helvetica", "B", 15)
    pdf.set_text_color(*BRAND_COLOR)
    pdf.ln(2)
    pdf.cell(0, 9, sp("LAPORAN PERHITUNGAN STRUKTUR"), ln=True, align="C")
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, sp("Kapasitas Aksial dan Lentur Kolom Beton Bertulang"), ln=True, align="C")
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(*GRAY)
    pdf.cell(0, 5, sp("Asumsi: Kolom Pendek (Short Column) - Tanpa Kelangsingan"), ln=True, align="C")
    pdf.cell(0, 5, sp("Referensi: SNI 2847:2019 (ACI 318-14)"), ln=True, align="C")
    pdf.cell(0, 5, sp(f"Proyek: {nama_proyek}   |   Tanggal: {datetime.datetime.now().strftime('%d %B %Y')}"),
             ln=True, align="C")
    pdf.ln(5)
    pdf.set_draw_color(*BRAND_COLOR); pdf.set_line_width(0.6)
    pdf.line(25, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)

    # 1. Data Input
    pdf.section_title(sp("1.  DATA INPUT PENAMPANG KOLOM"))
    data_input = [
        ("fc",       f"{fc:.1f} MPa",    "Kuat tekan beton"),
        ("fy",       f"{fy:.0f} MPa",    "Kuat leleh tulangan"),
        ("b",        f"{b:.0f} mm",      "Lebar kolom"),
        ("h",        f"{h:.0f} mm",      "Tinggi kolom"),
        ("selimut",  f"{cover:.0f} mm",  "Selimut ke pusat tulangan"),
        ("D-tul",    f"D{D_tul:.0f}",    "Diameter tulangan longitudinal"),
        ("n-total",  f"{n_total} btg",   f"{n_b} sisi lebar + {n_h} sisi tinggi"),
    ]
    for simb, nilai, ket in data_input:
        pdf.set_x(28)
        pdf.set_font("Courier", "B", 9.5)
        pdf.cell(22, 5, sp(f"{simb:<10}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(32, 5, sp(f"=  {nilai}"), ln=False)
        pdf.set_font("Helvetica", "I", 8.5)
        pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, sp(f"({ket})"), ln=True)
        pdf.set_text_color(0, 0, 0)
    pdf.ln(4)

    # 2. Analisa Perhitungan
    pdf.section_title(sp("2.  ANALISA PERHITUNGAN"))
    for s in steps:
        if pdf.get_y() > 240:
            pdf.add_page(); pdf.watermark()
        pdf.set_font("Helvetica", "B", 10)
        pdf.set_text_color(*BRAND_COLOR)
        pdf.set_x(25)
        pdf.cell(0, 6, sp(f"{s['no']}  {s['judul']}"), ln=True)
        pdf.set_font("Helvetica", "I", 8)
        pdf.set_text_color(*GRAY)
        pdf.set_x(28)
        pdf.cell(0, 4, sp(f"[{s['ref']}]"), ln=True)
        pdf.set_text_color(0, 0, 0)
        for baris in s["isi"].split("\n"):
            is_result = (baris.strip().startswith("-->") or
                         "[OK]" in baris or "[TIDAK OK]" in baris)
            if is_result:
                pdf.mono_line(baris, bold=True,
                              color=OK_COLOR if s["ok"] else FAIL_COLOR)
            else:
                pdf.mono_line(baris if baris.strip() else "")
        pdf.ln(2)

    # 3. Rangkuman
    if pdf.get_y() > 210:
        pdf.add_page(); pdf.watermark()
    pdf.section_title(sp("3.  RANGKUMAN HASIL"))
    rangkuman = [
        ("Beta-1",       f"{R['beta1']:.4f}",              ""),
        ("Ag",           f"{R['Ag']:.0f} mm2",             "Luas penampang bruto"),
        ("Ast",          f"{R['Ast']:.2f} mm2",            "Luas tulangan total"),
        ("Rho-g",        f"{R['rho_g']*100:.4f}%",         "Rasio tulangan"),
        ("Pn0",          f"{R['Pn0']:.2f} kN",             "Aksial murni nominal"),
        ("Phi.Pn,max",   f"{R['phiPn_max']:.2f} kN",       "Aksial maks rencana"),
        ("Mn0",          f"{R['Mn0']:.3f} kN.m",           "Lentur murni nominal"),
        ("Phi.Mn0",      f"{R['phiMn0']:.3f} kN.m",        "Lentur murni rencana"),
        ("Pb",           f"{R['Pb']:.2f} kN",              "Aksial balanced"),
        ("Mb",           f"{R['Mb']:.3f} kN.m",            "Lentur balanced"),
        ("Phi.Pb",       f"{R['phiPb']:.2f} kN",           ""),
        ("Phi.Mb",       f"{R['phiMb']:.3f} kN.m",         ""),
    ]
    for simb, nilai, ket in rangkuman:
        pdf.set_x(28)
        pdf.set_font("Courier", "B", 9.5)
        pdf.cell(26, 5, sp(f"{simb:<14}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(38, 5, sp(f"=  {nilai}"), ln=False)
        if ket:
            pdf.set_font("Helvetica", "I", 8.5); pdf.set_text_color(*GRAY)
            pdf.cell(0, 5, sp(f"({ket})"), ln=True); pdf.set_text_color(0, 0, 0)
        else:
            pdf.ln()
    pdf.ln(4)

    # 4. Kontrol
    pdf.section_title(sp("4.  KONTROL PENAMPANG"))
    teks_ctrl = f"Rho-g = {R['rho_g']*100:.4f}%  dalam batas 1% s/d 8%"
    ok_ctrl   = R["ok_rho"]
    tanda     = "[OK]" if ok_ctrl else "[TIDAK OK]"
    pdf.set_x(28)
    pdf.set_font("Courier", "B", 9.5)
    pdf.set_text_color(*(OK_COLOR if ok_ctrl else FAIL_COLOR))
    pdf.cell(0, 5.5, sp(f"{teks_ctrl}   --> {tanda}"), ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    kes_ok = R["ok_rho"]
    kes    = (
        "KESIMPULAN: Penampang kolom OK - Rasio tulangan memenuhi syarat SNI 2847:2019. "
        "Cek titik (Pu, Mu) terhadap diagram interaksi."
        if kes_ok else
        "KESIMPULAN: Penampang kolom PERLU DIREVISI - Rasio tulangan di luar batas SNI."
    )
    pdf.set_x(25)
    pdf.set_font("Helvetica", "B", 10.5)
    pdf.set_text_color(*(OK_COLOR if kes_ok else FAIL_COLOR))
    pdf.multi_cell(0, 6, sp(kes))
    pdf.set_text_color(0, 0, 0)

    buf = io.BytesIO()
    pdf.output(buf); buf.seek(0)
    return buf


# ============================================================
# UI - HEADER
# ============================================================
st.markdown('<p class="main-title"> Kapasitas Aksial & Lentur Kolom Beton Bertulang</p>',
            unsafe_allow_html=True)
st.markdown(
    '<p class="sub-title">Referensi: SNI 2847:2019 (setara ACI 318-14) '
    '| Asumsi: Kolom Pendek (Short Column) - tanpa kelangsingan</p>',
    unsafe_allow_html=True)

st.markdown(
    '<div class="asumsi-box">'
    '<b>Asumsi & Batasan Modul Ini:</b><br>'
    '1. Kolom pendek (short column) - efek kelangsingan diabaikan.<br>'
    '2. Penampang persegi, tulangan simetris 4 sisi.<br>'
    '3. Output berupa 3 titik kontrol diagram interaksi: '
    'Aksial Murni (A), Balanced (B), dan Lentur Murni (C).<br>'
    '4. Untuk verifikasi lengkap, plot titik (Pu, Mu) terhadap diagram interaksi.'
    '</div>',
    unsafe_allow_html=True)

st.markdown('<hr class="divider">', unsafe_allow_html=True)

# ============================================================
# UI - INPUT
# ============================================================
col_inp, col_out = st.columns([1, 2], gap="large")

with col_inp:
    st.markdown("###  Data Input Kolom")

    nama_proyek = st.text_input(
        "Nama Proyek (untuk header laporan)",
        value="Laporan Analisa Struktur",
    )

    st.markdown("**Material**")
    fc = st.number_input("f'c - Kuat tekan beton (MPa)",
                         min_value=17.0, max_value=100.0, value=30.0, step=1.0, format="%.1f")
    fy = st.number_input("fy - Kuat leleh tulangan (MPa)",
                         min_value=240.0, max_value=600.0, value=400.0, step=10.0, format="%.0f")

    st.markdown("**Geometri Penampang**")
    cb2, ch2 = st.columns(2)
    with cb2:
        b = st.number_input("b (mm) - Lebar", min_value=150.0, max_value=3000.0,
                            value=400.0, step=25.0, format="%.0f")
    with ch2:
        h = st.number_input("h (mm) - Tinggi", min_value=150.0, max_value=3000.0,
                            value=400.0, step=25.0, format="%.0f")
    cover = st.number_input(
        "Selimut ke pusat tulangan (mm)",
        min_value=30.0, max_value=150.0, value=60.0, step=5.0, format="%.0f",
        help="= tebal selimut + diameter sengkang + setengah diameter tulangan longitudinal"
    )

    st.markdown("**Tulangan Longitudinal - Simetris 4 Sisi**")
    pilih_d = st.selectbox("Diameter tulangan", options=list(DIAM_OPTS.keys()), index=4)
    D_tul, Ab_per = DIAM_OPTS[pilih_d]
    st.caption(f"Luas per batang: {Ab_per:.2f} mm2")

    st.markdown("Jumlah tulangan per sisi:")
    cn_b, cn_h = st.columns(2)
    with cn_b:
        n_b = st.number_input(
            "Sisi lebar (atas & bawah)",
            min_value=2, max_value=20, value=3, step=1,
            help="Termasuk tulangan sudut. Min = 2 (hanya sudut)"
        )
    with cn_h:
        n_h = st.number_input(
            "Sisi tinggi (kiri & kanan)",
            min_value=0, max_value=20, value=1, step=1,
            help="Tulangan antara dua sudut, tidak termasuk sudut"
        )

    # Hitung total otomatis
    # sudut 4 buah (terhitung di n_b: 2 sudut per sisi lebar x 2 sisi = 4)
    n_total = n_b * 2 + n_h * 2
    Ast_total = n_total * Ab_per
    st.markdown(f"""
    <div style='background:#f0f4ff;border-radius:8px;padding:10px 14px;
                font-size:.85rem;border-left:3px solid #3f51b5;margin-top:4px'>
    <b>Total tulangan : {n_total} batang</b><br>
    Sisi lebar atas  : {n_b} btg &nbsp;|&nbsp; Sisi lebar bawah : {n_b} btg<br>
    Sisi tinggi kiri : {n_h} btg &nbsp;|&nbsp; Sisi tinggi kanan: {n_h} btg<br>
    <b>Ast total = {Ast_total:.1f} mm2</b>
    </div>
    """, unsafe_allow_html=True)

    if n_total < 4:
        st.warning(" Jumlah tulangan minimum untuk kolom adalah 4 batang.")

    st.markdown("")
    tombol = st.button(" HITUNG KAPASITAS KOLOM",
                       use_container_width=True, type="primary")

    with st.expander(" Tabel luas tulangan (mm2)"):
        st.markdown("""
        | O | 1 | 2 | 3 | 4 | 6 | 8 |
        |---|---|---|---|---|---|---|
        | D13 | 132.7 | 265 | 398 | 531 | 796 | 1061 |
        | D16 | 201.1 | 402 | 603 | 804 | 1206 | 1608 |
        | D19 | 283.5 | 567 | 851 | 1134 | 1701 | 2268 |
        | D22 | 380.1 | 760 | 1140 | 1520 | 2281 | 3041 |
        | D25 | 490.9 | 982 | 1473 | 1964 | 2945 | 3927 |
        | D29 | 660.5 | 1321 | 1981 | 2642 | 3963 | 5284 |
        | D32 | 804.2 | 1608 | 2413 | 3217 | 4825 | 6434 |
        """)


# ============================================================
# UI - HASIL
# ============================================================
with col_out:
    if tombol:
        # Validasi
        if n_total < 4:
            st.error(" Jumlah tulangan minimum adalah 4 batang!"); st.stop()
        if cover >= h / 2 or cover >= b / 2:
            st.error(" Selimut terlalu besar relatif terhadap dimensi kolom!"); st.stop()

        R     = hitung_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h)
        steps = buat_steps_kolom(fc, fy, b, h, cover, D_tul, n_total, n_b, n_h, R)

        # -- Metrik Utama -------------------------------------
        st.markdown("###  Hasil Utama")

        m1, m2, m3 = st.columns(3)
        for col, lbl, val, unt in [
            (m1, "Phi.Pn,max (Aksial Maks)", f"{R['phiPn_max']:.1f}", "kN"),
            (m2, "Phi.Mn0 (Lentur Murni)",   f"{R['phiMn0']:.2f}",   "kN.m"),
            (m3, "Rho-g (Rasio Tulangan)",    f"{R['rho_g']*100:.3f}", "%"),
        ]:
            with col:
                st.markdown(
                    f'<div class="metric-card">'
                    f'<div class="metric-lbl">{lbl}</div>'
                    f'<div class="metric-val">{val}</div>'
                    f'<div class="metric-unt">{unt}</div>'
                    f'</div>', unsafe_allow_html=True)

        m4, m5, m6 = st.columns(3)
        for col, lbl, val, unt in [
            (m4, "Phi.Pb (Aksial Balanced)",  f"{R['phiPb']:.1f}",    "kN"),
            (m5, "Phi.Mb (Lentur Balanced)",  f"{R['phiMb']:.2f}",    "kN.m"),
            (m6, "Ast Total",                  f"{R['Ast']:.0f}",      "mm2"),
        ]:
            with col:
                st.markdown(
                    f'<div class="metric-card">'
                    f'<div class="metric-lbl">{lbl}</div>'
                    f'<div class="metric-val">{val}</div>'
                    f'<div class="metric-unt">{unt}</div>'
                    f'</div>', unsafe_allow_html=True)
        st.markdown("")

        # -- Status Rasio Tulangan -----------------------------
        if R["ok_rho"]:
            st.markdown(
                '<div class="result-ok">[OK] RASIO TULANGAN OK - '
                f'Rho-g = {R["rho_g"]*100:.4f}% memenuhi syarat 1% s/d 8% (SNI 2847:2019)</div>',
                unsafe_allow_html=True)
        else:
            msg = ("Rho-g terlalu kecil (< 1%) - tambah tulangan"
                   if R["rho_g"] < 0.01 else
                   "Rho-g terlalu besar (> 8%) - kurangi tulangan atau perbesar penampang")
            st.markdown(
                f'<div class="result-fail">[X] RASIO TULANGAN TIDAK OK - {msg}</div>',
                unsafe_allow_html=True)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)

        # -- Diagram Interaksi (tabel 3 titik) ----------------
        st.markdown("###  Titik Kontrol Diagram Interaksi")
        df_di = pd.DataFrame({
            "Titik": ["A - Aksial Murni", "B - Balanced", "C - Lentur Murni"],
            "Pn (kN)":    [f"{R['Pn0']:.2f}", f"{R['Pb']:.2f}", "0"],
            "Mn (kN.m)":  ["0", f"{R['Mb']:.3f}", f"{R['Mn0']:.3f}"],
            "Phi":        [f"{R['phi_c']}", f"{R['phi_b']}", f"{R['phi_f']:.4f}"],
            "Phi.Pn (kN)": [f"{R['phiPn_max']:.2f}", f"{R['phiPb']:.2f}", "0"],
            "Phi.Mn (kN.m)":["0", f"{R['phiMb']:.3f}", f"{R['phiMn0']:.3f}"],
        })
        st.dataframe(df_di, use_container_width=True, hide_index=True)
        st.caption(
            "Titik (Pu, Mu) dari analisa struktur harus berada **di dalam** kurva "
            "yang menghubungkan ketiga titik di atas agar penampang aman."
        )

        st.markdown('<hr class="divider">', unsafe_allow_html=True)

        # -- Langkah Perhitungan -------------------------------
        st.markdown("###  Urutan Perhitungan")
        for s in steps:
            warna = "#2e7d32" if s["ok"] else "#c62828"
            tanda = "v" if s["ok"] else "x"
            st.markdown(
                f'<div class="step-box" style="border-left-color:{warna}">'
                f'<div class="ref-badge">{s["ref"]}</div><br>'
                f'<div class="step-hdr">{s["no"]} - {s["judul"]} &nbsp; {tanda}</div>'
                f'<pre style="margin:0;font-size:.82rem;white-space:pre-wrap;'
                f'font-family:monospace">{s["isi"]}</pre></div>',
                unsafe_allow_html=True)

        # -- Tabel Rangkuman -----------------------------------
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("###  Tabel Rangkuman")
        rho_s = "[OK] OK" if R["ok_rho"] else "[X] Tidak OK"
        df_rng = pd.DataFrame({
            "Parameter": [
                "Beta-1", "Ag (mm2)", "Ast (mm2)", "Rho-g (%)",
                "Pn0 (kN)", "Phi.Pn,max (kN)",
                "Mn0 (kN.m)", "Phi.Mn0 (kN.m)",
                "Pb (kN)", "Mb (kN.m)", "Phi.Pb (kN)", "Phi.Mb (kN.m)",
            ],
            "Nilai": [
                f"{R['beta1']:.4f}",
                f"{R['Ag']:.0f}",
                f"{R['Ast']:.2f}",
                f"{R['rho_g']*100:.4f}",
                f"{R['Pn0']:.2f}",
                f"{R['phiPn_max']:.2f}",
                f"{R['Mn0']:.3f}",
                f"{R['phiMn0']:.3f}",
                f"{R['Pb']:.2f}",
                f"{R['Mb']:.3f}",
                f"{R['phiPb']:.2f}",
                f"{R['phiMb']:.3f}",
            ],
            "Status": [
                "-", "-", "-", rho_s,
                "-", "-", "-", "-",
                "-", "-", "-", "-",
            ],
        })
        st.dataframe(df_rng, use_container_width=True, hide_index=True)

        # ====================================================
        # TOMBOL DOWNLOAD
        # ====================================================
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("###  Download Laporan")

        nama_base = (f"Laporan_Kolom_fc{int(fc)}_fy{int(fy)}"
                     f"_b{int(b)}x{int(h)}_n{n_total}D{int(D_tul)}")

        word_buf = create_word_kolom(
            fc, fy, b, h, cover, D_tul, n_total, n_b, n_h,
            R, steps, nama_proyek)
        pdf_buf  = create_pdf_kolom(
            fc, fy, b, h, cover, D_tul, n_total, n_b, n_h,
            R, steps, nama_proyek)

        dl_w, dl_p = st.columns(2)
        with dl_w:
            st.download_button(
                label="  Download Laporan Word (.docx)",
                data=word_buf,
                file_name=f"{nama_base}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                help="Format natural, siap diedit di Microsoft Word.",
            )
        with dl_p:
            st.download_button(
                label="  Download Laporan PDF (.pdf)",
                data=pdf_buf,
                file_name=f"{nama_base}.pdf",
                mime="application/pdf",
                use_container_width=True,
                help="Layout formal dengan watermark LADOSI ENGINEERING.",
            )
        st.caption(
            "File Word dapat diedit lebih lanjut. "
            "File PDF sudah dilengkapi watermark dan siap untuk laporan resmi."
        )

    else:
        st.info("  Isi data di panel kiri, lalu klik **HITUNG KAPASITAS KOLOM**")
        st.markdown("""
        **Yang akan dihitung (Kolom Pendek, Simetris 4 Sisi):**
        1. Faktor Beta1 (blok tegangan Whitney)
        2. Data tulangan longitudinal dan Ast
        3. Rasio tulangan longitudinal Rho-g (kontrol 1%-8%)
        4. Kapasitas aksial murni Pn0 dan Phi.Pn,max
        5. Kapasitas lentur murni Mn0 dan Phi.Mn0
        6. Kondisi balanced: Pb, Mb, Phi.Pb, Phi.Mb
        7. Tiga titik kontrol diagram interaksi (P-M)

        **Output tersedia dalam dua format:**
        -  Word (.docx) - format natural, bisa diedit
        -  PDF (.pdf)  - formal + watermark LADOSI ENGINEERING
        """)

# Footer
st.markdown('<hr class="divider">', unsafe_allow_html=True)
st.markdown(
    "<p style='text-align:center;font-size:.75rem;color:#aaa'>"
    "Referensi: SNI 2847:2019 | ACI 318-14 | "
    "Untuk keperluan profesional - verifikasi mandiri tetap diperlukan"
    "</p>",
    unsafe_allow_html=True,
)
