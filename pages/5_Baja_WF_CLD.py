"""
Perhitungan Kapasitas Penampang Baja Profil WF
Referensi: SNI 1729:2020 & AISC 360-16 | Metode: LRFD & ASD
File: pages/4_Baja_WF.py
"""

import streamlit as st
import pandas as pd
import numpy as np

from io import BytesIO
import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from fpdf import FPDF
# ============================================================
# FUNGSI PERHITUNGAN
# ============================================================

def hitung_properties(d, bf, tw, tf, r):
    """Menghitung properties penampang WF."""
    hw = d - 2 * tf          # Tinggi badan bersih
    h  = d - 2 * (tf + r)   # Tinggi badan bersih (untuk k = tf + r)

    # Luas penampang
    A_sayap = 2 * bf * tf
    A_badan = hw * tw
    # Luas fillet (perkiraan: 2 * 0.8584 * r^2 per sisi → 4 sudut)
    A_fillet = 4 * 0.8584 * r**2
    A = A_sayap + A_badan + A_fillet

    # Inersia Ix
    Ix_sayap = 2 * (bf * tf**3 / 12 + bf * tf * ((d - tf) / 2)**2)
    Ix_badan = tw * hw**3 / 12
    Ix_fillet = 4 * (0.1098 * r**4 + 0.8584 * r**2 * (d/2 - tf - 0.4714*r)**2)
    Ix = Ix_sayap + Ix_badan + Ix_fillet

    # Inersia Iy
    Iy_sayap = 2 * (tf * bf**3 / 12)
    Iy_badan = hw * tw**3 / 12
    Iy = Iy_sayap + Iy_badan

    # Modulus elastis
    Sx = Ix / (d / 2)
    Sy = Iy / (bf / 2)

    # Modulus plastis Zx (WF simetris)
    # Zx = A/2 * jarak antar titik berat setengah penampang
    # Pendekatan praktis untuk profil WF:
    ybar_atas = (bf * tf * (d/2 - tf/2) + hw/2 * tw * hw/4) / (bf * tf + hw/2 * tw)
    Zx = 2 * (bf * tf * (d/2 - tf/2) + tw * (hw/2)**2 / 2)

    # Zy
    Zy = tf * bf**2 / 2 + tw**2 * hw / 4

    # Radius girasi
    rx = np.sqrt(Ix / A)
    ry = np.sqrt(Iy / A)

    # Konstanta torsi J (Saint-Venant)
    # J ≈ (1/3) * [2*bf*tf^3 + hw*tw^3] (tanpa koreksi fillet untuk simplifikasi)
    J = (1/3) * (2 * bf * tf**3 + hw * tw**3)

    # Konstanta warping Cw
    ho = d - tf  # jarak antar titik berat sayap
    Cw = (Iy * ho**2) / 4

    return {
        "hw": hw, "h": h, "A": A,
        "Ix": Ix, "Iy": Iy,
        "Sx": Sx, "Sy": Sy,
        "Zx": Zx, "Zy": Zy,
        "rx": rx, "ry": ry,
        "J": J, "Cw": Cw, "ho": ho
    }


def klasifikasi_penampang(bf, tw, tf, hw, Fy, E):
    """Klasifikasi kompak/non-kompak/langsing untuk sayap dan badan."""
    # Sayap (flanges in flexure): λ = bf/(2*tf)
    lam_f  = bf / (2 * tf)
    lpf    = 0.38 * np.sqrt(E / Fy)           # batas kompak
    lrf    = 1.0  * np.sqrt(E / Fy)           # batas non-kompak (AISC T.B4.1b)

    if lam_f <= lpf:
        klas_f = "Kompak"
    elif lam_f <= lrf:
        klas_f = "Non-Kompak"
    else:
        klas_f = "Langsing"

    # Badan (web in flexure): λ = hw/tw
    lam_w  = hw / tw
    lpw    = 3.76 * np.sqrt(E / Fy)
    lrw    = 5.70 * np.sqrt(E / Fy)

    if lam_w <= lpw:
        klas_w = "Kompak"
    elif lam_w <= lrw:
        klas_w = "Non-Kompak"
    else:
        klas_w = "Langsing"

    return {
        "lam_f": lam_f, "lpf": lpf, "lrf": lrf, "klas_f": klas_f,
        "lam_w": lam_w, "lpw": lpw, "lrw": lrw, "klas_w": klas_w
    }


def hitung_momen(d, bf, tw, tf, r, Fy, Fu, E, Lb, Cb, props, klas):
    """Hitung kapasitas momen Mn per AISC 360-16 Chapter F."""
    Zx  = props["Zx"]
    Sx  = props["Sx"]
    Iy  = props["Iy"]
    J   = props["J"]
    Cw  = props["Cw"]
    ho  = props["ho"]
    ry  = props["ry"]

    # Mp dan Mr
    Mp = Fy * Zx
    Mr = 0.7 * Fy * Sx

    # rts² = √(Iy*Cw) / Sx
    rts = np.sqrt(np.sqrt(Iy * Cw) / Sx)

    # c = 1 untuk profil WF simetri ganda
    c = 1.0

    # Lp = 1.76 * ry * √(E/Fy)
    Lp = 1.76 * ry * np.sqrt(E / Fy)

    # Lr = 1.95 * rts * E/(0.7*Fy) * √(J*c/(Sx*ho) + √((J*c/(Sx*ho))²+6.76*(0.7*Fy/E)²))
    term1 = J * c / (Sx * ho)
    term2 = np.sqrt(term1**2 + 6.76 * (0.7 * Fy / E)**2)
    Lr = 1.95 * rts * (E / (0.7 * Fy)) * np.sqrt(term1 + term2)

    # Tentukan kondisi & Mn
    if Lb <= Lp:
        kondisi = "Plastis (Lb ≤ Lp)"
        Mn = Mp
        Fcr = None
    elif Lb <= Lr:
        kondisi = "Inelastis (Lp < Lb ≤ Lr)"
        Mn = Cb * (Mp - (Mp - Mr) * ((Lb - Lp) / (Lr - Lp)))
        Mn = min(Mn, Mp)
        Fcr = None
    else:
        kondisi = "Elastis (Lb > Lr)"
        # Fcr = Cb*π²*E / (Lb/rts)² * √(1 + 0.078*J*c/(Sx*ho)*(Lb/rts)²)
        ratio_lt = Lb / rts
        Fcr = (Cb * np.pi**2 * E / ratio_lt**2) * np.sqrt(1 + 0.078 * J * c / (Sx * ho) * ratio_lt**2)
        Mn = min(Fcr * Sx, Mp)
        kondisi = "Elastis (Lb > Lr)"

    # Kapasitas akhir (tidak lebih dari Mp untuk kasus non-kompak penampang)
    # Untuk sayap non-kompak: Mn perlu dikurangi (FLB), tapi di sini fokus LTB
    # Catatan: jika sayap non-kompak, perlu cek FLB; di sini asumsi kompak/LTB govern
    Mn_final = min(Mn, Mp)

    return {
        "Mp": Mp, "Mr": Mr, "Lp": Lp, "Lr": Lr,
        "rts": rts, "ho": ho, "c": c,
        "Fcr": Fcr, "kondisi": kondisi,
        "Mn": Mn_final
    }


def hitung_geser(d, tw, tf, hw, Fy, E):
    """Hitung kapasitas geser Vn per AISC 360-16 Chapter G."""
    Aw = d * tw
    kv = 5.34  # untuk badan tanpa pengaku transversal (a/h > 3)

    ratio = hw / tw
    batas_cv1 = 2.24 * np.sqrt(E / Fy)

    # Cv1
    if ratio <= batas_cv1:
        Cv1 = 1.0
        catatan_cv1 = "hw/tw ≤ 2.24√(E/Fy) → Cv1 = 1.0"
    else:
        # Cv1 dari kv
        batas1 = 1.10 * np.sqrt(kv * E / Fy)
        batas2 = 1.37 * np.sqrt(kv * E / Fy)
        if ratio <= batas1:
            Cv1 = 1.0
            catatan_cv1 = f"hw/tw ≤ 1.10√(kv·E/Fy)={batas1:.1f} → Cv1=1.0"
        elif ratio <= batas2:
            Cv1 = batas1 / ratio
            catatan_cv1 = f"1.10√(kv·E/Fy) < hw/tw ≤ 1.37√ → Cv1={Cv1:.3f}"
        else:
            Cv1 = 1.51 * kv * E / (ratio**2 * Fy)
            catatan_cv1 = f"hw/tw > 1.37√(kv·E/Fy) → Cv1={Cv1:.3f}"

    Vn = 0.6 * Fy * Aw * Cv1

    return {
        "Aw": Aw, "kv": kv, "ratio": ratio,
        "batas_cv1": batas_cv1,
        "Cv1": Cv1, "catatan_cv1": catatan_cv1,
        "Vn": Vn
    }


# ============================================================
# HELPER: buat tabel DataFrame
# ============================================================

def buat_df(rows):
    """Buat DataFrame dari list of (Parameter, Nilai, Satuan, Keterangan)."""
    df = pd.DataFrame(rows, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])
    return df


def format_val(v, dec=3):
    if v is None:
        return "-"
    if abs(v) >= 1e9:
        return f"{v/1e9:.{dec}f} ×10⁹"
    if abs(v) >= 1e6:
        return f"{v/1e6:.{dec}f} ×10⁶"
    if abs(v) >= 1e3:
        return f"{v/1e3:.{dec}f} ×10³"
    return f"{v:.{dec}f}"
# ── SANITASI UNICODE → ASCII (untuk fpdf2) ─────────────────
def sanitasi_pdf(teks: str) -> str:
    """
    Ganti karakter Unicode/Yunani yang tidak didukung font PDF standar
    menjadi representasi ASCII yang terbaca.
    """
    if not isinstance(teks, str):
        teks = str(teks)
    tabel = [
        # Simbol Yunani & matematika
        ("φ",  "phi"),   ("Φ",  "Phi"),
        ("λ",  "lambda"),("Λ",  "Lambda"),
        ("Ω",  "Omega"), ("ω",  "omega"),
        ("π",  "pi"),    ("Π",  "Pi"),
        ("α",  "alpha"), ("β",  "beta"),
        ("γ",  "gamma"), ("δ",  "delta"),
        ("σ",  "sigma"), ("ε",  "epsilon"),
        ("μ",  "mu"),    ("ρ",  "rho"),
        # Superscript
        ("²",  "^2"),    ("³",  "^3"),
        ("⁴",  "^4"),    ("⁶",  "^6"),
        ("⁹",  "^9"),
        # Operator & tanda khusus
        ("×",  "x"),     ("·",  "."),
        ("√",  "sqrt"),  ("≤",  "<="),
        ("≥",  ">="),    ("≠",  "!="),
        ("→",  "->"),    ("←",  "<-"),
        ("−",  "-"),     ("—",  "--"),
        ("≈",  "~="),
        # Emoji angka (subheader tabel)
        ("1️⃣", "1."), ("2️⃣", "2."), ("3️⃣", "3."),
        ("4️⃣", "4."), ("5️⃣", "5."),
        # Simbol satuan
        ("°",  "deg"),
    ]
    for asli, ganti in tabel:
        teks = teks.replace(asli, ganti)
    # Hapus sisa karakter non-latin1 yang mungkin lolos
    teks = teks.encode("latin-1", errors="replace").decode("latin-1")
    return teks
 
 
# ── GENERATOR WORD (.docx) ─────────────────────────────────
def generate_word(
    daftar_tabel: list,          # list of (judul_str, dataframe)
    info_input: dict,            # dict parameter input untuk header laporan
) -> BytesIO:
    """
    Buat laporan Word lengkap dengan loop otomatis atas semua tabel.
    Mengembalikan BytesIO siap pakai untuk st.download_button.
    """
    doc = Document()
 
    # ── Margin halaman ──────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
 
    # ── Header Laporan ──────────────────────────────────────
    judul = doc.add_heading("LAPORAN PERHITUNGAN", level=0)
    judul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    judul.runs[0].font.size = Pt(16)
    judul.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
 
    sub = doc.add_heading("Kapasitas Penampang Baja Profil WF", level=1)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
 
    ref = doc.add_paragraph("Referensi: SNI 1729:2020 & AISC 360-16  |  Metode: LRFD & ASD")
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref.runs[0].font.size = Pt(9)
    ref.runs[0].font.italic = True
 
    tgl = doc.add_paragraph(f"Tanggal: {datetime.date.today().strftime('%d %B %Y')}")
    tgl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tgl.runs[0].font.size = Pt(9)
 
    doc.add_paragraph()
 
    # ── Tabel Ringkasan Input ───────────────────────────────
    doc.add_heading("Data Input", level=2)
    tbl_input = doc.add_table(rows=1, cols=2)
    tbl_input.style = "Table Grid"
    tbl_input.alignment = WD_TABLE_ALIGNMENT.LEFT
 
    hdr = tbl_input.rows[0].cells
    hdr[0].text = "Parameter"
    hdr[1].text = "Nilai"
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
    for k, v in info_input.items():
        row = tbl_input.add_row().cells
        row[0].text = k
        row[1].text = v
 
    doc.add_paragraph()
 
    # ── Loop Otomatis Semua Tabel Hasil ─────────────────────
    for judul_tbl, df in daftar_tabel:
        # Heading tabel
        h = doc.add_heading(judul_tbl, level=2)
        h.runs[0].font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
 
        cols = list(df.columns)
        n_cols = len(cols)
        n_rows = len(df)
 
        # Buat tabel Word: 1 baris header + n_rows baris data
        tbl = doc.add_table(rows=1 + n_rows, cols=n_cols)
        tbl.style = "Table Grid"
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
 
        # Set lebar kolom (total ~16 cm untuk margin 2.5 cm kiri-kanan)
        lebar = [Cm(3.5), Cm(4.5), Cm(1.8), Cm(6.2)]  # Parameter|Nilai|Satuan|Keterangan
        for i, cell in enumerate(tbl.rows[0].cells):
            cell.width = lebar[i] if i < len(lebar) else Cm(3)
 
        # Baris header kolom
        for ci, nama_col in enumerate(cols):
            cell = tbl.rows[0].cells[ci]
            cell.text = nama_col
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            # Background biru tua via XML shading
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1A3A5C")
            tcPr.append(shd)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
 
        # Baris data
        for ri, (_, row_data) in enumerate(df.iterrows()):
            word_row = tbl.rows[ri + 1]
            for ci, val in enumerate(row_data):
                cell = word_row.cells[ci]
                cell.text = str(val)
                run = cell.paragraphs[0].runs[0]
                run.font.size = Pt(8.5)
                # Baris bergantian warna (zebra)
                if ri % 2 == 1:
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "EEF3F8")
                    tcPr.append(shd)
                cell.paragraphs[0].alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if ci in (1, 2)
                    else WD_ALIGN_PARAGRAPH.LEFT
                )
 
        doc.add_paragraph()   # spasi antar tabel
 
    # ── Footer Catatan ──────────────────────────────────────
    doc.add_heading("Catatan & Asumsi", level=2)
    catatan_list = [
        "Properties Penampang: Luas fillet diperhitungkan dengan pendekatan 0.8584.r^2 per sudut (4 sudut total).",
        "Klasifikasi: Menggunakan Tabel B4.1b AISC 360-16 (elemen dalam lentur).",
        "Kapasitas Momen (LTB): Menggunakan Pasal F2 AISC 360-16 untuk profil WF kompak simetris ganda.",
        "Kapasitas Geser: Menggunakan Pasal G2.1 AISC 360-16. kv = 5.34 (badan tanpa pengaku, a/h > 3.0).",
        "Warping constant (Cw): Cw = Iy.ho^2/4 (pendekatan untuk WF simetris ganda).",
        "Torsional constant (J): Tanpa koreksi fillet (konservatif).",
        "Semua input dalam N dan mm; output dalam N.mm dan kN.m / kN.",
        "Dokumen ini dibuat otomatis. Selalu verifikasi dengan software resmi untuk keperluan desain.",
    ]
    for c in catatan_list:
        p = doc.add_paragraph(c, style="List Bullet")
        p.runs[0].font.size = Pt(8.5)
 
    # ── Simpan ke BytesIO ───────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
 
 
# ── GENERATOR PDF (fpdf2, Landscape) ──────────────────────
class PDFLaporan(FPDF):
    """Subclass FPDF dengan header & footer kustom."""
 
    def __init__(self, info_input: dict):
        super().__init__(orientation="L", unit="mm", format="A4")
        self.info_input = info_input
        self.set_auto_page_break(auto=True, margin=15)
        self.set_margins(left=15, top=15, right=15)
 
    def header(self):
        self.set_font("Helvetica", "B", 11)
        self.set_text_color(26, 58, 92)   # biru tua
        self.cell(0, 7, "LAPORAN PERHITUNGAN - KAPASITAS PENAMPANG BAJA PROFIL WF",
                  align="C", new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(80, 80, 80)
        self.cell(0, 5, "Referensi: SNI 1729:2020 & AISC 360-16  |  Metode: LRFD & ASD",
                  align="C", new_x="LMARGIN", new_y="NEXT")
        self.ln(2)
        # Garis pembatas header
        self.set_draw_color(26, 58, 92)
        self.set_line_width(0.5)
        self.line(self.l_margin, self.get_y(),
                  self.w - self.r_margin, self.get_y())
        self.ln(3)
        self.set_text_color(0, 0, 0)
 
    def footer(self):
        self.set_y(-12)
        self.set_font("Helvetica", "I", 7)
        self.set_text_color(120, 120, 120)
        tgl = datetime.date.today().strftime("%d %B %Y")
        self.cell(0, 5,
                  f"Halaman {self.page_no()} | Dicetak: {tgl} | "
                  "Verifikasi dengan software resmi sebelum digunakan untuk desain.",
                  align="C")
 
 
def generate_pdf(
    daftar_tabel: list,
    info_input: dict,
) -> BytesIO:
    """
    Buat laporan PDF landscape lengkap dengan loop otomatis atas semua tabel.
    Mengembalikan BytesIO siap pakai untuk st.download_button.
    """
    pdf = PDFLaporan(info_input)
    pdf.add_page()
 
    # ── Lebar total area cetak (A4 landscape: 297mm - 30mm margin) ──
    W = pdf.w - pdf.l_margin - pdf.r_margin   # ~267 mm
 
    # Proporsi lebar 4 kolom: Parameter | Nilai | Satuan | Keterangan
    lebar_kol = [W * 0.14, W * 0.22, W * 0.09, W * 0.55]
 
    # ── Tabel Ringkasan Input ───────────────────────────────
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(26, 58, 92)
    pdf.cell(0, 6, "DATA INPUT", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(1)
 
    pdf.set_font("Helvetica", "B", 8)
    pdf.set_fill_color(26, 58, 92)
    pdf.set_text_color(255, 255, 255)
    w_k = W * 0.40
    w_v = W * 0.60
    pdf.cell(w_k, 6, "Parameter", border=1, fill=True, align="C")
    pdf.cell(w_v, 6, "Nilai", border=1, fill=True, align="C",
             new_x="LMARGIN", new_y="NEXT")
 
    pdf.set_font("Helvetica", "", 8)
    pdf.set_text_color(0, 0, 0)
    for idx, (k, v) in enumerate(info_input.items()):
        fill = (idx % 2 == 1)
        pdf.set_fill_color(238, 243, 248) if fill else pdf.set_fill_color(255, 255, 255)
        pdf.cell(w_k, 5.5, sanitasi_pdf(k), border=1, fill=fill)
        pdf.cell(w_v, 5.5, sanitasi_pdf(v), border=1, fill=fill,
                 new_x="LMARGIN", new_y="NEXT")
 
    pdf.ln(5)
 
    # ── Loop Otomatis Semua Tabel Hasil ─────────────────────
    for judul_tbl, df in daftar_tabel:
        # Cek apakah cukup ruang untuk setidaknya judul + header + 2 baris data
        tinggi_min = 6 + 6 + 2 * 5.5
        if pdf.get_y() + tinggi_min > pdf.h - pdf.b_margin:
            pdf.add_page()
 
        # Judul tabel
        pdf.set_font("Helvetica", "B", 9)
        pdf.set_text_color(26, 58, 92)
        pdf.cell(0, 6, sanitasi_pdf(judul_tbl),
                 new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)
 
        cols = list(df.columns)
 
        # Header kolom
        pdf.set_font("Helvetica", "B", 7.5)
        pdf.set_fill_color(26, 58, 92)
        pdf.set_text_color(255, 255, 255)
        for ci, nama_col in enumerate(cols):
            pdf.cell(lebar_kol[ci], 6, sanitasi_pdf(nama_col),
                     border=1, fill=True, align="C")
        pdf.ln()
 
        # Baris data
        pdf.set_font("Helvetica", "", 7.5)
        pdf.set_text_color(0, 0, 0)
        for ri, (_, row_data) in enumerate(df.iterrows()):
            # Cek page break sebelum setiap baris
            if pdf.get_y() + 5.5 > pdf.h - pdf.b_margin:
                pdf.add_page()
                # Ulang header kolom setelah page break
                pdf.set_font("Helvetica", "I", 7)
                pdf.set_text_color(100, 100, 100)
                pdf.cell(0, 4,
                         f"(lanjutan) {sanitasi_pdf(judul_tbl)}",
                         new_x="LMARGIN", new_y="NEXT")
                pdf.set_font("Helvetica", "B", 7.5)
                pdf.set_fill_color(26, 58, 92)
                pdf.set_text_color(255, 255, 255)
                for ci, nama_col in enumerate(cols):
                    pdf.cell(lebar_kol[ci], 6, sanitasi_pdf(nama_col),
                             border=1, fill=True, align="C")
                pdf.ln()
                pdf.set_font("Helvetica", "", 7.5)
                pdf.set_text_color(0, 0, 0)
 
            fill = (ri % 2 == 1)
            pdf.set_fill_color(238, 243, 248) if fill else pdf.set_fill_color(255, 255, 255)
            vals = list(row_data)
            for ci, val in enumerate(vals):
                txt = sanitasi_pdf(str(val))
                align = "C" if ci in (1, 2) else "L"
                pdf.cell(lebar_kol[ci], 5.5, txt,
                         border=1, fill=fill, align=align)
            pdf.ln()
 
        pdf.ln(4)   # spasi antar tabel
 
    # ── Catatan ─────────────────────────────────────────────
    if pdf.get_y() + 40 > pdf.h - pdf.b_margin:
        pdf.add_page()
 
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_text_color(26, 58, 92)
    pdf.cell(0, 6, "CATATAN & ASUMSI", new_x="LMARGIN", new_y="NEXT")
 
    pdf.set_font("Helvetica", "", 7.5)
    pdf.set_text_color(0, 0, 0)
    catatan_list = [
        "Properties Penampang: Luas fillet diperhitungkan dengan pendekatan 0.8584 x r^2 per sudut (4 sudut total).",
        "Klasifikasi: Menggunakan Tabel B4.1b AISC 360-16 (elemen dalam lentur).",
        "Kapasitas Momen (LTB): Menggunakan Pasal F2 AISC 360-16 untuk profil WF kompak simetris ganda.",
        "Kapasitas Geser: Menggunakan Pasal G2.1 AISC 360-16. kv = 5.34 (badan tanpa pengaku, a/h > 3.0).",
        "Warping constant (Cw): Cw = Iy.ho^2/4. Torsional constant (J): Tanpa koreksi fillet (konservatif).",
        "Semua input dalam N dan mm; output dalam N.mm dan kN.m / kN.",
        "Dokumen ini dibuat otomatis. Selalu verifikasi dengan software resmi untuk desain.",
    ]
    for c in catatan_list:
        pdf.cell(4, 5.5, "-", border=0)
        pdf.multi_cell(W - 4, 5.5, sanitasi_pdf(c), border=0)
 
    # ── Output ke BytesIO ────────────────────────────────────
    buf = BytesIO()
    buf.write(pdf.output())
    buf.seek(0)
    return buf

# ============================================================
# MAIN APP
# ============================================================

st.title("🔩 Kapasitas Penampang Baja Profil WF")
st.caption("Referensi: **SNI 1729:2020** & **AISC 360-16** | Metode: **LRFD** & **ASD**")
st.markdown("---")

# ── SIDEBAR INPUT ──────────────────────────────────────────
st.sidebar.header("📐 Input Data")

st.sidebar.subheader("Material")
Fy = st.sidebar.number_input("Fy – Tegangan Leleh (MPa)", value=250.0, step=5.0, min_value=100.0)
Fu = st.sidebar.number_input("Fu – Tegangan Tarik (MPa)", value=410.0, step=5.0, min_value=100.0)
E  = st.sidebar.number_input("E – Modulus Elastisitas (MPa)", value=200000.0, step=1000.0)

st.sidebar.subheader("Dimensi Profil WF")
d  = st.sidebar.number_input("d  – Tinggi Total (mm)",        value=400.0, step=5.0)
bf = st.sidebar.number_input("bf – Lebar Sayap (mm)",         value=200.0, step=5.0)
tw = st.sidebar.number_input("tw – Tebal Badan (mm)",         value=8.0,   step=0.5)
tf = st.sidebar.number_input("tf – Tebal Sayap (mm)",         value=13.0,  step=0.5)
r  = st.sidebar.number_input("r  – Radius Fillet (mm)",       value=16.0,  step=1.0)

st.sidebar.subheader("Parameter Tekuk Lateral")
Lb = st.sidebar.number_input("Lb – Panjang Tidak Terbreis (mm)", value=3000.0, step=100.0)
Cb = st.sidebar.number_input("Cb – Faktor Modifikasi Momen",     value=1.0,    step=0.05, min_value=1.0)

st.sidebar.markdown("---")
st.sidebar.info("Faktor ketahanan:\n- **LRFD Momen:** φ = 0.90\n- **LRFD Geser:** φ = 1.00\n- **ASD Momen:** Ω = 1.67\n- **ASD Geser:** Ω = 1.50")

# ── PERHITUNGAN ────────────────────────────────────────────
props = hitung_properties(d, bf, tw, tf, r)
klas  = klasifikasi_penampang(bf, tw, tf, props["hw"], Fy, E)
mom   = hitung_momen(d, bf, tw, tf, r, Fy, Fu, E, Lb, Cb, props, klas)
gsr   = hitung_geser(d, tw, tf, props["hw"], Fy, E)

phi_m   = 0.90
Omega_m = 1.67
phi_v   = 1.00
Omega_v = 1.50

# ── RINGKASAN KAPASITAS (ATAS) ─────────────────────────────
col1, col2, col3, col4 = st.columns(4)
col1.metric("φMn (LRFD)", f"{phi_m * mom['Mn']/1e6:.2f} kN·m")
col2.metric("Mn/Ω (ASD)",  f"{mom['Mn']/(Omega_m*1e6):.2f} kN·m")
col3.metric("φVn (LRFD)", f"{phi_v * gsr['Vn']/1e3:.2f} kN")
col4.metric("Vn/Ω (ASD)",  f"{gsr['Vn']/(Omega_v*1e3):.2f} kN")

st.markdown("---")

# ── TABEL 1: PROPERTIES PENAMPANG ─────────────────────────
st.subheader("1️⃣ Properties Penampang")

rows_prop = [
    ("hw", f"{props['hw']:.1f}", "mm", "hw = d − 2·tf  (tinggi badan bersih)"),
    ("A",  f"{props['A']:.2f}",  "mm²","A = 2·bf·tf + hw·tw + koreksi fillet"),
    ("Ix", f"{props['Ix']/1e6:.4f} ×10⁶", "mm⁴", "Ix terhadap sumbu kuat"),
    ("Iy", f"{props['Iy']/1e6:.4f} ×10⁶", "mm⁴", "Iy terhadap sumbu lemah"),
    ("Sx", f"{props['Sx']/1e3:.3f} ×10³",  "mm³", "Sx = Ix / (d/2)"),
    ("Sy", f"{props['Sy']/1e3:.3f} ×10³",  "mm³", "Sy = Iy / (bf/2)"),
    ("Zx", f"{props['Zx']/1e3:.3f} ×10³",  "mm³", "Zx = modulus plastis sumbu kuat"),
    ("rx", f"{props['rx']:.2f}",  "mm", "rx = √(Ix/A)"),
    ("ry", f"{props['ry']:.2f}",  "mm", "ry = √(Iy/A)"),
    ("J",  f"{props['J']/1e3:.3f} ×10³",   "mm⁴", "J ≈ (1/3)[2·bf·tf³ + hw·tw³]  (torsi St. Venant)"),
    ("Cw", f"{props['Cw']/1e9:.4f} ×10⁹",  "mm⁶", "Cw = Iy·ho²/4  (konstanta warping)"),
    ("ho", f"{props['ho']:.1f}",  "mm", "ho = d − tf  (jarak antar titik berat sayap)"),
]
df1 = pd.DataFrame(rows_prop, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])
st.dataframe(df1, use_container_width=True, hide_index=True)

# ── TABEL 2: KLASIFIKASI PENAMPANG ─────────────────────────
st.subheader("2️⃣ Klasifikasi Penampang")

rows_klas = [
    ("λf  (sayap)",    f"{klas['lam_f']:.3f}",  "–", "λf = bf / (2·tf)"),
    ("λpf (kompak)",   f"{klas['lpf']:.3f}",    "–", "λpf = 0.38·√(E/Fy)"),
    ("λrf (non-komp)", f"{klas['lrf']:.3f}",    "–", "λrf = 1.00·√(E/Fy)"),
    ("Klasifikasi Sayap", klas['klas_f'],        "–", f"λf={'<' if klas['lam_f']<=klas['lpf'] else '≤' if klas['lam_f']<=klas['lrf'] else '>'} batas {'λpf' if klas['lam_f']<=klas['lpf'] else 'λrf' if klas['lam_f']<=klas['lrf'] else 'λrf'}"),
    ("λw  (badan)",    f"{klas['lam_w']:.3f}",  "–", "λw = hw / tw"),
    ("λpw (kompak)",   f"{klas['lpw']:.3f}",    "–", "λpw = 3.76·√(E/Fy)"),
    ("λrw (non-komp)", f"{klas['lrw']:.3f}",    "–", "λrw = 5.70·√(E/Fy)"),
    ("Klasifikasi Badan", klas['klas_w'],        "–", f"λw={'<' if klas['lam_w']<=klas['lpw'] else '≤' if klas['lam_w']<=klas['lrw'] else '>'} batas {'λpw' if klas['lam_w']<=klas['lpw'] else 'λrw' if klas['lam_w']<=klas['lrw'] else 'λrw'}"),
]
df2 = pd.DataFrame(rows_klas, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])

def warnai_klas(val):
    if val == "Kompak":
        return "background-color: #d4edda; color: #155724"
    elif val == "Non-Kompak":
        return "background-color: #fff3cd; color: #856404"
    elif val == "Langsing":
        return "background-color: #f8d7da; color: #721c24"
    return ""

st.dataframe(
    df2.style.map(warnai_klas, subset=["Nilai"]),
    use_container_width=True, hide_index=True
)

# ── TABEL 3: PARAMETER TEKUK LATERAL ──────────────────────
st.subheader("3️⃣ Parameter Tekuk Lateral (LTB)")

Fcr_str = f"{mom['Fcr']/1e6:.4f} ×10⁶ N·mm" if mom['Fcr'] is not None else "N/A (bukan zona elastis)"

rows_ltb = [
    ("Mp",      f"{mom['Mp']/1e6:.4f} ×10⁶", "N·mm",  "Mp = Fy · Zx"),
    ("Mr",      f"{mom['Mr']/1e6:.4f} ×10⁶", "N·mm",  "Mr = 0.7 · Fy · Sx"),
    ("rts",     f"{mom['rts']:.3f}",           "mm",    "rts = [√(Iy·Cw)/Sx]^0.5"),
    ("ho",      f"{mom['ho']:.1f}",            "mm",    "ho = d − tf"),
    ("c",       f"{mom['c']:.1f}",             "–",     "c = 1.0 (WF simetris ganda)"),
    ("Lp",      f"{mom['Lp']:.1f}",            "mm",    "Lp = 1.76 · ry · √(E/Fy)"),
    ("Lr",      f"{mom['Lr']:.1f}",            "mm",    "Lr = 1.95·rts·(E/0.7Fy)·√(J·c/(Sx·ho)+...)"),
    ("Lb (input)", f"{Lb:.1f}",               "mm",    "Panjang tanpa breis lateral"),
    ("Cb (input)", f"{Cb:.2f}",               "–",     "Faktor modifikasi momen"),
    ("Fcr",     Fcr_str,                       "MPa",   "Fcr = Cb·π²·E/(Lb/rts)² · √(1+0.078·J·c/(Sx·ho)·(Lb/rts)²)"),
    ("Kondisi", mom['kondisi'],                "–",     f"Lb={Lb:.0f} vs Lp={mom['Lp']:.0f} vs Lr={mom['Lr']:.0f} mm"),
]
df3 = pd.DataFrame(rows_ltb, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])

def warnai_kondisi(val):
    if "Plastis" in str(val):
        return "background-color: #d4edda; color: #155724"
    elif "Inelastis" in str(val):
        return "background-color: #fff3cd; color: #856404"
    elif "Elastis" in str(val):
        return "background-color: #f8d7da; color: #721c24"
    return ""

st.dataframe(
    df3.style.map(warnai_kondisi, subset=["Nilai"]),
    use_container_width=True, hide_index=True
)

# ── TABEL 4: KAPASITAS MOMEN ───────────────────────────────
st.subheader("4️⃣ Kapasitas Momen")

phi_Mn   = phi_m * mom["Mn"]
Mn_Omega = mom["Mn"] / Omega_m

rows_mom = [
    ("Mn",          f"{mom['Mn']/1e6:.4f} ×10⁶", "N·mm",  f"Mn = f(kondisi LTB) — {mom['kondisi']}"),
    ("Mn",          f"{mom['Mn']/1e6/1000:.3f}",  "kN·m",  "konversi kN·m"),
    ("φ (LRFD)",    f"{phi_m:.2f}",                "–",     "SNI 1729:2020 Pasal F1 / AISC F1"),
    ("φ·Mn (LRFD)", f"{phi_Mn/1e6:.4f} ×10⁶",     "N·mm",  "Kapasitas desain LRFD"),
    ("φ·Mn (LRFD)", f"{phi_Mn/1e6/1000:.3f}",      "kN·m",  "konversi kN·m ← **GUNAKAN INI**"),
    ("Ω (ASD)",     f"{Omega_m:.2f}",              "–",     "SNI 1729:2020 Pasal F1 / AISC F1"),
    ("Mn/Ω (ASD)",  f"{Mn_Omega/1e6:.4f} ×10⁶",   "N·mm",  "Kapasitas ijin ASD"),
    ("Mn/Ω (ASD)",  f"{Mn_Omega/1e6/1000:.3f}",    "kN·m",  "konversi kN·m ← **GUNAKAN INI**"),
]
df4 = pd.DataFrame(rows_mom, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])
st.dataframe(df4, use_container_width=True, hide_index=True)

# ── TABEL 5: KAPASITAS GESER ───────────────────────────────
st.subheader("5️⃣ Kapasitas Geser")

phi_Vn   = phi_v * gsr["Vn"]
Vn_Omega = gsr["Vn"] / Omega_v

rows_gsr = [
    ("Aw",          f"{gsr['Aw']:.2f}",            "mm²",  "Aw = d · tw  (luas geser badan)"),
    ("kv",          f"{gsr['kv']:.2f}",             "–",    "kv = 5.34  (tanpa pengaku transversal, a/h > 3)"),
    ("hw/tw",       f"{gsr['ratio']:.3f}",          "–",    "Rasio kelangsingan badan"),
    ("2.24√(E/Fy)", f"{gsr['batas_cv1']:.3f}",      "–",    "Batas untuk Cv1 = 1.0 (AISC G2.1(a))"),
    ("Cv1",         f"{gsr['Cv1']:.4f}",            "–",    gsr["catatan_cv1"]),
    ("Vn",          f"{gsr['Vn']/1e3:.4f} ×10³",   "N",    "Vn = 0.6 · Fy · Aw · Cv1"),
    ("Vn",          f"{gsr['Vn']/1e3:.3f}",         "kN",   "konversi kN"),
    ("φ (LRFD)",    f"{phi_v:.2f}",                 "–",    "AISC G2.1 / SNI 1729:2020"),
    ("φ·Vn (LRFD)", f"{phi_Vn/1e3:.3f}",            "kN",   "Kapasitas geser desain LRFD ← **GUNAKAN INI**"),
    ("Ω (ASD)",     f"{Omega_v:.2f}",               "–",    "AISC G2.1 / SNI 1729:2020"),
    ("Vn/Ω (ASD)",  f"{Vn_Omega/1e3:.3f}",          "kN",   "Kapasitas geser ijin ASD ← **GUNAKAN INI**"),
]
df5 = pd.DataFrame(rows_gsr, columns=["Parameter", "Nilai", "Satuan", "Keterangan / Rumus"])
st.dataframe(df5, use_container_width=True, hide_index=True)

# ── CATATAN FOOTER ─────────────────────────────────────────
st.markdown("---")
with st.expander("📋 Catatan & Asumsi Perhitungan"):
    st.markdown("""
    **Referensi Utama:**
    - SNI 1729:2020 – Spesifikasi untuk Bangunan Gedung Baja Struktural
    - AISC 360-16 – Specification for Structural Steel Buildings

    **Asumsi & Batasan:**
    1. **Properties Penampang**: Luas fillet diperhitungkan dengan pendekatan ≈ 0.8584·r² per sudut (4 sudut total).
    2. **Klasifikasi**: Menggunakan Tabel B4.1b AISC 360-16 (elemen dalam lentur).
    3. **Kapasitas Momen (LTB)**: Menggunakan Pasal F2 AISC 360-16 untuk profil WF kompak simetris ganda.
       - Jika penampang **non-kompak** (FLB), diperlukan cek tambahan Pasal F3.
    4. **Kapasitas Geser**: Menggunakan Pasal G2.1 AISC 360-16.
       - kv = 5.34 (badan tanpa pengaku, a/h > 3.0).
    5. **Warping constant (Cw)**: Cw = Iy·ho²/4 (pendekatan untuk WF simetris ganda).
    6. **Torsional constant (J)**: Tanpa koreksi fillet (konservatif).

    **Konversi Satuan:**
    - 1 kN·m = 10⁶ N·mm
    - Semua input dalam **N** dan **mm**, output dalam **N·mm** dan **kN·m/kN**
    """)

st.caption("Dibuat dengan Python + Streamlit | Untuk keperluan engineering, selalu verifikasi dengan software resmi.")
st.markdown("---")
st.subheader("⬇️ Download Laporan")
 
# Susun daftar tabel: (judul_string, dataframe) — loop otomatis
daftar_tabel_laporan = [
    ("1. Properties Penampang",         df1),
    ("2. Klasifikasi Penampang",         df2),
    ("3. Parameter Tekuk Lateral (LTB)", df3),
    ("4. Kapasitas Momen",              df4),
    ("5. Kapasitas Geser",              df5),
]
 
# Info input untuk header laporan
info_input_laporan = {
    "Fy - Tegangan Leleh":          f"{Fy:.0f} MPa",
    "Fu - Tegangan Tarik":          f"{Fu:.0f} MPa",
    "E  - Modulus Elastisitas":     f"{E:.0f} MPa",
    "d  - Tinggi Total":            f"{d:.0f} mm",
    "bf - Lebar Sayap":             f"{bf:.0f} mm",
    "tw - Tebal Badan":             f"{tw:.1f} mm",
    "tf - Tebal Sayap":             f"{tf:.1f} mm",
    "r  - Radius Fillet":           f"{r:.1f} mm",
    "Lb - Panjang Tdk Terbreis":    f"{Lb:.0f} mm",
    "Cb - Faktor Modifikasi Momen": f"{Cb:.2f}",
    "phi.Mn (LRFD)":                f"{phi_m * mom['Mn'] / 1e6:.3f} kN.m",
    "Mn/Omega (ASD)":               f"{mom['Mn'] / (Omega_m * 1e6):.3f} kN.m",
    "phi.Vn (LRFD)":                f"{phi_v * gsr['Vn'] / 1e3:.3f} kN",
    "Vn/Omega (ASD)":               f"{gsr['Vn'] / (Omega_v * 1e3):.3f} kN",
}
 
col_dl1, col_dl2 = st.columns(2)
 
with col_dl1:
    if st.button("📄 Generate & Download Word (.docx)"):
        with st.spinner("Membuat dokumen Word..."):
            word_buf = generate_word(daftar_tabel_laporan, info_input_laporan)
        nama_file_word = f"WF_d{int(d)}xbf{int(bf)}_SNI1729.docx"
        st.download_button(
            label="💾 Klik untuk Simpan Word",
            data=word_buf,
            file_name=nama_file_word,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
 
with col_dl2:
    if st.button("📑 Generate & Download PDF (.pdf)"):
        with st.spinner("Membuat dokumen PDF..."):
            pdf_buf = generate_pdf(daftar_tabel_laporan, info_input_laporan)
        nama_file_pdf = f"WF_d{int(d)}xbf{int(bf)}_SNI1729.pdf"
        st.download_button(
            label="💾 Klik untuk Simpan PDF",
            data=pdf_buf,
            file_name=nama_file_pdf,
            mime="application/pdf",
        )
 

