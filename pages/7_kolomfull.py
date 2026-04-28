# calculations.py
"""
Modul perhitungan struktur kolom beton bertulang berdasarkan SNI 2847:2019.
Hanya berisi fungsi logika murni, tanpa antarmuka Streamlit.
"""

import math
import numpy as np
from typing import Dict, List, Tuple, Any, Optional

# ----------------------------------------------------------------------
# CONSTANTS AND CONVERSIONS
# ----------------------------------------------------------------------
MM_TO_M = 1e-3
MPA_TO_KPA = 1000.0  # 1 MPa = 1000 kPa, but we stay in N & mm
N_TO_KN = 1e-3
NMM_TO_KNM = 1e-6

# ----------------------------------------------------------------------
# HELPER: VALIDATION
# ----------------------------------------------------------------------
def validate_inputs(data: Dict[str, Any]) -> List[str]:
    """Mengembalikan daftar pesan error jika input tidak valid."""
    errors = []
    required = [
        'fc', 'fy', 'Es', 'b', 'h', 'cover', 'ds', 'D', 'n_b', 'n_h',
        'Lu', 'frame_type', 'Pu', 'M1', 'M2', 'curvature'
    ]
    for key in required:
        if key not in data:
            errors.append(f"Input '{key}' tidak ditemukan.")
    if errors:
        return errors

    # Nilai positif
    if data['fc'] <= 0:
        errors.append("fc' harus > 0 MPa")
    if data['fy'] <= 0:
        errors.append("fy harus > 0 MPa")
    if data['Es'] <= 0:
        errors.append("Es harus > 0 MPa")
    if data['b'] <= 0:
        errors.append("Lebar b harus > 0 mm")
    if data['h'] <= 0:
        errors.append("Tinggi h harus > 0 mm")
    if data['cover'] < 0:
        errors.append("Selimut beton tidak boleh negatif")
    if data['ds'] < 0:
        errors.append("Diameter sengkang tidak boleh negatif")
    if data['D'] <= 0:
        errors.append("Diameter tulangan D harus > 0 mm")
    if data['n_b'] < 2:
        errors.append("Jumlah tulangan sisi b minimal 2")
    if data['n_h'] < 2:
        errors.append("Jumlah tulangan sisi h minimal 2")
    if data['Lu'] <= 0:
        errors.append("Panjang tak tertahan Lu harus > 0 mm")
    if data['frame_type'] not in ('braced', 'unbraced'):
        errors.append("Kondisi rangka harus 'braced' atau 'unbraced'")
    if data['curvature'] not in ('single', 'double'):
        errors.append("Kelengkungan harus 'single' atau 'double'")
    if data.get('beta_dns', 0.6) < 0:
        errors.append("beta_dns tidak boleh negatif")
    if data.get('L_col_lower', 0) < 0:
        errors.append("Panjang kolom bawah tidak boleh negatif")
    if data.get('L_col_upper', 0) < 0:
        errors.append("Panjang kolom atas tidak boleh negatif")
    if data.get('L_beam', 0) <= 0:
        errors.append("Panjang balok L_beam harus > 0 mm")

    # Cek tulangan total
    n_b = data['n_b']
    n_h = data['n_h']
    n_total = 2 * n_b + 2 * (n_h - 2)
    if n_total < 4:
        errors.append("Jumlah tulangan total minimal 4 buah.")

    return errors


# ----------------------------------------------------------------------
# PROPERTIES PENAMPANG & TULANGAN
# ----------------------------------------------------------------------
def calc_properties(data: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Menghitung properti penampang kolom.
    Mengembalikan list of dict siap DataFrame [Item/Uraian, Nilai, Satuan, Keterangan].
    """
    steps = []

    b = data['b']               # mm
    h = data['h']               # mm
    fc = data['fc']             # MPa
    fy = data['fy']             # MPa
    Es = data['Es']             # MPa
    cover = data['cover']       # mm
    ds = data['ds']             # mm
    D = data['D']               # mm
    n_b = data['n_b']
    n_h = data['n_h']

    # Luas bruto dan inersia
    Ag = b * h
    Ig = (1.0 / 12.0) * b * h**3
    Ec = 4700.0 * math.sqrt(fc)

    steps.append({
        'Item/Uraian': 'Ag',
        'Nilai': f"{Ag:.1f}",
        'Satuan': 'mm²',
        'Keterangan': f"Ag = b × h = {b} × {h} = {Ag:.1f} mm²"
    })
    steps.append({
        'Item/Uraian': 'Ig',
        'Nilai': f"{Ig:.2f}",
        'Satuan': 'mm⁴',
        'Keterangan': f"Ig = 1/12 × b × h³ = 1/12 × {b} × {h}³ = {Ig:.2f} mm⁴"
    })
    steps.append({
        'Item/Uraian': 'Ec',
        'Nilai': f"{Ec:.2f}",
        'Satuan': 'MPa',
        'Keterangan': f"Ec = 4700 × √fc' = 4700 × √{fc} = {Ec:.2f} MPa  [SNI 2847:2019 Pasal 19.2.2.1]"
    })

    # Selimut efektif dan posisi tulangan
    d = h - cover - ds - D / 2.0          # jarak dari serat tekan terluar ke pusat tulangan tarik
    d_prime = cover + ds + D / 2.0        # jarak dari serat tekan terluar ke pusat tulangan tekan

    steps.append({
        'Item/Uraian': 'd',
        'Nilai': f"{d:.2f}",
        'Satuan': 'mm',
        'Keterangan': f"d = h - c - Øs - D/2 = {h} - {cover} - {ds} - {D}/2 = {d:.2f} mm"
    })
    steps.append({
        'Item/Uraian': "d'",
        'Nilai': f"{d_prime:.2f}",
        'Satuan': 'mm',
        'Keterangan': f"d' = c + Øs + D/2 = {cover} + {ds} + {D}/2 = {d_prime:.2f} mm"
    })

    # Luas tulangan total
    A_bar = math.pi * D**2 / 4.0
    n_total = 2 * n_b + 2 * (n_h - 2)
    Ast = n_total * A_bar
    rho_g = (Ast / Ag) * 100.0           # dalam persen

    steps.append({
        'Item/Uraian': 'A_bar',
        'Nilai': f"{A_bar:.2f}",
        'Satuan': 'mm²',
        'Keterangan': f"A_bar = π × D² / 4 = π × {D}² / 4 = {A_bar:.2f} mm²"
    })
    steps.append({
        'Item/Uraian': 'n_total',
        'Nilai': f"{n_total}",
        'Satuan': 'bh',
        'Keterangan': f"n_total = 2×n_b + 2×(n_h - 2) = 2×{n_b} + 2×({n_h}-2) = {n_total} bh"
    })
    steps.append({
        'Item/Uraian': 'Ast',
        'Nilai': f"{Ast:.2f}",
        'Satuan': 'mm²',
        'Keterangan': f"Ast = n_total × A_bar = {n_total} × {A_bar:.2f} = {Ast:.2f} mm²"
    })
    steps.append({
        'Item/Uraian': 'ρg',
        'Nilai': f"{rho_g:.4f}",
        'Satuan': '%',
        'Keterangan': f"ρg = Ast / Ag × 100 = {Ast:.2f} / {Ag:.1f} × 100 = {rho_g:.4f}%"
    })

    # Cek rasio tulangan [SNI 2847:2019 Pasal 10.6.1]
    rho_min = 1.0
    rho_max = 8.0
    status = "OK" if rho_min <= rho_g <= rho_max else "TIDAK OK"
    steps.append({
        'Item/Uraian': 'Cek ρg',
        'Nilai': f"{rho_g:.4f}%",
        'Satuan': '%',
        'Keterangan': f"ρmin = {rho_min}%, ρmax = {rho_max}% -> {status}  [SNI 2847:2019 Pasal 10.6.1]"
    })

    return steps


# ----------------------------------------------------------------------
# KELANGSINGAN & KEKAKUAN
# ----------------------------------------------------------------------
def calc_slenderness(data: Dict[str, Any], prop: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Menghitung parameter kelangsingan, faktor panjang efektif k,
    dan memeriksa apakah kolom langsing atau pendek.
    """
    steps = []

    b = data['b']
    h = data['h']
    fc = data['fc']
    Lu = data['Lu']
    frame_type = data['frame_type']
    curvature = data['curvature']
    beta_dns = data.get('beta_dns', 0.6)

    # Dimensi balok dan kolom untuk perhitungan Ψ
    b_beam = data.get('b_beam', 300)
    h_beam = data.get('h_beam', 500)
    L_beam = data.get('L_beam', 6000)
    L_col_upper = data.get('L_col_upper', 3500)
    L_col_lower = data.get('L_col_lower', 3500)
    # Anggap dimensi kolom atas & bawah sama dengan kolom yang ditinjau (400x500)
    b_col = b
    h_col = h

    Ec = 4700.0 * math.sqrt(fc)
    # Inersia bruto
    Ig_col = (1.0/12.0) * b_col * h_col**3
    Ig_beam = (1.0/12.0) * b_beam * h_beam**3

    # Kekakuan lentur efektif EI = 0.4 Ec Ig / (1+βdns) [SNI 2847:2019 Pasal 6.6.3.1.1]
    EI_col = 0.4 * Ec * Ig_col / (1.0 + beta_dns)
    EI_beam = 0.4 * Ec * Ig_beam / (1.0 + beta_dns)

    # Faktor Ψ di ujung atas dan bawah untuk kolom yang ditinjau
    # Ψ = Σ(EI/L)_kolom / Σ(EI/L)_balok
    if L_col_upper == 0:
        Psi_A = 0.0  # fixed di atas
    else:
        EI_L_upper = EI_col / L_col_upper  # kolom atas
        Psi_A = (EI_col / Lu + EI_col / L_col_upper) / (2 * EI_beam / L_beam)

    if L_col_lower == 0:
        Psi_B = 0.0
    else:
        EI_L_lower = EI_col / L_col_lower
        Psi_B = (EI_col / Lu + EI_col / L_col_lower) / (2 * EI_beam / L_beam)

    steps.append({
        'Item/Uraian': 'Ig kolom',
        'Nilai': f"{Ig_col:.2f}",
        'Satuan': 'mm⁴',
        'Keterangan': f"Ig = 1/12 × {b_col} × {h_col}³ = {Ig_col:.2f} mm⁴"
    })
    steps.append({
        'Item/Uraian': 'Ig balok',
        'Nilai': f"{Ig_beam:.2f}",
        'Satuan': 'mm⁴',
        'Keterangan': f"Ig_balok = 1/12 × {b_beam} × {h_beam}³ = {Ig_beam:.2f} mm⁴"
    })
    steps.append({
        'Item/Uraian': 'EI kolom',
        'Nilai': f"{EI_col:.2f}",
        'Satuan': 'N·mm²',
        'Keterangan': f"EI = 0.4 × Ec × Ig / (1+βdns) = 0.4 × {Ec:.2f} × {Ig_col:.2f} / (1+{beta_dns}) = {EI_col:.2f} N·mm²"
    })
    steps.append({
        'Item/Uraian': 'EI balok',
        'Nilai': f"{EI_beam:.2f}",
        'Satuan': 'N·mm²',
        'Keterangan': f"EI_balok = 0.4 × Ec × Ig_balok / (1+βdns) = {EI_beam:.2f} N·mm²"
    })
    steps.append({
        'Item/Uraian': 'Ψ_A (atas)',
        'Nilai': f"{Psi_A:.4f}",
        'Satuan': '-',
        'Keterangan': f"Ψ_A = (Σ EI_col/L) / (Σ EI_beam/L) = ({EI_col:.2f}/{Lu}+{EI_col:.2f}/{L_col_upper}) / (2×{EI_beam:.2f}/{L_beam}) = {Psi_A:.4f}"
    })
    steps.append({
        'Item/Uraian': 'Ψ_B (bawah)',
        'Nilai': f"{Psi_B:.4f}",
        'Satuan': '-',
        'Keterangan': f"Ψ_B = (Σ EI_col/L) / (Σ EI_beam/L) = ({EI_col:.2f}/{Lu}+{EI_col:.2f}/{L_col_lower}) / (2×{EI_beam:.2f}/{L_beam}) = {Psi_B:.4f}"
    })

    # Faktor panjang efektif k
    if frame_type == 'braced':
        # Non-sway frame: k = min(0.7+0.05(ψA+ψB), 0.85+0.05*ψ_min) ≤ 1.0
        k1 = 0.7 + 0.05 * (Psi_A + Psi_B)
        psi_min = min(Psi_A, Psi_B)
        k2 = 0.85 + 0.05 * psi_min
        k = min(k1, k2)
        if k > 1.0:
            k = 1.0
        ref = "[SNI 2847:2019 Pasal 6.6.4.3.1]"
    else:  # unbraced (sway)
        # Sway frame: ψm = rata-rata ψA dan ψB
        psi_m = (Psi_A + Psi_B) / 2.0
        if psi_m < 2:
            k = (20.0 - psi_m) / 20.0 * math.sqrt(1.0 + psi_m)
        else:
            k = 0.9 * math.sqrt(1.0 + psi_m)
        ref = "[SNI 2847:2019 Pasal 6.6.4.3.2]"

    steps.append({
        'Item/Uraian': 'Faktor k',
        'Nilai': f"{k:.4f}",
        'Satuan': '-',
        'Keterangan': f"k (frame {frame_type}) = {k:.4f}  {ref}"
    })

    # Radius of girasi dan rasio kelangsingan
    r = 0.3 * h   # sesuai SNI untuk kolom persegi
    kLu_r = (k * Lu) / r

    steps.append({
        'Item/Uraian': 'r (radius girasi)',
        'Nilai': f"{r:.2f}",
        'Satuan': 'mm',
        'Keterangan': f"r = 0.3 × h = 0.3 × {h} = {r:.2f} mm"
    })
    steps.append({
        'Item/Uraian': 'k·Lu / r',
        'Nilai': f"{kLu_r:.4f}",
        'Satuan': '-',
        'Keterangan': f"kLu/r = {k:.4f} × {Lu} / {r:.2f} = {kLu_r:.4f}"
    })

    # Batas kelangsingan
    if frame_type == 'braced':
        # M1/M2 dengan tanda
        M1 = data['M1']   # kNm
        M2 = data['M2']
        if curvature == 'single':
            M1_M2 = M1 / M2
        else:
            M1_M2 = - M1 / M2   # negative for double curvature
        limit = 34.0 - 12.0 * M1_M2
        if limit > 40.0:
            limit = 40.0
        slenderness_check = "Langsing" if kLu_r > limit else "Pendek"
        steps.append({
            'Item/Uraian': 'Batas kLu/r (non-sway)',
            'Nilai': f"{limit:.4f}",
            'Satuan': '-',
            'Keterangan': f"34 - 12×(M1/M2) = 34 - 12×({M1_M2:.4f}) = {limit:.4f} (max 40)  [SNI 2847:2019 Pasal 6.6.4.3.1]"
        })
    else:
        limit = 22.0
        slenderness_check = "Langsing" if kLu_r > limit else "Pendek"
        steps.append({
            'Item/Uraian': 'Batas kLu/r (sway)',
            'Nilai': f"{limit:.4f}",
            'Satuan': '-',
            'Keterangan': f"22  [SNI 2847:2019 Pasal 6.6.4.3.2]"
        })

    steps.append({
        'Item/Uraian': 'Status Kolom',
        'Nilai': slenderness_check,
        'Satuan': '-',
        'Keterangan': f"kLu/r = {kLu_r:.4f} vs batas {limit:.4f} → {slenderness_check}"
    })

    return steps


# ----------------------------------------------------------------------
# MOMEN PEMBESARAN (MAGNIFICATION) untuk non-sway frame
# ----------------------------------------------------------------------
def calc_moment_magnification(data: Dict[str, Any], prop: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Menghitung faktor pembesaran momen δns untuk kolom braced.
    Untuk unbraced hanya memberi catatan.
    """
    steps = []
    frame_type = data['frame_type']

    if frame_type != 'braced':
        steps.append({
            'Item/Uraian': 'Pembesaran Momen',
            'Nilai': 'N/A',
            'Satuan': '-',
            'Keterangan': 'Pembesaran momen untuk sway frame memerlukan analisis orde dua terpisah.'
        })
        return steps

    # Data yang diperlukan
    fc = data['fc']
    b = data['b']
    h = data['h']
    fy = data['fy']
    Es = data['Es']
    cover = data['cover']
    ds = data['ds']
    D = data['D']
    Lu = data['Lu']
    k = data.get('k', 1.0)   # sebelumnya harus dihitung, kita akan ambil dari state
    M1 = data['M1'] * 1e6       # kNm -> Nmm
    M2 = data['M2'] * 1e6
    Pu = data['Pu'] * 1e3       # kN -> N
    curvature = data['curvature']
    beta_dns = data.get('beta_dns', 0.6)

    Ec = 4700.0 * math.sqrt(fc)
    # Properti yang sudah dihitung: d, d', Ast (kita hitung ulang untuk kemandirian)
    d = h - cover - ds - D/2.0
    d_prime = cover + ds + D/2.0
    A_bar = math.pi * D**2 / 4.0
    n_b = data['n_b']
    n_h = data['n_h']
    n_total = 2 * n_b + 2 * (n_h - 2)
    Ast = n_total * A_bar
    # layer 1 dan 3 masing-masing n_b bar, layer 2 = 2 bar
    As1 = n_b * A_bar
    As3 = As1  # simetri
    # momen inersia tulangan terhadap pusat penampang
    y_centroid = h / 2.0
    d1 = d_prime
    d3 = h - d_prime
    Ise = As1 * (d1 - y_centroid)**2 + As3 * (d3 - y_centroid)**2
    # seharusnya simetri sehingga 2*As1*(h/2 - d')^2
    Ig = (1.0/12.0) * b * h**3

    # Kekakuan lentur untuk Pc [SNI 2847:2019 Pasal 6.6.4.5.1]
    # EI = (0.2 Ec Ig + Es Ise) / (1 + βdns)
    EI = (0.2 * Ec * Ig + Es * Ise) / (1.0 + beta_dns)

    # Pc = π² EI / (k Lu)²
    Pc = (math.pi**2 * EI) / (k * Lu)**2

    # Cm = 0.6 + 0.4 (M1/M2) dengan tanda sesuai kelengkungan
    if curvature == 'single':
        ratio = M1 / M2
    else:
        ratio = - M1 / M2
    Cm = 0.6 + 0.4 * ratio
    if Cm < 0.4:
        Cm = 0.4

    # δns = Cm / (1 - Pu/(0.75 Pc)) ≥ 1.0
    delta_ns = Cm / (1.0 - Pu / (0.75 * Pc))
    if delta_ns < 1.0:
        delta_ns = 1.0

    # Momen rencana terbesar diperbesar
    Mc = delta_ns * M2   # Nmm

    steps.append({
        'Item/Uraian': 'EI (untuk Pc)',
        'Nilai': f"{EI:.2f}",
        'Satuan': 'N·mm²',
        'Keterangan': f"EI = (0.2 Ec Ig + Es Ise)/(1+βdns) = (0.2×{Ec:.2f}×{Ig:.2f} + {Es}×{Ise:.2f})/(1+{beta_dns}) = {EI:.2f}"
    })
    steps.append({
        'Item/Uraian': 'Pc',
        'Nilai': f"{Pc*N_TO_KN:.2f}",
        'Satuan': 'kN',
        'Keterangan': f"Pc = π² EI / (k Lu)² = π²×{EI:.2f} / ({k:.4f}×{Lu})² = {Pc*N_TO_KN:.2f} kN"
    })
    steps.append({
        'Item/Uraian': 'Cm',
        'Nilai': f"{Cm:.4f}",
        'Satuan': '-',
        'Keterangan': f"Cm = 0.6 + 0.4(M1/M2) = 0.6 + 0.4×({ratio:.4f}) = {Cm:.4f}  [SNI 2847:2019 Pasal 6.6.4.5.1]"
    })
    steps.append({
        'Item/Uraian': 'δns',
        'Nilai': f"{delta_ns:.4f}",
        'Satuan': '-',
        'Keterangan': f"δns = Cm / (1 - Pu/(0.75 Pc)) = {Cm:.4f} / (1 - {Pu*N_TO_KN:.2f}/(0.75×{Pc*N_TO_KN:.2f})) = {delta_ns:.4f}"
    })
    steps.append({
        'Item/Uraian': 'M2 (asli)',
        'Nilai': f"{M2*NMM_TO_KNM:.4f}",
        'Satuan': 'kN·m',
        'Keterangan': f"M2 = {M2*NMM_TO_KNM:.4f} kN·m"
    })
    steps.append({
        'Item/Uraian': 'Mc = δns × M2',
        'Nilai': f"{Mc*NMM_TO_KNM:.4f}",
        'Satuan': 'kN·m',
        'Keterangan': f"Mc = {delta_ns:.4f} × {M2*NMM_TO_KNM:.4f} = {Mc*NMM_TO_KNM:.4f} kN·m"
    })

    # Simpan ke data untuk pemeriksaan kapasitas
    data['Mc'] = Mc   # Nmm
    data['EI_for_Pc'] = EI
    data['Pc'] = Pc
    data['delta_ns'] = delta_ns

    return steps


# ----------------------------------------------------------------------
# DIAGRAM INTERAKSI P-M
# ----------------------------------------------------------------------
def calc_interaction_diagram(data: Dict[str, Any]) -> List[Dict[str, str]]:
    """
    Membangkitkan 52 titik diagram interaksi P-M sesuai SNI 2847:2019.
    """
    steps = []

    fc = data['fc']              # MPa
    fy = data['fy']              # MPa
    Es = data['Es']              # MPa
    b = data['b']                # mm
    h = data['h']                # mm
    cover = data['cover']        # mm
    ds = data['ds']              # mm
    D = data['D']                # mm
    n_b = data['n_b']
    n_h = data['n_h']

    # Properti
    Ag = b * h
    A_bar = math.pi * D**2 / 4.0
    n_total = 2 * n_b + 2 * (n_h - 2)
    Ast = n_total * A_bar

    # Lapisan tulangan (3 layer)
    d_prime = cover + ds + D/2.0
    # layer 1: tekan (atas)
    y1 = d_prime
    As1 = n_b * A_bar
    # layer 2: tengah
    y2 = h / 2.0
    As2 = 2.0 * A_bar    # dua bar di sisi h pada setengah tinggi
    # layer 3: tarik (bawah)
    y3 = h - d_prime
    As3 = n_b * A_bar
    layers = [
        {'y': y1, 'As': As1},
        {'y': y2, 'As': As2},
        {'y': y3, 'As': As3},
    ]

    # Regangan dan tegangan leleh
    ecu = 0.003
    ey = fy / Es

    # Faktor β1 [SNI 2847:2019 Pasal 22.2.2.4.3]
    if fc <= 28.0:
        beta1 = 0.85
    else:
        beta1 = 0.85 - 0.05 * (fc - 28.0) / 7.0
        if beta1 < 0.65:
            beta1 = 0.65

    # Titik diagram
    c_h_values = np.linspace(0.02, 1.2, 50)
    points = []  # list of dicts

    for c_h in c_h_values:
        c = c_h * h
        a = beta1 * c
        if a > h:
            a = h

        # Regangan tarik terjauh (lapisan bawah)
        dt = y3
        et = ecu * (dt - c) / c   # positif = tarik

        # Regangan dan gaya tiap lapis
        Fs_total = 0.0
        Ms_total = 0.0
        for layer in layers:
            yi = layer['y']
            Asi = layer['As']
            esi = ecu * (c - yi) / c  # positif = tekan
            fs = esi * Es
            if fs > fy:
                fs = fy
            elif fs < -fy:
                fs = -fy
            Fi = fs * Asi
            Fs_total += Fi
            Ms_total += Fi * (h/2.0 - yi)

        # Gaya beton
        Cc = 0.85 * fc * a * b
        # Momen terhadap pusat
        Mc = Cc * (h/2.0 - a/2.0)

        Pn = Cc + Fs_total  # N
        Mn = Mc + Ms_total  # Nmm

        # Faktor φ berdasarkan εt
        if et <= ey:
            phi = 0.65
        elif et >= 0.005:
            phi = 0.9
        else:
            phi = 0.65 + (et - ey) * (0.25 / (0.005 - ey))

        phiPn = phi * Pn
        phiMn = phi * Mn

        points.append({
            'Titik': len(points)+1,
            'c/h': f"{c_h:.4f}",
            'c (mm)': f"{c:.2f}",
            'Pn (kN)': f"{Pn*N_TO_KN:.4f}",
            'Mn (kN·m)': f"{Mn*NMM_TO_KNM:.4f}",
            'φ': f"{phi:.3f}",
            'φPn (kN)': f"{phiPn*N_TO_KN:.4f}",
            'φMn (kN·m)': f"{phiMn*NMM_TO_KNM:.4f}",
            'εt': f"{et:.6f}",
        })

    # Titik 51: Tekan Murni (Po dengan faktor reduksi)
    Po = 0.85 * fc * (Ag - Ast) + fy * Ast
    Pn_max = 0.80 * Po   # SNI 2847:2019 Pasal 22.4.2.1
    points.append({
        'Titik': 51,
        'c/h': '∞',
        'c (mm)': '∞',
        'Pn (kN)': f"{Pn_max*N_TO_KN:.4f}",
        'Mn (kN·m)': '0.0000',
        'φ': '0.650',
        'φPn (kN)': f"{(0.65*Pn_max)*N_TO_KN:.4f}",
        'φMn (kN·m)': '0.0000',
        'εt': 'N/A',
    })

    # Titik 52: Tarik Murni
    Tn = - fy * Ast
    points.append({
        'Titik': 52,
        'c/h': '-∞',
        'c (mm)': '-∞',
        'Pn (kN)': f"{Tn*N_TO_KN:.4f}",
        'Mn (kN·m)': '0.0000',
        'φ': '0.900',
        'φPn (kN)': f"{(0.9*Tn)*N_TO_KN:.4f}",
        'φMn (kN·m)': '0.0000',
        'εt': 'N/A',
    })

    return points


# ----------------------------------------------------------------------
# PEMERIKSAAN KAPASITAS TERHADAP BEBAN
# ----------------------------------------------------------------------
def check_capacity(data: Dict[str, Any], interaction_points: List[Dict]) -> List[Dict[str, str]]:
    """
    Memeriksa apakah (Pu, Mc) berada di dalam diagram interaksi.
    """
    steps = []
    Pu = data['Pu']  # kN
    Mc = data.get('Mc', 0.0) * NMM_TO_KNM  # Nmm -> kN·m
    frame_type = data['frame_type']

    if frame_type == 'braced':
        Mu = Mc
    else:
        # Untuk unbraced, momen desain tidak dihitung secara lengkap
        Mu = data['M2']  # placeholder
        steps.append({
            'Item/Uraian': 'Peringatan',
            'Nilai': 'Unbraced',
            'Satuan': '-',
            'Keterangan': 'Pembesaran momen orde dua tidak dihitung, digunakan M2 langsung.'
        })

    # Ambil titik interaksi, konversi ke float
    phiPn_vals = [float(p['φPn (kN)']) for p in interaction_points]
    phiMn_vals = [float(p['φMn (kN·m)']) for p in interaction_points]

    # Cari kapasitas momen untuk Pu yang diberikan dengan interpolasi linier
    # Urutkan berdasarkan φPn
    sorted_indices = np.argsort(phiPn_vals)
    phiPn_sorted = np.array(phiPn_vals)[sorted_indices]
    phiMn_sorted = np.array(phiMn_vals)[sorted_indices]

    # Cari dua titik terdekat
    idx = np.searchsorted(phiPn_sorted, Pu)
    if idx == 0:
        # Pu terlalu rendah, gunakan dua titik pertama
        p1, m1 = phiPn_sorted[0], phiMn_sorted[0]
        p2, m2 = phiPn_sorted[1], phiMn_sorted[1]
    elif idx >= len(phiPn_sorted):
        # Pu di atas semua, ambil titik teratas (tekan murni)
        p1, m1 = phiPn_sorted[-2], phiMn_sorted[-2]
        p2, m2 = phiPn_sorted[-1], phiMn_sorted[-1]
    else:
        p1, m1 = phiPn_sorted[idx-1], phiMn_sorted[idx-1]
        p2, m2 = phiPn_sorted[idx], phiMn_sorted[idx]

    if p1 != p2:
        Mn_capacity = m1 + (m2 - m1) * (Pu - p1) / (p2 - p1)
    else:
        Mn_capacity = m1

    ratio = Mu / Mn_capacity if Mn_capacity != 0 else float('inf')
    status = "OK" if ratio <= 1.0 else "OVER"

    steps.append({
        'Item/Uraian': 'Pu (aksial terfaktor)',
        'Nilai': f"{Pu:.4f}",
        'Satuan': 'kN',
        'Keterangan': f"Pu = {Pu:.4f} kN"
    })
    steps.append({
        'Item/Uraian': 'Mu (momen desain)',
        'Nilai': f"{Mu:.4f}",
        'Satuan': 'kN·m',
        'Keterangan': f"Mu = {Mu:.4f} kN·m"
    })
    steps.append({
        'Item/Uraian': 'φMn kapasitas pada Pu',
        'Nilai': f"{Mn_capacity:.4f}",
        'Satuan': 'kN·m',
        'Keterangan': f"Interpolasi dari diagram interaksi pada Pn={Pu:.4f} kN → φMn = {Mn_capacity:.4f} kN·m"
    })
    steps.append({
        'Item/Uraian': 'Rasio (Mu/φMn)',
        'Nilai': f"{ratio:.4f}",
        'Satuan': '-',
        'Keterangan': f"Rasio = {Mu:.4f} / {Mn_capacity:.4f} = {ratio:.4f}"
    })
    steps.append({
        'Item/Uraian': 'Status Kapasitas',
        'Nilai': status,
        'Satuan': '-',
        'Keterangan': f"Rasio {ratio:.4f} → {status}"
    })

    return steps


# ----------------------------------------------------------------------
# FUNGSI UTAMA (dipanggil dari app.py)
# ----------------------------------------------------------------------
def run_calculation(input_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    Melakukan seluruh perhitungan kolom.
    Input: dict dengan semua parameter.
    Output: dict dengan status, pesan error (jika ada), dan tabel hasil.
    """
    # Validasi awal
    errors = validate_inputs(input_data)
    if errors:
        return {'status': 'error', 'errors': errors}

    # Siapkan data tambahan
    if 'beta_dns' not in input_data:
        input_data['beta_dns'] = 0.6
    if 'b_beam' not in input_data:
        input_data['b_beam'] = 300
    if 'h_beam' not in input_data:
        input_data['h_beam'] = 500
    if 'L_beam' not in input_data:
        input_data['L_beam'] = 6000
    if 'L_col_upper' not in input_data:
        input_data['L_col_upper'] = 3500
    if 'L_col_lower' not in input_data:
        input_data['L_col_lower'] = 3500

    # Properti
    prop_table = calc_properties(input_data)

    # Cek rasio tulangan sebagai penjagaan ekstra
    rho_g = 0.0
    for row in prop_table:
        if row['Item/Uraian'] == 'ρg':
            rho_g = float(row['Nilai'].replace('%', ''))
            break
    if not (1.0 <= rho_g <= 8.0):
        return {'status': 'error', 'errors': [f"Rasio tulangan ρg={rho_g:.2f}% di luar batas 1% - 8%."]}

    # Kelangsingan (memerlukan perhitungan k, jadi kita hitung di sini dan simpan k)
    slenderness_table = calc_slenderness(input_data, {})
    # Ambil k dari tabel slenderness
    k = None
    for row in slenderness_table:
        if row['Item/Uraian'] == 'Faktor k':
            k = float(row['Nilai'])
            break
    if k is None:
        k = 1.0
    input_data['k'] = k   # tambahkan ke input untuk dipakai selanjutnya

    # Pembesaran momen untuk braced frame (atau catatan untuk unbraced)
    magn_table = calc_moment_magnification(input_data, {})

    # Diagram interaksi
    interaction_points = calc_interaction_diagram(input_data)

    # Pemeriksaan kapasitas
    capacity_table = check_capacity(input_data, interaction_points)

    return {
        'status': 'ok',
        'tables': {
            'properties': prop_table,
            'slenderness': slenderness_table,
            'magnification': magn_table,
            'interaction': interaction_points,
            'capacity': capacity_table,
        }
    }
# ======================================================================
# BAGIAN UI STREAMLIT (ANTARMUKA)
# Tambahkan kode ini di bagian paling bawah file 7_kolomfull.py
# ======================================================================

import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from fpdf import FPDF
import datetime
def sanitize_for_pdf(text):
    """Mengganti simbol khusus agar tidak menyebabkan UnicodeEncodeError di FPDF2"""
    replacements = {
        'φ': 'phi', 'λ': 'lambda', '√': 'sqrt', '²': '^2', '³': '^3',
        'Σ': 'Sigma', 'π': 'pi', 'ρ': 'rho', 'ε': 'epsilon', '≤': '<=', '≥': '>='
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    return text

@st.cache_data
def create_pdf_report(tables, project_name="Project Kolom", engineer="Engineer Name"):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # --- Header ---
    pdf.set_font("Arial", 'B', 12)
    pdf.cell(0, 10, f"LAPORAN PERHITUNGAN STRUKTUR: {project_name}", ln=True, align='C')
    pdf.set_font("Arial", '', 10)
    pdf.cell(0, 5, f"Engineer: {engineer} | Tanggal: {datetime.date.today()}", ln=True, align='C')
    pdf.line(10, 27, 200, 27)
    pdf.ln(10)

    # --- Content ---
    for title, data in tables.items():
        if title == 'interaction': continue # Lewati tabel 52 titik di PDF agar tidak terlalu panjang
        
        pdf.set_font("Arial", 'B', 11)
        pdf.cell(0, 10, title.upper().replace('_', ' '), ln=True)
        
        # Table Header
        pdf.set_font("Arial", 'B', 9)
        pdf.set_fill_color(240, 240, 240)
        pdf.cell(45, 8, "Item", 1, 0, 'C', True)
        pdf.cell(30, 8, "Nilai", 1, 0, 'C', True)
        pdf.cell(20, 8, "Satuan", 1, 0, 'C', True)
        pdf.cell(95, 8, "Keterangan", 1, 1, 'C', True)
        
        # Table Body
        pdf.set_font("Arial", '', 8)
        for row in data:
            pdf.cell(45, 7, sanitize_for_pdf(str(row['Item/Uraian'])), 1)
            pdf.cell(30, 7, sanitize_for_pdf(str(row['Nilai'])), 1)
            pdf.cell(20, 7, sanitize_for_pdf(str(row['Satuan'])), 1)
            pdf.multi_cell(95, 7, sanitize_for_pdf(str(row['Keterangan'])), 1)
        pdf.ln(5)
    
    return pdf.output(dest='S')

@st.cache_data
def create_word_report(tables, project_name="Project Kolom", engineer="Engineer Name"):
    doc = Document()
    
    # Header Section
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Laporan Struktur - {project_name} | {engineer}"
    header_para.style.font.size = Pt(9)

    doc.add_heading(f'Hasil Perhitungan Kolom Beton', 0)
    
    for title, data in tables.items():
        if title == 'interaction': continue
        doc.add_heading(title.title().replace('_', ' '), level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Item'
        hdr_cells[1].text = 'Nilai'
        hdr_cells[2].text = 'Satuan'
        hdr_cells[3].text = 'Keterangan'
        
        for row_data in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(row_data['Item/Uraian'])
            row_cells[1].text = str(row_data['Nilai'])
            row_cells[2].text = str(row_data['Satuan'])
            row_cells[3].text = str(row_data['Keterangan'])
    
    target = BytesIO()
    doc.save(target)
    return target.getvalue()
def main():
    st.title("Kapasitas Kolom Beton Bertulang (P-M)")
    st.markdown("**Berdasarkan SNI 2847:2019**")

    # Setup Sidebar untuk Input
    st.sidebar.header("Input Parameter")

    st.sidebar.subheader("1. Data Material")
    fc = st.sidebar.number_input("Mutu Beton fc' (MPa)", value=30.0, step=1.0)
    fy = st.sidebar.number_input("Mutu Baja fy (MPa)", value=400.0, step=10.0)
    Es = st.sidebar.number_input("Modulus Elastisitas Es (MPa)", value=200000.0, step=1000.0)

    st.sidebar.subheader("2. Dimensi Kolom")
    b = st.sidebar.number_input("Lebar b (mm)", value=400.0, step=10.0)
    h = st.sidebar.number_input("Tinggi h (mm)", value=500.0, step=10.0)
    cover = st.sidebar.number_input("Selimut Beton (mm)", value=40.0, step=1.0)
    ds = st.sidebar.number_input("Diameter Sengkang Øs (mm)", value=10.0, step=1.0)

    st.sidebar.subheader("3. Tulangan Longitudinal")
    D = st.sidebar.number_input("Diameter Tulangan Utama D (mm)", value=22.0, step=1.0)
    n_b = st.sidebar.number_input("Jumlah Tulangan Sisi Lebar (b)", value=4, step=1)
    n_h = st.sidebar.number_input("Jumlah Tulangan Sisi Tinggi (h)", value=3, step=1)

    st.sidebar.subheader("4. Data Panjang & Rangka")
    Lu = st.sidebar.number_input("Panjang Tak Tertahan Lu (mm)", value=6000.0, step=100.0)
    frame_type = st.sidebar.selectbox("Kondisi Rangka", ["braced", "unbraced"])
    curvature = st.sidebar.selectbox("Kelengkungan", ["single", "double"])

    st.sidebar.subheader("5. Beban Terfaktor")
    Pu = st.sidebar.number_input("Gaya Aksial Pu (kN)", value=1500.0, step=10.0)
    M1 = st.sidebar.number_input("Momen Ujung 1 M1 (kN.m) [Kecil]", value=80.0, step=5.0)
    M2 = st.sidebar.number_input("Momen Ujung 2 M2 (kN.m) [Besar]", value=150.0, step=5.0)

    # Tombol Eksekusi
    if st.sidebar.button("Hitung Kapasitas", type="primary"):
        # Mapping input UI ke dictionary yang diminta oleh mesin kalkulasi
        input_data = {
            'fc': fc, 'fy': fy, 'Es': Es,
            'b': b, 'h': h, 'cover': cover, 'ds': ds,
            'D': D, 'n_b': n_b, 'n_h': n_h,
            'Lu': Lu, 'frame_type': frame_type, 'curvature': curvature,
            'Pu': Pu, 'M1': M1, 'M2': M2,
            # Parameter kekakuan elemen (bisa dibuat interaktif nanti, saat ini pakai default)
            'beta_dns': 0.6, 'b_beam': 300, 'h_beam': 500, 'L_beam': 6000,
            'L_col_upper': 3500, 'L_col_lower': 3500
        }

        with st.spinner('Menghitung kapasitas dan membuat diagram...'):
            result = run_calculation(input_data)

        if result['status'] == 'error':
            st.error("Terjadi Kesalahan Input:")
            for err in result['errors']:
                st.warning(f"- {err}")
        else:
            st.success("Perhitungan Selesai!")
            tables = result['tables']
if result['status'] == 'ok':
            st.success("Perhitungan Selesai!")
            tables = result['tables']

            # --- TAMBAHKAN MENU DOWNLOAD DI SINI ---
            st.divider()
            st.subheader("📥 Download Laporan")
            
            # Input tambahan untuk Header Laporan
            col_doc1, col_doc2 = st.columns(2)
            with col_doc1:
                proj_name = st.text_input("Nama Proyek", "Proyek Gedung A")
            with col_doc2:
                eng_name = st.text_input("Nama Engineer", "Istiyono")

            col_btn1, col_btn2 = st.columns(2)
            
            with col_btn1:
                pdf_data = create_pdf_report(tables, proj_name, eng_name)
                st.download_button(
                    label="Download PDF Report",
                    data=pdf_data,
                    file_name=f"Laporan_Kolom_{proj_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True
                )
            
            with col_btn2:
                word_data = create_word_report(tables, proj_name, eng_name)
                st.download_button(
                    label="Download Word Report",
                    data=word_data,
                    file_name=f"Laporan_Kolom_{proj_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            st.divider()
            # --- AKHIR MENU DOWNLOAD ---
    
            # Layout 2 Kolom untuk Tabel Hasil
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Properties Penampang")
                st.dataframe(pd.DataFrame(tables['properties']), use_container_width=True, hide_index=True)

                st.subheader("Cek Kelangsingan")
                st.dataframe(pd.DataFrame(tables['slenderness']), use_container_width=True, hide_index=True)

            with col2:
                st.subheader("Pembesaran Momen")
                st.dataframe(pd.DataFrame(tables['magnification']), use_container_width=True, hide_index=True)

                st.subheader("Cek Kapasitas")
                # Beri warna merah jika OVER, hijau jika OK
                df_cap = pd.DataFrame(tables['capacity'])
                st.dataframe(df_cap, use_container_width=True, hide_index=True)

            # Plot Diagram Interaksi
            st.divider()
            st.subheader("Diagram Interaksi P-M")
            
            df_int = pd.DataFrame(tables['interaction'])
            phi_Mn = pd.to_numeric(df_int['φMn (kN·m)'], errors='coerce')
            phi_Pn = pd.to_numeric(df_int['φPn (kN)'], errors='coerce')

            # Ambil nilai Mu untuk plotting titik beban aktual
            Mu_str = [row['Nilai'] for row in tables['capacity'] if row['Item/Uraian'] == 'Mu (momen desain)']
            Mu_val = float(Mu_str[0]) if Mu_str else M2

            fig, ax = plt.subplots(figsize=(8, 6))
            ax.plot(phi_Mn, phi_Pn, 'b-', linewidth=2.5, label='Kapasitas (φPn, φMn)')
            ax.plot(Mu_val, Pu, 'ro', markersize=8, label=f'Beban Aktual (Mu={Mu_val:.1f}, Pu={Pu:.1f})')
            
            ax.set_xlabel('Momen Rencana - φMn (kN.m)')
            ax.set_ylabel('Aksial Rencana - φPn (kN)')
            ax.grid(True, linestyle='--', alpha=0.7)
            ax.legend()
            ax.axhline(0, color='black', linewidth=1)
            ax.axvline(0, color='black', linewidth=1)
            
            st.pyplot(fig)

            # Tabel 52 Titik P-M
            with st.expander("Tampilkan Data 52 Titik Diagram Interaksi P-M"):
                st.dataframe(df_int, use_container_width=True, hide_index=True)

if __name__ == "__main__":
    main()
