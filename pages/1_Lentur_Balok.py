"""
=============================================================
HALAMAN 1 - EVALUASI KAPASITAS LENTUR, GESER & TORSI BALOK BETON BERTULANG
                (Tulangan Rangkap - Strain Compatibility)
Referensi : SNI 2847:2019 (ACI 318-14)
Framework : Streamlit (multipage)
Output    : Word (.docx) & PDF (.pdf) + Watermark
Session   : st.session_state untuk persistensi hasil

TAHAP PENGEMBANGAN:
  [v] TAHAP 0 : Lentur Tulangan Rangkap (b1 - b10)
  [v] TAHAP 1 : Geser SNI 2847:2019      (b11 - b18)  +  Timestamp
  [v] TAHAP 2 : Torsi SNI 2847:2019      (b19 - b28)
=============================================================
"""

import math
import io
import datetime
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

# ============================================================
# KONFIGURASI HALAMAN
# ============================================================
st.set_page_config(
    page_title="Evaluasi Lentur, Geser & Torsi Balok | SNI 2847:2019",
    page_icon="🏗",
    layout="wide",
)

st.markdown("""
<style>
  .main-title{font-size:1.6rem;font-weight:600;color:#1a3c5e;margin-bottom:0}
  .sub-title{font-size:.9rem;color:#666;margin-bottom:1.5rem}
  .step-box{background:#f8f9fa;border-left:4px solid #1a3c5e;border-radius:0 8px 8px 0;
            padding:12px 16px;margin-bottom:10px;font-family:monospace;font-size:.85rem}
  .step-hdr{font-weight:700;color:#1a3c5e;font-size:.8rem;
            text-transform:uppercase;letter-spacing:.5px;margin-bottom:4px}
  .result-ok  {background:#e8f5e9;border-left:4px solid #2e7d32;color:#1b5e20;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .result-warn{background:#fff8e1;border-left:4px solid #f9a825;color:#5d4037;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .result-fail{background:#ffebee;border-left:4px solid #c62828;color:#b71c1c;
               padding:12px 16px;border-radius:8px;font-weight:600;margin-bottom:8px}
  .metric-card{background:white;border:1px solid #e0e0e0;border-radius:10px;
               padding:14px;text-align:center}
  .metric-lbl{font-size:.75rem;color:#888;margin-bottom:4px}
  .metric-val{font-size:1.5rem;font-weight:700;color:#1a3c5e}
  .metric-unt{font-size:.75rem;color:#aaa}
  .ref-badge{display:inline-block;background:#e3f2fd;color:#1565c0;
             font-size:.7rem;padding:2px 8px;border-radius:20px;margin-bottom:6px}
  .group-hdr{font-size:.85rem;font-weight:700;color:#1a3c5e;letter-spacing:.5px;
             text-transform:uppercase;margin-top:.6rem;margin-bottom:.3rem}
  hr.divider{border:none;border-top:2px solid #e0e0e0;margin:1.5rem 0}
</style>
""", unsafe_allow_html=True)


# ============================================================
# SANITASI STRING UNTUK PDF (Latin-1 only)
# ============================================================
_UNICODE_MAP = {
    "\u2014": "-",    "\u2013": "-",    "\u2019": "'",   "\u2018": "'",
    "\u201c": '"',    "\u201d": '"',    "\u00b2": "2",   "\u00b3": "3",
    "\u00b0": " deg", "\u00d7": "x",   "\u2265": ">=",  "\u2264": "<=",
    "\u2260": "!=",   "\u221a": "sqrt", "\u03c6": "Phi", "\u03b2": "Beta",
    "\u03b5": "et",   "\u03c1": "Rho",  "\u03bc": "mu",  "\u2022": "-",
    "\u2192": "->",   "\u00b7": ".",    "\u00e9": "e",   "\u00e8": "e",
    "\u00e0": "a",    "\u2550": "=",    "\u2500": "-",
    "\u03bb": "lambda",
}

def sp(teks: str) -> str:
    """Sanitasi string agar aman untuk fpdf2 (Latin-1 only)."""
    if not isinstance(teks, str):
        teks = str(teks)
    for ch, repl in _UNICODE_MAP.items():
        teks = teks.replace(ch, repl)
    return teks.encode("latin-1", errors="replace").decode("latin-1")


# ============================================================
# LIBRARY DIAMETER TULANGAN (mm)
# ============================================================
DIAMETER_LIST = [10, 13, 16, 19, 22, 25, 29, 32]

def luas_batang(d_mm: float) -> float:
    """Luas penampang satu batang tulangan (mm2)."""
    return math.pi * d_mm ** 2 / 4.0


# ============================================================
# FUNGSI HITUNG GEOMETRI TULANGAN OTOMATIS
# ============================================================
def hitung_lapis_tarik(h, cc, ds, n1, db1, n2, db2, spasi=25.0):
    lapis = []
    if n1 > 0 and db1 > 0:
        y1 = h - cc - ds - db1 / 2.0
        As1 = n1 * luas_batang(db1)
        lapis.append(dict(no=1, n=n1, db=db1, As=As1, y=y1))
    if n2 > 0 and db2 > 0:
        if n1 > 0 and db1 > 0:
            y_top_lapis1 = h - cc - ds - db1
            y2 = y_top_lapis1 - spasi - db2 / 2.0
        else:
            y2 = h - cc - ds - db2 / 2.0
        As2 = n2 * luas_batang(db2)
        lapis.append(dict(no=2, n=n2, db=db2, As=As2, y=y2))

    if not lapis:
        return [], 0.0, 0.0

    As_total = sum(L["As"] for L in lapis)
    d_total  = sum(L["As"] * L["y"] for L in lapis) / As_total
    return lapis, d_total, As_total


def hitung_lapis_tekan(cc, ds, n1, db1, n2, db2, spasi=25.0):
    lapis = []
    if n1 > 0 and db1 > 0:
        y1 = cc + ds + db1 / 2.0
        As1 = n1 * luas_batang(db1)
        lapis.append(dict(no=1, n=n1, db=db1, As=As1, y=y1))
    if n2 > 0 and db2 > 0:
        if n1 > 0 and db1 > 0:
            y_bot_lapis1 = cc + ds + db1
            y2 = y_bot_lapis1 + spasi + db2 / 2.0
        else:
            y2 = cc + ds + db2 / 2.0
        As2 = n2 * luas_batang(db2)
        lapis.append(dict(no=2, n=n2, db=db2, As=As2, y=y2))

    if not lapis:
        return [], 0.0, 0.0

    As_total = sum(L["As"] for L in lapis)
    d_total  = sum(L["As"] * L["y"] for L in lapis) / As_total
    return lapis, d_total, As_total


# ============================================================
# FUNGSI EVALUASI LENTUR (STRAIN COMPATIBILITY)
# ============================================================
def hitung_beta1(fc):
    """Beta-1 SNI 2847:2019 Pasal 22.2.2.4.3"""
    if fc <= 28:
        return 0.85, f"fc = {fc} MPa <= 28 MPa  -->  Beta-1 = 0.85"
    elif fc >= 56:
        return 0.65, f"fc = {fc} MPa >= 56 MPa  -->  Beta-1 = 0.65"
    else:
        b = max(0.65, min(0.85, 0.85 - 0.05 * (fc - 28) / 7))
        return b, (f"Beta-1 = 0.85 - 0.05 x (fc - 28) / 7\n"
                   f"       = 0.85 - 0.05 x ({fc} - 28) / 7")


def gaya_dalam(c, fc, fy, b, beta1, lapis_tarik, lapis_tekan):
    Es      = 200_000.0
    eps_cu  = 0.003
    eps_y   = fy / Es
    a       = beta1 * c

    Cc = 0.85 * fc * a * b

    Cs_total = 0.0
    info_tekan = []
    for L in lapis_tekan:
        eps_s = eps_cu * (c - L["y"]) / c
        if eps_s >= eps_y:
            fs = fy
        elif eps_s <= -eps_y:
            fs = -fy
        else:
            fs = eps_s * Es
        if L["y"] <= a:
            fs_eff = fs - 0.85 * fc
        else:
            fs_eff = fs
        F = L["As"] * fs_eff
        Cs_total += F
        info_tekan.append(dict(no=L["no"], y=L["y"], As=L["As"],
                               eps=eps_s, fs=fs, F=F))

    T_total = 0.0
    info_tarik = []
    for L in lapis_tarik:
        eps_s = eps_cu * (L["y"] - c) / c
        if eps_s >= eps_y:
            fs = fy
        elif eps_s <= -eps_y:
            fs = -fy
        else:
            fs = eps_s * Es
        F = L["As"] * fs
        T_total += F
        info_tarik.append(dict(no=L["no"], y=L["y"], As=L["As"],
                               eps=eps_s, fs=fs, F=F))

    residu = Cc + Cs_total - T_total
    return dict(
        a=a, c=c, Cc=Cc, Cs=Cs_total, T=T_total,
        info_tekan=info_tekan, info_tarik=info_tarik,
        residu=residu, eps_y=eps_y,
    )


def cari_c_keseimbangan(fc, fy, b, beta1, lapis_tarik, lapis_tekan, h):
    """Iterasi bisection untuk c sehingga residu (Cc+Cs-T) = 0."""
    lo, hi = 0.001, h * 1.5
    for _ in range(200):
        mid = (lo + hi) / 2.0
        r = gaya_dalam(mid, fc, fy, b, beta1, lapis_tarik, lapis_tekan)["residu"]
        r_lo = gaya_dalam(lo, fc, fy, b, beta1, lapis_tarik, lapis_tekan)["residu"]
        if abs(r) < 1e-3:
            return mid
        if r_lo * r < 0:
            hi = mid
        else:
            lo = mid
    return (lo + hi) / 2.0


def evaluasi_balok(fc, fy, b, h, cc_sel, ds, lapis_tarik, lapis_tekan,
                   d_tarik, As_tarik, d_tekan, As_tekan, Mu):
    """Evaluasi penuh penampang LENTUR dengan strain compatibility."""
    R = {}
    beta1, beta1_cara = hitung_beta1(fc)
    R["beta1"]      = beta1
    R["beta1_cara"] = beta1_cara

    c = cari_c_keseimbangan(fc, fy, b, beta1, lapis_tarik, lapis_tekan, h)
    G = gaya_dalam(c, fc, fy, b, beta1, lapis_tarik, lapis_tekan)

    a       = G["a"]
    Cc      = G["Cc"]
    Cs      = G["Cs"]
    T       = G["T"]
    info_tk = G["info_tekan"]
    info_tr = G["info_tarik"]
    eps_y   = G["eps_y"]

    Mn_Nmm = Cc * (d_tarik - a / 2)
    for L in info_tk:
        Mn_Nmm += L["F"] * (d_tarik - L["y"])
    Mn = Mn_Nmm / 1_000_000.0

    if info_tr:
        et = max(L["eps"] for L in info_tr)
    else:
        et = 0.0

    if et >= 0.005:
        phi = 0.90
        phi_cara = "et >= 0.005  -->  Tension-controlled  -->  Phi = 0.90"
    elif et <= 0.002:
        phi = 0.65
        phi_cara = "et <= 0.002  -->  Compression-controlled  -->  Phi = 0.65"
    else:
        phi = 0.65 + (et - 0.002) * (250 / 3)
        phi_cara = (f"0.002 < et < 0.005  -->  Zona transisi\n"
                    f"Phi = 0.65 + (et - 0.002) x 250/3\n"
                    f"    = 0.65 + ({et:.5f} - 0.002) x 83.333")

    phiMn = phi * Mn
    DC    = Mu / phiMn if phiMn > 0 else float("inf")

    rho      = As_tarik / (b * d_tarik)
    rho_mA   = 0.25 * math.sqrt(fc) / fy
    rho_mB   = 1.4 / fy
    rho_min  = max(rho_mA, rho_mB)
    rho_bal  = (0.85 * beta1 * fc / fy) * (600 / (600 + fy))
    rho_max  = 0.75 * rho_bal

    R.update(dict(
        c=c, a=a, Cc=Cc, Cs=Cs, T=T,
        info_tarik=info_tr, info_tekan=info_tk,
        eps_y=eps_y, et=et, phi=phi, phi_cara=phi_cara,
        Mn=Mn, phiMn=phiMn, Mu=Mu, DC=DC,
        rho=rho, rho_mA=rho_mA, rho_mB=rho_mB,
        rho_min=rho_min, rho_bal=rho_bal, rho_max=rho_max,
        ok_rho_min=(rho >= rho_min),
        ok_rho_max=(rho <= rho_max),
        ok_et=(et >= 0.004),
        ok_dc=(DC <= 1.0),
        d_tarik=d_tarik, d_tekan=d_tekan,
        As_tarik=As_tarik, As_tekan=As_tekan,
    ))
    return R


# ============================================================
# FUNGSI EVALUASI GESER (SNI 2847:2019)  -- TAHAP 1
# ============================================================
def hitung_geser(fc, fyt, b, d_aktual, Vu, ds, s_seng, n_kaki, lambda_=1.0):
    G = {}

    Phi_v = 0.75
    Vc_N  = 0.17 * lambda_ * math.sqrt(fc) * b * d_aktual
    Vc    = Vc_N / 1000.0

    half_phiVc = 0.5 * Phi_v * Vc
    phiVc      = Phi_v * Vc
    if Vu <= half_phiVc:
        klas_seng = "tidak_perlu"
        ket_klas  = (f"Vu = {Vu:.2f} kN <= 0.5 Phi.Vc = {half_phiVc:.2f} kN  "
                     f"-->  secara teori tidak perlu sengkang, "
                     f"tetap pasang sengkang minimum.")
    elif Vu <= phiVc:
        klas_seng = "minimum"
        ket_klas  = (f"0.5 Phi.Vc = {half_phiVc:.2f} kN < Vu = {Vu:.2f} kN <= "
                     f"Phi.Vc = {phiVc:.2f} kN  -->  pasang sengkang minimum.")
    else:
        klas_seng = "perlu_hitung"
        ket_klas  = (f"Vu = {Vu:.2f} kN > Phi.Vc = {phiVc:.2f} kN  "
                     f"-->  perlu menghitung tulangan geser (Vs).")

    Vs_perlu = max(0.0, Vu / Phi_v - Vc)
    Vs_max   = 0.66 * math.sqrt(fc) * b * d_aktual / 1000.0
    ok_Vs_max = (Vs_perlu <= Vs_max)

    Av_pasang = n_kaki * (math.pi / 4.0) * ds ** 2

    Vs_batas_spasi = 0.33 * math.sqrt(fc) * b * d_aktual / 1000.0
    if Vs_perlu <= Vs_batas_spasi:
        s_max = min(d_aktual / 2.0, 600.0)
        rumus_smax = (f"Vs_perlu = {Vs_perlu:.2f} kN <= 0.33 sqrt(fc) b d / 1000 "
                      f"= {Vs_batas_spasi:.2f} kN\n"
                      f"  -->  s_max = min(d/2 , 600 mm) = "
                      f"min({d_aktual/2.0:.1f} , 600) = {s_max:.1f} mm")
    else:
        s_max = min(d_aktual / 4.0, 300.0)
        rumus_smax = (f"Vs_perlu = {Vs_perlu:.2f} kN > 0.33 sqrt(fc) b d / 1000 "
                      f"= {Vs_batas_spasi:.2f} kN\n"
                      f"  -->  s_max = min(d/4 , 300 mm) = "
                      f"min({d_aktual/4.0:.1f} , 300) = {s_max:.1f} mm")
    ok_spasi = (s_seng <= s_max)

    AvS_minA = 0.062 * math.sqrt(fc) * b / fyt
    AvS_minB = 0.35 * b / fyt
    AvS_min  = max(AvS_minA, AvS_minB)
    Av_min   = AvS_min * s_seng
    ok_Av    = (Av_pasang >= Av_min)

    Vs_aktual  = (Av_pasang * fyt * d_aktual / s_seng) / 1000.0
    Vs_efektif = min(Vs_aktual, Vs_max)
    Vn_efektif = Vc + Vs_efektif
    PhiVn_efektif = Phi_v * Vn_efektif
    DC_geser   = Vu / PhiVn_efektif if PhiVn_efektif > 0 else float("inf")
    ok_dc      = (DC_geser <= 1.0)
    ok_total   = ok_dc and ok_Vs_max and ok_Av and ok_spasi

    G.update(dict(
        Phi_v=Phi_v, lambda_=lambda_,
        Vc_N=Vc_N, Vc=Vc, half_phiVc=half_phiVc, phiVc=phiVc,
        klas_seng=klas_seng, ket_klas=ket_klas,
        Vs_perlu=Vs_perlu, Vs_max=Vs_max, ok_Vs_max=ok_Vs_max,
        Av_pasang=Av_pasang,
        AvS_minA=AvS_minA, AvS_minB=AvS_minB, AvS_min=AvS_min,
        Av_min=Av_min, ok_Av=ok_Av,
        s_max=s_max, rumus_smax=rumus_smax, ok_spasi=ok_spasi,
        Vs_aktual=Vs_aktual, Vs_efektif=Vs_efektif,
        Vn_efektif=Vn_efektif, PhiVn_efektif=PhiVn_efektif,
        DC_geser=DC_geser, ok_dc=ok_dc, ok_total=ok_total,
        Vu=Vu,
    ))
    return G


# ============================================================
# FUNGSI EVALUASI TORSI (SNI 2847:2019 Pasal 22.7) -- TAHAP 2
# ============================================================
def hitung_torsi(fc, fy_long, fyt, b, h, cc_sel, ds, s_seng,
                 Tu, Vu, Vc, d_aktual, Av_pasang_geser,
                 tipe_torsi, db_long_torsi, lambda_=1.0):
    """
    Evaluasi torsi balok beton bertulang sesuai SNI 2847:2019 Pasal 22.7.

    Parameter
    ---------
    fc            : kuat tekan beton (MPa)
    fy_long       : kuat leleh tulangan longitudinal torsi (MPa)
    fyt           : kuat leleh sengkang (MPa)
    b, h          : dimensi penampang (mm)
    cc_sel        : tebal selimut bersih (mm)
    ds            : diameter sengkang (mm)
    s_seng        : jarak sengkang terpasang (mm)
    Tu            : momen torsi ultimit (kN.m)
    Vu            : gaya geser ultimit (kN) -- dari hasil geser
    Vc            : kapasitas geser beton (kN) -- dari hasil geser
    d_aktual      : tinggi efektif (mm) -- dari hasil lentur
    Av_pasang_geser: Av terpasang untuk geser (mm2) -- dari hasil geser
    tipe_torsi    : "Equilibrium" atau "Compatibility"
    db_long_torsi : diameter tulangan longitudinal torsi (mm)
    lambda_       : faktor beton ringan (1.0 normal)

    Returns
    -------
    dict hasil evaluasi torsi
    """
    T = {}
    Phi_t = 0.75   # SNI 2847:2019 Pasal 21.2.1

    # -- b19: Properti penampang torsi --
    Acp = b * h                              # mm2
    pcp = 2.0 * (b + h)                     # mm

    x1  = b - 2.0 * cc_sel - ds             # mm - dimensi pendek sengkang (c to c)
    y1  = h - 2.0 * cc_sel - ds             # mm - dimensi panjang sengkang (c to c)
    Aoh = x1 * y1                           # mm2
    ph  = 2.0 * (x1 + y1)                  # mm
    Ao  = 0.85 * Aoh                        # mm2

    # -- b20: Faktor reduksi --
    # Phi_t = 0.75 (sudah di atas)

    # -- b21: Batas ambang torsi (Tth) SNI Pasal 22.7.4 --
    Tth_Nmm = 0.083 * lambda_ * math.sqrt(fc) * (Acp ** 2 / pcp)   # N.mm
    Tth      = Tth_Nmm / 1_000_000.0                                 # kN.m
    phi_Tth  = Phi_t * Tth

    Tu_Nmm = Tu * 1_000_000.0   # konversi kN.m -> N.mm
    abaikan_torsi = (Tu <= phi_Tth)

    # -- b22: Klasifikasi & Tu desain --
    Tcr_Nmm = 0.33 * lambda_ * math.sqrt(fc) * (Acp ** 2 / pcp)   # N.mm
    Tcr      = Tcr_Nmm / 1_000_000.0                               # kN.m
    phi_Tcr  = Phi_t * Tcr

    if tipe_torsi == "Compatibility" and Tu > phi_Tcr:
        Tu_desain     = phi_Tcr   # kN.m -- boleh direduksi
        Tu_desain_Nmm = phi_Tcr * 1_000_000.0
        catatan_tu    = (f"Torsi Kompatibilitas: Tu ({Tu:.3f} kN.m) > Phi_t x Tcr "
                         f"({phi_Tcr:.3f} kN.m)\n"
                         f"  -->  Tu_desain direduksi = Phi_t x Tcr = {Tu_desain:.3f} kN.m")
    else:
        Tu_desain     = Tu        # kN.m
        Tu_desain_Nmm = Tu_Nmm
        if tipe_torsi == "Equilibrium":
            catatan_tu = (f"Torsi Keseimbangan: tidak boleh redistribusi\n"
                          f"  -->  Tu_desain = Tu = {Tu_desain:.3f} kN.m")
        else:
            catatan_tu = (f"Torsi Kompatibilitas: Tu ({Tu:.3f} kN.m) <= Phi_t x Tcr "
                          f"({phi_Tcr:.3f} kN.m)\n"
                          f"  -->  Tu_desain = Tu = {Tu_desain:.3f} kN.m (tidak direduksi)")

    # -- b23: Cek dimensi penampang (SNI Pasal 22.7.7.1) --
    # sqrt( (Vu/(b*d))^2 + (Tu*ph/(1.7*Aoh^2))^2 ) <= Phi_v*(Vc/(b*d) + 0.66*sqrt(fc))
    # Vu dalam N, Tu dalam N.mm, b dan d dalam mm
    Vu_N       = Vu * 1000.0
    Vc_N_val   = Vc * 1000.0
    Phi_v      = 0.75

    term_geser = Vu_N / (b * d_aktual)                               # N/mm2 = MPa
    term_torsi = Tu_desain_Nmm * ph / (1.7 * Aoh ** 2)              # MPa
    lhs_dim    = math.sqrt(term_geser ** 2 + term_torsi ** 2)       # MPa

    rhs_dim    = Phi_v * (Vc_N_val / (b * d_aktual) + 0.66 * math.sqrt(fc))  # MPa
    ok_dimensi = (lhs_dim <= rhs_dim)
    DC_dim     = lhs_dim / rhs_dim if rhs_dim > 0 else float("inf")

    # -- b24: At/s untuk torsi (SNI Pasal 22.7.6.1) --
    # theta = 45 deg -> cot(theta) = 1.0
    # At/s = Tu_desain / (Phi_t x 2 x Ao x fyt x cot(theta))
    cot_theta = 1.0
    At_per_s  = Tu_desain_Nmm / (Phi_t * 2.0 * Ao * fyt * cot_theta)   # mm2/mm

    # -- b25: Sengkang gabungan geser + torsi (SNI Pasal 9.6.4.2) --
    # (Av/s) geser dari tulangan terpasang
    Av_per_s_pasang = Av_pasang_geser / s_seng                # mm2/mm
    # Total kebutuhan per mm = Av/s + 2*At/s
    Avt_per_s_perlu = Av_per_s_pasang + 2.0 * At_per_s       # mm2/mm
    Avt_per_s_ada   = Av_pasang_geser / s_seng                # mm2/mm (yg terpasang untuk geser)

    # Cek apakah perlu tambah sengkang atau spasi diperketat
    Avt_pasang      = Av_pasang_geser        # luas sengkang terpasang (mm2)
    Avt_perlu_abs   = Avt_per_s_perlu * s_seng   # mm2 yang dibutuhkan pada s yg sama

    # Spasi yg diperlukan jika pakai sengkang yg sama
    if Avt_per_s_perlu > 0:
        s_perlu_gabungan = Av_pasang_geser / Avt_per_s_perlu
    else:
        s_perlu_gabungan = s_seng
    ok_sengkang_gabungan = (Av_pasang_geser >= Avt_perlu_abs)

    # -- Av+2At minimum (SNI Pasal 9.6.4.2) --
    AvtS_minA = 0.062 * math.sqrt(fc) * b / fyt    # mm2/mm
    AvtS_minB = 0.35 * b / fyt                     # mm2/mm
    AvtS_min  = max(AvtS_minA, AvtS_minB)
    Avt_min   = AvtS_min * s_seng                  # mm2
    ok_Avt_min = (Av_pasang_geser >= Avt_min)

    # -- b26: Spasi maksimum sengkang torsi (SNI Pasal 9.7.6.3.3) --
    s_max_torsi = min(ph / 8.0, 300.0)
    ok_spasi_torsi = (s_seng <= s_max_torsi)

    # -- b27: Tulangan longitudinal torsi Al (SNI Pasal 22.7.6.1) --
    Al = At_per_s * ph * (fyt / fy_long) * (cot_theta ** 2)   # mm2

    # Al minimum (SNI Pasal 9.6.4.3)
    Al_minA = (0.42 * math.sqrt(fc) * Acp / fy_long
               - At_per_s * ph * (fyt / fy_long))
    Al_minB = (0.42 * math.sqrt(fc) * Acp / fy_long
               - (0.175 * b / fyt) * ph * (fyt / fy_long))
    Al_min  = max(Al_minA, Al_minB, 0.0)   # tidak boleh negatif

    Al_pakai = max(Al, Al_min)

    # Jumlah batang longitudinal
    Ab_long   = luas_batang(db_long_torsi)   # mm2 per batang
    n_batang  = math.ceil(Al_pakai / Ab_long) if Ab_long > 0 else 0

    # -- b28: DC Torsi --
    # Phi.Tn = Phi_t x 2 x Ao x At_per_s x fyt x cot_theta x (1/ph) x ph
    # Cara lebih langsung: Phi.Tn = Phi_t x 2 x Ao x (Al/ph) x fy_long / cot^2
    # Lebih tepat: hitung Tn dari kapasitas
    # Tn = 2 x Ao x At/s x fyt x cot(theta) [SNI 22.7.6.1]
    # tapi At/s di sini sudah dari Tu_desain, jadi DC = Tu_desain / (Phi_t * Tn_cap)
    # Gunakan sengkang yang terpasang untuk Phi.Tn kapasitas:
    At_per_s_pasang = Av_pasang_geser / (2.0 * s_seng)   # 1/2 Av utk torsi (1 kaki per sisi)
    # Sebenarnya untuk closed stirrup, setiap kaki = At, jadi:
    At_per_s_cap    = Av_pasang_geser / s_seng            # konservatif, total kaki / s
    Tn_cap          = 2.0 * Ao * At_per_s_cap * fyt * cot_theta / 1_000_000.0   # kN.m
    PhiTn_cap       = Phi_t * Tn_cap
    DC_torsi        = Tu_desain / PhiTn_cap if PhiTn_cap > 0 else float("inf")
    ok_DC_torsi     = (DC_torsi <= 1.0)

    ok_torsi_total = (ok_dimensi and ok_sengkang_gabungan and ok_Avt_min
                      and ok_spasi_torsi and ok_DC_torsi)

    T.update(dict(
        Phi_t=Phi_t, lambda_=lambda_,
        Acp=Acp, pcp=pcp, x1=x1, y1=y1, Aoh=Aoh, ph=ph, Ao=Ao,
        Tth=Tth, phi_Tth=phi_Tth, Tth_Nmm=Tth_Nmm,
        Tcr=Tcr, phi_Tcr=phi_Tcr,
        abaikan_torsi=abaikan_torsi,
        Tu=Tu, Tu_desain=Tu_desain, Tu_desain_Nmm=Tu_desain_Nmm,
        tipe_torsi=tipe_torsi, catatan_tu=catatan_tu,
        term_geser=term_geser, term_torsi=term_torsi,
        lhs_dim=lhs_dim, rhs_dim=rhs_dim,
        ok_dimensi=ok_dimensi, DC_dim=DC_dim,
        cot_theta=cot_theta, At_per_s=At_per_s,
        Av_per_s_pasang=Av_per_s_pasang,
        Avt_per_s_perlu=Avt_per_s_perlu,
        Avt_perlu_abs=Avt_perlu_abs,
        Avt_pasang=Avt_pasang,
        s_perlu_gabungan=s_perlu_gabungan,
        ok_sengkang_gabungan=ok_sengkang_gabungan,
        AvtS_minA=AvtS_minA, AvtS_minB=AvtS_minB,
        AvtS_min=AvtS_min, Avt_min=Avt_min, ok_Avt_min=ok_Avt_min,
        s_max_torsi=s_max_torsi, ok_spasi_torsi=ok_spasi_torsi,
        Al=Al, Al_minA=Al_minA, Al_minB=Al_minB,
        Al_min=Al_min, Al_pakai=Al_pakai,
        db_long_torsi=db_long_torsi, Ab_long=Ab_long, n_batang=n_batang,
        At_per_s_cap=At_per_s_cap, Tn_cap=Tn_cap,
        PhiTn_cap=PhiTn_cap, DC_torsi=DC_torsi,
        ok_DC_torsi=ok_DC_torsi, ok_torsi_total=ok_torsi_total,
        Vu=Vu, Vc=Vc, d_aktual=d_aktual,
        fy_long=fy_long, fyt=fyt,
    ))
    return T


# ============================================================
# BUAT STEP-STEP LAPORAN LENTUR (b1 - b10)
# ============================================================
def buat_steps_balok(fc, fy, b, h, cc_sel, ds, lapis_tarik, lapis_tekan, R):
    eps_y   = R["eps_y"]
    ok_rho  = R["ok_rho_min"] and R["ok_rho_max"]
    Mu      = R["Mu"]

    s1 = dict(
        no="b1.", ref="SNI 2847:2019 Pasal 22.2.2.4.3",
        judul="Faktor Beta-1",
        isi=f"{R['beta1_cara']}\n  -->  Beta-1 = {R['beta1']:.4f}",
        ok=True,
    )

    teks_lapis = ""
    for L in lapis_tarik:
        teks_lapis += (f"  Tarik Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                       f"As = {L['As']:.1f} mm2   y (dari atas) = {L['y']:.2f} mm\n")
    for L in lapis_tekan:
        teks_lapis += (f"  Tekan Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                       f"As' = {L['As']:.1f} mm2   y (dari atas) = {L['y']:.2f} mm\n")

    s2 = dict(
        no="b2.", ref="SNI 2847:2019 Pasal 26.6.2 & geom. penampang",
        judul="Posisi tiap lapis tulangan (dari serat tekan atas)",
        isi=(
            f"Penentuan posisi y = h - cc - ds - db/2 (tarik bawah)\n"
            f"                 y = cc + ds + db/2 (tekan atas)\n\n"
            f"{teks_lapis.rstrip()}\n\n"
            f"d-aktual  (tarik) = {R['d_tarik']:.2f} mm  (titik berat As-tarik)\n"
            f"d'-aktual (tekan) = {R['d_tekan']:.2f} mm  (titik berat As-tekan)"
        ),
        ok=True,
    )

    s3 = dict(
        no="b3.", ref="SNI 2847:2019 Pasal 22.2.1.1 - eps_cu = 0.003",
        judul="Iterasi sumbu netral c (bisection, 200 iterasi)",
        isi=(
            f"Cari c sehingga Cc + Cs - T = 0\n"
            f"  a     = Beta-1 x c = {R['beta1']:.4f} x {R['c']:.2f} = {R['a']:.2f} mm\n"
            f"  Cc    = 0.85 x fc x a x b = 0.85 x {fc} x {R['a']:.2f} x {b} "
            f"= {R['Cc']:,.0f} N\n"
            f"  c     = {R['c']:.4f} mm  -->  [OK]"
        ),
        ok=True,
    )

    teks_eps = ""
    for L in R["info_tekan"]:
        status = "leleh" if abs(L["eps"]) >= eps_y else "elastis"
        teks_eps += (f"  Tekan Lapis-{L['no']}  (y={L['y']:.1f} mm):\n"
                     f"    eps-s' = 0.003 x (c - y) / c = "
                     f"0.003 x ({R['c']:.2f} - {L['y']:.1f}) / {R['c']:.2f} "
                     f"= {L['eps']:+.5f}  ({status})\n"
                     f"    fs'    = {L['fs']:+.2f} MPa\n")
    for i, L in enumerate(R["info_tarik"]):
        notasi = "et" if L["no"] == 1 else f"eps-s{L['no']}"
        status = "leleh" if abs(L["eps"]) >= eps_y else "elastis"
        teks_eps += (f"  Tarik Lapis-{L['no']} (y={L['y']:.1f} mm):\n"
                     f"    {notasi:<6} = 0.003 x (y - c) / c = "
                     f"0.003 x ({L['y']:.1f} - {R['c']:.2f}) / {R['c']:.2f} "
                     f"= {L['eps']:+.5f}  ({status})\n"
                     f"    fs     = {L['fs']:+.2f} MPa\n")

    s4 = dict(
        no="b4.", ref="SNI 2847:2019 Pasal 22.2.1.1 - Strain linear",
        judul="Regangan & tegangan baja per lapis",
        isi=(
            f"Tegangan tiap lapis: jika |eps| < eps-y -> fs = eps x 200000\n"
            f"                     jika |eps| >= eps-y -> fs = +/- fy\n\n"
            f"{teks_eps.rstrip()}"
        ),
        ok=True,
    )

    teks_F = ""
    for L in R["info_tekan"]:
        teks_F += (f"  Cs Lapis-{L['no']} = As' x (fs' - 0.85 fc)" if L["y"] <= R["a"]
                   else f"  Cs Lapis-{L['no']} = As' x fs'")
        teks_F += f" = {L['F']:+,.0f} N\n"
    for L in R["info_tarik"]:
        teks_F += f"  T  Lapis-{L['no']} = As x fs = {L['F']:+,.0f} N\n"

    s5 = dict(
        no="b5.", ref="SNI 2847:2019 Pasal 22.2.2 - Distribusi tegangan",
        judul="Gaya tekan beton (Cc), gaya baja tekan (Cs), gaya tarik (T)",
        isi=(
            f"Cc = 0.85 x fc x a x b\n"
            f"   = 0.85 x {fc} x {R['a']:.2f} x {b}\n"
            f"   = {R['Cc']:,.0f} N  =  {R['Cc']/1000:.2f} kN\n\n"
            f"{teks_F}"
            f"\nCek keseimbangan: Cc + Cs - T = "
            f"{R['Cc']:,.0f} + {R['Cs']:,.0f} - {R['T']:,.0f} = "
            f"{R['Cc']+R['Cs']-R['T']:,.0f} N  (~ 0, OK)"
        ),
        ok=True,
    )

    s6 = dict(
        no="b6.", ref="SNI 2847:2019 Pasal 9.6.1 & 21.2.2",
        judul="Rasio tulangan tarik (Rho) - kontrol min & max",
        isi=(
            f"Rho     = As / (b x d) = {R['As_tarik']:.1f} / "
            f"({b:.0f} x {R['d_tarik']:.2f}) = {R['rho']*100:.4f}%\n"
            f"Rho-min = max(0.25 sqrt(fc)/fy , 1.4/fy)\n"
            f"        = max({R['rho_mA']*100:.4f}% , {R['rho_mB']*100:.4f}%) "
            f"= {R['rho_min']*100:.4f}%\n"
            f"Rho-bal = 0.85 Beta-1 fc/fy x 600/(600+fy) = {R['rho_bal']*100:.4f}%\n"
            f"Rho-max = 0.75 Rho-bal = {R['rho_max']*100:.4f}%\n\n"
            f"Kontrol Rho >= Rho-min : "
            f"{'[OK]' if R['ok_rho_min'] else '[TIDAK OK]'}\n"
            f"Kontrol Rho <= Rho-max : "
            f"{'[OK]' if R['ok_rho_max'] else '[TIDAK OK]'}"
        ),
        ok=ok_rho,
    )

    if R["et"] >= 0.005:
        klas = "et >= 0.005  -->  Tension-controlled  [OK]"
    elif R["et"] >= 0.004:
        klas = "0.004 <= et < 0.005  -->  Zona transisi  [PERLU TINJAUAN]"
    else:
        klas = "et < 0.004  -->  Tidak memenuhi syarat  [TIDAK OK]"

    s7 = dict(
        no="b7.", ref="SNI 2847:2019 Pasal 21.2.2",
        judul="Regangan tarik terjauh (et) - klasifikasi penampang",
        isi=(
            f"et = regangan lapis tarik paling bawah\n"
            f"   = {R['et']:.5f}\n"
            f"{klas}"
        ),
        ok=R["ok_et"],
    )

    s8 = dict(
        no="b8.", ref="SNI 2847:2019 Tabel 21.2.2",
        judul="Faktor reduksi kekuatan (Phi)",
        isi=f"{R['phi_cara']}\n  -->  Phi = {R['phi']:.4f}",
        ok=R["ok_et"],
    )

    teks_mom = (
        f"Mn = Cc x (d - a/2) + Sum[ Cs_i x (d - y_i) ]\n"
        f"     (referensi: titik berat tulangan tarik d = {R['d_tarik']:.2f} mm)\n\n"
        f"  Cc x (d - a/2) = {R['Cc']:,.0f} x ({R['d_tarik']:.2f} - {R['a']/2:.2f})\n"
        f"                 = {R['Cc']*(R['d_tarik']-R['a']/2):,.0f} N.mm\n"
    )
    for L in R["info_tekan"]:
        teks_mom += (f"  Cs Lapis-{L['no']} x (d - y_i) = "
                     f"{L['F']:+,.0f} x ({R['d_tarik']:.2f} - {L['y']:.1f}) "
                     f"= {L['F']*(R['d_tarik']-L['y']):+,.0f} N.mm\n")

    s9 = dict(
        no="b9.", ref="SNI 2847:2019 Pasal 22.3.2",
        judul="Momen nominal (Mn) dan momen rencana (Phi.Mn)",
        isi=(
            f"{teks_mom}\n"
            f"Mn     = {R['Mn']:.3f} kN.m\n"
            f"Phi.Mn = Phi x Mn = {R['phi']:.4f} x {R['Mn']:.3f} = "
            f"{R['phiMn']:.3f} kN.m"
        ),
        ok=True,
    )

    if R["ok_dc"]:
        ket_dc = "AMAN  --  Phi.Mn >= Mu"
    else:
        ket_dc = "TIDAK AMAN  --  Phi.Mn < Mu  (penampang perlu diperbesar / tulangan ditambah)"

    s10 = dict(
        no="b10.", ref="SNI 2847:2019 Pasal 9.5.1.1 - Mu <= Phi.Mn",
        judul="D/C Ratio (Demand-to-Capacity) - LENTUR",
        isi=(
            f"D/C = Mu / Phi.Mn\n"
            f"    = {Mu:.3f} / {R['phiMn']:.3f}\n"
            f"    = {R['DC']:.3f}\n\n"
            f"{ket_dc}"
        ),
        ok=R["ok_dc"],
    )

    return [s1, s2, s3, s4, s5, s6, s7, s8, s9, s10]


# ============================================================
# BUAT STEP-STEP LAPORAN GESER (b11 - b18)
# ============================================================
def buat_steps_geser(fc, fyt, b, d_aktual, Vu, ds, s_seng, n_kaki, G):
    s11 = dict(
        no="b11.", ref="SNI 2847:2019 Pasal 21.2.1",
        judul="Faktor reduksi geser (Phi_v)",
        isi=(
            f"Untuk evaluasi geser balok beton bertulang:\n"
            f"  -->  Phi_v = {G['Phi_v']:.2f}"
        ),
        ok=True,
    )

    s12 = dict(
        no="b12.", ref="SNI 2847:2019 Pasal 22.5.5.1",
        judul="Kapasitas geser beton (Vc)",
        isi=(
            f"Vc = 0.17 x lambda x sqrt(fc) x b x d        (Newton)\n"
            f"   = 0.17 x {G['lambda_']:.2f} x sqrt({fc:.1f}) x {b:.0f} x {d_aktual:.2f}\n"
            f"   = 0.17 x {G['lambda_']:.2f} x {math.sqrt(fc):.4f} x "
            f"{b:.0f} x {d_aktual:.2f}\n"
            f"   = {G['Vc_N']:,.0f} N\n"
            f"   = {G['Vc']:.2f} kN\n\n"
            f"Phi.Vc       = {G['Phi_v']:.2f} x {G['Vc']:.2f} = {G['phiVc']:.2f} kN\n"
            f"0.5 Phi.Vc   = {G['half_phiVc']:.2f} kN"
        ),
        ok=True,
    )

    s13 = dict(
        no="b13.", ref="SNI 2847:2019 Pasal 9.6.3 & 22.5.1.1",
        judul="Cek kebutuhan sengkang",
        isi=(
            f"Kategori berdasarkan Vu:\n"
            f"  Jika Vu <= 0.5 Phi.Vc            -->  tidak perlu sengkang\n"
            f"  Jika 0.5 Phi.Vc < Vu <= Phi.Vc   -->  pasang sengkang minimum\n"
            f"  Jika Vu > Phi.Vc                 -->  hitung Vs\n\n"
            f"{G['ket_klas']}"
        ),
        ok=True,
    )

    s14 = dict(
        no="b14.", ref="SNI 2847:2019 Pasal 22.5.1.2",
        judul="Vs perlu & cek Vs maksimum",
        isi=(
            f"Vs_perlu = Vu/Phi_v - Vc\n"
            f"         = {Vu:.2f}/{G['Phi_v']:.2f} - {G['Vc']:.2f}\n"
            f"         = {G['Vs_perlu']:.2f} kN\n\n"
            f"Vs_max   = 0.66 x sqrt(fc) x b x d / 1000\n"
            f"         = 0.66 x {math.sqrt(fc):.4f} x {b:.0f} x {d_aktual:.2f} / 1000\n"
            f"         = {G['Vs_max']:.2f} kN\n\n"
            f"Kontrol Vs_perlu ({G['Vs_perlu']:.2f}) <= Vs_max ({G['Vs_max']:.2f})  "
            f"-->  {'[OK]' if G['ok_Vs_max'] else '[TIDAK OK] - Penampang harus diperbesar'}"
        ),
        ok=G["ok_Vs_max"],
    )

    s15 = dict(
        no="b15.", ref="SNI 2847:2019 Pasal 22.5.10.5.3",
        judul="Luas sengkang terpasang (Av) per spasi",
        isi=(
            f"Av_pasang = n_kaki x (pi/4 x ds^2)\n"
            f"          = {n_kaki} x (pi/4 x {ds:.0f}^2)\n"
            f"          = {n_kaki} x {math.pi/4*ds**2:.2f}\n"
            f"          = {G['Av_pasang']:.2f} mm2\n\n"
            f"  --> Sengkang terpasang : {n_kaki} kaki D{int(ds)} - "
            f"jarak {s_seng:.0f} mm"
        ),
        ok=True,
    )

    s16 = dict(
        no="b16.", ref="SNI 2847:2019 Pasal 9.6.3.3 & 9.7.6.2.2",
        judul="Av minimum & spasi maksimum sengkang",
        isi=(
            f"-- Av minimum --\n"
            f"Av_min/s = max( 0.062 sqrt(fc) b/fyt , 0.35 b/fyt )\n"
            f"         = max( 0.062 x {math.sqrt(fc):.4f} x {b:.0f}/{fyt:.0f} ,\n"
            f"                0.35 x {b:.0f}/{fyt:.0f} )\n"
            f"         = max( {G['AvS_minA']:.4f} , {G['AvS_minB']:.4f} ) "
            f"= {G['AvS_min']:.4f} mm2/mm\n"
            f"Av_min   = (Av_min/s) x s_pasang = {G['AvS_min']:.4f} x {s_seng:.0f}\n"
            f"         = {G['Av_min']:.2f} mm2\n\n"
            f"Kontrol Av_pasang ({G['Av_pasang']:.2f}) >= Av_min ({G['Av_min']:.2f})  "
            f"-->  {'[OK]' if G['ok_Av'] else '[TIDAK OK]'}\n\n"
            f"-- Spasi maksimum --\n"
            f"{G['rumus_smax']}\n\n"
            f"Kontrol s_pasang ({s_seng:.0f} mm) <= s_max ({G['s_max']:.1f} mm)  "
            f"-->  {'[OK]' if G['ok_spasi'] else '[TIDAK OK]'}"
        ),
        ok=(G["ok_Av"] and G["ok_spasi"]),
    )

    s17 = dict(
        no="b17.", ref="SNI 2847:2019 Pasal 22.5.10.5.3",
        judul="Vs aktual & kapasitas geser rencana (Phi.Vn)",
        isi=(
            f"Vs_aktual = Av_pasang x fyt x d / s_pasang / 1000\n"
            f"          = {G['Av_pasang']:.2f} x {fyt:.0f} x {d_aktual:.2f} / "
            f"{s_seng:.0f} / 1000\n"
            f"          = {G['Vs_aktual']:.2f} kN\n"
            + (f"  (Vs_aktual > Vs_max -> Vs efektif dibatasi = "
               f"{G['Vs_efektif']:.2f} kN)\n"
               if G['Vs_aktual'] > G['Vs_max'] else "") +
            f"\nVn        = Vc + Vs_efektif\n"
            f"          = {G['Vc']:.2f} + {G['Vs_efektif']:.2f}\n"
            f"          = {G['Vn_efektif']:.2f} kN\n\n"
            f"Phi.Vn    = Phi_v x Vn = {G['Phi_v']:.2f} x {G['Vn_efektif']:.2f}\n"
            f"          = {G['PhiVn_efektif']:.2f} kN"
        ),
        ok=True,
    )

    if G["ok_dc"]:
        ket_dc = "AMAN  --  Phi.Vn >= Vu"
    else:
        ket_dc = ("TIDAK AMAN  --  Phi.Vn < Vu  "
                  "(perbesar penampang / rapatkan sengkang / tambah kaki)")

    s18 = dict(
        no="b18.", ref="SNI 2847:2019 Pasal 9.5.1.1 - Vu <= Phi.Vn",
        judul="D/C Ratio (Demand-to-Capacity) - GESER",
        isi=(
            f"D/C = Vu / Phi.Vn\n"
            f"    = {Vu:.2f} / {G['PhiVn_efektif']:.2f}\n"
            f"    = {G['DC_geser']:.3f}\n\n"
            f"{ket_dc}"
        ),
        ok=G["ok_dc"],
    )

    return [s11, s12, s13, s14, s15, s16, s17, s18]


# ============================================================
# BUAT STEP-STEP LAPORAN TORSI (b19 - b28) -- TAHAP 2
# ============================================================
def buat_steps_torsi(fc, fyt, fy_long, b, h, cc_sel, ds, s_seng,
                     Tu, Vu, d_aktual, T):
    """Susun langkah perhitungan torsi b19 s/d b28."""

    if T["abaikan_torsi"]:
        s_abaikan = dict(
            no="b19-b28.", ref="SNI 2847:2019 Pasal 22.7.4",
            judul="EVALUASI TORSI -- DILEWATI",
            isi=(
                f"Tu = {Tu:.3f} kN.m\n"
                f"Phi_t x Tth = {T['Phi_t']:.2f} x {T['Tth']:.3f} = {T['phi_Tth']:.3f} kN.m\n\n"
                f"Karena Tu ({Tu:.3f}) <= Phi_t x Tth ({T['phi_Tth']:.3f}) kN.m\n"
                f"  -->  [DIABAIKAN] Efek torsi dapat diabaikan sesuai SNI 2847:2019 Pasal 22.7.4"
            ),
            ok=True,
        )
        return [s_abaikan]

    # Pre-compute strings yang butuh conditional agar tidak ada nested f-string
    _ok_dim    = "[OK] Dimensi penampang mencukupi" if T["ok_dimensi"] else "[TIDAK OK] Penampang harus diperbesar!"
    _ok_seng   = "[OK]" if T["ok_sengkang_gabungan"] else ("[TIDAK OK] Perlu s <= " + str(round(T["s_perlu_gabungan"])) + " mm")
    _ok_avt    = "[OK]" if T["ok_Avt_min"] else "[TIDAK OK]"
    _ok_spasi  = "[OK]" if T["ok_spasi_torsi"] else "[TIDAK OK]"
    _ok_dc_t   = "AMAN  --  Phi.Tn >= Tu_desain" if T["ok_DC_torsi"] else "TIDAK AMAN  --  Phi.Tn < Tu_desain (perkecil spasi / tambah sengkang)"
    _ok_dim_r  = "[OK]" if T["ok_dimensi"] else "[TIDAK OK]"
    _ok_seng_r = "[OK]" if T["ok_sengkang_gabungan"] else "[TIDAK OK]"
    _ok_avt_r  = "[OK]" if T["ok_Avt_min"] else "[TIDAK OK]"
    _ok_spasi_r= "[OK]" if T["ok_spasi_torsi"] else "[TIDAK OK]"
    _ok_dc_r   = "[OK]" if T["ok_DC_torsi"] else "[TIDAK OK]"
    _abaik_str = "[DIABAIKAN]" if T["abaikan_torsi"] else "[HITUNG TORSI] Tu > Phi_t x Tth"

    # -- b19: Properti penampang --
    s19 = dict(
        no="b19.", ref="SNI 2847:2019 Pasal 22.7",
        judul="Properti penampang torsi",
        isi=(
            f"Acp = b x h = {b:.0f} x {h:.0f} = {T['Acp']:,.0f} mm2\n"
            f"pcp = 2 x (b + h) = 2 x ({b:.0f} + {h:.0f}) = {T['pcp']:.0f} mm\n\n"
            f"Dimensi sengkang tertutup (c to c):\n"
            f"  x1 = b - 2 cc - ds = {b:.0f} - 2x{cc_sel:.0f} - {ds:.0f} "
            f"= {T['x1']:.1f} mm\n"
            f"  y1 = h - 2 cc - ds = {h:.0f} - 2x{cc_sel:.0f} - {ds:.0f} "
            f"= {T['y1']:.1f} mm\n\n"
            f"Aoh = x1 x y1 = {T['x1']:.1f} x {T['y1']:.1f} = {T['Aoh']:,.1f} mm2\n"
            f"ph  = 2 x (x1 + y1) = 2 x ({T['x1']:.1f} + {T['y1']:.1f}) "
            f"= {T['ph']:.1f} mm\n"
            f"Ao  = 0.85 x Aoh = 0.85 x {T['Aoh']:,.1f} = {T['Ao']:,.1f} mm2"
        ),
        ok=True,
    )

    # -- b20: Faktor reduksi --
    s20 = dict(
        no="b20.", ref="SNI 2847:2019 Pasal 21.2.1",
        judul="Faktor reduksi torsi (Phi_t)",
        isi=(
            f"Untuk evaluasi torsi balok beton bertulang:\n"
            f"  -->  Phi_t = {T['Phi_t']:.2f}"
        ),
        ok=True,
    )

    # -- b21: Batas ambang torsi --
    s21 = dict(
        no="b21.", ref="SNI 2847:2019 Pasal 22.7.4",
        judul="Batas ambang torsi (Tth)",
        isi=(
            f"Tth = 0.083 x lambda x sqrt(fc) x (Acp^2 / pcp)    (N.mm)\n"
            f"    = 0.083 x {T['lambda_']:.2f} x sqrt({fc:.1f}) x "
            f"({T['Acp']:,.0f}^2 / {T['pcp']:.0f})\n"
            f"    = 0.083 x {T['lambda_']:.2f} x {math.sqrt(fc):.4f} x "
            f"{T['Acp']**2/T['pcp']:,.0f}\n"
            f"    = {T['Tth_Nmm']:,.0f} N.mm = {T['Tth']:.4f} kN.m\n\n"
            f"Phi_t x Tth = {T['Phi_t']:.2f} x {T['Tth']:.4f} = {T['phi_Tth']:.4f} kN.m\n\n"
            f"Tu = {Tu:.3f} kN.m\n"
            f"Kontrol Tu ({Tu:.3f}) vs Phi_t x Tth ({T['phi_Tth']:.4f})\n"
            f"  -->  {_abaik_str}"
        ),
        ok=True,
    )

    # -- b22: Klasifikasi & Tu desain --
    s22 = dict(
        no="b22.", ref="SNI 2847:2019 Pasal 22.7.3",
        judul="Klasifikasi torsi & Tu desain",
        isi=(
            f"Tipe torsi : {T['tipe_torsi']}\n"
            f"Tcr = 0.33 x lambda x sqrt(fc) x (Acp^2 / pcp)\n"
            f"    = 0.33 x {T['lambda_']:.2f} x {math.sqrt(fc):.4f} x "
            f"{T['Acp']**2/T['pcp']:,.0f} / 1e6\n"
            f"    = {T['Tcr']:.4f} kN.m\n"
            f"Phi_t x Tcr = {T['Phi_t']:.2f} x {T['Tcr']:.4f} = {T['phi_Tcr']:.4f} kN.m\n\n"
            f"{T['catatan_tu']}\n"
            f"  -->  Tu_desain = {T['Tu_desain']:.4f} kN.m"
        ),
        ok=True,
    )

    # -- b23: Cek dimensi penampang --
    s23 = dict(
        no="b23.", ref="SNI 2847:2019 Pasal 22.7.7.1",
        judul="Cek dimensi penampang (kombinasi geser + torsi)",
        isi=(
            f"Syarat: sqrt( (Vu/(b*d))^2 + (Tu*ph/(1.7*Aoh^2))^2 ) <= Phi_v*(Vc/(b*d) + 0.66*sqrt(fc))\n\n"
            f"Term geser  = Vu / (b x d) = {T['Vu']*1000:.0f} / ({b:.0f} x {d_aktual:.2f})\n"
            f"            = {T['term_geser']:.4f} MPa\n\n"
            f"Term torsi  = Tu_desain x ph / (1.7 x Aoh^2)\n"
            f"            = {T['Tu_desain_Nmm']:,.0f} x {T['ph']:.1f} / (1.7 x {T['Aoh']:,.1f}^2)\n"
            f"            = {T['term_torsi']:.4f} MPa\n\n"
            f"LHS = sqrt( {T['term_geser']:.4f}^2 + {T['term_torsi']:.4f}^2 )\n"
            f"    = {T['lhs_dim']:.4f} MPa\n\n"
            f"RHS = Phi_v x (Vc/(b*d) + 0.66 sqrt(fc))\n"
            f"    = 0.75 x ({T['Vc']*1000:.0f}/({b:.0f}x{d_aktual:.2f}) + 0.66x{math.sqrt(fc):.4f})\n"
            f"    = {T['rhs_dim']:.4f} MPa\n\n"
            f"D/C dimensi = {T['DC_dim']:.3f}\n"
            f"  -->  {_ok_dim}"
        ),
        ok=T["ok_dimensi"],
    )

    # -- b24: At/s untuk torsi --
    s24 = dict(
        no="b24.", ref="SNI 2847:2019 Pasal 22.7.6.1",
        judul="Tulangan sengkang tertutup untuk torsi (At/s)",
        isi=(
            f"At/s = Tu_desain / (Phi_t x 2 x Ao x fyt x cot(theta))\n"
            f"     (theta = 45 deg, cot(theta) = 1.0)\n\n"
            f"At/s = {T['Tu_desain_Nmm']:,.0f} / ({T['Phi_t']:.2f} x 2 x "
            f"{T['Ao']:,.1f} x {fyt:.0f} x {T['cot_theta']:.1f})\n"
            f"     = {T['At_per_s']:.6f} mm2/mm"
        ),
        ok=True,
    )

    # -- b25: Sengkang gabungan --
    s25 = dict(
        no="b25.", ref="SNI 2847:2019 Pasal 9.6.4.2",
        judul="Sengkang gabungan geser + torsi (Av + 2At)",
        isi=(
            f"Av/s terpasang (geser) = Av_pasang / s\n"
            f"                      = {T['Avt_pasang']:.2f} / {s_seng:.0f}\n"
            f"                      = {T['Av_per_s_pasang']:.4f} mm2/mm\n\n"
            f"At/s perlu (torsi)    = {T['At_per_s']:.6f} mm2/mm\n"
            f"2 x At/s              = {2*T['At_per_s']:.6f} mm2/mm\n\n"
            f"(Av + 2At)/s perlu    = {T['Av_per_s_pasang']:.4f} + {2*T['At_per_s']:.6f}\n"
            f"                      = {T['Avt_per_s_perlu']:.6f} mm2/mm\n\n"
            f"(Av + 2At) perlu pada s = {s_seng:.0f} mm\n"
            f"  = {T['Avt_per_s_perlu']:.6f} x {s_seng:.0f} = {T['Avt_perlu_abs']:.2f} mm2\n\n"
            f"Av_pasang (ada)       = {T['Avt_pasang']:.2f} mm2\n"
            f"  -->  {_ok_seng}\n\n"
            f"-- Av+2At minimum (SNI 9.6.4.2) --\n"
            f"(Av+2At)_min/s = max(0.062 sqrt(fc) b/fyt , 0.35 b/fyt)\n"
            f"               = max({T['AvtS_minA']:.4f} , {T['AvtS_minB']:.4f}) "
            f"= {T['AvtS_min']:.4f} mm2/mm\n"
            f"(Av+2At)_min   = {T['AvtS_min']:.4f} x {s_seng:.0f} = {T['Avt_min']:.2f} mm2\n"
            f"Kontrol Av_pasang ({T['Avt_pasang']:.2f}) >= (Av+2At)_min ({T['Avt_min']:.2f})  "
            f"-->  {_ok_avt}"
        ),
        ok=(T["ok_sengkang_gabungan"] and T["ok_Avt_min"]),
    )

    # -- b26: Spasi maksimum & kontrol --
    s26 = dict(
        no="b26.", ref="SNI 2847:2019 Pasal 9.7.6.3.3",
        judul="Spasi maksimum sengkang torsi",
        isi=(
            f"s_max_torsi = min(ph/8 , 300 mm)\n"
            f"            = min({T['ph']:.1f}/8 , 300)\n"
            f"            = min({T['ph']/8:.1f} , 300)\n"
            f"            = {T['s_max_torsi']:.1f} mm\n\n"
            f"s_pasang = {s_seng:.0f} mm\n"
            f"Kontrol s_pasang ({s_seng:.0f}) <= s_max_torsi ({T['s_max_torsi']:.1f})  "
            f"-->  {_ok_spasi}"
        ),
        ok=T["ok_spasi_torsi"],
    )

    # -- b27: Tulangan longitudinal torsi --
    s27 = dict(
        no="b27.", ref="SNI 2847:2019 Pasal 22.7.6.1 & 9.6.4.3",
        judul="Tulangan longitudinal torsi (Al)",
        isi=(
            f"Al = (At/s) x ph x (fyt/fy_long) x cot^2(theta)\n"
            f"   = {T['At_per_s']:.6f} x {T['ph']:.1f} x ({fyt:.0f}/{T['fy_long']:.0f}) x 1.0\n"
            f"   = {T['Al']:.2f} mm2\n\n"
            f"Al_min (SNI 9.6.4.3):\n"
            f"  Al_minA = 0.42 sqrt(fc) Acp/fy_long - (At/s) ph (fyt/fy_long)\n"
            f"          = 0.42 x {math.sqrt(fc):.4f} x {T['Acp']:,.0f}/{T['fy_long']:.0f} "
            f"- {T['At_per_s']:.6f} x {T['ph']:.1f} x ({fyt:.0f}/{T['fy_long']:.0f})\n"
            f"          = {T['Al_minA']:.2f} mm2\n"
            f"  Al_minB = 0.42 sqrt(fc) Acp/fy_long - (0.175b/fyt) ph (fyt/fy_long)\n"
            f"          = 0.42 x {math.sqrt(fc):.4f} x {T['Acp']:,.0f}/{T['fy_long']:.0f} "
            f"- (0.175x{b:.0f}/{fyt:.0f}) x {T['ph']:.1f} x ({fyt:.0f}/{T['fy_long']:.0f})\n"
            f"          = {T['Al_minB']:.2f} mm2\n"
            f"  Al_min  = max(Al_minA, Al_minB, 0) = {T['Al_min']:.2f} mm2\n\n"
            f"Al_pakai  = max(Al , Al_min) = max({T['Al']:.2f} , {T['Al_min']:.2f})\n"
            f"          = {T['Al_pakai']:.2f} mm2\n\n"
            f"Diameter longitudinal torsi dipilih: D{int(T['db_long_torsi'])}\n"
            f"  Ab_long = pi/4 x {T['db_long_torsi']:.0f}^2 = {T['Ab_long']:.2f} mm2\n"
            f"  n_batang = ceil(Al_pakai / Ab_long) = ceil({T['Al_pakai']:.2f} / {T['Ab_long']:.2f})\n"
            f"           = {T['n_batang']} batang D{int(T['db_long_torsi'])}\n\n"
            f"Rekomendasi distribusi: minimal 1 batang di tiap sudut sengkang tertutup,\n"
            f"  sisanya distribusikan di sisi panjang dengan spasi max 300 mm."
        ),
        ok=True,
    )

    # -- b28: DC Torsi & Kesimpulan --
    s28 = dict(
        no="b28.", ref="SNI 2847:2019 Pasal 9.5.1.2",
        judul="D/C Ratio (Demand-to-Capacity) - TORSI & Kesimpulan",
        isi=(
            f"Kapasitas Phi.Tn (dari sengkang terpasang):\n"
            f"  At/s_cap = Av_pasang / s = {T['Avt_pasang']:.2f} / {s_seng:.0f} "
            f"= {T['At_per_s_cap']:.4f} mm2/mm\n"
            f"  Tn_cap = 2 x Ao x At/s_cap x fyt x cot(theta) / 1e6\n"
            f"         = 2 x {T['Ao']:,.1f} x {T['At_per_s_cap']:.4f} x {fyt:.0f} x 1.0 / 1e6\n"
            f"         = {T['Tn_cap']:.4f} kN.m\n"
            f"  Phi.Tn = {T['Phi_t']:.2f} x {T['Tn_cap']:.4f} = {T['PhiTn_cap']:.4f} kN.m\n\n"
            f"D/C = Tu_desain / Phi.Tn\n"
            f"    = {T['Tu_desain']:.4f} / {T['PhiTn_cap']:.4f}\n"
            f"    = {T['DC_torsi']:.3f}\n\n"
            f"{_ok_dc_t}\n\n"
            f"-- Rangkuman kontrol torsi --\n"
            f"Dimensi penampang    : {_ok_dim_r}\n"
            f"Sengkang gabungan    : {_ok_seng_r}\n"
            f"(Av+2At) minimum     : {_ok_avt_r}\n"
            f"Spasi max torsi      : {_ok_spasi_r}\n"
            f"D/C torsi <= 1.0     : {_ok_dc_r}"
        ),
        ok=T["ok_torsi_total"],
    )

    return [s19, s20, s21, s22, s23, s24, s25, s26, s27, s28]




# ============================================================
# VISUALISASI PENAMPANG (matplotlib) -- UPDATED dengan Aoh
# ============================================================
def gambar_penampang(b, h, cc_sel, ds, lapis_tarik, lapis_tekan,
                     c=None, torsi_data=None):
    """Gambar cross-section balok proporsional. torsi_data = dict T dari hitung_torsi."""
    fig, ax = plt.subplots(figsize=(5.5, 6.5))

    # Penampang beton
    ax.add_patch(patches.Rectangle((0, 0), b, h, fill=True,
                                   facecolor="#e8e8e8", edgecolor="black",
                                   linewidth=1.8))

    # Sengkang (selimut)
    sx, sy, sw, sh = cc_sel, cc_sel, b - 2*cc_sel, h - 2*cc_sel
    ax.add_patch(patches.Rectangle((sx, sy), sw, sh, fill=False,
                                   edgecolor="#1a3c5e", linewidth=1.2,
                                   linestyle="--"))

    # Area Aoh (jika ada data torsi)
    if torsi_data is not None and not torsi_data.get("abaikan_torsi", True):
        x1 = torsi_data["x1"]
        y1_t = torsi_data["y1"]
        # Aoh dimulai dari tengah sengkang
        xoh = cc_sel + ds / 2.0
        yoh = cc_sel + ds / 2.0
        ax.add_patch(patches.Rectangle((xoh, yoh), x1, y1_t,
                                       fill=True,
                                       facecolor="#ff9800", alpha=0.12,
                                       edgecolor="#e65100", linewidth=1.0,
                                       linestyle=":"))
        ax.annotate(f"Aoh = {x1:.0f} x {y1_t:.0f} = {torsi_data['Aoh']:,.0f} mm2",
                    xy=(b / 2, yoh + y1_t / 2), fontsize=7.5,
                    ha="center", va="center", color="#e65100",
                    bbox=dict(boxstyle="round,pad=0.2", facecolor="white", alpha=0.7))

        # Plot tulangan longitudinal torsi di sudut
        if torsi_data.get("n_batang", 0) > 0:
            db_lt = torsi_data["db_long_torsi"]
            r_lt  = db_lt / 2.0
            corners = [
                (xoh, yoh),
                (xoh + x1, yoh),
                (xoh, yoh + y1_t),
                (xoh + x1, yoh + y1_t),
            ]
            for cx_, cy_ in corners:
                ax.add_patch(patches.Circle((cx_, cy_), r_lt,
                                            facecolor="#9e9e9e",
                                            edgecolor="black", linewidth=0.6))

    # Tulangan tarik
    for L in lapis_tarik:
        y_plot = h - L["y"]
        n = L["n"]
        x_left  = cc_sel + ds + L["db"] / 2
        x_right = b - cc_sel - ds - L["db"] / 2
        if n > 1:
            xs = [x_left + i * (x_right - x_left) / (n - 1) for i in range(n)]
        else:
            xs = [(x_left + x_right) / 2]
        for x in xs:
            ax.add_patch(patches.Circle((x, y_plot), L["db"] / 2,
                                        facecolor="#1a3c5e",
                                        edgecolor="black", linewidth=0.6))
        ax.annotate(f"{n}D{int(L['db'])}",
                    xy=(b + 8, y_plot), fontsize=8,
                    va="center", color="#1a3c5e")

    # Tulangan tekan
    for L in lapis_tekan:
        y_plot = h - L["y"]
        n = L["n"]
        x_left  = cc_sel + ds + L["db"] / 2
        x_right = b - cc_sel - ds - L["db"] / 2
        if n > 1:
            xs = [x_left + i * (x_right - x_left) / (n - 1) for i in range(n)]
        else:
            xs = [(x_left + x_right) / 2]
        for x in xs:
            ax.add_patch(patches.Circle((x, y_plot), L["db"] / 2,
                                        facecolor="#c62828",
                                        edgecolor="black", linewidth=0.6))
        ax.annotate(f"{n}D{int(L['db'])}",
                    xy=(b + 8, y_plot), fontsize=8,
                    va="center", color="#c62828")

    # Garis netral
    if c is not None:
        y_c_plot = h - c
        ax.axhline(y=y_c_plot, color="#f9a825", linewidth=1.4,
                   linestyle="-.", alpha=0.85)
        ax.annotate(f"  garis netral c = {c:.1f} mm",
                    xy=(0, y_c_plot), fontsize=8,
                    va="bottom", color="#7a5800")

    ax.annotate(f"b = {int(b)} mm", xy=(b/2, -h*0.05),
                fontsize=9, ha="center", color="black")
    ax.annotate(f"h = {int(h)} mm", xy=(-b*0.10, h/2),
                fontsize=9, ha="center", color="black", rotation=90)

    # Legend sederhana jika ada torsi
    if torsi_data is not None and not torsi_data.get("abaikan_torsi", True):
        legend_elems = [
            patches.Patch(facecolor="#ff9800", alpha=0.3, edgecolor="#e65100",
                          linestyle=":", label="Area Aoh (torsi)"),
            patches.Patch(facecolor="#9e9e9e", edgecolor="black",
                          label=f"Tul. long. torsi D{int(torsi_data['db_long_torsi'])}"),
        ]
        ax.legend(handles=legend_elems, loc="upper right", fontsize=7,
                  framealpha=0.8)

    pad = max(b, h) * 0.18
    ax.set_xlim(-pad, b + pad * 1.6)
    ax.set_ylim(-pad, h + pad * 0.6)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.set_title("Penampang Balok (skala proporsional)",
                 fontsize=10, color="#1a3c5e", pad=10)
    fig.tight_layout()
    return fig


def fig_to_png_bytes(fig):
    """Convert matplotlib figure ke PNG bytes."""
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    buf.seek(0)
    return buf


# ============================================================
# GENERATOR WORD (.docx) -- UPDATED dengan Torsi
# ============================================================
def create_word_balok(fc, fy, b, h, cc_sel, ds, Mu, lapis_tarik, lapis_tekan,
                      R, steps, nama_proyek, png_buf,
                      G=None, steps_geser=None, geser_inputs=None,
                      Tor=None, steps_torsi=None, torsi_inputs=None,
                      timestamp_str=None):
    doc = Document()

    # -- Pengaturan margin --
    for sect in doc.sections:
        sect.top_margin    = Cm(2.5)
        sect.bottom_margin = Cm(2.5)
        sect.left_margin   = Cm(3.0)
        sect.right_margin  = Cm(2.5)

    # -- Gaya helper --
    def par(teks="", bold=False, italic=False, size=11, indent=0.0,
            align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(space_after)
        p.alignment = align
        if indent:
            p.paragraph_format.left_indent = Cm(indent)
        r = p.add_run(teks)
        r.bold = bold; r.italic = italic
        r.font.size = Pt(size); r.font.name = "Calibri"
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    def judul_utama(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(teks); r.bold = True; r.font.size = Pt(14)
        r.font.name = "Calibri"; r.font.color.rgb = RGBColor(26, 60, 94)

    def subjudul(teks):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(teks); r.bold = True; r.font.size = Pt(11.5)
        r.font.name = "Calibri"; r.font.color.rgb = RGBColor(26, 60, 94)
        pf = p.paragraph_format
        pf.border_bottom = None

    def garis():
        par("", space_after=2)
        p = doc.add_paragraph()
        p.paragraph_format.border_bottom = None
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(4)

    def tulis_steps(steps_list):
        for s in steps_list:
            ok_col = (0x1B, 0x5E, 0x20) if s["ok"] else (0xB7, 0x1C, 0x1C)
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(1)
            r1 = p.add_run(f"{s['no']} {s['judul']}")
            r1.bold = True; r1.font.size = Pt(10.5); r1.font.name = "Calibri"
            r1.font.color.rgb = RGBColor(26, 60, 94)

            p2 = doc.add_paragraph()
            p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(0)
            p2.paragraph_format.left_indent = Cm(0.4)
            r2 = p2.add_run(f"[{s['ref']}]")
            r2.italic = True; r2.font.size = Pt(9); r2.font.name = "Calibri"
            r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

            for baris in s["isi"].split("\n"):
                is_result = (baris.strip().startswith("-->") or
                             "[OK]" in baris or "[TIDAK OK]" in baris or
                             "AMAN" in baris or "[DIABAIKAN]" in baris)
                p3 = doc.add_paragraph()
                p3.paragraph_format.space_before = Pt(0)
                p3.paragraph_format.space_after  = Pt(0)
                p3.paragraph_format.left_indent  = Cm(0.5)
                r3 = p3.add_run(baris if baris.strip() else "")
                r3.font.name = "Courier New"; r3.font.size = Pt(9.5)
                if is_result:
                    r3.bold = True
                    r3.font.color.rgb = RGBColor(*ok_col)
            par(space_after=4)

    if timestamp_str is None:
        timestamp_str = datetime.datetime.now().strftime('%d %B %Y  %H:%M')

    # Header
    judul_utama("LAPORAN PERHITUNGAN STRUKTUR")
    judul_utama("Evaluasi Kapasitas Lentur, Geser & Torsi Balok Beton Bertulang")
    par("(Tulangan Rangkap - Strain Compatibility)",
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55))
    par("Referensi: SNI 2847:2019 (ACI 318-14)",
        italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x55, 0x55, 0x55))
    par(f"Proyek: {nama_proyek}   |   Tanggal: {timestamp_str}",
        size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x33, 0x33, 0x33), space_after=8)
    par(space_after=4)

    # =========== A. DATA INPUT ===========
    subjudul("A.  DATA INPUT PENAMPANG")
    par(f"Proyek    : {nama_proyek}", size=10, indent=0.5, space_after=2)
    par(f"Tanggal   : {timestamp_str}", size=10, indent=0.5, space_after=6)

    # Material
    par("Material:", bold=True, size=10, indent=0.5, space_after=2)
    for simb, nilai in [
        ("f'c", f"{fc:.1f} MPa"), ("fy", f"{fy:.0f} MPa"),
        ("fyt", f"{geser_inputs['fyt']:.0f} MPa" if geser_inputs else "-"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        r = p.add_run(f"{simb:<6} = {nilai}")
        r.font.size = Pt(10); r.font.name = "Courier New"

    par("Geometri:", bold=True, size=10, indent=0.5, space_after=2)
    for simb, nilai in [
        ("b",  f"{b:.0f} mm"), ("h",  f"{h:.0f} mm"),
        ("cc", f"{cc_sel:.0f} mm"), ("ds", f"{ds:.0f} mm"),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        r = p.add_run(f"{simb:<6} = {nilai}")
        r.font.size = Pt(10); r.font.name = "Courier New"

    par("Gaya Dalam:", bold=True, size=10, indent=0.5, space_after=2)
    gd_items = [("Mu", f"{R['Mu']:.2f} kN.m")]
    if G is not None: gd_items.append(("Vu", f"{G['Vu']:.2f} kN"))
    if Tor is not None: gd_items.append(("Tu", f"{Tor['Tu']:.2f} kN.m"))
    for simb, nilai in gd_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        r = p.add_run(f"{simb:<6} = {nilai}")
        r.font.size = Pt(10); r.font.name = "Courier New"

    par("Tulangan TARIK:", bold=True, size=10, indent=0.5, space_after=2)
    for L in lapis_tarik:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(1.0)
        r = p.add_run(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                      f"As = {L['As']:.1f} mm2   y = {L['y']:.1f} mm")
        r.font.size = Pt(10); r.font.name = "Courier New"
    par(f"  As-total = {R['As_tarik']:.1f} mm2  |  d-aktual = {R['d_tarik']:.2f} mm",
        bold=True, size=10, indent=1.0, space_after=4)

    par("Tulangan TEKAN:", bold=True, size=10, indent=0.5, space_after=2)
    if lapis_tekan:
        for L in lapis_tekan:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent  = Cm(1.0)
            r = p.add_run(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                          f"As' = {L['As']:.1f} mm2   y = {L['y']:.1f} mm")
            r.font.size = Pt(10); r.font.name = "Courier New"
        par(f"  As'-total = {R['As_tekan']:.1f} mm2  |  d'-aktual = {R['d_tekan']:.2f} mm",
            bold=True, size=10, indent=1.0, space_after=4)
    else:
        par("  (Tidak ada tulangan tekan)", size=10, indent=1.0,
            color=(0x55, 0x55, 0x55), space_after=4)

    par(space_after=4)
    par("Visualisasi Penampang :", bold=True, size=10, indent=0.5, space_after=4)
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(png_buf, width=Cm(9))
    par(space_after=6)

    # =========== B. ANALISA LENTUR ===========
    subjudul("B.  ANALISA PERHITUNGAN - LENTUR")
    par("Urutan perhitungan lentur mengacu pada SNI 2847:2019.",
        size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
    tulis_steps(steps)
    par(space_after=6)

    # =========== C. RANGKUMAN ===========
    subjudul("C.  RANGKUMAN HASIL")
    rangkuman = [
        ("Beta-1",  f"{R['beta1']:.4f}",        ""),
        ("c",       f"{R['c']:.2f} mm",          "Kedalaman sumbu netral"),
        ("a",       f"{R['a']:.2f} mm",          "Kedalaman blok tegangan"),
        ("Cc",      f"{R['Cc']/1000:.2f} kN",    "Gaya tekan beton"),
        ("Cs",      f"{R['Cs']/1000:.2f} kN",    "Gaya tekan baja total"),
        ("T",       f"{R['T']/1000:.2f} kN",     "Gaya tarik baja total"),
        ("et",      f"{R['et']:.5f}",            "Regangan tarik terjauh"),
        ("Phi",     f"{R['phi']:.4f}",           "Faktor reduksi lentur"),
        ("Rho",     f"{R['rho']*100:.4f}%",      "Rasio tulangan tarik"),
        ("Rho-min", f"{R['rho_min']*100:.4f}%",  "Batas minimum"),
        ("Rho-max", f"{R['rho_max']*100:.4f}%",  "Batas maksimum"),
        ("Mn",      f"{R['Mn']:.3f} kN.m",       "Momen nominal"),
        ("Phi.Mn",  f"{R['phiMn']:.3f} kN.m",    "Momen rencana kapasitas"),
        ("Mu",      f"{R['Mu']:.3f} kN.m",       "Momen rencana ultimit"),
        ("D/C-Lentur", f"{R['DC']:.3f}",         "Demand-to-Capacity Lentur"),
    ]
    if G is not None:
        rangkuman.extend([
            ("Phi_v",      f"{G['Phi_v']:.2f}",             "Faktor reduksi geser"),
            ("Vc",         f"{G['Vc']:.2f} kN",             "Kapasitas geser beton"),
            ("Vs",         f"{G['Vs_efektif']:.2f} kN",     "Kapasitas geser sengkang"),
            ("Vn",         f"{G['Vn_efektif']:.2f} kN",     "Kapasitas geser nominal"),
            ("Phi.Vn",     f"{G['PhiVn_efektif']:.2f} kN",  "Kapasitas geser rencana"),
            ("Vu",         f"{G['Vu']:.2f} kN",             "Geser ultimit"),
            ("Av_pasang",  f"{G['Av_pasang']:.1f} mm2",     "Av sengkang terpasang"),
            ("Av_min",     f"{G['Av_min']:.1f} mm2",        "Av sengkang minimum"),
            ("D/C-Geser",  f"{G['DC_geser']:.3f}",          "Demand-to-Capacity Geser"),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        rangkuman.extend([
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",        "Batas ambang torsi"),
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",         "Torsi ultimit"),
            ("Tu_desain",  f"{Tor['Tu_desain']:.4f} kN.m",  "Tu desain (setelah reduksi)"),
            ("Phi.Tn",     f"{Tor['PhiTn_cap']:.4f} kN.m",  "Kapasitas torsi rencana"),
            ("D/C-Torsi",  f"{Tor['DC_torsi']:.3f}",        "Demand-to-Capacity Torsi"),
            ("At/s",       f"{Tor['At_per_s']:.6f} mm2/mm", "Kebutuhan sengkang torsi"),
            ("Al_pakai",   f"{Tor['Al_pakai']:.2f} mm2",    "Tul. longitudinal torsi"),
            ("n_batang",   f"{Tor['n_batang']} D{int(Tor['db_long_torsi'])}", "Batang longitudinal torsi"),
            ("s_max_torsi",f"{Tor['s_max_torsi']:.1f} mm",  "Spasi maks sengkang torsi"),
        ])
    elif Tor is not None and Tor["abaikan_torsi"]:
        rangkuman.extend([
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",         "Torsi ultimit"),
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",        "Batas ambang torsi"),
            ("D/C-Torsi",  "DIABAIKAN",                      "Tu < Phi_t x Tth"),
        ])

    for simb, nilai, ket in rangkuman:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        r1 = p.add_run(f"{simb:<14}"); r1.bold = True
        r1.font.size = Pt(10); r1.font.name = "Courier New"
        r2 = p.add_run(f"=  {nilai:<24}  {ket}")
        r2.font.size = Pt(10); r2.font.name = "Courier New"
    par(space_after=6)

    # =========== D. KESIMPULAN ===========
    subjudul("D.  KESIMPULAN")

    cek_list = [
        (f"Rho-min = {R['rho_min']*100:.4f}%  <=  Rho = {R['rho']*100:.4f}%", R["ok_rho_min"]),
        (f"Rho-max = {R['rho_max']*100:.4f}%  >=  Rho = {R['rho']*100:.4f}%", R["ok_rho_max"]),
        (f"et      = {R['et']:.5f}  >=  0.004",                                R["ok_et"]),
        (f"D/C-Lentur = {R['DC']:.3f}  <=  1.000",                             R["ok_dc"]),
    ]
    if G is not None:
        cek_list.extend([
            (f"D/C-Geser  = {G['DC_geser']:.3f}  <=  1.000",                 G["ok_dc"]),
            (f"Vs_perlu   = {G['Vs_perlu']:.2f}  <=  Vs_max = {G['Vs_max']:.2f} kN",
             G["ok_Vs_max"]),
            (f"Av_pasang  = {G['Av_pasang']:.1f}  >=  Av_min = {G['Av_min']:.1f} mm2",
             G["ok_Av"]),
            (f"s_pasang   = {geser_inputs['s_seng']:.0f}  <=  s_max = {G['s_max']:.1f} mm",
             G["ok_spasi"]),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        cek_list.extend([
            (f"D/C-Torsi  = {Tor['DC_torsi']:.3f}  <=  1.000",               Tor["ok_DC_torsi"]),
            (f"Dimensi penampang (cek geser+torsi)",                           Tor["ok_dimensi"]),
            (f"Av+2At terpasang  >= Av+2At minimum",                          Tor["ok_Avt_min"]),
            (f"s_pasang   = {geser_inputs['s_seng']:.0f}  <=  s_max_torsi = {Tor['s_max_torsi']:.1f} mm",
             Tor["ok_spasi_torsi"]),
        ])

    for teks_k, ok_k in cek_list:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent  = Cm(0.5)
        r_k = p.add_run(f"{teks_k}   --> {'[OK]' if ok_k else '[TIDAK OK]'}")
        r_k.font.name = "Courier New"; r_k.font.size = Pt(10); r_k.bold = True
        r_k.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if ok_k
                              else RGBColor(0xB7, 0x1C, 0x1C))
    par(space_after=4)

    # Status gabungan
    status_lentur = "AMAN" if R["ok_dc"] else "TIDAK AMAN"
    status_geser  = ("AMAN" if (G is not None and G["ok_total"]) else
                     ("TIDAK AMAN" if G is not None else "belum dievaluasi"))
    if Tor is None:
        status_torsi = "belum dievaluasi"
        dc_torsi_str = "-"
        aman_torsi = True
    elif Tor["abaikan_torsi"]:
        status_torsi = "DIABAIKAN"
        dc_torsi_str = f"Tu={Tor['Tu']:.3f} < phi_t x Tth={Tor['phi_Tth']:.3f} kN.m"
        aman_torsi = True
    else:
        status_torsi = "AMAN" if Tor["ok_torsi_total"] else "TIDAK AMAN"
        dc_torsi_str = f"D/C = {Tor['DC_torsi']:.3f}"
        aman_torsi = Tor["ok_torsi_total"]

    aman_total = R["ok_dc"] and (G is None or G["ok_total"]) and aman_torsi

    kes_lines = [
        f"Status Lentur  : {status_lentur}  (D/C = {R['DC']:.3f})",
        f"Status Geser   : {status_geser}  (D/C = {G['DC_geser']:.3f})" if G else "Status Geser   : belum dievaluasi",
        f"Status Torsi   : {status_torsi}  ({dc_torsi_str})",
        "",
        "KESIMPULAN AKHIR: Penampang " + ("AMAN" if aman_total else "TIDAK AMAN") +
        " secara LENTUR + GESER + TORSI",
    ]
    for ln in kes_lines:
        p_kes = doc.add_paragraph()
        p_kes.paragraph_format.space_before = Pt(2); p_kes.paragraph_format.space_after = Pt(2)
        r_kes = p_kes.add_run(ln)
        r_kes.bold = True; r_kes.font.size = Pt(10.5); r_kes.font.name = "Calibri"
        if "AMAN" in ln or "AMAN" == status_torsi:
            r_kes.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        if "TIDAK AMAN" in ln:
            r_kes.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    par(space_after=6)

    # =========== E. ANALISA GESER ===========
    if G is not None and steps_geser:
        subjudul("E.  ANALISA PERHITUNGAN - GESER")
        par("Urutan perhitungan geser mengacu pada SNI 2847:2019.",
            size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
        tulis_steps(steps_geser)
        par(space_after=6)

        if G["ok_total"]:
            kes_g = (f"KESIMPULAN GESER : AMAN  |  D/C-Geser = {G['DC_geser']:.3f} <= 1.0  |  "
                     f"Av, spasi, dan Vs_max memenuhi syarat.")
            ok_g = True
        elif G["ok_dc"]:
            kes_g = (f"KESIMPULAN GESER : Kapasitas terpenuhi (D/C = {G['DC_geser']:.3f}), "
                     f"namun ada syarat detailing yang tidak terpenuhi.")
            ok_g = False
        else:
            kes_g = (f"KESIMPULAN GESER : TIDAK AMAN  |  D/C-Geser = {G['DC_geser']:.3f} > 1.0  |  "
                     f"Vu = {G['Vu']:.2f} kN > Phi.Vn = {G['PhiVn_efektif']:.2f} kN.")
            ok_g = False

        p_kg = doc.add_paragraph()
        p_kg.paragraph_format.space_before = Pt(6); p_kg.paragraph_format.space_after = Pt(4)
        r_kg = p_kg.add_run(kes_g); r_kg.bold = True; r_kg.font.size = Pt(10.5)
        r_kg.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if ok_g
                               else RGBColor(0xB7, 0x1C, 0x1C))

    # =========== F. ANALISA TORSI ===========
    if Tor is not None and steps_torsi:
        subjudul("F.  ANALISA PERHITUNGAN - TORSI")
        par("Urutan perhitungan torsi mengacu pada SNI 2847:2019 Pasal 22.7.",
            size=10, italic=True, color=(0x55, 0x55, 0x55), space_after=8)
        tulis_steps(steps_torsi)
        par(space_after=6)

        if Tor["abaikan_torsi"]:
            kes_t = (f"KESIMPULAN TORSI : DIABAIKAN  |  "
                     f"Tu = {Tor['Tu']:.3f} kN.m <= Phi_t x Tth = {Tor['phi_Tth']:.3f} kN.m")
            ok_t = True
        elif Tor["ok_torsi_total"]:
            kes_t = (f"KESIMPULAN TORSI : AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} <= 1.0  |  "
                     f"Semua syarat torsi terpenuhi.")
            ok_t = True
        else:
            kes_t = (f"KESIMPULAN TORSI : TIDAK AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} > 1.0  |  "
                     f"Periksa dimensi, sengkang, dan tulangan longitudinal.")
            ok_t = False

        p_kt = doc.add_paragraph()
        p_kt.paragraph_format.space_before = Pt(6); p_kt.paragraph_format.space_after = Pt(4)
        r_kt = p_kt.add_run(kes_t); r_kt.bold = True; r_kt.font.size = Pt(10.5)
        r_kt.font.color.rgb = (RGBColor(0x1B, 0x5E, 0x20) if ok_t
                               else RGBColor(0xB7, 0x1C, 0x1C))

    garis()
    par("Referensi: SNI 2847:2019 | ACI 318-14  --  "
        "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan.",
        size=8, italic=True, color=(0x99, 0x99, 0x99),
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf


# ============================================================
# GENERATOR PDF (.pdf) -- UPDATED dengan Torsi
# ============================================================
WATERMARK_TEXT = "DIHASILKAN OLEH: LADOSI ENGINEERING"
BRAND_COLOR    = (26, 60, 94)
OK_COLOR       = (27, 94, 32)
FAIL_COLOR     = (183, 28, 28)
GRAY           = (120, 120, 120)


class LaporanBalokPDF(FPDF):
    def __init__(self, nama_proyek):
        super().__init__()
        self.nama_proyek = sp(nama_proyek)
        self.set_margins(25, 25, 20)
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.8)
        self.line(25, 15, 190, 15)
        self.set_font("Helvetica", "B", 9); self.set_text_color(*BRAND_COLOR)
        self.set_xy(25, 17)
        self.cell(0, 5, sp("LAPORAN EVALUASI LENTUR, GESER & TORSI BALOK  |  SNI 2847:2019"),
                  ln=False, align="L")
        self.set_font("Helvetica", "", 8); self.set_text_color(*GRAY)
        self.set_xy(25, 17)
        self.cell(0, 5, sp(f"Proyek: {self.nama_proyek}"), ln=False, align="R")
        self.ln(10)

    def footer(self):
        self.set_y(-18)
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.4)
        self.line(25, self.get_y(), 190, self.get_y())
        self.set_font("Helvetica", "I", 7.5); self.set_text_color(*GRAY)
        self.cell(0, 6,
            sp("Referensi: SNI 2847:2019 | ACI 318-14 - "
               "Untuk keperluan profesional, verifikasi mandiri tetap diperlukan."),
            align="C")
        self.set_xy(25, self.get_y())
        self.set_font("Helvetica", "", 7.5)
        self.cell(0, 6, sp(f"Halaman {self.page_no()}"), align="R")

    def watermark(self):
        self.set_font("Helvetica", "B", 28); self.set_text_color(210, 215, 220)
        xc, yc = self.w / 2, self.h / 2
        with self.rotation(40, xc, yc):
            self.set_xy(xc - 65, yc - 6)
            self.cell(130, 12, sp(WATERMARK_TEXT), align="C")
        self.set_text_color(0, 0, 0)

    def section_title(self, teks):
        self.set_font("Helvetica", "B", 11); self.set_text_color(*BRAND_COLOR)
        self.ln(4); self.cell(0, 7, sp(teks), ln=True)
        self.set_draw_color(*BRAND_COLOR); self.set_line_width(0.4)
        self.line(self.get_x(), self.get_y(), 190, self.get_y())
        self.ln(3); self.set_text_color(0, 0, 0)

    def mono_line(self, teks, bold=False, color=None):
        self.set_font("Courier", "B" if bold else "", 9)
        if color: self.set_text_color(*color)
        self.set_x(28); self.multi_cell(0, 4.5, sp(teks))
        self.set_text_color(0, 0, 0)


def _tulis_steps_pdf(pdf, steps_, ok_color=OK_COLOR, fail_color=FAIL_COLOR):
    """Helper: tulis daftar steps ke PDF dengan auto-add page."""
    for s in steps_:
        if pdf.get_y() > 240:
            pdf.add_page(); pdf.watermark()
        pdf.set_font("Helvetica", "B", 10); pdf.set_text_color(*BRAND_COLOR); pdf.set_x(25)
        pdf.cell(0, 6, sp(f"{s['no']}  {s['judul']}"), ln=True)
        pdf.set_font("Helvetica", "I", 8); pdf.set_text_color(*GRAY); pdf.set_x(28)
        pdf.cell(0, 4, sp(f"[{s['ref']}]"), ln=True); pdf.set_text_color(0, 0, 0)
        for baris in s["isi"].split("\n"):
            is_result = (baris.strip().startswith("-->") or
                         "[OK]" in baris or "[TIDAK OK]" in baris or
                         "AMAN" in baris or "[DIABAIKAN]" in baris)
            pdf.mono_line(baris if baris.strip() else "",
                          bold=is_result,
                          color=(ok_color if s["ok"] else fail_color) if is_result else None)
        pdf.ln(2)


def create_pdf_balok(fc, fy, b, h, cc_sel, ds, Mu, lapis_tarik, lapis_tekan,
                     R, steps, nama_proyek, png_buf,
                     G=None, steps_geser=None, geser_inputs=None,
                     Tor=None, steps_torsi=None, torsi_inputs=None,
                     timestamp_str=None):
    pdf = LaporanBalokPDF(nama_proyek)
    pdf.add_page(); pdf.watermark()

    if timestamp_str is None:
        timestamp_str = datetime.datetime.now().strftime('%d %B %Y  %H:%M')

    # === Header laporan ===
    pdf.set_font("Helvetica", "B", 15); pdf.set_text_color(*BRAND_COLOR); pdf.ln(2)
    pdf.cell(0, 9, sp("LAPORAN PERHITUNGAN STRUKTUR"), ln=True, align="C")
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, sp("Evaluasi Kapasitas Lentur, Geser & Torsi Balok Beton Bertulang"),
             ln=True, align="C")
    pdf.set_font("Helvetica", "I", 9); pdf.set_text_color(*GRAY)
    pdf.cell(0, 5, sp("(Tulangan Rangkap - Strain Compatibility)"), ln=True, align="C")
    pdf.cell(0, 5, sp("Referensi: SNI 2847:2019 (ACI 318-14)"), ln=True, align="C")
    pdf.cell(0, 5, sp(f"Proyek: {nama_proyek}   |   Tanggal: {timestamp_str}"),
             ln=True, align="C")
    pdf.ln(6); pdf.set_draw_color(*BRAND_COLOR); pdf.set_line_width(0.6)
    pdf.line(25, pdf.get_y(), 190, pdf.get_y()); pdf.ln(6); pdf.set_text_color(0, 0, 0)

    # =========== A. DATA INPUT ===========
    pdf.section_title("A.  DATA INPUT PENAMPANG")
    rows = [
        ("Mu",   f"{Mu:.2f} kN.m",   "Momen rencana"),
        ("fc",   f"{fc:.1f} MPa",    "Kuat tekan beton"),
        ("fy",   f"{fy:.0f} MPa",    "Kuat leleh tulangan lentur"),
        ("b",    f"{b:.0f} mm",      "Lebar balok"),
        ("h",    f"{h:.0f} mm",      "Tinggi total balok"),
        ("cc",   f"{cc_sel:.0f} mm", "Tebal selimut bersih"),
        ("ds",   f"{ds:.0f} mm",     "Diameter sengkang"),
    ]
    if geser_inputs is not None:
        rows.extend([
            ("Vu",     f"{geser_inputs['Vu']:.2f} kN",    "Gaya geser ultimit"),
            ("fyt",    f"{geser_inputs['fyt']:.0f} MPa",  "Kuat leleh sengkang"),
            ("s",      f"{geser_inputs['s_seng']:.0f} mm","Jarak sengkang"),
            ("n_kaki", f"{geser_inputs['n_kaki']:d}",     "Jumlah kaki sengkang"),
        ])
    if torsi_inputs is not None:
        rows.extend([
            ("Tu",          f"{torsi_inputs['Tu']:.2f} kN.m",    "Momen torsi ultimit"),
            ("tipe_torsi",  f"{torsi_inputs['tipe_torsi']}",     "Tipe torsi"),
            ("db_long_torsi",f"D{int(torsi_inputs['db_long_torsi'])} mm","Dia. tul. long. torsi"),
        ])
    for simb, nilai, ket in rows:
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.cell(22, 5, sp(f"{simb:<10}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(36, 5, sp(f"=  {nilai}"), ln=False)
        pdf.set_font("Helvetica", "I", 8.5); pdf.set_text_color(*GRAY)
        pdf.cell(0, 5, sp(f"({ket})"), ln=True); pdf.set_text_color(0, 0, 0)

    pdf.ln(2)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp("Tulangan TARIK (bawah):"), ln=True); pdf.set_text_color(0, 0, 0)
    for L in lapis_tarik:
        pdf.set_x(30); pdf.set_font("Courier", "", 9)
        pdf.cell(0, 4.5,
                 sp(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                    f"As = {L['As']:.1f} mm2   y = {L['y']:.1f} mm"),
                 ln=True)
    pdf.set_x(30); pdf.set_font("Courier", "B", 9)
    pdf.cell(0, 4.5,
             sp(f"As-total = {R['As_tarik']:.1f} mm2  |  d-aktual = {R['d_tarik']:.2f} mm"),
             ln=True)

    pdf.ln(1)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp("Tulangan TEKAN (atas):"), ln=True); pdf.set_text_color(0, 0, 0)
    if lapis_tekan:
        for L in lapis_tekan:
            pdf.set_x(30); pdf.set_font("Courier", "", 9)
            pdf.cell(0, 4.5,
                     sp(f"Lapis-{L['no']}: {L['n']} D{int(L['db'])}  "
                        f"As' = {L['As']:.1f} mm2   y = {L['y']:.1f} mm"),
                     ln=True)
        pdf.set_x(30); pdf.set_font("Courier", "B", 9)
        pdf.cell(0, 4.5,
                 sp(f"As'-total = {R['As_tekan']:.1f} mm2  |  d'-aktual = {R['d_tekan']:.2f} mm"),
                 ln=True)
    else:
        pdf.set_x(30); pdf.set_font("Helvetica", "I", 9); pdf.set_text_color(*GRAY)
        pdf.cell(0, 4.5, sp("(Tidak ada tulangan tekan)"), ln=True)
        pdf.set_text_color(0, 0, 0)

    pdf.ln(3)
    pdf.set_x(25); pdf.set_font("Helvetica", "B", 9.5); pdf.set_text_color(*BRAND_COLOR)
    pdf.cell(0, 5, sp("Visualisasi Penampang:"), ln=True); pdf.set_text_color(0, 0, 0)
    img_path = "/tmp/_penampang.png"
    with open(img_path, "wb") as fimg:
        fimg.write(png_buf.getvalue())
    pdf.image(img_path, x=70, y=pdf.get_y() + 1, w=70)
    pdf.ln(82)

    # =========== B. ANALISA LENTUR ===========
    if pdf.get_y() > 220:
        pdf.add_page(); pdf.watermark()
    pdf.section_title("B.  ANALISA PERHITUNGAN - LENTUR")
    _tulis_steps_pdf(pdf, steps)

    # =========== C. RANGKUMAN ===========
    if pdf.get_y() > 200:
        pdf.add_page(); pdf.watermark()
    pdf.section_title("C.  RANGKUMAN HASIL")

    rangkuman = [
        ("Beta-1",     f"{R['beta1']:.4f}",        "Sumbu netral"),
        ("c",          f"{R['c']:.2f} mm",          "Sumbu netral"),
        ("a",          f"{R['a']:.2f} mm",          "Blok tegangan"),
        ("Cc",         f"{R['Cc']/1000:.2f} kN",    "Gaya tekan beton"),
        ("Cs",         f"{R['Cs']/1000:.2f} kN",    "Gaya tekan baja"),
        ("T",          f"{R['T']/1000:.2f} kN",     "Gaya tarik baja"),
        ("et",         f"{R['et']:.5f}",            "Regangan tarik"),
        ("Phi",        f"{R['phi']:.4f}",           "Faktor reduksi lentur"),
        ("Rho",        f"{R['rho']*100:.4f}%",      "Rasio tulangan tarik"),
        ("Rho-min",    f"{R['rho_min']*100:.4f}%",  "Batas minimum"),
        ("Rho-max",    f"{R['rho_max']*100:.4f}%",  "Batas maksimum"),
        ("Mn",         f"{R['Mn']:.3f} kN.m",       "Momen nominal"),
        ("Phi.Mn",     f"{R['phiMn']:.3f} kN.m",    "Momen rencana kapasitas"),
        ("Mu",         f"{R['Mu']:.3f} kN.m",       "Momen rencana ultimit"),
        ("D/C-Lentur", f"{R['DC']:.3f}",            "DC Lentur"),
    ]
    if G is not None:
        rangkuman.extend([
            ("Phi_v",     f"{G['Phi_v']:.2f}",            "Faktor reduksi geser"),
            ("Vc",        f"{G['Vc']:.2f} kN",            "Geser beton"),
            ("Vs",        f"{G['Vs_efektif']:.2f} kN",    "Geser sengkang"),
            ("Vn",        f"{G['Vn_efektif']:.2f} kN",    "Geser nominal"),
            ("Phi.Vn",    f"{G['PhiVn_efektif']:.2f} kN", "Geser rencana"),
            ("Vu",        f"{G['Vu']:.2f} kN",            "Geser ultimit"),
            ("Av_pasang", f"{G['Av_pasang']:.1f} mm2",    "Sengkang terpasang"),
            ("Av_min",    f"{G['Av_min']:.1f} mm2",       "Sengkang minimum"),
            ("D/C-Geser", f"{G['DC_geser']:.3f}",         "DC Geser"),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        rangkuman.extend([
            ("Phi_t",      f"{Tor['Phi_t']:.2f}",               "Faktor reduksi torsi"),
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",            "Batas ambang torsi"),
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",             "Torsi ultimit"),
            ("Tu_desain",  f"{Tor['Tu_desain']:.4f} kN.m",      "Tu desain"),
            ("Phi.Tn",     f"{Tor['PhiTn_cap']:.4f} kN.m",      "Kapasitas torsi"),
            ("D/C-Torsi",  f"{Tor['DC_torsi']:.3f}",            "DC Torsi"),
            ("At/s",       f"{Tor['At_per_s']:.6f} mm2/mm",     "Kebutuhan sengkang torsi"),
            ("Al_pakai",   f"{Tor['Al_pakai']:.2f} mm2",        "Tul. longitudinal torsi"),
            ("n_tul_torsi",f"{Tor['n_batang']} D{int(Tor['db_long_torsi'])}", "Batang long. torsi"),
            ("s_max_torsi",f"{Tor['s_max_torsi']:.1f} mm",      "Spasi maks sengkang torsi"),
        ])
    elif Tor is not None and Tor["abaikan_torsi"]:
        rangkuman.extend([
            ("Tu",         f"{Tor['Tu']:.3f} kN.m",    "Torsi ultimit"),
            ("Tth",        f"{Tor['Tth']:.4f} kN.m",   "Batas ambang torsi"),
            ("D/C-Torsi",  "DIABAIKAN",                 "Tu < phi_t x Tth"),
        ])

    for simb, nilai, ket in rangkuman:
        if pdf.get_y() > 260:
            pdf.add_page(); pdf.watermark()
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.cell(28, 5, sp(f"{simb:<13}"), ln=False)
        pdf.set_font("Courier", "", 9.5)
        pdf.cell(38, 5, sp(f"=  {nilai}"), ln=False)
        if ket:
            pdf.set_font("Helvetica", "I", 8.5); pdf.set_text_color(*GRAY)
            pdf.cell(0, 5, sp(f"({ket})"), ln=True); pdf.set_text_color(0, 0, 0)
        else:
            pdf.ln()
    pdf.ln(4)

    # =========== D. KESIMPULAN ===========
    if pdf.get_y() > 220:
        pdf.add_page(); pdf.watermark()
    pdf.section_title("D.  KESIMPULAN")
    cek_list = [
        (f"Rho-min    = {R['rho_min']*100:.4f}%  <=  Rho = {R['rho']*100:.4f}%", R["ok_rho_min"]),
        (f"Rho-max    = {R['rho_max']*100:.4f}%  >=  Rho = {R['rho']*100:.4f}%", R["ok_rho_max"]),
        (f"et         = {R['et']:.5f}  >=  0.004",                                R["ok_et"]),
        (f"D/C-Lentur = {R['DC']:.3f}  <=  1.000",                                R["ok_dc"]),
    ]
    if G is not None:
        cek_list.extend([
            (f"D/C-Geser  = {G['DC_geser']:.3f}  <=  1.000",                  G["ok_dc"]),
            (f"Vs_perlu   = {G['Vs_perlu']:.2f}  <=  Vs_max = {G['Vs_max']:.2f} kN",
             G["ok_Vs_max"]),
            (f"Av_pasang  = {G['Av_pasang']:.1f}  >=  Av_min = {G['Av_min']:.1f} mm2",
             G["ok_Av"]),
            (f"s_pasang   = {geser_inputs['s_seng']:.0f}  <=  s_max = {G['s_max']:.1f} mm",
             G["ok_spasi"]),
        ])
    if Tor is not None and not Tor["abaikan_torsi"]:
        cek_list.extend([
            (f"D/C-Torsi  = {Tor['DC_torsi']:.3f}  <=  1.000",               Tor["ok_DC_torsi"]),
            (f"Cek dimensi penampang (SNI 22.7.7.1)",                          Tor["ok_dimensi"]),
            (f"Av+2At >= (Av+2At)_min  (SNI 9.6.4.2)",                        Tor["ok_Avt_min"]),
            (f"s_pasang = {geser_inputs['s_seng']:.0f} <= s_max_torsi = {Tor['s_max_torsi']:.1f} mm",
             Tor["ok_spasi_torsi"]),
        ])

    for teks_k, ok_k in cek_list:
        if pdf.get_y() > 260:
            pdf.add_page(); pdf.watermark()
        pdf.set_x(28); pdf.set_font("Courier", "B", 9.5)
        pdf.set_text_color(*(OK_COLOR if ok_k else FAIL_COLOR))
        pdf.cell(0, 5.5, sp(f"{teks_k}   --> {'[OK]' if ok_k else '[TIDAK OK]'}"), ln=True)
        pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    # Status gabungan
    status_lentur = "AMAN" if R["ok_dc"] else "TIDAK AMAN"
    status_geser  = ("AMAN" if (G is not None and G["ok_total"])
                     else ("TIDAK AMAN" if G is not None else "belum dievaluasi"))
    if Tor is None:
        status_torsi = "belum dievaluasi"
        dc_torsi_str = "-"
        aman_torsi = True
    elif Tor["abaikan_torsi"]:
        status_torsi = "DIABAIKAN"
        dc_torsi_str = f"Tu={Tor['Tu']:.3f} < phi_t x Tth={Tor['phi_Tth']:.3f} kN.m"
        aman_torsi = True
    else:
        status_torsi = "AMAN" if Tor["ok_torsi_total"] else "TIDAK AMAN"
        dc_torsi_str = f"D/C = {Tor['DC_torsi']:.3f}"
        aman_torsi = Tor["ok_torsi_total"]

    aman_total = R["ok_dc"] and (G is None or G["ok_total"]) and aman_torsi

    kes_lines = [
        f"Status Lentur  : {status_lentur}  (D/C = {R['DC']:.3f})",
        (f"Status Geser   : {status_geser}  (D/C = {G['DC_geser']:.3f})"
         if G else "Status Geser   : belum dievaluasi"),
        f"Status Torsi   : {status_torsi}  ({dc_torsi_str})",
        "",
        ("KESIMPULAN AKHIR: Penampang " +
         ("AMAN" if aman_total else "TIDAK AMAN") +
         " secara LENTUR + GESER + TORSI"),
    ]
    for ln_t in kes_lines:
        if not ln_t:
            pdf.ln(2); continue
        ok_line = ("AMAN" in ln_t and "TIDAK" not in ln_t)
        fail_line = "TIDAK AMAN" in ln_t
        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        if fail_line:
            pdf.set_text_color(*FAIL_COLOR)
        elif ok_line:
            pdf.set_text_color(*OK_COLOR)
        else:
            pdf.set_text_color(*GRAY)
        pdf.multi_cell(0, 6, sp(ln_t))
        pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    # =========== E. ANALISA GESER ===========
    if G is not None and steps_geser:
        if pdf.get_y() > 220:
            pdf.add_page(); pdf.watermark()
        pdf.section_title("E.  ANALISA PERHITUNGAN - GESER")
        _tulis_steps_pdf(pdf, steps_geser)

        if G["ok_total"]:
            kes_g = (f"KESIMPULAN GESER : AMAN  |  "
                     f"D/C-Geser = {G['DC_geser']:.3f} <= 1.0  |  "
                     f"Av, spasi, dan Vs_max memenuhi syarat.")
            ok_g = True
        elif G["ok_dc"]:
            kes_g = (f"KESIMPULAN GESER : Kapasitas terpenuhi "
                     f"(D/C = {G['DC_geser']:.3f}), namun ada syarat detailing yang tidak terpenuhi.")
            ok_g = False
        else:
            kes_g = (f"KESIMPULAN GESER : TIDAK AMAN  |  "
                     f"D/C-Geser = {G['DC_geser']:.3f} > 1.0  |  "
                     f"Vu = {G['Vu']:.2f} kN > Phi.Vn = {G['PhiVn_efektif']:.2f} kN.")
            ok_g = False

        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*(OK_COLOR if ok_g else FAIL_COLOR))
        pdf.multi_cell(0, 6, sp(kes_g)); pdf.set_text_color(0, 0, 0)

    # =========== F. ANALISA TORSI ===========
    if Tor is not None and steps_torsi:
        if pdf.get_y() > 220:
            pdf.add_page(); pdf.watermark()
        pdf.section_title("F.  ANALISA PERHITUNGAN - TORSI")
        _tulis_steps_pdf(pdf, steps_torsi)

        if Tor["abaikan_torsi"]:
            kes_t = (f"KESIMPULAN TORSI : DIABAIKAN  |  "
                     f"Tu = {Tor['Tu']:.3f} kN.m <= Phi_t x Tth = {Tor['phi_Tth']:.3f} kN.m")
            ok_t = True
        elif Tor["ok_torsi_total"]:
            kes_t = (f"KESIMPULAN TORSI : AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} <= 1.0  |  "
                     f"Semua syarat torsi terpenuhi.")
            ok_t = True
        else:
            kes_t = (f"KESIMPULAN TORSI : TIDAK AMAN  |  D/C-Torsi = {Tor['DC_torsi']:.3f} > 1.0  |  "
                     f"Periksa dimensi, sengkang, dan tulangan longitudinal.")
            ok_t = False

        pdf.set_x(25); pdf.set_font("Helvetica", "B", 10.5)
        pdf.set_text_color(*(OK_COLOR if ok_t else FAIL_COLOR))
        pdf.multi_cell(0, 6, sp(kes_t)); pdf.set_text_color(0, 0, 0)

    buf = io.BytesIO(); pdf.output(buf); buf.seek(0)
    return buf


# ============================================================
# HELPER SESSION STATE
# ============================================================
def _init_state():
    defaults = {
        "balok_hasil":         None,
        "balok_steps":         None,
        "balok_geser_hasil":   None,
        "balok_geser_steps":   None,
        "balok_geser_inputs":  None,
        "balok_torsi_hasil":   None,    # TAHAP 2
        "balok_torsi_steps":   None,    # TAHAP 2
        "balok_torsi_inputs":  None,    # TAHAP 2
        "balok_word":          None,
        "balok_pdf":           None,
        "balok_fig":           None,
        "balok_last_inputs":   {},
        "balok_nama_file":     "",
        "balok_timestamp":     "",
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v


def _snapshot_inputs(d):
    return dict(d)


def _inputs_changed(snap_now):
    last = st.session_state.balok_last_inputs
    if not last:
        return False
    return any(snap_now.get(k) != last.get(k) for k in snap_now)


# ============================================================
# UI
# ============================================================
_init_state()

st.markdown('<p class="main-title">Evaluasi Kapasitas Lentur, Geser & Torsi Balok Beton Bertulang</p>',
            unsafe_allow_html=True)
st.markdown(
    '<p class="sub-title">Tulangan Rangkap | Strain Compatibility | '
    'Referensi: SNI 2847:2019 (setara ACI 318-14)</p>',
    unsafe_allow_html=True)
st.markdown('<hr class="divider">', unsafe_allow_html=True)

col_inp, col_out = st.columns([1, 2], gap="large")

with col_inp:
    st.markdown("### Data Input")

    nama_proyek = st.text_input(
        "Nama Proyek (untuk header laporan)",
        value="Laporan Analisa Struktur",
        key="balok_inp_nama",
    )

    st.markdown('<div class="group-hdr">Gaya Dalam</div>', unsafe_allow_html=True)
    Mu = st.number_input("Mu - Momen rencana (kN.m)",
                         min_value=0.0, max_value=10000.0, value=200.0,
                         step=5.0, format="%.2f", key="balok_inp_mu")
    Vu = st.number_input("Vu - Gaya geser ultimit (kN)",
                         min_value=0.0, max_value=5000.0, value=80.0,
                         step=5.0, format="%.2f", key="balok_inp_vu")
    Tu = st.number_input("Tu - Momen torsi ultimit (kN.m)",
                         min_value=0.0, max_value=2000.0, value=0.0,
                         step=1.0, format="%.2f", key="balok_inp_tu",
                         help="Isi 0 jika tidak ada torsi. Evaluasi akan otomatis berjalan.")

    st.markdown('<div class="group-hdr">Material</div>', unsafe_allow_html=True)
    fc = st.number_input("f'c - Kuat tekan beton (MPa)",
                         min_value=17.0, max_value=100.0, value=30.0,
                         step=1.0, format="%.1f", key="balok_inp_fc")
    fy = st.number_input("fy - Kuat leleh tulangan lentur (MPa)",
                         min_value=240.0, max_value=600.0, value=400.0,
                         step=10.0, format="%.0f", key="balok_inp_fy")
    fyt = st.number_input("fyt - Kuat leleh sengkang (MPa)",
                          min_value=240.0, max_value=420.0, value=240.0,
                          step=10.0, format="%.0f", key="balok_inp_fyt",
                          help="Umumnya 240 MPa untuk BJTP, atau samakan dengan fy bila pakai BJTS")

    st.markdown('<div class="group-hdr">Geometri</div>', unsafe_allow_html=True)
    cb, ch_ = st.columns(2)
    with cb:
        b = st.number_input("b (mm)", min_value=100.0, max_value=2000.0,
                            value=300.0, step=10.0, format="%.0f", key="balok_inp_b")
    with ch_:
        h = st.number_input("h (mm)", min_value=100.0, max_value=5000.0,
                            value=500.0, step=10.0, format="%.0f", key="balok_inp_h")
    cc_sel = st.number_input("cc (mm) - Tebal selimut bersih",
                             min_value=15.0, max_value=100.0, value=40.0,
                             step=5.0, format="%.0f", key="balok_inp_cc")
    ds_dia = st.number_input("ds (mm) - Diameter sengkang",
                             min_value=6.0, max_value=16.0, value=10.0,
                             step=1.0, format="%.0f", key="balok_inp_ds")

    st.markdown('<div class="group-hdr">Sengkang</div>', unsafe_allow_html=True)
    cs1, cs2 = st.columns(2)
    with cs1:
        s_seng = st.number_input("s - Jarak sengkang (mm)",
                                 min_value=50.0, max_value=600.0, value=150.0,
                                 step=10.0, format="%.0f", key="balok_inp_s")
    with cs2:
        n_kaki = st.number_input("Jumlah kaki (n_kaki)",
                                 min_value=2, max_value=6, value=2,
                                 step=1, key="balok_inp_nkaki",
                                 help="Default 2 kaki. Untuk balok lebar bisa 4 kaki.")

    st.markdown('<div class="group-hdr">Input Torsi Tambahan</div>', unsafe_allow_html=True)
    tipe_torsi = st.radio(
        "Tipe Torsi",
        options=["Equilibrium", "Compatibility"],
        index=0,
        key="balok_inp_tipe_torsi",
        help="Equilibrium: tidak boleh redistribusi. Compatibility: boleh direduksi ke Phi.Tcr.",
        horizontal=True,
    )
    db_long_torsi = st.selectbox(
        "Diameter tul. longitudinal torsi (mm)",
        options=[10, 13, 16, 19],
        index=1,
        key="balok_inp_db_long_torsi",
        help="Diameter tulangan longitudinal tambahan untuk menahan torsi.",
    )

    st.markdown('<div class="group-hdr">Tulangan TARIK (bawah)</div>', unsafe_allow_html=True)
    st.caption("Lapis 1 = paling bawah, Lapis 2 = di atas Lapis 1 (spasi bersih 25 mm)")
    ct1, ct2 = st.columns(2)
    with ct1:
        n_t1  = st.number_input("Lapis-1: Jumlah", min_value=0, max_value=20,
                                value=4, step=1, key="balok_inp_nt1")
        n_t2  = st.number_input("Lapis-2: Jumlah", min_value=0, max_value=20,
                                value=0, step=1, key="balok_inp_nt2")
    with ct2:
        db_t1 = st.selectbox("Lapis-1: Diameter (mm)", DIAMETER_LIST,
                             index=4, key="balok_inp_dbt1")
        db_t2 = st.selectbox("Lapis-2: Diameter (mm)", DIAMETER_LIST,
                             index=4, key="balok_inp_dbt2")

    st.markdown('<div class="group-hdr">Tulangan TEKAN (atas)</div>', unsafe_allow_html=True)
    st.caption("Lapis 1 = paling atas, Lapis 2 = di bawah Lapis 1 (spasi bersih 25 mm)")
    cc1, cc2 = st.columns(2)
    with cc1:
        n_c1  = st.number_input("Lapis-1: Jumlah", min_value=0, max_value=20,
                                value=2, step=1, key="balok_inp_nc1")
        n_c2  = st.number_input("Lapis-2: Jumlah", min_value=0, max_value=20,
                                value=0, step=1, key="balok_inp_nc2")
    with cc2:
        db_c1 = st.selectbox("Lapis-1: Diameter (mm)", DIAMETER_LIST,
                             index=3, key="balok_inp_dbc1")
        db_c2 = st.selectbox("Lapis-2: Diameter (mm)", DIAMETER_LIST,
                             index=3, key="balok_inp_dbc2")

    st.markdown("")
    tombol = st.button("HITUNG EVALUASI LENTUR, GESER & TORSI",
                       use_container_width=True, type="primary",
                       key="balok_btn_hitung")

    with st.expander("Tabel luas tulangan (mm2)"):
        st.markdown("""
        | D | 1 | 2 | 3 | 4 | 5 | 6 |
        |---|---|---|---|---|---|---|
        | D10 | 78.5 | 157 | 236 | 314 | 393 | 471 |
        | D13 | 132.7 | 265 | 398 | 531 | 663 | 796 |
        | D16 | 201.1 | 402 | 603 | 804 | 1005 | 1206 |
        | D19 | 283.5 | 567 | 851 | 1134 | 1418 | 1701 |
        | D22 | 380.1 | 760 | 1140 | 1520 | 1901 | 2281 |
        | D25 | 490.9 | 982 | 1473 | 1964 | 2454 | 2945 |
        | D29 | 660.5 | 1321 | 1981 | 2642 | 3302 | 3963 |
        | D32 | 804.2 | 1608 | 2413 | 3217 | 4021 | 4825 |
        """)


# ============================================================
# PROSES HITUNG
# ============================================================
if tombol:
    valid = True
    if n_t1 == 0 and n_t2 == 0:
        st.error("Harus ada minimal 1 lapis tulangan tarik!")
        valid = False
    if Mu <= 0:
        st.error("Mu harus lebih besar dari 0!")
        valid = False
    if Vu < 0:
        st.error("Vu tidak boleh negatif!")
        valid = False
    if Tu < 0:
        st.error("Tu tidak boleh negatif!")
        valid = False

    if valid:
        now = datetime.datetime.now()
        ts_file    = now.strftime("%Y%m%d_%H%M")
        ts_display = now.strftime("%d %B %Y  %H:%M")

        lapis_tarik, d_tarik, As_tarik = hitung_lapis_tarik(
            h, cc_sel, ds_dia, n_t1, db_t1, n_t2, db_t2, spasi=25.0)
        lapis_tekan, d_tekan, As_tekan = hitung_lapis_tekan(
            cc_sel, ds_dia, n_c1, db_c1, n_c2, db_c2, spasi=25.0)

        if d_tarik >= h or d_tarik <= 0:
            st.error(f"d-aktual ({d_tarik:.1f} mm) tidak valid. "
                     f"Cek selimut, sengkang, dan diameter tulangan.")
        else:
            # === LENTUR ===
            R = evaluasi_balok(fc, fy, b, h, cc_sel, ds_dia,
                               lapis_tarik, lapis_tekan,
                               d_tarik, As_tarik, d_tekan, As_tekan, Mu)
            steps = buat_steps_balok(fc, fy, b, h, cc_sel, ds_dia,
                                     lapis_tarik, lapis_tekan, R)

            # === GESER ===
            G = hitung_geser(fc=fc, fyt=fyt, b=b, d_aktual=d_tarik,
                             Vu=Vu, ds=ds_dia, s_seng=s_seng,
                             n_kaki=int(n_kaki), lambda_=1.0)
            steps_geser = buat_steps_geser(fc, fyt, b, d_tarik,
                                           Vu, ds_dia, s_seng,
                                           int(n_kaki), G)
            geser_inputs = dict(Vu=Vu, fyt=fyt, s_seng=s_seng,
                                n_kaki=int(n_kaki))

            # === TORSI === (sequential: ambil d, Vc, Vu dari hasil Lentur & Geser)
            Tor = hitung_torsi(
                fc=fc, fy_long=fy, fyt=fyt,
                b=b, h=h, cc_sel=cc_sel, ds=ds_dia, s_seng=s_seng,
                Tu=Tu, Vu=Vu, Vc=G["Vc"], d_aktual=d_tarik,
                Av_pasang_geser=G["Av_pasang"],
                tipe_torsi=tipe_torsi, db_long_torsi=float(db_long_torsi),
                lambda_=1.0,
            )
            steps_torsi = buat_steps_torsi(
                fc, fyt, fy, b, h, cc_sel, ds_dia, s_seng,
                Tu, Vu, d_tarik, Tor
            )
            torsi_inputs = dict(
                Tu=Tu, tipe_torsi=tipe_torsi,
                db_long_torsi=float(db_long_torsi)
            )

            # Gambar penampang (dengan Aoh jika ada torsi aktif)
            fig = gambar_penampang(b, h, cc_sel, ds_dia,
                                   lapis_tarik, lapis_tekan, c=R["c"],
                                   torsi_data=Tor)
            png_buf = fig_to_png_bytes(fig)

            w_buf = create_word_balok(
                fc, fy, b, h, cc_sel, ds_dia, Mu,
                lapis_tarik, lapis_tekan,
                R, steps, nama_proyek, png_buf,
                G=G, steps_geser=steps_geser, geser_inputs=geser_inputs,
                Tor=Tor, steps_torsi=steps_torsi, torsi_inputs=torsi_inputs,
                timestamp_str=ts_display,
            )
            png_buf.seek(0)
            p_buf = create_pdf_balok(
                fc, fy, b, h, cc_sel, ds_dia, Mu,
                lapis_tarik, lapis_tekan,
                R, steps, nama_proyek, png_buf,
                G=G, steps_geser=steps_geser, geser_inputs=geser_inputs,
                Tor=Tor, steps_torsi=steps_torsi, torsi_inputs=torsi_inputs,
                timestamp_str=ts_display,
            )

            snap = dict(
                nama_proyek=nama_proyek,
                Mu=Mu, Vu=Vu, Tu=Tu,
                fc=fc, fy=fy, fyt=fyt,
                b=b, h=h, cc=cc_sel, ds=ds_dia,
                s_seng=s_seng, n_kaki=int(n_kaki),
                tipe_torsi=tipe_torsi, db_long_torsi=float(db_long_torsi),
                n_t1=n_t1, db_t1=db_t1, n_t2=n_t2, db_t2=db_t2,
                n_c1=n_c1, db_c1=db_c1, n_c2=n_c2, db_c2=db_c2,
            )

            st.session_state.balok_hasil        = R
            st.session_state.balok_steps        = steps
            st.session_state.balok_geser_hasil  = G
            st.session_state.balok_geser_steps  = steps_geser
            st.session_state.balok_geser_inputs = geser_inputs
            st.session_state.balok_torsi_hasil  = Tor
            st.session_state.balok_torsi_steps  = steps_torsi
            st.session_state.balok_torsi_inputs = torsi_inputs
            st.session_state.balok_word         = w_buf.getvalue()
            st.session_state.balok_pdf          = p_buf.getvalue()
            st.session_state.balok_fig          = fig
            st.session_state.balok_last_inputs  = snap
            st.session_state.balok_timestamp    = ts_display

            # Format nama file: fc, Mu, Vu, Tu
            st.session_state.balok_nama_file    = (
                f"Laporan_Lentur_Geser_Torsi_fc{int(fc)}_Mu{int(Mu)}"
                f"_Vu{int(Vu)}_Tu{int(Tu)}_{ts_file}"
            )


# ============================================================
# UI - HASIL
# ============================================================
with col_out:
    if st.session_state.balok_hasil is None:
        st.info("Isi data di panel kiri, lalu klik HITUNG EVALUASI LENTUR, GESER & TORSI")
        st.markdown("""
        **Yang akan dihitung:**

        **LENTUR (Strain Compatibility):**
        - **b1.** Faktor Beta-1
        - **b2.** Posisi tiap lapis tulangan & titik berat (d-aktual, d'-aktual)
        - **b3.** Iterasi sumbu netral c
        - **b4.** Regangan & tegangan setiap lapis (et, eps-s', dst.)
        - **b5.** Gaya dalam Cc, Cs, T
        - **b6.** Rasio tulangan Rho (kontrol min & max)
        - **b7.** Klasifikasi penampang
        - **b8.** Faktor reduksi Phi
        - **b9.** Mn dan Phi.Mn
        - **b10.** D/C Ratio Lentur

        **GESER (SNI 2847:2019):**
        - **b11.** Faktor reduksi geser Phi_v
        - **b12.** Vc (kapasitas geser beton)
        - **b13.** Cek kebutuhan sengkang
        - **b14.** Vs perlu & cek Vs maksimum
        - **b15.** Av terpasang
        - **b16.** Av minimum & spasi maksimum
        - **b17.** Vs aktual & Phi.Vn
        - **b18.** D/C Ratio Geser

        **TORSI (SNI 2847:2019 Pasal 22.7):**
        - **b19.** Properti penampang torsi (Acp, pcp, Aoh, ph, Ao)
        - **b20.** Faktor reduksi Phi_t
        - **b21.** Batas ambang torsi (Tth)
        - **b22.** Klasifikasi torsi & Tu desain
        - **b23.** Cek dimensi penampang (geser + torsi)
        - **b24.** Tulangan sengkang torsi (At/s)
        - **b25.** Sengkang gabungan geser + torsi (Av+2At)
        - **b26.** Spasi maksimum & kontrol
        - **b27.** Tulangan longitudinal torsi (Al)
        - **b28.** D/C Torsi & kesimpulan

        **Output tersedia dalam dua format:**
        - Word (.docx) - format natural, bisa diedit
        - PDF (.pdf) - formal + watermark LADOSI ENGINEERING
        """)
    else:
        R           = st.session_state.balok_hasil
        steps       = st.session_state.balok_steps
        G           = st.session_state.balok_geser_hasil
        steps_geser = st.session_state.balok_geser_steps
        geser_inputs= st.session_state.balok_geser_inputs
        Tor         = st.session_state.balok_torsi_hasil
        steps_torsi = st.session_state.balok_torsi_steps
        fig         = st.session_state.balok_fig

        # Soft warning jika input berubah
        snap_now = dict(
            nama_proyek=nama_proyek,
            Mu=Mu, Vu=Vu, Tu=Tu,
            fc=fc, fy=fy, fyt=fyt,
            b=b, h=h, cc=cc_sel, ds=ds_dia,
            s_seng=s_seng, n_kaki=int(n_kaki),
            tipe_torsi=tipe_torsi, db_long_torsi=float(db_long_torsi),
            n_t1=n_t1, db_t1=db_t1, n_t2=n_t2, db_t2=db_t2,
            n_c1=n_c1, db_c1=db_c1, n_c2=n_c2, db_c2=db_c2,
        )
        if _inputs_changed(snap_now):
            st.warning(
                "Perhatian: Data input telah diubah. "
                "Hasil dan file laporan di bawah masih menggunakan "
                "data perhitungan sebelumnya. Klik HITUNG kembali untuk update."
            )

        # ---- Metrik Utama - Baris 1: Lentur ----
        st.markdown("### Hasil Utama - Lentur")
        m1, m2, m3, m4 = st.columns(4)
        for col, lbl, val, unt in [
            (m1, "Phi.Mn (Kapasitas)", f"{R['phiMn']:.2f}", "kN.m"),
            (m2, "Mu (Demand)",        f"{R['Mu']:.2f}",    "kN.m"),
            (m3, "D/C Lentur",         f"{R['DC']:.3f}",    "-"),
            (m4, "Regangan et",        f"{R['et']:.4f}",    "-"),
        ]:
            with col:
                st.markdown(
                    f'<div class="metric-card">'
                    f'<div class="metric-lbl">{lbl}</div>'
                    f'<div class="metric-val">{val}</div>'
                    f'<div class="metric-unt">{unt}</div>'
                    f'</div>', unsafe_allow_html=True)
        st.markdown("")

        # ---- Metrik Utama - Baris 2: Geser ----
        st.markdown("### Hasil Utama - Geser")
        g1, g2, g3, g4 = st.columns(4)
        for col, lbl, val, unt in [
            (g1, "Phi.Vn (Kapasitas)", f"{G['PhiVn_efektif']:.2f}", "kN"),
            (g2, "Vu (Demand)",        f"{G['Vu']:.2f}",            "kN"),
            (g3, "D/C Geser",          f"{G['DC_geser']:.3f}",      "-"),
            (g4, "Av pasang / Av min", f"{G['Av_pasang']:.0f} / {G['Av_min']:.0f}", "mm2"),
        ]:
            with col:
                st.markdown(
                    f'<div class="metric-card">'
                    f'<div class="metric-lbl">{lbl}</div>'
                    f'<div class="metric-val">{val}</div>'
                    f'<div class="metric-unt">{unt}</div>'
                    f'</div>', unsafe_allow_html=True)
        st.markdown("")

        # ---- Metrik Utama - Baris 3: Torsi ----
        st.markdown("### Hasil Utama - Torsi")
        t1, t2, t3, t4 = st.columns(4)
        if Tor is not None and not Tor["abaikan_torsi"]:
            tor_metrics = [
                (t1, "Phi.Tn (Kapasitas)", f"{Tor['PhiTn_cap']:.3f}", "kN.m"),
                (t2, "Tu (Demand)",         f"{Tor['Tu_desain']:.3f}", "kN.m"),
                (t3, "D/C Torsi",           f"{Tor['DC_torsi']:.3f}",  "-"),
                (t4, "Al perlu",            f"{Tor['Al_pakai']:.1f}",  "mm2"),
            ]
        else:
            dc_str = "DIABAIKAN" if (Tor and Tor["abaikan_torsi"]) else "-"
            tor_metrics = [
                (t1, "Phi.Tn (Kapasitas)", "-",       "kN.m"),
                (t2, "Tu (Demand)",         f"{Tu:.3f}", "kN.m"),
                (t3, "D/C Torsi",           dc_str,    "-"),
                (t4, "Al perlu",            "-",        "mm2"),
            ]
        for col, lbl, val, unt in tor_metrics:
            with col:
                st.markdown(
                    f'<div class="metric-card">'
                    f'<div class="metric-lbl">{lbl}</div>'
                    f'<div class="metric-val">{val}</div>'
                    f'<div class="metric-unt">{unt}</div>'
                    f'</div>', unsafe_allow_html=True)
        st.markdown("")

        # ---- Banner Status Lentur ----
        if R["ok_dc"] and R["et"] >= 0.005 and R["ok_rho_min"] and R["ok_rho_max"]:
            st.markdown(
                f'<div class="result-ok">[OK] LENTUR AMAN - D/C = {R["DC"]:.3f} <= 1.0 | '
                f'Tension-controlled | Memenuhi seluruh syarat SNI 2847:2019</div>',
                unsafe_allow_html=True)
        elif R["ok_dc"]:
            catatan = []
            if R["et"] < 0.005 and R["et"] >= 0.004: catatan.append("zona transisi")
            if not R["ok_rho_min"]: catatan.append("Rho < Rho-min")
            if not R["ok_rho_max"]: catatan.append("Rho > Rho-max")
            if not R["ok_et"]:      catatan.append("et < 0.004")
            note = " | ".join(catatan) if catatan else "perlu tinjauan"
            st.markdown(
                f'<div class="result-warn">[!] LENTUR AMAN secara D/C ({R["DC"]:.3f}) '
                f'namun perlu perhatian: {note}</div>',
                unsafe_allow_html=True)
        else:
            st.markdown(
                f'<div class="result-fail">[X] LENTUR TIDAK AMAN - D/C = {R["DC"]:.3f} > 1.0 | '
                f'Mu = {R["Mu"]:.2f} kN.m > Phi.Mn = {R["phiMn"]:.2f} kN.m</div>',
                unsafe_allow_html=True)

        # ---- Banner Status Geser ----
        if G["ok_total"]:
            st.markdown(
                f'<div class="result-ok">[OK] GESER AMAN - D/C = {G["DC_geser"]:.3f} <= 1.0 | '
                f'Av, spasi, dan Vs_max memenuhi syarat SNI 2847:2019</div>',
                unsafe_allow_html=True)
        elif G["ok_dc"]:
            cat_g = []
            if not G["ok_Vs_max"]: cat_g.append("Vs_perlu > Vs_max")
            if not G["ok_Av"]:     cat_g.append("Av < Av_min")
            if not G["ok_spasi"]:  cat_g.append("s > s_max")
            note_g = " | ".join(cat_g) if cat_g else "perlu tinjauan detailing"
            st.markdown(
                f'<div class="result-warn">[!] GESER kapasitas terpenuhi '
                f'(D/C = {G["DC_geser"]:.3f}) namun: {note_g}</div>',
                unsafe_allow_html=True)
        else:
            st.markdown(
                f'<div class="result-fail">[X] GESER TIDAK AMAN - '
                f'D/C = {G["DC_geser"]:.3f} > 1.0 | '
                f'Vu = {G["Vu"]:.2f} kN > Phi.Vn = {G["PhiVn_efektif"]:.2f} kN</div>',
                unsafe_allow_html=True)

        # ---- Banner Status Torsi ----
        if Tor is not None and Tor["abaikan_torsi"]:
            st.markdown(
                f'<div class="result-ok">[DIABAIKAN] TORSI - '
                f'Tu = {Tor["Tu"]:.3f} kN.m <= Phi_t x Tth = {Tor["phi_Tth"]:.3f} kN.m | '
                f'Efek torsi dapat diabaikan (SNI 22.7.4)</div>',
                unsafe_allow_html=True)
        elif Tor is not None and Tor["ok_torsi_total"]:
            st.markdown(
                f'<div class="result-ok">[OK] TORSI AMAN - D/C = {Tor["DC_torsi"]:.3f} <= 1.0 | '
                f'Semua syarat torsi terpenuhi</div>',
                unsafe_allow_html=True)
        elif Tor is not None and not Tor["ok_torsi_total"]:
            prob_t = []
            if not Tor["ok_dimensi"]:           prob_t.append("Dimensi tidak cukup")
            if not Tor["ok_sengkang_gabungan"]: prob_t.append("Sengkang perlu diperketat")
            if not Tor["ok_Avt_min"]:           prob_t.append("Av+2At < minimum")
            if not Tor["ok_spasi_torsi"]:       prob_t.append(f"s > s_max_torsi ({Tor['s_max_torsi']:.0f} mm)")
            if not Tor["ok_DC_torsi"]:          prob_t.append(f"D/C = {Tor['DC_torsi']:.3f} > 1.0")
            note_t = " | ".join(prob_t) if prob_t else "periksa semua kontrol"
            st.markdown(
                f'<div class="result-fail">[X] TORSI TIDAK AMAN - {note_t}</div>',
                unsafe_allow_html=True)

        # ---- Banner Gabungan ----
        status_lentur = "AMAN" if R["ok_dc"] else "TIDAK AMAN"
        status_geser  = "AMAN" if G["ok_total"] else "TIDAK AMAN"
        if Tor is None:
            status_torsi = "belum dievaluasi"
            aman_torsi   = True
        elif Tor["abaikan_torsi"]:
            status_torsi = "DIABAIKAN"
            aman_torsi   = True
        else:
            status_torsi = "AMAN" if Tor["ok_torsi_total"] else "TIDAK AMAN"
            aman_torsi   = Tor["ok_torsi_total"]

        aman_total = R["ok_dc"] and G["ok_total"] and aman_torsi
        css_cls = "result-ok" if aman_total else "result-fail"
        st.markdown(
            f'<div class="{css_cls}">STATUS GABUNGAN | '
            f'Lentur: {status_lentur} | Geser: {status_geser} | '
            f'Torsi: {status_torsi}</div>',
            unsafe_allow_html=True)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)

        # ---- A. DATA INPUT (Visualisasi) ----
        st.markdown("### A.  Data Input - Visualisasi Penampang")
        cvis1, cvis2 = st.columns([1, 1])
        with cvis1:
            st.pyplot(fig, use_container_width=True)
        with cvis2:
            st.markdown(f"""
            **Material:**
            - f'c = {fc:.1f} MPa
            - fy  = {fy:.0f} MPa  (lentur)
            - fyt = {fyt:.0f} MPa  (sengkang)

            **Geometri:**
            - b x h = {int(b)} x {int(h)} mm
            - cc = {int(cc_sel)} mm  |  ds = {int(ds_dia)} mm

            **Tulangan TARIK:**
            - As-total = {R['As_tarik']:.1f} mm2
            - d-aktual = {R['d_tarik']:.2f} mm

            **Tulangan TEKAN:**
            - As'-total = {R['As_tekan']:.1f} mm2
            - d'-aktual = {R['d_tekan']:.2f} mm

            **Sengkang:**
            - {int(n_kaki)} kaki D{int(ds_dia)} - {int(s_seng)} mm

            **Gaya Dalam:**
            - Mu = {Mu:.2f} kN.m
            - Vu = {Vu:.2f} kN
            - Tu = {Tu:.2f} kN.m  (tipe: {tipe_torsi})
            """)

        st.markdown('<hr class="divider">', unsafe_allow_html=True)

        # ---- B. ANALISA LENTUR ----
        st.markdown("### B.  Analisa Perhitungan - Lentur")
        for s in steps:
            warna = "#2e7d32" if s["ok"] else "#c62828"
            tanda = "[OK]" if s["ok"] else "[X]"
            st.markdown(
                f'<div class="step-box" style="border-left-color:{warna}">'
                f'<div class="ref-badge">{s["ref"]}</div><br>'
                f'<div class="step-hdr">{s["no"]} {s["judul"]} &nbsp; {tanda}</div>'
                f'<pre style="margin:0;font-size:.82rem;white-space:pre-wrap;'
                f'font-family:monospace">{s["isi"]}</pre></div>',
                unsafe_allow_html=True)

        # ---- C. RANGKUMAN ----
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("### C.  Rangkuman Hasil")
        et_s  = ("[OK] OK" if R["et"] >= 0.005
                 else "[!] Transisi" if R["et"] >= 0.004 else "[X] Tidak OK")
        rh_s  = "[OK] OK" if (R["ok_rho_min"] and R["ok_rho_max"]) else "[X] Tidak OK"
        dc_s  = "[OK] AMAN" if R["ok_dc"] else "[X] TIDAK AMAN"
        dcg_s = "[OK] AMAN" if G["ok_dc"] else "[X] TIDAK AMAN"
        av_s  = "[OK] OK"   if G["ok_Av"]  else "[X] Av < Av_min"
        sp_s  = "[OK] OK"   if G["ok_spasi"] else "[X] s > s_max"

        params_rangkuman = [
            # Lentur
            "Beta-1","c (mm)","a (mm)","Cc (kN)","Cs (kN)","T (kN)",
            "et","Phi","Rho (%)","Rho-min (%)","Rho-max (%)",
            "Mn (kN.m)","Phi.Mn (kN.m)","Mu (kN.m)","D/C Lentur",
            # Geser
            "Phi_v","Vc (kN)","Vs (kN)","Vn (kN)","Phi.Vn (kN)","Vu (kN)",
            "Av_pasang (mm2)","Av_min (mm2)","s_max (mm)","D/C Geser",
        ]
        nilai_rangkuman = [
            f"{R['beta1']:.4f}", f"{R['c']:.2f}", f"{R['a']:.2f}",
            f"{R['Cc']/1000:.2f}", f"{R['Cs']/1000:.2f}", f"{R['T']/1000:.2f}",
            f"{R['et']:.5f}", f"{R['phi']:.4f}",
            f"{R['rho']*100:.4f}", f"{R['rho_min']*100:.4f}", f"{R['rho_max']*100:.4f}",
            f"{R['Mn']:.3f}", f"{R['phiMn']:.3f}", f"{R['Mu']:.3f}", f"{R['DC']:.3f}",
            f"{G['Phi_v']:.2f}", f"{G['Vc']:.2f}", f"{G['Vs_efektif']:.2f}",
            f"{G['Vn_efektif']:.2f}", f"{G['PhiVn_efektif']:.2f}", f"{G['Vu']:.2f}",
            f"{G['Av_pasang']:.1f}", f"{G['Av_min']:.1f}", f"{G['s_max']:.1f}",
            f"{G['DC_geser']:.3f}",
        ]
        status_rangkuman = [
            "-","-","-","-","-","-",
            et_s,"-", rh_s,"-","-","-","-","-", dc_s,
            "-","-","-","-","-","-", av_s, "-", sp_s, dcg_s,
        ]

        if Tor is not None and not Tor["abaikan_torsi"]:
            params_rangkuman += [
                "Phi_t","Tth (kN.m)","Tu (kN.m)","Tu_desain (kN.m)",
                "Phi.Tn (kN.m)","D/C Torsi",
                "At/s (mm2/mm)","Al_pakai (mm2)",
                f"n_batang D{int(Tor['db_long_torsi'])}","s_max_torsi (mm)",
            ]
            dct_s = "[OK] AMAN" if Tor["ok_DC_torsi"] else "[X] TIDAK AMAN"
            nilai_rangkuman += [
                f"{Tor['Phi_t']:.2f}", f"{Tor['Tth']:.4f}", f"{Tor['Tu']:.3f}",
                f"{Tor['Tu_desain']:.4f}", f"{Tor['PhiTn_cap']:.4f}",
                f"{Tor['DC_torsi']:.3f}", f"{Tor['At_per_s']:.6f}",
                f"{Tor['Al_pakai']:.2f}", f"{Tor['n_batang']}",
                f"{Tor['s_max_torsi']:.1f}",
            ]
            status_rangkuman += ["-","-","-","-","-", dct_s,"-","-","-","-"]
        elif Tor is not None and Tor["abaikan_torsi"]:
            params_rangkuman += ["Tu (kN.m)","Tth (kN.m)","D/C Torsi"]
            nilai_rangkuman  += [f"{Tor['Tu']:.3f}", f"{Tor['Tth']:.4f}", "DIABAIKAN"]
            status_rangkuman += ["-","-","[DIABAIKAN]"]

        df = pd.DataFrame({
            "Parameter": params_rangkuman,
            "Nilai":     nilai_rangkuman,
            "Status":    status_rangkuman,
        })
        st.dataframe(df, use_container_width=True, hide_index=True)

        # ---- D. KESIMPULAN ----
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("### D.  Kesimpulan")

        cek_ringkas_items = [
            (f"Rho-min ({R['rho_min']*100:.4f}%) <= Rho ({R['rho']*100:.4f}%)", "[OK]" if R["ok_rho_min"] else "[TIDAK OK]"),
            (f"Rho-max ({R['rho_max']*100:.4f}%) >= Rho ({R['rho']*100:.4f}%)", "[OK]" if R["ok_rho_max"] else "[TIDAK OK]"),
            (f"et ({R['et']:.5f}) >= 0.004",                                     "[OK]" if R["ok_et"]      else "[TIDAK OK]"),
            (f"D/C-Lentur ({R['DC']:.3f}) <= 1.000",                             "[OK] AMAN" if R["ok_dc"] else "[X] TIDAK AMAN"),
            (f"D/C-Geser ({G['DC_geser']:.3f}) <= 1.000",                        "[OK] AMAN" if G["ok_dc"] else "[X] TIDAK AMAN"),
            (f"Vs_perlu ({G['Vs_perlu']:.2f}) <= Vs_max ({G['Vs_max']:.2f}) kN", "[OK]" if G["ok_Vs_max"] else "[TIDAK OK]"),
            (f"Av_pasang ({G['Av_pasang']:.1f}) >= Av_min ({G['Av_min']:.1f}) mm2","[OK]" if G["ok_Av"] else "[TIDAK OK]"),
            (f"s_pasang ({s_seng:.0f}) <= s_max ({G['s_max']:.1f}) mm",           "[OK]" if G["ok_spasi"] else "[TIDAK OK]"),
        ]
        if Tor is not None and not Tor["abaikan_torsi"]:
            cek_ringkas_items.extend([
                (f"D/C-Torsi ({Tor['DC_torsi']:.3f}) <= 1.000",                  "[OK] AMAN" if Tor["ok_DC_torsi"] else "[X] TIDAK AMAN"),
                ("Cek dimensi penampang (SNI 22.7.7.1)",                           "[OK]" if Tor["ok_dimensi"] else "[TIDAK OK]"),
                ("Av+2At terpasang >= Av+2At minimum",                             "[OK]" if Tor["ok_Avt_min"] else "[TIDAK OK]"),
                (f"s_pasang ({s_seng:.0f}) <= s_max_torsi ({Tor['s_max_torsi']:.1f}) mm",
                 "[OK]" if Tor["ok_spasi_torsi"] else "[TIDAK OK]"),
            ])
        elif Tor is not None and Tor["abaikan_torsi"]:
            cek_ringkas_items.append(
                (f"Tu ({Tor['Tu']:.3f}) <= Phi_t x Tth ({Tor['phi_Tth']:.3f}) kN.m",
                 "[DIABAIKAN]")
            )

        cek_ringkas = pd.DataFrame({
            "Kontrol": [x[0] for x in cek_ringkas_items],
            "Status":  [x[1] for x in cek_ringkas_items],
        })
        st.dataframe(cek_ringkas, use_container_width=True, hide_index=True)

        # ---- E. ANALISA GESER ----
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("### E.  Analisa Perhitungan - Geser")
        for s in steps_geser:
            warna = "#2e7d32" if s["ok"] else "#c62828"
            tanda = "[OK]" if s["ok"] else "[X]"
            st.markdown(
                f'<div class="step-box" style="border-left-color:{warna}">'
                f'<div class="ref-badge">{s["ref"]}</div><br>'
                f'<div class="step-hdr">{s["no"]} {s["judul"]} &nbsp; {tanda}</div>'
                f'<pre style="margin:0;font-size:.82rem;white-space:pre-wrap;'
                f'font-family:monospace">{s["isi"]}</pre></div>',
                unsafe_allow_html=True)

        # ---- F. ANALISA TORSI ----
        if steps_torsi:
            st.markdown('<hr class="divider">', unsafe_allow_html=True)
            st.markdown("### F.  Analisa Perhitungan - Torsi")
            for s in steps_torsi:
                warna = "#2e7d32" if s["ok"] else "#c62828"
                tanda = "[OK]" if s["ok"] else "[X]"
                st.markdown(
                    f'<div class="step-box" style="border-left-color:{warna}">'
                    f'<div class="ref-badge">{s["ref"]}</div><br>'
                    f'<div class="step-hdr">{s["no"]} {s["judul"]} &nbsp; {tanda}</div>'
                    f'<pre style="margin:0;font-size:.82rem;white-space:pre-wrap;'
                    f'font-family:monospace">{s["isi"]}</pre></div>',
                    unsafe_allow_html=True)

        # ---- Download ----
        st.markdown('<hr class="divider">', unsafe_allow_html=True)
        st.markdown("### Download Laporan")
        st.caption(f"Timestamp perhitungan: {st.session_state.balok_timestamp}")
        nama_f = st.session_state.balok_nama_file
        dl_w, dl_p = st.columns(2)
        with dl_w:
            st.download_button(
                label="Download Laporan Word (.docx)",
                data=st.session_state.balok_word,
                file_name=f"{nama_f}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="balok_dl_word",
            )
        with dl_p:
            st.download_button(
                label="Download Laporan PDF (.pdf)",
                data=st.session_state.balok_pdf,
                file_name=f"{nama_f}.pdf",
                mime="application/pdf",
                use_container_width=True,
                key="balok_dl_pdf",
            )
        st.caption(
            "File Word dapat diedit lebih lanjut. "
            "File PDF sudah dilengkapi watermark dan visualisasi penampang, "
            "siap untuk laporan resmi."
        )

# Footer
st.markdown('<hr class="divider">', unsafe_allow_html=True)
st.markdown(
    "<p style='text-align:center;font-size:.75rem;color:#aaa'>"
    "Referensi: SNI 2847:2019 | ACI 318-14 | "
    "Untuk keperluan profesional - verifikasi mandiri tetap diperlukan"
    "</p>",
    unsafe_allow_html=True,
)
