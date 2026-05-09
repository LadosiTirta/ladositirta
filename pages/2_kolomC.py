"""
app.py - Perhitungan Kapasitas Kolom Beton Bertulang
Diagram Interaksi P-M | SNI 2847:2019
Author: Ladosi
v2.0 - Fixed: diagram interaksi, cek kapasitas, TypeError phi_Pn_kapasitas
"""

import streamlit as st
import numpy as np
import math
import io
import pandas as pd

# ============================================================
# BAGIAN 1: LOGIKA KALKULASI
# ============================================================

def hitung_properties(b, h, c, dia_s, D, n_b, n_h, fc, fy, Es):
    Ag = b * h
    Ig = (1.0/12.0) * b * h**3
    Ec = 4700.0 * math.sqrt(fc)
    d_prime = c + dia_s + D / 2.0
    d = h - d_prime
    n_total = 2 * n_b + 2 * (n_h - 2)
    Ast = n_total * math.pi * D**2 / 4.0
    rho_g = Ast / Ag * 100.0
    cek_rasio = "OK" if 1.0 <= rho_g <= 8.0 else "NOT OK"
    return {
        "Ag": Ag, "Ig": Ig, "Ec": Ec,
        "d": d, "d_prime": d_prime,
        "n_total": n_total, "Ast": Ast,
        "rho_g": rho_g, "cek_rasio": cek_rasio
    }


def hitung_kekakuan_elemen(fc,
                            bb_al, hb_al, Lb_al,
                            bb_ar, hb_ar, Lb_ar,
                            bk_a,  hk_a,  Lk_a,
                            bb_bl, hb_bl, Lb_bl,
                            bb_br, hb_br, Lb_br,
                            bk_b,  hk_b,  Lk_b,
                            b_kol, h_kol, Lu):
    Ec = 4700.0 * math.sqrt(fc)

    def ei_l(b_e, h_e, L_e):
        if L_e <= 0:
            return 0.0
        return Ec * (1.0/12.0 * b_e * h_e**3) / L_e

    v = {
        "bal_atas_kiri":   ei_l(bb_al, hb_al, Lb_al),
        "bal_atas_kanan":  ei_l(bb_ar, hb_ar, Lb_ar),
        "kol_atas":        ei_l(bk_a,  hk_a,  Lk_a),
        "bal_bawah_kiri":  ei_l(bb_bl, hb_bl, Lb_bl),
        "bal_bawah_kanan": ei_l(bb_br, hb_br, Lb_br),
        "kol_bawah":       ei_l(bk_b,  hk_b,  Lk_b),
        "kol_ditinjau":    ei_l(b_kol, h_kol, Lu),
    }

    sum_kol_atas    = v["kol_atas"]  + v["kol_ditinjau"]
    sum_bal_atas    = v["bal_atas_kiri"] + v["bal_atas_kanan"]
    sum_kol_bawah   = v["kol_bawah"] + v["kol_ditinjau"]
    sum_bal_bawah   = v["bal_bawah_kiri"] + v["bal_bawah_kanan"]

    psi_A = sum_kol_atas  / sum_bal_atas  if sum_bal_atas  > 0 else 10.0
    psi_B = sum_kol_bawah / sum_bal_bawah if sum_bal_bawah > 0 else 10.0

    v.update({
        "sum_kol_atas": sum_kol_atas, "sum_bal_atas": sum_bal_atas,
        "sum_kol_bawah": sum_kol_bawah, "sum_bal_bawah": sum_bal_bawah,
        "psi_A": psi_A, "psi_B": psi_B,
    })
    return v


def hitung_k_kelangsingan(psi_A, psi_B, kondisi, Lu, h, M1, M2):
    psi_m = (psi_A + psi_B) / 2.0

    k_br1 = min(0.7 + 0.05 * (psi_A + psi_B), 1.0)
    k_br2 = min(0.85 + 0.05 * min(psi_A, psi_B), 1.0)
    k_braced = min(k_br1, k_br2)

    if psi_m < 2.0:
        k_unbraced = (20.0 - psi_m) / 20.0 * math.sqrt(1.0 + psi_m)
    else:
        k_unbraced = 0.9 * math.sqrt(1.0 + psi_m)

    k = k_braced if kondisi == "Braced" else k_unbraced
    r = 0.3 * h
    kLu = k * Lu
    rasio = kLu / r

    if kondisi == "Braced":
        batas = 34.0 - 12.0 * (M1 / M2) if abs(M2) > 0 else 34.0
    else:
        batas = 22.0

    klasifikasi = "Short Column" if rasio <= batas else "Slender Column"
    return {
        "k_braced": k_braced, "k_unbraced": k_unbraced,
        "k": k, "psi_m": psi_m,
        "r": r, "kLu": kLu, "rasio": rasio,
        "batas": batas, "klasifikasi": klasifikasi
    }


def hitung_pembesaran(fc, Ig, k, Lu, Pu, M2, M1, beta_dns):
    Ec = 4700.0 * math.sqrt(fc)
    EI_eff = (0.4 * Ec * Ig) / (1.0 + beta_dns)
    kLu = k * Lu
    Pc = (math.pi**2 * EI_eff) / (kLu**2) / 1000.0  # kN
    Cm = max(0.6 + 0.4 * (M1 / M2), 0.4) if abs(M2) > 0 else 1.0
    denom = 1.0 - Pu / (0.75 * Pc)
    if denom <= 0:
        denom = 1e-6
    delta_ns = max(Cm / denom, 1.0)
    Mc = delta_ns * abs(M2)
    return {
        "Ec": Ec, "EI_eff": EI_eff, "Pc": Pc,
        "Cm": Cm, "beta_dns": beta_dns,
        "delta_ns": delta_ns, "Mc": Mc
    }


def susun_layers(h, d_prime, n_b, n_h, D):
    """3 layer tulangan: tekan, tengah, tarik."""
    A_bar = math.pi * D**2 / 4.0
    n_layer1 = n_b
    n_layer2 = 2 * (n_h - 2)
    n_layer3 = n_b
    return [
        {"nama": "Layer 1 (Tekan)",  "yi": d_prime,       "n": n_layer1, "A": n_layer1 * A_bar},
        {"nama": "Layer 2 (Tengah)", "yi": h / 2.0,       "n": n_layer2, "A": n_layer2 * A_bar},
        {"nama": "Layer 3 (Tarik)",  "yi": h - d_prime,   "n": n_layer3, "A": n_layer3 * A_bar},
    ]


def phi_dinamis(eps_t, eps_y):
    if eps_t >= 0.005:
        return 0.90
    elif eps_t <= eps_y:
        return 0.65
    else:
        return 0.65 + (eps_t - eps_y) / (0.005 - eps_y) * 0.25


def hitung_diagram_interaksi(b, h, fc, fy, Es, layers, n_titik=50):
    """
    Hitung 50 titik c/h + titik 51 (tekan murni) + titik 52 (tarik murni).
    Acuan: SNI 2847:2019.
    Mn dihitung terhadap TITIK TENGAH penampang.
    """
    eps_cu = 0.003
    eps_y  = fy / Es

    # beta1
    if fc <= 28.0:
        beta1 = 0.85
    else:
        beta1 = max(0.85 - 0.05 * (fc - 28.0) / 7.0, 0.65)

    Ag  = b * h
    Ast = sum(lyr["A"] for lyr in layers)

    hasil = []
    ch_values = np.linspace(0.02, 1.20, n_titik)

    for i, ch in enumerate(ch_values):
        c = ch * h
        a = min(beta1 * c, h)

        # Gaya beton (tekan positif)
        Cc = 0.85 * fc * b * a / 1000.0  # kN
        # Lengan Cc terhadap tengah penampang
        arm_Cc = h / 2.0 - a / 2.0       # mm

        # Tulangan
        Psteel = 0.0
        Msteel = 0.0
        for lyr in layers:
            yi   = lyr["yi"]
            Ai   = lyr["A"]
            # Regangan tulangan
            eps_si = eps_cu * (c - yi) / c
            # Tegangan (dibatasi fy)
            fs_i   = max(min(eps_si * Es, fy), -fy)
            # Net: kurangi tegangan beton pada zona tekan
            if yi <= a:
                fs_net = fs_i - 0.85 * fc
            else:
                fs_net = fs_i
            Fi = fs_net * Ai / 1000.0     # kN
            Psteel += Fi
            # Lengan terhadap tengah penampang (positif ke atas)
            arm_i = h / 2.0 - yi         # mm
            Msteel += Fi * arm_i / 1000.0  # kN.m

        Pn = Cc + Psteel                  # kN  (+ = tekan)
        Mn = Cc * arm_Cc / 1000.0 + Msteel   # kN.m (selalu positif sisi tekan)
        Mn = abs(Mn)

        # Regangan tarik terluar (layer 3)
        d_tarik = layers[-1]["yi"]
        eps_t   = eps_cu * (d_tarik - c) / c

        phi    = phi_dinamis(eps_t, eps_y)
        phi_Pn = phi * Pn
        phi_Mn = phi * Mn

        hasil.append({
            "No":     i + 1,
            "c/h":    round(ch, 4),
            "c":      round(c, 2),
            "a":      round(a, 2),
            "eps_t":  round(eps_t, 5),
            "Pn":     round(Pn, 2),
            "Mn":     round(Mn, 2),
            "phi":    round(phi, 2),
            "phi_Pn": round(phi_Pn, 2),
            "phi_Mn": round(phi_Mn, 2),
        })

    # Titik 51: Tekan Murni (Po)
    Po = (0.85 * fc * (Ag - Ast) + fy * Ast) / 1000.0
    phi_51 = 0.65
    hasil.append({
        "No": 51, "c/h": "inf", "c": "inf", "a": "-", "eps_t": "-",
        "Pn": round(Po, 2), "Mn": 0,
        "phi": phi_51, "phi_Pn": round(phi_51 * Po, 2), "phi_Mn": 0,
    })

    # Titik 52: Tarik Murni (To)
    To = -fy * Ast / 1000.0
    phi_52 = 0.90
    hasil.append({
        "No": 52, "c/h": 0.0, "c": 0, "a": "-", "eps_t": "-",
        "Pn": round(To, 2), "Mn": 0,
        "phi": phi_52, "phi_Pn": round(phi_52 * To, 2), "phi_Mn": 0,
    })

    return hasil, beta1, eps_y


def data_grafik(hasil_interaksi):
    """Kembalikan list (phi_Mn, phi_Pn) untuk plotting - urut dari tekan murni ke tarik murni."""
    pts = []
    for row in hasil_interaksi:
        try:
            pts.append((float(row["phi_Mn"]), float(row["phi_Pn"])))
        except Exception:
            pass
    # Urut berdasarkan phi_Pn descending (dari atas ke bawah)
    pts.sort(key=lambda x: -x[1])
    return pts


def cek_kapasitas(hasil_interaksi, Pu, Mu):
    """
    Cek apakah (Pu, Mu) berada di dalam kurva interaksi.
    Menggunakan metode: titik di dalam polygon tertutup kurva interaksi.
    Juga hitung phi_Pn_kap dan phi_Mn_kap via interpolasi.
    """
    # Kumpulkan titik numerik kurva
    pts = []
    for row in hasil_interaksi:
        try:
            pts.append((float(row["phi_Mn"]), float(row["phi_Pn"])))
        except Exception:
            pass

    if not pts:
        return {"phi_Pn_kapasitas": None, "phi_Mn_kapasitas": None,
                "ratio_Pu": None, "ratio_Mu": None, "status": "ERROR"}

    # --- Cari phi_Mn_kapasitas saat phi_Pn = Pu ---
    # Kelompokkan titik berdasarkan phi_Pn naik
    pts_by_pn = sorted(pts, key=lambda x: x[1])  # urut phi_Pn ascending
    phi_Mn_kap = None
    # Cari interval phi_Pn yang mengapit Pu, ambil phi_Mn maksimum di sekitar itu
    # Metode lebih robust: untuk setiap level Pu, cari phi_Mn pada kurva
    # Kurva interaksi punya dua sisi: sisi kanan (phi_Mn naik s/d max) dan sisi kiri (phi_Mn turun)
    # Kita cari semua pasangan (phi_Mn, phi_Pn) yang phi_Pn mengapit Pu
    candidates_Mn = []
    for j in range(len(pts_by_pn) - 1):
        p1n, p1p = pts_by_pn[j][0], pts_by_pn[j][1]
        p2n, p2p = pts_by_pn[j+1][0], pts_by_pn[j+1][1]
        if min(p1p, p2p) <= Pu <= max(p1p, p2p):
            if abs(p2p - p1p) > 1e-9:
                t = (Pu - p1p) / (p2p - p1p)
                mn_interp = p1n + t * (p2n - p1n)
                candidates_Mn.append(mn_interp)
    phi_Mn_kap = max(candidates_Mn) if candidates_Mn else None

    # --- Cari phi_Pn_kapasitas saat phi_Mn = Mu ---
    pts_by_mn = sorted(pts, key=lambda x: x[0])  # urut phi_Mn ascending
    candidates_Pn = []
    for j in range(len(pts_by_mn) - 1):
        p1n, p1p = pts_by_mn[j][0], pts_by_mn[j][1]
        p2n, p2p = pts_by_mn[j+1][0], pts_by_mn[j+1][1]
        if min(p1n, p2n) <= Mu <= max(p1n, p2n):
            if abs(p2n - p1n) > 1e-9:
                t = (Mu - p1n) / (p2n - p1n)
                pn_interp = p1p + t * (p2p - p1p)
                candidates_Pn.append(pn_interp)
    # Ambil nilai phi_Pn kapasitas yang relevan (terbesar yang masih di sisi tekan > Pu)
    phi_Pn_kap = max(candidates_Pn) if candidates_Pn else None

    # --- Hitung rasio ---
    ratio_Pu = (Pu / phi_Pn_kap) if (phi_Pn_kap and abs(phi_Pn_kap) > 1e-6) else None
    ratio_Mu = (Mu / phi_Mn_kap) if (phi_Mn_kap and abs(phi_Mn_kap) > 1e-6) else None

    # --- Status ---
    ok_Pu = (ratio_Pu is None) or (ratio_Pu <= 1.0)
    ok_Mu = (ratio_Mu is None) or (ratio_Mu <= 1.0)
    status = "OK - AMAN" if (ok_Pu and ok_Mu) else "NOT OK - TIDAK AMAN"

    return {
        "phi_Pn_kapasitas": phi_Pn_kap,
        "phi_Mn_kapasitas": phi_Mn_kap,
        "ratio_Pu": ratio_Pu,
        "ratio_Mu": ratio_Mu,
        "status": status,
    }


# ============================================================

# ============================================================
# BAGIAN 3: GENERATOR LAPORAN PROFESIONAL
# ============================================================

def _buat_grafik_interaksi(hasil_interaksi, Pu, Mu_desain, cek):
    """Buat grafik diagram interaksi, return bytes PNG."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np

    pts = []
    for row in hasil_interaksi:
        try:
            pts.append((float(row["phi_Mn"]), float(row["phi_Pn"])))
        except Exception:
            pass
    pts.sort(key=lambda x: -x[1])

    phi_Mn_k = [p[0] for p in pts]
    phi_Pn_k = [p[1] for p in pts]

    fig, ax = plt.subplots(figsize=(10, 9))
    fig.patch.set_facecolor("white")

    # Isi area kurva
    ax.fill(phi_Mn_k, phi_Pn_k, alpha=0.08, color="royalblue")
    # Kurva
    ax.plot(phi_Mn_k, phi_Pn_k, "b-o", lw=2.2, ms=4,
            label="Kurva Interaksi φPn-φMn (52 titik)")

    # Titik beban
    is_ok = "NOT" not in cek["status"]
    warna = "red"
    ax.plot(Mu_desain, Pu, marker="*", ms=16, color=warna, zorder=5,
            label=f"Beban Aktual  Pu={Pu:.0f} kN | Mu={Mu_desain:.2f} kN.m")

    # Garis horizontal & vertikal dari titik beban
    ax.axhline(Pu,       color="gray", lw=0.8, ls="--", alpha=0.6)
    ax.axvline(Mu_desain, color="gray", lw=0.8, ls="--", alpha=0.6)
    ax.axhline(0, color="black", lw=0.9, ls="-")
    ax.axvline(0, color="black", lw=0.9, ls="-")

    # Anotasi status
    status_txt = cek["status"]
    clr_txt = "#1a7a1a" if is_ok else "#c0392b"
    ax.annotate(
        status_txt,
        xy=(Mu_desain, Pu),
        xytext=(Mu_desain + max(phi_Mn_k) * 0.05, Pu + max(phi_Pn_k) * 0.05),
        fontsize=12, color=clr_txt, fontweight="bold",
        arrowprops=dict(arrowstyle="->", color=clr_txt, lw=1.5),
        bbox=dict(boxstyle="round,pad=0.3", facecolor="white",
                  edgecolor=clr_txt, alpha=0.9),
    )

    ax.set_xlabel("φMn  (kN.m)", fontsize=12, fontweight="bold")
    ax.set_ylabel("φPn  (kN)", fontsize=12, fontweight="bold")
    ax.set_title("DIAGRAM INTERAKSI P-M  (φPn  vs  φMn)\nSNI 2847:2019",
                 fontsize=13, fontweight="bold", pad=14)
    ax.legend(fontsize=10, loc="upper right")
    ax.grid(True, alpha=0.25, lw=0.8)
    ax.tick_params(labelsize=10)
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight",
                facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


@st.cache_data
def buat_word(R, nama_eng, nama_prj, tgl_lpr, Mu_desain):
    """Laporan Word profesional: step-by-step + tabel + diagram interaksi."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import copy

    props        = R["props"]
    kekakuan     = R["kekakuan"]
    kelangsingan = R["kelangsingan"]
    pembesaran   = R["pembesaran"]
    layers       = R["layers"]
    hasil_interaksi = R["hasil_interaksi"]
    cek          = R["cek"]
    inp          = R["inp"]
    is_slender   = R["is_slender"]

    def fv(v, d=2):
        return f"{v:,.{d}f}" if v is not None else "-"

    doc = Document()
    sec = doc.sections[0]
    sec.page_width     = Cm(21)
    sec.page_height    = Cm(29.7)
    sec.left_margin    = Cm(2.5)
    sec.right_margin   = Cm(2.0)
    sec.top_margin     = Cm(2.5)
    sec.bottom_margin  = Cm(2.0)

    BIRU    = RGBColor(0, 70, 127)
    ABU     = RGBColor(80, 80, 80)
    HITAM   = RGBColor(0, 0, 0)
    HIJAU   = RGBColor(0, 120, 0)
    MERAH   = RGBColor(180, 0, 0)

    # ---- helper: heading berwarna ----
    def h1(txt):
        p = doc.add_heading("", level=1)
        run = p.add_run(txt)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(12)
        # shading biru
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "00467F")
        pPr.append(shd)
        return p

    def h2(txt):
        p = doc.add_heading("", level=2)
        run = p.add_run(txt)
        run.font.color.rgb = BIRU
        run.font.bold = True
        run.font.size = Pt(11)
        return p

    def h3(txt):
        p = doc.add_paragraph()
        run = p.add_run(txt)
        run.font.color.rgb = ABU
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.italic = True
        return p

    # ---- helper: paragraf step-by-step (Courier) ----
    def step(lines_list):
        """lines_list: list of strings, masing-masing 1 baris rumus."""
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        full = "\n".join(lines_list)
        run = p.add_run(full)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = HITAM

    # ---- helper: tabel berformat ----
    def tbl(headers, rows, shade_header=True):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        # Header
        for i, h_txt in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.text = h_txt
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Shading biru header
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "00467F")
            tc_pr.append(shd)
        # Baris data
        for ri, row_data in enumerate(rows):
            rc = t.add_row().cells
            fill = "EEF3F9" if ri % 2 == 0 else "FFFFFF"
            for i, val in enumerate(row_data):
                rc[i].text = str(val)
                for p in rc[i].paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                tc_pr = rc[i]._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  fill)
                tc_pr.append(shd)
        return t

    def garis():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pb = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "00467F")
        pb.append(bot)
        pPr.append(pb)

    def status_box(txt, ok=True):
        p = doc.add_paragraph()
        run = p.add_run(f"  {'[OK - AMAN]' if ok else '[NOT OK - TIDAK AMAN]'}  {txt}")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        fill = "1a7a1a" if ok else "c0392b"
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill)
        pPr.append(shd)

    # ==============================================================
    # COVER
    # ==============================================================
    doc.add_paragraph()
    doc.add_paragraph()

    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_logo.add_run("LAPORAN PERHITUNGAN STRUKTUR")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = BIRU

    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("KAPASITAS KOLOM BETON BERTULANG")
    r2.bold = True; r2.font.size = Pt(15); r2.font.color.rgb = HITAM

    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Diagram Interaksi P-M  |  SNI 2847:2019")
    r3.font.size = Pt(12); r3.font.color.rgb = ABU

    garis()
    doc.add_paragraph()

    # Info proyek
    tbl_cover = [
        ["Nama Proyek",  nama_prj],
        ["Dibuat Oleh",  nama_eng],
        ["Tanggal",      tgl_lpr],
        ["Standar",      "SNI 2847:2019"],
        ["Status Akhir", cek["status"]],
    ]
    t_cv = doc.add_table(rows=len(tbl_cover), cols=2)
    t_cv.style = "Table Grid"
    for ri, (k_txt, v_txt) in enumerate(tbl_cover):
        cells = t_cv.rows[ri].cells
        cells[0].text = k_txt
        cells[1].text = v_txt
        for p in cells[0].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(10)
        for p in cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if ri == 4:  # Status
                    r.bold = True
                    r.font.color.rgb = HIJAU if "AMAN" in v_txt and "NOT" not in v_txt else MERAH

    doc.add_page_break()

    # ==============================================================
    # BAB 1: DATA INPUT
    # ==============================================================
    h1("BAB 1  —  DATA INPUT PARAMETER")
    doc.add_paragraph()

    h2("1.1  Material")
    tbl(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
        [
            ["Kuat Tekan Beton",         "fc'",   f"{inp['fc']:.1f}",    "MPa",  "Mutu beton karakteristik"],
            ["Tegangan Leleh Baja",       "fy",    f"{inp['fy']:.1f}",    "MPa",  "Tulangan longitudinal"],
            ["Modulus Elastisitas Baja",  "Es",    f"{inp['Es']:.0f}",    "MPa",  "SNI 2847:2019"],
        ]
    )
    doc.add_paragraph()

    h2("1.2  Dimensi Kolom")
    tbl(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
        [
            ["Lebar Kolom",     "b",  f"{inp['b']:.0f}",     "mm",  "Dimensi arah x"],
            ["Tinggi Kolom",    "h",  f"{inp['h']:.0f}",     "mm",  "Dimensi arah y"],
            ["Selimut Beton",   "c",  f"{inp['c_sel']:.0f}", "mm",  "Jarak tepi ke tul. terluar"],
            ["Dia. Sengkang",   "Os", f"{inp['dia_s']:.0f}", "mm",  "Tulangan sengkang"],
        ]
    )
    doc.add_paragraph()

    h2("1.3  Tulangan Longitudinal")
    tbl(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
        [
            ["Diameter Tulangan",    "D",      f"{inp['D']:.0f}",          "mm",   "Tulangan utama"],
            ["Tul. Sisi Lebar (b)",  "n_b",    f"{inp['n_b']}",            "buah", "Per sisi b"],
            ["Tul. Sisi Tinggi (h)", "n_h",    f"{inp['n_h']}",            "buah", "Termasuk sudut"],
            ["Total Tulangan",       "n_total", f"{props['n_total']}",      "buah", "2xn_b + 2x(n_h-2)"],
            ["Luas Total Tulangan",  "Ast",    f"{props['Ast']:,.2f}",      "mm2",  "n_total x pi x D2/4"],
            ["Rasio Tulangan",       "rho_g",  f"{props['rho_g']:.2f}",    "%",    f"OK: 1% <= rho <= 8%"],
        ]
    )
    doc.add_paragraph()

    h2("1.4  Panjang & Kondisi Rangka")
    tbl(
        ["Parameter", "Nilai", "Satuan", "Keterangan"],
        [
            ["Panjang Tak Tertahan", f"{inp['Lu']:.0f}", "mm",  "Antar restraint"],
            ["Kondisi Rangka",       inp["kondisi_rangka"], "-", "Braced / Unbraced"],
            ["Kelengkungan",         inp["kelengkungan"],  "-",  "Single / Double"],
        ]
    )
    doc.add_paragraph()

    h2("1.5  Kekakuan Elemen Penghubung")
    tbl(
        ["Elemen", "b (mm)", "h (mm)", "L (mm)", "I = bh3/12 (mm4)", "EI/L (N.mm)"],
        [
            ["Balok Atas - Kiri",   f"{inp['bb_al']:.0f}", f"{inp['hb_al']:.0f}", f"{inp['Lb_al']:.0f}",
             f"{(1/12)*inp['bb_al']*inp['hb_al']**3:,.2f}", f"{kekakuan['bal_atas_kiri']:,.2f}"],
            ["Balok Atas - Kanan",  f"{inp['bb_ar']:.0f}", f"{inp['hb_ar']:.0f}", f"{inp['Lb_ar']:.0f}",
             f"{(1/12)*inp['bb_ar']*inp['hb_ar']**3:,.2f}", f"{kekakuan['bal_atas_kanan']:,.2f}"],
            ["Kolom Atas",          f"{inp['bk_a']:.0f}",  f"{inp['hk_a']:.0f}",  f"{inp['Lk_a']:.0f}",
             f"{(1/12)*inp['bk_a']*inp['hk_a']**3:,.2f}", f"{kekakuan['kol_atas']:,.2f}"],
            ["Balok Bawah - Kiri",  f"{inp['bb_bl']:.0f}", f"{inp['hb_bl']:.0f}", f"{inp['Lb_bl']:.0f}",
             f"{(1/12)*inp['bb_bl']*inp['hb_bl']**3:,.2f}", f"{kekakuan['bal_bawah_kiri']:,.2f}"],
            ["Balok Bawah - Kanan", f"{inp['bb_br']:.0f}", f"{inp['hb_br']:.0f}", f"{inp['Lb_br']:.0f}",
             f"{(1/12)*inp['bb_br']*inp['hb_br']**3:,.2f}", f"{kekakuan['bal_bawah_kanan']:,.2f}"],
            ["Kolom Bawah",         f"{inp['bk_b']:.0f}",  f"{inp['hk_b']:.0f}",  f"{inp['Lk_b']:.0f}",
             f"{(1/12)*inp['bk_b']*inp['hk_b']**3:,.2f}", f"{kekakuan['kol_bawah']:,.2f}"],
            ["Kolom Ditinjau",      f"{inp['b']:.0f}",     f"{inp['h']:.0f}",     f"{inp['Lu']:.0f}",
             f"{(1/12)*inp['b']*inp['h']**3:,.2f}", f"{kekakuan['kol_ditinjau']:,.2f}"],
        ]
    )
    doc.add_paragraph()

    h2("1.6  Beban Terfaktor")
    tbl(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
        [
            ["Gaya Aksial Terfaktor",  "Pu", f"{inp['Pu']:.2f}",      "kN",   "Gaya tekan terfaktor"],
            ["Momen Ujung 1 (lebih kecil)", "M1", f"{inp['M1']:.2f}", "kN.m", "Absolut"],
            ["Momen Ujung 2 (lebih besar)", "M2", f"{inp['M2']:.2f}", "kN.m", "Absolut"],
            ["Momen Desain Akhir",     "Mc", f"{Mu_desain:.4f}",      "kN.m", "Setelah pembesaran (jika Slender)"],
        ]
    )

    doc.add_page_break()

    # ==============================================================
    # BAB 2: PERHITUNGAN STEP-BY-STEP
    # ==============================================================
    h1("BAB 2  —  PROSES PERHITUNGAN  (STEP-BY-STEP)")
    doc.add_paragraph()

    # --- 2.1 Properties ---
    h2("2.1  Properti Penampang  (SNI 2847:2019 Pasal 19.2.2)")
    garis()

    h3("Luas Penampang Bruto (Ag) :")
    step([
        f"  Ag  = b  x  h",
        f"      = {inp['b']:.0f}  x  {inp['h']:.0f}",
        f"      = {props['Ag']:,.2f}  mm2",
    ])

    h3("Momen Inersia Penampang (Ig) :")
    step([
        f"  Ig  = (1/12) x b x h^3",
        f"      = (1/12) x {inp['b']:.0f} x {inp['h']:.0f}^3",
        f"      = {props['Ig']:,.2f}  mm4",
    ])

    h3("Modulus Elastisitas Beton (Ec) :")
    step([
        f"  Ec  = 4700 x SQRT(fc')",
        f"      = 4700 x SQRT({inp['fc']:.1f})",
        f"      = {props['Ec']:,.2f}  MPa",
    ])

    h3("Jarak Tulangan Tekan dari Tepi (d') :")
    step([
        f"  d'  = c + Os + D/2",
        f"      = {inp['c_sel']:.0f} + {inp['dia_s']:.0f} + {inp['D']:.0f}/2",
        f"      = {props['d_prime']:.2f}  mm",
    ])

    h3("Tinggi Efektif (d) :")
    step([
        f"  d   = h - d'",
        f"      = {inp['h']:.0f} - {props['d_prime']:.2f}",
        f"      = {props['d']:.2f}  mm",
    ])

    h3("Total Tulangan & Luas (Ast) :")
    step([
        f"  n_total = 2 x n_b  +  2 x (n_h - 2)",
        f"          = 2 x {inp['n_b']}  +  2 x ({inp['n_h']} - 2)",
        f"          = {props['n_total']}  buah",
        f"",
        f"  Ast     = n_total x pi x D^2 / 4",
        f"          = {props['n_total']} x 3.14159 x {inp['D']:.0f}^2 / 4",
        f"          = {props['Ast']:,.2f}  mm2",
    ])

    h3("Rasio Tulangan (rho_g)  [SNI 2847:2019 Ps.10.6.1] :")
    step([
        f"  rho_g = Ast / Ag x 100",
        f"        = {props['Ast']:,.2f} / {props['Ag']:,.2f} x 100",
        f"        = {props['rho_g']:.2f}%",
        f"",
        f"  Syarat: 1% <= rho_g <= 8%",
        f"  Hasil : 1% <= {props['rho_g']:.2f}% <= 8%  -->  {props['cek_rasio']}",
    ])
    doc.add_paragraph()

    # --- 2.2 Kekakuan & Psi ---
    h2("2.2  Kekakuan Elemen & Faktor Psi  (SNI 2847:2019 Pasal 6.6.4.4)")
    garis()

    h3("Rumus EI/L tiap elemen :")
    step([
        f"  EI/L = Ec x (b x h^3 / 12) / L",
        f"",
        f"  Ec  = {props['Ec']:,.2f}  MPa",
        f"",
        f"  EI/L Balok Atas Kiri  = {props['Ec']:,.2f} x ({inp['bb_al']:.0f}x{inp['hb_al']:.0f}^3/12) / {inp['Lb_al']:.0f}",
        f"                        = {kekakuan['bal_atas_kiri']:,.2f}  N.mm",
        f"",
        f"  EI/L Balok Atas Kanan = {props['Ec']:,.2f} x ({inp['bb_ar']:.0f}x{inp['hb_ar']:.0f}^3/12) / {inp['Lb_ar']:.0f}",
        f"                        = {kekakuan['bal_atas_kanan']:,.2f}  N.mm",
        f"",
        f"  EI/L Kolom Atas       = {props['Ec']:,.2f} x ({inp['bk_a']:.0f}x{inp['hk_a']:.0f}^3/12) / {inp['Lk_a']:.0f}",
        f"                        = {kekakuan['kol_atas']:,.2f}  N.mm",
        f"",
        f"  EI/L Kolom Ditinjau   = {props['Ec']:,.2f} x ({inp['b']:.0f}x{inp['h']:.0f}^3/12) / {inp['Lu']:.0f}",
        f"                        = {kekakuan['kol_ditinjau']:,.2f}  N.mm",
        f"",
        f"  EI/L Balok Bawah Kiri  = {kekakuan['bal_bawah_kiri']:,.2f}  N.mm",
        f"  EI/L Balok Bawah Kanan = {kekakuan['bal_bawah_kanan']:,.2f}  N.mm",
        f"  EI/L Kolom Bawah       = {kekakuan['kol_bawah']:,.2f}  N.mm",
    ])

    h3("Faktor Kekakuan Psi (PSI) :")
    step([
        f"  Psi_A = Sigma(EI/L)_kolom_atas / Sigma(EI/L)_balok_atas",
        f"",
        f"  Sigma(EI/L) Kolom Atas  = EI/L_kolom_atas + EI/L_kolom_ditinjau",
        f"                          = {kekakuan['kol_atas']:,.2f} + {kekakuan['kol_ditinjau']:,.2f}",
        f"                          = {kekakuan['sum_kol_atas']:,.2f}  N.mm",
        f"",
        f"  Sigma(EI/L) Balok Atas  = EI/L_bal_kiri + EI/L_bal_kanan",
        f"                          = {kekakuan['bal_atas_kiri']:,.2f} + {kekakuan['bal_atas_kanan']:,.2f}",
        f"                          = {kekakuan['sum_bal_atas']:,.2f}  N.mm",
        f"",
        f"  Psi_A = {kekakuan['sum_kol_atas']:,.2f} / {kekakuan['sum_bal_atas']:,.2f}",
        f"        = {kekakuan['psi_A']:.4f}",
        f"",
        f"  Psi_B = {kekakuan['sum_kol_bawah']:,.2f} / {kekakuan['sum_bal_bawah']:,.2f}",
        f"        = {kekakuan['psi_B']:.4f}",
        f"",
        f"  Psi_m = (Psi_A + Psi_B) / 2",
        f"        = ({kekakuan['psi_A']:.4f} + {kekakuan['psi_B']:.4f}) / 2",
        f"        = {kelangsingan['psi_m']:.4f}",
    ])
    doc.add_paragraph()

    # --- 2.3 Faktor k & Kelangsingan ---
    h2("2.3  Faktor Panjang Efektif (k) & Kelangsingan  (SNI 2847:2019 Ps.6.6.4.4 & 6.2.5)")
    garis()

    h3("Faktor k :")
    step([
        f"  Kondisi Rangka : {inp['kondisi_rangka']}",
        f"",
        f"  k_Braced :",
        f"    Cara 1 : k = 0.7 + 0.05 x (Psi_A + Psi_B)",
        f"           = 0.7 + 0.05 x ({kekakuan['psi_A']:.4f} + {kekakuan['psi_B']:.4f})",
        f"           = {0.7 + 0.05*(kekakuan['psi_A']+kekakuan['psi_B']):.4f}  -->  min(nilai, 1.0) = {min(0.7+0.05*(kekakuan['psi_A']+kekakuan['psi_B']),1.0):.4f}",
        f"",
        f"    Cara 2 : k = 0.85 + 0.05 x Psi_min",
        f"           = 0.85 + 0.05 x {min(kekakuan['psi_A'],kekakuan['psi_B']):.4f}",
        f"           = {0.85+0.05*min(kekakuan['psi_A'],kekakuan['psi_B']):.4f}  -->  min(nilai, 1.0) = {min(0.85+0.05*min(kekakuan['psi_A'],kekakuan['psi_B']),1.0):.4f}",
        f"",
        f"    k_Braced = min(Cara1, Cara2) = {kelangsingan['k_braced']:.6f}",
        f"",
        f"  k_Unbraced :",
        f"    Psi_m = {kelangsingan['psi_m']:.4f}  {'< 2 --> k=(20-Psi_m)/20 x SQRT(1+Psi_m)' if kelangsingan['psi_m']<2 else '>= 2 --> k=0.9 x SQRT(1+Psi_m)'}",
        f"    k_Unbraced = {kelangsingan['k_unbraced']:.6f}",
        f"",
        f"  k DIPAKAI ({inp['kondisi_rangka']}) = {kelangsingan['k']:.6f}",
    ])

    h3("Radius Girasi & Panjang Efektif :")
    step([
        f"  r   = 0.3 x h   (kolom persegi - SNI 2847:2019)",
        f"      = 0.3 x {inp['h']:.0f}",
        f"      = {kelangsingan['r']:.2f}  mm",
        f"",
        f"  k x Lu = {kelangsingan['k']:.6f} x {inp['Lu']:.0f}",
        f"         = {kelangsingan['kLu']:.2f}  mm",
    ])

    h3("Cek Kelangsingan :")
    step([
        f"  Rasio Kelangsingan = k x Lu / r",
        f"                     = {kelangsingan['kLu']:.2f} / {kelangsingan['r']:.2f}",
        f"                     = {kelangsingan['rasio']:.2f}",
        f"",
        f"  Batas Kelangsingan ({inp['kondisi_rangka']}) :",
        (f"    = 34 - 12 x (M1/M2)"
         f"\n    = 34 - 12 x ({inp['M1']:.0f}/{inp['M2']:.0f})"
         f"\n    = {kelangsingan['batas']:.2f}") if inp["kondisi_rangka"]=="Braced"
         else f"    = 22 (Unbraced)",
        f"",
        f"  Rasio = {kelangsingan['rasio']:.2f}  {'<=' if kelangsingan['rasio']<=kelangsingan['batas'] else '>'}  Batas = {kelangsingan['batas']:.2f}",
        f"",
        f"  KLASIFIKASI : {kelangsingan['klasifikasi']}",
    ])
    doc.add_paragraph()

    # --- 2.4 Pembesaran Momen ---
    h2("2.4  Pembesaran Momen  (SNI 2847:2019 Pasal 6.6.4)")
    garis()

    if is_slender and inp["kondisi_rangka"] == "Braced":
        h3("Kekakuan Efektif (EI)eff  [SNI Pers. 6.6.4.4.4] :")
        step([
            f"  (EI)eff = (0.4 x Ec x Ig) / (1 + beta_dns)",
            f"          = (0.4 x {pembesaran['Ec']:,.2f} x {props['Ig']:,.2f})",
            f"            / (1 + {inp['beta_dns']:.2f})",
            f"          = {pembesaran['EI_eff']:,.2f}  N.mm2",
        ])

        h3("Beban Kritis Euler (Pc)  [SNI Pers. 6.6.4.4.2] :")
        step([
            f"  Pc  = (pi^2 x (EI)eff) / (k x Lu)^2",
            f"      = (pi^2 x {pembesaran['EI_eff']:,.2f})",
            f"        / ({kelangsingan['k']:.6f} x {inp['Lu']:.0f})^2",
            f"      = {pembesaran['Pc']:,.2f}  kN",
        ])

        h3("Faktor Cm  [SNI Pers. 6.6.4.5.3] :")
        step([
            f"  Cm  = 0.6 + 0.4 x (M1/M2)  [>= 0.4]",
            f"      = 0.6 + 0.4 x ({inp['M1']:.0f}/{inp['M2']:.0f})",
            f"      = {0.6+0.4*(inp['M1']/inp['M2']):.4f}  -->  dipakai = {pembesaran['Cm']:.4f}",
        ])

        h3("Faktor Pembesaran delta_ns  [SNI Pers. 6.6.4.5.2] :")
        step([
            f"  delta_ns = Cm / (1 - Pu / (0.75 x Pc))  [>= 1.0]",
            f"           = {pembesaran['Cm']:.4f} / (1 - {inp['Pu']:.0f} / (0.75 x {pembesaran['Pc']:.2f}))",
            f"           = {pembesaran['delta_ns']:.4f}",
        ])

        h3("Momen Desain Akhir (Mc) :")
        step([
            f"  Mc  = delta_ns x M2",
            f"      = {pembesaran['delta_ns']:.4f} x {inp['M2']:.0f}",
            f"      = {pembesaran['Mc']:.4f}  kN.m",
        ])
    else:
        p_sk = doc.add_paragraph()
        run_sk = p_sk.add_run(
            f"  Kolom diklasifikasikan sebagai SHORT COLUMN\n"
            f"  (kLu/r = {kelangsingan['rasio']:.2f} <= batas = {kelangsingan['batas']:.2f})\n"
            f"  --> Tidak diperlukan pembesaran momen.\n"
            f"  Momen Desain Mu = M2 = {abs(inp['M2']):.2f} kN.m"
        )
        run_sk.font.name = "Courier New"
        run_sk.font.size = Pt(9)

    doc.add_paragraph()

    # --- 2.5 Parameter Diagram Interaksi ---
    h2("2.5  Parameter Diagram Interaksi P-M  (SNI 2847:2019 Pasal 22.2.2)")
    garis()

    beta1 = R["beta1"]
    eps_y = R["eps_y"]
    step([
        f"  Faktor beta1 :",
        f"    fc' = {inp['fc']:.1f} MPa  {'<= 28 MPa --> beta1 = 0.85' if inp['fc']<=28 else '> 28 MPa --> beta1 = 0.85 - 0.05 x (fc-28)/7'}",
        f"    beta1 = {beta1:.6f}",
        f"",
        f"  Regangan Ultimit Beton :",
        f"    eps_cu = 0.003  (SNI 2847:2019 Ps.22.2.2.1)",
        f"",
        f"  Regangan Leleh Baja :",
        f"    eps_y = fy / Es = {inp['fy']:.0f} / {inp['Es']:.0f} = {eps_y:.4f}",
        f"",
        f"  Layer Tulangan :",
        f"    Layer 1 (Tekan)  : yi = {layers[0]['yi']:.2f} mm  | A = {layers[0]['A']:,.2f} mm2 | n = {layers[0]['n']} bh",
        f"    Layer 2 (Tengah) : yi = {layers[1]['yi']:.2f} mm  | A = {layers[1]['A']:,.2f} mm2 | n = {layers[1]['n']} bh",
        f"    Layer 3 (Tarik)  : yi = {layers[2]['yi']:.2f} mm  | A = {layers[2]['A']:,.2f} mm2 | n = {layers[2]['n']} bh",
    ])
    doc.add_paragraph()

    # --- 2.6 Tabel 52 Titik ---
    h2("2.6  Tabel Diagram Interaksi P-M  (52 Titik)")
    garis()
    p_note = doc.add_paragraph()
    p_note.add_run(
        "Iterasi c/h dari 0.02 s/d 1.20 (50 titik) + Titik 51 (Tekan Murni Po) + Titik 52 (Tarik Murni To)"
    ).font.size = Pt(9)

    rows_int = []
    for row in hasil_interaksi:
        def sv(k, d=2):
            try:
                return f"{float(row[k]):,.{d}f}"
            except Exception:
                return str(row.get(k, "-"))
        rows_int.append([
            str(row["No"]),
            str(row["c/h"]),
            sv("c"),
            sv("a"),
            sv("eps_t", 5),
            sv("Pn"),
            sv("Mn"),
            str(row["phi"]),
            sv("phi_Pn"),
            sv("phi_Mn"),
        ])
    tbl(
        ["No", "c/h", "c(mm)", "a(mm)", "eps_t", "Pn(kN)", "Mn(kN.m)", "phi", "phiPn(kN)", "phiMn(kN.m)"],
        rows_int
    )

    doc.add_page_break()

    # ==============================================================
    # BAB 3: DIAGRAM INTERAKSI (GAMBAR)
    # ==============================================================
    h1("BAB 3  —  DIAGRAM INTERAKSI P-M")
    doc.add_paragraph()

    p_desc = doc.add_paragraph()
    p_desc.add_run(
        f"Diagram berikut menampilkan kurva interaksi phi.Pn vs phi.Mn berdasarkan "
        f"52 titik iterasi. Titik beban aktual (Pu={inp['Pu']:.0f} kN, "
        f"Mu={Mu_desain:.2f} kN.m) ditandai dengan bintang merah."
    ).font.size = Pt(10)

    # Sisipkan gambar
    try:
        png_bytes = _buat_grafik_interaksi(hasil_interaksi, inp["Pu"], Mu_desain, cek)
        img_stream = io.BytesIO(png_bytes)
        doc.add_picture(img_stream, width=Inches(5.8))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Grafik tidak dapat dimuat: {e}]")

    doc.add_paragraph()
    p_legend = doc.add_paragraph()
    p_legend.paragraph_format.left_indent = Cm(2)
    run_leg = p_legend.add_run(
        "Keterangan :\n"
        "  - Kurva biru       : Diagram interaksi phi.Pn - phi.Mn (52 titik)\n"
        "  - Bintang merah    : Titik beban aktual (Pu, Mu)\n"
        "  - Garis putus-putus: Proyeksi Pu dan Mu pada kurva\n"
        "  - Area biru muda   : Zona AMAN (di dalam kurva)\n"
    )
    run_leg.font.name = "Courier New"
    run_leg.font.size = Pt(9)

    doc.add_page_break()

    # ==============================================================
    # BAB 4: KESIMPULAN
    # ==============================================================
    h1("BAB 4  —  KESIMPULAN & STATUS AKHIR")
    doc.add_paragraph()

    h2("4.1  Ringkasan Kapasitas")
    garis()

    is_ok_final = "NOT" not in cek["status"]
    tbl(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Status"],
        [
            ["Klasifikasi Kolom",       "-",       kelangsingan["klasifikasi"],     "-",    "-"],
            ["Gaya Aksial Terfaktor",   "Pu",      f"{inp['Pu']:.2f}",             "kN",   "INPUT"],
            ["Momen Desain",            "Mc/Mu",   f"{Mu_desain:.4f}",             "kN.m", "INPUT"],
            ["Kap. Aksial (pd Mu)",    "phi.Pn",  fv(cek["phi_Pn_kapasitas"]),    "kN",   "-"],
            ["Kap. Momen (pd Pu)",     "phi.Mn",  fv(cek["phi_Mn_kapasitas"]),    "kN.m", "-"],
            ["Rasio Pu / phi.Pn",       "-",       fv(cek["ratio_Pu"]),            "-",    "OK" if (cek["ratio_Pu"] or 0)<=1 else "NOT OK"],
            ["Rasio Mu / phi.Mn",       "-",       fv(cek["ratio_Mu"]),            "-",    "OK" if (cek["ratio_Mu"] or 0)<=1 else "NOT OK"],
            ["STATUS AKHIR",            "-",       cek["status"],                  "-",    cek["status"]],
        ]
    )

    doc.add_paragraph()
    h2("4.2  Verifikasi")
    garis()

    step([
        f"  Pu    = {inp['Pu']:.2f}  kN",
        f"  Mc    = {Mu_desain:.4f}  kN.m",
        f"",
        f"  phi.Pn kapasitas (pada Mu) = {fv(cek['phi_Pn_kapasitas'])}  kN",
        f"  phi.Mn kapasitas (pada Pu) = {fv(cek['phi_Mn_kapasitas'])}  kN.m",
        f"",
        f"  Rasio Pu / phi.Pn = {inp['Pu']:.2f} / {fv(cek['phi_Pn_kapasitas'])}",
        f"                    = {fv(cek['ratio_Pu'])}  -->  {'<= 1.0  OK' if (cek['ratio_Pu'] or 0)<=1 else '> 1.0  NOT OK'}",
        f"",
        f"  Rasio Mu / phi.Mn = {Mu_desain:.4f} / {fv(cek['phi_Mn_kapasitas'])}",
        f"                    = {fv(cek['ratio_Mu'])}  -->  {'<= 1.0  OK' if (cek['ratio_Mu'] or 0)<=1 else '> 1.0  NOT OK'}",
    ])

    doc.add_paragraph()
    status_box(f"  Titik beban {'DI DALAM' if is_ok_final else 'DI LUAR'} kurva interaksi", ok=is_ok_final)
    doc.add_paragraph()

    p_catatan = doc.add_paragraph()
    p_catatan.paragraph_format.left_indent = Cm(1)
    run_cat = p_catatan.add_run(
        "Catatan :\n"
        f"  1. Perhitungan mengacu pada SNI 2847:2019\n"
        f"  2. Diagram Interaksi dihitung dengan 52 titik\n"
        f"     (50 titik iterasi c/h + Tekan Murni + Tarik Murni)\n"
        f"  3. Faktor reduksi phi dinamis (0.65 s/d 0.90)\n"
        f"     sesuai regangan tarik tulangan terluar\n"
        f"  4. Dibuat oleh  : {nama_eng}\n"
        f"  5. Tanggal      : {tgl_lpr}\n"
        f"  6. Proyek       : {nama_prj}"
    )
    run_cat.font.size = Pt(9)
    run_cat.font.name = "Courier New"

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@st.cache_data
def buat_pdf(R, nama_eng, nama_prj, tgl_lpr, Mu_desain):
    """Laporan PDF profesional: step-by-step + tabel + diagram interaksi."""
    from fpdf import FPDF

    props           = R["props"]
    kekakuan        = R["kekakuan"]
    kelangsingan    = R["kelangsingan"]
    pembesaran      = R["pembesaran"]
    layers          = R["layers"]
    hasil_interaksi = R["hasil_interaksi"]
    cek             = R["cek"]
    inp             = R["inp"]
    is_slender      = R["is_slender"]
    beta1           = R["beta1"]
    eps_y           = R["eps_y"]

    def sc(text):
        subs = {
            "φ":"phi","Φ":"Phi","²":"2","³":"3","√":"sqrt",
            "·":".","Ø":"O","ε":"eps","β":"beta","δ":"delta",
            "ρ":"rho","Ψ":"Psi","ψ":"psi","≤":"<=","≥":">=",
            "×":"x","π":"pi","∞":"inf","–":"-","—":"-",
            "\u2019":"'","\u2018":"'","\u201c":'"',"\u201d":'"',
        }
        for k, v in subs.items():
            text = text.replace(k, v)
        return text.encode("latin-1", errors="replace").decode("latin-1")

    def fv(v, d=2):
        return f"{v:,.{d}f}" if v is not None else "-"

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 8)
            self.set_fill_color(0, 70, 127)
            self.set_text_color(255, 255, 255)
            self.rect(0, 0, 210, 10, "F")
            self.set_xy(0, 1)
            self.cell(0, 8,
                sc(f"LAPORAN KAPASITAS KOLOM BETON BERTULANG | SNI 2847:2019  |  {nama_prj}"),
                align="C")
            self.set_text_color(0, 0, 0)
            self.ln(8)

        def footer(self):
            self.set_y(-12)
            self.set_draw_color(0, 70, 127)
            self.line(10, self.get_y(), 200, self.get_y())
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8,
                sc(f"{nama_eng}  |  {tgl_lpr}  |  Halaman {self.page_no()}"),
                align="C")
            self.set_text_color(0, 0, 0)

        def bab(self, txt):
            self.set_font("Helvetica", "B", 11)
            self.set_fill_color(0, 70, 127)
            self.set_text_color(255, 255, 255)
            self.cell(0, 8, sc(txt), fill=True, ln=True)
            self.set_text_color(0, 0, 0)
            self.ln(2)

        def sub(self, txt):
            self.set_font("Helvetica", "B", 10)
            self.set_text_color(0, 70, 127)
            self.cell(0, 6, sc(txt), ln=True)
            self.set_text_color(0, 0, 0)
            self.ln(1)

        def sub2(self, txt):
            self.set_font("Helvetica", "BI", 9)
            self.set_text_color(60, 60, 60)
            self.cell(0, 5, sc(txt), ln=True)
            self.set_text_color(0, 0, 0)

        def garis(self):
            self.set_draw_color(0, 70, 127)
            self.line(self.get_x(), self.get_y(),
                      self.get_x() + 186, self.get_y())
            self.ln(2)

        def rumus(self, lines):
            self.set_font("Courier", "", 8.5)
            self.set_fill_color(245, 248, 255)
            x0 = self.get_x() + 5
            for line in lines:
                self.set_x(x0)
                self.cell(0, 5, sc(line), ln=True, fill=True)
            self.ln(2)

        def status_box(self, txt, ok=True):
            self.set_font("Helvetica", "B", 11)
            if ok:
                self.set_fill_color(26, 122, 26)
            else:
                self.set_fill_color(192, 57, 43)
            self.set_text_color(255, 255, 255)
            self.cell(0, 10, sc(txt), fill=True, align="C", ln=True)
            self.set_text_color(0, 0, 0)
            self.ln(3)

        def tbl_head(self, cols, ws):
            self.set_font("Helvetica", "B", 8)
            self.set_fill_color(0, 70, 127)
            self.set_text_color(255, 255, 255)
            for col, w in zip(cols, ws):
                self.cell(w, 6, sc(col), border=1, fill=True, align="C")
            self.ln()
            self.set_text_color(0, 0, 0)

        def tbl_row(self, vals, ws, shade=False):
            self.set_font("Helvetica", "", 7.5)
            if shade:
                self.set_fill_color(238, 243, 249)
            else:
                self.set_fill_color(255, 255, 255)
            for val, w in zip(vals, ws):
                self.cell(w, 5.5, sc(str(val)), border=1, fill=True)
            self.ln()

    pdf = PDF()
    pdf.set_auto_page_break(True, 15)
    pdf.set_margins(12, 14, 12)

    # ==============================================================
    # COVER
    # ==============================================================
    pdf.add_page()
    pdf.ln(18)
    pdf.set_font("Helvetica", "B", 22)
    pdf.set_text_color(0, 70, 127)
    pdf.cell(0, 14, "LAPORAN PERHITUNGAN STRUKTUR", align="C", ln=True)
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(0, 0, 0)
    pdf.cell(0, 10, "KAPASITAS KOLOM BETON BERTULANG", align="C", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.set_text_color(80, 80, 80)
    pdf.cell(0, 7, "Diagram Interaksi P-M  |  SNI 2847:2019", align="C", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(10)
    pdf.set_draw_color(0, 70, 127)
    pdf.set_line_width(0.8)
    pdf.line(20, pdf.get_y(), 190, pdf.get_y())
    pdf.ln(8)

    for k_txt, v_txt in [
        ("Nama Proyek",  nama_prj),
        ("Dibuat Oleh",  nama_eng),
        ("Tanggal",      tgl_lpr),
        ("Standar",      "SNI 2847:2019"),
    ]:
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(55, 8, sc(k_txt + " :"), align="R")
        pdf.set_font("Helvetica", "", 11)
        pdf.cell(0, 8, sc(v_txt), ln=True)

    pdf.ln(10)
    is_ok_final = "NOT" not in cek["status"]
    pdf.status_box(f"  STATUS AKHIR :  {cek['status']}", ok=is_ok_final)

    # ==============================================================
    # BAB 1: DATA INPUT
    # ==============================================================
    pdf.add_page()
    pdf.bab("BAB 1  -  DATA INPUT PARAMETER")

    pdf.sub("1.1  Material")
    pdf.garis()
    ws = [70, 22, 35, 22, 37]
    pdf.tbl_head(["Parameter","Simbol","Nilai","Satuan","Keterangan"], ws)
    for ri, r in enumerate([
        ["Kuat Tekan Beton",        "fc'",  f"{inp['fc']:.1f}",  "MPa", "Mutu beton karakteristik"],
        ["Tegangan Leleh Baja",     "fy",   f"{inp['fy']:.1f}",  "MPa", "Tulangan longitudinal"],
        ["Modulus Elastisitas Baja","Es",   f"{inp['Es']:.0f}", "MPa", "SNI 2847:2019"],
    ]):
        pdf.tbl_row(r, ws, ri%2==0)
    pdf.ln(4)

    pdf.sub("1.2  Dimensi Kolom")
    pdf.garis()
    ws2 = [70, 22, 35, 22, 37]
    pdf.tbl_head(["Parameter","Simbol","Nilai","Satuan","Keterangan"], ws2)
    for ri, r in enumerate([
        ["Lebar Kolom",   "b",  f"{inp['b']:.0f}",     "mm", "Dimensi arah x"],
        ["Tinggi Kolom",  "h",  f"{inp['h']:.0f}",     "mm", "Dimensi arah y"],
        ["Selimut Beton", "c",  f"{inp['c_sel']:.0f}", "mm", "Ke tulangan terluar"],
        ["Dia. Sengkang", "Os", f"{inp['dia_s']:.0f}", "mm", "Sengkang"],
        ["Dia. Tulangan", "D",  f"{inp['D']:.0f}",     "mm", "Tulangan utama"],
        ["Tul. Sisi b",   "n_b",f"{inp['n_b']}",       "bh", "Per sisi b"],
        ["Tul. Sisi h",   "n_h",f"{inp['n_h']}",       "bh", "Incl. sudut"],
        ["Total Tul.",    "nt", f"{props['n_total']}",  "bh", "2xnb + 2x(nh-2)"],
        ["Lu",            "Lu", f"{inp['Lu']:.0f}",    "mm", "Panjang tak tertahan"],
    ]):
        pdf.tbl_row(r, ws2, ri%2==0)
    pdf.ln(4)

    pdf.sub("1.3  Kekakuan Elemen Penghubung")
    pdf.garis()
    ws3 = [44, 14, 14, 18, 40, 36]
    pdf.tbl_head(["Elemen","b(mm)","h(mm)","L(mm)","I=bh3/12 (mm4)","EI/L (N.mm)"], ws3)
    for ri, (nm, b_e, h_e, L_e, ei_val) in enumerate([
        ("Balok Atas Kiri",   inp["bb_al"],inp["hb_al"],inp["Lb_al"],kekakuan["bal_atas_kiri"]),
        ("Balok Atas Kanan",  inp["bb_ar"],inp["hb_ar"],inp["Lb_ar"],kekakuan["bal_atas_kanan"]),
        ("Kolom Atas",        inp["bk_a"], inp["hk_a"], inp["Lk_a"], kekakuan["kol_atas"]),
        ("Balok Bawah Kiri",  inp["bb_bl"],inp["hb_bl"],inp["Lb_bl"],kekakuan["bal_bawah_kiri"]),
        ("Balok Bawah Kanan", inp["bb_br"],inp["hb_br"],inp["Lb_br"],kekakuan["bal_bawah_kanan"]),
        ("Kolom Bawah",       inp["bk_b"], inp["hk_b"], inp["Lk_b"], kekakuan["kol_bawah"]),
        ("Kolom Ditinjau",    inp["b"],    inp["h"],    inp["Lu"],   kekakuan["kol_ditinjau"]),
    ]):
        I_e = (1/12)*b_e*h_e**3
        pdf.tbl_row([nm, f"{b_e:.0f}", f"{h_e:.0f}", f"{L_e:.0f}",
                     f"{I_e:,.0f}", f"{ei_val:,.2f}"], ws3, ri%2==0)
    pdf.ln(4)

    pdf.sub("1.4  Beban Terfaktor")
    pdf.garis()
    ws4 = [70, 22, 35, 22, 37]
    pdf.tbl_head(["Parameter","Simbol","Nilai","Satuan","Keterangan"], ws4)
    for ri, r in enumerate([
        ["Gaya Aksial",   "Pu", f"{inp['Pu']:.2f}",  "kN",   "Terfaktor"],
        ["Momen Ujung 1", "M1", f"{inp['M1']:.2f}",  "kN.m", "Lebih kecil"],
        ["Momen Ujung 2", "M2", f"{inp['M2']:.2f}",  "kN.m", "Lebih besar"],
        ["Momen Desain",  "Mc", f"{Mu_desain:.4f}",  "kN.m", "Setelah pembesaran"],
    ]):
        pdf.tbl_row(r, ws4, ri%2==0)

    # ==============================================================
    # BAB 2: PERHITUNGAN
    # ==============================================================
    pdf.add_page()
    pdf.bab("BAB 2  -  PROSES PERHITUNGAN STEP-BY-STEP")

    pdf.sub("2.1  Properti Penampang  (SNI 2847:2019 Ps.19.2.2)")
    pdf.garis()
    pdf.sub2("Luas Penampang Bruto (Ag) :")
    pdf.rumus([
        f"  Ag  = b x h",
        f"      = {inp['b']:.0f} x {inp['h']:.0f}",
        f"      = {props['Ag']:,.2f}  mm2",
    ])
    pdf.sub2("Momen Inersia (Ig) :")
    pdf.rumus([
        f"  Ig  = (1/12) x b x h^3",
        f"      = (1/12) x {inp['b']:.0f} x {inp['h']:.0f}^3",
        f"      = {props['Ig']:,.2f}  mm4",
    ])
    pdf.sub2("Modulus Elastisitas Beton (Ec) :")
    pdf.rumus([
        f"  Ec  = 4700 x SQRT(fc') = 4700 x SQRT({inp['fc']:.1f}) = {props['Ec']:,.2f}  MPa",
    ])
    pdf.sub2("Selimut efektif d' dan d :")
    pdf.rumus([
        f"  d'  = c + Os + D/2 = {inp['c_sel']:.0f} + {inp['dia_s']:.0f} + {inp['D']:.0f}/2 = {props['d_prime']:.2f}  mm",
        f"  d   = h - d' = {inp['h']:.0f} - {props['d_prime']:.2f} = {props['d']:.2f}  mm",
    ])
    pdf.sub2("Total tulangan & Ast :")
    pdf.rumus([
        f"  n_total = 2xn_b + 2x(n_h-2) = 2x{inp['n_b']} + 2x({inp['n_h']}-2) = {props['n_total']}  bh",
        f"  Ast     = n_total x pi x D^2/4",
        f"          = {props['n_total']} x 3.14159 x {inp['D']:.0f}^2/4",
        f"          = {props['Ast']:,.2f}  mm2",
        f"  rho_g   = Ast/Ag x 100 = {props['Ast']:,.2f}/{props['Ag']:,.2f} x 100 = {props['rho_g']:.2f}%  --> {props['cek_rasio']}",
    ])

    pdf.sub("2.2  Kekakuan Elemen & Faktor Psi  (SNI 2847:2019 Ps.6.6.4.4)")
    pdf.garis()
    pdf.rumus([
        f"  EI/L = Ec x (b x h^3/12) / L    [Ec = {props['Ec']:,.2f} MPa]",
        f"",
        f"  EI/L Balok Atas Kiri   = {kekakuan['bal_atas_kiri']:,.2f}  N.mm",
        f"  EI/L Balok Atas Kanan  = {kekakuan['bal_atas_kanan']:,.2f}  N.mm",
        f"  EI/L Kolom Atas        = {kekakuan['kol_atas']:,.2f}  N.mm",
        f"  EI/L Kolom Ditinjau    = {kekakuan['kol_ditinjau']:,.2f}  N.mm",
        f"  EI/L Balok Bawah Kiri  = {kekakuan['bal_bawah_kiri']:,.2f}  N.mm",
        f"  EI/L Balok Bawah Kanan = {kekakuan['bal_bawah_kanan']:,.2f}  N.mm",
        f"  EI/L Kolom Bawah       = {kekakuan['kol_bawah']:,.2f}  N.mm",
        f"",
        f"  Psi_A = SUM(EI/L)_kol_atas / SUM(EI/L)_bal_atas",
        f"        = {kekakuan['sum_kol_atas']:,.2f} / {kekakuan['sum_bal_atas']:,.2f}",
        f"        = {kekakuan['psi_A']:.4f}",
        f"",
        f"  Psi_B = {kekakuan['sum_kol_bawah']:,.2f} / {kekakuan['sum_bal_bawah']:,.2f}",
        f"        = {kekakuan['psi_B']:.4f}",
        f"",
        f"  Psi_m = ({kekakuan['psi_A']:.4f} + {kekakuan['psi_B']:.4f}) / 2 = {kelangsingan['psi_m']:.4f}",
    ])

    pdf.sub("2.3  Faktor k & Kelangsingan  (SNI 2847:2019 Ps.6.6.4.4 & 6.2.5)")
    pdf.garis()
    pdf.rumus([
        f"  k_Braced  Cara1 = 0.7 + 0.05x(Psi_A+Psi_B) = {min(0.7+0.05*(kekakuan['psi_A']+kekakuan['psi_B']),1.0):.4f}",
        f"  k_Braced  Cara2 = 0.85 + 0.05xPsi_min      = {min(0.85+0.05*min(kekakuan['psi_A'],kekakuan['psi_B']),1.0):.4f}",
        f"  k_Braced        = min(Cara1,Cara2)          = {kelangsingan['k_braced']:.6f}",
        f"  k_Unbraced                                  = {kelangsingan['k_unbraced']:.6f}",
        f"  k DIPAKAI ({inp['kondisi_rangka']:7s})              = {kelangsingan['k']:.6f}",
        f"",
        f"  r    = 0.3 x h = 0.3 x {inp['h']:.0f} = {kelangsingan['r']:.2f}  mm",
        f"  kxLu = {kelangsingan['k']:.6f} x {inp['Lu']:.0f} = {kelangsingan['kLu']:.2f}  mm",
        f"  kLu/r= {kelangsingan['kLu']:.2f} / {kelangsingan['r']:.2f} = {kelangsingan['rasio']:.2f}",
        f"",
        f"  Batas ({inp['kondisi_rangka']}) = 34 - 12x(M1/M2) = 34 - 12x({inp['M1']:.0f}/{inp['M2']:.0f}) = {kelangsingan['batas']:.2f}"
        if inp["kondisi_rangka"]=="Braced" else f"  Batas (Unbraced) = 22",
        f"",
        f"  kLu/r = {kelangsingan['rasio']:.2f}  {'<=' if kelangsingan['rasio']<=kelangsingan['batas'] else '>'}  Batas = {kelangsingan['batas']:.2f}",
        f"  --> {kelangsingan['klasifikasi']}",
    ])

    if is_slender and inp["kondisi_rangka"] == "Braced":
        pdf.sub("2.4  Pembesaran Momen  (SNI 2847:2019 Ps.6.6.4)")
        pdf.garis()
        pdf.rumus([
            f"  (EI)eff = 0.4 x Ec x Ig / (1 + beta_dns)       [SNI Pers.6.6.4.4.4]",
            f"          = 0.4 x {pembesaran['Ec']:,.0f} x {props['Ig']:,.0f}",
            f"            / (1 + {inp['beta_dns']:.2f})",
            f"          = {pembesaran['EI_eff']:,.2f}  N.mm2",
            f"",
            f"  Pc      = pi^2 x (EI)eff / (kxLu)^2             [SNI Pers.6.6.4.4.2]",
            f"          = pi^2 x {pembesaran['EI_eff']:,.2f}",
            f"            / ({kelangsingan['k']:.4f} x {inp['Lu']:.0f})^2",
            f"          = {pembesaran['Pc']:,.2f}  kN",
            f"",
            f"  Cm      = 0.6 + 0.4 x (M1/M2)  [>= 0.4]         [SNI Pers.6.6.4.5.3]",
            f"          = 0.6 + 0.4 x ({inp['M1']:.0f}/{inp['M2']:.0f})",
            f"          = {0.6+0.4*(inp['M1']/inp['M2']):.4f}  --> dipakai = {pembesaran['Cm']:.4f}",
            f"",
            f"  delta_ns= Cm / (1 - Pu/(0.75xPc))  [>= 1.0]     [SNI Pers.6.6.4.5.2]",
            f"          = {pembesaran['Cm']:.4f} / (1 - {inp['Pu']:.0f}/(0.75x{pembesaran['Pc']:.2f}))",
            f"          = {pembesaran['delta_ns']:.4f}",
            f"",
            f"  Mc      = delta_ns x M2",
            f"          = {pembesaran['delta_ns']:.4f} x {inp['M2']:.0f}",
            f"          = {pembesaran['Mc']:.4f}  kN.m",
        ])
    else:
        pdf.rumus([
            f"  Short Column -> tidak perlu pembesaran momen",
            f"  Mu = M2 = {abs(inp['M2']):.2f}  kN.m",
        ])

    pdf.sub("2.5  Parameter Diagram Interaksi")
    pdf.garis()
    pdf.rumus([
        f"  beta1   = {'0.85' if inp['fc']<=28 else '0.85 - 0.05x(fc-28)/7'} = {beta1:.6f}",
        f"  eps_cu  = 0.003  (SNI 2847:2019 Ps.22.2.2.1)",
        f"  eps_y   = fy/Es = {inp['fy']:.0f}/{inp['Es']:.0f} = {eps_y:.4f}",
        f"",
        f"  Layer 1 (Tekan)  : yi={layers[0]['yi']:.2f} mm | n={layers[0]['n']} bh | A={layers[0]['A']:,.2f} mm2",
        f"  Layer 2 (Tengah) : yi={layers[1]['yi']:.2f} mm | n={layers[1]['n']} bh | A={layers[1]['A']:,.2f} mm2",
        f"  Layer 3 (Tarik)  : yi={layers[2]['yi']:.2f} mm | n={layers[2]['n']} bh | A={layers[2]['A']:,.2f} mm2",
    ])

    # Tabel 52 titik
    pdf.sub("2.6  Tabel Diagram Interaksi (52 Titik)")
    pdf.garis()
    wt = [10, 17, 18, 18, 17, 22, 22, 10, 22, 22]
    pdf.tbl_head(["No","c/h","c(mm)","a(mm)","eps_t","Pn(kN)","Mn(kN.m)","phi","phiPn","phiMn"], wt)

    for ri, row in enumerate(hasil_interaksi):
        def sv(k, d=2):
            try:
                return f"{float(row[k]):,.{d}f}"
            except Exception:
                return str(row.get(k, "-"))
        pdf.tbl_row([
            str(row["No"]), str(row["c/h"]),
            sv("c"), sv("a"), sv("eps_t", 5),
            sv("Pn"), sv("Mn"), str(row["phi"]),
            sv("phi_Pn"), sv("phi_Mn"),
        ], wt, ri % 2 == 0)

    # ==============================================================
    # BAB 3: DIAGRAM INTERAKSI
    # ==============================================================
    pdf.add_page()
    pdf.bab("BAB 3  -  DIAGRAM INTERAKSI P-M")
    pdf.ln(3)

    try:
        png_bytes = _buat_grafik_interaksi(hasil_interaksi, inp["Pu"], Mu_desain, cek)
        img_buf = io.BytesIO(png_bytes)
        import tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            tmp.write(png_bytes)
            tmp_path = tmp.name
        pdf.image(tmp_path, x=15, y=pdf.get_y(), w=175)
        os.unlink(tmp_path)
        pdf.ln(138)
    except Exception as e:
        pdf.set_font("Helvetica", "", 9)
        pdf.cell(0, 8, sc(f"[Grafik tidak dapat ditampilkan: {e}]"), ln=True)

    pdf.set_font("Courier", "", 8.5)
    pdf.set_fill_color(245, 248, 255)
    for line in [
        "  Keterangan :",
        "  - Kurva biru       : Diagram interaksi phi.Pn - phi.Mn (52 titik)",
        "  - Bintang merah    : Titik beban aktual (Pu, Mu)",
        "  - Garis putus-putus: Proyeksi Pu dan Mu pada kurva",
        "  - Area biru muda   : Zona AMAN (di dalam kurva)",
    ]:
        pdf.cell(0, 5, sc(line), ln=True, fill=True)

    # ==============================================================
    # BAB 4: KESIMPULAN
    # ==============================================================
    pdf.add_page()
    pdf.bab("BAB 4  -  KESIMPULAN & STATUS AKHIR")
    pdf.ln(3)

    pdf.sub("4.1  Ringkasan Kapasitas")
    pdf.garis()
    ws5 = [58, 22, 35, 20, 51]
    pdf.tbl_head(["Parameter","Simbol","Nilai","Satuan","Status"], ws5)
    for ri, r in enumerate([
        ["Klasifikasi Kolom",     "-",      kelangsingan["klasifikasi"],     "-",    "-"],
        ["Gaya Aksial Pu",        "Pu",     f"{inp['Pu']:.2f}",             "kN",   "INPUT"],
        ["Momen Desain Mc/Mu",    "Mc",     f"{Mu_desain:.4f}",             "kN.m", "INPUT"],
        ["Kap. Aksial (pd Mu)",   "phi.Pn", fv(cek["phi_Pn_kapasitas"]),   "kN",   "-"],
        ["Kap. Momen  (pd Pu)",   "phi.Mn", fv(cek["phi_Mn_kapasitas"]),   "kN.m", "-"],
        ["Rasio Pu / phi.Pn",     "-",      fv(cek["ratio_Pu"]),            "-",    "OK" if (cek["ratio_Pu"] or 0)<=1 else "NOT OK"],
        ["Rasio Mu / phi.Mn",     "-",      fv(cek["ratio_Mu"]),            "-",    "OK" if (cek["ratio_Mu"] or 0)<=1 else "NOT OK"],
        ["STATUS AKHIR",          "-",      cek["status"],                  "-",    cek["status"]],
    ]):
        pdf.tbl_row(r, ws5, ri%2==0)

    pdf.ln(5)
    pdf.sub("4.2  Verifikasi")
    pdf.garis()
    pdf.rumus([
        f"  Pu   = {inp['Pu']:.2f}  kN",
        f"  Mc   = {Mu_desain:.4f}  kN.m",
        f"",
        f"  phi.Pn kapasitas (saat Mu={Mu_desain:.2f} kN.m) = {fv(cek['phi_Pn_kapasitas'])}  kN",
        f"  phi.Mn kapasitas (saat Pu={inp['Pu']:.2f} kN)   = {fv(cek['phi_Mn_kapasitas'])}  kN.m",
        f"",
        f"  Rasio Pu / phi.Pn = {inp['Pu']:.2f} / {fv(cek['phi_Pn_kapasitas'])}",
        f"                    = {fv(cek['ratio_Pu'])}  -->  {'<= 1.0  --> OK' if (cek['ratio_Pu'] or 0)<=1 else '> 1.0   --> NOT OK'}",
        f"",
        f"  Rasio Mu / phi.Mn = {Mu_desain:.4f} / {fv(cek['phi_Mn_kapasitas'])}",
        f"                    = {fv(cek['ratio_Mu'])}  -->  {'<= 1.0  --> OK' if (cek['ratio_Mu'] or 0)<=1 else '> 1.0   --> NOT OK'}",
        f"",
        f"  Titik beban {'DI DALAM' if is_ok_final else 'DI LUAR'} kurva interaksi",
    ])

    pdf.ln(5)
    pdf.status_box(f"  STATUS AKHIR  :  {cek['status']}", ok=is_ok_final)

    pdf.ln(5)
    pdf.set_font("Courier", "", 8.5)
    for line in [
        "  Catatan :",
        f"  1. Perhitungan mengacu SNI 2847:2019",
        f"  2. Diagram 52 titik (50 iter. c/h + Tekan Murni + Tarik Murni)",
        f"  3. Faktor phi dinamis 0.65 s/d 0.90 (sesuai regangan tarik terluar)",
        f"  4. Dibuat oleh  : {nama_eng}",
        f"  5. Tanggal      : {tgl_lpr}",
        f"  6. Proyek       : {nama_prj}",
    ]:
        pdf.cell(0, 5, sc(line), ln=True)

    return bytes(pdf.output())



# BAGIAN 2: UI STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Kolom Beton Bertulang | SNI 2847:2019",
    page_icon="🏛️",
    layout="wide",
)

st.title("🏛️ Perhitungan Kapasitas Kolom Beton Bertulang")
st.markdown("**Diagram Interaksi P-M | SNI 2847:2019** | Dibuat oleh: Ladosi")
st.divider()

# ---- SIDEBAR ----
with st.sidebar:
    st.header("📋 Data Input")

    st.subheader("1. Material")
    fc    = st.number_input("Mutu Beton fc' (MPa)", value=30.0,     step=1.0,    min_value=17.0)
    fy    = st.number_input("Mutu Baja fy (MPa)",   value=400.0,    step=10.0,   min_value=200.0)
    Es    = st.number_input("Modulus Baja Es (MPa)", value=200000.0, step=1000.0, min_value=100000.0)

    st.subheader("2. Dimensi Kolom")
    b      = st.number_input("Lebar b (mm)",          value=400.0, step=10.0, min_value=100.0)
    h      = st.number_input("Tinggi h (mm)",          value=500.0, step=10.0, min_value=100.0)
    c_sel  = st.number_input("Selimut Beton c (mm)",   value=40.0,  step=5.0,  min_value=20.0)
    dia_s  = st.number_input("Diameter Sengkang Øs (mm)", value=10.0, step=2.0, min_value=6.0)

    st.subheader("3. Tulangan Longitudinal")
    D   = st.number_input("Diameter Tulangan D (mm)", value=22.0, step=2.0,  min_value=10.0)
    n_b = st.number_input("Tulangan Sisi b (n_b)",    value=4,    step=1,    min_value=2)
    n_h = st.number_input("Tulangan Sisi h (n_h, incl. sudut)", value=3, step=1, min_value=2)

    st.subheader("4. Panjang & Kondisi Rangka")
    Lu             = st.number_input("Panjang Tak Tertahan Lu (mm)", value=6000.0, step=100.0, min_value=500.0)
    kondisi_rangka = st.selectbox("Kondisi Rangka", ["Braced", "Unbraced"])
    kelengkungan   = st.selectbox("Kelengkungan", ["Single", "Double"])

    st.subheader("5. Kekakuan Elemen Penghubung")
    st.markdown("**Balok Atas** *(0 jika tidak ada)*")
    bb_al = st.number_input("Bal. Atas-Kiri b (mm)",  value=300.0, key="bbal_b")
    hb_al = st.number_input("Bal. Atas-Kiri h (mm)",  value=500.0, key="bbal_h")
    Lb_al = st.number_input("Bal. Atas-Kiri L (mm)",  value=6000.0, key="bbal_L")
    bb_ar = st.number_input("Bal. Atas-Kanan b (mm)", value=300.0, key="bbar_b")
    hb_ar = st.number_input("Bal. Atas-Kanan h (mm)", value=500.0, key="bbar_h")
    Lb_ar = st.number_input("Bal. Atas-Kanan L (mm)", value=6000.0, key="bbar_L")

    st.markdown("**Kolom Atas** *(L=0 jika lantai atap)*")
    bk_a = st.number_input("Kol. Atas b (mm)", value=400.0, key="ka_b")
    hk_a = st.number_input("Kol. Atas h (mm)", value=500.0, key="ka_h")
    Lk_a = st.number_input("Kol. Atas L (mm) [0=tidak ada]", value=3500.0, key="ka_L")

    st.markdown("**Balok Bawah** *(0 jika tidak ada)*")
    bb_bl = st.number_input("Bal. Bawah-Kiri b (mm)",  value=300.0, key="bbbl_b")
    hb_bl = st.number_input("Bal. Bawah-Kiri h (mm)",  value=500.0, key="bbbl_h")
    Lb_bl = st.number_input("Bal. Bawah-Kiri L (mm)",  value=6000.0, key="bbbl_L")
    bb_br = st.number_input("Bal. Bawah-Kanan b (mm)", value=300.0, key="bbbr_b")
    hb_br = st.number_input("Bal. Bawah-Kanan h (mm)", value=500.0, key="bbbr_h")
    Lb_br = st.number_input("Bal. Bawah-Kanan L (mm)", value=6000.0, key="bbbr_L")

    st.markdown("**Kolom Bawah** *(L=0 jika langsung pondasi)*")
    bk_b = st.number_input("Kol. Bawah b (mm)", value=400.0, key="kb_b")
    hk_b = st.number_input("Kol. Bawah h (mm)", value=500.0, key="kb_h")
    Lk_b = st.number_input("Kol. Bawah L (mm) [0=pondasi]", value=3500.0, key="kb_L")
    st.caption("ℹ️ Kolom bawah L=0 → ΨB = 0 (jepit di pondasi)")

    st.subheader("6. Beban Terfaktor")
    Pu  = st.number_input("Gaya Aksial Pu (kN)",             value=1500.0, step=10.0)
    M1  = st.number_input("Momen Ujung 1 M1 (kN.m) - lebih kecil", value=80.0, step=5.0)
    M2  = st.number_input("Momen Ujung 2 M2 (kN.m) - lebih besar", value=150.0, step=5.0, min_value=0.1)

    st.subheader("7. Faktor Tambahan")
    beta_dns = st.number_input("βdns (rasio beban tetap/total)", value=0.60, step=0.05,
                               min_value=0.0, max_value=1.0)

    st.divider()
    btn_hitung = st.button("🔴 HITUNG SEKARANG", use_container_width=True, type="primary")


# ============================================================
# PROSES KALKULASI
# ============================================================

if btn_hitung:
    # Tanda M1: Double curvature = berlawanan arah (positif), Single = searah (negatif untuk rumus Cm)
    M1_signed = M1 if kelengkungan == "Double" else -M1

    props = hitung_properties(b, h, c_sel, dia_s, D, n_b, n_h, fc, fy, Es)

    kekakuan = hitung_kekakuan_elemen(
        fc,
        bb_al, hb_al, Lb_al,
        bb_ar, hb_ar, Lb_ar,
        bk_a,  hk_a,  Lk_a,
        bb_bl, hb_bl, Lb_bl,
        bb_br, hb_br, Lb_br,
        bk_b,  hk_b,  Lk_b,
        b,     h,     Lu,
    )

    kelangsingan = hitung_k_kelangsingan(
        kekakuan["psi_A"], kekakuan["psi_B"],
        kondisi_rangka, Lu, h, M1_signed, M2,
    )

    pembesaran = hitung_pembesaran(
        fc, props["Ig"], kelangsingan["k"],
        Lu, Pu, M2, M1_signed, beta_dns,
    )

    is_slender = (kelangsingan["klasifikasi"] == "Slender Column")
    Mu_desain  = pembesaran["Mc"] if (is_slender and kondisi_rangka == "Braced") else abs(M2)

    layers = susun_layers(h, props["d_prime"], n_b, n_h, D)

    hasil_interaksi, beta1, eps_y = hitung_diagram_interaksi(b, h, fc, fy, Es, layers)

    cek = cek_kapasitas(hasil_interaksi, Pu, Mu_desain)

    st.session_state["R"] = {
        "props":           props,
        "kekakuan":        kekakuan,
        "kelangsingan":    kelangsingan,
        "pembesaran":      pembesaran,
        "layers":          layers,
        "hasil_interaksi": hasil_interaksi,
        "cek":             cek,
        "beta1":           beta1,
        "eps_y":           eps_y,
        "Mu_desain":       Mu_desain,
        "is_slender":      is_slender,
        "inp": {
            "fc": fc, "fy": fy, "Es": Es,
            "b": b, "h": h, "c_sel": c_sel, "dia_s": dia_s,
            "D": D, "n_b": n_b, "n_h": n_h,
            "Lu": Lu, "kondisi_rangka": kondisi_rangka, "kelengkungan": kelengkungan,
            "Pu": Pu, "M1": M1, "M2": M2, "beta_dns": beta_dns,
            # Elemen kekakuan
            "bb_al": bb_al, "hb_al": hb_al, "Lb_al": Lb_al,
            "bb_ar": bb_ar, "hb_ar": hb_ar, "Lb_ar": Lb_ar,
            "bk_a":  bk_a,  "hk_a":  hk_a,  "Lk_a":  Lk_a,
            "bb_bl": bb_bl, "hb_bl": hb_bl, "Lb_bl": Lb_bl,
            "bb_br": bb_br, "hb_br": hb_br, "Lb_br": Lb_br,
            "bk_b":  bk_b,  "hk_b":  hk_b,  "Lk_b":  Lk_b,
        },
    }


# ============================================================
# TAMPILKAN HASIL
# ============================================================

if "R" not in st.session_state:
    st.info("👈 Isi data input di sidebar, lalu klik **HITUNG SEKARANG** [ POJOK KIRI ATAS LOGO " >> " ]")
    st.stop()

R            = st.session_state["R"]
props        = R["props"]
kekakuan     = R["kekakuan"]
kelangsingan = R["kelangsingan"]
pembesaran   = R["pembesaran"]
layers       = R["layers"]
hasil_interaksi = R["hasil_interaksi"]
cek          = R["cek"]
beta1        = R["beta1"]
eps_y        = R["eps_y"]
Mu_desain    = R["Mu_desain"]
is_slender   = R["is_slender"]
inp          = R["inp"]

# --- STATUS BANNER ---
is_ok = ("NOT" not in cek["status"])
banner_color = "#1a7a1a" if is_ok else "#c0392b"
st.markdown(
    f'<div style="background:{banner_color};padding:16px;border-radius:10px;text-align:center;">'
    f'<h2 style="color:white;margin:0;">🏛️ {cek["status"]}</h2></div>',
    unsafe_allow_html=True,
)
st.write("")

tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "📐 Properties", "📏 Kelangsingan", "📈 Diagram Interaksi",
    "✅ Cek Kapasitas", "📄 Laporan",
])


# ============================================================
# TAB 1: PROPERTIES
# ============================================================
with tab1:
    st.subheader("7. Properti Penampang (SNI 2847:2019 Pasal 19.2.2)")
    col_a, col_b = st.columns(2)

    with col_a:
        df_p = pd.DataFrame({
            "Parameter": ["Ag", "Ig", "Ec", "d'", "d", "n_total", "Ast", "ρg"],
            "Rumus": [
                "b × h",
                "(1/12) × b × h³",
                "4700 × √fc'",
                "c + Øs + D/2",
                "h - d'",
                "2×n_b + 2×(n_h-2)",
                "n_total × π × D²/4",
                "Ast/Ag × 100",
            ],
            "Nilai": [
                f"{props['Ag']:,.2f}",
                f"{props['Ig']:,.2f}",
                f"{props['Ec']:,.2f}",
                f"{props['d_prime']:.2f}",
                f"{props['d']:.2f}",
                f"{props['n_total']}",
                f"{props['Ast']:,.2f}",
                f"{props['rho_g']:.2f}",
            ],
            "Satuan": ["mm²", "mm⁴", "MPa", "mm", "mm", "bh", "mm²", "%"],
        })
        st.dataframe(df_p, use_container_width=True, hide_index=True)

        cek_color = "green" if props["cek_rasio"] == "OK" else "red"
        st.markdown(
            f"**Cek Rasio Tulangan (SNI Ps.10.6.1):** "
            f"1% ≤ ρg = **{props['rho_g']:.2f}%** ≤ 8% → "
            f"<span style='color:{cek_color};font-weight:bold;'>{props['cek_rasio']}</span>",
            unsafe_allow_html=True,
        )

    with col_b:
        st.code(
            f"Ag  = {inp['b']:.0f} × {inp['h']:.0f}\n"
            f"    = {props['Ag']:,.2f} mm²\n\n"
            f"Ig  = (1/12) × {inp['b']:.0f} × {inp['h']:.0f}³\n"
            f"    = {props['Ig']:,.2f} mm⁴\n\n"
            f"Ec  = 4700 × √{inp['fc']:.0f}\n"
            f"    = {props['Ec']:,.2f} MPa\n\n"
            f"d'  = {inp['c_sel']:.0f} + {inp['dia_s']:.0f} + {inp['D']:.0f}/2\n"
            f"    = {props['d_prime']:.2f} mm\n\n"
            f"d   = {inp['h']:.0f} - {props['d_prime']:.2f}\n"
            f"    = {props['d']:.2f} mm\n\n"
            f"n   = 2×{inp['n_b']} + 2×({inp['n_h']}-2) = {props['n_total']} bh\n\n"
            f"Ast = {props['n_total']} × π × {inp['D']:.0f}²/4\n"
            f"    = {props['Ast']:,.2f} mm²\n\n"
            f"ρg  = {props['Ast']:.2f}/{props['Ag']:.2f} × 100\n"
            f"    = {props['rho_g']:.2f}%  → {props['cek_rasio']}",
            language="text",
        )

    st.divider()
    st.subheader("9. Kekakuan Elemen (EI/L) — SNI 2847:2019 Pasal 6.6.3")
    df_kek = pd.DataFrame({
        "Elemen": [
            "Balok Atas - Kiri", "Balok Atas - Kanan", "Kolom Atas",
            "Balok Bawah - Kiri", "Balok Bawah - Kanan", "Kolom Bawah",
            "Kolom Ditinjau",
        ],
        "b (mm)": [inp["bb_al"], inp["bb_ar"], inp["bk_a"],
                   inp["bb_bl"], inp["bb_br"], inp["bk_b"], inp["b"]],
        "h (mm)": [inp["hb_al"], inp["hb_ar"], inp["hk_a"],
                   inp["hb_bl"], inp["hb_br"], inp["hk_b"], inp["h"]],
        "L (mm)": [inp["Lb_al"], inp["Lb_ar"], inp["Lk_a"],
                   inp["Lb_bl"], inp["Lb_br"], inp["Lk_b"], inp["Lu"]],
        "EI/L (N.mm)": [
            f"{kekakuan['bal_atas_kiri']:,.2f}",
            f"{kekakuan['bal_atas_kanan']:,.2f}",
            f"{kekakuan['kol_atas']:,.2f}",
            f"{kekakuan['bal_bawah_kiri']:,.2f}",
            f"{kekakuan['bal_bawah_kanan']:,.2f}",
            f"{kekakuan['kol_bawah']:,.2f}",
            f"{kekakuan['kol_ditinjau']:,.2f}",
        ],
    })
    st.dataframe(df_kek, use_container_width=True, hide_index=True)


# ============================================================
# TAB 2: KELANGSINGAN
# ============================================================
with tab2:
    st.subheader("10. Faktor Ψ (Psi) — SNI 2847:2019 Pasal 6.6.4.4")
    col_a, col_b = st.columns(2)
    with col_a:
        df_psi = pd.DataFrame({
            "Parameter": [
                "Σ(EI/L) Kolom Atas", "Σ(EI/L) Balok Atas", "ΨA",
                "Σ(EI/L) Kolom Bawah", "Σ(EI/L) Balok Bawah", "ΨB", "Ψm",
            ],
            "Nilai": [
                f"{kekakuan['sum_kol_atas']:,.2f}",
                f"{kekakuan['sum_bal_atas']:,.2f}",
                f"{kekakuan['psi_A']:.4f}",
                f"{kekakuan['sum_kol_bawah']:,.2f}",
                f"{kekakuan['sum_bal_bawah']:,.2f}",
                f"{kekakuan['psi_B']:.4f}",
                f"{kelangsingan['psi_m']:.4f}",
            ],
            "Satuan": ["N.mm", "N.mm", "-", "N.mm", "N.mm", "-", "-"],
        })
        st.dataframe(df_psi, use_container_width=True, hide_index=True)

    with col_b:
        st.code(
            f"ΨA = Σ(EI/L)kolom / Σ(EI/L)balok\n"
            f"   = {kekakuan['sum_kol_atas']:,.2f}\n"
            f"     / {kekakuan['sum_bal_atas']:,.2f}\n"
            f"   = {kekakuan['psi_A']:.4f}\n\n"
            f"ΨB = {kekakuan['sum_kol_bawah']:,.2f}\n"
            f"     / {kekakuan['sum_bal_bawah']:,.2f}\n"
            f"   = {kekakuan['psi_B']:.4f}\n\n"
            f"Ψm = (ΨA + ΨB) / 2\n"
            f"   = ({kekakuan['psi_A']:.4f} + {kekakuan['psi_B']:.4f}) / 2\n"
            f"   = {kelangsingan['psi_m']:.4f}",
            language="text",
        )

    st.divider()
    st.subheader("11 & 12. Faktor k dan Cek Kelangsingan — SNI 2847:2019 Pasal 6.6.4.4 & 6.2.5")
    col_a, col_b = st.columns(2)
    with col_a:
        df_k = pd.DataFrame({
            "Parameter": [
                "k (Braced)", "k (Unbraced)", "k Dipakai",
                "r = 0.3h", "k×Lu", "k×Lu/r",
                "Batas Kelangsingan", "Klasifikasi",
            ],
            "Nilai": [
                f"{kelangsingan['k_braced']:.6f}",
                f"{kelangsingan['k_unbraced']:.6f}",
                f"{kelangsingan['k']:.6f}",
                f"{kelangsingan['r']:.2f} mm",
                f"{kelangsingan['kLu']:.2f} mm",
                f"{kelangsingan['rasio']:.2f}",
                f"{kelangsingan['batas']:.2f}",
                kelangsingan["klasifikasi"],
            ],
        })
        st.dataframe(df_k, use_container_width=True, hide_index=True)

    with col_b:
        st.code(
            f"k_Braced  = min(\n"
            f"  0.7 + 0.05(ΨA+ΨB) = {0.7+0.05*(kekakuan['psi_A']+kekakuan['psi_B']):.4f} → ≤1.0\n"
            f"  0.85 + 0.05×Ψmin  = {0.85+0.05*min(kekakuan['psi_A'],kekakuan['psi_B']):.4f} → ≤1.0\n"
            f") = {kelangsingan['k_braced']:.6f}\n\n"
            f"r   = 0.3 × {inp['h']:.0f} = {kelangsingan['r']:.2f} mm\n\n"
            f"k×Lu = {kelangsingan['k']:.6f} × {inp['Lu']:.0f}\n"
            f"     = {kelangsingan['kLu']:.2f} mm\n\n"
            f"k×Lu/r = {kelangsingan['kLu']:.2f} / {kelangsingan['r']:.2f}\n"
            f"       = {kelangsingan['rasio']:.2f}\n\n"
            f"Batas ({inp['kondisi_rangka']}) = {kelangsingan['batas']:.2f}\n"
            f"→ {kelangsingan['klasifikasi']}",
            language="text",
        )

    st.divider()
    st.subheader("13. Pembesaran Momen — SNI 2847:2019 Pasal 6.6.4")
    if is_slender and inp["kondisi_rangka"] == "Braced":
        col_a, col_b = st.columns(2)
        with col_a:
            df_pem = pd.DataFrame({
                "Parameter": ["(EI)eff", "Pc", "Cm", "βdns", "δns", "Mc"],
                "Nilai": [
                    f"{pembesaran['EI_eff']:,.2f}",
                    f"{pembesaran['Pc']:,.2f}",
                    f"{pembesaran['Cm']:.4f}",
                    f"{pembesaran['beta_dns']:.2f}",
                    f"{pembesaran['delta_ns']:.4f}",
                    f"{pembesaran['Mc']:,.4f}",
                ],
                "Satuan": ["N.mm²", "kN", "-", "-", "-", "kN.m"],
                "Referensi SNI": [
                    "Pers. 6.6.4.4.4", "Pers. 6.6.4.4.2",
                    "Pers. 6.6.4.5.3", "-",
                    "Pers. 6.6.4.5.2", "Mc = δns×M2",
                ],
            })
            st.dataframe(df_pem, use_container_width=True, hide_index=True)
        with col_b:
            st.code(
                f"(EI)eff = 0.4×Ec×Ig / (1+βdns)\n"
                f"        = 0.4×{pembesaran['Ec']:,.0f}×{props['Ig']:,.0f}\n"
                f"          / (1+{pembesaran['beta_dns']:.2f})\n"
                f"        = {pembesaran['EI_eff']:,.2f} N.mm²\n\n"
                f"Pc      = π²×(EI)eff / (k×Lu)²\n"
                f"        = {pembesaran['Pc']:,.2f} kN\n\n"
                f"Cm      = 0.6 + 0.4×(M1/M2)\n"
                f"        = 0.6 + 0.4×({inp['M1']:.0f}/{inp['M2']:.0f})\n"
                f"        = {pembesaran['Cm']:.4f}  [≥0.4]\n\n"
                f"δns     = Cm / (1 - Pu/(0.75×Pc))\n"
                f"        = {pembesaran['Cm']:.4f} / (1 - {inp['Pu']:.0f}/(0.75×{pembesaran['Pc']:.2f}))\n"
                f"        = {pembesaran['delta_ns']:.4f}  [≥1.0]\n\n"
                f"Mc      = δns × M2\n"
                f"        = {pembesaran['delta_ns']:.4f} × {inp['M2']:.0f}\n"
                f"        = {pembesaran['Mc']:,.4f} kN.m",
                language="text",
            )
    else:
        st.info(
            f"**Short Column** → Tidak perlu pembesaran momen.\n\n"
            f"Momen Desain Mu = M2 = **{abs(inp['M2']):.2f} kN.m**"
        )


# ============================================================
# TAB 3: DIAGRAM INTERAKSI
# ============================================================
with tab3:
    st.subheader("14 & 15. Parameter & Layer Tulangan")
    col_a, col_b = st.columns(2)

    with col_a:
        df_layer = pd.DataFrame({
            "Layer": [lyr["nama"] for lyr in layers],
            "yi dari tepi tekan (mm)": [lyr["yi"] for lyr in layers],
            "Jarak dari tengah (mm)": [round(lyr["yi"] - h/2, 2) for lyr in layers],
            "Jumlah (bh)": [lyr["n"] for lyr in layers],
            "Luas (mm²)": [f"{lyr['A']:,.2f}" for lyr in layers],
        })
        st.dataframe(df_layer, use_container_width=True, hide_index=True)

    with col_b:
        st.info(
            f"**β1** = {beta1:.6f}  (fc'={inp['fc']} MPa)\n\n"
            f"**εcu** = 0.003  (SNI 2847:2019 Ps. 22.2.2.1)\n\n"
            f"**εy** = fy/Es = {inp['fy']:.0f}/{inp['Es']:.0f} = {eps_y:.4f}\n\n"
            f"**d'** = {props['d_prime']:.2f} mm  |  **d** = {props['d']:.2f} mm\n\n"
            f"**Momen Desain Mu** = {Mu_desain:.4f} kN.m"
        )

    st.divider()
    st.subheader("16. Tabel Diagram Interaksi P-M (52 Titik) — SNI 2847:2019")
    df_int = pd.DataFrame(hasil_interaksi)
    st.dataframe(df_int, use_container_width=True, hide_index=True, height=500)

    st.divider()
    st.subheader("20. Diagram Interaksi P-M (φPn vs φMn)")

    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    pts = data_grafik(hasil_interaksi)
    phi_Mn_k = [p[0] for p in pts]
    phi_Pn_k = [p[1] for p in pts]

    fig, ax = plt.subplots(figsize=(9, 8))
    ax.plot(phi_Mn_k, phi_Pn_k, "b-o", lw=2, ms=3, label="Kurva φPn-φMn")
    ax.plot(
        Mu_desain, inp["Pu"], "r*", ms=14,
        label=f"Beban Aktual (Mu={Mu_desain:.2f} kN.m, Pu={inp['Pu']:.0f} kN)",
    )
    ax.axhline(0, color="k", lw=0.7, ls="--")
    ax.axvline(0, color="k", lw=0.7, ls="--")
    ax.set_xlabel("φMn (kN.m)", fontsize=12)
    ax.set_ylabel("φPn (kN)", fontsize=12)
    ax.set_title("DIAGRAM INTERAKSI P-M\nSNI 2847:2019", fontsize=13, fontweight="bold")
    ax.legend(fontsize=10)
    ax.grid(True, alpha=0.3)

    stxt_color = "green" if is_ok else "red"
    ax.annotate(
        cek["status"],
        xy=(Mu_desain, inp["Pu"]),
        xytext=(Mu_desain + 5, inp["Pu"] + 150),
        fontsize=11, color=stxt_color, fontweight="bold",
        arrowprops=dict(arrowstyle="->", color=stxt_color),
    )
    st.pyplot(fig)
    plt.close()


# ============================================================
# TAB 4: CEK KAPASITAS
# ============================================================
with tab4:
    st.subheader("19. Ringkasan & Cek Kapasitas — SNI 2847:2019")

    def fmt_val(v, decimals=2):
        if v is None:
            return "-"
        return f"{v:.{decimals}f}"

    ratio_Pu_val = cek["ratio_Pu"]
    ratio_Mu_val = cek["ratio_Mu"]
    phi_Pn_kap   = cek["phi_Pn_kapasitas"]
    phi_Mn_kap   = cek["phi_Mn_kapasitas"]

    ok_pn = "-" if ratio_Pu_val is None else ("OK" if ratio_Pu_val <= 1.0 else "NOT OK")
    ok_mn = "-" if ratio_Mu_val is None else ("OK" if ratio_Mu_val <= 1.0 else "NOT OK")

    df_cek = pd.DataFrame({
        "Parameter": [
            "Klasifikasi Kolom",
            "Gaya Aksial (Pu)",
            "Momen Desain (Mc/M2)",
            "φPn Kapasitas (pada Mu=Mc)",
            "φMn Kapasitas (pada Pu)",
            "Rasio Pu / φPn",
            "Rasio Mu / φMn",
            "STATUS AKHIR",
        ],
        "Nilai": [
            kelangsingan["klasifikasi"],
            f"{inp['Pu']:.2f}",
            f"{Mu_desain:.4f}",
            fmt_val(phi_Pn_kap),
            fmt_val(phi_Mn_kap),
            fmt_val(ratio_Pu_val),
            fmt_val(ratio_Mu_val),
            cek["status"],
        ],
        "Satuan": ["-", "kN", "kN.m", "kN", "kN.m", "-", "-", "-"],
        "Status": ["-", "OK", "OK", "-", "-", ok_pn, ok_mn, cek["status"]],
    })
    st.dataframe(df_cek, use_container_width=True, hide_index=True)

    st.write("")
    if is_ok:
        pn_txt = f"φPn = {phi_Pn_kap:.2f} kN" if phi_Pn_kap else "φPn = (kurva di sisi tarik)"
        mn_txt = f"φMn = {phi_Mn_kap:.2f} kN.m" if phi_Mn_kap else "φMn = tidak terhitung"
        st.success(
            f"### ✅ {cek['status']}\n\n"
            f"Titik beban **(Pu={inp['Pu']:.0f} kN, Mu={Mu_desain:.2f} kN.m)** "
            f"berada **DI DALAM** kurva interaksi φPn-φMn.\n\n"
            f"- Pu = {inp['Pu']:.0f} kN  |  {pn_txt}\n"
            f"- Mu = {Mu_desain:.2f} kN.m  |  {mn_txt}"
        )
    else:
        st.error(
            f"### ❌ {cek['status']}\n\n"
            f"Titik beban berada **DI LUAR** kurva interaksi.\n\n"
            f"Perlu revisi: perbesar dimensi atau tambah tulangan!"
        )


# ============================================================
# TAB 5: LAPORAN
# ============================================================
with tab5:
    st.subheader("📄 Generator Laporan Profesional")
    nama_eng = st.text_input("Nama Engineer", value="Ladosi")
    nama_prj = st.text_input("Nama Proyek", value="Perhitungan Kolom Beton Bertulang")
    tgl_lpr  = st.text_input("Tanggal Laporan", value="2025")

    col_w, col_p = st.columns(2)
    with col_w:
        if st.button("📄 Buat Laporan Word", use_container_width=True):
            try:
                docx_bytes = buat_word(R, nama_eng, nama_prj, tgl_lpr, Mu_desain)
                st.download_button(
                    "⬇️ Download Word (.docx)",
                    data=docx_bytes,
                    file_name="laporan_kolom.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as ex:
                st.error(f"Error Word: {ex}")

    with col_p:
        if st.button("📋 Buat Laporan PDF", use_container_width=True):
            try:
                pdf_bytes = buat_pdf(R, nama_eng, nama_prj, tgl_lpr, Mu_desain)
                st.download_button(
                    "⬇️ Download PDF (.pdf)",
                    data=pdf_bytes,
                    file_name="laporan_kolom.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as ex:
                st.error(f"Error PDF: {ex}")


# ============================================================
