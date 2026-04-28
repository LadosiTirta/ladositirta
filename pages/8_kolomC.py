"""
app.py - Perhitungan Kapasitas Kolom Beton Bertulang
Diagram Interaksi P-M | SNI 2847:2019
Author: Ladosi
"""

import streamlit as st
import numpy as np
import math
import io

# ============================================================
# BAGIAN 1: LOGIKA KALKULASI
# ============================================================

def hitung_properties(b, h, c, dia_s, D, n_b, n_h, fc, fy, Es):
    """Hitung properti penampang kolom."""
    Ag = b * h
    Ig = (1/12) * b * h**3
    Ec = 4700 * math.sqrt(fc)
    d_prime = c + dia_s + D/2
    d = h - d_prime
    n_total = 2*n_b + 2*(n_h - 2)
    Ast = n_total * math.pi * D**2 / 4
    rho_g = Ast / Ag * 100
    rho_min = 1.0
    rho_max = 8.0
    cek_rasio = "OK" if rho_min <= rho_g <= rho_max else "NOT OK"
    return {
        "Ag": Ag, "Ig": Ig, "Ec": Ec, "d": d, "d_prime": d_prime,
        "n_total": n_total, "Ast": Ast, "rho_g": rho_g,
        "rho_min": rho_min, "rho_max": rho_max, "cek_rasio": cek_rasio
    }


def hitung_kekakuan_elemen(fc, Ig, Ec,
                            bb_al, hb_al, Lb_al,
                            bb_ar, hb_ar, Lb_ar,
                            bk_a, hk_a, Lk_a,
                            bb_bl, hb_bl, Lb_bl,
                            bb_br, hb_br, Lb_br,
                            bk_b, hk_b, Lk_b,
                            b_kolom, h_kolom, Lu):
    """Hitung EI/L setiap elemen untuk faktor psi."""
    def EI_L(b_e, h_e, L_e, is_beam):
        if L_e <= 0:
            return 0.0
        I_e = (1/12) * b_e * h_e**3
        Ec_e = 4700 * math.sqrt(fc)
        return Ec_e * I_e / L_e

    Ig_kolom = (1/12) * b_kolom * h_kolom**3
    Ec_kolom = 4700 * math.sqrt(fc)

    EI_L_bal_atas_kiri   = EI_L(bb_al, hb_al, Lb_al, True)
    EI_L_bal_atas_kanan  = EI_L(bb_ar, hb_ar, Lb_ar, True)
    EI_L_kol_atas        = EI_L(bk_a, hk_a, Lk_a, False)
    EI_L_bal_bawah_kiri  = EI_L(bb_bl, hb_bl, Lb_bl, True)
    EI_L_bal_bawah_kanan = EI_L(bb_br, hb_br, Lb_br, True)
    EI_L_kol_bawah       = EI_L(bk_b, hk_b, Lk_b, False)
    EI_L_kol_ditinjau    = Ec_kolom * Ig_kolom / Lu

    sum_kol_atas   = EI_L_kol_atas + EI_L_kol_ditinjau
    sum_balok_atas = EI_L_bal_atas_kiri + EI_L_bal_atas_kanan

    sum_kol_bawah   = EI_L_kol_ditinjau + EI_L_kol_bawah
    sum_balok_bawah = EI_L_bal_bawah_kiri + EI_L_bal_bawah_kanan

    psi_A = sum_kol_atas / sum_balok_atas if sum_balok_atas > 0 else 10.0
    psi_B = sum_kol_bawah / sum_balok_bawah if sum_balok_bawah > 0 else 10.0

    return {
        "EI_L_bal_atas_kiri":   EI_L_bal_atas_kiri,
        "EI_L_bal_atas_kanan":  EI_L_bal_atas_kanan,
        "EI_L_kol_atas":        EI_L_kol_atas,
        "EI_L_bal_bawah_kiri":  EI_L_bal_bawah_kiri,
        "EI_L_bal_bawah_kanan": EI_L_bal_bawah_kanan,
        "EI_L_kol_bawah":       EI_L_kol_bawah,
        "EI_L_kol_ditinjau":    EI_L_kol_ditinjau,
        "sum_kol_atas":   sum_kol_atas,
        "sum_balok_atas": sum_balok_atas,
        "sum_kol_bawah":   sum_kol_bawah,
        "sum_balok_bawah": sum_balok_bawah,
        "psi_A": psi_A,
        "psi_B": psi_B,
    }


def hitung_k_dan_kelangsingan(psi_A, psi_B, kondisi_rangka, Lu, h_kolom, M1, M2):
    """Hitung faktor k, radius girasi, dan cek kelangsingan."""
    psi_m = (psi_A + psi_B) / 2

    # Faktor k - Braced (ambil minimum dari dua rumus)
    k_braced_1 = 0.7 + 0.05 * (psi_A + psi_B)
    k_braced_1 = min(k_braced_1, 1.0)
    k_braced_2 = 0.85 + 0.05 * min(psi_A, psi_B)
    k_braced_2 = min(k_braced_2, 1.0)
    k_braced = min(k_braced_1, k_braced_2)

    # Faktor k - Unbraced
    if psi_m < 2:
        k_unbraced = (20 - psi_m) / 20 * math.sqrt(1 + psi_m)
    else:
        k_unbraced = 0.9 * math.sqrt(1 + psi_m)

    k = k_braced if kondisi_rangka == "Braced" else k_unbraced

    r = 0.3 * h_kolom  # SNI 2847:2019
    kLu = k * Lu
    rasio_kelangsingan = kLu / r

    # Batas kelangsingan
    if kondisi_rangka == "Braced":
        batas = 34 - 12 * (M1 / M2) if M2 != 0 else 34
    else:
        batas = 22

    klasifikasi = "Short Column" if rasio_kelangsingan <= batas else "Slender Column"

    return {
        "k_braced": k_braced,
        "k_unbraced": k_unbraced,
        "k": k,
        "psi_m": psi_m,
        "r": r,
        "kLu": kLu,
        "rasio_kelangsingan": rasio_kelangsingan,
        "batas": batas,
        "klasifikasi": klasifikasi
    }


def hitung_pembesaran_momen(fc, Ig, k, Lu, Pu, M2, M1, beta_dns, kondisi_rangka):
    """Hitung pembesaran momen untuk Slender Column - Braced Frame."""
    Ec = 4700 * math.sqrt(fc)
    EI_eff = (0.4 * Ec * Ig) / (1 + beta_dns)
    Pc = (math.pi**2 * EI_eff) / (k * Lu)**2 / 1000  # kN
    Cm = max(0.6 + 0.4 * (M1 / M2), 0.4) if M2 != 0 else 1.0
    denom = 1 - Pu / (0.75 * Pc)
    if denom <= 0:
        denom = 0.001
    delta_ns = max(Cm / denom, 1.0)
    Mc = delta_ns * M2
    return {
        "Ec": Ec,
        "EI_eff": EI_eff,
        "Pc": Pc,
        "Cm": Cm,
        "beta_dns": beta_dns,
        "delta_ns": delta_ns,
        "Mc": Mc
    }


def hitung_layer_tulangan(h, d_prime, n_b, n_h, D):
    """Hitung posisi dan luas tiap layer tulangan."""
    n_total = 2*n_b + 2*(n_h - 2)
    A_bar = math.pi * D**2 / 4

    # Layer 1 (Tekan): d'
    # Layer 2 (Tengah): h/2
    # Layer 3 (Tarik): h - d'
    n_layer1 = n_b   # tulangan di sisi tekan
    n_layer2 = 2*(n_h - 2)  # tulangan tengah (sisi kiri-kanan)
    n_layer3 = n_b   # tulangan di sisi tarik

    y1 = d_prime          # dari tepi tekan
    y2 = h / 2
    y3 = h - d_prime      # d

    layers = [
        {"nama": "Layer 1 (Tekan)", "yi": y1, "jarak_tengah": y1 - h/2, "n": n_layer1, "A": n_layer1 * A_bar},
        {"nama": "Layer 2 (Tengah)", "yi": y2, "jarak_tengah": 0.0, "n": n_layer2, "A": n_layer2 * A_bar},
        {"nama": "Layer 3 (Tarik)", "yi": y3, "jarak_tengah": y3 - h/2, "n": n_layer3, "A": n_layer3 * A_bar},
    ]
    return layers


def hitung_phi_dinamis(eps_t, eps_y):
    """Faktor reduksi phi dinamis berdasarkan regangan tarik."""
    if eps_t >= 0.005:
        return 0.90
    elif eps_t <= eps_y:
        return 0.65
    else:
        # Interpolasi linear antara 0.65 dan 0.90
        return 0.65 + (eps_t - eps_y) / (0.005 - eps_y) * (0.90 - 0.65)


def hitung_diagram_interaksi(b, h, fc, fy, Es, layers, n_titik=50):
    """Hitung 52 titik diagram interaksi P-M."""
    eps_cu = 0.003
    eps_y = fy / Es

    # beta1
    if fc <= 28:
        beta1 = 0.85
    else:
        beta1 = max(0.85 - 0.05 * (fc - 28) / 7, 0.65)

    hasil = []
    ch_values = np.linspace(0.02, 1.20, n_titik)

    for i, ch in enumerate(ch_values):
        c = ch * h
        a = beta1 * c
        a = min(a, h)

        # Gaya beton tekan
        Cc = 0.85 * fc * b * a / 1000  # kN

        # Gaya tulangan tiap layer
        F_layers = []
        for layer in layers:
            yi = layer["yi"]
            eps_si = eps_cu * (c - yi) / c if c > 0 else 0
            eps_si = max(min(eps_si, eps_y), -eps_y)
            # Regangan aktual (tanpa batas):
            eps_si_actual = eps_cu * (c - yi) / c if c > 0 else 0
            fs_i = max(min(eps_si_actual * Es, fy), -fy)
            # Kurangi tegangan beton yang sudah dihitung di zone tekan
            if yi <= a:
                fs_i_net = fs_i - 0.85 * fc
            else:
                fs_i_net = fs_i
            F_i = fs_i_net * layer["A"] / 1000  # kN
            F_layers.append({"fs": fs_si_actual if False else eps_si_actual * Es, "F": F_i, "eps": eps_si_actual})

        # Regangan tarik terluar (layer terjauh dari tepi tekan)
        d_tarik = layers[-1]["yi"]
        eps_t = eps_cu * (d_tarik - c) / c if c > 0 else 0

        # Pn dan Mn (terhadap titik tengah penampang)
        Pn = Cc + sum(lyr["F"] for lyr in F_layers)
        xc = h/2 - a/2  # eksentrisitas gaya beton dari tengah
        Mn_cc = Cc * xc / 1000  # kN.m

        Mn_steel = sum(F_layers[j]["F"] * (h/2 - layers[j]["yi"]) / 1000
                       for j in range(len(layers)))
        Mn = Mn_cc - Mn_steel
        Mn = abs(Mn)

        phi = hitung_phi_dinamis(eps_t, eps_y)
        phi_Pn = phi * Pn
        phi_Mn = phi * Mn

        hasil.append({
            "No": i+1, "c/h": round(ch, 4), "c": round(c, 2),
            "a": round(a, 2), "eps_t": round(eps_t, 5),
            "Pn": round(Pn, 2), "Mn": round(Mn, 2),
            "phi": round(phi, 2), "phi_Pn": round(phi_Pn, 2),
            "phi_Mn": round(phi_Mn, 2)
        })

    # Titik 51: Tekan Murni (Po)
    # Po = 0.85*fc*(Ag-Ast) + fy*Ast
    Ag = b * h
    Ast = sum(lyr["A"] for lyr in layers)
    Po = (0.85 * fc * (Ag - Ast) + fy * Ast) / 1000  # kN
    phi_po = 0.65
    hasil.append({
        "No": 51, "c/h": "∞", "c": "∞", "a": "-", "eps_t": "-",
        "Pn": round(Po, 2), "Mn": 0,
        "phi": phi_po, "phi_Pn": round(phi_po * Po, 2), "phi_Mn": 0
    })

    # Titik 52: Tarik Murni (To)
    To = -fy * Ast / 1000  # kN (negatif = tarik)
    phi_to = 0.90
    hasil.append({
        "No": 52, "c/h": 0.0, "c": 0, "a": "-", "eps_t": "-",
        "Pn": round(To, 2), "Mn": 0,
        "phi": phi_to, "phi_Pn": round(phi_to * To, 2), "phi_Mn": 0
    })

    return hasil, beta1, eps_y


def cek_kapasitas(hasil_interaksi, Pu, Mu):
    """Cek apakah titik beban (Pu, Mu) berada di dalam kurva interaksi."""
    # Kumpulkan titik numerik
    phi_Pn_list = []
    phi_Mn_list = []
    for row in hasil_interaksi:
        try:
            phi_Pn_list.append(float(row["phi_Pn"]))
            phi_Mn_list.append(float(row["phi_Mn"]))
        except:
            pass

    # Cari titik di kurva dengan Pn mendekati Pu
    # Gunakan interpolasi sederhana: temukan phi_Mn_max saat phi_Pn >= Pu
    # Metode: cek apakah titik beban di dalam polygon kurva
    # Simplifikasi: cari phi_Mn pada phi_Pn = Pu
    pairs = sorted(zip(phi_Pn_list, phi_Mn_list), key=lambda x: x[0])

    # Interpolasi phi_Mn_kapasitas pada Pu
    phi_Mn_kapasitas = None
    for idx in range(len(pairs)-1):
        p1, m1 = pairs[idx]
        p2, m2 = pairs[idx+1]
        if p1 <= Pu <= p2:
            if p2 != p1:
                t = (Pu - p1) / (p2 - p1)
                phi_Mn_kapasitas = m1 + t * (m2 - m1)
            else:
                phi_Mn_kapasitas = max(m1, m2)
            break

    # Cari phi_Pn_kapasitas saat phi_Mn = Mu
    phi_Pn_kapasitas = None
    pairs_mn = sorted(zip(phi_Mn_list, phi_Pn_list), key=lambda x: x[0])
    for idx in range(len(pairs_mn)-1):
        m1, p1 = pairs_mn[idx]
        m2, p2 = pairs_mn[idx+1]
        if m1 <= Mu <= m2:
            if m2 != m1:
                t = (Mu - m1) / (m2 - m1)
                phi_Pn_kapasitas = p1 + t * (p2 - p1)
            else:
                phi_Pn_kapasitas = max(p1, p2)
            break

    ratio_Pu = Pu / phi_Pn_kapasitas if phi_Pn_kapasitas and phi_Pn_kapasitas != 0 else None
    ratio_Mu = Mu / phi_Mn_kapasitas if phi_Mn_kapasitas and phi_Mn_kapasitas != 0 else None

    status = "OK - AMAN"
    if ratio_Pu is not None and ratio_Mu is not None:
        if ratio_Pu > 1.0 or ratio_Mu > 1.0:
            status = "NOT OK - TIDAK AMAN"

    return {
        "phi_Mn_kapasitas": phi_Mn_kapasitas,
        "phi_Pn_kapasitas": phi_Pn_kapasitas,
        "ratio_Pu": ratio_Pu,
        "ratio_Mu": ratio_Mu,
        "status": status
    }


def susun_data_grafik(hasil_interaksi):
    """Urutkan data kurva untuk grafik (dari tekan murni ke tarik murni)."""
    numerik = []
    for row in hasil_interaksi:
        try:
            numerik.append((float(row["phi_Mn"]), float(row["phi_Pn"]),
                            str(row["c/h"])))
        except:
            pass
    # Urut dari phi_Pn besar ke kecil (atas ke bawah)
    numerik.sort(key=lambda x: -x[1])
    return numerik


# ============================================================
# BAGIAN 2: UI STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Kolom Beton Bertulang | SNI 2847:2019",
    page_icon="🏛️",
    layout="wide"
)

st.title("🏛️ Perhitungan Kapasitas Kolom Beton Bertulang")
st.markdown("**Diagram Interaksi P-M | SNI 2847:2019** | Dibuat oleh: Ladosi")
st.divider()

# --- SIDEBAR INPUT ---
with st.sidebar:
    st.header("📋 Data Input")

    st.subheader("1. Material")
    fc  = st.number_input("Mutu Beton fc' (MPa)", value=30.0, step=1.0)
    fy  = st.number_input("Mutu Baja fy (MPa)", value=400.0, step=10.0)
    Es  = st.number_input("Modulus Elastisitas Baja Es (MPa)", value=200000.0, step=1000.0)

    st.subheader("2. Dimensi Kolom")
    b      = st.number_input("Lebar Kolom b (mm)", value=400.0, step=10.0)
    h      = st.number_input("Tinggi Kolom h (mm)", value=500.0, step=10.0)
    c_sel  = st.number_input("Selimut Beton c (mm)", value=40.0, step=5.0)
    dia_s  = st.number_input("Diameter Sengkang Øs (mm)", value=10.0, step=2.0)

    st.subheader("3. Tulangan Longitudinal")
    D   = st.number_input("Diameter Tulangan D (mm)", value=22.0, step=2.0)
    n_b = st.number_input("Jumlah Tulangan Sisi b (n_b)", value=4, step=1, min_value=2)
    n_h = st.number_input("Jumlah Tulangan Sisi h (n_h, incl. sudut)", value=3, step=1, min_value=2)

    st.subheader("4. Data Panjang & Kondisi")
    Lu             = st.number_input("Panjang Tak Tertahan Lu (mm)", value=6000.0, step=100.0)
    kondisi_rangka = st.selectbox("Kondisi Rangka", ["Braced", "Unbraced"])
    kelengkungan   = st.selectbox("Kelengkungan", ["Single", "Double"])

    st.subheader("5. Kekakuan Elemen Penghubung")
    st.markdown("**Balok Atas** *(isi 0 jika tidak ada)*")
    bb_al  = st.number_input("Balok Atas-Kiri: b (mm)", value=300.0)
    hb_al  = st.number_input("Balok Atas-Kiri: h (mm)", value=500.0)
    Lb_al  = st.number_input("Balok Atas-Kiri: L (mm)", value=6000.0)
    bb_ar  = st.number_input("Balok Atas-Kanan: b (mm)", value=300.0)
    hb_ar  = st.number_input("Balok Atas-Kanan: h (mm)", value=500.0)
    Lb_ar  = st.number_input("Balok Atas-Kanan: L (mm)", value=6000.0)

    st.markdown("**Kolom Atas** *(isi 0 jika tidak ada / atap)*")
    bk_a = st.number_input("Kolom Atas: b (mm)", value=400.0)
    hk_a = st.number_input("Kolom Atas: h (mm)", value=500.0)
    Lk_a = st.number_input("Kolom Atas: L (mm) [0 jika tidak ada]", value=3500.0)
    st.caption("ℹ️ Jika tidak ada kolom atas (lantai atap), isi L=0")

    st.markdown("**Balok Bawah** *(isi 0 jika tidak ada)*")
    bb_bl  = st.number_input("Balok Bawah-Kiri: b (mm)", value=300.0)
    hb_bl  = st.number_input("Balok Bawah-Kiri: h (mm)", value=500.0)
    Lb_bl  = st.number_input("Balok Bawah-Kiri: L (mm)", value=6000.0)
    bb_br  = st.number_input("Balok Bawah-Kanan: b (mm)", value=300.0)
    hb_br  = st.number_input("Balok Bawah-Kanan: h (mm)", value=500.0)
    Lb_br  = st.number_input("Balok Bawah-Kanan: L (mm)", value=6000.0)

    st.markdown("**Kolom Bawah** *(isi 0 jika langsung pondasi)*")
    bk_b = st.number_input("Kolom Bawah: b (mm)", value=400.0)
    hk_b = st.number_input("Kolom Bawah: h (mm)", value=500.0)
    Lk_b = st.number_input("Kolom Bawah: L (mm) [0 jika pondasi]", value=3500.0)
    st.caption("ℹ️ Jika bawah langsung pondasi (sendi/jepit), isi L=0 → ψ=0 (jepit) atau biarkan besar untuk sendi")

    st.subheader("6. Beban Terfaktor")
    Pu = st.number_input("Gaya Aksial Pu (kN)", value=1500.0, step=10.0)
    M1 = st.number_input("Momen Ujung 1 M1 (kN.m) - lebih kecil", value=80.0, step=5.0)
    M2 = st.number_input("Momen Ujung 2 M2 (kN.m) - lebih besar", value=150.0, step=5.0)

    st.subheader("7. Faktor Tambahan")
    beta_dns = st.number_input("Faktor βdns (rasio beban tetap)", value=0.60, step=0.05, min_value=0.0, max_value=1.0)

    st.divider()
    hitung_btn = st.button("🔴 HITUNG SEKARANG", use_container_width=True, type="primary")

# ============================================================
# PROSES KALKULASI & TAMPILAN HASIL
# ============================================================

if hitung_btn:
    # --- Validasi input dasar ---
    if M2 == 0:
        st.error("M2 tidak boleh 0!")
        st.stop()

    # Sesuaikan tanda M1 untuk Single/Double curvature
    M1_signed = M1 if kelengkungan == "Double" else -M1

    # 1. Properties
    props = hitung_properties(b, h, c_sel, dia_s, D, n_b, n_h, fc, fy, Es)

    # 2. Kekakuan Elemen
    kekakuan = hitung_kekakuan_elemen(
        fc, props["Ig"], props["Ec"],
        bb_al, hb_al, Lb_al,
        bb_ar, hb_ar, Lb_ar,
        bk_a, hk_a, Lk_a,
        bb_bl, hb_bl, Lb_bl,
        bb_br, hb_br, Lb_br,
        bk_b, hk_b, Lk_b,
        b, h, Lu
    )

    # 3. K dan Kelangsingan
    kelangsingan = hitung_k_dan_kelangsingan(
        kekakuan["psi_A"], kekakuan["psi_B"],
        kondisi_rangka, Lu, h, M1_signed, M2
    )

    # 4. Pembesaran Momen
    pembesaran = hitung_pembesaran_momen(
        fc, props["Ig"], kelangsingan["k"], Lu, Pu, M2, M1_signed, beta_dns, kondisi_rangka
    )
    Mu_desain = pembesaran["Mc"] if kelangsingan["klasifikasi"] == "Slender Column" and kondisi_rangka == "Braced" else M2

    # 5. Layer Tulangan
    layers = hitung_layer_tulangan(h, props["d_prime"], n_b, n_h, D)

    # 6. Diagram Interaksi
    hasil_interaksi, beta1, eps_y = hitung_diagram_interaksi(b, h, fc, fy, Es, layers)

    # 7. Cek Kapasitas
    cek = cek_kapasitas(hasil_interaksi, Pu, Mu_desain)

    # Simpan ke session_state
    st.session_state["hasil"] = {
        "props": props, "kekakuan": kekakuan, "kelangsingan": kelangsingan,
        "pembesaran": pembesaran, "layers": layers,
        "hasil_interaksi": hasil_interaksi, "cek": cek,
        "beta1": beta1, "eps_y": eps_y,
        "Mu_desain": Mu_desain,
        # Input snapshot
        "inp": {
            "fc": fc, "fy": fy, "Es": Es, "b": b, "h": h,
            "c_sel": c_sel, "dia_s": dia_s, "D": D, "n_b": n_b, "n_h": n_h,
            "Lu": Lu, "kondisi_rangka": kondisi_rangka, "kelengkungan": kelengkungan,
            "Pu": Pu, "M1": M1, "M2": M2, "beta_dns": beta_dns,
        }
    }

# Tampilkan hasil jika ada di session_state
if "hasil" in st.session_state:
    R = st.session_state["hasil"]
    props        = R["props"]
    kekakuan     = R["kekakuan"]
    kelangsingan = R["kelangsingan"]
    pembesaran   = R["pembesaran"]
    layers       = R["layers"]
    hasil_interaksi = R["hasil_interaksi"]
    cek          = R["cek"]
    beta1        = R["beta1"]
    eps_y        = R["eps_y"]
    Mu_desain    = R["Mu_desain"]
    inp          = R["inp"]

    # ---- STATUS BANNER ----
    status_color = "green" if "OK" in cek["status"] and "NOT" not in cek["status"] else "red"
    st.markdown(f"""
    <div style="background:{status_color};padding:15px;border-radius:10px;text-align:center;">
        <h2 style="color:white;margin:0;">🏛️ {cek['status']}</h2>
    </div>
    """, unsafe_allow_html=True)
    st.write("")

    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📐 Properties", "📏 Kelangsingan", "📈 Diagram Interaksi",
        "✅ Cek Kapasitas", "📄 Laporan"
    ])

    # ============================================================
    # TAB 1: PROPERTIES
    # ============================================================
    with tab1:
        st.subheader("7. Properti Penampang Kolom (SNI 2847:2019)")
        col1, col2 = st.columns(2)
        with col1:
            data_props = {
                "Parameter": ["Luas Penampang Bruto (Ag)", "Momen Inersia (Ig)",
                              "Modulus Elastisitas Beton (Ec)", "Tinggi Efektif (d)",
                              "Tinggi Efektif Tekan (d')", "Total Tulangan (n_total)",
                              "Luas Tulangan Total (Ast)", "Rasio Tulangan (ρg)"],
                "Nilai": [
                    f"{props['Ag']:,.2f}", f"{props['Ig']:,.2f}",
                    f"{props['Ec']:,.2f}", f"{props['d']:,.2f}",
                    f"{props['d_prime']:,.2f}", f"{props['n_total']}",
                    f"{props['Ast']:,.2f}", f"{props['rho_g']:,.2f}"
                ],
                "Satuan": ["mm²", "mm⁴", "MPa", "mm", "mm", "buah", "mm²", "%"],
                "Keterangan": [
                    "b × h",
                    "(1/12) × b × h³",
                    "4700 × √fc' (SNI Ps. 19.2.2)",
                    "h - c - Øs - D/2",
                    "c + Øs + D/2",
                    "2×n_b + 2×(n_h-2)",
                    "n_total × π × D² / 4",
                    "Ast / Ag × 100"
                ]
            }
            import pandas as pd
            df_props = pd.DataFrame(data_props)
            st.dataframe(df_props, use_container_width=True, hide_index=True)

        with col2:
            st.markdown("**Step-by-Step Perhitungan:**")
            st.code(f"""
Ag  = b × h = {inp['b']:.0f} × {inp['h']:.0f} = {props['Ag']:,.2f} mm²

Ig  = (1/12) × b × h³
    = (1/12) × {inp['b']:.0f} × {inp['h']:.0f}³
    = {props['Ig']:,.2f} mm⁴

Ec  = 4700 × √fc'
    = 4700 × √{inp['fc']:.0f}
    = {props['Ec']:,.2f} MPa

d'  = c + Øs + D/2
    = {inp['c_sel']:.0f} + {inp['dia_s']:.0f} + {inp['D']:.0f}/2
    = {props['d_prime']:,.2f} mm

d   = h - d'
    = {inp['h']:.0f} - {props['d_prime']:.2f}
    = {props['d']:,.2f} mm

n_total = 2×n_b + 2×(n_h-2)
        = 2×{inp['n_b']} + 2×({inp['n_h']}-2)
        = {props['n_total']} buah

Ast = n_total × π × D² / 4
    = {props['n_total']} × π × {inp['D']:.0f}² / 4
    = {props['Ast']:,.2f} mm²

ρg  = Ast/Ag × 100
    = {props['Ast']:.2f}/{props['Ag']:.2f} × 100
    = {props['rho_g']:.2f}% → {props['cek_rasio']}
            """)

        st.subheader("8. Cek Rasio Tulangan (SNI 2847:2019 Pasal 10.6.1)")
        st.info(f"ρmin = {props['rho_min']}% ≤ ρg = {props['rho_g']:.2f}% ≤ ρmax = {props['rho_max']}% → **{props['cek_rasio']}**")

        st.subheader("9. Perhitungan Kekakuan Elemen (EI/L) - SNI 2847:2019 Pasal 6.6.3")
        data_kekakuan = {
            "Elemen": [
                "Balok Atas - Kiri", "Balok Atas - Kanan", "Kolom Atas",
                "Balok Bawah - Kiri", "Balok Bawah - Kanan", "Kolom Bawah",
                "Kolom Ditinjau"
            ],
            "b (mm)": [bb_al, bb_ar, bk_a, bb_bl, bb_br, bk_b, inp["b"]],
            "h (mm)": [hb_al, hb_ar, hk_a, hb_bl, hb_br, hk_b, inp["h"]],
            "L (mm)": [Lb_al, Lb_ar, Lk_a, Lb_bl, Lb_br, Lk_b, inp["Lu"]],
            "EI/L (N.mm)": [
                f"{kekakuan['EI_L_bal_atas_kiri']:,.2f}",
                f"{kekakuan['EI_L_bal_atas_kanan']:,.2f}",
                f"{kekakuan['EI_L_kol_atas']:,.2f}",
                f"{kekakuan['EI_L_bal_bawah_kiri']:,.2f}",
                f"{kekakuan['EI_L_bal_bawah_kanan']:,.2f}",
                f"{kekakuan['EI_L_kol_bawah']:,.2f}",
                f"{kekakuan['EI_L_kol_ditinjau']:,.2f}"
            ]
        }
        import pandas as pd
        df_kekakuan = pd.DataFrame(data_kekakuan)
        st.dataframe(df_kekakuan, use_container_width=True, hide_index=True)

    # ============================================================
    # TAB 2: KELANGSINGAN
    # ============================================================
    with tab2:
        st.subheader("10. Faktor Kekakuan Ψ (Psi) - SNI 2847:2019 Pasal 6.6.4.4")
        col1, col2 = st.columns(2)
        with col1:
            data_psi = {
                "Parameter": [
                    "Σ(EI/L) Kolom Atas", "Σ(EI/L) Balok Atas", "ΨA",
                    "Σ(EI/L) Kolom Bawah", "Σ(EI/L) Balok Bawah", "ΨB",
                    "Ψm (rata-rata)"
                ],
                "Nilai": [
                    f"{kekakuan['sum_kol_atas']:,.2f}",
                    f"{kekakuan['sum_balok_atas']:,.2f}",
                    f"{kekakuan['psi_A']:.4f}",
                    f"{kekakuan['sum_kol_bawah']:,.2f}",
                    f"{kekakuan['sum_balok_bawah']:,.2f}",
                    f"{kekakuan['psi_B']:.4f}",
                    f"{kelangsingan['psi_m']:.4f}"
                ],
                "Satuan": ["N.mm", "N.mm", "-", "N.mm", "N.mm", "-", "-"]
            }
            import pandas as pd
            df_psi = pd.DataFrame(data_psi)
            st.dataframe(df_psi, use_container_width=True, hide_index=True)

        with col2:
            st.code(f"""
PSI ATAS (ΨA):
  ΨA = Σ(EI/L)kolom / Σ(EI/L)balok
     = {kekakuan['sum_kol_atas']:,.2f} / {kekakuan['sum_balok_atas']:,.2f}
     = {kekakuan['psi_A']:.4f}

PSI BAWAH (ΨB):
  ΨB = {kekakuan['sum_kol_bawah']:,.2f} / {kekakuan['sum_balok_bawah']:,.2f}
     = {kekakuan['psi_B']:.4f}

Ψm = (ΨA + ΨB) / 2
   = ({kekakuan['psi_A']:.4f} + {kekakuan['psi_B']:.4f}) / 2
   = {kelangsingan['psi_m']:.4f}
            """)

        st.subheader("11. Faktor Panjang Efektif (k) - SNI 2847:2019 Pasal 6.6.4.4")
        col1, col2 = st.columns(2)
        with col1:
            data_k = {
                "Parameter": [
                    "k (Braced)", "k (Unbraced)", "k Dipakai",
                    "Radius Girasi (r)", "Panjang Efektif (k×Lu)",
                    "Rasio Kelangsingan (k×Lu/r)"
                ],
                "Nilai": [
                    f"{kelangsingan['k_braced']:.6f}",
                    f"{kelangsingan['k_unbraced']:.6f}",
                    f"{kelangsingan['k']:.6f}",
                    f"{kelangsingan['r']:.2f}",
                    f"{kelangsingan['kLu']:.2f}",
                    f"{kelangsingan['rasio_kelangsingan']:.2f}"
                ],
                "Satuan": ["-", "-", "-", "mm", "mm", "-"]
            }
            import pandas as pd
            df_k = pd.DataFrame(data_k)
            st.dataframe(df_k, use_container_width=True, hide_index=True)

        with col2:
            st.code(f"""
k (Braced) = min(
  0.7 + 0.05×(ΨA+ΨB), 1.0
  0.85 + 0.05×Ψmin,   1.0)
  = {kelangsingan['k_braced']:.6f}

k (Unbraced):
  Ψm = {kelangsingan['psi_m']:.4f}
  {'k = (20-Ψm)/20×√(1+Ψm)' if kelangsingan['psi_m'] < 2 else 'k = 0.9×√(1+Ψm)'}
  = {kelangsingan['k_unbraced']:.6f}

r   = 0.3 × h = 0.3 × {inp['h']:.0f} = {kelangsingan['r']:.2f} mm

k×Lu = {kelangsingan['k']:.6f} × {inp['Lu']:.0f}
     = {kelangsingan['kLu']:.2f} mm

k×Lu/r = {kelangsingan['kLu']:.2f} / {kelangsingan['r']:.2f}
       = {kelangsingan['rasio_kelangsingan']:.2f}
            """)

        st.subheader("12. Cek Kelangsingan Kolom - SNI 2847:2019 Pasal 6.2.5")
        st.info(f"""
**Batas Kelangsingan ({inp['kondisi_rangka']}):** 
{f"34 - 12×(M1/M2) = 34 - 12×({abs(inp['M1'])}/{inp['M2']}) = {kelangsingan['batas']:.1f}" if inp['kondisi_rangka'] == "Braced" else "22"}

**Rasio Kelangsingan:** k×Lu/r = **{kelangsingan['rasio_kelangsingan']:.2f}**

**Klasifikasi: {kelangsingan['klasifikasi']}**
        """)

        if kelangsingan["klasifikasi"] == "Slender Column" and inp["kondisi_rangka"] == "Braced":
            st.subheader("13. Faktor Pembesaran Momen - SNI 2847:2019 Pasal 6.6.4")
            col1, col2 = st.columns(2)
            with col1:
                data_pem = {
                    "Parameter": [
                        "Kekakuan Efektif (EI)eff", "Beban Kritis Euler (Pc)",
                        "Faktor Cm", "Faktor βdns",
                        "Faktor Pembesaran (δns)", "Momen Desain (Mc)"
                    ],
                    "Nilai": [
                        f"{pembesaran['EI_eff']:,.2f}",
                        f"{pembesaran['Pc']:,.2f}",
                        f"{pembesaran['Cm']:.2f}",
                        f"{pembesaran['beta_dns']:.2f}",
                        f"{pembesaran['delta_ns']:.2f}",
                        f"{pembesaran['Mc']:,.4f}"
                    ],
                    "Satuan": ["N.mm²", "kN", "-", "-", "-", "kN.m"]
                }
                import pandas as pd
                df_pem = pd.DataFrame(data_pem)
                st.dataframe(df_pem, use_container_width=True, hide_index=True)

            with col2:
                st.code(f"""
(EI)eff = 0.4 × Ec × Ig / (1 + βdns)
        = 0.4 × {pembesaran['Ec']:,.2f} × {props['Ig']:,.2f}
          / (1 + {pembesaran['beta_dns']:.2f})
        = {pembesaran['EI_eff']:,.2f} N.mm²

Pc = π² × (EI)eff / (k×Lu)²
   = π² × {pembesaran['EI_eff']:,.2f}
     / ({kelangsingan['k']:.4f} × {inp['Lu']:.0f})²
   = {pembesaran['Pc']:,.2f} kN

Cm = 0.6 + 0.4 × (M1/M2)
   = 0.6 + 0.4 × ({inp['M1']:.0f}/{inp['M2']:.0f})
   = {pembesaran['Cm']:.2f} ≥ 0.4

δns = Cm / (1 - Pu/(0.75×Pc))
    = {pembesaran['Cm']:.2f} / (1 - {inp['Pu']:.0f}/(0.75×{pembesaran['Pc']:.2f}))
    = {pembesaran['delta_ns']:.2f} ≥ 1.0

Mc = δns × M2
   = {pembesaran['delta_ns']:.2f} × {inp['M2']:.0f}
   = {pembesaran['Mc']:,.4f} kN.m
                """)

    # ============================================================
    # TAB 3: DIAGRAM INTERAKSI
    # ============================================================
    with tab3:
        st.subheader("14 & 15. Parameter & Layer Tulangan")
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("**Layer Tulangan:**")
            data_layer = {
                "Layer": [lyr["nama"] for lyr in layers],
                "yi dari tepi tekan (mm)": [lyr["yi"] for lyr in layers],
                "Jarak dari Tengah (mm)": [f"{lyr['jarak_tengah']:,.2f}" for lyr in layers],
                "Jumlah": [lyr["n"] for lyr in layers],
                "Luas (mm²)": [f"{lyr['A']:,.2f}" for lyr in layers]
            }
            import pandas as pd
            df_layer = pd.DataFrame(data_layer)
            st.dataframe(df_layer, use_container_width=True, hide_index=True)

        with col2:
            st.info(f"""
**Parameter Diagram Interaksi:**
- β1 = {beta1:.6f} (fc'={inp['fc']} MPa)
- εcu = 0.003 (SNI 2847:2019 Ps. 22.2.2.1)
- εy = fy/Es = {inp['fy']:.0f}/{inp['Es']:.0f} = {eps_y:.4f}
- d' = {props['d_prime']:.2f} mm
- d = {props['d']:.2f} mm
- Jumlah Layer = 3
            """)

        st.subheader("16. Tabel Diagram Interaksi P-M (52 Titik) - SNI 2847:2019")
        import pandas as pd
        df_interaksi = pd.DataFrame(hasil_interaksi)
        st.dataframe(df_interaksi, use_container_width=True, hide_index=True,
                     height=500)

        st.subheader("20. Diagram Interaksi P-M (φPn vs φMn)")

        # Siapkan data grafik
        data_grafik = susun_data_grafik(hasil_interaksi)
        phi_Mn_kurva = [d[0] for d in data_grafik]
        phi_Pn_kurva = [d[1] for d in data_grafik]

        # Plot menggunakan matplotlib
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt

        fig, ax = plt.subplots(figsize=(10, 8))
        ax.plot(phi_Mn_kurva, phi_Pn_kurva, 'b-o', linewidth=2, markersize=3,
                label='Kurva Interaksi φPn-φMn')
        ax.plot(Mu_desain, inp["Pu"], 'r*', markersize=15,
                label=f'Beban Aktual (Mu={Mu_desain:.2f} kN.m, Pu={inp["Pu"]:.0f} kN)')
        ax.axhline(y=0, color='k', linewidth=0.5)
        ax.axvline(x=0, color='k', linewidth=0.5)
        ax.set_xlabel("φMn (kN.m)", fontsize=12)
        ax.set_ylabel("φPn (kN)", fontsize=12)
        ax.set_title("DIAGRAM INTERAKSI P-M (φPn vs φMn)\nSNI 2847:2019", fontsize=13, fontweight="bold")
        ax.legend(fontsize=10)
        ax.grid(True, alpha=0.3)

        # Anotasi status
        status_txt = cek["status"]
        color_txt = "green" if "AMAN" in status_txt else "red"
        ax.annotate(status_txt, xy=(Mu_desain, inp["Pu"]),
                    xytext=(Mu_desain + 10, inp["Pu"] + 100),
                    fontsize=11, color=color_txt, fontweight="bold",
                    arrowprops=dict(arrowstyle="->", color=color_txt))

        st.pyplot(fig)
        plt.close()

    # ============================================================
    # TAB 4: CEK KAPASITAS
    # ============================================================
    with tab4:
        st.subheader("19. Ringkasan & Cek Kapasitas")
        col1, col2 = st.columns(2)
        with col1:
            data_cek = {
                "Parameter": [
                    "Klasifikasi Kolom", "Gaya Aksial (Pu)",
                    "Momen Desain (Mc/M2)", "φPn Kapasitas (pada Pu)",
                    "φMn Kapasitas (pada Mu)", "Rasio Pu/φPn",
                    "Rasio Mu/φMn", "STATUS AKHIR"
                ],
                "Nilai": [
                    kelangsingan["klasifikasi"],
                    f"{inp['Pu']:.2f}",
                    f"{Mu_desain:.4f}",
                    f"{cek['phi_Pn_kapasitas']:.2f}" if cek["phi_Pn_kapasitas"] else "-",
                    f"{cek['phi_Mn_kapasitas']:.2f}" if cek["phi_Mn_kapasitas"] else "-",
                    f"{cek['ratio_Pu']:.2f}" if cek["ratio_Pu"] else "-",
                    f"{cek['ratio_Mu']:.2f}" if cek["ratio_Mu"] else "-",
                    cek["status"]
                ],
                "Satuan": ["-", "kN", "kN.m", "kN", "kN.m", "-", "-", "-"],
                "Status": [
                    "-", "OK", "OK",
                    "-", "-",
                    "OK" if (cek["ratio_Pu"] or 0) <= 1 else "NOT OK",
                    "OK" if (cek["ratio_Mu"] or 0) <= 1 else "NOT OK",
                    cek["status"]
                ]
            }
            import pandas as pd
            df_cek = pd.DataFrame(data_cek)
            st.dataframe(df_cek, use_container_width=True, hide_index=True)

        with col2:
            if "AMAN" in cek["status"]:
                st.success(f"""
### ✅ {cek['status']}

**Titik beban (Pu, Mu) berada DI DALAM kurva interaksi**

- Pu = {inp['Pu']:.0f} kN < φPn = {cek['phi_Pn_kapasitas']:.2f} kN
- Mu = {Mu_desain:.2f} kN.m < φMn = {cek['phi_Mn_kapasitas']:.2f} kN.m
                """)
            else:
                st.error(f"""
### ❌ {cek['status']}

**Titik beban (Pu, Mu) berada DI LUAR kurva interaksi**

- Perlu revisi dimensi atau tulangan!
                """)

    # ============================================================
    # TAB 5: LAPORAN
    # ============================================================
    with tab5:
        st.subheader("📄 Generator Laporan Profesional")
        nama_engineer = st.text_input("Nama Engineer", value="Ladosi")
        nama_proyek   = st.text_input("Nama Proyek", value="Perhitungan Kolom Beton Bertulang")
        tgl_laporan   = st.text_input("Tanggal Laporan", value="April 2025")

        col_dl1, col_dl2 = st.columns(2)

        with col_dl1:
            if st.button("📄 Generate Laporan Word (.docx)", use_container_width=True):
                try:
                    docx_bytes = buat_laporan_word(
                        R, nama_engineer, nama_proyek, tgl_laporan, Mu_desain
                    )
                    st.download_button(
                        "⬇️ Download Word",
                        data=docx_bytes,
                        file_name="laporan_kolom.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Error generate Word: {e}")

        with col_dl2:
            if st.button("📋 Generate Laporan PDF (.pdf)", use_container_width=True):
                try:
                    pdf_bytes = buat_laporan_pdf(
                        R, nama_engineer, nama_proyek, tgl_laporan, Mu_desain
                    )
                    st.download_button(
                        "⬇️ Download PDF",
                        data=pdf_bytes,
                        file_name="laporan_kolom.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as e:
                    st.error(f"Error generate PDF: {e}")


# ============================================================
# BAGIAN 3: GENERATOR LAPORAN
# ============================================================

@st.cache_data
def buat_laporan_word(R, nama_engineer, nama_proyek, tgl_laporan, Mu_desain):
    """Generate laporan Word Engineering Report."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io

    doc = Document()

    # --- Setup halaman ---
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2)

    styles = doc.styles

    def add_heading(text, level=1, color_rgb=(0, 70, 127)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color_rgb)
            run.font.bold = True
        return p

    def add_para(text, bold=False, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_table_2col(headers, rows_data):
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for para in hdr_cells[i].paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for row_data in rows_data:
            row_cells = table.add_row().cells
            for i, val in enumerate(row_data):
                row_cells[i].text = str(val)
                for para in row_cells[i].paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
        return table

    # ===================== COVER =====================
    doc.add_paragraph("")
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("LAPORAN PERHITUNGAN STRUKTUR")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0, 70, 127)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("KAPASITAS KOLOM BETON BERTULANG")
    r2.font.size = Pt(14)
    r2.font.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("DIAGRAM INTERAKSI P-M | SNI 2847:2019")
    r3.font.size = Pt(12)

    doc.add_paragraph("")
    info_data = [
        ("Proyek", nama_proyek),
        ("Engineer", nama_engineer),
        ("Tanggal", tgl_laporan),
        ("Standar", "SNI 2847:2019"),
    ]
    for k, v in info_data:
        p_i = doc.add_paragraph()
        p_i.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_i.add_run(f"{k}: ").bold = True
        p_i.add_run(v)

    doc.add_page_break()

    # ===================== BAB 1: DATA INPUT =====================
    add_heading("BAB 1: DATA INPUT PARAMETER", level=1)

    add_heading("1.1 Material", level=2)
    add_table_2col(
        ["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
        [
            ["Mutu Beton", "fc'", f"{R['inp']['fc']:.0f}", "MPa", "Kuat tekan karakteristik"],
            ["Mutu Baja", "fy", f"{R['inp']['fy']:.0f}", "MPa", "Tegangan leleh"],
            ["Modulus Elastisitas Baja", "Es", f"{R['inp']['Es']:.0f}", "MPa", "SNI 2847:2019"],
        ]
    )

    doc.add_paragraph("")
    add_heading("1.2 Dimensi Kolom", level=2)
    add_table_2col(
        ["Parameter", "Simbol", "Nilai", "Satuan"],
        [
            ["Lebar Kolom", "b", f"{R['inp']['b']:.0f}", "mm"],
            ["Tinggi Kolom", "h", f"{R['inp']['h']:.0f}", "mm"],
            ["Selimut Beton", "c", f"{R['inp']['c_sel']:.0f}", "mm"],
            ["Diameter Sengkang", "Os", f"{R['inp']['dia_s']:.0f}", "mm"],
        ]
    )

    doc.add_paragraph("")
    add_heading("1.3 Tulangan Longitudinal", level=2)
    add_table_2col(
        ["Parameter", "Simbol", "Nilai", "Satuan"],
        [
            ["Diameter Tulangan", "D", f"{R['inp']['D']:.0f}", "mm"],
            ["Jumlah Tulangan Sisi b", "n_b", f"{R['inp']['n_b']}", "buah"],
            ["Jumlah Tulangan Sisi h", "n_h", f"{R['inp']['n_h']}", "buah"],
            ["Total Tulangan", "n_total", f"{R['props']['n_total']}", "buah"],
        ]
    )

    doc.add_paragraph("")
    add_heading("1.4 Beban Terfaktor", level=2)
    add_table_2col(
        ["Parameter", "Simbol", "Nilai", "Satuan"],
        [
            ["Gaya Aksial Terfaktor", "Pu", f"{R['inp']['Pu']:.0f}", "kN"],
            ["Momen Ujung 1", "M1", f"{R['inp']['M1']:.0f}", "kN.m"],
            ["Momen Ujung 2", "M2", f"{R['inp']['M2']:.0f}", "kN.m"],
            ["Momen Desain", "Mc", f"{Mu_desain:.4f}", "kN.m"],
        ]
    )

    doc.add_page_break()

    # ===================== BAB 2: PERHITUNGAN =====================
    add_heading("BAB 2: PROSES PERHITUNGAN (STEP-BY-STEP)", level=1)

    add_heading("2.1 Properti Penampang", level=2)
    props = R["props"]
    inp   = R["inp"]
    calc_text = (
        f"Ag = b x h = {inp['b']:.0f} x {inp['h']:.0f} = {props['Ag']:,.2f} mm2\n"
        f"Ig = (1/12) x b x h3 = (1/12) x {inp['b']:.0f} x {inp['h']:.0f}3 = {props['Ig']:,.2f} mm4\n"
        f"Ec = 4700 x sqrt(fc') = 4700 x sqrt({inp['fc']:.0f}) = {props['Ec']:,.2f} MPa\n"
        f"d' = c + Os + D/2 = {inp['c_sel']:.0f} + {inp['dia_s']:.0f} + {inp['D']:.0f}/2 = {props['d_prime']:.2f} mm\n"
        f"d  = h - d' = {inp['h']:.0f} - {props['d_prime']:.2f} = {props['d']:.2f} mm\n"
        f"n_total = 2 x n_b + 2 x (n_h-2) = 2x{inp['n_b']} + 2x({inp['n_h']}-2) = {props['n_total']} buah\n"
        f"Ast = n_total x pi x D2/4 = {props['n_total']} x pi x {inp['D']:.0f}2/4 = {props['Ast']:,.2f} mm2\n"
        f"rho_g = Ast/Ag x 100 = {props['Ast']:.2f}/{props['Ag']:.2f} x 100 = {props['rho_g']:.2f}% -> {props['cek_rasio']}"
    )
    p_calc = doc.add_paragraph()
    run_calc = p_calc.add_run(calc_text)
    run_calc.font.name = "Courier New"
    run_calc.font.size = Pt(8.5)

    add_heading("2.2 Faktor Psi dan Kelangsingan", level=2)
    kekakuan     = R["kekakuan"]
    kelangsingan = R["kelangsingan"]
    calc2 = (
        f"PSI_A = {kekakuan['sum_kol_atas']:,.2f} / {kekakuan['sum_balok_atas']:,.2f} = {kekakuan['psi_A']:.4f}\n"
        f"PSI_B = {kekakuan['sum_kol_bawah']:,.2f} / {kekakuan['sum_balok_bawah']:,.2f} = {kekakuan['psi_B']:.4f}\n"
        f"k ({inp['kondisi_rangka']}) = {kelangsingan['k']:.6f}\n"
        f"r = 0.3 x h = 0.3 x {inp['h']:.0f} = {kelangsingan['r']:.2f} mm\n"
        f"k x Lu / r = {kelangsingan['k']:.4f} x {inp['Lu']:.0f} / {kelangsingan['r']:.2f} = {kelangsingan['rasio_kelangsingan']:.2f}\n"
        f"Batas = {kelangsingan['batas']:.2f} -> {kelangsingan['klasifikasi']}"
    )
    p2 = doc.add_paragraph()
    run2 = p2.add_run(calc2)
    run2.font.name = "Courier New"
    run2.font.size = Pt(8.5)

    pembesaran = R["pembesaran"]
    if kelangsingan["klasifikasi"] == "Slender Column":
        add_heading("2.3 Pembesaran Momen", level=2)
        calc3 = (
            f"(EI)eff = 0.4 x Ec x Ig / (1 + beta_dns)\n"
            f"        = 0.4 x {pembesaran['Ec']:,.2f} x {props['Ig']:,.2f} / (1 + {pembesaran['beta_dns']:.2f})\n"
            f"        = {pembesaran['EI_eff']:,.2f} N.mm2\n"
            f"Pc      = pi2 x (EI)eff / (k x Lu)2\n"
            f"        = {pembesaran['Pc']:,.2f} kN\n"
            f"Cm      = 0.6 + 0.4 x (M1/M2) = {pembesaran['Cm']:.2f}\n"
            f"delta_ns = Cm / (1 - Pu/(0.75 x Pc)) = {pembesaran['delta_ns']:.2f}\n"
            f"Mc      = delta_ns x M2 = {pembesaran['delta_ns']:.2f} x {inp['M2']:.0f} = {pembesaran['Mc']:.4f} kN.m"
        )
        p3 = doc.add_paragraph()
        run3 = p3.add_run(calc3)
        run3.font.name = "Courier New"
        run3.font.size = Pt(8.5)

    add_heading("2.4 Diagram Interaksi - Tabel 52 Titik", level=2)
    # Tabel ringkas (show 10 titik saja untuk Word agar tidak terlalu panjang)
    interaksi_rows = []
    for row in R["hasil_interaksi"]:
        try:
            interaksi_rows.append([
                str(row["No"]), str(row["c/h"]),
                f"{float(row['Pn']):,.2f}" if row["Pn"] != "-" else "-",
                f"{float(row['Mn']):,.2f}" if row["Mn"] != "-" else "-",
                str(row["phi"]),
                f"{float(row['phi_Pn']):,.2f}" if row["phi_Pn"] != "-" else "-",
                f"{float(row['phi_Mn']):,.2f}" if row["phi_Mn"] != "-" else "-",
            ])
        except:
            interaksi_rows.append([
                str(row["No"]), str(row["c/h"]),
                str(row.get("Pn", "-")), str(row.get("Mn", "-")),
                str(row.get("phi", "-")),
                str(row.get("phi_Pn", "-")), str(row.get("phi_Mn", "-"))
            ])

    add_table_2col(
        ["No", "c/h", "Pn (kN)", "Mn (kN.m)", "phi", "phiPn (kN)", "phiMn (kN.m)"],
        interaksi_rows
    )

    doc.add_page_break()

    # ===================== BAB 3: KESIMPULAN =====================
    add_heading("BAB 3: KESIMPULAN & HASIL AKHIR", level=1)
    cek = R["cek"]
    kesimpulan_rows = [
        ["Klasifikasi Kolom", kelangsingan["klasifikasi"], "-"],
        ["Gaya Aksial (Pu)", f"{inp['Pu']:.2f}", "kN"],
        ["Momen Desain (Mc)", f"{Mu_desain:.4f}", "kN.m"],
        ["phi.Pn Kapasitas", f"{cek['phi_Pn_kapasitas']:.2f}" if cek["phi_Pn_kapasitas"] else "-", "kN"],
        ["phi.Mn Kapasitas", f"{cek['phi_Mn_kapasitas']:.2f}" if cek["phi_Mn_kapasitas"] else "-", "kN.m"],
        ["Rasio Pu/phi.Pn", f"{cek['ratio_Pu']:.2f}" if cek["ratio_Pu"] else "-", "-"],
        ["Rasio Mu/phi.Mn", f"{cek['ratio_Mu']:.2f}" if cek["ratio_Mu"] else "-", "-"],
        ["STATUS AKHIR", cek["status"], "-"],
    ]
    add_table_2col(["Parameter", "Nilai", "Satuan"], kesimpulan_rows)

    doc.add_paragraph("")
    p_note = doc.add_paragraph()
    p_note.add_run("Catatan: ").bold = True
    p_note.add_run(
        "Perhitungan berdasarkan SNI 2847:2019. "
        "Diagram Interaksi dibuat dengan 52 titik (50 titik c/h + tekan murni + tarik murni). "
        f"Dibuat oleh: {nama_engineer} | {tgl_laporan}"
    )

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


@st.cache_data
def buat_laporan_pdf(R, nama_engineer, nama_proyek, tgl_laporan, Mu_desain):
    """Generate laporan PDF Engineering Report."""
    from fpdf import FPDF
    import io

    def sanitasi(text):
        """Ganti karakter unicode agar kompatibel latin-1."""
        replacements = {
            '\u03c6': 'phi', '\u03a6': 'Phi',
            '\u00b2': '2',   '\u00b3': '3',
            '\u221a': 'sqrt', '\u00b7': '.',
            '\u00d8': 'O',   '\u2248': '~',
            '\u03b5': 'eps', '\u03b2': 'beta',
            '\u03b4': 'delta', '\u03c1': 'rho',
            '\u03a8': 'Psi', '\u03c8': 'psi',
            '\u2264': '<=',  '\u2265': '>=',
            '\u00d7': 'x',   '\u03c0': 'pi',
            '\u221e': 'inf',
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        # Encode/decode untuk buang karakter non-latin-1 sisanya
        return text.encode('latin-1', errors='replace').decode('latin-1')

    props        = R["props"]
    kekakuan     = R["kekakuan"]
    kelangsingan = R["kelangsingan"]
    pembesaran   = R["pembesaran"]
    cek          = R["cek"]
    inp          = R["inp"]

    class PDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "B", 9)
            self.set_text_color(0, 70, 127)
            self.cell(0, 8, sanitasi("LAPORAN KAPASITAS KOLOM BETON BERTULANG | SNI 2847:2019"), align="C")
            self.ln(5)
            self.set_draw_color(0, 70, 127)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(3)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 7)
            self.set_text_color(128, 128, 128)
            self.cell(0, 10, sanitasi(f"Engineer: {nama_engineer} | {tgl_laporan} | Hal. {self.page_no()}"), align="C")
            self.set_text_color(0, 0, 0)

        def section_title(self, text, level=1):
            if level == 1:
                self.set_font("Helvetica", "B", 12)
                self.set_fill_color(0, 70, 127)
                self.set_text_color(255, 255, 255)
                self.cell(0, 8, sanitasi(text), fill=True, ln=True)
                self.set_text_color(0, 0, 0)
            else:
                self.set_font("Helvetica", "B", 10)
                self.set_text_color(0, 70, 127)
                self.cell(0, 6, sanitasi(text), ln=True)
                self.set_text_color(0, 0, 0)
            self.ln(1)

        def kv_row(self, label, value, unit="", keterangan=""):
            self.set_font("Helvetica", "", 8.5)
            self.cell(75, 5.5, sanitasi(label), border="TB")
            self.cell(40, 5.5, sanitasi(str(value)), border="TB")
            self.cell(20, 5.5, sanitasi(unit), border="TB")
            self.cell(0,  5.5, sanitasi(keterangan), border="TB", ln=True)

        def tabel_header(self, cols, widths):
            self.set_font("Helvetica", "B", 8)
            self.set_fill_color(220, 230, 241)
            for col, w in zip(cols, widths):
                self.cell(w, 6, sanitasi(col), border=1, fill=True)
            self.ln()

        def tabel_row(self, vals, widths):
            self.set_font("Helvetica", "", 7.5)
            for val, w in zip(vals, widths):
                self.cell(w, 5.5, sanitasi(str(val)), border=1)
            self.ln()

    pdf = PDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(12, 15, 12)

    # ---- COVER ----
    pdf.add_page()
    pdf.ln(20)
    pdf.set_font("Helvetica", "B", 20)
    pdf.set_text_color(0, 70, 127)
    pdf.cell(0, 12, "LAPORAN PERHITUNGAN STRUKTUR", align="C", ln=True)
    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "KAPASITAS KOLOM BETON BERTULANG", align="C", ln=True)
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "DIAGRAM INTERAKSI P-M | SNI 2847:2019", align="C", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(15)

    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 7, "Proyek :", align="R")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, sanitasi(nama_proyek), ln=True)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 7, "Engineer :", align="R")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, sanitasi(nama_engineer), ln=True)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 7, "Tanggal :", align="R")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, sanitasi(tgl_laporan), ln=True)
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(50, 7, "Standar :", align="R")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, "SNI 2847:2019", ln=True)

    # ---- BAB 1 ----
    pdf.add_page()
    pdf.section_title("BAB 1: DATA INPUT PARAMETER", level=1)

    pdf.section_title("1.1 Material", level=2)
    pdf.tabel_header(["Parameter", "Simbol", "Nilai", "Satuan", "Keterangan"],
                     [60, 20, 30, 20, 56])
    pdf.tabel_row(["Mutu Beton", "fc'", f"{inp['fc']:.0f}", "MPa", "Kuat tekan karakteristik"],
                  [60, 20, 30, 20, 56])
    pdf.tabel_row(["Mutu Baja", "fy", f"{inp['fy']:.0f}", "MPa", "Tegangan leleh tulangan"],
                  [60, 20, 30, 20, 56])
    pdf.tabel_row(["Modulus Elastisitas Baja", "Es", f"{inp['Es']:.0f}", "MPa", "SNI 2847:2019"],
                  [60, 20, 30, 20, 56])
    pdf.ln(3)

    pdf.section_title("1.2 Dimensi Kolom", level=2)
    pdf.tabel_header(["Parameter", "Simbol", "Nilai", "Satuan"], [70, 25, 40, 51])
    for row in [
        ["Lebar Kolom", "b", f"{inp['b']:.0f}", "mm"],
        ["Tinggi Kolom", "h", f"{inp['h']:.0f}", "mm"],
        ["Selimut Beton", "c", f"{inp['c_sel']:.0f}", "mm"],
        ["Diameter Sengkang", "Os", f"{inp['dia_s']:.0f}", "mm"],
        ["Diameter Tulangan", "D", f"{inp['D']:.0f}", "mm"],
        ["Jumlah Tulangan Sisi b", "n_b", f"{inp['n_b']}", "buah"],
        ["Jumlah Tulangan Sisi h", "n_h", f"{inp['n_h']}", "buah"],
    ]:
        pdf.tabel_row(row, [70, 25, 40, 51])
    pdf.ln(3)

    pdf.section_title("1.3 Beban Terfaktor", level=2)
    pdf.tabel_header(["Parameter", "Simbol", "Nilai", "Satuan"], [70, 25, 40, 51])
    for row in [
        ["Gaya Aksial Terfaktor", "Pu", f"{inp['Pu']:.0f}", "kN"],
        ["Momen Ujung 1", "M1", f"{inp['M1']:.0f}", "kN.m"],
        ["Momen Ujung 2 (Design)", "M2", f"{inp['M2']:.0f}", "kN.m"],
        ["Momen Desain (Mc)", "Mc", f"{Mu_desain:.4f}", "kN.m"],
    ]:
        pdf.tabel_row(row, [70, 25, 40, 51])
    pdf.ln(3)

    # ---- BAB 2 ----
    pdf.add_page()
    pdf.section_title("BAB 2: PROSES PERHITUNGAN (STEP-BY-STEP)", level=1)

    pdf.section_title("2.1 Properti Penampang (SNI 2847:2019)", level=2)
    calc_lines = [
        f"Ag  = b x h = {inp['b']:.0f} x {inp['h']:.0f} = {props['Ag']:,.2f} mm2",
        f"Ig  = (1/12) x b x h3 = {props['Ig']:,.2f} mm4",
        f"Ec  = 4700 x sqrt({inp['fc']:.0f}) = {props['Ec']:,.2f} MPa  [SNI Ps.19.2.2]",
        f"d'  = c + Os + D/2 = {inp['c_sel']:.0f}+{inp['dia_s']:.0f}+{inp['D']:.0f}/2 = {props['d_prime']:.2f} mm",
        f"d   = h - d' = {inp['h']:.0f} - {props['d_prime']:.2f} = {props['d']:.2f} mm",
        f"n_t = 2x{inp['n_b']} + 2x({inp['n_h']}-2) = {props['n_total']} buah",
        f"Ast = {props['n_total']} x pi x {inp['D']:.0f}2/4 = {props['Ast']:,.2f} mm2",
        f"rho = Ast/Ag x 100 = {props['rho_g']:.2f}% -> {props['cek_rasio']} [1%-8%]",
    ]
    pdf.set_font("Courier", "", 8.5)
    for line in calc_lines:
        pdf.cell(0, 5, sanitasi(line), ln=True)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.ln(2)

    pdf.section_title("2.2 Kekakuan Elemen & Faktor Psi (SNI Ps.6.6.4.4)", level=2)
    psi_lines = [
        f"EI/L Balok Atas  Kiri  = {kekakuan['EI_L_bal_atas_kiri']:,.2f} N.mm",
        f"EI/L Balok Atas  Kanan = {kekakuan['EI_L_bal_atas_kanan']:,.2f} N.mm",
        f"EI/L Kolom Atas        = {kekakuan['EI_L_kol_atas']:,.2f} N.mm",
        f"EI/L Kolom Ditinjau    = {kekakuan['EI_L_kol_ditinjau']:,.2f} N.mm",
        f"PSI_A = {kekakuan['sum_kol_atas']:,.2f} / {kekakuan['sum_balok_atas']:,.2f} = {kekakuan['psi_A']:.4f}",
        f"PSI_B = {kekakuan['sum_kol_bawah']:,.2f} / {kekakuan['sum_balok_bawah']:,.2f} = {kekakuan['psi_B']:.4f}",
        f"k ({inp['kondisi_rangka']}) = {kelangsingan['k']:.6f}",
        f"r = 0.3 x h = {kelangsingan['r']:.2f} mm",
        f"k x Lu / r = {kelangsingan['rasio_kelangsingan']:.2f}  Batas = {kelangsingan['batas']:.2f}",
        f"Klasifikasi: {kelangsingan['klasifikasi']}",
    ]
    pdf.set_font("Courier", "", 8.5)
    for line in psi_lines:
        pdf.cell(0, 5, sanitasi(line), ln=True)
    pdf.set_font("Helvetica", "", 8.5)
    pdf.ln(2)

    if kelangsingan["klasifikasi"] == "Slender Column":
        pdf.section_title("2.3 Pembesaran Momen (SNI Ps.6.6.4)", level=2)
        pem_lines = [
            f"(EI)eff = 0.4 x Ec x Ig / (1+beta_dns)",
            f"       = 0.4 x {pembesaran['Ec']:,.0f} x {props['Ig']:,.0f} / (1+{pembesaran['beta_dns']:.2f})",
            f"       = {pembesaran['EI_eff']:,.2f} N.mm2",
            f"Pc     = pi2 x (EI)eff / (k x Lu)2 = {pembesaran['Pc']:,.2f} kN",
            f"Cm     = 0.6 + 0.4 x (M1/M2) = {pembesaran['Cm']:.2f}  [min 0.4]",
            f"delta  = Cm/(1-Pu/(0.75xPc)) = {pembesaran['delta_ns']:.2f}  [min 1.0]",
            f"Mc     = delta x M2 = {pembesaran['delta_ns']:.2f} x {inp['M2']:.0f} = {pembesaran['Mc']:.4f} kN.m",
        ]
        pdf.set_font("Courier", "", 8.5)
        for line in pem_lines:
            pdf.cell(0, 5, sanitasi(line), ln=True)
        pdf.set_font("Helvetica", "", 8.5)
        pdf.ln(2)

    pdf.section_title("2.4 Tabel Diagram Interaksi P-M (52 Titik)", level=2)
    cols_int = ["No", "c/h", "Pn(kN)", "Mn(kN.m)", "phi", "phiPn(kN)", "phiMn(kN.m)"]
    widths_int = [10, 18, 25, 26, 12, 25, 26]
    pdf.tabel_header(cols_int, widths_int)
    for row in R["hasil_interaksi"]:
        try:
            pdf.tabel_row([
                str(row["No"]), str(row["c/h"]),
                f"{float(row['Pn']):,.2f}" if row["Pn"] not in ["-", 0] else str(row["Pn"]),
                f"{float(row['Mn']):,.2f}" if row["Mn"] not in ["-", 0] else str(row["Mn"]),
                str(row["phi"]),
                f"{float(row['phi_Pn']):,.2f}" if row["phi_Pn"] not in ["-"] else "-",
                f"{float(row['phi_Mn']):,.2f}" if row["phi_Mn"] not in ["-"] else "-",
            ], widths_int)
        except:
            pdf.tabel_row([str(row.get(k, "-")) for k in
                          ["No", "c/h", "Pn", "Mn", "phi", "phi_Pn", "phi_Mn"]],
                         widths_int)

    # ---- BAB 3 ----
    pdf.add_page()
    pdf.section_title("BAB 3: KESIMPULAN & HASIL AKHIR", level=1)

    pdf.section_title("3.1 Ringkasan Kapasitas", level=2)
    cols_k = ["Parameter", "Nilai", "Satuan", "Status"]
    widths_k = [75, 40, 25, 46]
    pdf.tabel_header(cols_k, widths_k)
    kesimpulan = [
        ["Klasifikasi Kolom", kelangsingan["klasifikasi"], "-", "-"],
        ["Gaya Aksial (Pu)", f"{inp['Pu']:.2f}", "kN", "INPUT"],
        ["Momen Desain (Mc)", f"{Mu_desain:.4f}", "kN.m", "INPUT"],
        ["phi.Pn Kapasitas", f"{cek['phi_Pn_kapasitas']:.2f}" if cek['phi_Pn_kapasitas'] else "-", "kN", "-"],
        ["phi.Mn Kapasitas", f"{cek['phi_Mn_kapasitas']:.2f}" if cek['phi_Mn_kapasitas'] else "-", "kN.m", "-"],
        ["Rasio Pu/phi.Pn", f"{cek['ratio_Pu']:.2f}" if cek['ratio_Pu'] else "-", "-",
         "OK" if (cek['ratio_Pu'] or 0) <= 1 else "NOT OK"],
        ["Rasio Mu/phi.Mn", f"{cek['ratio_Mu']:.2f}" if cek['ratio_Mu'] else "-", "-",
         "OK" if (cek['ratio_Mu'] or 0) <= 1 else "NOT OK"],
        ["STATUS AKHIR", cek["status"], "-", cek["status"]],
    ]
    for row in kesimpulan:
        pdf.tabel_row(row, widths_k)

    pdf.ln(5)
    pdf.section_title("3.2 Catatan & Referensi", level=2)
    catatan = [
        "- Perhitungan berdasarkan SNI 2847:2019",
        "- Diagram Interaksi dibuat dengan 52 titik (50 titik c/h + tekan murni + tarik murni)",
        "- Faktor reduksi phi dinamis (0.65 s/d 0.90) sesuai regangan tarik tulangan",
        "- Pembesaran momen diterapkan untuk Slender Column pada Braced Frame",
        f"- Dibuat oleh: {nama_engineer} | {tgl_laporan}",
    ]
    pdf.set_font("Helvetica", "", 9)
    for cat in catatan:
        pdf.cell(0, 6, sanitasi(cat), ln=True)

    pdf_bytes = pdf.output()
    return bytes(pdf_bytes)


# ============================================================
# FOOTER
# ============================================================
st.divider()
st.markdown(
    "<div style='text-align:center;color:gray;font-size:11px;'>"
    "🏛️ Kolom Beton Bertulang | SNI 2847:2019 | Ladosi | "
    "Precast Concrete Structural Engineering</div>",
    unsafe_allow_html=True
)
